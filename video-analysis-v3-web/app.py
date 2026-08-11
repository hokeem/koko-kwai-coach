#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Public-facing web UI for video-analysis-v3."""
from __future__ import annotations

import base64
import csv
import errno
import hashlib
import html
import http.client
import io
import json
import os
import re
import shutil
import secrets
import subprocess
import sys
import threading
import time
import urllib.parse
import urllib.error
import urllib.request
import xml.etree.ElementTree as ET
import zipfile
from collections import Counter, deque
from concurrent.futures import ThreadPoolExecutor, as_completed
from datetime import datetime, timedelta, timezone
from http import HTTPStatus
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path
from typing import Any
from uuid import uuid4
from PIL import Image, ImageDraw


PORT = int(os.environ.get("PORT", 8310))
PIPELINE_TIMEOUT_SEC = int(os.environ.get("VIDEO_ANALYSIS_PIPELINE_TIMEOUT_SEC", "720"))
MAX_CONCURRENT_ANALYSES = max(1, int(os.environ.get("VIDEO_ANALYSIS_MAX_CONCURRENT_JOBS", "1")))
RUNNING_TASK_STALE_SEC = int(os.environ.get("VIDEO_ANALYSIS_RUNNING_STALE_SEC", "900"))
REVIEW_TASK_STALE_SEC = int(os.environ.get("VIDEO_ANALYSIS_REVIEW_STALE_SEC", "900"))
PROCESSLESS_TASK_STALE_SEC = int(os.environ.get("VIDEO_ANALYSIS_PROCESSLESS_STALE_SEC", "180"))
RESTORE_PENDING_MAX_AGE_SEC = int(os.environ.get("VIDEO_ANALYSIS_RESTORE_PENDING_MAX_AGE_SEC", "1800"))
WATCHDOG_INTERVAL_SEC = int(os.environ.get("VIDEO_ANALYSIS_WATCHDOG_INTERVAL_SEC", "15"))
SOURCE_VIDEO_RETENTION_DAYS = int(os.environ.get("VIDEO_ANALYSIS_SOURCE_RETENTION_DAYS", "2"))
RAW_ARTIFACT_RETENTION_DAYS = int(os.environ.get("VIDEO_ANALYSIS_RAW_RETENTION_DAYS", "14"))
MAX_CONCURRENT_FILTERS = max(1, int(os.environ.get("VIDEO_FILTER_MAX_CONCURRENT_JOBS", "1")))
FILTER_USE_LLM = str(os.environ.get("VIDEO_FILTER_USE_LLM", "1")).strip().lower() in {"1", "true", "yes", "on"}
FILTER_DURATION_MIN_SEC = int(os.environ.get("VIDEO_FILTER_DURATION_MIN_SEC", "30"))
FILTER_DURATION_MAX_SEC = int(os.environ.get("VIDEO_FILTER_DURATION_MAX_SEC", "120"))
FILTER_AUDIO_MAX_MB = int(os.environ.get("VIDEO_FILTER_AUDIO_MAX_MB", "18"))
MAX_CONCURRENT_TRANSLATIONS = max(1, int(os.environ.get("VIDEO_TRANSLATION_MAX_CONCURRENT_JOBS", "1")))
SYNC_LIBRARY_ON_STARTUP = str(os.environ.get("VIDEO_ANALYSIS_SYNC_LIBRARY_ON_STARTUP", "0")).strip().lower() in {"1", "true", "yes", "on"}
GEMINI_TRANSIENT_RETRY_ATTEMPTS = max(1, int(os.environ.get("VIDEO_ANALYSIS_TRANSIENT_RETRY_ATTEMPTS", "3")))
GEMINI_TRANSIENT_RETRY_DELAYS = [
    max(1, int(value))
    for value in re.split(r"[, ]+", os.environ.get("VIDEO_ANALYSIS_TRANSIENT_RETRY_DELAYS", "45,120").strip())
    if value.strip().isdigit()
] or [45, 120]


def stage_timeout(name: str, default: int) -> int:
    env_name = f"VIDEO_ANALYSIS_STAGE_TIMEOUT_{name.upper()}"
    return int(os.environ.get(env_name, str(default)))


STAGE_TIMEOUTS_SEC = {
    "download": stage_timeout("download", 180),
    "media_prep": stage_timeout("media_prep", 60),
    "gemini_analysis": stage_timeout("gemini_analysis", 240),
    "v2_analysis": stage_timeout("v2_analysis", 360),
    "consistency_audit": stage_timeout("consistency_audit", 120),
    "targeted_recheck": stage_timeout("targeted_recheck", 240),
    "arbitration": stage_timeout("arbitration", 90),
    "final_output": stage_timeout("final_output", 240),
}
BASE = Path(__file__).resolve().parent
REPO_ROOT = BASE.parent
SKILL_ROOT = REPO_ROOT / "skills" / "video-analysis-v3"
V2_SKILL_ROOT = REPO_ROOT / "skills" / "video-analysis-v2-skill"
AUTO_ANALYZE = SKILL_ROOT / "scripts" / "hybrid_v2_pipeline.py"
TRANSCREATE_VIDEO = SKILL_ROOT / "scripts" / "transcreate_video.py"
DATA_ROOT = Path(os.environ.get("VIDEO_ANALYSIS_WEB_DATA_DIR", str(BASE / "data"))).expanduser()
JOBS_FILE = DATA_ROOT / "jobs.json"
TRANSLATION_JOBS_FILE = DATA_ROOT / "translation_jobs.json"
RESULTS_ROOT = DATA_ROOT / "results"
LIBRARY_FILE = DATA_ROOT / "script_library.json"
ERROR_CASE_LIBRARY_FILE = DATA_ROOT / "error_case_library.json"
CREATOR_ONLINE_LIBRARY_FILE = DATA_ROOT / "creator_online_library.json"
CREATOR_THUMBNAIL_CACHE_FILE = DATA_ROOT / "creator_thumbnail_cache.json"
CREATOR_SUBMISSIONS_FILE = DATA_ROOT / "creator_submissions.json"
CREATOR_SYNC_META_FILE = DATA_ROOT / "creator_sync_meta.json"
CREATOR_IMPORT_JOBS_FILE = DATA_ROOT / "creator_import_jobs.json"
CREATOR_ADMIN_SCRIPTS_CACHE_FILE = DATA_ROOT / "creator_admin_scripts_cache.json"
CREATOR_ADMIN_STATE_FILE = DATA_ROOT / "creator_admin_state.json"
CREATOR_LIBRARY_SOURCE_URL = os.environ.get("CREATOR_LIBRARY_SOURCE_URL", "https://koko-kwai-coach.onrender.com/api/library")
CREATOR_LIBRARY_SYNC_INTERVAL_SEC = int(os.environ.get("CREATOR_LIBRARY_SYNC_INTERVAL_SEC", "86400"))
CREATOR_CENTER_SYNC_URL = os.environ.get("CREATOR_CENTER_SYNC_URL", "https://kokocomedy.com/api/creator/sync-library")
CREATOR_CENTER_BASE_URL = os.environ.get("CREATOR_CENTER_BASE_URL", "https://kokocomedy.com").rstrip("/")
CREATOR_ADMIN_PASSWORD = os.environ.get("KOKO_CREATOR_ADMIN_PASSWORD", "koko")
CREATOR_ADMIN_AUTH_COOKIE = "koko_creator_admin_auth"
CREATOR_REMOTE_ADMIN_COOKIE = "koko_creator_admin"
CREATOR_IMPORT_MAX_WORKERS = max(1, int(os.environ.get("CREATOR_IMPORT_MAX_WORKERS", "3")))
CREATOR_IMPORT_IMAGE_RETRY_ATTEMPTS = max(1, int(os.environ.get("CREATOR_IMPORT_IMAGE_RETRY_ATTEMPTS", "2")))
STORYBOARD_LOCAL_ONLY = str(os.environ.get("VIDEO_ANALYSIS_STORYBOARD_LOCAL_ONLY", "0")).strip().lower() in {"1", "true", "yes", "on"}


def creator_script_share_url(entry_id: object) -> str:
    script_id = str(entry_id or "").strip()
    if not re.fullmatch(r"[0-9a-f]{32}", script_id):
        return ""
    return f"{CREATOR_CENTER_BASE_URL}/script/{script_id}"
FILTER_JOBS_FILE = DATA_ROOT / "filter_jobs.json"
FILTER_CACHE_ROOT = DATA_ROOT / "filter_cache"
VISION_MODELS_DIR = DATA_ROOT / "vision_models"


def load_env_file(path: Path) -> None:
    if not path.exists():
        return
    for raw_line in path.read_text(encoding="utf-8", errors="ignore").splitlines():
        line = raw_line.strip()
        if not line or line.startswith("#") or "=" not in line:
            continue
        key, value = line.split("=", 1)
        key = key.strip()
        value = value.strip().strip("\"'")
        if key and key not in os.environ:
            os.environ[key] = value


load_env_file(REPO_ROOT / ".env.local")
load_env_file(BASE / ".env.local")
GOOGLE_API_KEY = os.environ.get("GOOGLE_API_KEY", "")
FILTER_USE_LLM = str(os.environ.get("VIDEO_FILTER_USE_LLM", "1")).strip().lower() in {"1", "true", "yes", "on"}
FILTER_DURATION_MIN_SEC = int(os.environ.get("VIDEO_FILTER_DURATION_MIN_SEC", str(FILTER_DURATION_MIN_SEC)))
FILTER_DURATION_MAX_SEC = int(os.environ.get("VIDEO_FILTER_DURATION_MAX_SEC", str(FILTER_DURATION_MAX_SEC)))
FILTER_AUDIO_MAX_MB = int(os.environ.get("VIDEO_FILTER_AUDIO_MAX_MB", str(FILTER_AUDIO_MAX_MB)))
ERROR_CASE_PASSWORD = "kwai666"
ERROR_CASE_AUTH_COOKIE = "koko_error_case_auth"
ASSETS_ROOT = BASE / "assets"
HERO_WORDMARK = ASSETS_ROOT / "kwai-wordmark.svg"
KWAI_FAVICON = ASSETS_ROOT / "kwai-favicon.svg"
STUDIO_HERO_VIDEO = ASSETS_ROOT / "studio-hero-video.png"
STUDIO_TITLE_ART = ASSETS_ROOT / "studio-title-art.png"
STUDIO_HERO_BANNER_SHALLOW = ASSETS_ROOT / "studio-hero-banner-shallow.png"
FAVICON_LINKS = """<link rel="icon" type="image/svg+xml" href="/favicon.svg?v=kwai1">
  <link rel="shortcut icon" href="/favicon.ico?v=kwai1">"""


def parse_model_candidates(*groups: str) -> list[str]:
    ordered: list[str] = []
    seen: set[str] = set()
    retired = {"gemini-2.0-flash", "gemini-2.0-flash-lite"}
    for group in groups:
        for name in re.split(r"[,;\s]+", str(group or "")):
            value = name.strip()
            if not value or value in seen or value in retired:
                continue
            seen.add(value)
            ordered.append(value)
    return ordered


STABLE_VIDEO_MODELS = [
    "gemini-2.5-flash-lite",
    "gemini-2.5-flash",
    "gemini-2.5-pro",
    "gemini-3-flash-preview",
]
MODEL_CANDIDATES = parse_model_candidates(
    ",".join(STABLE_VIDEO_MODELS),
    os.environ.get("VIDEO_ANALYSIS_MODEL", ""),
)
IMAGE_MODEL_CANDIDATES = parse_model_candidates(
    "gemini-3.1-flash-image,gemini-2.5-flash-image",
    os.environ.get("VIDEO_ANALYSIS_IMAGE_MODEL", ""),
)
BEIJING_TZ = timezone(timedelta(hours=8))
SOURCE_VIDEO_NAME = "source.mp4"
STORYBOARD_PROMPT_FILE = "storyboard_prompt.txt"
STORYBOARD_METADATA_FILE = "storyboard_cover.json"
STORYBOARD_PREVIEW_BASENAME = "storyboard_preview"
STORYBOARD_COVER_BASENAME = "storyboard_cover"
RAW_ARTIFACT_NAMES = {
    "primary_analysis_raw_gemini.json",
    "v2_local_raw_gemini.json",
    "comparison_raw_gemini.json",
    "logic_audit_raw_gemini.json",
    "supplement_raw_gemini.json",
    "arbitration_raw_gemini.json",
    "audio_multiview_raw_gemini.json",
    "final_refine_raw_gemini.json",
    "review_plan_raw_gemini.json",
    "review_video_recheck_raw_gemini.json",
    "review_refine_raw_gemini.json",
    "analysis_raw_gemini.json",
    "observations_raw_gemini.json",
}

job_lock = threading.Lock()
jobs: dict[str, dict[str, Any]] = {}
job_queue: deque[str] = deque()
queued_job_ids: set[str] = set()
queue_condition = threading.Condition()
analysis_slots = threading.BoundedSemaphore(MAX_CONCURRENT_ANALYSES)
active_processes_lock = threading.Lock()
active_processes: dict[str, subprocess.Popen[str]] = {}
cancelled_item_ids_lock = threading.Lock()
cancelled_item_ids: set[str] = set()
filter_jobs_lock = threading.Lock()
filter_jobs: dict[str, dict[str, Any]] = {}
filter_queue: deque[str] = deque()
queued_filter_job_ids: set[str] = set()
filter_queue_condition = threading.Condition()
translation_jobs_lock = threading.Lock()
translation_jobs: dict[str, dict[str, Any]] = {}
translation_queue: deque[str] = deque()
queued_translation_job_ids: set[str] = set()
translation_queue_condition = threading.Condition()

if (SKILL_ROOT / "scripts").exists():
    sys.path.insert(0, str(SKILL_ROOT / "scripts"))

try:
    from docx import Document
except Exception:
    Document = None

try:
    import cv2  # type: ignore
except Exception:
    cv2 = None

try:
    from hybrid_v2_pipeline import (
        PRIMARY_FALLBACK_MODELS,
        SUPPLEMENT_FALLBACK_MODELS,
        enforce_chinese_dialogue_translation,
        run_text_json_prompt_with_fallback,
        run_video_json_prompt_with_fallback,
        unique_models,
    )
except Exception:
    PRIMARY_FALLBACK_MODELS = STABLE_VIDEO_MODELS[:]
    SUPPLEMENT_FALLBACK_MODELS = STABLE_VIDEO_MODELS[:]
    def enforce_chinese_dialogue_translation(script_json: dict, key: str, models: list[str]) -> dict:
        return script_json
    run_text_json_prompt_with_fallback = None
    run_video_json_prompt_with_fallback = None
    def unique_models(*names: str) -> list[str]:
        seen: set[str] = set()
        out: list[str] = []
        for name in names:
            value = str(name or "").strip()
            if value and value not in seen:
                seen.add(value)
                out.append(value)
        return out


def now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


def format_beijing_time(value: object) -> str:
    text = str(value or "").strip()
    if not text:
        return ""
    try:
        normalized = text.replace("Z", "+00:00")
        dt = datetime.fromisoformat(normalized)
        if dt.tzinfo is None:
            dt = dt.replace(tzinfo=timezone.utc)
        return dt.astimezone(BEIJING_TZ).strftime("%Y-%m-%d %H:%M:%S")
    except Exception:
        compact = text[:19].replace("T", " ")
        return compact


def parse_iso_datetime(value: object) -> datetime | None:
    text = str(value or "").strip()
    if not text:
        return None
    try:
        normalized = text.replace("Z", "+00:00")
        dt = datetime.fromisoformat(normalized)
        if dt.tzinfo is None:
            dt = dt.replace(tzinfo=timezone.utc)
        return dt
    except Exception:
        return None


def html_escape(value: object) -> str:
    return html.escape(str(value or ""), quote=True)


def read_json_file(path: Path, *, default: Any, backup_on_error: bool = True) -> Any:
    if not path.exists():
        return default
    try:
        raw = path.read_text(encoding="utf-8")
    except Exception:
        return default
    if not raw.strip():
        if backup_on_error:
            try:
                broken = path.with_suffix(path.suffix + f".empty-{int(time.time())}.bak")
                path.replace(broken)
            except Exception:
                pass
        return default
    try:
        return json.loads(raw)
    except Exception:
        if backup_on_error:
            try:
                broken = path.with_suffix(path.suffix + f".broken-{int(time.time())}.bak")
                path.replace(broken)
            except Exception:
                pass
        return default


def write_text_atomic(path: Path, content: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    temp_path = path.parent / f"{path.name}.{uuid4().hex}.tmp"
    temp_path.write_text(content, encoding="utf-8")
    temp_path.replace(path)


def write_json_atomic(path: Path, payload: Any) -> None:
    write_text_atomic(path, json.dumps(payload, ensure_ascii=False, indent=2))


def env_flag(name: str, default: str = "0") -> bool:
    return str(os.environ.get(name, default)).strip().lower() in {"1", "true", "yes", "on"}


def is_no_space_error(exc: Exception) -> bool:
    return isinstance(exc, OSError) and getattr(exc, "errno", None) == errno.ENOSPC


def log_runtime_warning(event: str, message: str, **extra: Any) -> None:
    payload = {"level": "warning", "event": event, "message": message}
    if extra:
        payload.update(extra)
    try:
        print(json.dumps(payload, ensure_ascii=False))
    except Exception:
        print(f"[warning] {event}: {message}")


def log_runtime_info(event: str, message: str, **extra: Any) -> None:
    payload = {"level": "info", "event": event, "message": message}
    if extra:
        payload.update(extra)
    try:
        print(json.dumps(payload, ensure_ascii=False))
    except Exception:
        print(f"[info] {event}: {message}")


def best_timestamp_from_values(*values: object) -> datetime | None:
    for value in values:
        dt = parse_iso_datetime(value)
        if dt:
            return dt
    return None


def collect_cleanup_metadata() -> dict[str, dict[str, Any]]:
    metadata: dict[str, dict[str, Any]] = {}
    for job_id, job in jobs.items():
        items = job.get("items") or []
        if items:
            for item in items:
                item_id = str(item.get("id") or "").strip()
                if not item_id:
                    continue
                metadata[item_id] = {
                    "status": str(item.get("status") or "").strip(),
                    "reviewed": bool(item.get("reviewed")) or str(item.get("review_status") or "").strip() == "completed",
                    "edited": bool(item.get("edited")),
                    "in_library": bool(item.get("saved_to_library_at")) or library_entry_exists(item_id),
                    "saved_to_library_at": best_timestamp_from_values(item.get("saved_to_library_at")),
                    "updated_at": best_timestamp_from_values(
                        item.get("completed_at"),
                        item.get("updated_at"),
                        item.get("created_at"),
                    ),
                }
        else:
            metadata[job_id] = {
                "status": str(job.get("status") or "").strip(),
                "reviewed": bool(job.get("reviewed")) or str(job.get("review_status") or "").strip() == "completed",
                "edited": bool(job.get("edited")),
                "in_library": bool(job.get("saved_to_library_at")) or library_entry_exists(job_id),
                "saved_to_library_at": best_timestamp_from_values(job.get("saved_to_library_at")),
                "updated_at": best_timestamp_from_values(
                    job.get("completed_at"),
                    job.get("updated_at"),
                    job.get("created_at"),
                ),
            }
    return metadata


def source_video_cleanup_reason(info: dict[str, Any], output_dir: Path, now_dt: datetime) -> str:
    saved_at = info.get("saved_to_library_at")
    if isinstance(saved_at, datetime) or bool(info.get("in_library")):
        return "saved_to_library"
    updated_at = info.get("updated_at")
    if not isinstance(updated_at, datetime):
        try:
            updated_at = datetime.fromtimestamp(output_dir.stat().st_mtime, timezone.utc)
        except Exception:
            updated_at = now_dt
    age_days = max(0.0, (now_dt - updated_at).total_seconds() / 86400.0)
    if age_days > SOURCE_VIDEO_RETENTION_DAYS:
        return "unsaved_retention_expired"
    return ""


def delete_source_video_if_allowed(item_id: str, *, reason: str = "saved_to_library") -> bool:
    source_path = RESULTS_ROOT / item_id / SOURCE_VIDEO_NAME
    if not source_path.exists():
        return False
    try:
        source_path.unlink()
        log_runtime_info("source_video_deleted", "Removed source video after it was no longer needed.", path=str(source_path), reason=reason)
        return True
    except FileNotFoundError:
        return False
    except Exception as exc:
        log_runtime_warning("source_video_delete_failed", "Could not remove source video.", path=str(source_path), reason=reason, error=str(exc))
        return False


def cleanup_old_results(*, now_dt: datetime | None = None) -> dict[str, int]:
    RESULTS_ROOT.mkdir(parents=True, exist_ok=True)
    now_dt = now_dt or datetime.now(timezone.utc)
    metadata = collect_cleanup_metadata()
    cleaned_source = 0
    cleaned_raw = 0
    skipped_running = 0
    for output_dir in RESULTS_ROOT.iterdir():
        if not output_dir.is_dir():
            continue
        item_id = output_dir.name
        info = metadata.get(item_id) or {}
        status = str(info.get("status") or "").strip()
        if status in {"queued", "running"}:
            skipped_running += 1
            continue
        updated_at = info.get("updated_at")
        if not isinstance(updated_at, datetime):
            try:
                updated_at = datetime.fromtimestamp(output_dir.stat().st_mtime, timezone.utc)
            except Exception:
                updated_at = now_dt

        source_path = output_dir / SOURCE_VIDEO_NAME
        source_cleanup_reason = source_video_cleanup_reason(info, output_dir, now_dt)
        if source_path.exists() and source_cleanup_reason:
            try:
                source_path.unlink()
                cleaned_source += 1
            except FileNotFoundError:
                pass
            except Exception as exc:
                log_runtime_warning(
                    "cleanup_source_failed",
                    "Could not remove expired source video.",
                    path=str(source_path),
                    reason=source_cleanup_reason,
                    error=str(exc),
                )

        age_days = max(0.0, (now_dt - updated_at).total_seconds() / 86400.0)
        if age_days > RAW_ARTIFACT_RETENTION_DAYS:
            for name in RAW_ARTIFACT_NAMES:
                raw_path = output_dir / name
                if not raw_path.exists():
                    continue
                try:
                    raw_path.unlink()
                    cleaned_raw += 1
                except FileNotFoundError:
                    pass
                except Exception as exc:
                    log_runtime_warning(
                        "cleanup_raw_failed",
                        "Could not remove expired raw artifact.",
                        path=str(raw_path),
                        error=str(exc),
                    )
    summary = {
        "cleaned_source_videos": cleaned_source,
        "cleaned_raw_artifacts": cleaned_raw,
        "skipped_running_dirs": skipped_running,
    }
    log_runtime_info("cleanup_results_complete", "Finished automatic cleanup of old result artifacts.", **summary)
    return summary


def emergency_cleanup_result_artifacts() -> dict[str, int]:
    RESULTS_ROOT.mkdir(parents=True, exist_ok=True)
    metadata = collect_cleanup_metadata()
    cleaned_source = 0
    cleaned_raw = 0
    skipped_running = 0
    for output_dir in RESULTS_ROOT.iterdir():
        if not output_dir.is_dir():
            continue
        info = metadata.get(output_dir.name) or {}
        if str(info.get("status") or "").strip() in {"queued", "running"}:
            skipped_running += 1
            continue
        source_path = output_dir / SOURCE_VIDEO_NAME
        source_cleanup_reason = source_video_cleanup_reason(info, output_dir, datetime.now(timezone.utc))
        if source_path.exists() and source_cleanup_reason:
            try:
                source_path.unlink()
                cleaned_source += 1
            except Exception as exc:
                log_runtime_warning(
                    "emergency_cleanup_source_failed",
                    "Could not remove source video during emergency cleanup.",
                    path=str(source_path),
                    reason=source_cleanup_reason,
                    error=str(exc),
                )
        for name in RAW_ARTIFACT_NAMES:
            raw_path = output_dir / name
            if not raw_path.exists():
                continue
            try:
                raw_path.unlink()
                cleaned_raw += 1
            except Exception as exc:
                log_runtime_warning("emergency_cleanup_raw_failed", "Could not remove raw artifact during emergency cleanup.", path=str(raw_path), error=str(exc))
    summary = {
        "cleaned_source_videos": cleaned_source,
        "cleaned_raw_artifacts": cleaned_raw,
        "skipped_running_dirs": skipped_running,
    }
    log_runtime_warning("emergency_cleanup_results_complete", "Finished emergency cleanup after disk pressure.", **summary)
    return summary


def collect_tracked_result_ids() -> set[str]:
    tracked: set[str] = set()

    for entry in load_library_entries():
        entry_id = str((entry or {}).get("entry_id") or (entry or {}).get("id") or "").strip()
        if re.fullmatch(r"[0-9a-f]{32}", entry_id):
            tracked.add(entry_id)

    with job_lock:
        for job_id, job in jobs.items():
            if re.fullmatch(r"[0-9a-f]{32}", str(job_id or "").strip()):
                tracked.add(str(job_id).strip())
            for item in job.get("items") or []:
                item_id = str((item or {}).get("id") or "").strip()
                if re.fullmatch(r"[0-9a-f]{32}", item_id):
                    tracked.add(item_id)

    with translation_jobs_lock:
        for job_id in translation_jobs.keys():
            value = str(job_id or "").strip()
            if re.fullmatch(r"[0-9a-f]{32}", value):
                tracked.add(value)

    return tracked


def cleanup_orphan_result_dirs() -> dict[str, int]:
    RESULTS_ROOT.mkdir(parents=True, exist_ok=True)
    tracked_ids = collect_tracked_result_ids()
    deleted_dirs = 0
    freed_bytes = 0
    skipped_non_job_dirs = 0

    for output_dir in RESULTS_ROOT.iterdir():
        if not output_dir.is_dir():
            continue
        item_id = str(output_dir.name or "").strip()
        if not re.fullmatch(r"[0-9a-f]{32}", item_id):
            skipped_non_job_dirs += 1
            continue
        if item_id in tracked_ids:
            continue
        dir_bytes = 0
        for path in output_dir.rglob("*"):
            if path.is_file():
                try:
                    dir_bytes += path.stat().st_size
                except FileNotFoundError:
                    continue
        shutil.rmtree(output_dir, ignore_errors=True)
        deleted_dirs += 1
        freed_bytes += dir_bytes

    summary = {
        "deleted_dirs": deleted_dirs,
        "freed_mb": round(freed_bytes / 1024 / 1024, 2),
        "tracked_ids": len(tracked_ids),
        "skipped_non_job_dirs": skipped_non_job_dirs,
    }
    if deleted_dirs:
        log_runtime_warning("orphan_results_cleaned", "Removed orphaned result directories that are no longer tracked by the library or jobs.", **summary)
    else:
        log_runtime_info("orphan_results_clean", "No orphaned result directories needed cleanup.", **summary)
    return summary


def load_jobs() -> None:
    global jobs
    DATA_ROOT.mkdir(parents=True, exist_ok=True)
    RESULTS_ROOT.mkdir(parents=True, exist_ok=True)
    jobs = read_json_file(JOBS_FILE, default={})
    if not isinstance(jobs, dict):
        jobs = {}
    try:
        emergency_cleanup_result_artifacts()
    except Exception as exc:
        log_runtime_warning("emergency_cleanup_results_skipped", "Emergency result cleanup failed during startup.", error=str(exc))
    try:
        backfill_completed_jobs()
    except Exception as exc:
        if is_no_space_error(exc):
            log_runtime_warning("backfill_jobs_skipped", "Skipped completed job backfill because the persistent disk is full.", path=str(JOBS_FILE), error=str(exc))
        else:
            raise
    try:
        cleanup_old_results()
    except Exception as exc:
        log_runtime_warning("cleanup_results_skipped", "Automatic result cleanup failed during startup.", error=str(exc))
    log_runtime_info("library_sync_startup_skipped", "Script library sync is manual-only.")


def save_jobs() -> None:
    write_json_atomic(JOBS_FILE, jobs)


def load_filter_jobs() -> None:
    global filter_jobs
    DATA_ROOT.mkdir(parents=True, exist_ok=True)
    filter_jobs = read_json_file(FILTER_JOBS_FILE, default={})
    if not isinstance(filter_jobs, dict):
        filter_jobs = {}


def save_filter_jobs() -> None:
    write_json_atomic(FILTER_JOBS_FILE, filter_jobs)


def load_translation_jobs() -> None:
    global translation_jobs
    DATA_ROOT.mkdir(parents=True, exist_ok=True)
    RESULTS_ROOT.mkdir(parents=True, exist_ok=True)
    translation_jobs = read_json_file(TRANSLATION_JOBS_FILE, default={})
    if not isinstance(translation_jobs, dict):
        translation_jobs = {}


def save_translation_jobs() -> None:
    write_json_atomic(TRANSLATION_JOBS_FILE, translation_jobs)


def enqueue_job(job_id: str) -> None:
    with queue_condition:
        if job_id in queued_job_ids:
            return
        queued_job_ids.add(job_id)
        job_queue.append(job_id)
        queue_condition.notify()


def register_active_process(item_id: str, proc: subprocess.Popen[str]) -> None:
    with active_processes_lock:
        active_processes[item_id] = proc


def unregister_active_process(item_id: str) -> None:
    with active_processes_lock:
        active_processes.pop(item_id, None)


def mark_item_cancelled(item_id: str) -> None:
    with cancelled_item_ids_lock:
        cancelled_item_ids.add(item_id)


def clear_item_cancelled(item_id: str) -> None:
    with cancelled_item_ids_lock:
        cancelled_item_ids.discard(item_id)


def is_item_cancelled(item_id: str) -> bool:
    with cancelled_item_ids_lock:
        return item_id in cancelled_item_ids


def has_active_process(item_id: str) -> bool:
    with active_processes_lock:
        proc = active_processes.get(item_id)
        return bool(proc and proc.poll() is None)


def item_output_ready(item: dict[str, Any]) -> bool:
    item_id = str(item.get("id") or "").strip()
    if not item_id:
        return False
    if item.get("saved_to_library_at"):
        return True
    if item.get("result_json") and (
        str(item.get("html_url") or "").strip()
        or str(item.get("report_url") or "").strip()
        or str(item.get("docx_url") or "").strip()
    ):
        return True
    output_dir = RESULTS_ROOT / item_id
    final_json = output_dir / "script_table.json"
    final_html = output_dir / "script_table.html"
    final_docx = output_dir / "script_export.docx"
    product_report = output_dir / "product_report.html"
    if final_json.exists() and final_html.exists():
        return True
    if final_json.exists() and (final_docx.exists() or product_report.exists()):
        return True
    if final_json.exists() and library_entry_exists(item_id):
        return True
    return False


def hydrate_item_from_outputs(item: dict[str, Any]) -> bool:
    item_id = str(item.get("id") or "").strip()
    if not item_id:
        return False
    output_dir = RESULTS_ROOT / item_id
    script_json_path = output_dir / "script_table.json"
    script_html_path = output_dir / "script_table.html"
    docx_path = output_dir / "script_export.docx"
    report_path = output_dir / "product_report.html"
    evidence_path = output_dir / "evidence_bundle.json"
    if not script_json_path.exists():
        return False
    script_json = read_json(script_json_path) or read_json(output_dir / "analysis_result.json")
    if not script_json:
        return False
    item["result_json"] = item.get("result_json") or script_json
    item["zh_result_json"] = item.get("zh_result_json") or script_json
    item["original_result_json"] = item.get("original_result_json") or script_json
    if script_html_path.exists():
        item["html_url"] = item.get("html_url") or f"/results/{item_id}/script_table.html"
        item["zh_html_url"] = item.get("zh_html_url") or item["html_url"]
    if docx_path.exists():
        item["docx_url"] = item.get("docx_url") or f"/results/{item_id}/{docx_path.name}"
        item["zh_docx_url"] = item.get("zh_docx_url") or item["docx_url"]
    if report_path.exists():
        item["report_url"] = item.get("report_url") or f"/results/{item_id}/product_report.html"
    if evidence_path.exists():
        item["evidence_url"] = item.get("evidence_url") or f"/results/{item_id}/evidence_bundle.json"
    item["artifacts"] = item.get("artifacts") or summarize_artifacts(item_id, output_dir)
    item["display_language"] = item.get("display_language") or "zh"
    item["title"] = item.get("title") or script_json.get("title") or "Video Script"
    item["status"] = "completed"
    item["stage"] = "completed"
    item["stage_message"] = "Completed."
    item["completed_at"] = item.get("completed_at") or now_iso()
    item["updated_at"] = now_iso()
    return True


def recompute_job_status(job_id: str) -> None:
    job = jobs.get(job_id)
    if not job:
        return
    items = job.get("items") or []
    if not items:
        return
    changed = False
    for item in items:
        status = str(item.get("status") or "").strip()
        if status in {"queued", "running"} and item_output_ready(item):
            changed = hydrate_item_from_outputs(item) or changed
    statuses = [str(item.get("status") or "").strip() for item in items]
    review_statuses = [str(item.get("review_status") or "").strip() for item in items]
    completed_count = sum(1 for status in statuses if status == "completed")
    failed_count = sum(1 for status in statuses if status == "failed")
    previous_status = str(job.get("status") or "").strip()
    previous_stage = str(job.get("stage") or "").strip()
    previous_message = str(job.get("stage_message") or "").strip()
    if any(status == "running" for status in statuses) or any(status == "running" for status in review_statuses):
        job["status"] = "running"
    elif any(status == "queued" for status in statuses):
        job["status"] = "queued"
    elif completed_count == len(items):
        job["status"] = "completed"
        job["completed_at"] = job.get("completed_at") or now_iso()
    elif completed_count > 0 and completed_count + failed_count == len(items):
        job["status"] = "completed"
        job["completed_at"] = job.get("completed_at") or now_iso()
    else:
        job["status"] = "failed"
        job["completed_at"] = now_iso()
    if job["status"] == "completed":
        job["stage"] = "completed"
        job["stage_message"] = f"Completed {completed_count}/{len(items)} items. Failed {failed_count}."
    elif job["status"] == "failed":
        job["stage"] = "failed"
        job["stage_message"] = job.get("stage_message") or "All batch items failed."
    job["updated_at"] = now_iso()
    if (
        changed
        or previous_status != job.get("status")
        or previous_stage != job.get("stage")
        or previous_message != job.get("stage_message")
    ):
        save_jobs()


def library_entry_exists(entry_id: str) -> bool:
    target = str(entry_id or "").strip()
    if not target or not LIBRARY_FILE.exists():
        return False
    try:
        data = load_library_entries()
    except Exception:
        return False
    for entry in data:
        if str((entry or {}).get("entry_id") or "").strip() == target:
            return True
    return False


def reconcile_stale_jobs() -> None:
    now_dt = datetime.now(timezone.utc)
    changed = False
    killed_item_ids: list[str] = []
    requeue_job_ids: list[str] = []
    with queue_condition:
        queued_snapshot = set(job_queue)
        queued_job_ids.intersection_update(queued_snapshot)
    with job_lock:
        for job_id, job in jobs.items():
            items = job.get("items") or []
            job_changed = False
            for item in items:
                item_id = str(item.get("id") or "")
                status = str(item.get("status") or "").strip()
                updated_at = parse_iso_datetime(item.get("updated_at") or item.get("created_at"))
                if status == "running" and updated_at:
                    stale_for = (now_dt - updated_at).total_seconds()
                    if item_output_ready(item):
                        hydrate_item_from_outputs(item)
                        changed = True
                        job_changed = True
                    elif not has_active_process(item_id) and stale_for > PROCESSLESS_TASK_STALE_SEC:
                        item["status"] = "failed"
                        item["stage"] = "failed"
                        item["stage_message"] = "Worker stopped unexpectedly."
                        item["error"] = "后台执行中断，任务已自动停止。"
                        item["completed_at"] = now_iso()
                        item["updated_at"] = now_iso()
                        changed = True
                        job_changed = True
                    elif stale_for > RUNNING_TASK_STALE_SEC:
                        item["status"] = "failed"
                        item["stage"] = "failed"
                        item["stage_message"] = "Stopped after no progress."
                        item["error"] = "任务长时间没有进展，已自动停止。"
                        item["completed_at"] = now_iso()
                        item["updated_at"] = now_iso()
                        mark_item_cancelled(item_id)
                        killed_item_ids.append(item_id)
                        changed = True
                        job_changed = True
                review_status = str(item.get("review_status") or "").strip()
                if review_status == "running" and updated_at:
                    stale_for = (now_dt - updated_at).total_seconds()
                    if stale_for > REVIEW_TASK_STALE_SEC:
                        item["review_status"] = "failed"
                        item["review_stage"] = "failed"
                        item["review_message"] = "复盘长时间没有进展，已自动停止。"
                        item["updated_at"] = now_iso()
                        mark_item_cancelled(item_id)
                        changed = True
                        job_changed = True
            has_running = any(str(item.get("status") or "").strip() == "running" for item in items)
            has_running_review = any(str(item.get("review_status") or "").strip() == "running" for item in items)
            has_queued = any(str(item.get("status") or "").strip() == "queued" for item in items)
            for item in items:
                if str(item.get("status") or "").strip() in {"queued", "running"} and item_output_ready(item):
                    if hydrate_item_from_outputs(item):
                        changed = True
                        job_changed = True
            if has_queued and not has_running and not has_running_review and job_id not in queued_snapshot:
                job["status"] = "queued"
                job["stage"] = "queued"
                job["stage_message"] = "Recovered queued task."
                job["updated_at"] = now_iso()
                requeue_job_ids.append(job_id)
                changed = True
                job_changed = True
            if job_changed:
                recompute_job_status(job_id)
        if changed:
            save_jobs()
    for job_id in requeue_job_ids:
        enqueue_job(job_id)
    if killed_item_ids:
        with active_processes_lock:
            for item_id in killed_item_ids:
                proc = active_processes.get(item_id)
                if proc and proc.poll() is None:
                    try:
                        proc.kill()
                    except Exception:
                        pass


def watchdog_loop() -> None:
    while True:
        try:
            reconcile_stale_jobs()
        except Exception as exc:
            log_runtime_warning("watchdog_failed", "Background task watchdog failed.", error=str(exc))
        time.sleep(max(5, WATCHDOG_INTERVAL_SEC))


def stop_all_tasks() -> dict[str, int]:
    stopped_jobs = 0
    stopped_items = 0
    stopped_reviews = 0
    killed_item_ids: list[str] = []
    with queue_condition:
        job_queue.clear()
        queued_job_ids.clear()
    with job_lock:
        for job_id, job in jobs.items():
            items = job.get("items") or []
            job_touched = False
            for item in items:
                item_id = str(item.get("id") or "")
                if str(item.get("status") or "") in {"queued", "running"}:
                    item["status"] = "failed"
                    item["stage"] = "failed"
                    item["stage_message"] = "Stopped manually."
                    item["error"] = "任务已手动停止。"
                    item["completed_at"] = now_iso()
                    item["updated_at"] = now_iso()
                    mark_item_cancelled(item_id)
                    killed_item_ids.append(item_id)
                    stopped_items += 1
                    job_touched = True
                if str(item.get("review_status") or "") == "running":
                    item["review_status"] = "failed"
                    item["review_stage"] = "failed"
                    item["review_message"] = "复盘已手动停止。"
                    item["updated_at"] = now_iso()
                    mark_item_cancelled(item_id)
                    stopped_reviews += 1
                    job_touched = True
            if job_touched:
                recompute_job_status(job_id)
                job["error"] = "任务已手动停止。"
                stopped_jobs += 1
        save_jobs()
    with active_processes_lock:
        for item_id in killed_item_ids:
            proc = active_processes.get(item_id)
            if proc and proc.poll() is None:
                try:
                    proc.kill()
                except Exception:
                    pass
    return {
        "stopped_jobs": stopped_jobs,
        "stopped_items": stopped_items,
        "stopped_reviews": stopped_reviews,
    }


def restore_pending_jobs_to_queue() -> None:
    pending_job_ids: list[str] = []
    cleared_stale = 0
    changed = False
    now_dt = datetime.now(timezone.utc)
    with job_lock:
        for job_id, job in jobs.items():
            items = job.get("items") or []
            if not items:
                continue
            pending = False
            completed_count = 0
            failed_count = 0
            for item in items:
                status = str(item.get("status") or "").strip()
                if status == "completed":
                    completed_count += 1
                    continue
                if status == "failed":
                    failed_count += 1
                    continue
                updated_at = parse_iso_datetime(item.get("updated_at") or item.get("created_at") or job.get("updated_at") or job.get("created_at"))
                age_sec = (now_dt - updated_at).total_seconds() if updated_at else RESTORE_PENDING_MAX_AGE_SEC + 1
                if age_sec > RESTORE_PENDING_MAX_AGE_SEC:
                    item["status"] = "failed"
                    item["stage"] = "failed"
                    item["stage_message"] = "Cleared stale queued task during service startup."
                    item["error"] = "旧任务在服务重启前已长时间停留在队列/运行状态，已自动清理。"
                    item["completed_at"] = now_iso()
                    item["updated_at"] = now_iso()
                    failed_count += 1
                    cleared_stale += 1
                    changed = True
                    continue
                pending = True
                if status != "queued":
                    item["status"] = "queued"
                    item["stage"] = "queued"
                    item["stage_message"] = "Queued after service restart."
                    item["updated_at"] = now_iso()
                    changed = True
            if pending:
                job["status"] = "queued"
                job["stage"] = "queued"
                job["stage_message"] = "Queued after service restart."
                job["updated_at"] = now_iso()
                pending_job_ids.append(job_id)
                changed = True
            elif items and completed_count == len(items):
                if job.get("status") != "completed":
                    job["status"] = "completed"
                    job["stage"] = "completed"
                    job["stage_message"] = f"Completed {completed_count}/{len(items)} items. Failed {failed_count}."
                    changed = True
            elif items and failed_count == len(items):
                if job.get("status") != "failed":
                    job["status"] = "failed"
                    job["stage"] = "failed"
                    job["stage_message"] = "All batch items failed."
                    changed = True
        if cleared_stale:
            log_runtime_warning(
                "stale_pending_jobs_cleared",
                "Cleared stale analysis tasks instead of restoring them to the queue.",
                cleared_items=cleared_stale,
                max_age_sec=RESTORE_PENDING_MAX_AGE_SEC,
            )
        if changed:
            try:
                save_jobs()
            except Exception as exc:
                if not is_no_space_error(exc):
                    raise
                log_runtime_warning(
                    "restore_pending_jobs_save_skipped",
                    "Skipped saving restored analysis jobs because the persistent disk is full.",
                    error=str(exc),
                )
    for job_id in pending_job_ids:
        enqueue_job(job_id)


def restore_pending_filter_jobs_to_queue() -> None:
    pending_job_ids: list[str] = []
    changed = False
    with filter_jobs_lock:
        for job_id, job in filter_jobs.items():
            items = job.get("items") or []
            if not items:
                continue
            pending = False
            for item in items:
                status = str(item.get("status") or "").strip()
                if status in {"completed", "failed"}:
                    continue
                pending = True
                if status != "queued":
                    item["status"] = "queued"
                    item["stage"] = "queued"
                    item["stage_message"] = "Queued after service restart."
                    item["updated_at"] = now_iso()
                    changed = True
            if pending:
                job["status"] = "queued"
                job["stage"] = "queued"
                job["stage_message"] = "Queued after service restart."
                job["updated_at"] = now_iso()
                pending_job_ids.append(job_id)
                changed = True
        if changed:
            try:
                save_filter_jobs()
            except Exception as exc:
                if not is_no_space_error(exc):
                    raise
                log_runtime_warning(
                    "restore_pending_filter_jobs_save_skipped",
                    "Skipped saving restored filter jobs because the persistent disk is full.",
                    error=str(exc),
                )
    for job_id in pending_job_ids:
        with filter_queue_condition:
            if job_id not in queued_filter_job_ids:
                queued_filter_job_ids.add(job_id)
                filter_queue.append(job_id)
                filter_queue_condition.notify()


def restore_pending_translation_jobs_to_queue() -> None:
    pending_job_ids: list[str] = []
    changed = False
    with translation_jobs_lock:
        for job_id, job in translation_jobs.items():
            status = str(job.get("status") or "").strip()
            if status in {"completed", "failed"}:
                continue
            if status != "queued":
                job["status"] = "queued"
                job["stage"] = "queued"
                job["stage_message"] = "Queued after service restart."
                job["updated_at"] = now_iso()
                changed = True
            pending_job_ids.append(job_id)
        if changed:
            save_translation_jobs()
    for job_id in pending_job_ids:
        with translation_queue_condition:
            if job_id not in queued_translation_job_ids:
                queued_translation_job_ids.add(job_id)
                translation_queue.append(job_id)
                translation_queue_condition.notify()


def job_worker_loop() -> None:
    while True:
        try:
            with queue_condition:
                while not job_queue:
                    queue_condition.wait()
                job_id = job_queue.popleft()
                queued_job_ids.discard(job_id)
            with job_lock:
                job_exists = job_id in jobs
            if not job_exists:
                log_runtime_warning("analysis_worker_skipped_missing_job", "Skipped queued analysis job because it no longer exists.", job_id=job_id)
                continue
            with analysis_slots:
                run_job_batch(job_id)
        except Exception as exc:
            log_runtime_warning("analysis_worker_loop_error", "Analysis worker recovered after an unexpected error.", error=str(exc))
            time.sleep(1)


def start_job_workers() -> None:
    for index in range(MAX_CONCURRENT_ANALYSES):
        thread = threading.Thread(target=job_worker_loop, name=f"koko-job-worker-{index+1}", daemon=True)
        thread.start()


def start_watchdog() -> None:
    threading.Thread(target=watchdog_loop, name="koko-task-watchdog", daemon=True).start()


def read_json(path: Path) -> dict[str, Any]:
    data = read_json_file(path, default={})
    return data if isinstance(data, dict) else {}


def load_library_entries() -> list[dict[str, Any]]:
    data = read_json_file(LIBRARY_FILE, default=[])
    if not isinstance(data, list):
        return []
    for entry in data:
        if not isinstance(entry, dict):
            continue
        content_type_text = " ".join(str(entry.get(key) or "") for key in ["title", "whole_video_summary", "summary", "content_type_reasoning", "video_url"])
        entry["content_type"] = normalize_creator_content_type(entry.get("content_type"), content_type_text)
        source = str(entry.get("content_type_source") or "").strip().lower()
        entry["content_type_source"] = source if source in {"auto", "manual"} else "auto"
        storyboard_url = library_storyboard_cover_url(entry)
        if storyboard_url:
            entry["storyboard_cover_url"] = storyboard_url
        if not entry.get("duration_bucket"):
            entry.update(creator_duration_fields(str(entry.get("entry_id") or "")))
        if str(entry.get("location_tag") or "").strip() not in LOCATION_TAGS:
            entry.update(infer_location_tag_fields(entry))
        elif not entry.get("location_tag_pt"):
            entry["location_tag_pt"] = LOCATION_TAG_PT.get(str(entry.get("location_tag") or "").strip(), "")
        share_url = creator_script_share_url(entry.get("entry_id"))
        if share_url:
            entry["creator_share_url"] = share_url
    return data


def library_entry_by_id(entry_id: str) -> dict[str, Any] | None:
    target = str(entry_id or "").strip()
    if not target:
        return None
    for entry in load_library_entries():
        if str((entry or {}).get("entry_id") or "").strip() == target:
            return dict(entry)
    return None


def infer_library_display_language(entry: dict[str, Any], zh_script: dict[str, Any], pt_script: dict[str, Any]) -> str:
    explicit = str(entry.get("display_language") or "").strip().lower()
    if explicit in {"zh", "pt"}:
        return explicit
    html_url = str(entry.get("html_url") or "").strip()
    pt_html_url = str(entry.get("pt_html_url") or "").strip()
    zh_html_url = str(entry.get("zh_html_url") or "").strip()
    if pt_script and html_url and pt_html_url and html_url == pt_html_url:
        return "pt"
    if zh_script and html_url and zh_html_url and html_url == zh_html_url:
        return "zh"
    if pt_script and not zh_script:
        return "pt"
    return "zh"


def save_library_entries(entries: list[dict[str, Any]]) -> bool:
    try:
        write_json_atomic(LIBRARY_FILE, entries)
        return True
    except Exception as exc:
        if is_no_space_error(exc):
            log_runtime_warning(
                "library_write_skipped",
                "Could not save script library because the persistent disk is full.",
                path=str(LIBRARY_FILE),
                entries=len(entries),
            )
            return False
        raise


def load_error_case_entries() -> list[dict[str, Any]]:
    data = read_json_file(ERROR_CASE_LIBRARY_FILE, default=[])
    if not isinstance(data, list):
        return []
    return [entry for entry in data if isinstance(entry, dict)]


def save_error_case_entries(entries: list[dict[str, Any]]) -> bool:
    try:
        write_json_atomic(ERROR_CASE_LIBRARY_FILE, entries)
        return True
    except Exception as exc:
        if is_no_space_error(exc):
            log_runtime_warning(
                "error_case_write_skipped",
                "Could not save error case library because the persistent disk is full.",
                path=str(ERROR_CASE_LIBRARY_FILE),
                entries=len(entries),
            )
            return False
        raise


def split_video_urls(raw: str) -> list[str]:
    return unique_urls_from_values(extract_http_urls(raw))


def sanitize_analysis_prompt(value: Any, limit: int = 1200) -> str:
    prompt = re.sub(r"\s+", " ", str(value or "").strip())
    return prompt[:limit]


def extract_http_urls(raw: str) -> list[str]:
    text = str(raw or "")
    candidates = re.findall(r"https?://[^\s<>'\"）)\]}]+", text, flags=re.IGNORECASE)
    cleaned: list[str] = []
    for candidate in candidates:
        value = candidate.strip().rstrip(".,;!?)]}>")
        if value:
            cleaned.append(value)
    return cleaned


def unique_urls_from_values(values: list[str]) -> list[str]:
    seen: set[str] = set()
    ordered: list[str] = []
    for value in values:
        url = str(value or "").strip()
        if not url or url in seen:
            continue
        seen.add(url)
        ordered.append(url)
    return ordered


def text_from_tabular_rows(rows: list[list[str]]) -> str:
    return "\n".join("\t".join(cell for cell in row if cell) for row in rows if any(cell for cell in row))


def parse_csv_like_bytes(blob: bytes, *, delimiter: str | None = None) -> str:
    for encoding in ("utf-8", "utf-8-sig", "gb18030", "latin-1"):
        try:
            text = blob.decode(encoding)
            break
        except Exception:
            text = ""
    if not text:
        return ""
    sample = text[:4096]
    if delimiter is None:
        try:
            dialect = csv.Sniffer().sniff(sample, delimiters=",\t;")
            delimiter = dialect.delimiter
        except Exception:
            delimiter = ","
    try:
        reader = csv.reader(io.StringIO(text), delimiter=delimiter)
        rows = [[str(cell or "").strip() for cell in row] for row in reader]
        return text_from_tabular_rows(rows)
    except Exception:
        return text


def parse_xlsx_bytes(blob: bytes) -> str:
    ns = {"main": "http://schemas.openxmlformats.org/spreadsheetml/2006/main"}
    shared_strings: list[str] = []
    rows: list[list[str]] = []
    with zipfile.ZipFile(io.BytesIO(blob)) as archive:
        if "xl/sharedStrings.xml" in archive.namelist():
            root = ET.fromstring(archive.read("xl/sharedStrings.xml"))
            for si in root.findall("main:si", ns):
                text_parts = [node.text or "" for node in si.findall(".//main:t", ns)]
                shared_strings.append("".join(text_parts).strip())
        worksheet_names = sorted(name for name in archive.namelist() if name.startswith("xl/worksheets/sheet") and name.endswith(".xml"))
        for worksheet_name in worksheet_names:
            root = ET.fromstring(archive.read(worksheet_name))
            for row in root.findall(".//main:sheetData/main:row", ns):
                cells: list[str] = []
                for cell in row.findall("main:c", ns):
                    cell_type = cell.attrib.get("t", "")
                    value_node = cell.find("main:v", ns)
                    inline_node = cell.find("main:is", ns)
                    value = ""
                    if cell_type == "inlineStr" and inline_node is not None:
                        value = "".join(node.text or "" for node in inline_node.findall(".//main:t", ns)).strip()
                    elif value_node is not None and value_node.text is not None:
                        raw_value = value_node.text.strip()
                        if cell_type == "s":
                            try:
                                value = shared_strings[int(raw_value)]
                            except Exception:
                                value = raw_value
                        else:
                            value = raw_value
                    cells.append(value)
                rows.append(cells)
    return text_from_tabular_rows(rows)


def xlsx_col_to_index(ref: str) -> int:
    letters = re.sub(r"[^A-Z]", "", str(ref or "").upper())
    total = 0
    for char in letters:
        total = total * 26 + (ord(char) - ord("A") + 1)
    return max(0, total - 1)


def parse_xlsx_workbook_rows(blob: bytes) -> list[dict[str, Any]]:
    ns = {
        "main": "http://schemas.openxmlformats.org/spreadsheetml/2006/main",
        "rel": "http://schemas.openxmlformats.org/package/2006/relationships",
        "docrel": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    }
    shared_strings: list[str] = []
    sheets: list[tuple[str, str]] = []
    with zipfile.ZipFile(io.BytesIO(blob)) as archive:
        names = set(archive.namelist())
        if "xl/sharedStrings.xml" in names:
            root = ET.fromstring(archive.read("xl/sharedStrings.xml"))
            for si in root.findall("main:si", ns):
                text_parts = [node.text or "" for node in si.findall(".//main:t", ns)]
                shared_strings.append("".join(text_parts).strip())
        rels: dict[str, str] = {}
        if "xl/_rels/workbook.xml.rels" in names:
            rel_root = ET.fromstring(archive.read("xl/_rels/workbook.xml.rels"))
            for rel in rel_root.findall("rel:Relationship", ns):
                target = rel.attrib.get("Target") or ""
                if target and not target.startswith("/"):
                    target = "xl/" + target.lstrip("/")
                rels[rel.attrib.get("Id") or ""] = target
        if "xl/workbook.xml" in names:
            wb_root = ET.fromstring(archive.read("xl/workbook.xml"))
            for sheet in wb_root.findall(".//main:sheet", ns):
                sheet_name = sheet.attrib.get("name") or "Sheet"
                rel_id = sheet.attrib.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id") or ""
                sheet_path = rels.get(rel_id) or ""
                if sheet_path in names:
                    sheets.append((sheet_name, sheet_path))
        if not sheets:
            sheets = [
                (Path(name).stem, name)
                for name in sorted(names)
                if name.startswith("xl/worksheets/sheet") and name.endswith(".xml")
            ]
        parsed: list[dict[str, Any]] = []
        for sheet_name, worksheet_name in sheets:
            root = ET.fromstring(archive.read(worksheet_name))
            rows: list[list[str]] = []
            for row in root.findall(".//main:sheetData/main:row", ns):
                cells: dict[int, str] = {}
                for cell in row.findall("main:c", ns):
                    col_index = xlsx_col_to_index(cell.attrib.get("r") or "")
                    cell_type = cell.attrib.get("t", "")
                    value_node = cell.find("main:v", ns)
                    inline_node = cell.find("main:is", ns)
                    value = ""
                    if cell_type == "inlineStr" and inline_node is not None:
                        value = "".join(node.text or "" for node in inline_node.findall(".//main:t", ns)).strip()
                    elif value_node is not None and value_node.text is not None:
                        raw_value = value_node.text.strip()
                        if cell_type == "s":
                            try:
                                value = shared_strings[int(raw_value)]
                            except Exception:
                                value = raw_value
                        else:
                            value = raw_value
                    cells[col_index] = value
                width = max(cells.keys(), default=-1) + 1
                rows.append([str(cells.get(idx, "") or "").strip() for idx in range(width)])
            parsed.append({"sheet": sheet_name, "rows": rows})
    return parsed


def compact_cell_text(value: object) -> str:
    return re.sub(r"\s+", " ", str(value or "").strip())


def normalized_pt_label(value: object) -> str:
    text = compact_cell_text(value).lower()
    text = re.sub(r"[：:*()（）]", " ", text)
    text = re.sub(r"\s+", " ", text).strip()
    replacements = str.maketrans("áàãâéêíóôõúç", "aaaaeeiooouc")
    return text.translate(replacements)


def row_cell(row: list[str], index: int) -> str:
    return compact_cell_text(row[index]) if 0 <= index < len(row) else ""


def row_value_after(row: list[str], index: int) -> str:
    values: list[str] = []
    seen: set[str] = set()
    for cell in row[index + 1:]:
        value = compact_cell_text(cell)
        if not value or value in seen:
            continue
        seen.add(value)
        values.append(value)
    return " ".join(values).strip()


def split_import_cards(text: str, fallback_label: str) -> list[dict[str, str]]:
    chunks = [compact_cell_text(part) for part in re.split(r"(?=\d+\.\s*)", text or "") if compact_cell_text(part)]
    cards: list[dict[str, str]] = []
    for raw in chunks or [compact_cell_text(text)]:
        item = re.sub(r"^\d+\.\s*", "", raw).strip()
        if not item:
            continue
        if ":" in item:
            label, body = item.split(":", 1)
        elif "—" in item:
            label, body = item.split("—", 1)
        else:
            words = item.split()
            label = " ".join(words[:4]) or fallback_label
            body = " ".join(words[4:]) or item
        cards.append({"label": compact_cell_text(label)[:80] or fallback_label, "text": compact_cell_text(body) or item})
    return cards[:6]


def imported_title_from_summary(summary: str) -> str:
    title = re.sub(r"^O vídeo\s+(mostra|começa|registra|apresenta)\s+", "", compact_cell_text(summary), flags=re.I)
    return title[:112].rstrip(" ,.;") or "Roteiro importado"


def parse_creator_script_tables_from_xlsx(blob: bytes) -> list[dict[str, Any]]:
    workbook = parse_xlsx_workbook_rows(blob)
    scripts: list[dict[str, Any]] = []
    for sheet in workbook:
        rows = sheet.get("rows") or []
        r = 0
        while r < len(rows):
            row = rows[r]
            start_col = next((idx for idx, cell in enumerate(row) if normalized_pt_label(cell) == "video original"), -1)
            if start_col < 0:
                r += 1
                continue
            metadata: dict[str, str] = {}
            table_header_row = -1
            table_cols: dict[str, int] = {}
            cursor = r
            while cursor < len(rows):
                current = rows[cursor]
                labels = [normalized_pt_label(cell) for cell in current]
                if cursor > r and any(label == "video original" for label in labels):
                    break
                for idx, label in enumerate(labels):
                    if label == "video original":
                        metadata["video_url"] = row_value_after(current, idx)
                    elif label == "conteudo principal":
                        metadata["summary"] = row_value_after(current, idx)
                    elif label == "pontos principais":
                        metadata["points"] = row_value_after(current, idx)
                    elif label.startswith("partes que podem ser adaptadas"):
                        metadata["adaptable"] = row_value_after(current, idx)
                if "tempo" in labels and any(label in {"imagem", "acoes"} or label.startswith("dialog") for label in labels):
                    table_header_row = cursor
                    for idx, label in enumerate(labels):
                        if label == "tempo":
                            table_cols["time"] = idx
                        elif label == "imagem":
                            table_cols["visual_content"] = idx
                        elif label == "acoes":
                            table_cols["action"] = idx
                        elif label.startswith("dialog"):
                            table_cols["dialogue"] = idx
                    cursor += 1
                    break
                cursor += 1
            if table_header_row < 0:
                r += 1
                continue
            script_rows: list[dict[str, str]] = []
            cursor = table_header_row + 1
            while cursor < len(rows):
                current = rows[cursor]
                labels = [normalized_pt_label(cell) for cell in current]
                if any(label == "video original" for label in labels):
                    break
                time_value = row_cell(current, table_cols.get("time", -1))
                visual = row_cell(current, table_cols.get("visual_content", -1))
                action = row_cell(current, table_cols.get("action", -1))
                dialogue = row_cell(current, table_cols.get("dialogue", -1))
                if not any([time_value, visual, action, dialogue]):
                    cursor += 1
                    continue
                if normalized_pt_label(time_value) in {"tempo", "conteudo principal", "pontos principais"}:
                    break
                script_rows.append({
                    "time": time_value,
                    "visual_content": visual,
                    "action": action,
                    "dialogue": dialogue,
                })
                cursor += 1
            if metadata.get("video_url") and metadata.get("summary") and script_rows:
                stable_key = f"{metadata.get('video_url')}|{metadata.get('summary')}|{len(script_rows)}"
                item_id = hashlib.md5(stable_key.encode("utf-8")).hexdigest()
                summary = metadata.get("summary", "")
                scripts.append({
                    "id": item_id,
                    "sheet": sheet.get("sheet") or "",
                    "row": r + 1,
                    "video_url": metadata.get("video_url", ""),
                    "script": {
                        "title": imported_title_from_summary(summary),
                        "whole_video_summary": summary,
                        "core_viral_points": split_import_cards(metadata.get("points", ""), "Ponto-chave"),
                        "replaceable_parts": split_import_cards(metadata.get("adaptable", ""), "Plano de substituição"),
                        "rows": script_rows,
                    },
                })
            r = max(cursor, r + 1)
    deduped: list[dict[str, Any]] = []
    seen: set[str] = set()
    for item in scripts:
        if item["id"] in seen:
            continue
        deduped.append(item)
        seen.add(item["id"])
    return deduped


def extract_urls_from_uploaded_file(filename: str, file_b64: str) -> list[str]:
    name = str(filename or "").strip().lower()
    raw_b64 = str(file_b64 or "").strip()
    if not raw_b64:
        return []
    if "," in raw_b64 and raw_b64.startswith("data:"):
        raw_b64 = raw_b64.split(",", 1)[1]
    blob = base64.b64decode(raw_b64)
    text = ""
    if name.endswith(".xlsx"):
        text = parse_xlsx_bytes(blob)
    elif name.endswith(".tsv"):
        text = parse_csv_like_bytes(blob, delimiter="\t")
    elif name.endswith(".csv"):
        text = parse_csv_like_bytes(blob, delimiter=",")
    else:
        text = parse_csv_like_bytes(blob)
    return unique_urls_from_values(extract_http_urls(text))


def kwai_candidate_from_url(url: str) -> bool:
    try:
        parsed = urllib.parse.urlparse(str(url or "").strip())
    except Exception:
        return False
    host = (parsed.hostname or "").lower()
    return host.endswith("kwai.com") or host.endswith("k.kwai.com")


def fetch_remote_text(url: str, *, timeout: int = 20) -> str:
    req = urllib.request.Request(
        url,
        headers={
            "User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0 Safari/537.36",
            "Accept-Language": "en-US,en;q=0.9,zh-CN;q=0.8",
        },
    )
    with urllib.request.urlopen(req, timeout=timeout) as response:
        charset = response.headers.get_content_charset() or "utf-8"
        body = response.read()
    return body.decode(charset, errors="ignore")


def search_meta_tag(html_text: str, patterns: list[str]) -> str:
    for key in patterns:
        regex = rf'<meta[^>]+(?:property|name)=["\']{re.escape(key)}["\'][^>]+content=["\']([^"\']+)["\']'
        match = re.search(regex, html_text, flags=re.IGNORECASE)
        if match:
            return html.unescape(match.group(1).strip())
    return ""


def parse_embedded_json_ld(html_text: str) -> dict[str, Any]:
    for match in re.finditer(r'<script[^>]+type=["\']application/ld\+json["\'][^>]*>(.*?)</script>', html_text, flags=re.IGNORECASE | re.DOTALL):
        raw = html.unescape(match.group(1).strip())
        if not raw:
            continue
        try:
            data = json.loads(raw)
        except Exception:
            continue
        if isinstance(data, dict) and (data.get("@type") == "VideoObject" or data.get("contentUrl") or data.get("transcript")):
            return data
    return {}


def fetch_kwai_light_metadata(url: str) -> dict[str, Any]:
    html_text = fetch_remote_text(url)
    json_ld = parse_embedded_json_ld(html_text)
    creator = json_ld.get("creator") or {}
    creator_main = creator.get("mainEntity") if isinstance(creator, dict) else {}
    transcript = str(json_ld.get("transcript") or "").strip()
    title = str(json_ld.get("name") or search_meta_tag(html_text, ["og:title", "title"]) or "").strip()
    description = str(json_ld.get("description") or search_meta_tag(html_text, ["description", "og:description"]) or "").strip()
    thumbnail = ""
    thumb_list = json_ld.get("thumbnailUrl")
    if isinstance(thumb_list, list) and thumb_list:
        thumbnail = str(thumb_list[0] or "").strip()
    elif isinstance(thumb_list, str):
        thumbnail = thumb_list.strip()
    if not thumbnail:
        thumbnail = search_meta_tag(html_text, ["og:image", "twitter:image"])
    creator_description = ""
    if isinstance(creator_main, dict):
        creator_description = str(creator_main.get("description") or "").strip()
    return {
        "source_url": url,
        "title": title,
        "description": description,
        "transcript": transcript,
        "thumbnail_url": thumbnail,
        "content_url": str(json_ld.get("contentUrl") or "").strip(),
        "creator_name": str((creator_main or {}).get("name") or "").strip() if isinstance(creator_main, dict) else "",
        "creator_handle": str((creator_main or {}).get("alternateName") or "").strip() if isinstance(creator_main, dict) else "",
        "creator_description": creator_description,
        "duration": str(json_ld.get("duration") or "").strip(),
        "genre": json_ld.get("genre") if isinstance(json_ld.get("genre"), list) else [],
    }


GENDER_PROTO_URL = "https://raw.githubusercontent.com/spmallick/learnopencv/master/AgeGender/gender_deploy.prototxt"
GENDER_MODEL_URL = "https://raw.githubusercontent.com/GilLevi/AgeGenderDeepLearning/master/models/gender_net.caffemodel"
FACE_PROTO_URL = "https://raw.githubusercontent.com/opencv/opencv/master/samples/dnn/face_detector/deploy.prototxt"
FACE_MODEL_URL = "https://raw.githubusercontent.com/opencv/opencv_3rdparty/dnn_samples_face_detector_20180205_fp16/res10_300x300_ssd_iter_140000_fp16.caffemodel"
GENDER_MODEL_MEAN_VALUES = (78.4263377603, 87.7689143744, 114.895847746)
GENDER_LABELS = ["male", "female"]
_gender_net_lock = threading.Lock()
_gender_net_cache: Any = None
_face_net_lock = threading.Lock()
_face_net_cache: Any = None
_haar_face_lock = threading.Lock()
_haar_face_cache: Any = None
_haar_profile_lock = threading.Lock()
_haar_profile_cache: Any = None


def parse_duration_seconds(duration_text: object) -> float:
    text = str(duration_text or "").strip()
    if not text:
        return 0.0
    match = re.fullmatch(r"PT(?:(\d+)H)?(?:(\d+)M)?(?:(\d+(?:\.\d+)?)S)?", text)
    if not match:
        return 0.0
    hours = float(match.group(1) or 0)
    minutes = float(match.group(2) or 0)
    seconds = float(match.group(3) or 0)
    return hours * 3600 + minutes * 60 + seconds


CREATOR_DURATION_LABELS = {
    "dur_1_20": {"pt": "1-20 s", "zh": "1-20 秒"},
    "dur_20_60": {"pt": "20 s-1 min", "zh": "20 秒-1 分钟"},
    "dur_60_120": {"pt": "1-2 min", "zh": "1-2 分钟"},
    "dur_120_plus": {"pt": "Mais de 2 min", "zh": "2 分钟以上"},
}

LOCATION_TAG_OPTIONS = [
    {"zh": "室内房间", "pt": "Cômodo interno"},
    {"zh": "乡村院子", "pt": "Quintal / rural"},
    {"zh": "工地", "pt": "Obra / construção"},
    {"zh": "酒馆", "pt": "Bar / boteco"},
    {"zh": "超市", "pt": "Supermercado"},
    {"zh": "药店", "pt": "Farmácia"},
    {"zh": "房屋内外结合", "pt": "Casa: interno + externo"},
]
LOCATION_TAG_PT = {item["zh"]: item["pt"] for item in LOCATION_TAG_OPTIONS}
LOCATION_TAGS = {item["zh"] for item in LOCATION_TAG_OPTIONS}


def location_has_word(text: str, term: str) -> bool:
    if re.search(r"[\u4e00-\u9fff]", term):
        return term in text
    return bool(re.search(rf"(?<![\wÀ-ÿ]){re.escape(term)}(?![\wÀ-ÿ])", text, flags=re.I))


def location_score(text: str, terms: list[str]) -> int:
    return sum(1 for term in terms if location_has_word(text, term))


def flatten_script_text(value: Any, limit: int = 40000) -> str:
    chunks: list[str] = []

    def walk(item: Any) -> None:
        if len(" ".join(chunks)) > limit:
            return
        if isinstance(item, dict):
            for child in item.values():
                walk(child)
        elif isinstance(item, list):
            for child in item:
                walk(child)
        elif isinstance(item, str):
            chunks.append(item)

    walk(value)
    return " ".join(chunks)[:limit]


def library_script_json(entry_id: object) -> dict[str, Any]:
    output_dir = RESULTS_ROOT / str(entry_id or "")
    return (
        read_json(output_dir / "script_table_pt.json")
        or read_json(output_dir / "script_table.json")
        or read_json(output_dir / "analysis_result.json")
        or {}
    )


def infer_location_tag_fields(entry: dict[str, Any], script_json: dict[str, Any] | None = None) -> dict[str, str]:
    current = str(entry.get("location_tag") or "").strip()
    if current in LOCATION_TAGS:
        return {
            "location_tag": current,
            "location_tag_pt": str(entry.get("location_tag_pt") or LOCATION_TAG_PT.get(current) or ""),
            "location_tag_confidence": str(entry.get("location_tag_confidence") or "manual"),
            "location_tag_source": str(entry.get("location_tag_source") or "existing"),
            "location_tag_reasoning": str(entry.get("location_tag_reasoning") or "Existing location tag."),
        }
    text = " ".join(
        [
            str(entry.get("title") or ""),
            str(entry.get("whole_video_summary") or ""),
            str(entry.get("summary") or ""),
            str(entry.get("content_type_reasoning") or ""),
            flatten_script_text(script_json or {}),
        ]
    ).lower()
    terms = {
        "室内房间": ["quarto", "sala", "cozinha", "banheiro", "cama", "sofá", "sofa", "mesa", "tapete", "lavanderia", "casa", "apartamento", "hospital", "consultório", "consultorio", "室内", "房间", "卧室", "客厅", "厨房", "浴室", "卫生间", "床", "沙发", "洗衣房", "医院", "诊所"],
        "乡村院子": ["quintal", "roça", "roca", "rural", "fazenda", "sítio", "sitio", "terreiro", "milho", "plantação", "plantacao", "horta", "campo", "curral", "galinha", "院子", "乡村", "农村", "农田", "玉米地", "田地", "菜地", "后院", "农家"],
        "工地": ["obra", "construção", "construcao", "pedreiro", "cimento", "tijolo", "andaime", "reforma", "capacete", "martelo", "工地", "施工", "建筑", "装修", "水泥", "砖", "脚手架", "安全帽"],
        "酒馆": ["bar", "boteco", "pub", "balcão do bar", "balcao do bar", "mesa de bar", "bebida no bar", "cerveja no bar", "酒馆", "酒吧", "吧台", "啤酒馆"],
        "超市": ["supermercado", "mercado", "mercearia", "caixa do mercado", "carrinho de compras", "prateleira", "corredor do mercado", "compras no mercado", "超市", "便利店", "商店", "货架", "收银台", "购物车"],
        "药店": ["farmácia", "farmacia", "balcão da farmácia", "balcao da farmacia", "atendente da farmácia", "atendente da farmacia", "药店", "药房"],
    }
    scores = {label: location_score(text, values) for label, values in terms.items()}
    indoor = scores.get("室内房间", 0)
    outdoor = (
        scores.get("乡村院子", 0)
        + location_score(text, ["rua", "portão", "portao", "porta de casa", "fora de casa", "sai de casa", "calçada", "calcada", "estrada", "门口", "大门", "户外", "屋外", "出门", "街道"])
    )
    if indoor >= 2 and outdoor >= 2:
        tag = "房屋内外结合"
        confidence = "high" if indoor + outdoor >= 5 else "medium"
        reason = f"室内信号 {indoor} 个，室外/门口信号 {outdoor} 个。"
    else:
        priority = ["药店", "工地", "酒馆", "超市", "乡村院子", "室内房间"]
        tag = max(priority, key=lambda label: (scores.get(label, 0), -priority.index(label)))
        best = scores.get(tag, 0)
        if best <= 0:
            tag = "室内房间"
            confidence = "low"
            reason = "未识别到明确地点词，按常见室内短剧兜底。"
        else:
            confidence = "high" if best >= 3 else "medium" if best >= 2 else "low"
            reason = f"命中 {tag} 地点词 {best} 个。"
    return {
        "location_tag": tag,
        "location_tag_pt": LOCATION_TAG_PT.get(tag, ""),
        "location_tag_confidence": confidence,
        "location_tag_source": "summary+storyboard",
        "location_tag_reasoning": reason,
    }


def parse_script_timecode_seconds(value: object) -> float:
    text = str(value or "").strip()
    if not text:
        return 0.0
    nums = [int(part) for part in re.findall(r"\d+", text)]
    if len(nums) >= 3:
        return float(nums[-3] * 3600 + nums[-2] * 60 + nums[-1])
    if len(nums) >= 2:
        return float(nums[-2] * 60 + nums[-1])
    if nums:
        return float(nums[-1])
    return 0.0


def creator_duration_bucket(seconds: float) -> str:
    if seconds <= 0:
        return ""
    if seconds <= 20:
        return "dur_1_20"
    if seconds <= 60:
        return "dur_20_60"
    if seconds <= 120:
        return "dur_60_120"
    return "dur_120_plus"


def script_duration_seconds(script_json: dict[str, Any]) -> float:
    rows = script_json.get("segments") or script_json.get("rows") or script_json.get("script_table") or []
    if not isinstance(rows, list):
        return 0.0
    values = [
        parse_script_timecode_seconds(row.get("time") or row.get("tempo") or row.get("Tempo") or "")
        for row in rows
        if isinstance(row, dict)
    ]
    return max(values) if values else 0.0


def creator_duration_fields(entry_id: str, script_json: dict[str, Any] | None = None) -> dict[str, Any]:
    if script_json is None:
        output_dir = RESULTS_ROOT / str(entry_id or "")
        script_json = read_json(output_dir / "script_table_pt.json") or read_json(output_dir / "script_table.json")
    duration_seconds = script_duration_seconds(script_json or {})
    duration_bucket = creator_duration_bucket(duration_seconds)
    return {
        "duration_seconds": round(duration_seconds, 2) if duration_seconds > 0 else 0,
        "duration_bucket": duration_bucket,
        "duration_label_pt": CREATOR_DURATION_LABELS.get(duration_bucket, {}).get("pt", ""),
        "duration_label_zh": CREATOR_DURATION_LABELS.get(duration_bucket, {}).get("zh", ""),
    }


def download_url_to_file(url: str, target: Path) -> None:
    target.parent.mkdir(parents=True, exist_ok=True)
    request = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
    with urllib.request.urlopen(request, timeout=60) as response, target.open("wb") as handle:
        shutil.copyfileobj(response, handle)


def ensure_gender_net() -> Any:
    if cv2 is None:
        raise RuntimeError("cv2 not available")
    global _gender_net_cache
    with _gender_net_lock:
        if _gender_net_cache is not None:
            return _gender_net_cache
        proto_path = VISION_MODELS_DIR / "gender_deploy.prototxt"
        model_path = VISION_MODELS_DIR / "gender_net.caffemodel"
        if not proto_path.exists():
            download_url_to_file(GENDER_PROTO_URL, proto_path)
        if not model_path.exists():
            download_url_to_file(GENDER_MODEL_URL, model_path)
        _gender_net_cache = cv2.dnn.readNet(str(model_path), str(proto_path))
        return _gender_net_cache


def ensure_face_net() -> Any:
    if cv2 is None:
        raise RuntimeError("cv2 not available")
    global _face_net_cache
    with _face_net_lock:
        if _face_net_cache is not None:
            return _face_net_cache
        proto_path = VISION_MODELS_DIR / "face_deploy.prototxt"
        model_path = VISION_MODELS_DIR / "face_res10_fp16.caffemodel"
        if not proto_path.exists():
            download_url_to_file(FACE_PROTO_URL, proto_path)
        if not model_path.exists():
            download_url_to_file(FACE_MODEL_URL, model_path)
        _face_net_cache = cv2.dnn.readNet(str(model_path), str(proto_path))
        return _face_net_cache


def ensure_haar_face_cascade() -> Any:
    if cv2 is None:
        raise RuntimeError("cv2 not available")
    global _haar_face_cache
    with _haar_face_lock:
        if _haar_face_cache is not None:
            return _haar_face_cache
        cascade = cv2.CascadeClassifier(cv2.data.haarcascades + "haarcascade_frontalface_default.xml")
        if cascade.empty():
            raise RuntimeError("failed to load frontal face cascade")
        _haar_face_cache = cascade
        return _haar_face_cache


def ensure_haar_profile_cascade() -> Any:
    if cv2 is None:
        raise RuntimeError("cv2 not available")
    global _haar_profile_cache
    with _haar_profile_lock:
        if _haar_profile_cache is not None:
            return _haar_profile_cache
        cascade = cv2.CascadeClassifier(cv2.data.haarcascades + "haarcascade_profileface.xml")
        if cascade.empty():
            raise RuntimeError("failed to load profile face cascade")
        _haar_profile_cache = cascade
        return _haar_profile_cache


def detect_faces_dnn(image: Any, *, min_confidence: float = 0.45, min_face_size: int = 24) -> list[tuple[int, int, int, int]]:
    if cv2 is None or image is None:
        return []
    net = ensure_face_net()
    (height, width) = image.shape[:2]
    blob = cv2.dnn.blobFromImage(cv2.resize(image, (300, 300)), 1.0, (300, 300), (104.0, 177.0, 123.0))
    net.setInput(blob)
    detections = net.forward()
    faces: list[tuple[int, int, int, int]] = []
    for i in range(detections.shape[2]):
        confidence = float(detections[0, 0, i, 2])
        if confidence < min_confidence:
            continue
        box = detections[0, 0, i, 3:7] * [width, height, width, height]
        x1, y1, x2, y2 = box.astype(int)
        x1 = max(0, x1)
        y1 = max(0, y1)
        x2 = min(width, x2)
        y2 = min(height, y2)
        if (x2 - x1) < min_face_size or (y2 - y1) < min_face_size:
            continue
        faces.append((x1, y1, x2, y2))
    return faces


def dedupe_face_boxes(boxes: list[tuple[int, int, int, int]], iou_threshold: float = 0.35) -> list[tuple[int, int, int, int]]:
    if not boxes:
        return []
    kept: list[tuple[int, int, int, int]] = []
    for box in sorted(boxes, key=lambda item: (item[2] - item[0]) * (item[3] - item[1]), reverse=True):
        x1, y1, x2, y2 = box
        keep = True
        for kx1, ky1, kx2, ky2 in kept:
            inter_x1 = max(x1, kx1)
            inter_y1 = max(y1, ky1)
            inter_x2 = min(x2, kx2)
            inter_y2 = min(y2, ky2)
            inter_w = max(0, inter_x2 - inter_x1)
            inter_h = max(0, inter_y2 - inter_y1)
            inter = inter_w * inter_h
            area_a = max(1, (x2 - x1) * (y2 - y1))
            area_b = max(1, (kx2 - kx1) * (ky2 - ky1))
            union = area_a + area_b - inter
            iou = inter / union if union > 0 else 0.0
            if iou >= iou_threshold:
                keep = False
                break
        if keep:
            kept.append(box)
    return kept


def detect_faces_haar(image: Any, *, min_face_size: int = 20) -> list[tuple[int, int, int, int]]:
    if cv2 is None or image is None:
        return []
    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    frontal = ensure_haar_face_cascade()
    profile = ensure_haar_profile_cascade()
    boxes: list[tuple[int, int, int, int]] = []
    for raw_boxes in (
        frontal.detectMultiScale(gray, scaleFactor=1.1, minNeighbors=3, minSize=(min_face_size, min_face_size)),
        profile.detectMultiScale(gray, scaleFactor=1.1, minNeighbors=3, minSize=(min_face_size, min_face_size)),
    ):
        for x, y, w, h in raw_boxes:
            boxes.append((int(x), int(y), int(x + w), int(y + h)))
    return dedupe_face_boxes(boxes)


def detect_faces_combined(image: Any, *, dnn_confidence: float = 0.3, min_face_size: int = 18) -> list[tuple[int, int, int, int]]:
    boxes = detect_faces_dnn(image, min_confidence=dnn_confidence, min_face_size=min_face_size)
    boxes.extend(detect_faces_haar(image, min_face_size=min_face_size))
    return dedupe_face_boxes(boxes)


def ffmpeg_input_options(referer: str = "https://www.kwai.com/") -> list[str]:
    headers = (
        "User-Agent: Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
        "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0 Safari/537.36\r\n"
        f"Referer: {referer}\r\n"
    )
    return [
        "-hide_banner",
        "-loglevel",
        "error",
        "-reconnect",
        "1",
        "-reconnect_streamed",
        "1",
        "-reconnect_delay_max",
        "5",
        "-headers",
        headers,
    ]


def clean_ffmpeg_error(stderr: str, stdout: str = "") -> str:
    text = (stderr or stdout or "").strip()
    if not text:
        return "ffmpeg exited without error output"
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    ignored_prefixes = (
        "ffmpeg version ",
        "built with ",
        "configuration:",
        "libav",
        "libsw",
        "libpostproc",
    )
    useful = [
        line
        for line in lines
        if not line.startswith(ignored_prefixes)
        and not re.fullmatch(r"[A-Za-z0-9_]+ +\d+\.\s*\d+\.\d+.*", line)
    ]
    cleaned = "\n".join(useful or lines[-8:])
    cleaned = re.sub(r"\s+", " ", cleaned).strip()
    if len(cleaned) > 900:
        cleaned = cleaned[:900].rstrip() + "..."
    return cleaned or "ffmpeg failed"


def ensure_filter_source_video(content_url: str, cache_dir: Path) -> Path:
    cache_dir.mkdir(parents=True, exist_ok=True)
    source_path = cache_dir / "source.mp4"
    if source_path.exists() and source_path.stat().st_size > 0:
        return source_path
    download_url_to_file(content_url, source_path)
    if not source_path.exists() or source_path.stat().st_size <= 0:
        raise RuntimeError("临时源视频下载失败。")
    return source_path


def ffmpeg_input_args(source: str, *, remote: bool) -> list[str]:
    if remote:
        return [*ffmpeg_input_options(), "-i", source]
    return ["-i", source]


def extract_keyframes_from_source(ffmpeg_bin: str, source: str, duration_seconds: float, out_dir: Path, *, remote: bool) -> list[Path]:
    duration = duration_seconds if duration_seconds > 1 else 8.0
    timestamps = [0.2, max(duration * 0.5, 0.4), max(duration - 0.4, 0.6)]
    names = ["start.jpg", "middle.jpg", "end.jpg"]
    frames: list[Path] = []
    last_error = ""
    for ts, name in zip(timestamps, names):
        out_path = out_dir / name
        cmd = [
            ffmpeg_bin,
            "-y",
            "-ss",
            f"{max(ts, 0):.2f}",
            *ffmpeg_input_args(source, remote=remote),
            "-frames:v",
            "1",
            "-vf",
            "scale=360:-1",
            str(out_path),
        ]
        result = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
        if result.returncode == 0 and out_path.exists() and out_path.stat().st_size > 0:
            frames.append(out_path)
        elif result.returncode != 0:
            last_error = clean_ffmpeg_error(result.stderr, result.stdout)
    if not frames and last_error:
        raise RuntimeError(f"关键帧抽取失败：{last_error}")
    return frames


def extract_remote_keyframes(content_url: str, duration_seconds: float, out_dir: Path) -> list[Path]:
    if not content_url:
        return []
    try:
        import imageio_ffmpeg  # type: ignore
    except Exception as exc:
        raise RuntimeError(f"imageio_ffmpeg unavailable: {exc}") from exc
    ffmpeg_bin = imageio_ffmpeg.get_ffmpeg_exe()
    out_dir.mkdir(parents=True, exist_ok=True)
    try:
        return extract_keyframes_from_source(ffmpeg_bin, content_url, duration_seconds, out_dir, remote=True)
    except Exception as remote_exc:
        source_path = ensure_filter_source_video(content_url, out_dir)
        try:
            return extract_keyframes_from_source(ffmpeg_bin, str(source_path), duration_seconds, out_dir, remote=False)
        except Exception as local_exc:
            raise RuntimeError(f"{remote_exc}；临时下载后仍失败：{local_exc}") from local_exc


def mime_type_for_path(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in {".jpg", ".jpeg"}:
        return "image/jpeg"
    if suffix == ".png":
        return "image/png"
    if suffix == ".webp":
        return "image/webp"
    if suffix == ".mp3":
        return "audio/mpeg"
    if suffix == ".wav":
        return "audio/wav"
    if suffix in {".m4a", ".mp4"}:
        return "audio/mp4"
    return "application/octet-stream"


def extract_text_from_gemini_response(raw: dict[str, Any]) -> str:
    parts: list[str] = []
    for candidate in raw.get("candidates") or []:
        content = candidate.get("content") or {}
        for part in content.get("parts") or []:
            text = str(part.get("text") or "").strip()
            if text:
                parts.append(text)
    return "\n".join(parts).strip()


def parse_json_object_from_text(text: str) -> dict[str, Any]:
    cleaned = str(text or "").strip()
    cleaned = re.sub(r"^```(?:json)?", "", cleaned, flags=re.IGNORECASE).strip()
    cleaned = re.sub(r"```$", "", cleaned).strip()
    try:
        value = json.loads(cleaned)
        return value if isinstance(value, dict) else {}
    except Exception:
        pass
    match = re.search(r"\{.*\}", cleaned, flags=re.DOTALL)
    if not match:
        return {}
    value = json.loads(match.group(0))
    return value if isinstance(value, dict) else {}


def run_gemini_file_json_prompt(files: list[Path], payload: dict[str, Any], prompt: str, label: str) -> tuple[dict[str, Any], dict[str, Any], str]:
    if not GOOGLE_API_KEY:
        raise RuntimeError(f"{label} requires GOOGLE_API_KEY")
    models = unique_models(os.environ.get("VIDEO_FILTER_MODEL", "gemini-2.5-flash-lite"), *PRIMARY_FALLBACK_MODELS)
    last_error: Exception | None = None
    for model in models:
        try:
            parts: list[dict[str, Any]] = []
            for path in files:
                if not path.exists() or path.stat().st_size <= 0:
                    continue
                parts.append(
                    {
                        "inline_data": {
                            "mime_type": mime_type_for_path(path),
                            "data": base64.b64encode(path.read_bytes()).decode("ascii"),
                        }
                    }
                )
            parts.extend(
                [
                    {"text": prompt},
                    {"text": json.dumps(payload, ensure_ascii=False)},
                ]
            )
            body = {
                "contents": [{"parts": parts}],
            }
            url = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent"
            req = urllib.request.Request(
                url,
                data=json.dumps(body).encode("utf-8"),
                headers={"Content-Type": "application/json", "x-goog-api-key": GOOGLE_API_KEY},
            )
            try:
                with urllib.request.urlopen(req, timeout=240) as response:
                    raw = json.loads(response.read().decode("utf-8"))
            except urllib.error.HTTPError as exc:
                detail = exc.read().decode("utf-8", "replace")
                raise RuntimeError(f"Gemini filter HTTP {exc.code}: {detail}") from exc
            text = extract_text_from_gemini_response(raw)
            return parse_json_object_from_text(text), raw, model
        except Exception as exc:
            last_error = exc
            if not is_retryable_filter_error(str(exc)):
                break
    raise RuntimeError(f"{label} failed across models {models}: {last_error}") from last_error


def is_retryable_filter_error(text: str) -> bool:
    hay = str(text or "").upper()
    return any(token in hay for token in ["HTTP 503", "UNAVAILABLE", "RESOURCE_EXHAUSTED", "TIMED OUT", "JSON", "EOF", "CONNECTION RESET"])


def extract_audio_from_source(ffmpeg_bin: str, source: str, out_path: Path, *, remote: bool) -> Path:
    cmd = [
        ffmpeg_bin,
        "-y",
        *ffmpeg_input_args(source, remote=remote),
        "-vn",
        "-ac",
        "1",
        "-ar",
        "16000",
        "-b:a",
        "48k",
        str(out_path),
    ]
    result = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True, timeout=240)
    if result.returncode != 0 or not out_path.exists() or out_path.stat().st_size <= 0:
        raise RuntimeError(f"完整音频提取失败：{clean_ffmpeg_error(result.stderr, result.stdout)}")
    return out_path


def extract_remote_audio(content_url: str, cache_dir: Path) -> Path:
    if not content_url:
        raise RuntimeError("content url missing")
    try:
        import imageio_ffmpeg  # type: ignore
    except Exception as exc:
        raise RuntimeError(f"imageio_ffmpeg unavailable: {exc}") from exc
    ffmpeg_bin = imageio_ffmpeg.get_ffmpeg_exe()
    cache_dir.mkdir(parents=True, exist_ok=True)
    out_path = cache_dir / "full_audio.mp3"
    try:
        extract_audio_from_source(ffmpeg_bin, content_url, out_path, remote=True)
    except Exception as remote_exc:
        source_path = ensure_filter_source_video(content_url, cache_dir)
        try:
            extract_audio_from_source(ffmpeg_bin, str(source_path), out_path, remote=False)
        except Exception as local_exc:
            raise RuntimeError(f"{remote_exc}；临时下载后仍失败：{local_exc}") from local_exc
    max_bytes = max(1, FILTER_AUDIO_MAX_MB) * 1024 * 1024
    if out_path.stat().st_size > max_bytes:
        raise RuntimeError(f"audio file too large for filter transcription: {out_path.stat().st_size} bytes")
    try:
        source_path = cache_dir / "source.mp4"
        if source_path.exists():
            source_path.unlink()
    except Exception:
        pass
    return out_path


def transcribe_filter_audio(metadata: dict[str, Any], cache_dir: Path) -> dict[str, Any]:
    page_transcript = str(metadata.get("transcript") or "").strip()
    if page_transcript:
        return {
            "available": True,
            "source": "page_transcript",
            "full_transcript": page_transcript,
            "language": "",
            "dialogue_summary": page_transcript[:500],
            "speaker_hints": [],
            "audio_form": "page transcript",
            "confidence": "medium",
            "audio_path": "",
            "error": "",
        }
    if not GOOGLE_API_KEY:
        raise RuntimeError("GOOGLE_API_KEY missing; cannot transcribe remote audio.")
    audio_path = extract_remote_audio(str(metadata.get("content_url") or "").strip(), cache_dir)
    prompt = """
你是 Koko 的筛选前置音频转写器。请完整理解这段短视频音频，优先输出人物对白、旁白和能帮助判断剧情的信息。

严格返回 JSON：
{
  "full_transcript": "完整音频转写；听不清处标注[听不清]；不同说话人可用人物A/人物B/旁白",
  "language": "主要语言",
  "dialogue_summary": "一句话总结音频发生了什么",
  "speaker_hints": ["最多6条：说话人数量、关系、角色线索"],
  "audio_form": "dialogue/monologue/voiceover/BGM_only/mixed/unknown",
  "confidence": "high/medium/low"
}
不要编造听不到的台词；如果只有音乐或环境音，也要明确说明。
""".strip()
    payload = {
        "source_url": metadata.get("source_url") or "",
        "title": metadata.get("title") or "",
        "duration": metadata.get("duration") or "",
    }
    result, _, model = run_gemini_file_json_prompt([audio_path], payload, prompt, "filter audio transcription")
    transcript = str(result.get("full_transcript") or "").strip()
    if not transcript:
        raise RuntimeError("filter audio transcription returned empty transcript")
    return {
        "available": True,
        "source": "gemini_audio",
        "full_transcript": transcript,
        "language": str(result.get("language") or "").strip(),
        "dialogue_summary": str(result.get("dialogue_summary") or "").strip(),
        "speaker_hints": [str(item or "").strip() for item in (result.get("speaker_hints") or []) if str(item or "").strip()][:6],
        "audio_form": str(result.get("audio_form") or "unknown").strip(),
        "confidence": str(result.get("confidence") or "medium").strip().lower(),
        "model_used": model,
        "audio_path": str(audio_path),
        "error": "",
    }


def count_thumbnail_faces(thumbnail_url: str, cache_dir: Path) -> dict[str, Any]:
    if cv2 is None or not thumbnail_url:
        return {"available": False, "face_count": 0}
    cache_dir.mkdir(parents=True, exist_ok=True)
    thumb_path = cache_dir / "thumbnail.webp"
    try:
        if not thumb_path.exists():
            download_url_to_file(thumbnail_url, thumb_path)
        image = cv2.imread(str(thumb_path))
        if image is None:
            return {"available": False, "face_count": 0}
        faces = detect_faces_combined(image, dnn_confidence=0.3, min_face_size=18)
        return {"available": True, "face_count": len(faces), "path": str(thumb_path)}
    except Exception:
        return {"available": False, "face_count": 0}


def detect_gender_presence_from_frames(content_url: str, duration_text: object, cache_dir: Path) -> dict[str, Any]:
    if not content_url:
        return {"available": False, "reason": "content url missing", "bucket": "low", "signals": [], "score_boost": 0}
    duration_seconds = parse_duration_seconds(duration_text)
    frames = extract_remote_keyframes(content_url, duration_seconds, cache_dir)
    if not frames:
        return {"available": False, "reason": "frames unavailable", "bucket": "low", "signals": [], "score_boost": 0}
    if cv2 is None:
        return {
            "available": False,
            "reason": "cv2 unavailable; keyframes extracted without face detection",
            "bucket": "low",
            "signals": ["已抽取开头、中间、结尾三张关键帧，但当前环境无法做人脸/性别检测。"],
            "score_boost": 0,
            "frame_count": len(frames),
            "inspected_frames": 0,
            "face_total": 0,
            "male_count": 0,
            "female_count": 0,
            "pair_frames": 0,
            "has_both": False,
            "max_faces_single_frame": 0,
            "frame_summaries": [{"name": path.name, "face_count": 0, "male_count": 0, "female_count": 0, "is_pair_frame": False} for path in frames],
            "frame_paths": [str(path) for path in frames],
        }
    gender_net = ensure_gender_net()
    male_count = 0
    female_count = 0
    pair_frames = 0
    inspected_frames = 0
    face_total = 0
    max_faces_single_frame = 0
    frame_summaries: list[dict[str, Any]] = []
    for frame_path in frames:
        image = cv2.imread(str(frame_path))
        if image is None:
            continue
        inspected_frames += 1
        faces = detect_faces_combined(image, dnn_confidence=0.3, min_face_size=18)
        frame_male = 0
        frame_female = 0
        face_total += len(faces)
        max_faces_single_frame = max(max_faces_single_frame, len(faces))
        for (x1, y1, x2, y2) in faces:
            face = image[y1:y2, x1:x2]
            if face.size == 0:
                continue
            blob = cv2.dnn.blobFromImage(
                cv2.resize(face, (227, 227)),
                1.0,
                (227, 227),
                GENDER_MODEL_MEAN_VALUES,
                swapRB=False,
            )
            gender_net.setInput(blob)
            preds = gender_net.forward()[0]
            label = GENDER_LABELS[int(preds.argmax())]
            if label == "male":
                male_count += 1
                frame_male += 1
            elif label == "female":
                female_count += 1
                frame_female += 1
        if len(faces) == 2 and frame_male == 1 and frame_female == 1:
            pair_frames += 1
        frame_summaries.append(
            {
                "name": frame_path.name,
                "face_count": len(faces),
                "male_count": frame_male,
                "female_count": frame_female,
                "is_pair_frame": len(faces) == 2 and frame_male == 1 and frame_female == 1,
            }
        )
    has_both = male_count > 0 and female_count > 0
    signals: list[str] = []
    score_boost = 0
    bucket = "low"
    if pair_frames >= 1:
        signals.append("至少一帧满足双人一男一女主场景")
        score_boost += 8
        bucket = "high"
    elif max_faces_single_frame >= 3:
        signals.append(f"关键帧出现三人及以上主场景（单帧最多 {max_faces_single_frame} 张脸）")
        bucket = "low"
    elif has_both:
        signals.append(f"跨帧检测到男女都出现过（男 {male_count} / 女 {female_count}），但没有稳定双人主场景")
        bucket = "low"
    elif male_count > 0 or female_count > 0:
        signals.append(f"关键帧只稳定检测到单一性别（男 {male_count} / 女 {female_count}）")
        bucket = "low"
    else:
        signals.append("关键帧里未稳定识别出可用人脸")
    return {
        "available": True,
        "frame_count": len(frames),
        "inspected_frames": inspected_frames,
        "face_total": face_total,
        "male_count": male_count,
        "female_count": female_count,
        "pair_frames": pair_frames,
        "has_both": has_both,
        "max_faces_single_frame": max_faces_single_frame,
        "frame_summaries": frame_summaries,
        "bucket": bucket,
        "signals": signals,
        "score_boost": score_boost,
        "frame_paths": [str(path) for path in frames],
    }


COUPLE_KEYWORDS = {
    "high": [
        "casal", "marido", "esposa", "husband", "wife", "老公", "老婆", "丈夫", "妻子", "夫妻",
        "namorado", "namorada", "boyfriend", "girlfriend", "伴侣", "情侣", "婚姻", "casamento",
    ],
    "context": [
        "dia das mães", "mãe", "ciúmes", "ciume", "família", "family", "humor", "relationship",
        "house", "casa", "cozinha", "kitchen", "present", "gift", "carro", "car", "争吵", "吵架",
        "礼物", "家务", "卧室", "沙发", "妈妈", "爸爸", "母亲节",
    ],
}

EXPLICIT_SPOUSE_TERMS = {
    "casal", "marido", "esposa", "husband", "wife", "老公", "老婆", "丈夫", "妻子", "夫妻"
}


def has_duo_creator_signal(metadata: dict[str, Any]) -> bool:
    creator_name = str(metadata.get("creator_name") or "").strip().lower()
    creator_handle = str(metadata.get("creator_handle") or "").strip().lower()
    if not creator_name:
        return False
    duo_markers = [" e ", " & ", " and ", "+", "/"]
    if any(marker in creator_name for marker in duo_markers):
        return True
    if any(marker in creator_handle for marker in duo_markers):
        return True
    return False


def has_explicit_spouse_signal(metadata: dict[str, Any]) -> bool:
    combined = "\n".join(
        str(metadata.get(key) or "")
        for key in ("title", "description", "transcript", "creator_description", "creator_name")
    ).lower()
    return any(term in combined for term in EXPLICIT_SPOUSE_TERMS)


def score_couple_candidate(metadata: dict[str, Any], visual: dict[str, Any] | None = None) -> dict[str, Any]:
    combined = "\n".join(
        str(metadata.get(key) or "")
        for key in ("title", "description", "transcript", "creator_description", "creator_name", "creator_handle")
    ).lower()
    score = 0
    reasons: list[str] = []
    for keyword in COUPLE_KEYWORDS["high"]:
        if keyword.lower() in combined:
            score += 3
            reasons.append(f"命中强关系词：{keyword}")
    for keyword in COUPLE_KEYWORDS["context"]:
        if keyword.lower() in combined:
            score += 1
            reasons.append(f"命中场景词：{keyword}")
    title = str(metadata.get("title") or "")
    transcript = str(metadata.get("transcript") or "")
    if title and transcript and title.strip() != transcript.strip():
        score += 1
        reasons.append("页面标题与转写都可用")
    if metadata.get("thumbnail_url"):
        score += 1
        reasons.append("存在封面图，可供后续视觉扩展")
    if "casal" in str(metadata.get("creator_description") or "").lower():
        score += 2
        reasons.append("账号简介包含 casal")
    visual_bucket = "low"
    if visual:
        visual_bucket = str(visual.get("bucket") or "low").strip().lower()
        score += int(visual.get("score_boost") or 0)
        reasons.extend(
            str(signal or "").strip()
            for signal in (visual.get("signals") or [])
            if str(signal or "").strip()
        )
    if score >= 8:
        bucket = "high"
    elif score >= 4:
        bucket = "medium"
    else:
        bucket = "low"
    return {"score": score, "bucket": bucket, "reasons": reasons, "visual_bucket": visual_bucket}


def classify_couple_candidate(metadata: dict[str, Any], heuristic: dict[str, Any], visual: dict[str, Any] | None = None) -> dict[str, Any]:
    visual_data = visual or {}
    max_faces_single_frame = int(visual_data.get("max_faces_single_frame") or 0)
    pair_frames = int(visual_data.get("pair_frames") or 0)
    has_both = bool(visual_data.get("has_both"))
    thumbnail_face_count = int((visual_data.get("thumbnail_faces") or {}).get("face_count") or 0)
    explicit_spouse = has_explicit_spouse_signal(metadata)
    duo_creator = has_duo_creator_signal(metadata)
    cover_not_crowded = thumbnail_face_count == 0 or thumbnail_face_count <= 2
    stable_duo_scene = cover_not_crowded and has_both and max_faces_single_frame <= 2 and (
        pair_frames >= 2 or (pair_frames >= 1 and 1 <= thumbnail_face_count <= 2)
    )
    if stable_duo_scene:
        return {
            "bucket": "high",
            "confidence": "high",
            "reason": "关键帧里出现了稳定的双人一男一女主场景。",
            "signals": [str(signal or "").strip() for signal in (visual_data.get("signals") or []) if str(signal or "").strip()][:4]
            or ["关键帧满足稳定双人一男一女主场景"],
            "used_llm": False,
        }
    if explicit_spouse:
        return {
            "bucket": "high",
            "confidence": "medium",
            "reason": "标题或转写里出现了明确的配偶关系词，判为夫妻候选。",
            "signals": ["命中明确配偶关系词"] + (heuristic.get("reasons") or [])[:3],
            "used_llm": False,
        }
    if duo_creator and 1 <= thumbnail_face_count <= 2:
        return {
            "bucket": "high",
            "confidence": "medium",
            "reason": "账号明显是双人组合，且封面未出现三人以上，判为夫妻候选。",
            "signals": [f"账号名呈现双人组合", f"封面检测到 {thumbnail_face_count} 张脸"],
            "used_llm": False,
        }
    if not FILTER_USE_LLM or not GOOGLE_API_KEY or run_text_json_prompt_with_fallback is None:
        return {
            "bucket": "low",
            "confidence": "high" if max_faces_single_frame >= 1 or thumbnail_face_count >= 1 else "medium",
            "reason": "未满足双人一男一女主场景，也没有命中明确的夫妻候选信号。",
            "signals": (
                [str(signal or "").strip() for signal in (visual_data.get("signals") or []) if str(signal or "").strip()][:2]
                + (heuristic.get("reasons") or [])[:2]
            )[:4],
            "used_llm": False,
        }
    prompt = f"""
你是 Koko 的视频筛选器。目标是判断一个 Kwai 视频是否“可能属于夫妻类型”，这里的夫妻类型包括：
- 明确是夫妻/老公老婆/丈夫妻子
- 或非常像固定亲密伴侣关系的家庭喜剧视频

请严格返回 JSON，字段：
{{
  "bucket": "high" | "medium" | "low",
  "confidence": "high" | "medium" | "low",
  "reason": "一句中文解释",
  "signals": ["最多4条中文信号"]
}}

判断标准：
- high：高度像夫妻/伴侣关系主导的视频
- medium：疑似夫妻/伴侣，但证据不够稳
- low：不太像夫妻类型，或无法判断

不要编造不存在的画面，只能基于页面公开信息判断。

视频链接：{metadata.get("source_url") or ""}
标题：{metadata.get("title") or ""}
简介：{metadata.get("description") or ""}
转写：{metadata.get("transcript") or ""}
账号名：{metadata.get("creator_name") or ""}
账号简介：{metadata.get("creator_description") or ""}
启发式结果：{json.dumps(heuristic, ensure_ascii=False)}
关键帧视觉结果：{json.dumps(visual_data, ensure_ascii=False)}
""".strip()
    models = unique_models(os.environ.get("VIDEO_FILTER_MODEL", "gemini-2.5-flash-lite"), *PRIMARY_FALLBACK_MODELS)
    payload = {
        "source_url": metadata.get("source_url") or "",
        "title": metadata.get("title") or "",
        "description": metadata.get("description") or "",
        "transcript": metadata.get("transcript") or "",
        "creator_name": metadata.get("creator_name") or "",
        "creator_description": metadata.get("creator_description") or "",
        "heuristic": heuristic,
        "visual": visual_data,
        "schema": {
            "type": "object",
            "properties": {
                "bucket": {"type": "string"},
                "confidence": {"type": "string"},
                "reason": {"type": "string"},
                "signals": {"type": "array", "items": {"type": "string"}},
            },
            "required": ["bucket", "confidence", "reason", "signals"],
        },
    }
    result, _, _ = run_text_json_prompt_with_fallback(
        payload,
        GOOGLE_API_KEY,
        models,
        prompt,
        "couple candidate classification",
    )
    bucket = str(result.get("bucket") or heuristic.get("bucket") or "low").strip().lower()
    if bucket not in {"high", "medium", "low"}:
        bucket = heuristic.get("bucket") or "low"
    confidence = str(result.get("confidence") or "medium").strip().lower()
    if confidence not in {"high", "medium", "low"}:
        confidence = "medium"
    reason = str(result.get("reason") or "").strip() or "已基于页面信息完成轻量判断。"
    signals = [str(signal or "").strip() for signal in (result.get("signals") or []) if str(signal or "").strip()][:4]
    return {
        "bucket": bucket,
        "confidence": confidence,
        "reason": reason,
        "signals": signals or (heuristic.get("reasons") or [])[:4],
        "used_llm": True,
    }


def filter_check(passed: bool, reason: str, *, confidence: str = "high", evidence: list[str] | None = None, **extra: Any) -> dict[str, Any]:
    payload = {
        "passed": bool(passed),
        "confidence": confidence if confidence in {"high", "medium", "low"} else "medium",
        "reason": str(reason or "").strip(),
        "evidence": [str(item or "").strip() for item in (evidence or []) if str(item or "").strip()][:6],
    }
    payload.update(extra)
    return payload


def build_filter_evidence_bundle(metadata: dict[str, Any], audio: dict[str, Any], visual: dict[str, Any]) -> dict[str, Any]:
    duration_seconds = parse_duration_seconds(metadata.get("duration"))
    frame_paths = [str(path or "").strip() for path in (visual.get("frame_paths") or []) if str(path or "").strip()]
    bundle = {
        "metadata": {
            "source_url": metadata.get("source_url") or "",
            "title": metadata.get("title") or "",
            "description": metadata.get("description") or "",
            "creator_name": metadata.get("creator_name") or "",
            "creator_handle": metadata.get("creator_handle") or "",
            "creator_description": metadata.get("creator_description") or "",
            "duration": metadata.get("duration") or "",
            "duration_seconds": duration_seconds,
            "thumbnail_url": metadata.get("thumbnail_url") or "",
            "content_url_available": bool(metadata.get("content_url")),
            "genre": metadata.get("genre") or [],
        },
        "audio": {
            "available": bool(audio.get("available")),
            "source": audio.get("source") or "",
            "full_transcript": audio.get("full_transcript") or "",
            "language": audio.get("language") or "",
            "dialogue_summary": audio.get("dialogue_summary") or "",
            "speaker_hints": audio.get("speaker_hints") or [],
            "audio_form": audio.get("audio_form") or "unknown",
            "confidence": audio.get("confidence") or "low",
            "error": audio.get("error") or "",
        },
        "frames": {
            "frame_paths": frame_paths,
            "frame_count": int(visual.get("frame_count") or len(frame_paths) or 0),
            "inspected_frames": int(visual.get("inspected_frames") or 0),
            "frame_summaries": visual.get("frame_summaries") or [],
            "face_total": int(visual.get("face_total") or 0),
            "max_faces_single_frame": int(visual.get("max_faces_single_frame") or 0),
            "male_count": int(visual.get("male_count") or 0),
            "female_count": int(visual.get("female_count") or 0),
            "pair_frames": int(visual.get("pair_frames") or 0),
            "has_both": bool(visual.get("has_both")),
            "thumbnail_faces": visual.get("thumbnail_faces") or {},
            "signals": visual.get("signals") or [],
            "error": visual.get("reason") or "",
        },
    }
    return bundle


def duration_gate_from_metadata(metadata: dict[str, Any]) -> dict[str, Any]:
    duration_seconds = parse_duration_seconds(metadata.get("duration"))
    if duration_seconds <= 0:
        return filter_check(
            False,
            "未能从页面公开信息确认视频时长。",
            confidence="low",
            duration_seconds=duration_seconds,
            min_seconds=FILTER_DURATION_MIN_SEC,
            max_seconds=FILTER_DURATION_MAX_SEC,
        )
    passed = FILTER_DURATION_MIN_SEC <= duration_seconds <= FILTER_DURATION_MAX_SEC
    reason = (
        f"视频时长 {duration_seconds:.0f}s，在 {FILTER_DURATION_MIN_SEC}-{FILTER_DURATION_MAX_SEC}s 范围内。"
        if passed
        else f"视频时长 {duration_seconds:.0f}s，不在 {FILTER_DURATION_MIN_SEC}-{FILTER_DURATION_MAX_SEC}s 范围内。"
    )
    return filter_check(
        passed,
        reason,
        confidence="high",
        duration_seconds=duration_seconds,
        min_seconds=FILTER_DURATION_MIN_SEC,
        max_seconds=FILTER_DURATION_MAX_SEC,
    )


def normalize_llm_filter_check(raw: Any, fallback_reason: str) -> dict[str, Any]:
    if not isinstance(raw, dict):
        return filter_check(False, fallback_reason, confidence="low")
    passed = bool(raw.get("passed"))
    confidence = str(raw.get("confidence") or "medium").strip().lower()
    reason = str(raw.get("reason") or fallback_reason).strip()
    evidence = raw.get("evidence") if isinstance(raw.get("evidence"), list) else []
    return filter_check(passed, reason, confidence=confidence, evidence=evidence)


def fallback_story_filter_decision(evidence_bundle: dict[str, Any], duration_check: dict[str, Any]) -> dict[str, Any]:
    frames = evidence_bundle.get("frames") or {}
    audio = evidence_bundle.get("audio") or {}
    transcript = str(audio.get("full_transcript") or "").strip()
    speaker_hints = audio.get("speaker_hints") or []
    max_faces = int(frames.get("max_faces_single_frame") or 0)
    pair_frames = int(frames.get("pair_frames") or 0)
    multi_passed = max_faces >= 2 or pair_frames >= 1 or len(speaker_hints) >= 2
    story_terms = ["误会", "冲突", "反转", "争吵", "发现", "被骗", "解释", "结局", "plot", "conflict", "twist"]
    story_passed = len(transcript) >= 80 and any(term.lower() in transcript.lower() for term in story_terms)
    multi_check = filter_check(
        multi_passed,
        "基于关键帧人脸数和音频说话人线索做本地兜底判断。" if multi_passed else "缺少两个以上可区分人物/角色的稳定证据。",
        confidence="low",
        evidence=[*speaker_hints, *[str(item) for item in (frames.get("signals") or [])]][:6],
    )
    story_check = filter_check(
        story_passed,
        "转写里出现基本剧情推进线索。" if story_passed else "缺少可稳定判断剧情结构的模型结果。",
        confidence="low",
        evidence=[],
    )
    final_passed = bool(duration_check.get("passed") and multi_check.get("passed") and story_check.get("passed"))
    return {
        "duration_check": duration_check,
        "multi_character_check": multi_check,
        "story_check": story_check,
        "final_result": "passed" if final_passed else "rejected",
        "bucket": "high" if final_passed else "low",
        "confidence": "low",
        "reason": "三轮规则均通过。" if final_passed else "三轮规则未全部通过。",
        "signals": [
            f"时长：{'通过' if duration_check.get('passed') else '不通过'}",
            f"多人物：{'通过' if multi_check.get('passed') else '不通过'}",
            f"剧情：{'通过' if story_check.get('passed') else '不通过'}",
        ],
        "used_llm": False,
    }


def classify_story_candidate(evidence_bundle: dict[str, Any]) -> dict[str, Any]:
    metadata = evidence_bundle.get("metadata") or {}
    duration_check = duration_gate_from_metadata({"duration": metadata.get("duration") or ""})
    if not FILTER_USE_LLM or not GOOGLE_API_KEY:
        raise RuntimeError("LLM筛选未启用或 GOOGLE_API_KEY 缺失，不能完成多人物和剧情判断。")
    frame_paths = [Path(path) for path in ((evidence_bundle.get("frames") or {}).get("frame_paths") or []) if str(path or "").strip()]
    prompt = f"""
你是 Koko 的视频筛选器。你会收到一个 Kwai 外部链接整理出的证据包，以及三张关键帧图片（开头、中间、结尾）。

筛选目标：判断这条视频是否值得进入后续视频拆解。必须按三轮分别判断：
1. 时长是否在 {FILTER_DURATION_MIN_SEC}-{FILTER_DURATION_MAX_SEC} 秒内。
2. 整个视频是否出现两个以上可区分的人物/角色。夫妻、兄弟、朋友、一人分饰多角、电话对端、旁白与画面角色都可以算，但必须有证据。
3. 整个视频是否拥有剧情：至少有起因/目标、冲突或误会、推进、结果/反转/包袱之一；纯展示、跳舞、无情节口播、单句段子不算。

请基于完整音频转写、页面公开信息和三张关键帧判断。不要编造看不到/听不到的内容。

严格返回 JSON：
{{
  "duration_check": {{"passed": true, "confidence": "high/medium/low", "reason": "一句话", "evidence": ["证据"]}},
  "multi_character_check": {{"passed": true, "confidence": "high/medium/low", "reason": "一句话", "evidence": ["证据"]}},
  "story_check": {{"passed": true, "confidence": "high/medium/low", "reason": "一句话", "evidence": ["证据"]}},
  "overall_confidence": "high/medium/low",
  "reason": "最终一句话解释"
}}
""".strip()
    payload = {
        "metadata": metadata,
        "audio": evidence_bundle.get("audio") or {},
        "frames": {
            key: value
            for key, value in (evidence_bundle.get("frames") or {}).items()
            if key != "frame_paths"
        },
        "hard_duration_check": duration_check,
    }
    result, _, model = run_gemini_file_json_prompt(frame_paths[:3], payload, prompt, "story candidate classification")
    multi_check = normalize_llm_filter_check(result.get("multi_character_check"), "未能稳定判断是否有两个以上可区分人物/角色。")
    story_check = normalize_llm_filter_check(result.get("story_check"), "未能稳定判断是否有剧情结构。")
    llm_duration = normalize_llm_filter_check(result.get("duration_check"), duration_check.get("reason") or "")
    duration_check["llm_reason"] = llm_duration.get("reason") or ""
    duration_check["llm_passed"] = bool(llm_duration.get("passed"))
    final_passed = bool(duration_check.get("passed") and multi_check.get("passed") and story_check.get("passed"))
    confidence = str(result.get("overall_confidence") or "medium").strip().lower()
    if confidence not in {"high", "medium", "low"}:
        confidence = "medium"
    return {
        "duration_check": duration_check,
        "multi_character_check": multi_check,
        "story_check": story_check,
        "final_result": "passed" if final_passed else "rejected",
        "bucket": "high" if final_passed else "low",
        "confidence": confidence,
        "reason": str(result.get("reason") or ("三轮规则均通过。" if final_passed else "三轮规则未全部通过。")).strip(),
        "signals": [
            f"时长：{'通过' if duration_check.get('passed') else '不通过'}",
            f"多人物：{'通过' if multi_check.get('passed') else '不通过'}",
            f"剧情：{'通过' if story_check.get('passed') else '不通过'}",
        ],
        "model_used": model,
        "used_llm": True,
    }


def parse_video_display_name(url: object, index: int = 0) -> str:
    fallback = f"视频 {index + 1}"
    text = str(url or "").strip()
    if not text:
        return fallback
    try:
        parsed = urllib.parse.urlparse(text)
        parts = [part for part in parsed.path.split("/") if part]
        handle = next((part for part in parts if part.startswith("@")), "")
        video_id = ""
        if "video" in parts:
            video_index = parts.index("video")
            if video_index + 1 < len(parts):
                video_id = parts[video_index + 1]
        if handle and video_id:
            return f"{handle} / {video_id}"
        if handle:
            return handle
        if video_id:
            return f"video / {video_id}"
        hostname = (parsed.hostname or "").replace("www.", "")
        return hostname or fallback
    except Exception:
        return fallback


def summarize_job_focus(job: dict[str, Any]) -> dict[str, Any]:
    items = job.get("items") or []
    active_item = next((item for item in items if str(item.get("status") or "").strip() == "running"), None)
    if active_item:
        index = int(active_item.get("index") or 0)
        return {
            "job_id": str(job.get("id") or "").strip(),
            "item_id": str(active_item.get("id") or "").strip(),
            "kind": "analysis",
            "title": active_item.get("title") or parse_video_display_name(active_item.get("video_url"), index),
            "video_url": active_item.get("video_url") or "",
            "stage": active_item.get("stage") or "",
            "stage_message": active_item.get("stage_message") or "",
            "item_index": index,
        }
    review_item = next((item for item in items if str(item.get("review_status") or "").strip() == "running"), None)
    if review_item:
        index = int(review_item.get("index") or 0)
        return {
            "job_id": str(job.get("id") or "").strip(),
            "item_id": str(review_item.get("id") or "").strip(),
            "kind": "review",
            "title": review_item.get("title") or parse_video_display_name(review_item.get("video_url"), index),
            "video_url": review_item.get("video_url") or "",
            "stage": review_item.get("review_stage") or "review",
            "stage_message": review_item.get("review_message") or "",
            "item_index": index,
        }
    pending_item = next((item for item in items if str(item.get("status") or "").strip() == "queued"), None)
    if pending_item:
        index = int(pending_item.get("index") or 0)
        return {
            "job_id": str(job.get("id") or "").strip(),
            "item_id": str(pending_item.get("id") or "").strip(),
            "kind": "analysis",
            "title": pending_item.get("title") or parse_video_display_name(pending_item.get("video_url"), index),
            "video_url": pending_item.get("video_url") or "",
            "stage": pending_item.get("stage") or "",
            "stage_message": pending_item.get("stage_message") or "",
            "item_index": index,
        }
    return {
        "job_id": str(job.get("id") or "").strip(),
        "item_id": "",
        "kind": "analysis",
        "title": parse_video_display_name(job.get("video_url") or "", 0),
        "video_url": job.get("video_url") or "",
        "stage": job.get("stage") or "",
        "stage_message": job.get("stage_message") or "",
        "item_index": 0,
    }


def build_system_queue_snapshot(current_job_id: str | None = None) -> dict[str, Any]:
    active_workloads: list[dict[str, Any]] = []
    queued_jobs: list[dict[str, Any]] = []
    current_job_position: int | None = None
    current_job_ahead = 0
    with job_lock:
        queued_order = list(job_queue)
        for job in jobs.values():
            items = job.get("items") or []
            is_active = any(
                str(item.get("status") or "").strip() == "running" or str(item.get("review_status") or "").strip() == "running"
                for item in items
            )
            if is_active:
                active_workloads.append(summarize_job_focus(job))
        for idx, job_id in enumerate(queued_order):
            job = jobs.get(job_id)
            if not job:
                continue
            focus = summarize_job_focus(job)
            queue_position = len(active_workloads) + idx + 1
            queued_jobs.append(
                {
                    "job_id": job_id,
                    "title": focus["title"],
                    "video_url": focus["video_url"],
                    "stage": focus["stage"],
                    "stage_message": focus["stage_message"],
                    "queue_position": queue_position,
                }
            )
            if current_job_id and job_id == current_job_id:
                current_job_position = queue_position
                current_job_ahead = max(0, queue_position - 1)
        if current_job_id and current_job_position is None:
            for idx, focus in enumerate(active_workloads):
                if focus.get("job_id") == current_job_id:
                    current_job_position = idx + 1
                    current_job_ahead = idx
                    break
    return {
        "running_count": len(active_workloads),
        "queued_count": len(queued_jobs),
        "active_workloads": active_workloads,
        "queued_jobs": queued_jobs,
        "current_job_position": current_job_position,
        "current_job_ahead": current_job_ahead,
    }


CONTENT_TYPE_RULES: list[tuple[str, list[str]]] = [
    ("夫妻暧昧", ["出轨", "暧昧", "好色", "黄段子", "撬墙角", "第三者", "偷看", "吃醋", "抓奸", "情人"]),
    ("夫妻整蛊/冲突", ["夫妻", "妻子", "丈夫", "老公", "老婆", "情侣", "吵架", "欺骗", "算计", "妻管严", "整蛊", "反转"]),
    ("家庭整蛊", ["家庭", "妈妈", "爸爸", "母亲", "父亲", "儿子", "女儿", "孩子", "亲戚", "婆婆"]),
    ("朋友整蛊", ["朋友", "同事", "闺蜜", "兄弟", "老板", "员工", "路人", "街头", "公共场景", "世界杯", "偷手机", "便利店", "骗局", "赖账", "偷吃", "偷懒"]),
]
CONTENT_TYPE_RULES_PT: list[tuple[str, list[str]]] = [
    ("夫妻暧昧", ["traição", "infiel", "amante", "outra mulher", "outro homem", "ciúme", "ciume", "seduz", "paquera", "íntim", "intim", "beijo", "mulher bonita"]),
    ("夫妻整蛊/冲突", ["esposa", "marido", "casal", "namorado", "namorada", "briga", "discute", "reclama", "conflito", "finge", "mentira", "segredo", "pegadinha"]),
    ("家庭整蛊", ["família", "familia", "mãe", "mae", "pai", "filho", "filha", "criança", "crianca", "sogra", "irmão", "irmao", "irmã", "irma"]),
    ("朋友整蛊", ["amigo", "amiga", "colega", "chefe", "funcionário", "funcionario", "cliente", "vizinho", "rua", "público", "publico", "loja", "conveniência", "conveniencia", "copa do mundo", "celular", "golpe", "pegadinha", "dinheiro"]),
]

ALLOWED_CONTENT_TYPES = {label for label, _ in CONTENT_TYPE_RULES}
DEFAULT_CONTENT_TYPE = "朋友整蛊"
LIBRARY_FILTER_LABELS = [
    "夫妻整蛊/冲突",
    "夫妻暧昧",
    "家庭整蛊",
    "朋友整蛊",
]
CONTENT_TYPE_CHOICE_TEXT = "、".join(LIBRARY_FILTER_LABELS)


def has_word_signal(text: str, terms: list[str]) -> bool:
    return any(re.search(rf"(?<![\wÀ-ÿ]){re.escape(term)}(?![\wÀ-ÿ])", text, flags=re.I) for term in terms)


def normalize_creator_content_type(value: object, text: str = "") -> str:
    current = str(value or "").strip()
    lowered = str(text or "").lower()
    legacy_flirt = {"夫妻暧昧", "夫妻出轨", "夫妻好色", "夫妻黄段子", "撬墙角", "Relacionamento de casal"}
    legacy_couple = {"夫妻整蛊/冲突", "夫妻吵架", "夫妻欺骗", "夫妻算计", "妻管严", "夫妻整蛊", "夫妻关系", "夫妻/情侣", "夫妻情感"}
    legacy_family = {"家庭/亲子"}
    legacy_friends = {
        "朋友整蛊",
        "整蛊",
        "整蛊恶搞",
        "骗局反转",
        "赖账",
        "赖账/金钱冲突",
        "骗子",
        "偷奸耍滑",
        "偷吃东西",
        "偷吃/偷懒/耍小聪明",
        "Popular",
        "Golpe e reviravolta",
        "Pegadinha",
        "Esperteza cotidiana",
        "待分类",
        "热门",
        "",
    }
    family_cn_terms = ["妈妈", "爸爸", "母亲", "父亲", "孩子", "儿子", "女儿", "小孩", "宝宝", "亲戚", "婆婆", "岳母", "兄弟", "姐妹"]
    family_pt_terms = ["mãe", "mae", "pai", "filho", "filha", "criança", "crianca", "crianças", "criancas", "bebê", "bebe", "sogra", "sogro", "irmão", "irmao", "irmã", "irma"]
    has_family_text = any(term in lowered for term in family_cn_terms) or any(
        re.search(rf"(?<![\wÀ-ÿ]){re.escape(term)}(?![\wÀ-ÿ])", lowered, flags=re.I)
        for term in family_pt_terms
    )
    flirt_terms = ["trai", "infiel", "amante", "ciúme", "ciume", "暧昧", "出轨", "好色", "撬墙角", "mulher bonita", "namorado", "namorada", "paquera", "beijo"]
    couple_terms = ["marido", "esposa", "casal", "夫妻", "妻子", "丈夫", "情侣"]
    has_flirt = current in legacy_flirt or any(term in lowered for term in flirt_terms)
    has_couple = current in legacy_couple or any(term in lowered for term in couple_terms)
    if current in legacy_family or has_family_text:
        return "家庭整蛊"
    if has_flirt:
        return "夫妻暧昧"
    if has_couple:
        return "夫妻整蛊/冲突"
    if current in ALLOWED_CONTENT_TYPES and current != "家庭整蛊":
        return current
    if current in legacy_friends:
        return "朋友整蛊"
    if has_family_text:
        return "家庭整蛊"
    return DEFAULT_CONTENT_TYPE

CONTENT_TYPE_CLASSIFY_PROMPT = f"""你是一个短视频脚本分类器。

你会收到已经整理完成的脚本信息，尤其是：
1. 标题
2. 故事梗概（whole_video_summary）
3. 包袱机制原因
4. 可选的替换方案
5. 脚本表最后时间码推导出的时长信息
6. 可选的路由说明

你的任务不是改写脚本，而是根据“最终语义”从固定分类白名单里选一个最合适的类型。

分类原则：
1. 只能从四个类型里选：{CONTENT_TYPE_CHOICE_TEXT}。
2. `夫妻整蛊/冲突`：主角关系是夫妻、情侣、男女朋友、丈夫妻子，核心是吵架、误会、欺骗、算计、整蛊、反转、日常冲突。
3. `夫妻暧昧`：主角关系仍然是夫妻/情侣，但核心笑点是暧昧、出轨、第三者、吃醋、好色、亲密误会、黄段子、撬墙角。
4. `家庭整蛊`：父母、孩子、兄弟姐妹、亲戚、婆媳等家庭成员是主轴。
5. `朋友整蛊`：朋友、同事、老板员工、顾客、邻居、路人、公共场景、街头骗局、偷手机、便利店、世界杯等非夫妻/非家庭脚本都归到这里。
6. 如果证据不够或无法归类，也选 `朋友整蛊`，不要输出其他标签。
7. 时长信息只作为辅助理解脚本结构，不允许输出时间作为分类标签。

输出严格 JSON：
{{
  "content_type": "白名单中的一个值",
  "confidence": "high/medium/low",
  "reasoning": "一句话说明为什么是这个类，重点讲语义依据"
}}
"""


def detect_content_type(script: dict[str, Any], bundle: dict[str, Any] | None = None) -> str:
    return detect_content_type_decision(script, bundle).get("content_type") or DEFAULT_CONTENT_TYPE


def classify_content_type_with_llm(
    script: dict[str, Any],
    bundle: dict[str, Any] | None,
    key: str,
    models: list[str],
) -> dict[str, Any] | None:
    if not key or run_text_json_prompt_with_fallback is None:
        return None
    routing = (bundle or {}).get("routing") or script.get("type_router") or {}
    duration_seconds = script_duration_seconds(script)
    duration_bucket = creator_duration_bucket(duration_seconds)
    payload = {
        "title": script.get("title") or "",
        "whole_video_summary": script.get("whole_video_summary") or "",
        "summary": script.get("summary") or "",
        "mechanism_reason": ((script.get("mechanism") or {}).get("reason") or ""),
        "duration_seconds": round(duration_seconds, 2) if duration_seconds > 0 else 0,
        "duration_bucket": duration_bucket,
        "duration_label_pt": CREATOR_DURATION_LABELS.get(duration_bucket, {}).get("pt", ""),
        "key_points": script.get("key_points") or script.get("points") or [],
        "replaceable_parts": script.get("replaceable_parts") or [],
        "script_rows": (script.get("rows") or script.get("script_table") or [])[:12] if isinstance(script.get("rows") or script.get("script_table") or [], list) else [],
        "routing": {
            "primary_type": routing.get("primary_type") or "",
            "subtype_guess": routing.get("subtype_guess") or "",
            "reasoning_summary": routing.get("reasoning_summary") or "",
        },
        "allowed_content_types": LIBRARY_FILTER_LABELS,
    }
    try:
        result, _, _ = run_text_json_prompt_with_fallback(
            payload,
            key,
            models,
            CONTENT_TYPE_CLASSIFY_PROMPT,
            "content type classification",
        )
    except Exception:
        return None
    if not isinstance(result, dict):
        return None
    content_type = str(result.get("content_type") or "").strip()
    if content_type not in ALLOWED_CONTENT_TYPES and content_type != DEFAULT_CONTENT_TYPE:
        content_type = normalize_creator_content_type(content_type)
    confidence = str(result.get("confidence") or "").strip().lower()
    if confidence not in {"high", "medium", "low"}:
        confidence = "low"
    return {
        "content_type": content_type or DEFAULT_CONTENT_TYPE,
        "reasoning": str(result.get("reasoning") or "").strip(),
        "confidence": confidence,
        "source": "auto",
    }


def keyword_fallback_content_type(script: dict[str, Any], bundle: dict[str, Any] | None = None) -> str:
    routing = (bundle or {}).get("routing") or script.get("type_router") or {}
    primary = routing.get("primary_type") or ""
    subtype = routing.get("subtype_guess") or ""
    for candidate in [subtype, primary]:
        if candidate:
            text = str(candidate)
            if text in ALLOWED_CONTENT_TYPES:
                return text
    text = " ".join(
        str(x or "")
        for x in [
            routing.get("reasoning_summary"),
            script.get("whole_video_summary"),
            script.get("content_summary"),
            script.get("summary"),
            script.get("title"),
            (script.get("mechanism") or {}).get("reason"),
            (script.get("type_router") or {}).get("reasoning_summary"),
            " ".join(str(item or "") for item in script.get("key_points") or []),
            " ".join(str(item or "") for item in script.get("replaceable_parts") or []),
        ]
    )
    for label, keywords in CONTENT_TYPE_RULES:
        hit_count = sum(1 for word in keywords if word in text)
        if hit_count >= 2:
            return label
    lowered = text.lower()
    scored: list[tuple[int, str]] = []
    for label, keywords in CONTENT_TYPE_RULES_PT:
        hit_count = sum(1 for word in keywords if word in lowered)
        if hit_count:
            scored.append((hit_count, label))
    if scored:
        scored.sort(key=lambda item: (item[0], 1 if item[1].startswith("夫妻") else 0), reverse=True)
        return scored[0][1]
    return DEFAULT_CONTENT_TYPE


def detect_content_type_decision(
    script: dict[str, Any],
    bundle: dict[str, Any] | None = None,
    *,
    existing_type: str = "",
    existing_source: str = "",
    use_llm: bool = True,
    use_keyword_fallback: bool = True,
) -> dict[str, Any]:
    normalized_existing_type = str(existing_type or "").strip()
    normalized_existing_source = str(existing_source or "").strip().lower()
    if normalized_existing_source == "manual" and (normalized_existing_type in ALLOWED_CONTENT_TYPES or normalized_existing_type == DEFAULT_CONTENT_TYPE):
        return {
            "content_type": normalized_existing_type or DEFAULT_CONTENT_TYPE,
            "content_type_source": "manual",
            "content_type_reasoning": "Manual override",
            "content_type_confidence": "manual",
        }
    if use_llm:
        llm_result = classify_content_type_with_llm(script, bundle, GOOGLE_API_KEY, unique_models(*MODEL_CANDIDATES))
        if llm_result:
            content_type = llm_result.get("content_type") or DEFAULT_CONTENT_TYPE
            confidence = llm_result.get("confidence") or "low"
            if content_type != DEFAULT_CONTENT_TYPE or confidence in {"high", "medium"}:
                return {
                    "content_type": content_type,
                    "content_type_source": "auto",
                    "content_type_reasoning": llm_result.get("reasoning") or "LLM semantic classification",
                    "content_type_confidence": confidence,
                }
    if not use_keyword_fallback:
        return {
            "content_type": DEFAULT_CONTENT_TYPE,
            "content_type_source": "auto",
            "content_type_reasoning": "LLM classification unavailable or inconclusive; keyword fallback disabled for Creator import.",
            "content_type_confidence": "low",
        }
    content_type = keyword_fallback_content_type(script, bundle)
    return {
        "content_type": content_type,
        "content_type_source": "auto",
        "content_type_reasoning": "Keyword/routing fallback",
        "content_type_confidence": "low" if content_type == DEFAULT_CONTENT_TYPE else "medium",
    }


def detect_content_type_for_output(
    output_dir: Path,
    script: dict[str, Any] | None = None,
    bundle: dict[str, Any] | None = None,
    *,
    existing_type: str = "",
    existing_source: str = "",
) -> str:
    return detect_content_type_decision_for_output(
        output_dir,
        script,
        bundle,
        existing_type=existing_type,
        existing_source=existing_source,
    ).get("content_type") or DEFAULT_CONTENT_TYPE


def detect_content_type_decision_for_output(
    output_dir: Path,
    script: dict[str, Any] | None = None,
    bundle: dict[str, Any] | None = None,
    *,
    existing_type: str = "",
    existing_source: str = "",
    use_llm: bool = True,
) -> dict[str, Any]:
    script_json = script or read_json(output_dir / "script_table.json") or {}
    bundle_json = bundle or read_json(output_dir / "evidence_bundle.json") or {}
    type_router = read_json(output_dir / "type_router.json") or {}
    if type_router and not bundle_json.get("routing"):
        bundle_json = {**bundle_json, "routing": type_router}
    if type_router and not script_json.get("type_router"):
        script_json = {**script_json, "type_router": type_router}
    return detect_content_type_decision(
        script_json,
        bundle_json,
        existing_type=existing_type,
        existing_source=existing_source,
        use_llm=use_llm,
    )


def choose_script_rows(script: dict[str, Any]) -> list[dict[str, Any]]:
    rows = script.get("rows")
    if isinstance(rows, list) and rows:
        return [row for row in rows if isinstance(row, dict)]
    rows = script.get("synthesized_segments")
    if isinstance(rows, list) and rows:
        return [row for row in rows if isinstance(row, dict)]
    return []


def fill_text(value: Any, fallback: str = "无") -> str:
    text = str(value or "").strip()
    return text if text else fallback


def normalize_replaceable_parts(value: Any, fallback: list[dict[str, str]] | None = None) -> list[dict[str, str]]:
    items = value if isinstance(value, list) else []
    normalized = [
        {
            "label": fill_text((item or {}).get("label") or (item or {}).get("title") or (item or {}).get("name"), "替换方案"),
            "text": fill_text((item or {}).get("text") or (item or {}).get("description") or (item or {}).get("value"), "无"),
        }
        for item in items
        if isinstance(item, dict)
    ]
    if normalized:
        return normalized[:8]
    return json.loads(json.dumps(fallback or [{"label": "替换方案", "text": "无"}], ensure_ascii=False))


def normalize_core_viral_points(value: Any, fallback: list[dict[str, str]] | None = None) -> list[dict[str, str]]:
    items = value if isinstance(value, list) else []
    normalized = [
        {
            "label": fill_text((item or {}).get("label") or (item or {}).get("title") or (item or {}).get("name"), "爆点"),
            "text": fill_text((item or {}).get("text") or (item or {}).get("description") or (item or {}).get("value"), "无"),
        }
        for item in items
        if isinstance(item, dict) and str((item or {}).get("text") or (item or {}).get("description") or (item or {}).get("value") or "").strip()
    ]
    if normalized:
        return normalized[:8]
    return json.loads(json.dumps(fallback or [], ensure_ascii=False))


HEAVY_REVIEW_FAILURE_LAYERS = {
    "story_spine",
    "primary_analysis",
    "entity_mapping",
}

REVIEW_MODE_PARTIAL = "partial"
REVIEW_MODE_FULL = "full"
REVIEW_SCRIPT_KEYS = [
    "title",
    "route",
    "audio_information_score",
    "source_url",
    "whole_video_summary",
    "replaceable_parts",
    "rows",
    "mechanism",
]
REVIEW_CORE_CHANGE_KEYS = [
    "whole_video_summary",
    "rows",
    "mechanism",
]


def normalize_review_mode(value: Any) -> str:
    text = str(value or "").strip().lower()
    if text in {"full", "complete", "完全错误", "完全重做"}:
        return REVIEW_MODE_FULL
    return REVIEW_MODE_PARTIAL


def should_run_review_video_recheck(review_plan: dict[str, Any]) -> tuple[bool, str]:
    if not bool(review_plan.get("needs_video_recheck")):
        return False, "Review plan says the prior evidence is enough."
    layer = str(review_plan.get("likely_failure_layer") or "").strip()
    if layer in HEAVY_REVIEW_FAILURE_LAYERS:
        return True, f"Heavy review triggered for {layer}."
    return False, f"Skipping video recheck for lighter failure layer: {layer or 'unknown'}."


def extract_review_script_payload(payload: Any) -> dict[str, Any]:
    if not isinstance(payload, dict):
        return {}
    for key in ["corrected_script", "script", "script_table", "result_json"]:
        nested = payload.get(key)
        if isinstance(nested, dict):
            extracted = extract_review_script_payload(nested)
            if extracted:
                return extracted
    if any(key in payload for key in REVIEW_SCRIPT_KEYS):
        return payload
    return {}


def stable_json_fingerprint(value: Any) -> str:
    return json.dumps(value, ensure_ascii=False, sort_keys=True, separators=(",", ":"))


def review_script_changed(before: dict[str, Any], after: dict[str, Any], keys: list[str]) -> bool:
    for key in keys:
        if stable_json_fingerprint((before or {}).get(key)) != stable_json_fingerprint((after or {}).get(key)):
            return True
    return False


REVIEW_PLAN_PROMPT = """你是一个短视频分析复盘诊断器。

输入里会包含：
1. 旧的脚本结果
2. 旧的主分析、补证据、音频侧产物
3. 人类审阅者用自然语言写的反馈，说明这次分析哪里错了

你的任务不是直接重写脚本，而是先做“复盘诊断”：
- 判断这次错误更像是故事主轴理解错误、人物/关系映射错误、关键物体错误、因果链错误，还是仅仅需要局部复核
- 判断旧证据是否仍可用
- 判断是否需要重新让 Gemini 回看视频
- 如果需要回看，给出最值得复核的时间窗和关注点

要求：
- 把人类反馈当作“高优先级纠错假设”，但不是绝对真理
- 不要替旧答案辩护
- 优先寻找最省资源的纠正路径
- 如果你认为不需要重新看视频，也要解释原因

输出严格 JSON：
{
  "problem_summary": "一句话概括人类指出的问题",
  "likely_failure_layer": "final_refine/primary_analysis/supplement/type_router/entity_mapping/story_spine/unknown",
  "needs_video_recheck": true,
  "needs_structural_rewrite": true,
  "reasoning": "为什么这样判断",
  "focus_windows": [
    {
      "start": "00:20",
      "end": "00:55",
      "reason": "这里决定了故事主轴"
    }
  ],
  "focus_entities": [
    "需要重点确认的人物、关系、道具或动作"
  ],
  "correction_goal": "纠偏后最应该纠正成什么方向",
  "confidence": "low/medium/high"
}
"""


REVIEW_VIDEO_PROMPT = """你现在是在做一次“复盘回看”，不是从零分析整条视频。

背景：
- 之前这条视频已经被分析过，但人类审阅者指出核心理解可能错了
- 你需要带着这个反馈，重新确认关键证据

要求：
- 不要为旧答案辩护
- 把人类反馈当成一个需要验证的纠错假设
- 只关注：人物关系、关键道具、关键动作、关键对白、故事主轴
- 如果某个关键点仍无法证实，要明确说不确定
- 如果给出了 focus windows，请优先围绕这些窗口回看

输出严格 JSON：
{
  "feedback_hypothesis": "人类反馈想纠正的核心点",
  "verification_result": "confirmed/partially_confirmed/rejected/inconclusive",
  "corrected_story_spine": "基于回看后，更稳妥的一句话故事主轴",
  "evidence_findings": [
    {
      "time": "00:20-00:55",
      "finding": "这一段真正发生了什么",
      "supports_feedback": true,
      "confidence": "low/medium/high"
    }
  ],
  "entity_corrections": [
    {
      "wrong": "之前模型误写的东西",
      "correct": "更可能正确的理解",
      "evidence": "为什么"
    }
  ],
  "rewrite_notes": [
    "后续重写脚本时必须遵守的纠偏说明"
  ]
}
"""


REVIEW_REFINE_PROMPT = """你是一个“复盘重做”脚本整理器。

输入里会有：
1. 旧脚本结果
2. 人类反馈
3. review_plan（复盘诊断）
4. 可选的 review_video_recheck（带着纠错假设重新回看视频后的结果）

你的任务：
- 不要机械修几个字
- 要根据人类指出的核心问题，对脚本做结构性纠偏
- 如果 review_video_recheck 提供了更可靠的故事主轴或实体纠正，优先使用
- 如果 review_video_recheck 不充分，也要基于 review_plan 做保守改写
- 必须保持 v2 script_table.json 结构
- `dialogue_or_audio` 是最终 HTML 里“对话”部分的直接来源，必须继续保持中文 1:1 直译：只翻译，不改写，不润色，不概括，不补解释，不合并句子，不删减语气词
- 如果某个实体仍不确定，宁可保守，也不要继续沿用明显错误的旧判断

成品写法 SOP（必须遵守）：
- `title` 必须写成完整的故事钩子句，不要只写前半段设定。要尽量把“人物 + 冲突/伪装 + 最终落点/反转”放在同一句里，优先使用 `不料`、`却`、`反而`、`最终`、`结果` 等连接词把结局说满。
- `whole_video_summary` 要按“起因 -> 推进 -> 对质/证据 -> 最终落点”来写，重点落在具体剧情和最终结果，不要用抽象的“揭示了复杂关系/人性弱点/社会判断”来替代真正的结局描述。
- 背后原因可以写，但必须嵌在具体人物动机里，例如“他为了维护面子选择否认真相”，而不是另起一段空泛说教。
- 如果人物关系在复盘后已经足够稳定，`title` 和 `whole_video_summary` 优先用自然角色称呼（丈夫/妻子/邻居等），不要为了机械一致性继续保留 `男性A/女性A`。
- `core_viral_points` 重新保留，但要用新的简洁格式来写：每一项都是一句能直接说明“这条视频为什么成立/为什么抓人”的短爆点，不要写成长报告。
- `replaceable_parts` 必须写成可以直接套用的替换方案，不是建议。每一项都要包含“替换成什么人物/场景/道具/冲突”和“替换后脚本主轴如何变化”。
- `mechanism.items[*].text` 继续保持具体，尤其 `背后原因` 必须落在这条视频里的真实人物心理和关系机制上。
- 如果旧版本里有空泛的抽象收束句，请在复盘版里收掉，改成更具体的剧情归纳和结局描述。

输出严格 JSON：
{
  "title": "视频总结归纳 + 脚本表",
  "route": "audio-sop 或 keyframe-sop",
  "audio_information_score": "0/10 到 10/10",
  "source_url": "原视频链接",
  "whole_video_summary": "纠偏后的完整总结",
  "replaceable_parts": [
    {"label": "替换方案名", "text": "直接可执行的替换方案：把原脚本中的哪些人物/场景/道具/冲突替换成什么，并说明替换后的故事主轴。"}
  ],
  "rows": [
    {
      "source_url": "原视频链接",
      "time": "00:00-00:15",
      "visual_content": "这一段整体看到了什么",
      "action": "这一段动作如何推进",
      "dialogue_or_audio": "中文 1:1 直译，不改写，不概括，不润色",
      "integrated_summary": "如果有必要可补充"
    }
  ],
  "mechanism": {
    "title": "包袱机制",
    "items": [
      {"label": "铺垫", "text": "..."},
      {"label": "违和点", "text": "..."},
      {"label": "反转点", "text": "..."},
      {"label": "笑点落点", "text": "..."},
      {"label": "背后原因", "text": "..."}
    ]
  }
}
"""


PARTIAL_REVIEW_REFINE_PROMPT = REVIEW_REFINE_PROMPT + """

本次模式：部分错误。

人类反馈指出的是“当前脚本大体方向可用，但漏掉/误判了一个关键点”。你必须以 current_script 为基础做全局修订：
- 重点修正人类指出的核心错误点，并让这个纠偏贯穿 title、whole_video_summary、replaceable_parts、rows、mechanism
- 允许保留旧脚本里仍然正确的时间段、图片引用、分镜顺序和视觉字段
- 不要只改标题，不要只改概述，不要返回局部补丁
- 输出必须是完整 v2 script_table.json 同格式对象
"""


FULL_REVIEW_REFINE_PROMPT = REVIEW_REFINE_PROMPT + """

本次模式：完全错误。

人类反馈指出的是“当前故事主轴或核心理解整体错误”。你必须把 review_video_recheck 当作新的高优先级证据来重建脚本：
- title、whole_video_summary、replaceable_parts、rows、mechanism 都要围绕新的故事主轴重新整理
- 可以保留 confirmed 的时间段和图片引用，但不要沿用旧故事主轴
- 如果回看证据不充分，要明确采用更保守的剧情表达，不要把不确定内容写死
- 输出必须是完整 v2 script_table.json 同格式对象
"""


CHAT_SCRIPT_EDIT_PROMPT = """你是 Koko 的脚本修稿助手，负责根据用户的自然语言反馈，直接修改当前短视频拆解脚本。

输入里会包含：
1. current_script：当前已经生成的完整 script_table.json
2. user_message：用户这一次想让你修改的内容
3. conversation：同一条脚本的历史修稿对话
4. edit_mode：minor、major 或 replace

你的权限边界：
- 只能修改当前脚本 JSON 内容，不要修改文件名、任务状态、视频文件或脚本库状态
- 可以全局修改 title、whole_video_summary、replaceable_parts、rows、mechanism
- 如果用户只指出一个小问题，也要让这个纠正自然贯穿相关标题、总结、替换方案和分镜，不要只机械改一个词
- 如果用户要求大改，要以 current_script 为基础重组故事主轴，但不要假装重新看过视频
- 如果 edit_mode 是 replace，说明用户选择了一个替换方案。你必须把这个替换方案当作新的创作约束，直接重写 title、whole_video_summary、replaceable_parts、rows、mechanism，让整份脚本变成替换后的新版本；不要只解释方案。
- 如果用户的问题需要重新看视频才能确认，请在 assistant_message 里说明“需要重新看视频”，并仍然尽量做保守文本修订
- 不要删除 rows 里的图片引用或时间顺序
- `dialogue_or_audio` 必须保持中文 1:1 直译风格：只翻译，不改写，不润色，不概括，不补解释，不合并句子

输出严格 JSON，必须包含：
{
  "assistant_message": "像聊天一样告诉用户你改了什么；简短、具体",
  "change_summary": [
    "修改点1",
    "修改点2"
  ],
  "script": {
    "title": "修改后的标题",
    "route": "保留或修正后的 route",
    "audio_information_score": "保留或修正后的分数",
    "source_url": "原视频链接",
    "whole_video_summary": "修改后的完整总结",
    "replaceable_parts": [
      {"label": "替换方案名", "text": "直接可执行的替换方案"}
    ],
    "rows": [
      {
        "source_url": "原视频链接",
        "time": "00:00-00:15",
        "visual_content": "这一段整体看到了什么",
        "action": "这一段动作如何推进",
        "dialogue_or_audio": "中文 1:1 直译",
        "integrated_summary": "可选补充"
      }
    ],
    "mechanism": {
      "title": "包袱机制",
      "items": [
        {"label": "铺垫", "text": "..."},
        {"label": "违和点", "text": "..."},
        {"label": "反转点", "text": "..."},
        {"label": "笑点落点", "text": "..."},
        {"label": "背后原因", "text": "..."}
      ]
    }
  }
}
"""


REPLACEMENT_PLAN_REFRESH_PROMPT = """你是 Koko 的脚本替换方案生成器。

你会收到一份已经存在的短视频脚本 JSON。你的任务不是重写整份脚本，而是基于当前最新脚本内容，重新整理“替换方案”。

要求：
- 只输出 `replaceable_parts`
- 最终会同步到 Koko Creator 给创作者看，所以必须短、直、可拍摄
- 给 1-3 项即可，每一项都必须是“可以直接拿去改整份脚本”的替换元素，不是泛泛建议
- 优先使用 `人物关系`、`冲突事项`、`场景`、`道具/诱因`、`结尾反转` 这类短标签
- `text` 用一句话明确写出替换成什么，以及替换后主冲突怎么变
- 必须和当前脚本标题、整体梗概、分镜脚本保持一致，不能沿用旧主轴
- 最多 3 条，不要写长段分析
- 用中文输出

输出严格 JSON：
{
  "replaceable_parts": [
    {
      "label": "替换方案名",
      "text": "直接可执行的替换方案"
    }
  ]
}
"""


CORE_VIRAL_POINTS_REFRESH_PROMPT = """你是 Koko 的短视频爆点整理助手。

你会收到一份已经存在的短视频脚本 JSON。你的任务不是重写整份脚本，而是基于当前最新脚本内容，重新整理“核心爆点”。

要求：
- 只输出 `core_viral_points`
- 最终会同步到 Koko Creator 给创作者看，所以必须短、直、可拍摄
- 每一项都必须是短句式爆点，不要写成长段分析
- 要直接说明这条视频为什么抓人、反差点在哪、笑点/冲突为什么成立
- 每项尽量一句话说完
- 结合当前标题、整体梗概、分镜脚本来写，不能沿用旧主轴
- 给 1-3 条即可，最多 3 条
- 用中文输出

输出严格 JSON：
{
  "core_viral_points": [
    {
      "label": "爆点名",
      "text": "一句短爆点"
    }
  ]
}
"""


ERROR_CASE_REVIEW_PROMPT = """你是一个短视频分析系统的“错误案例复盘器”。

你会收到一次已经完成的“复盘重做”案例，里面包括：
1. 原始脚本
2. 用户指出的核心问题（review_feedback）
3. 复盘诊断（review_plan）
4. 可选的视频回看结果（review_video_recheck）
5. 主链当时实际走过的流程快照（process_trace）
6. 修正后的最新脚本

你的任务不是再写脚本，而是为“错误案例库”输出一份结构化复盘，重点回答：
- 这次最核心到底错在哪
- 更偏向是哪一层出了问题：
  - gemini_analysis
  - v2_analysis
  - consistency_audit
  - targeted_recheck
  - arbitration
  - final_output
  - unknown
- 为什么主流程当时没有拦住这个问题
- 从这次实际执行链路看，哪一步“没找到”、哪一步“没触发”、哪一步“误判放行”

要求：
1. 重点讲“流程上错在哪”，不是只复述剧情。
2. 必须结合 process_trace 和 review_plan 来分析。
3. 如果某层其实执行了，但结果没有起到纠偏作用，也要明确指出。
4. 不要生成代码建议，不要改 prompt，只做错误归因和案例沉淀。
5. 输出严格 JSON。

输出格式：
{
  "core_issue": "一句话说明这次原始脚本的核心错误",
  "primary_failure_layer": "gemini_analysis/v2_analysis/consistency_audit/targeted_recheck/arbitration/final_output/unknown",
  "flow_failure_summary": "从执行流程角度看，这次为什么会错",
  "difference_summary": "原始脚本和修正后脚本的关键差异",
  "missed_or_weak_links": [
    {
      "layer": "具体层名",
      "status": "not_run/ran_but_missed/incorrect_pass/insufficient_signal",
      "reason": "为什么这层没有拦住问题"
    }
  ],
  "evidence_notes": [
    "和 review_plan/process_trace 对应的关键观察"
  ],
  "preventive_notes": [
    "以后遇到相似 case 时应重点关注的流程风险点"
  ],
  "confidence": "low/medium/high"
}
"""

PORTUGUESE_TRANSLATION_PROMPT = """你会收到一个脚本 JSON。你的任务不是重写，不是总结，不是润色，而是把所有面向用户展示的中文文本严格直译成葡萄牙语（pt-BR）。

硬规则：
1. 必须保留原 JSON 结构、字段名、数组长度、对象层级。
2. 只翻译“值里的自然语言文本”，不要改字段名。
3. 所有翻译必须尽量 1:1 忠实直译，不要概括，不要补充解释，不要删减，不要美化。
4. `dialogue_or_audio` 必须做逐句忠实翻译，不能保留中文，不能改写成摘要。
5. 时间、URL、文件路径、frame 路径、source_url、数字、布尔值保持原样。
6. `无`、空值占位等，翻成最直接的葡语占位，例如 `Sem conteúdo`，但不要删除字段。
7. 输出必须是可解析 JSON，不要带 Markdown。
"""


def language_copy(locale: str) -> dict[str, str]:
    if locale == "pt":
        return {
            "title": "Roteiro do vídeo",
            "source_url": "Link de origem",
            "content_type": "Tipo de conteúdo",
            "summary": "Resumo geral do vídeo",
            "mechanism": "Mecanismo",
            "rows": "Linhas do roteiro",
            "row": "Linha",
            "dialogue": "Diálogo/narração",
        }
    return {
        "title": "Video Script",
        "source_url": "Source URL",
        "content_type": "Content Type",
        "summary": "Whole Video Summary",
        "mechanism": "Mechanism",
        "rows": "Script Rows",
        "row": "Row",
        "dialogue": "Dialogue / audio",
    }


def translate_script_to_portuguese(script_json: dict[str, Any], key: str, models: list[str]) -> dict[str, Any]:
    if not key or run_text_json_prompt_with_fallback is None:
        raise RuntimeError("Portuguese translation helpers are unavailable.")
    payload = json.loads(json.dumps(script_json or {}, ensure_ascii=False))
    translated, _, _ = run_text_json_prompt_with_fallback(
        payload,
        key,
        models,
        PORTUGUESE_TRANSLATION_PROMPT,
        "portuguese translation",
    )
    merged = json.loads(json.dumps(script_json or {}, ensure_ascii=False))
    for field in ["title", "whole_video_summary", "core_viral_points", "replaceable_parts", "rows", "mechanism"]:
        if field in translated:
            merged[field] = translated[field]
    merged["display_language"] = "pt"
    return merged


def write_script_docx(output_dir: Path, script: dict[str, Any], video_url: str, *, suffix: str = "", locale: str = "zh") -> Path | None:
    if Document is None:
        return None
    labels = language_copy(locale)
    path = output_dir / f"script_export{suffix}.docx"
    try:
        doc = Document()
        doc.add_heading(script.get("title") or labels["title"], 0)
        doc.add_paragraph(f"{labels['source_url']}: {video_url}")
        existing_content_type = str(script.get("content_type") or "").strip()
        doc.add_paragraph(f"{labels['content_type']}: {existing_content_type if existing_content_type else keyword_fallback_content_type(script)}")
        summary = script.get("whole_video_summary") or ""
        if summary:
            doc.add_heading(labels["summary"], level=1)
            doc.add_paragraph(summary)
        mechanism = script.get("mechanism") or {}
        if mechanism:
            doc.add_heading(labels["mechanism"], level=1)
            for key in ["name", "reason", "backfire_point", "story_question"]:
                value = mechanism.get(key)
                if value:
                    doc.add_paragraph(f"{key}: {value}")
        rows = choose_script_rows(script)
        if rows:
            doc.add_heading(labels["rows"], level=1)
            for idx, row in enumerate(rows, start=1):
                title = row.get("time") or row.get("start") or f"{labels['row']} {idx}"
                doc.add_heading(f"{idx}. {title}", level=2)
                for key in ["visual_content", "action", "dialogue_or_audio", "integrated_summary", "logic_status"]:
                    value = row.get(key)
                    if value:
                        doc.add_paragraph(f"{key}: {value}")
        doc.save(path)
        return path if path.exists() else None
    except Exception:
        return None


def append_library_entry(entry: dict[str, Any]) -> bool:
    with job_lock:
        entries = load_library_entries()
        entries = [existing for existing in entries if existing.get("entry_id") != entry.get("entry_id")]
        entries.insert(0, entry)
        return save_library_entries(entries[:500])


def append_creator_online_entry(entry: dict[str, Any]) -> bool:
    with job_lock:
        entries = read_json_file(CREATOR_ONLINE_LIBRARY_FILE, default=[])
        if not isinstance(entries, list):
            entries = []
        entries = [existing for existing in entries if isinstance(existing, dict) and existing.get("entry_id") != entry.get("entry_id")]
        entries.insert(0, entry)
        write_json_atomic(CREATOR_ONLINE_LIBRARY_FILE, entries[:500])
        return True


def save_creator_direct_import(payload: dict[str, Any]) -> dict[str, Any]:
    entry = payload.get("entry") if isinstance(payload.get("entry"), dict) else {}
    entry_id = str(entry.get("entry_id") or payload.get("entry_id") or "").strip()
    if not re.fullmatch(r"[0-9a-f]{32}", entry_id):
        raise ValueError("Invalid script id.")
    script_json = payload.get("script_json") if isinstance(payload.get("script_json"), dict) else {}
    html_content = str(payload.get("html_content") or "").strip()
    if not html_content:
        raise ValueError("Missing script HTML content.")
    output_dir = RESULTS_ROOT / entry_id
    output_dir.mkdir(parents=True, exist_ok=True)
    html_path = output_dir / "script_table_pt.html"
    html_path.write_text(html_content, encoding="utf-8")
    if script_json:
        write_json_atomic(output_dir / "script_table_pt.json", script_json)
    cover_url = ""
    cover_b64 = str(payload.get("cover_b64") or "").strip()
    cover_mime = str(payload.get("cover_mime") or "image/png").strip()
    if cover_b64:
        if "," in cover_b64 and cover_b64.startswith("data:"):
            cover_b64 = cover_b64.split(",", 1)[1]
        cover_bytes = base64.b64decode(cover_b64)
        cover_suffix = guess_extension_from_mime(cover_mime)
        cover_name = STORYBOARD_COVER_BASENAME + cover_suffix
        (output_dir / cover_name).write_bytes(cover_bytes)
        preview_name = STORYBOARD_PREVIEW_BASENAME + cover_suffix
        (output_dir / preview_name).write_bytes(cover_bytes)
        save_storyboard_state(entry_id, preview_name=preview_name, cover_name=cover_name, model=str(payload.get("cover_model") or "imported"))
        cover_url = f"/results/{entry_id}/{cover_name}"
    raw_content_type = str(entry.get("content_type") or DEFAULT_CONTENT_TYPE).strip()
    raw_content_source = str(entry.get("content_type_source") or "").strip().lower()
    if raw_content_source == "manual" and raw_content_type in ALLOWED_CONTENT_TYPES and raw_content_type != DEFAULT_CONTENT_TYPE:
        content_type_decision = {
            "content_type": raw_content_type,
            "content_type_source": "manual",
            "content_type_reasoning": str(entry.get("content_type_reasoning") or "Manual import selection."),
            "content_type_confidence": str(entry.get("content_type_confidence") or "manual"),
        }
    else:
        classify_script = {
            **script_json,
            "title": entry.get("title") or script_json.get("title") or "Roteiro importado",
            "whole_video_summary": entry.get("whole_video_summary") or script_json.get("whole_video_summary") or "",
            "summary": entry.get("summary") or script_json.get("summary") or "",
        }
        content_type_decision = detect_content_type_decision(
            classify_script,
            None,
            existing_type="" if raw_content_type in {"", DEFAULT_CONTENT_TYPE} else raw_content_type,
            existing_source="" if raw_content_type in {"", DEFAULT_CONTENT_TYPE} else raw_content_source,
            use_llm=True,
            use_keyword_fallback=False,
        )
    duration_seconds = script_duration_seconds(script_json)
    duration_bucket = creator_duration_bucket(duration_seconds)
    imported_entry = {
        "entry_id": entry_id,
        "parent_job_id": str(entry.get("parent_job_id") or f"creator_import_{entry_id}"),
        "created_at": str(entry.get("created_at") or now_iso()),
        "saved_at": str(entry.get("saved_at") or now_iso()),
        "video_url": str(entry.get("video_url") or payload.get("video_url") or ""),
        "title": str(entry.get("title") or script_json.get("title") or "Roteiro importado"),
        "content_type": str(content_type_decision.get("content_type") or DEFAULT_CONTENT_TYPE),
        "content_type_source": str(content_type_decision.get("content_type_source") or "auto"),
        "content_type_reasoning": str(content_type_decision.get("content_type_reasoning") or "Imported from Creator admin Excel."),
        "content_type_confidence": str(content_type_decision.get("content_type_confidence") or "medium"),
        "duration_seconds": round(duration_seconds, 2) if duration_seconds > 0 else 0,
        "duration_bucket": duration_bucket,
        "duration_label_pt": CREATOR_DURATION_LABELS.get(duration_bucket, {}).get("pt", ""),
        "duration_label_zh": CREATOR_DURATION_LABELS.get(duration_bucket, {}).get("zh", ""),
        "whole_video_summary": str(entry.get("whole_video_summary") or script_json.get("whole_video_summary") or ""),
        "html_url": f"/results/{entry_id}/script_table_pt.html",
        "pt_html_url": f"/results/{entry_id}/script_table_pt.html",
        "zh_html_url": f"/results/{entry_id}/script_table_pt.html",
        "preview_image_url": cover_url,
        "source": "creator_direct_import",
        "published": bool(entry.get("published", True)),
    }
    imported_entry.update(infer_location_tag_fields(imported_entry, script_json))
    append_creator_online_entry(imported_entry)
    return {"ok": True, "entry": imported_entry, "share_url": f"/script/{entry_id}"}


def load_creator_import_jobs() -> dict[str, Any]:
    data = read_json_file(CREATOR_IMPORT_JOBS_FILE, default={})
    return data if isinstance(data, dict) else {}


def save_creator_import_jobs(data: dict[str, Any]) -> None:
    write_json_atomic(CREATOR_IMPORT_JOBS_FILE, data)


def update_creator_import_job(import_id: str, **changes: Any) -> dict[str, Any]:
    with job_lock:
        data = load_creator_import_jobs()
        job = data.get(import_id) if isinstance(data.get(import_id), dict) else {}
        job.update(changes)
        job["updated_at"] = now_iso()
        data[import_id] = job
        save_creator_import_jobs(data)
        return json.loads(json.dumps(job, ensure_ascii=False))


def public_creator_import_job(import_id: str) -> dict[str, Any] | None:
    job = load_creator_import_jobs().get(import_id)
    return job if isinstance(job, dict) else None


def imported_creator_entry(
    item_id: str,
    script_json: dict[str, Any],
    video_url: str,
    variant: dict[str, Any],
    *,
    content_type: str = DEFAULT_CONTENT_TYPE,
    content_type_source: str = "manual",
    content_type_reasoning: str = "Imported from standard Portuguese Excel script table.",
    content_type_confidence: str = "high",
) -> dict[str, Any]:
    content_type = content_type if content_type in ALLOWED_CONTENT_TYPES else DEFAULT_CONTENT_TYPE
    duration_seconds = script_duration_seconds(script_json)
    duration_bucket = creator_duration_bucket(duration_seconds)
    entry = {
        "entry_id": item_id,
        "parent_job_id": f"creator_import_{item_id}",
        "created_at": now_iso(),
        "video_url": video_url,
        "title": script_json.get("title") or "Roteiro importado",
        "content_type": content_type,
        "content_type_source": str(content_type_source or "manual"),
        "content_type_reasoning": str(content_type_reasoning or "Imported from standard Portuguese Excel script table."),
        "content_type_confidence": str(content_type_confidence or "high"),
        "duration_seconds": round(duration_seconds, 2) if duration_seconds > 0 else 0,
        "duration_bucket": duration_bucket,
        "duration_label_pt": CREATOR_DURATION_LABELS.get(duration_bucket, {}).get("pt", ""),
        "duration_label_zh": CREATOR_DURATION_LABELS.get(duration_bucket, {}).get("zh", ""),
        "whole_video_summary": script_json.get("whole_video_summary") or "",
        "html_url": variant.get("html_url") or "",
        "pt_html_url": variant.get("html_url") or "",
        "pt_docx_url": variant.get("docx_url") or "",
        "preview_image_url": library_preview_image_url(item_id, script_json, RESULTS_ROOT / item_id),
        "source": "creator_excel_import",
        "saved_at": now_iso(),
    }
    entry.update(infer_location_tag_fields(entry, script_json))
    return entry


def process_creator_import_script(item: dict[str, Any], *, content_type: str = DEFAULT_CONTENT_TYPE) -> dict[str, Any]:
    item_id = str(item.get("id") or uuid4().hex)
    video_url = str(item.get("video_url") or "").strip()
    script_json = json.loads(json.dumps(item.get("script") or {}, ensure_ascii=False))
    output_dir = RESULTS_ROOT / item_id
    output_dir.mkdir(parents=True, exist_ok=True)
    script_json["display_language"] = "pt"
    variant = generate_script_variant_outputs(output_dir, item_id, script_json, video_url, locale="pt")
    assets = generate_storyboard_assets(item_id, output_dir, script_json)
    script_json["storyboard_cover_url"] = assets.get("cover_url") or ""
    variant = generate_script_variant_outputs(output_dir, item_id, script_json, video_url, locale="pt")
    final_script_json = variant.get("script_json") or script_json
    if content_type == DEFAULT_CONTENT_TYPE:
        content_type_decision = detect_content_type_decision(
            final_script_json,
            None,
            existing_type="",
            existing_source="",
            use_llm=True,
            use_keyword_fallback=False,
        )
    else:
        content_type_decision = {
            "content_type": content_type if content_type in ALLOWED_CONTENT_TYPES else DEFAULT_CONTENT_TYPE,
            "content_type_source": "manual",
            "content_type_reasoning": "Manual import selection",
            "content_type_confidence": "manual",
        }
    entry = imported_creator_entry(
        item_id,
        final_script_json,
        video_url,
        variant,
        content_type=str(content_type_decision.get("content_type") or DEFAULT_CONTENT_TYPE),
        content_type_source=str(content_type_decision.get("content_type_source") or "manual"),
        content_type_reasoning=str(content_type_decision.get("content_type_reasoning") or ""),
        content_type_confidence=str(content_type_decision.get("content_type_confidence") or ""),
    )
    append_library_entry(entry)
    center_import = push_creator_import_to_center(entry, final_script_json, output_dir)
    return {
        "ok": True,
        "id": item_id,
        "video_url": video_url,
        "title": final_script_json.get("title") or script_json.get("title") or "",
        "sheet": item.get("sheet") or "",
        "row": item.get("row") or "",
        "status": "imported" if center_import.get("ok") else "local_imported_remote_failed",
        "share_url": f"{CREATOR_CENTER_BASE_URL}/script/{item_id}",
        "html_url": entry.get("pt_html_url") or entry.get("html_url") or "",
        "preview_image_url": entry.get("preview_image_url") or "",
        "content_type": entry.get("content_type") or DEFAULT_CONTENT_TYPE,
        "content_type_reasoning": entry.get("content_type_reasoning") or "",
        "center_import": center_import,
        "storyboard_cover_url": assets.get("cover_url") or "",
        "storyboard_prompt_model": assets.get("prompt_model") or "",
        "storyboard_image_model": assets.get("image_model") or "",
    }


def process_creator_import_job(import_id: str) -> None:
    job = public_creator_import_job(import_id)
    if not job:
        return
    scripts = job.get("scripts") if isinstance(job.get("scripts"), list) else []
    total = len(scripts)
    worker_count = min(max(1, CREATOR_IMPORT_MAX_WORKERS), max(1, total))
    results: list[dict[str, Any]] = []
    imported_count = 0
    failed_count = 0
    update_creator_import_job(
        import_id,
        status="running",
        stage="importing",
        message=f"开始导入脚本，worker={worker_count}",
        started_at=now_iso(),
    )
    for item in scripts:
        item_id = str(item.get("id") or uuid4().hex)
        script_json = json.loads(json.dumps(item.get("script") or {}, ensure_ascii=False))
        results.append({
            "id": item_id,
            "video_url": str(item.get("video_url") or "").strip(),
            "title": script_json.get("title") or "",
            "sheet": item.get("sheet") or "",
            "row": item.get("row") or "",
            "status": "queued",
            "share_url": f"{CREATOR_CENTER_BASE_URL}/script/{item_id}",
        })
    update_creator_import_job(import_id, current_index=0, imported_count=0, failed_count=0, results=results)
    job_content_type = str(job.get("content_type") or DEFAULT_CONTENT_TYPE).strip()
    future_map: dict[Any, int] = {}
    with ThreadPoolExecutor(max_workers=worker_count, thread_name_prefix=f"creator-import-{import_id[:6]}") as executor:
        for index, item in enumerate(scripts):
            future = executor.submit(process_creator_import_script, item, content_type=job_content_type)
            future_map[future] = index
        for completed_count, future in enumerate(as_completed(future_map), start=1):
            index = future_map[future]
            item_result = results[index]
            try:
                payload = future.result()
                if payload.get("ok"):
                    imported_count += 1
                    item_result.update(payload)
                else:
                    failed_count += 1
                    item_result.update(status="failed", error=friendly_error(str(payload.get("error") or "导入失败。")))
            except Exception as exc:
                failed_count += 1
                item_result.update(status="failed", error=friendly_error(str(exc)))
            update_creator_import_job(
                import_id,
                current_index=completed_count,
                imported_count=imported_count,
                failed_count=failed_count,
                results=results,
                message=f"已处理 {completed_count}/{total} 条",
            )
    sync_result = trigger_creator_center_sync()
    final_status = "completed" if imported_count and not failed_count else ("partial" if imported_count else "failed")
    update_creator_import_job(
        import_id,
        status=final_status,
        stage="completed",
        completed_at=now_iso(),
        imported_count=imported_count,
        failed_count=failed_count,
        results=results,
        sync_result=sync_result,
        message=f"导入完成：成功 {imported_count} 条，失败 {failed_count} 条。",
    )


def start_creator_script_imports(
    source_name: str,
    scripts: list[dict[str, Any]],
    *,
    content_type: str = DEFAULT_CONTENT_TYPE,
    start_async: bool = True,
) -> dict[str, Any]:
    if not scripts:
        raise ValueError("没有识别到标准脚本表。请确认包含 Vídeo original / Conteúdo principal / Pontos principais / Tempo / Imagem / Ações / Diálogos。")
    import_id = uuid4().hex
    job = {
        "id": import_id,
        "filename": source_name,
        "status": "queued",
        "stage": "queued",
        "created_at": now_iso(),
        "updated_at": now_iso(),
        "total": len(scripts),
        "current_index": 0,
        "imported_count": 0,
        "failed_count": 0,
        "content_type": content_type if content_type in ALLOWED_CONTENT_TYPES else DEFAULT_CONTENT_TYPE,
        "scripts": scripts,
        "results": [],
        "message": f"已识别 {len(scripts)} 条脚本，等待导入。",
    }
    with job_lock:
        data = load_creator_import_jobs()
        data[import_id] = job
        save_creator_import_jobs(data)
    if start_async:
        threading.Thread(target=process_creator_import_job, args=(import_id,), name=f"creator-import-{import_id[:8]}", daemon=True).start()
    return job


def start_creator_excel_import(filename: str, file_b64: str, *, content_type: str = DEFAULT_CONTENT_TYPE) -> dict[str, Any]:
    raw_b64 = str(file_b64 or "").strip()
    if "," in raw_b64 and raw_b64.startswith("data:"):
        raw_b64 = raw_b64.split(",", 1)[1]
    if not raw_b64:
        raise ValueError("请上传一个 Excel 文件。")
    blob = base64.b64decode(raw_b64)
    if not str(filename or "").lower().endswith(".xlsx"):
        raise ValueError("目前只支持 .xlsx 格式。")
    scripts = parse_creator_script_tables_from_xlsx(blob)
    return start_creator_script_imports(filename, scripts, content_type=content_type, start_async=True)


def load_storyboard_state(item_id: str) -> dict[str, Any]:
    if not item_id:
        return {}
    output_dir = RESULTS_ROOT / item_id
    state = read_json(output_dir / STORYBOARD_METADATA_FILE) or {}
    preview_name = str(state.get("preview_name") or "").strip()
    cover_name = str(state.get("cover_name") or "").strip()
    prompt_text = str(state.get("prompt") or "").strip()
    preview_url = f"/results/{item_id}/{preview_name}" if preview_name and (output_dir / preview_name).exists() else ""
    cover_url = f"/results/{item_id}/{cover_name}" if cover_name and (output_dir / cover_name).exists() else ""
    return {
        "storyboard_prompt": prompt_text,
        "storyboard_preview_url": preview_url,
        "storyboard_cover_url": cover_url,
        "storyboard_updated_at": state.get("updated_at") or "",
    }


def save_storyboard_state(item_id: str, **changes: Any) -> dict[str, Any]:
    output_dir = RESULTS_ROOT / item_id
    output_dir.mkdir(parents=True, exist_ok=True)
    path = output_dir / STORYBOARD_METADATA_FILE
    state = read_json(path) or {}
    state.update(changes)
    state["updated_at"] = now_iso()
    write_json_atomic(path, state)
    return load_storyboard_state(item_id)


def library_preview_image_url(entry_id: str, script_json: dict[str, Any] | None = None, output_dir: Path | None = None) -> str:
    if not entry_id:
        return ""
    output_dir = output_dir or (RESULTS_ROOT / entry_id)
    storyboard_state = load_storyboard_state(entry_id)
    if storyboard_state.get("storyboard_cover_url"):
        return str(storyboard_state.get("storyboard_cover_url"))
    script_json = script_json or read_json(output_dir / "script_table.json") or read_json(output_dir / "analysis_result.json") or {}
    for row in script_json.get("rows") or []:
        for key in ("start_frame", "end_frame"):
            value = str(row.get(key) or "").strip()
            if not value:
                continue
            candidate = Path(value)
            if not candidate.is_absolute():
                candidate = output_dir / candidate
            try:
                candidate = candidate.resolve()
                base = output_dir.resolve()
            except FileNotFoundError:
                continue
            if candidate.exists() and candidate.is_file() and (candidate == base or base in candidate.parents):
                rel = candidate.relative_to(base).as_posix()
                return f"/results/{entry_id}/{rel}"
    return ""


def library_storyboard_cover_url(entry: dict[str, Any]) -> str:
    entry_id = str(entry.get("entry_id") or "").strip()
    explicit = str(entry.get("storyboard_cover_url") or entry.get("storyboard_image_url") or "").strip()
    if explicit:
        return explicit
    preview_url = str(entry.get("preview_image_url") or "").strip()
    if "storyboard_cover" in preview_url:
        return preview_url
    if not entry_id:
        return ""
    state_url = str(load_storyboard_state(entry_id).get("storyboard_cover_url") or "").strip()
    if state_url:
        return state_url
    output_dir = RESULTS_ROOT / entry_id
    for name in (
        "storyboard_cover.png",
        "storyboard_cover.jpg",
        "storyboard_cover.jpeg",
        "storyboard_cover.webp",
    ):
        if (output_dir / name).exists():
            return f"/results/{entry_id}/{name}"
    return ""


def hydrate_library_entry_preview(entry: dict[str, Any]) -> dict[str, Any]:
    if entry.get("preview_image_url"):
        return entry
    entry_id = str(entry.get("entry_id") or "").strip()
    if not entry_id:
        return entry
    preview_url = library_preview_image_url(entry_id)
    if preview_url:
        entry["preview_image_url"] = preview_url
    return entry


def delete_library_entries(entry_ids: list[str]) -> dict[str, Any]:
    valid_ids = []
    seen = set()
    for entry_id in entry_ids:
        entry_id = str(entry_id or "").strip()
        if not re.fullmatch(r"[0-9a-f]{32}", entry_id) or entry_id in seen:
            continue
        valid_ids.append(entry_id)
        seen.add(entry_id)
    deleted = []
    missing = []
    for entry_id in valid_ids:
        if delete_library_entry(entry_id):
            deleted.append(entry_id)
        else:
            missing.append(entry_id)
    return {"deleted": deleted, "missing": missing}


def delete_library_entry(entry_id: str) -> bool:
    removed = False
    with job_lock:
        entries = load_library_entries()
        next_entries = [entry for entry in entries if entry.get("entry_id") != entry_id]
        if len(next_entries) != len(entries):
            save_library_entries(next_entries)
            removed = True

        job_ids_to_delete: list[str] = []
        for job_id, job in list(jobs.items()):
            items = job.get("items") or []
            if items:
                filtered_items = [item for item in items if item.get("id") != entry_id]
                if len(filtered_items) != len(items):
                    removed = True
                    if filtered_items:
                        job["items"] = filtered_items
                        current_index = int(job.get("current_item_index") or 0)
                        job["current_item_index"] = max(0, min(current_index, len(filtered_items) - 1))
                        job["updated_at"] = now_iso()
                    else:
                        job_ids_to_delete.append(job_id)
            elif job_id == entry_id:
                removed = True
                job_ids_to_delete.append(job_id)

        for job_id in job_ids_to_delete:
            jobs.pop(job_id, None)

        if removed:
            save_jobs()

    output_dir = RESULTS_ROOT / entry_id
    if output_dir.exists():
        shutil.rmtree(output_dir, ignore_errors=True)
        removed = True
    return removed


def update_library_entry_content_type(entry_id: str, content_type: str) -> dict[str, Any] | None:
    if content_type not in ALLOWED_CONTENT_TYPES and content_type != DEFAULT_CONTENT_TYPE:
        raise RuntimeError("Unsupported content type.")
    updated_entry: dict[str, Any] | None = None
    with job_lock:
        entries = load_library_entries()
        found = False
        for entry in entries:
            if entry.get("entry_id") != entry_id:
                continue
            entry["content_type"] = content_type
            entry["content_type_source"] = "manual"
            entry["content_type_reasoning"] = "Manual override"
            entry["content_type_confidence"] = "manual"
            updated_entry = dict(entry)
            found = True
            break
        if not found:
            return None
        save_library_entries(entries)

        for job_id, job in jobs.items():
            if job.get("id") == entry_id:
                job["content_type"] = content_type
                job["content_type_source"] = "manual"
                job["content_type_reasoning"] = "Manual override"
                job["content_type_confidence"] = "manual"
                job["updated_at"] = now_iso()
            for item in job.get("items") or []:
                if item.get("id") != entry_id:
                    continue
                item["content_type"] = content_type
                item["content_type_source"] = "manual"
                item["content_type_reasoning"] = "Manual override"
                item["content_type_confidence"] = "manual"
                item["updated_at"] = now_iso()
                job["updated_at"] = now_iso()
        save_jobs()
    return updated_entry


def apply_manual_item_content_type(parent_job_id: str, item_index: int, content_type: object) -> None:
    raw_selected = str(content_type or "").strip()
    if not raw_selected:
        return
    selected = raw_selected if raw_selected in ALLOWED_CONTENT_TYPES else normalize_creator_content_type(raw_selected)
    if selected not in ALLOWED_CONTENT_TYPES:
        return
    with job_lock:
        if parent_job_id not in jobs:
            return
        job = jobs[parent_job_id]
        items = job.get("items") or []
        item = items[item_index] if 0 <= item_index < len(items) else None
        if item is not None:
            item["content_type"] = selected
            item["content_type_source"] = "manual"
            item["content_type_reasoning"] = "Manual selection before library save"
            item["content_type_confidence"] = "manual"
            item["updated_at"] = now_iso()
        if item is None or job.get("id") == item.get("id") or len(items) == 1:
            job["content_type"] = selected
            job["content_type_source"] = "manual"
            job["content_type_reasoning"] = "Manual selection before library save"
            job["content_type_confidence"] = "manual"
            job["updated_at"] = now_iso()
        save_jobs()


def generate_script_variant_outputs(output_dir: Path, item_id: str, script_json: dict[str, Any], video_url: str, *, locale: str) -> dict[str, Any]:
    if locale == "pt":
        json_name = "script_table_pt.json"
        html_name = "script_table_pt.html"
        suffix = "_pt"
    else:
        json_name = "script_table.json"
        html_name = "script_table.html"
        suffix = ""
    script_json = json.loads(json.dumps(script_json or {}, ensure_ascii=False))
    script_json["display_language"] = locale
    json_path = output_dir / json_name
    write_json_atomic(json_path, script_json)
    render_script = V2_SKILL_ROOT / "scripts" / "render_script_table.py"
    subprocess.run(
        [
            os.environ.get("PYTHON_BIN", "python3"),
            str(render_script),
            str(json_path),
            "--output",
            str(output_dir / html_name),
            "--locale",
            locale,
        ],
        check=True,
        capture_output=True,
        text=True,
    )
    docx_path = write_script_docx(output_dir, script_json, video_url, suffix=suffix, locale=locale)
    return {
        "script_json": script_json,
        "json_url": f"/results/{item_id}/{json_name}",
        "html_url": f"/results/{item_id}/{html_name}",
        "docx_url": f"/results/{item_id}/{docx_path.name}" if docx_path and docx_path.exists() else "",
    }


def should_try_next_model(error_text: str) -> bool:
    hay = (error_text or "").upper()
    return any(
        token in hay
        for token in [
            "HTTP 400",
            "HTTP 500",
            "FAILED_PRECONDITION",
            "HTTP 503",
            "INTERNAL",
            "UNAVAILABLE",
            "RESOURCE_EXHAUSTED",
            "TIMED OUT",
            "REMOTE END CLOSED CONNECTION WITHOUT RESPONSE",
            "EOF OCCURRED",
            "CONNECTION RESET",
            "BROKEN PIPE",
            "NO PARSEABLE JSON OBJECT FOUND",
            "JSON",
        ]
    )


def should_retry_transient_pipeline(error_text: str) -> bool:
    hay = (error_text or "").upper()
    return any(
        token in hay
        for token in [
            "HTTP 429",
            "HTTP 500",
            "HTTP 503",
            "INTERNAL",
            "UNAVAILABLE",
            "HIGH DEMAND",
            "RATE_LIMIT",
            "RATE LIMIT",
            "QUOTA",
            "OVERLOADED",
            "SERVER ERROR",
            "REMOTE END CLOSED CONNECTION WITHOUT RESPONSE",
            "CONNECTION RESET",
            "BROKEN PIPE",
        ]
    )


def gemini_retry_delay(attempt_index: int) -> int:
    if attempt_index <= 0:
        return 0
    return GEMINI_TRANSIENT_RETRY_DELAYS[min(attempt_index - 1, len(GEMINI_TRANSIENT_RETRY_DELAYS) - 1)]


def friendly_error(error_text: str) -> str:
    text = (error_text or "").strip()
    if not text:
        return "分析失败，未返回具体错误。"
    if "no space left on device" in text.lower() or "enospc" in text.lower():
        return "服务器存储空间不足，暂时无法创建新任务。请先清理历史结果或扩容后重试。"
    upper = text.upper()
    if (
        "HTTP 429" in text
        or "RESOURCE_EXHAUSTED" in upper
        or "RATE_LIMIT" in upper
        or "RATE LIMIT" in upper
        or "QUOTA" in upper
    ):
        return "Gemini 当前更像是触发了配额或限流，不一定是服务本身过载。请优先检查 API key 的额度、限速或计费状态。"
    if (
        "HTTP 500" in text
        or "HTTP 503" in text
        or "UNAVAILABLE" in text
        or "INTERNAL" in upper
        or "HIGH DEMAND" in upper
    ):
        return "Gemini 当前负载较高，系统已自动换模型并延迟重试，但这次仍未成功。请稍后重试。"
    if "FAILED_PRECONDITION" in text and "User location is not supported" in text:
        return "当前运行环境所在地区不支持这个 Gemini 接口。"
    if "timed out" in text.lower():
        return "请求超时，可能是目标视频或 Gemini 响应过慢。"
    if "remote end closed connection without response" in text.lower():
        return "Gemini 连接被远端中断了，这次已尝试自动换模型；如果仍失败，通常是接口瞬时不稳定。"
    return text


def extract_json_line(text: str) -> dict[str, Any] | None:
    for line in reversed([line.strip() for line in text.splitlines() if line.strip()]):
        if not line.startswith("{"):
            continue
        try:
            return json.loads(line)
        except json.JSONDecodeError:
            continue
    return None


def collect_observation_rows(observations: list[dict[str, Any]], limit: int = 24) -> list[dict[str, str]]:
    rows: list[dict[str, str]] = []
    for obs in observations[:limit]:
        people = []
        for person in obs.get("people") or []:
            if isinstance(person, dict):
                people.append(
                    " / ".join(
                        str(person.get(key, ""))
                        for key in ["id", "position", "visible_action"]
                        if person.get(key)
                    )
                )
            else:
                people.append(str(person))
        objects = []
        for obj in obs.get("objects") or []:
            if isinstance(obj, dict):
                objects.append(
                    " / ".join(
                        str(obj.get(key, ""))
                        for key in ["label", "position", "state"]
                        if obj.get(key)
                    )
                )
            else:
                objects.append(str(obj))
        rows.append(
            {
                "time": str(obs.get("time", "")),
                "scene": str(obs.get("visual_scene", "")),
                "people": "；".join(x for x in people if x),
                "objects": "；".join(x for x in objects if x),
                "audio": str(obs.get("audio", "")),
                "visible_text": str(obs.get("visible_text", "")),
                "uncertainty": str(obs.get("uncertainty", "")),
            }
        )
    return rows


def raw_text_preview(raw: dict[str, Any], limit: int = 2200) -> str:
    parts = raw.get("generate", {}).get("candidates", [{}])[0].get("content", {}).get("parts", [])
    text = "".join(part.get("text", "") for part in parts if isinstance(part, dict)).strip()
    if not text:
        return ""
    return text[:limit] + ("\n..." if len(text) > limit else "")


def summarize_artifacts(job_id: str, output_dir: Path) -> dict[str, str]:
    names = [
        "source.mp4",
        "source_metadata.json",
        "primary_v2_draft.json",
        "primary_analysis_raw_gemini.json",
        "type_router.json",
        "audio_multiview.json",
        "audio_multiview_raw_gemini.json",
        "supplement_evidence.json",
        "supplement_raw_gemini.json",
        "final_refine_raw_gemini.json",
        "case_memory_entry.json",
        "observations.json",
        "observations_raw_gemini.json",
        "analysis_raw_gemini.json",
        "analysis_result.json",
        "script_table.json",
        "script_table.html",
        "evidence_bundle.json",
        "product_report.html",
    ]
    out: dict[str, str] = {}
    for name in names:
        if (output_dir / name).exists():
            out[name] = f"/results/{job_id}/{name}"
    return out


def backfill_completed_jobs() -> None:
    changed = False
    for job_id, job in jobs.items():
        if job.get("status") != "completed":
            continue
        output_dir = RESULTS_ROOT / job_id
        artifacts = summarize_artifacts(job_id, output_dir)
        if not artifacts:
            continue
        if artifacts.get("product_report.html") and not job.get("report_url"):
            job["report_url"] = artifacts["product_report.html"]
            changed = True
        if artifacts.get("evidence_bundle.json") and not job.get("evidence_url"):
            job["evidence_url"] = artifacts["evidence_bundle.json"]
            changed = True
        if artifacts != (job.get("artifacts") or {}):
            job["artifacts"] = artifacts
            changed = True
    if changed:
        save_jobs()


def sync_library_from_jobs() -> None:
    existing_entries = {entry.get("entry_id"): entry for entry in load_library_entries()}
    for parent_job_id, job in jobs.items():
        items = job.get("items") or []
        if not items and job.get("status") == "completed":
            output_dir = RESULTS_ROOT / parent_job_id
            script_json = read_json(output_dir / "script_table.json") or read_json(output_dir / "analysis_result.json") or job.get("result_json") or {}
            if script_json:
                docx_path = write_script_docx(output_dir, script_json, job.get("video_url") or "")
                docx_url = f"/results/{parent_job_id}/{docx_path.name}" if docx_path and docx_path.exists() else ""
                entry = {
                    **detect_content_type_decision_for_output(
                        output_dir,
                        script_json,
                        read_json(output_dir / "evidence_bundle.json"),
                        existing_type=job.get("content_type") or "",
                        existing_source=job.get("content_type_source") or "",
                        use_llm=False,
                    ),
                    **creator_duration_fields(parent_job_id, script_json),
                    "entry_id": parent_job_id,
                    "parent_job_id": parent_job_id,
                    "created_at": job.get("completed_at") or job.get("updated_at") or now_iso(),
                    "video_url": job.get("video_url") or "",
                    "title": script_json.get("title") or "Video Script",
                    "whole_video_summary": script_json.get("whole_video_summary") or "",
                    "html_url": job.get("html_url") or f"/results/{parent_job_id}/script_table.html",
                    "report_url": job.get("report_url") or f"/results/{parent_job_id}/product_report.html",
                    "evidence_url": job.get("evidence_url") or f"/results/{parent_job_id}/evidence_bundle.json",
                    "docx_url": docx_url,
                    "preview_image_url": library_preview_image_url(parent_job_id, script_json, output_dir),
                }
                append_library_entry(entry)
                existing_entries[parent_job_id] = entry
        for item in items:
            if item.get("status") != "completed":
                continue
            output_dir = RESULTS_ROOT / item["id"]
            script_json = read_json(output_dir / "script_table.json") or read_json(output_dir / "analysis_result.json") or item.get("result_json") or {}
            if not script_json:
                continue
            if not item.get("docx_url"):
                docx_path = write_script_docx(output_dir, script_json, item.get("video_url") or "")
                if docx_path and docx_path.exists():
                    item["docx_url"] = f"/results/{item['id']}/{docx_path.name}"
            decision = detect_content_type_decision_for_output(
                output_dir,
                script_json,
                read_json(output_dir / "evidence_bundle.json"),
                existing_type=item.get("content_type") or "",
                existing_source=item.get("content_type_source") or "",
                use_llm=False,
            )
            item["content_type"] = decision["content_type"]
            item["content_type_source"] = decision["content_type_source"]
            item["content_type_reasoning"] = decision["content_type_reasoning"]
            item["content_type_confidence"] = decision["content_type_confidence"]
            item["title"] = item.get("title") or script_json.get("title") or "Video Script"
            if persist_library_entry(parent_job_id, item):
                item["saved_to_library_at"] = item.get("saved_to_library_at") or now_iso()
            existing_entries[item.get("id")] = item
    save_jobs()


def build_evidence_bundle(job_id: str, output_dir: Path, result_json: dict[str, Any]) -> dict[str, Any]:
    observations = read_json(output_dir / "observations.json")
    raw = read_json(output_dir / "observations_raw_gemini.json")
    analysis = read_json(output_dir / "analysis_result.json")
    script = analysis or read_json(output_dir / "script_table.json")
    type_router = read_json(output_dir / "type_router.json") or script.get("type_router") or {}
    case_memory_entry = read_json(output_dir / "case_memory_entry.json")
    similar_cases = script.get("similar_cases_used") or []

    usage = raw.get("generate", {}).get("usageMetadata", {})
    upload = raw.get("upload", {}).get("file", {})
    source_meta = script.get("source_metadata") or observations.get("source_metadata") or {}
    observation_rows = collect_observation_rows(observations.get("observations") or [], limit=30)
    segments = (script.get("synthesized_segments") or [])[:12]
    windows = (script.get("story_analysis") or {}).get("verification_windows") or []
    hypotheses = (script.get("story_analysis") or {}).get("mechanism_hypotheses") or []
    raw_preview = raw_text_preview(raw)

    bundle = {
        "job_id": job_id,
        "generated_at": now_iso(),
        "source": {
            "video_url": source_meta.get("source_url") or source_meta.get("input") or "",
            "title": source_meta.get("title") or "",
            "route": source_meta.get("route") or observations.get("analysis_route") or "",
            "local_video": source_meta.get("local_video") or "",
            "downloaded_bytes": source_meta.get("downloaded_bytes") or source_meta.get("size") or 0,
            "mime_type": source_meta.get("mime_type") or "",
        },
        "extraction": {
            "analysis_route": observations.get("analysis_route") or "",
            "gemini_model": observations.get("gemini_model") or result_json.get("model") or "",
            "analysis_model": script.get("analysis_model") or result_json.get("analysis_model") or "",
            "api_key_source": observations.get("api_key_source") or "",
            "duration_estimate": observations.get("duration_estimate") or "",
            "observation_interval_sec": observations.get("observation_interval_sec") or "",
            "observation_count": len(observations.get("observations") or []),
            "upload_file": {
                "name": upload.get("name") or "",
                "mime_type": upload.get("mimeType") or "",
                "size_bytes": upload.get("sizeBytes") or "",
                "state": upload.get("state") or "",
            },
            "usage": {
                "prompt_tokens": usage.get("promptTokenCount"),
                "candidate_tokens": usage.get("candidatesTokenCount"),
                "total_tokens": usage.get("totalTokenCount"),
                "thoughts_tokens": usage.get("thoughtsTokenCount"),
            },
        },
        "evidence": {
            "timeline_preview": observation_rows,
            "raw_gemini_preview": raw_preview,
            "verification_windows": windows[:14],
            "mechanism_hypotheses": hypotheses[:8],
        },
        "reasoning": {
            "whole_video_summary": script.get("whole_video_summary") or "",
            "safe_final_story": (script.get("story_analysis") or {}).get("safe_final_story") or "",
            "logic_quality": script.get("logic_quality") or "",
            "analysis_route": script.get("analysis_route") or "",
            "segments_preview": segments,
        },
        "routing": {
            "routing_mode": type_router.get("routing_mode") or "",
            "primary_type": type_router.get("primary_type") or "",
            "subtype_guess": type_router.get("subtype_guess") or "",
            "confidence": type_router.get("confidence") or "",
            "reasoning_summary": type_router.get("reasoning_summary") or "",
            "matched_templates": type_router.get("matched_templates") or [],
            "review_questions": type_router.get("review_questions") or [],
        },
        "case_memory": {
            "similar_cases_used": similar_cases[:3],
            "current_entry": case_memory_entry,
        },
        "artifacts": summarize_artifacts(job_id, output_dir),
    }
    return bundle


def render_template_cards(items: list[dict[str, Any]]) -> str:
    if not items:
        return '<div class="empty">这次没有命中明确模板，系统退回通用故事框架。</div>'
    cards = []
    for item in items:
        cards.append(
            "<article class='chip-card'>"
            f"<h4>{html_escape(item.get('title') or item.get('id') or '模板')}</h4>"
            f"<p>{html_escape(item.get('description') or '')}</p>"
            "</article>"
        )
    return "".join(cards)


def render_similar_cases(items: list[dict[str, Any]]) -> str:
    if not items:
        return '<div class="empty">这次没有命中可复用的历史案例，系统按通用框架独立分析。</div>'
    cards = []
    for item in items:
        cards.append(
            "<article class='segment-card'>"
            f"<div class='segment-range'>相似度 {html_escape(item.get('score') or '')}</div>"
            f"<h4>{html_escape(item.get('subtype_guess') or item.get('primary_type') or '历史案例')}</h4>"
            f"<p>{html_escape(item.get('whole_video_summary') or item.get('safe_final_story') or '')}</p>"
            f"<div class='meta-line'>route: {html_escape(item.get('route') or '')}</div>"
            "</article>"
        )
    return "".join(cards)


def render_hypothesis_cards(items: list[dict[str, Any]]) -> str:
    if not items:
        return '<div class="empty">这次没有产出机制假设。</div>'
    cards = []
    for item in items:
        cards.append(
            "<article class='chip-card'>"
            f"<h4>{html_escape(item.get('name') or '候选机制')}</h4>"
            f"<p>{html_escape(item.get('story_question') or '')}</p>"
            f"<div class='meta-line'>可能性：{html_escape(item.get('likelihood') or '')}</div>"
            "</article>"
        )
    return "".join(cards)


def render_windows(items: list[dict[str, Any]]) -> str:
    if not items:
        return '<div class="empty">没有建议复核窗口。</div>'
    rows = []
    for item in items:
        rows.append(
            "<tr>"
            f"<td>{html_escape(item.get('start'))} - {html_escape(item.get('end'))}</td>"
            f"<td>{html_escape(item.get('reason'))}</td>"
            "</tr>"
        )
    return "".join(rows)


def render_timeline(rows: list[dict[str, str]]) -> str:
    if not rows:
        return '<div class="empty">这次没有生成逐秒 observation。</div>'
    out = []
    for row in rows:
        out.append(
            "<tr>"
            f"<td>{html_escape(row['time'])}</td>"
            f"<td>{html_escape(row['scene'])}</td>"
            f"<td>{html_escape(row['people'])}</td>"
            f"<td>{html_escape(row['objects'])}</td>"
            f"<td>{html_escape(row['audio'])}</td>"
            f"<td>{html_escape(row['uncertainty'])}</td>"
            "</tr>"
        )
    return "".join(out)


def render_segments(items: list[dict[str, Any]]) -> str:
    if not items:
        return '<div class="empty">还没有合成后的脚本分段。</div>'
    blocks = []
    for item in items:
        blocks.append(
            "<article class='segment-card'>"
            f"<div class='segment-range'>{html_escape(item.get('start'))} - {html_escape(item.get('end'))}</div>"
            f"<h4>{html_escape(item.get('segment_role') or '段落')}</h4>"
            f"<p>{html_escape(item.get('integrated_summary') or '')}</p>"
            f"<div class='segment-meta'>逻辑状态：{html_escape(item.get('logic_status') or '')}</div>"
            "</article>"
        )
    return "".join(blocks)


def render_artifact_links(artifacts: dict[str, str]) -> str:
    if not artifacts:
        return '<div class="empty">没有找到过程产物文件。</div>'
    links = []
    for name, url in artifacts.items():
        links.append(f"<a class='artifact-link' href='{html_escape(url)}' target='_blank' rel='noreferrer'>{html_escape(name)}</a>")
    return "".join(links)


def render_product_report(bundle: dict[str, Any]) -> str:
    source = bundle.get("source", {})
    extraction = bundle.get("extraction", {})
    evidence = bundle.get("evidence", {})
    reasoning = bundle.get("reasoning", {})
    routing = bundle.get("routing", {})
    case_memory = bundle.get("case_memory", {})
    artifacts = bundle.get("artifacts", {})

    return f"""<!doctype html>
<html lang="zh-CN">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width,initial-scale=1">
  <title>视频证据报告</title>
  {FAVICON_LINKS}
  <style>
    :root {{
      --bg: #f5f4ee;
      --paper: rgba(255,255,255,.96);
      --ink: #1e293b;
      --muted: #526071;
      --line: rgba(30,41,59,.12);
      --brand: #0f766e;
      --warm: #d97706;
      --soft: #e7ecef;
      --shadow: 0 30px 80px rgba(15,23,42,.08);
    }}
    * {{ box-sizing: border-box; }}
    body {{
      margin: 0;
      font-family: "Avenir Next", "Segoe UI", sans-serif;
      background:
        radial-gradient(circle at top right, rgba(15,118,110,.10), transparent 24%),
        linear-gradient(180deg, #fbfaf6 0%, var(--bg) 100%);
      color: var(--ink);
    }}
    .wrap {{ max-width: 1320px; margin: 0 auto; padding: 36px 24px 72px; }}
    .hero {{
      display: grid;
      grid-template-columns: 1.4fr .8fr;
      gap: 18px;
      margin-bottom: 18px;
    }}
    .card {{
      background: var(--paper);
      border: 1px solid rgba(255,255,255,.9);
      border-radius: 28px;
      padding: 24px;
      box-shadow: var(--shadow);
    }}
    h1 {{
      margin: 0 0 10px;
      font-family: "Iowan Old Style", Georgia, serif;
      font-size: clamp(2.1rem, 3.8vw, 3.6rem);
      line-height: 1.02;
      letter-spacing: -.04em;
    }}
    h2 {{ margin: 0 0 12px; font-size: 24px; }}
    h3 {{ margin: 0 0 10px; font-size: 18px; }}
    h4 {{ margin: 0 0 8px; font-size: 16px; }}
    p {{ margin: 0; line-height: 1.8; color: var(--muted); }}
    .eyebrow {{
      display: inline-flex;
      align-items: center;
      gap: 8px;
      color: var(--brand);
      background: rgba(15,118,110,.08);
      border: 1px solid rgba(15,118,110,.16);
      border-radius: 999px;
      padding: 6px 12px;
      font-size: 12px;
      font-weight: 700;
      letter-spacing: .08em;
      text-transform: uppercase;
      margin-bottom: 18px;
    }}
    .lede {{ font-size: 16px; max-width: 56ch; }}
    .grid-3 {{
      display: grid;
      grid-template-columns: repeat(3, minmax(0, 1fr));
      gap: 18px;
      margin-top: 18px;
    }}
    .metric {{
      border-radius: 20px;
      border: 1px solid var(--line);
      background: rgba(255,255,255,.72);
      padding: 16px;
    }}
    .metric b {{ display: block; font-size: 28px; margin-bottom: 6px; color: var(--ink); }}
    .stack {{ display: grid; gap: 18px; margin-top: 18px; }}
    .two-up {{
      display: grid;
      grid-template-columns: 1fr 1fr;
      gap: 18px;
      margin-top: 18px;
    }}
    .chip-grid {{
      display: grid;
      grid-template-columns: repeat(auto-fit, minmax(220px, 1fr));
      gap: 12px;
    }}
    .chip-card, .segment-card {{
      border: 1px solid var(--line);
      border-radius: 18px;
      background: rgba(255,255,255,.8);
      padding: 16px;
    }}
    .meta-line {{
      margin-top: 8px;
      font-size: 13px;
      color: var(--muted);
    }}
    .segment-range {{
      display: inline-flex;
      margin-bottom: 10px;
      border-radius: 999px;
      padding: 5px 10px;
      background: rgba(217,119,6,.10);
      color: var(--warm);
      font-size: 12px;
      font-weight: 700;
      letter-spacing: .06em;
      text-transform: uppercase;
    }}
    .summary-box {{
      padding: 16px 18px;
      border-radius: 20px;
      background: #f9fafb;
      border: 1px solid var(--line);
      color: var(--ink);
      line-height: 1.85;
      white-space: pre-wrap;
    }}
    .artifact-row {{
      display: flex;
      flex-wrap: wrap;
      gap: 10px;
    }}
    .artifact-link {{
      text-decoration: none;
      color: var(--ink);
      background: #fff;
      border: 1px solid var(--line);
      border-radius: 999px;
      padding: 10px 14px;
      font-size: 14px;
    }}
    table {{
      width: 100%;
      border-collapse: collapse;
      table-layout: fixed;
      border-radius: 18px;
      overflow: hidden;
    }}
    th, td {{
      border: 1px solid var(--line);
      padding: 12px 10px;
      text-align: left;
      vertical-align: top;
      font-size: 14px;
      line-height: 1.7;
      word-break: break-word;
    }}
    th {{ background: #f8fafc; font-size: 13px; letter-spacing: .03em; text-transform: uppercase; color: var(--muted); }}
    pre {{
      margin: 0;
      white-space: pre-wrap;
      word-break: break-word;
      border-radius: 20px;
      background: #0f172a;
      color: #dbeafe;
      padding: 18px;
      font-size: 13px;
      line-height: 1.7;
      overflow: auto;
    }}
    .empty {{
      padding: 16px;
      border: 1px dashed var(--line);
      border-radius: 16px;
      color: var(--muted);
    }}
    a.inline-link {{ color: var(--brand); text-decoration: none; }}
    @media (max-width: 1040px) {{
      .hero, .grid-3, .two-up {{ grid-template-columns: 1fr; }}
    }}
  </style>
</head>
<body>
  <div class="wrap">
    <section class="hero">
      <article class="card">
        <div class="eyebrow">Evidence-First Video Analysis</div>
        <h1>这不是“让 Gemini 直接写结论”，而是先把视频拆成证据，再做推理。</h1>
        <p class="lede">当前这份报告按三层结构组织：第一层是 Gemini 的感知提取，第二层是结构化证据包，第三层才是脚本和机制判断。这样我们能看见模型到底看到了什么，又是如何被后续逻辑使用的。</p>
      </article>
      <article class="card">
        <h3>源视频</h3>
        <p><a class="inline-link" href="{html_escape(source.get('video_url'))}" target="_blank" rel="noreferrer">{html_escape(source.get('video_url'))}</a></p>
        <div class="stack">
          <div class="metric"><b>{html_escape(extraction.get('gemini_model') or 'unknown')}</b><span>Gemini 提取模型</span></div>
          <div class="metric"><b>{html_escape(extraction.get('analysis_model') or 'unknown')}</b><span>LLM 分析模型</span></div>
          <div class="metric"><b>{html_escape(extraction.get('observation_count') or 0)}</b><span>逐秒 observation 条数</span></div>
          <div class="metric"><b>{html_escape(extraction.get('duration_estimate') or 'n/a')}</b><span>Gemini 估计时长</span></div>
        </div>
      </article>
    </section>

    <section class="grid-3">
      <article class="card">
        <h3>感知层</h3>
        <p>Gemini 在这里主要负责把视频切成逐秒 observation，而不是直接下剧情结论。</p>
      </article>
      <article class="card">
        <h3>证据层</h3>
        <p>我们保留逐秒场景、人物动作、道具、音频、不确定点，以及建议复核窗口。</p>
      </article>
      <article class="card">
        <h3>推理层</h3>
        <p>最后才把这些证据压成机制假设、脚本分段和可分享的最终 HTML。</p>
      </article>
    </section>

    <section class="two-up">
      <article class="card">
        <h2>模板路由系统</h2>
        <div class="stack">
          <div class="metric"><b>{html_escape(routing.get('primary_type') or '通用故事')}</b><span>主类型</span></div>
          <div class="metric"><b>{html_escape(routing.get('subtype_guess') or '未细分')}</b><span>子类型判断</span></div>
          <div class="metric"><b>{html_escape(routing.get('routing_mode') or 'universal')}</b><span>路由模式</span></div>
          <div class="metric"><b>{html_escape(routing.get('confidence') or 'low')}</b><span>模板置信度</span></div>
        </div>
        <div class="summary-box" style="margin-top:14px;">{html_escape(routing.get('reasoning_summary') or '这次没有稳定命中模板，已退回通用故事框架。')}</div>
      </article>
      <article class="card">
        <h2>命中的模板</h2>
        <div class="chip-grid">{render_template_cards(routing.get('matched_templates') or [])}</div>
      </article>
    </section>

    <section class="two-up">
      <article class="card">
        <h2>第一层：Gemini 感知提取</h2>
        <div class="stack">
          <div class="metric"><b>{html_escape(extraction.get('analysis_route') or 'n/a')}</b><span>提取通道</span></div>
          <div class="metric"><b>{html_escape(extraction.get('upload_file', {}).get('state') or 'n/a')}</b><span>Gemini 文件状态</span></div>
          <div class="metric"><b>{html_escape(extraction.get('usage', {}).get('total_tokens') or 'n/a')}</b><span>本次总 token</span></div>
        </div>
      </article>
      <article class="card">
        <h2>原始产物</h2>
        <div class="artifact-row">{render_artifact_links(artifacts)}</div>
      </article>
    </section>

    <section class="card" style="margin-top:18px;">
      <h2>Gemini 原始返回摘录</h2>
      <p style="margin-bottom:14px;">这里展示的不是后处理结果，而是 Gemini 返回文本里的原始 JSON 片段。</p>
      <pre>{html_escape(evidence.get('raw_gemini_preview') or '这次没有可展示的 Gemini 原始文本。')}</pre>
    </section>

    <section class="card" style="margin-top:18px;">
      <h2>第二层：结构化证据包</h2>
      <p style="margin-bottom:14px;">这部分已经把 Gemini 的文本观察整理成可回溯的逐秒证据表。当前没有关键帧图片，是因为本机还没启用抽帧依赖，但时间、动作、道具、音频和不确定点已经能独立审阅。</p>
      <table>
        <thead>
          <tr>
            <th>时间</th>
            <th>场景</th>
            <th>人物动作</th>
            <th>道具</th>
            <th>音频</th>
            <th>不确定点</th>
          </tr>
        </thead>
        <tbody>{render_timeline(evidence.get('timeline_preview') or [])}</tbody>
      </table>
    </section>

    <section class="two-up">
      <article class="card">
        <h2>建议复核窗口</h2>
        <table>
          <thead><tr><th>时间窗</th><th>原因</th></tr></thead>
          <tbody>{render_windows(evidence.get('verification_windows') or [])}</tbody>
        </table>
      </article>
      <article class="card">
        <h2>候选机制假设</h2>
        <div class="chip-grid">{render_hypothesis_cards(evidence.get('mechanism_hypotheses') or [])}</div>
      </article>
    </section>

    <section class="two-up">
      <article class="card">
        <h2>案例沉淀系统</h2>
        <p style="margin-bottom:14px;">系统会把当前结果压成一个案例条目，并在整理终稿前检索相似案例，避免重复犯同一种误判。</p>
        <div class="chip-grid">{render_similar_cases(case_memory.get('similar_cases_used') or [])}</div>
      </article>
      <article class="card">
        <h2>当前沉淀条目</h2>
        <div class="summary-box">{html_escape((case_memory.get('current_entry') or {}).get('whole_video_summary') or '这次尚未写入案例摘要。')}</div>
      </article>
    </section>

    <section class="card" style="margin-top:18px;">
      <h2>第三层：推理与脚本输出</h2>
      <div class="stack">
        <div>
          <h3>整段总结</h3>
          <div class="summary-box">{html_escape(reasoning.get('whole_video_summary') or '暂无')}</div>
        </div>
        <div>
          <h3>保守版最终故事</h3>
          <div class="summary-box">{html_escape(reasoning.get('safe_final_story') or '暂无')}</div>
        </div>
      </div>
    </section>

    <section class="card" style="margin-top:18px;">
      <h2>合成后的脚本分段预览</h2>
      <div class="chip-grid">{render_segments(reasoning.get('segments_preview') or [])}</div>
      <div class="meta-line" style="margin-top:14px;">逻辑质量：{html_escape(reasoning.get('logic_quality') or 'n/a')} · 最终脚本 HTML：<a class="inline-link" href="{html_escape(artifacts.get('script_table.html', ''))}" target="_blank" rel="noreferrer">打开</a></div>
    </section>
  </div>
</body>
</html>"""


def collect_completed_script_records() -> list[dict[str, Any]]:
    records: list[dict[str, Any]] = []
    for job_id, job in jobs.items():
        items = job.get("items") or []
        if items:
            for item in items:
                if item.get("status") != "completed":
                    continue
                completed_dt = parse_iso_datetime(item.get("completed_at") or item.get("updated_at") or item.get("created_at"))
                if not completed_dt:
                    continue
                records.append(
                    {
                        "entry_id": item.get("id") or "",
                        "job_id": job_id,
                        "video_url": item.get("video_url") or "",
                        "title": item.get("title") or "",
                        "completed_at": completed_dt,
                        "completed_at_bj": completed_dt.astimezone(BEIJING_TZ).strftime("%Y-%m-%d %H:%M:%S"),
                        "reviewed": bool(item.get("reviewed")) or str(item.get("review_status") or "").strip() == "completed",
                        "edited": bool(item.get("edited")),
                    }
                )
            continue
        if job.get("status") != "completed":
            continue
        completed_dt = parse_iso_datetime(job.get("completed_at") or job.get("updated_at") or job.get("created_at"))
        if not completed_dt:
            continue
        records.append(
            {
                "entry_id": job_id,
                "job_id": job_id,
                "video_url": job.get("video_url") or "",
                "title": (job.get("result_json") or {}).get("title") or "",
                "completed_at": completed_dt,
                "completed_at_bj": completed_dt.astimezone(BEIJING_TZ).strftime("%Y-%m-%d %H:%M:%S"),
                "reviewed": bool(job.get("reviewed")) or str(job.get("review_status") or "").strip() == "completed",
                "edited": bool(job.get("edited")),
            }
        )
    records.sort(key=lambda item: item["completed_at"], reverse=True)
    return records


def build_stats_payload() -> dict[str, Any]:
    records = collect_completed_script_records()
    now_dt = datetime.now(timezone.utc)
    summary = {
        "last_24h": 0,
        "last_7d": 0,
        "last_30d": 0,
        "review_count_30d": 0,
        "edited_count_30d": 0,
        "all_time": len(records),
    }
    days_map: dict[str, dict[str, Any]] = {}
    for record in records:
        completed_dt: datetime = record["completed_at"]
        delta = now_dt - completed_dt
        if delta <= timedelta(hours=24):
            summary["last_24h"] += 1
        if delta <= timedelta(days=7):
            summary["last_7d"] += 1
        if delta <= timedelta(days=30):
            summary["last_30d"] += 1
            if record["reviewed"]:
                summary["review_count_30d"] += 1
            if record["edited"]:
                summary["edited_count_30d"] += 1
        day_key = completed_dt.astimezone(BEIJING_TZ).strftime("%Y-%m-%d")
        bucket = days_map.setdefault(
            day_key,
            {
                "date": day_key,
                "script_count": 0,
                "review_count": 0,
                "edited_count": 0,
                "items": [],
            },
        )
        bucket["script_count"] += 1
        if record["reviewed"]:
            bucket["review_count"] += 1
        if record["edited"]:
            bucket["edited_count"] += 1
        bucket["items"].append(
            {
                "video_url": record["video_url"],
                "title": record["title"],
                "completed_at_bj": record["completed_at_bj"],
                "reviewed": record["reviewed"],
                "edited": record["edited"],
            }
        )
    days = [days_map[key] for key in sorted(days_map.keys(), reverse=True)]
    return {"summary": summary, "days": days}


def stats_html() -> str:
    payload = build_stats_payload()
    summary = payload["summary"]
    days = payload["days"]
    summary_cards = "".join(
        [
            f"<article class='stats-summary-card'><span class='stats-summary-label'>最近 24 小时</span><strong>{summary['last_24h']}</strong><small>生成脚本数</small></article>",
            f"<article class='stats-summary-card'><span class='stats-summary-label'>最近 7 天</span><strong>{summary['last_7d']}</strong><small>生成脚本数</small></article>",
            f"<article class='stats-summary-card'><span class='stats-summary-label'>最近 30 天</span><strong>{summary['last_30d']}</strong><small>生成脚本数</small></article>",
            f"<article class='stats-summary-card'><span class='stats-summary-label'>最近 30 天</span><strong>{summary['review_count_30d']}</strong><small>复盘重做次数</small></article>",
            f"<article class='stats-summary-card'><span class='stats-summary-label'>最近 30 天</span><strong>{summary['edited_count_30d']}</strong><small>整稿编辑次数</small></article>",
            f"<article class='stats-summary-card'><span class='stats-summary-label'>累计</span><strong>{summary['all_time']}</strong><small>历史生成脚本数</small></article>",
        ]
    )
    day_cards = []
    for day in days:
        item_rows = []
        for idx, item in enumerate(day.get("items") or [], start=1):
            review_badge = "<span class='stats-badge yes'>已复盘</span>" if item.get("reviewed") else "<span class='stats-badge'>未复盘</span>"
            edit_badge = "<span class='stats-badge yes'>已整稿编辑</span>" if item.get("edited") else "<span class='stats-badge'>未整稿编辑</span>"
            item_rows.append(
                "<article class='stats-item-row'>"
                f"<div class='stats-item-index'>#{idx}</div>"
                "<div class='stats-item-copy'>"
                f"<a class='stats-link' href='{html_escape(item.get('video_url') or '')}' target='_blank' rel='noreferrer'>{html_escape(item.get('video_url') or '')}</a>"
                f"<div class='stats-item-meta'><span>{html_escape(item.get('completed_at_bj') or '')}</span>{review_badge}{edit_badge}</div>"
                "</div>"
                "</article>"
            )
        day_body_html = "".join(item_rows) or "<div class='stats-empty'>这一天没有可展示的脚本。</div>"
        day_cards.append(
            "<details class='stats-day-card'>"
            "<summary class='stats-day-summary'>"
            f"<div class='stats-day-title'>{html_escape(day.get('date') or '')}</div>"
            "<div class='stats-day-metrics'>"
            f"<span class='stats-day-chip'>生成 {day.get('script_count') or 0}</span>"
            f"<span class='stats-day-chip'>复盘 {day.get('review_count') or 0}</span>"
            f"<span class='stats-day-chip'>修改 {day.get('edited_count') or 0}</span>"
            "</div>"
            "</summary>"
            f"<div class='stats-day-body'>{day_body_html}</div>"
            "</details>"
        )
    day_cards_html = "".join(day_cards) or "<div class='stats-empty'>还没有历史脚本数据。</div>"
    return f"""<!doctype html>
<html lang="zh-CN">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width,initial-scale=1">
  <title>Koko Stats</title>
  {FAVICON_LINKS}
  <style>
    @import url('https://fonts.googleapis.com/css2?family=Readex+Pro:wght@300;400;500;600;700&display=swap');
    * {{ box-sizing: border-box; font-family: 'Readex Pro', -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif; }}
    body {{
      margin: 0;
      min-height: 100vh;
      color: #FF8200;
      background:
        radial-gradient(circle at 8% 10%, rgba(255,130,0,.46), transparent 30%),
        radial-gradient(circle at 86% 12%, rgba(249,115,0,.38), transparent 28%),
        radial-gradient(circle at 50% 42%, rgba(255,178,84,.20), transparent 36%),
        linear-gradient(180deg, #FFB15A 0%, #FFD9AF 34%, #FFF1E2 70%, #FFFFFF 100%);
    }}
    .stats-shell {{ padding: 24px; }}
    .stats-wrap {{
      width: min(1320px, 100%);
      margin: 0 auto;
      border-radius: 34px;
      overflow: hidden;
      border: 1px solid rgba(255,255,255,.72);
      box-shadow: 0 28px 80px rgba(249,115,0,.16);
      background:
        radial-gradient(circle at 12% 16%, rgba(255,130,0,.32), rgba(255,130,0,0) 22%),
        radial-gradient(circle at 86% 18%, rgba(249,115,0,.26), rgba(249,115,0,0) 22%),
        radial-gradient(circle at 70% 62%, rgba(255,244,232,.70), rgba(255,244,232,0) 24%),
        linear-gradient(180deg, rgba(255,207,146,.64) 0%, rgba(255,240,222,.52) 44%, rgba(255,255,255,.62) 100%);
      backdrop-filter: blur(26px);
      -webkit-backdrop-filter: blur(26px);
      padding: 28px;
    }}
    .stats-topbar {{ display:flex; align-items:flex-start; justify-content:space-between; gap:20px; flex-wrap:wrap; }}
    .stats-topbar h1 {{ margin:0; font-size:clamp(3rem, 7vw, 5.4rem); letter-spacing:-.08em; line-height:.88; }}
    .stats-topbar p {{ margin:14px 0 0; font-size:14px; line-height:1.8; color:#FF8200; opacity:.9; max-width:56ch; }}
    .action-link {{
      display:inline-flex; align-items:center; justify-content:center; text-decoration:none;
      border-radius:999px; padding:10px 14px; color:#FF8200; border:1px solid rgba(255,130,0,.18); background:rgba(255,255,255,.72);
      font-weight:700; font-size:13px; cursor:pointer;
      backdrop-filter: blur(14px);
      -webkit-backdrop-filter: blur(14px);
    }}
    .stats-grid {{
      display:grid;
      grid-template-columns:repeat(auto-fit, minmax(180px, 1fr));
      gap:18px;
      margin-top:32px;
    }}
    .stats-summary-card {{
      border:1px solid rgba(255,130,0,.16);
      border-radius:24px;
      background:rgba(255,255,255,.58);
      padding:18px;
      display:flex;
      flex-direction:column;
      gap:8px;
      box-shadow: 0 18px 42px rgba(249,115,0,.10);
      backdrop-filter: blur(18px);
      -webkit-backdrop-filter: blur(18px);
    }}
    .stats-summary-label {{ font-size:12px; font-weight:800; letter-spacing:.08em; text-transform:uppercase; opacity:.82; }}
    .stats-summary-card strong {{ font-size:40px; line-height:1; color:#1F1F1F; }}
    .stats-summary-card small {{ font-size:13px; line-height:1.6; color:#FF8200; }}
    .stats-section {{ margin-top:34px; }}
    .stats-section-head {{
      display:flex; align-items:flex-end; justify-content:space-between; gap:16px; flex-wrap:wrap; margin-bottom:14px;
    }}
    .stats-section-head h2 {{ margin:0; font-size:28px; letter-spacing:-.04em; color:#1F1F1F; }}
    .stats-section-head p {{ margin:0; font-size:13px; line-height:1.6; color:#FF8200; opacity:.88; }}
    .stats-day-list {{ display:flex; flex-direction:column; gap:16px; }}
    .stats-day-card {{
      border:1px solid rgba(255,130,0,.16);
      border-radius:24px;
      background:rgba(255,255,255,.62);
      overflow:hidden;
      box-shadow: 0 18px 42px rgba(249,115,0,.08);
      backdrop-filter: blur(18px);
      -webkit-backdrop-filter: blur(18px);
    }}
    .stats-day-summary {{
      list-style:none; cursor:pointer; padding:18px 20px; display:flex; align-items:center; justify-content:space-between; gap:16px; flex-wrap:wrap;
    }}
    .stats-day-summary::-webkit-details-marker {{ display:none; }}
    .stats-day-title {{ font-size:22px; font-weight:800; color:#1F1F1F; letter-spacing:-.03em; }}
    .stats-day-metrics {{ display:flex; flex-wrap:wrap; gap:10px; }}
    .stats-day-chip {{
      display:inline-flex; align-items:center; border-radius:999px; padding:8px 12px;
      font-size:12px; font-weight:700; color:#FF8200; background:rgba(255,255,255,.74); border:1px solid rgba(255,130,0,.16);
    }}
    .stats-day-body {{ padding:0 20px 20px; display:flex; flex-direction:column; gap:12px; }}
    .stats-item-row {{
      display:flex; gap:12px; align-items:flex-start;
      border:1px solid rgba(255,130,0,.12); border-radius:18px; background:rgba(255,255,255,.72);
      padding:14px;
    }}
    .stats-item-index {{
      min-width:40px; height:40px; border-radius:999px;
      display:inline-flex; align-items:center; justify-content:center;
      background:rgba(255,130,0,.12); color:#FF8200; font-size:12px; font-weight:800;
    }}
    .stats-item-copy {{ display:flex; flex-direction:column; gap:8px; min-width:0; }}
    .stats-link {{ color:#2962FF; text-decoration:none; word-break:break-all; font-size:14px; line-height:1.7; }}
    .stats-item-meta {{ display:flex; flex-wrap:wrap; gap:8px; align-items:center; font-size:12px; color:#FF8200; }}
    .stats-badge {{
      display:inline-flex; align-items:center; border-radius:999px; padding:6px 10px; font-size:11px; font-weight:700;
      color:#935F14; background:rgba(147,95,20,.12);
    }}
    .stats-badge.yes {{ color:#157347; background:rgba(21,115,71,.12); }}
    .stats-empty {{
      border:1px dashed rgba(255,130,0,.18); border-radius:18px; background:rgba(255,255,255,.56);
      padding:18px; font-size:14px; line-height:1.7; color:#FF8200;
    }}
    .creator-analytics {{ display:flex; flex-direction:column; gap:16px; }}
    .creator-analytics-grid {{ display:grid; grid-template-columns:repeat(auto-fit,minmax(160px,1fr)); gap:14px; }}
    .creator-metric {{ border:1px solid rgba(255,130,0,.14); border-radius:22px; background:rgba(255,255,255,.68); padding:16px; box-shadow:0 16px 36px rgba(249,115,0,.08); }}
    .creator-metric span {{ display:block; color:#FF8200; font-size:12px; font-weight:800; letter-spacing:.04em; }}
    .creator-metric strong {{ display:block; margin-top:8px; color:#1F1F1F; font-size:34px; line-height:1; }}
    .creator-row-grid {{ display:grid; grid-template-columns:1.1fr 1fr; gap:16px; }}
    .creator-panel {{ border:1px solid rgba(255,130,0,.14); border-radius:24px; background:rgba(255,255,255,.68); padding:16px; overflow:hidden; }}
    .creator-panel h3 {{ margin:0 0 12px; color:#1F1F1F; font-size:18px; letter-spacing:-.03em; }}
    .creator-table {{ display:flex; flex-direction:column; gap:8px; }}
    .creator-table-row {{ display:grid; grid-template-columns:minmax(0,1fr) auto; gap:10px; align-items:center; padding:11px 0; border-top:1px solid rgba(255,130,0,.10); color:#1F1F1F; }}
    .creator-table-row:first-child {{ border-top:0; }}
    .creator-table-row b {{ display:block; font-size:13px; line-height:1.35; overflow:hidden; text-overflow:ellipsis; white-space:nowrap; }}
    .creator-table-row small {{ display:block; margin-top:3px; color:#FF8200; opacity:.82; font-size:11px; line-height:1.45; }}
    .creator-table-row code {{ color:#1F1F1F; background:rgba(255,130,0,.10); border-radius:999px; padding:5px 8px; font-size:12px; font-weight:800; }}
    .creator-event-chips {{ display:flex; flex-wrap:wrap; gap:8px; }}
    .creator-event-chip {{ display:inline-flex; align-items:center; gap:6px; border-radius:999px; padding:8px 11px; color:#FF8200; background:rgba(255,255,255,.72); border:1px solid rgba(255,130,0,.14); font-size:12px; font-weight:800; }}
    .creator-event-chip em {{ color:#1F1F1F; font-style:normal; }}
    .creator-daily {{ display:grid; grid-template-columns:repeat(auto-fit,minmax(120px,1fr)); gap:10px; }}
    .creator-day {{ border:1px solid rgba(255,130,0,.12); border-radius:18px; background:rgba(255,255,255,.62); padding:12px; }}
    .creator-day b {{ color:#1F1F1F; font-size:13px; }}
    .creator-day span {{ display:block; color:#FF8200; font-size:12px; line-height:1.65; margin-top:6px; }}
    @media (max-width: 760px) {{
      .stats-shell {{ padding: 12px; }}
      .stats-wrap {{ padding: 18px; }}
      .stats-day-summary {{ align-items:flex-start; }}
      .stats-item-row {{ flex-direction:column; }}
      .creator-row-grid {{ grid-template-columns:1fr; }}
      .creator-table-row b {{ white-space:normal; }}
    }}
  </style>
</head>
<body>
  <main class="stats-shell">
    <section class="stats-wrap">
      <div class="stats-topbar">
        <div>
          <button class="action-link" id="back-home" type="button">← 返回 Koko</button>
          <h1>Stats</h1>
          <p>这里会按北京时间聚合已有历史任务数据，统计每天生成了多少脚本、分别是哪些链接，以及它们是否触发过复盘重做和整稿编辑。</p>
        </div>
      </div>
      <div class="stats-grid">{summary_cards}</div>
      <section class="stats-section">
        <div class="stats-section-head">
          <div>
            <h2>Creator 使用情况</h2>
            <p>来自 kokocomedy 的前端埋点，按巴西时间聚合，脱敏统计访客与创作者行为。</p>
          </div>
          <button class="action-link" id="refresh-creator-analytics" type="button">刷新使用数据</button>
        </div>
        <div class="creator-analytics" id="creator-analytics"><div class="stats-empty">正在读取 kokocomedy 使用情况...</div></div>
      </section>
      <section class="stats-section">
        <div class="stats-section-head">
          <h2>每日统计</h2>
          <p>按天展开查看具体脚本链接、生成时间以及后续操作情况</p>
        </div>
        <div class="stats-day-list">{day_cards_html}</div>
      </section>
    </section>
  </main>
  <script>
    const backHomeButton = document.getElementById("back-home");
    if (backHomeButton) {{
      backHomeButton.addEventListener("click", () => {{
        window.location.assign("/");
      }});
    }}
    const escStats = (value) => String(value || "").replace(/[&<>"']/g, (c) => ({{"&":"&amp;","<":"&lt;",">":"&gt;",'"':"&quot;","'":"&#39;"}}[c]));
    const metricCard = (label, value) => `<article class="creator-metric"><span>${{escStats(label)}}</span><strong>${{Number(value || 0).toLocaleString("zh-CN")}}</strong></article>`;
    const row = (title, meta, value) => `<div class="creator-table-row"><div><b>${{escStats(title || "未命名")}}</b><small>${{escStats(meta || "")}}</small></div><code>${{escStats(value)}}</code></div>`;
    function renderCreatorAnalytics(data) {{
      const root = document.getElementById("creator-analytics");
      if (!root) return;
      const summary = data.summary || {{}};
      const daily = data.daily || [];
      const creators = data.creators || [];
      const scripts = data.scripts || [];
      const eventCounts = data.event_counts || [];
      const recent = data.recent_events || [];
      root.innerHTML = `
        <div class="creator-analytics-grid">
          ${{metricCard("昨日活跃创作者", summary.yesterday_active_creators)}}
          ${{metricCard("昨日独立访客估算", summary.yesterday_unique_visitors)}}
          ${{metricCard("昨日行为事件", summary.yesterday_events)}}
          ${{metricCard("昨日回传", summary.yesterday_submissions)}}
          ${{metricCard("累计行为事件", summary.total_events)}}
          ${{metricCard("累计回传", summary.total_submissions)}}
        </div>
        <div class="creator-panel">
          <h3>最近 14 天趋势</h3>
          <div class="creator-daily">${{daily.slice(0,14).map(d => `<div class="creator-day"><b>${{escStats(d.date)}}</b><span>创作者 ${{d.active_creators || 0}} · 访客 ${{d.unique_visitors || 0}}<br>脚本 ${{d.scripts || 0}} · 事件 ${{d.events || 0}} · 回传 ${{d.submissions || 0}}</span></div>`).join("") || '<div class="stats-empty">暂无趋势数据。</div>'}}</div>
        </div>
        <div class="creator-row-grid">
          <section class="creator-panel"><h3>活跃创作者</h3><div class="creator-table">${{creators.slice(0,10).map(c => row(c.creator_name || c.creator_id, `打开 ${{c.detail_views || 0}} · 收藏 ${{c.favorites || 0}} · 日历 ${{c.calendar_adds || 0}} · 分享 ${{c.shares || 0}} · 回传 ${{c.submissions || 0}}`, c.scripts || c.events || 0)).join("") || '<div class="stats-empty">还没有创作者行为。</div>'}}</div></section>
          <section class="creator-panel"><h3>热门脚本</h3><div class="creator-table">${{scripts.slice(0,10).map(s => row(s.title || s.script_id, `曝光 ${{s.impressions || 0}} · 打开 ${{s.detail_views || 0}} · 收藏 ${{s.favorites || 0}} · 日历 ${{s.calendar_adds || 0}} · 分享 ${{s.shares || 0}} · 回传 ${{s.submissions || 0}}`, s.creators || s.events || 0)).join("") || '<div class="stats-empty">还没有脚本行为。</div>'}}</div></section>
        </div>
        <div class="creator-row-grid">
          <section class="creator-panel"><h3>事件类型</h3><div class="creator-event-chips">${{eventCounts.slice(0,18).map(e => `<span class="creator-event-chip">${{escStats(e.event)}} <em>${{e.count || 0}}</em></span>`).join("") || '<div class="stats-empty">还没有事件。</div>'}}</div></section>
          <section class="creator-panel"><h3>最近事件</h3><div class="creator-table">${{recent.slice(0,10).map(e => row(e.event, `${{e.creator_name || e.creator_id || "匿名"}} · ${{e.created_at || ""}}`, e.script_id ? "脚本" : "页面")).join("") || '<div class="stats-empty">还没有事件流水。</div>'}}</div></section>
        </div>`;
    }}
    async function loadCreatorAnalytics() {{
      const root = document.getElementById("creator-analytics");
      if (root) root.innerHTML = '<div class="stats-empty">正在读取 kokocomedy 使用情况...</div>';
      try {{
        const res = await fetch(`/api/creator-admin/analytics?days=14&_=${{Date.now()}}`);
        const data = await res.json();
        if (!res.ok) throw new Error(data.error || "读取失败");
        renderCreatorAnalytics(data);
      }} catch (err) {{
        if (root) root.innerHTML = `<div class="stats-empty">${{escStats(err.message || err)}}<br>如果你还没登录 Creator 运营后台，请先打开 /creator-admin 登录一次。</div>`;
      }}
    }}
    document.getElementById("refresh-creator-analytics")?.addEventListener("click", loadCreatorAnalytics);
    loadCreatorAnalytics();
  </script>
</body>
</html>"""


def build_error_case_process_trace(output_dir: Path) -> dict[str, Any]:
    media_probe = read_json(output_dir / "media_probe.json")
    primary = read_json(output_dir / "primary_v2_draft.json")
    v2_local = read_json(output_dir / "v2_local_result.json")
    comparison = read_json(output_dir / "comparison_report.json")
    logic = read_json(output_dir / "logic_audit.json")
    supplement = read_json(output_dir / "supplement_evidence.json")
    conflict = read_json(output_dir / "conflict_recheck.json")
    arbitration = read_json(output_dir / "arbitration_result.json")
    type_router = read_json(output_dir / "type_router.json")
    audio_multiview = read_json(output_dir / "audio_multiview.json")
    review_plan = read_json(output_dir / "review_plan.json")
    review_video = read_json(output_dir / "review_video_recheck.json")
    trace_steps = [
        {
            "step": "download",
            "status": "present" if (output_dir / "source.mp4").exists() else "missing",
            "note": "source.mp4 exists" if (output_dir / "source.mp4").exists() else "source.mp4 missing",
        },
        {
            "step": "media_prep",
            "status": "present" if media_probe else "missing",
            "note": f"duration={media_probe.get('duration_sec') or media_probe.get('duration') or 0}" if media_probe else "media_probe missing",
        },
        {
            "step": "gemini_analysis",
            "status": "present" if primary else "missing",
            "note": f"route={primary.get('route') or ''}, audio_score={primary.get('audio_information_score') or ''}" if primary else "primary draft missing",
        },
        {
            "step": "v2_analysis",
            "status": "present" if v2_local else "missing",
            "note": f"route={v2_local.get('route') or ''}, audio_score={v2_local.get('audio_information_score') or ''}" if v2_local else "v2 local result missing",
        },
        {
            "step": "consistency_audit",
            "status": "present" if (comparison or logic) else "missing",
            "note": (
                f"comparison={comparison.get('recommended_action') or ''}, "
                f"story_alignment={(comparison.get('story_spine_alignment') or {}).get('status') or ''}, "
                f"logic={logic.get('recommended_action') or ''}"
            ) if (comparison or logic) else "comparison and logic audit missing",
        },
        {
            "step": "targeted_recheck",
            "status": "present" if (supplement or conflict or audio_multiview) else "missing",
            "note": (
                f"supplement_windows={len((supplement.get('windows') or [])) if supplement else 0}, "
                f"conflict_recheck_skipped={conflict.get('skipped') if conflict else ''}, "
                f"audio_multiview_skipped={audio_multiview.get('skipped') if audio_multiview else ''}"
            ) if (supplement or conflict or audio_multiview) else "no targeted recheck artifacts",
        },
        {
            "step": "arbitration",
            "status": "present" if arbitration else "missing",
            "note": f"accepted_pipeline={arbitration.get('accepted_pipeline') or ''}" if arbitration else "arbitration result missing",
        },
        {
            "step": "review_plan",
            "status": "present" if review_plan else "missing",
            "note": f"failure_layer={review_plan.get('likely_failure_layer') or ''}, needs_video_recheck={review_plan.get('needs_video_recheck')}" if review_plan else "review plan missing",
        },
        {
            "step": "review_video_recheck",
            "status": "present" if review_video else "missing",
            "note": f"skipped={review_video.get('skipped') if review_video else ''}, verification={review_video.get('verification_result') if review_video else ''}" if review_video else "review video recheck missing",
        },
        {
            "step": "final_output",
            "status": "present" if (output_dir / 'script_table.json').exists() else "missing",
            "note": "final script exists" if (output_dir / 'script_table.json').exists() else "final script missing",
        },
    ]
    return {
        "steps": trace_steps,
        "routing": {
            "primary_type": type_router.get("primary_type") or "",
            "subtype_guess": type_router.get("subtype_guess") or "",
            "routing_mode": type_router.get("routing_mode") or "",
            "reasoning_summary": type_router.get("reasoning_summary") or "",
        },
        "comparison_report": comparison,
        "logic_audit": logic,
        "supplement": supplement,
        "conflict_recheck": conflict,
        "arbitration_result": arbitration,
        "audio_multiview": audio_multiview,
    }


def fallback_error_case_review(
    review_plan: dict[str, Any],
    process_trace: dict[str, Any],
    original_script: dict[str, Any],
    corrected_script: dict[str, Any],
    feedback: str,
) -> dict[str, Any]:
    layer = str(review_plan.get("likely_failure_layer") or "").strip() or "unknown"
    original_summary = str(original_script.get("whole_video_summary") or original_script.get("title") or "").strip()
    corrected_summary = str(corrected_script.get("whole_video_summary") or corrected_script.get("title") or "").strip()
    notes = []
    comparison = process_trace.get("comparison_report") or {}
    if comparison:
        story_status = ((comparison.get("story_spine_alignment") or {}).get("status") or "").strip()
        if story_status:
            notes.append(f"一致性审查里的故事主轴状态为 {story_status}。")
    logic = process_trace.get("logic_audit") or {}
    if logic:
        rec = str(logic.get("recommended_action") or "").strip()
        if rec:
            notes.append(f"逻辑审查建议动作为 {rec}。")
    confidence = str(review_plan.get("confidence") or "medium").strip().lower()
    if confidence not in {"low", "medium", "high"}:
        confidence = "medium"
    return {
        "core_issue": str(review_plan.get("problem_summary") or feedback or "原始脚本主干理解错误").strip(),
        "primary_failure_layer": layer,
        "flow_failure_summary": str(review_plan.get("reasoning") or "主流程中的关键层未能正确识别或拦截错误主轴。").strip(),
        "difference_summary": f"原始脚本偏向：{original_summary or '无'}；修正后脚本偏向：{corrected_summary or '无'}",
        "missed_or_weak_links": [
            {
                "layer": layer,
                "status": "ran_but_missed" if layer != "unknown" else "insufficient_signal",
                "reason": "复盘计划判断该层是最可能的失效点，说明它执行过但没有正确纠偏。"
            }
        ],
        "evidence_notes": notes or ["需要结合 review_plan 与流程快照人工查看。"],
        "preventive_notes": [str(review_plan.get("correction_goal") or "下次遇到相似案例时，应优先复核主轴、实体和因果链。").strip()],
        "confidence": confidence,
    }


def summarize_error_case_with_llm(
    *,
    output_dir: Path,
    parent_job_id: str,
    item_index: int,
    item_id: str,
    feedback: str,
    original_script: dict[str, Any],
    corrected_script: dict[str, Any],
    review_plan: dict[str, Any],
    review_video: dict[str, Any],
) -> dict[str, Any]:
    process_trace = build_error_case_process_trace(output_dir)
    video_url = ""
    with job_lock:
        job = jobs.get(parent_job_id) or {}
        items = job.get("items") or []
        if 0 <= item_index < len(items):
            video_url = str(items[item_index].get("video_url") or "")
    summary = None
    if GOOGLE_API_KEY and run_text_json_prompt_with_fallback is not None:
        payload = {
            "job_id": parent_job_id,
            "item_index": item_index,
            "item_id": item_id,
            "video_url": video_url,
            "review_feedback": feedback,
            "original_script": original_script,
            "corrected_script": corrected_script,
            "review_plan": review_plan,
            "review_video_recheck": review_video,
            "process_trace": process_trace,
        }
        try:
            result, _, _ = run_text_json_prompt_with_fallback(
                payload,
                GOOGLE_API_KEY,
                unique_models(*MODEL_CANDIDATES),
                ERROR_CASE_REVIEW_PROMPT,
                "error case review",
            )
            if isinstance(result, dict):
                summary = result
        except Exception:
            summary = None
    if not isinstance(summary, dict):
        summary = fallback_error_case_review(review_plan, process_trace, original_script, corrected_script, feedback)
    if str(summary.get("primary_failure_layer") or "").strip() not in {
        "gemini_analysis",
        "v2_analysis",
        "consistency_audit",
        "targeted_recheck",
        "arbitration",
        "final_output",
        "unknown",
    }:
        summary["primary_failure_layer"] = "unknown"
    confidence = str(summary.get("confidence") or "").strip().lower()
    if confidence not in {"low", "medium", "high"}:
        summary["confidence"] = "medium"
    return {
        "entry_id": item_id,
        "job_id": parent_job_id,
        "item_index": item_index,
        "video_url": video_url,
        "title": corrected_script.get("title") or original_script.get("title") or "",
        "updated_at": now_iso(),
        "original_script": original_script,
        "review_feedback": feedback,
        "corrected_script": corrected_script,
        "review_plan": review_plan,
        "review_video_recheck": review_video,
        "process_trace": process_trace,
        "learning_review": summary,
    }


def record_error_case_library_entry(entry: dict[str, Any]) -> None:
    if not entry:
        return
    entries = load_error_case_entries()
    existing = next((item for item in entries if str(item.get("entry_id") or "") == str(entry.get("entry_id") or "")), None)
    if existing:
        preserved_created_at = existing.get("created_at") or existing.get("first_recorded_at") or now_iso()
        entry["created_at"] = preserved_created_at
        entry["first_recorded_at"] = existing.get("first_recorded_at") or preserved_created_at
        entry["review_count"] = int(existing.get("review_count") or 1) + 1
        entries = [item for item in entries if str(item.get("entry_id") or "") != str(entry.get("entry_id") or "")]
    else:
        created_at = now_iso()
        entry["created_at"] = created_at
        entry["first_recorded_at"] = created_at
        entry["review_count"] = 1
    entries.insert(0, entry)
    save_error_case_entries(entries[:500])


def parse_cookie_header(header_value: str) -> dict[str, str]:
    cookies: dict[str, str] = {}
    for part in (header_value or "").split(";"):
        if "=" not in part:
            continue
        name, value = part.split("=", 1)
        cookies[name.strip()] = value.strip()
    return cookies


def has_error_case_access(handler: BaseHTTPRequestHandler) -> bool:
    cookies = parse_cookie_header(handler.headers.get("Cookie", ""))
    return cookies.get(ERROR_CASE_AUTH_COOKIE) == "1"


def has_creator_admin_access(handler: BaseHTTPRequestHandler) -> bool:
    cookies = parse_cookie_header(handler.headers.get("Cookie", ""))
    token = urllib.parse.unquote(cookies.get(CREATOR_ADMIN_AUTH_COOKIE) or "")
    remote_token = urllib.parse.unquote(cookies.get(CREATOR_REMOTE_ADMIN_COOKIE) or "")
    return (
        bool(token) and secrets.compare_digest(token, CREATOR_ADMIN_PASSWORD)
    ) or (
        bool(remote_token) and secrets.compare_digest(remote_token, CREATOR_ADMIN_PASSWORD)
    )


def load_creator_admin_scripts_cache() -> dict[str, Any] | None:
    data = read_json_file(CREATOR_ADMIN_SCRIPTS_CACHE_FILE, default={})
    return data if isinstance(data, dict) and isinstance(data.get("entries"), list) else None


def save_creator_admin_scripts_cache(payload: dict[str, Any]) -> None:
    if not isinstance(payload.get("entries"), list):
        return
    cached = dict(payload)
    cached["cached_at"] = now_iso()
    write_json_atomic(CREATOR_ADMIN_SCRIPTS_CACHE_FILE, cached)


def load_creator_admin_state() -> dict[str, Any]:
    state = read_json_file(CREATOR_ADMIN_STATE_FILE, default={})
    if not isinstance(state, dict):
        state = {}
    creators = state.get("creators")
    if not isinstance(creators, dict):
        creators = {}
    return {"version": 1, "creators": creators, "updated_at": state.get("updated_at") or ""}


def creator_admin_state_for(profile_id: str) -> dict[str, Any]:
    state = load_creator_admin_state()
    creator_state = state.get("creators", {}).get(profile_id)
    return creator_state if isinstance(creator_state, dict) else {}


def save_creator_admin_state_for(profile_id: str, patch: dict[str, Any]) -> dict[str, Any]:
    if not re.fullmatch(r"[0-9a-f]{32}", profile_id):
        raise ValueError("Invalid creator profile id.")
    state = load_creator_admin_state()
    creators = state.setdefault("creators", {})
    current = creators.get(profile_id)
    if not isinstance(current, dict):
        current = {}
    next_state = dict(current)
    if "metrics" in patch:
        next_state["metrics"] = patch.get("metrics") if isinstance(patch.get("metrics"), dict) else {}
    if "feeds" in patch:
        feeds = patch.get("feeds")
        next_state["feeds"] = feeds if isinstance(feeds, list) else []
    if "deleted_feed_keys" in patch:
        deleted = patch.get("deleted_feed_keys")
        next_state["deleted_feed_keys"] = deleted if isinstance(deleted, list) else []
    next_state["updated_at"] = now_iso()
    creators[profile_id] = next_state
    state["updated_at"] = next_state["updated_at"]
    write_json_atomic(CREATOR_ADMIN_STATE_FILE, state)
    return next_state


def creator_admin_remote_json(path: str, *, method: str = "GET", payload: dict[str, Any] | None = None, timeout: int = 45) -> tuple[int, dict[str, Any]]:
    if not path.startswith("/"):
        path = "/" + path
    data = None if method.upper() == "GET" else json.dumps(payload or {}, ensure_ascii=False).encode("utf-8")
    opener = urllib.request.build_opener(urllib.request.ProxyHandler({}))
    attempts = 3 if method.upper() == "GET" else 1
    last_error = "Creator admin returned a non-JSON response."
    for attempt in range(attempts):
        request = urllib.request.Request(
            CREATOR_CENTER_BASE_URL + path,
            data=data,
            method=method.upper(),
            headers={
                "Content-Type": "application/json",
                "Accept-Encoding": "identity",
                "Connection": "close",
                "User-Agent": "KokoCreatorOps/1.0",
                "Cookie": f"{CREATOR_REMOTE_ADMIN_COOKIE}={urllib.parse.quote(CREATOR_ADMIN_PASSWORD)}",
            },
        )
        try:
            with opener.open(request, timeout=timeout) as response:
                chunks: list[bytes] = []
                while True:
                    try:
                        chunk = response.read(32768)
                    except http.client.IncompleteRead as exc:
                        chunk = bytes(exc.partial or b"")
                    if not chunk:
                        break
                    chunks.append(chunk)
                raw_bytes = b"".join(chunks)
                expected_length = response.headers.get("Content-Length")
                if expected_length and len(raw_bytes) < int(expected_length):
                    last_error = f"Creator admin response was interrupted ({len(raw_bytes)}/{expected_length})."
                    time.sleep(0.25 * (attempt + 1))
                    continue
                raw = raw_bytes.decode(response.headers.get_content_charset() or "utf-8", errors="ignore")
                status = int(getattr(response, "status", 200) or 200)
        except urllib.error.HTTPError as exc:
            raw = exc.read().decode(exc.headers.get_content_charset() or "utf-8", errors="ignore")
            status = int(exc.code or 500)
        except http.client.IncompleteRead as exc:
            raw = bytes(exc.partial or b"").decode("utf-8", errors="ignore")
            status = 502
            if not raw.strip().endswith(("}", "]")):
                last_error = "Creator admin response was interrupted. Please retry."
                time.sleep(0.25 * (attempt + 1))
                continue
        except TimeoutError:
            return 504, {"ok": False, "error": "Creator 数据接口请求超时，请稍后重试。"}
        except Exception as exc:
            message = str(exc) or exc.__class__.__name__
            if "timed out" in message.lower() or "timeout" in message.lower():
                return 504, {"ok": False, "error": "Creator 数据接口请求超时，请稍后重试。"}
            return 502, {"ok": False, "error": f"Creator 数据接口暂时不可用：{message}"}
        try:
            data_obj = json.loads(raw) if raw.strip() else {}
        except Exception:
            last_error = raw.strip() or "Creator admin returned a non-JSON response."
            time.sleep(0.25 * (attempt + 1))
            continue
        if not isinstance(data_obj, dict):
            data_obj = {"response": data_obj}
        return status, data_obj
    return 502, {"ok": False, "error": last_error}


def error_cases_login_html(error_message: str = "") -> str:
    message_html = f"<div class='login-error'>{html_escape(error_message)}</div>" if error_message else ""
    return f"""<!doctype html>
<html lang="zh-CN">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width,initial-scale=1">
  <title>Koko Error Cases Login</title>
  {FAVICON_LINKS}
  <style>
    @import url('https://fonts.googleapis.com/css2?family=Readex+Pro:wght@300;400;500;600;700&display=swap');
    * {{ box-sizing: border-box; font-family: 'Readex Pro', -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif; }}
    body {{
      margin: 0;
      min-height: 100vh;
      color: #FF8200;
      background:
        radial-gradient(circle at 8% 10%, rgba(255,130,0,.46), transparent 30%),
        radial-gradient(circle at 86% 12%, rgba(249,115,0,.38), transparent 28%),
        radial-gradient(circle at 50% 42%, rgba(255,178,84,.20), transparent 36%),
        linear-gradient(180deg, #FFB15A 0%, #FFD9AF 34%, #FFF1E2 70%, #FFFFFF 100%);
      display:flex;
      align-items:center;
      justify-content:center;
      padding:24px;
    }}
    .login-wrap {{
      width:min(560px, 100%);
      border-radius:34px;
      border:1px solid rgba(255,255,255,.72);
      box-shadow: 0 28px 80px rgba(249,115,0,.16);
      background:
        radial-gradient(circle at 12% 16%, rgba(255,130,0,.32), rgba(255,130,0,0) 22%),
        radial-gradient(circle at 86% 18%, rgba(249,115,0,.26), rgba(249,115,0,0) 22%),
        linear-gradient(180deg, rgba(255,207,146,.64) 0%, rgba(255,240,222,.52) 44%, rgba(255,255,255,.62) 100%);
      backdrop-filter: blur(26px);
      -webkit-backdrop-filter: blur(26px);
      padding:28px;
    }}
    .action-link {{
      display:inline-flex; align-items:center; justify-content:center; text-decoration:none;
      border-radius:999px; padding:10px 14px; color:#FF8200; border:1px solid rgba(255,130,0,.18); background:rgba(255,255,255,.72);
      font-weight:700; font-size:13px; cursor:pointer; backdrop-filter: blur(14px); -webkit-backdrop-filter: blur(14px);
    }}
    h1 {{ margin:18px 0 8px; font-size:clamp(2.6rem, 8vw, 4.4rem); letter-spacing:-.08em; line-height:.92; color:#1F1F1F; }}
    p {{ margin:0; font-size:14px; line-height:1.8; color:#FF8200; opacity:.92; }}
    form {{ display:flex; flex-direction:column; gap:14px; margin-top:24px; }}
    input {{
      width:100%; border-radius:22px; border:1px solid rgba(255,130,0,.16); background:rgba(255,255,255,.78);
      padding:18px 20px; font-size:16px; color:#1F1F1F; outline:none;
      box-shadow: inset 0 1px 0 rgba(255,255,255,.75), 0 16px 36px rgba(249,115,0,.06);
    }}
    button.action-link {{ width:100%; padding:16px 20px; font-size:15px; }}
    .login-error {{
      margin-top:18px; border:1px solid rgba(220,53,69,.18); background:rgba(220,53,69,.08);
      color:#B42318; border-radius:18px; padding:14px 16px; font-size:14px; line-height:1.7;
    }}
  </style>
</head>
<body>
  <section class="login-wrap">
    <button class="action-link" id="back-home" type="button">← 返回 Koko</button>
    <h1>错误案例库</h1>
    <p>这是一个内部页面，用来查看复盘重做后自动沉淀的错误案例。请输入访问密码继续。</p>
    {message_html}
    <form method="post" action="/error-cases/login">
      <input type="password" name="password" placeholder="请输入访问密码" autocomplete="current-password" />
      <button class="action-link" type="submit">进入错误案例库</button>
    </form>
  </section>
  <script>
    const backHomeButton = document.getElementById("back-home");
    if (backHomeButton) {{
      backHomeButton.addEventListener("click", () => {{
        window.location.assign("/");
      }});
    }}
  </script>
</body>
</html>"""


def error_cases_html() -> str:
    entries = load_error_case_entries()
    cards: list[str] = []
    for index, entry in enumerate(entries, start=1):
        review = entry.get("learning_review") or {}
        weak_links = review.get("missed_or_weak_links") or []
        evidence_notes = review.get("evidence_notes") or []
        preventive_notes = review.get("preventive_notes") or []
        weak_links_html = "".join(
            f"<li><strong>{html_escape(item.get('layer') or 'unknown')}</strong> · {html_escape(item.get('status') or '')}<br>{html_escape(item.get('reason') or '')}</li>"
            for item in weak_links if isinstance(item, dict)
        ) or "<li>暂无结构化弱点说明。</li>"
        evidence_html = "".join(
            f"<li>{html_escape(note)}</li>"
            for note in evidence_notes if str(note or "").strip()
        ) or "<li>暂无补充观察。</li>"
        preventive_html = "".join(
            f"<li>{html_escape(note)}</li>"
            for note in preventive_notes if str(note or "").strip()
        ) or "<li>暂无预防说明。</li>"
        cards.append(
            "<details class='error-case-card'>"
            "<summary class='error-case-summary'>"
            f"<div class='error-case-index'>#{index}</div>"
            "<div class='error-case-head'>"
            f"<div class='error-case-title'>{html_escape(entry.get('title') or '未命名脚本')}</div>"
            f"<a class='error-case-link' href='{html_escape(entry.get('video_url') or '')}' target='_blank' rel='noreferrer'>{html_escape(entry.get('video_url') or '')}</a>"
            "</div>"
            "<div class='error-case-meta'>"
            f"<span class='error-chip'>{html_escape(review.get('primary_failure_layer') or 'unknown')}</span>"
            f"<span class='error-chip'>{html_escape(str(review.get('confidence') or 'medium').upper())}</span>"
            f"<span class='error-chip'>复盘 {int(entry.get('review_count') or 1)} 次</span>"
            "</div>"
            "</summary>"
            "<div class='error-case-body'>"
            "<section class='error-block'>"
            "<h3>核心问题</h3>"
            f"<p>{html_escape(review.get('core_issue') or '')}</p>"
            "</section>"
            "<section class='error-block'>"
            "<h3>流程上为什么会错</h3>"
            f"<p>{html_escape(review.get('flow_failure_summary') or '')}</p>"
            "</section>"
            "<section class='error-block'>"
            "<h3>原始脚本与用户反馈</h3>"
            f"<p><strong>用户反馈：</strong>{html_escape(entry.get('review_feedback') or '')}</p>"
            f"<p><strong>原始标题：</strong>{html_escape((entry.get('original_script') or {}).get('title') or '')}</p>"
            f"<p><strong>原始梗概：</strong>{html_escape((entry.get('original_script') or {}).get('whole_video_summary') or '')}</p>"
            "</section>"
            "<section class='error-block'>"
            "<h3>修正后差异</h3>"
            f"<p>{html_escape(review.get('difference_summary') or '')}</p>"
            f"<p><strong>修正后标题：</strong>{html_escape((entry.get('corrected_script') or {}).get('title') or '')}</p>"
            f"<p><strong>修正后梗概：</strong>{html_escape((entry.get('corrected_script') or {}).get('whole_video_summary') or '')}</p>"
            "</section>"
            "<section class='error-columns'>"
            f"<div class='error-column'><h3>薄弱环节</h3><ul>{weak_links_html}</ul></div>"
            f"<div class='error-column'><h3>证据观察</h3><ul>{evidence_html}</ul></div>"
            f"<div class='error-column'><h3>预防备注</h3><ul>{preventive_html}</ul></div>"
            "</section>"
            "<section class='error-block'>"
            "<h3>流程快照</h3>"
            "<div class='trace-grid'>"
            + "".join(
                f"<article class='trace-chip'><strong>{html_escape(step.get('step') or '')}</strong><span>{html_escape(step.get('status') or '')}</span><small>{html_escape(step.get('note') or '')}</small></article>"
                for step in ((entry.get("process_trace") or {}).get("steps") or []) if isinstance(step, dict)
            ) +
            "</div>"
            "</section>"
            "<section class='error-block'>"
            f"<small>首次记录：{html_escape(format_beijing_time(entry.get('first_recorded_at')))} · 最近更新：{html_escape(format_beijing_time(entry.get('updated_at')))}</small>"
            "</section>"
            "</div>"
            "</details>"
        )
    cards_html = "".join(cards) or "<div class='error-empty'>还没有错误案例。只有复盘重做成功后的脚本才会被自动记录在这里。</div>"
    return f"""<!doctype html>
<html lang="zh-CN">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width,initial-scale=1">
  <title>Koko Error Cases</title>
  {FAVICON_LINKS}
  <style>
    @import url('https://fonts.googleapis.com/css2?family=Readex+Pro:wght@300;400;500;600;700&display=swap');
    * {{ box-sizing: border-box; font-family: 'Readex Pro', -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif; }}
    body {{
      margin: 0;
      min-height: 100vh;
      color: #FF8200;
      background:
        radial-gradient(circle at 8% 10%, rgba(255,130,0,.46), transparent 30%),
        radial-gradient(circle at 86% 12%, rgba(249,115,0,.38), transparent 28%),
        radial-gradient(circle at 50% 42%, rgba(255,178,84,.20), transparent 36%),
        linear-gradient(180deg, #FFB15A 0%, #FFD9AF 34%, #FFF1E2 70%, #FFFFFF 100%);
    }}
    .error-shell {{ padding: 24px; }}
    .error-wrap {{
      width: min(1360px, 100%);
      margin: 0 auto;
      border-radius: 34px;
      overflow: hidden;
      border: 1px solid rgba(255,255,255,.72);
      box-shadow: 0 28px 80px rgba(249,115,0,.16);
      background:
        radial-gradient(circle at 12% 16%, rgba(255,130,0,.32), rgba(255,130,0,0) 22%),
        radial-gradient(circle at 86% 18%, rgba(249,115,0,.26), rgba(249,115,0,0) 22%),
        radial-gradient(circle at 70% 62%, rgba(255,244,232,.70), rgba(255,244,232,0) 24%),
        linear-gradient(180deg, rgba(255,207,146,.64) 0%, rgba(255,240,222,.52) 44%, rgba(255,255,255,.62) 100%);
      backdrop-filter: blur(26px);
      -webkit-backdrop-filter: blur(26px);
      padding: 28px;
    }}
    .action-link {{
      display:inline-flex; align-items:center; justify-content:center; text-decoration:none;
      border-radius:999px; padding:10px 14px; color:#FF8200; border:1px solid rgba(255,130,0,.18); background:rgba(255,255,255,.72);
      font-weight:700; font-size:13px; cursor:pointer; backdrop-filter: blur(14px); -webkit-backdrop-filter: blur(14px);
    }}
    .topbar {{ display:flex; align-items:flex-start; justify-content:space-between; gap:20px; flex-wrap:wrap; }}
    h1 {{ margin:18px 0 8px; font-size:clamp(2.8rem, 7vw, 5rem); letter-spacing:-.08em; line-height:.9; color:#1F1F1F; }}
    .lede {{ margin:0; font-size:14px; line-height:1.8; color:#FF8200; opacity:.92; max-width:68ch; }}
    .error-list {{ display:flex; flex-direction:column; gap:16px; margin-top:30px; }}
    .error-case-card {{
      border:1px solid rgba(255,130,0,.16);
      border-radius:24px;
      background:rgba(255,255,255,.62);
      overflow:hidden;
      box-shadow: 0 18px 42px rgba(249,115,0,.08);
      backdrop-filter: blur(18px);
      -webkit-backdrop-filter: blur(18px);
    }}
    .error-case-summary {{
      list-style:none; cursor:pointer; padding:18px 20px; display:grid; grid-template-columns:auto minmax(0,1fr) auto; gap:16px; align-items:start;
    }}
    .error-case-summary::-webkit-details-marker {{ display:none; }}
    .error-case-index {{
      min-width:44px; height:44px; border-radius:999px; display:inline-flex; align-items:center; justify-content:center;
      background:rgba(255,130,0,.12); color:#FF8200; font-size:13px; font-weight:800;
    }}
    .error-case-head {{ min-width:0; display:flex; flex-direction:column; gap:8px; }}
    .error-case-title {{ font-size:24px; font-weight:800; line-height:1.25; letter-spacing:-.04em; color:#1F1F1F; }}
    .error-case-link {{ color:#2962FF; text-decoration:none; word-break:break-all; font-size:14px; line-height:1.7; }}
    .error-case-meta {{ display:flex; flex-wrap:wrap; gap:8px; justify-content:flex-end; }}
    .error-chip {{
      display:inline-flex; align-items:center; border-radius:999px; padding:8px 12px;
      font-size:12px; font-weight:700; color:#FF8200; background:rgba(255,255,255,.74); border:1px solid rgba(255,130,0,.16);
    }}
    .error-case-body {{ padding:0 20px 20px; display:flex; flex-direction:column; gap:14px; }}
    .error-block {{
      border:1px solid rgba(255,130,0,.12); border-radius:18px; background:rgba(255,255,255,.72);
      padding:16px;
    }}
    .error-block h3, .error-column h3 {{ margin:0 0 10px; font-size:16px; color:#1F1F1F; letter-spacing:-.02em; }}
    .error-block p, .error-block small, .error-column li {{
      margin:0; font-size:14px; line-height:1.8; color:#FF8200;
    }}
    .error-columns {{ display:grid; grid-template-columns:repeat(3, minmax(0,1fr)); gap:14px; }}
    .error-column {{
      border:1px solid rgba(255,130,0,.12); border-radius:18px; background:rgba(255,255,255,.72); padding:16px;
    }}
    .error-column ul {{ margin:0; padding-left:18px; display:flex; flex-direction:column; gap:10px; }}
    .trace-grid {{ display:grid; grid-template-columns:repeat(auto-fit, minmax(170px, 1fr)); gap:10px; }}
    .trace-chip {{
      border:1px solid rgba(255,130,0,.12); border-radius:16px; background:rgba(255,255,255,.78);
      padding:12px; display:flex; flex-direction:column; gap:6px;
    }}
    .trace-chip strong {{ font-size:13px; color:#1F1F1F; }}
    .trace-chip span {{ font-size:12px; font-weight:700; color:#FF8200; }}
    .trace-chip small {{ font-size:12px; line-height:1.6; color:#935F14; }}
    .error-empty {{
      border:1px dashed rgba(255,130,0,.18); border-radius:18px; background:rgba(255,255,255,.56);
      padding:18px; font-size:14px; line-height:1.7; color:#FF8200; margin-top:24px;
    }}
    @media (max-width: 900px) {{
      .error-case-summary {{ grid-template-columns:1fr; }}
      .error-case-meta {{ justify-content:flex-start; }}
      .error-columns {{ grid-template-columns:1fr; }}
    }}
  </style>
</head>
<body>
  <main class="error-shell">
    <section class="error-wrap">
      <div class="topbar">
        <div>
          <button class="action-link" id="back-home" type="button">← 返回 Koko</button>
          <h1>错误案例库</h1>
          <p class="lede">这里只展示点击“复盘重做”后自动沉淀下来的重大错误案例。每条案例都会记录原始脚本、用户反馈、修正结果，以及当时实际走过的流程为什么没拦住问题。</p>
        </div>
      </div>
      <section class="error-list">{cards_html}</section>
    </section>
  </main>
  <script>
    const backHomeButton = document.getElementById("back-home");
    if (backHomeButton) {{
      backHomeButton.addEventListener("click", () => {{
        window.location.assign("/");
      }});
    }}
  </script>
</body>
</html>"""


def write_product_outputs(job_id: str, output_dir: Path, result_json: dict[str, Any]) -> dict[str, str]:
    bundle = build_evidence_bundle(job_id, output_dir, result_json)
    bundle_path = output_dir / "evidence_bundle.json"
    report_path = output_dir / "product_report.html"
    bundle_path.write_text(json.dumps(bundle, ensure_ascii=False, indent=2), encoding="utf-8")
    report_path.write_text(render_product_report(bundle), encoding="utf-8")
    return {
        "report_url": f"/results/{job_id}/product_report.html",
        "evidence_url": f"/results/{job_id}/evidence_bundle.json",
    }


def build_understanding_summary(script_json: dict[str, Any] | None, video_url: str = "") -> str:
    script = script_json if isinstance(script_json, dict) else {}
    candidates = [
        script.get("whole_video_summary"),
        script.get("content_summary"),
        script.get("summary"),
        (script.get("mechanism") or {}).get("summary") if isinstance(script.get("mechanism"), dict) else "",
        (script.get("story_analysis") or {}).get("safe_final_story") if isinstance(script.get("story_analysis"), dict) else "",
    ]
    for value in candidates:
        text = compact_cell_text(str(value or ""))
        if text:
            return text
    rows = script.get("rows")
    if isinstance(rows, list):
        fragments: list[str] = []
        for row in rows[:6]:
            if not isinstance(row, dict):
                continue
            fragments.extend(
                compact_cell_text(str(row.get(key) or ""))
                for key in ["visual_content", "action", "dialogue_or_audio", "integrated_summary"]
                if row.get(key)
            )
        text = compact_cell_text("；".join(fragment for fragment in fragments if fragment))
        if text:
            return text[:520]
    title = compact_cell_text(str(script.get("title") or ""))
    if title:
        return f"这个视频的核心内容是：{title}"
    if video_url:
        return f"已经完成视频理解，但没有提取到稳定的文字概要。原链接：{video_url}"
    return "已经完成视频理解，但没有提取到稳定的文字概要。"


def public_item_view(item: dict[str, Any]) -> dict[str, Any]:
    item_id = str(item.get("id") or "").strip()
    source_video_path = RESULTS_ROOT / item_id / SOURCE_VIDEO_NAME if item_id else Path()
    source_video_available = bool(item_id and source_video_path.exists())
    storyboard_state = load_storyboard_state(item_id) if item_id else {}
    return {
        "id": item.get("id"),
        "index": item.get("index"),
        "video_url": item.get("video_url"),
        "user_prompt": item.get("user_prompt") or "",
        "status": item.get("status"),
        "updated_at": item.get("updated_at"),
        "completed_at": item.get("completed_at"),
        "stage": item.get("stage") or "",
        "stage_message": item.get("stage_message") or "",
        "html_url": item.get("html_url") or "",
        "report_url": item.get("report_url") or "",
        "evidence_url": item.get("evidence_url") or "",
        "docx_url": item.get("docx_url") or "",
        "zh_docx_url": item.get("zh_docx_url") or item.get("docx_url") or "",
        "pt_docx_url": item.get("pt_docx_url") or "",
        "zh_html_url": item.get("zh_html_url") or item.get("html_url") or "",
        "pt_html_url": item.get("pt_html_url") or "",
        "error": item.get("error") or "",
        "artifacts": item.get("artifacts") or {},
        "result_json": item.get("result_json"),
        "zh_result_json": item.get("zh_result_json") or item.get("result_json"),
        "pt_result_json": item.get("pt_result_json"),
        "content_type": item.get("content_type") or "",
        "content_type_source": item.get("content_type_source") or "auto",
        "content_type_reasoning": item.get("content_type_reasoning") or "",
        "content_type_confidence": item.get("content_type_confidence") or "",
        "title": item.get("title") or "",
        "understanding_summary": item.get("understanding_summary")
        or build_understanding_summary(item.get("result_json") or item.get("zh_result_json"), item.get("video_url") or ""),
        "display_language": item.get("display_language") or "zh",
        "review_status": item.get("review_status") or "",
        "review_stage": item.get("review_stage") or "",
        "review_message": item.get("review_message") or "",
        "review_feedback": item.get("review_feedback") or "",
        "review_mode": normalize_review_mode(item.get("review_mode")),
        "chat_messages": item.get("chat_messages") if isinstance(item.get("chat_messages"), list) else [],
        "reviewed": bool(item.get("reviewed")),
        "edited": bool(item.get("edited")),
        "saved_to_library_at": item.get("saved_to_library_at") or "",
        "in_library": bool(item.get("saved_to_library_at")) or library_entry_exists(str(item.get("id") or "")),
        "source_video_available": source_video_available,
        "source_video_url": f"/results/{item_id}/{SOURCE_VIDEO_NAME}" if source_video_available else "",
        "storyboard_prompt": item.get("storyboard_prompt") or storyboard_state.get("storyboard_prompt") or "",
        "storyboard_preview_url": item.get("storyboard_preview_url") or storyboard_state.get("storyboard_preview_url") or "",
        "storyboard_cover_url": item.get("storyboard_cover_url") or storyboard_state.get("storyboard_cover_url") or "",
        "storyboard_updated_at": item.get("storyboard_updated_at") or storyboard_state.get("storyboard_updated_at") or "",
    }


def public_job_view(job: dict[str, Any]) -> dict[str, Any]:
    hydrated_items: list[dict[str, Any]] = []
    for item in job.get("items") or []:
        if item_output_ready(item) and (
            not item.get("result_json")
            or not item.get("html_url")
            or str(item.get("status") or "").strip() != "completed"
        ):
            hydrate_item_from_outputs(item)
        hydrated_items.append(item)
    items = [public_item_view(item) for item in hydrated_items]
    total_items = len(items)
    completed_items = sum(1 for item in items if item.get("status") == "completed")
    failed_items = sum(1 for item in items if item.get("status") == "failed")
    system_queue = build_system_queue_snapshot(str(job.get("id") or "").strip())
    payload = {
        "id": job["id"],
        "mode": job.get("mode") or ("batch" if total_items > 1 else "single"),
        "status": job["status"],
        "created_at": job.get("created_at"),
        "updated_at": job.get("updated_at"),
        "video_url": job.get("video_url"),
        "video_urls": job.get("video_urls") or [],
        "user_prompt": job.get("user_prompt") or "",
        "error": job.get("error"),
        "html_url": job.get("html_url"),
        "report_url": job.get("report_url"),
        "evidence_url": job.get("evidence_url"),
        "docx_url": job.get("docx_url"),
        "result_json": job.get("result_json"),
        "artifacts": job.get("artifacts") or {},
        "stage": job.get("stage") or "",
        "stage_message": job.get("stage_message") or "",
        "items": items,
        "total_items": total_items,
        "completed_items": completed_items,
        "failed_items": failed_items,
        "system_queue": system_queue,
    }
    if job.get("status") == "running":
        payload["message"] = job.get("stage_message") or "正在提取视频证据并生成脚本，请稍候。"
    elif job.get("status") == "queued":
        payload["message"] = "任务已进入队列。"
    elif job.get("status") == "completed":
        if total_items > 1:
            payload["message"] = f"已完成 {completed_items}/{total_items} 条脚本，失败 {failed_items} 条。"
        else:
            payload["message"] = "证据报告和最终脚本都已生成。"
    elif job.get("status") == "failed":
        payload["message"] = "分析失败，请检查错误信息。"
    return payload


def public_filter_item_view(item: dict[str, Any]) -> dict[str, Any]:
    visual = item.get("visual") or {}
    thumbnail_faces = visual.get("thumbnail_faces") or {}
    audio = item.get("audio") or {}
    checks = item.get("checks") or {}
    evidence_bundle = item.get("evidence_bundle") or {}
    return {
        "id": str(item.get("id") or "").strip(),
        "index": int(item.get("index") or 0),
        "video_url": item.get("video_url") or "",
        "display_name": item.get("display_name") or "",
        "status": item.get("status") or "",
        "stage": item.get("stage") or "",
        "stage_message": item.get("stage_message") or "",
        "bucket": item.get("bucket") or "",
        "confidence": item.get("confidence") or "",
        "reason": item.get("reason") or "",
        "signals": item.get("signals") or [],
        "score": item.get("score") or 0,
        "thumbnail_url": item.get("thumbnail_url") or "",
        "metadata": item.get("metadata") or {},
        "evidence_url": item.get("evidence_url") or "",
        "audio": {
            "available": bool(audio.get("available")),
            "source": audio.get("source") or "",
            "language": audio.get("language") or "",
            "dialogue_summary": audio.get("dialogue_summary") or "",
            "speaker_hints": audio.get("speaker_hints") or [],
            "audio_form": audio.get("audio_form") or "",
            "confidence": audio.get("confidence") or "",
            "full_transcript": audio.get("full_transcript") or "",
            "error": audio.get("error") or "",
        },
        "checks": {
            "duration_check": checks.get("duration_check") or {},
            "multi_character_check": checks.get("multi_character_check") or {},
            "story_check": checks.get("story_check") or {},
            "final_result": checks.get("final_result") or "",
        },
        "evidence_bundle": {
            "metadata": evidence_bundle.get("metadata") or {},
            "audio": {
                key: value
                for key, value in (evidence_bundle.get("audio") or {}).items()
                if key != "full_transcript"
            },
            "frames": evidence_bundle.get("frames") or {},
        },
        "visual": {
            "available": bool(visual.get("available")),
            "frame_count": int(visual.get("frame_count") or 0),
            "inspected_frames": int(visual.get("inspected_frames") or 0),
            "male_count": int(visual.get("male_count") or 0),
            "female_count": int(visual.get("female_count") or 0),
            "pair_frames": int(visual.get("pair_frames") or 0),
            "has_both": bool(visual.get("has_both")),
            "max_faces_single_frame": int(visual.get("max_faces_single_frame") or 0),
            "frame_summaries": visual.get("frame_summaries") or [],
            "thumbnail_faces": {
                "available": bool(thumbnail_faces.get("available")),
                "face_count": int(thumbnail_faces.get("face_count") or 0),
            },
        },
        "error": item.get("error") or "",
    }


def public_filter_job_view(job: dict[str, Any]) -> dict[str, Any]:
    items = [public_filter_item_view(item) for item in (job.get("items") or [])]
    matched = [item for item in items if item.get("bucket") == "high"]
    return {
        "id": str(job.get("id") or "").strip(),
        "status": job.get("status") or "",
        "created_at": job.get("created_at"),
        "updated_at": job.get("updated_at"),
        "stage": job.get("stage") or "",
        "stage_message": job.get("stage_message") or "",
        "input_count": len(items),
        "matched_count": len(matched),
        "items": items,
        "matched_links": [item.get("video_url") for item in matched if item.get("video_url")],
        "message": job.get("message") or "",
    }


def public_translation_job_view(job: dict[str, Any]) -> dict[str, Any]:
    return {
        "id": str(job.get("id") or "").strip(),
        "video_url": job.get("video_url") or "",
        "status": job.get("status") or "",
        "created_at": job.get("created_at"),
        "updated_at": job.get("updated_at"),
        "completed_at": job.get("completed_at") or "",
        "stage": job.get("stage") or "",
        "stage_message": job.get("stage_message") or "",
        "language": job.get("language") or "pt-BR",
        "subject_summary": job.get("subject_summary") or "",
        "original_audio_summary": job.get("original_audio_summary") or "",
        "portuguese_voiceover": job.get("portuguese_voiceover") or "",
        "translated_video_url": job.get("translated_video_url") or "",
        "audio_url": job.get("audio_url") or "",
        "metadata": job.get("metadata") or {},
        "error": job.get("error") or "",
    }


def create_filter_job(video_urls: list[str], *, source_label: str = "") -> dict[str, Any]:
    job_id = uuid4().hex
    items: list[dict[str, Any]] = []
    for index, video_url in enumerate(video_urls):
        items.append(
            {
                "id": uuid4().hex,
                "index": index,
                "video_url": video_url,
                "display_name": parse_video_display_name(video_url, index),
                "status": "queued",
                "stage": "queued",
                "stage_message": "Queued.",
                "created_at": now_iso(),
                "updated_at": now_iso(),
                "bucket": "",
                "confidence": "",
                "reason": "",
                "signals": [],
                "score": 0,
                "thumbnail_url": "",
                "metadata": {},
                "audio": {},
                "visual": {},
                "evidence_bundle": {},
                "checks": {},
                "error": "",
            }
        )
    job = {
        "id": job_id,
        "status": "queued",
        "created_at": now_iso(),
        "updated_at": now_iso(),
        "stage": "queued",
        "stage_message": "Queued.",
        "source_label": source_label,
        "items": items,
        "message": "",
    }
    with filter_jobs_lock:
        filter_jobs[job_id] = job
        save_filter_jobs()
    with filter_queue_condition:
        if job_id not in queued_filter_job_ids:
            queued_filter_job_ids.add(job_id)
            filter_queue.append(job_id)
            filter_queue_condition.notify()
    return public_filter_job_view(job)


def create_translation_job(video_url: str, *, language: str = "pt-BR") -> dict[str, Any]:
    job_id = uuid4().hex
    job = {
        "id": job_id,
        "video_url": video_url,
        "status": "queued",
        "created_at": now_iso(),
        "updated_at": now_iso(),
        "completed_at": "",
        "stage": "queued",
        "stage_message": "Queued.",
        "language": language or "pt-BR",
        "subject_summary": "",
        "original_audio_summary": "",
        "portuguese_voiceover": "",
        "translated_video_url": "",
        "audio_url": "",
        "metadata": {},
        "error": "",
    }
    with translation_jobs_lock:
        translation_jobs[job_id] = job
        save_translation_jobs()
    with translation_queue_condition:
        if job_id not in queued_translation_job_ids:
            queued_translation_job_ids.add(job_id)
            translation_queue.append(job_id)
            translation_queue_condition.notify()
    return public_translation_job_view(job)


def update_translation_job(job_id: str, **changes: Any) -> None:
    with translation_jobs_lock:
        job = translation_jobs.get(job_id)
        if not job:
            return
        job.update(changes)
        job["updated_at"] = now_iso()
        save_translation_jobs()


def run_translation_job(job_id: str) -> None:
    with translation_jobs_lock:
        job = dict(translation_jobs.get(job_id) or {})
    if not job:
        return
    output_dir = RESULTS_ROOT / job_id
    output_dir.mkdir(parents=True, exist_ok=True)
    update_translation_job(job_id, status="running", stage="download", stage_message="正在下载源视频。", error="")
    try:
        if not TRANSCREATE_VIDEO.exists():
            raise RuntimeError(f"Missing transcreation entrypoint: {TRANSCREATE_VIDEO}")
        command = [
            sys.executable,
            str(TRANSCREATE_VIDEO),
            str(job.get("video_url") or "").strip(),
            "--out",
            str(output_dir),
            "--language",
            str(job.get("language") or "pt-BR"),
        ]
        proc = subprocess.run(command, text=True, capture_output=True, timeout=PIPELINE_TIMEOUT_SEC)
        result_path = output_dir / "translation_result.json"
        result = read_json_file(result_path, default={}) if result_path.exists() else {}
        if proc.returncode != 0 or not result.get("ok"):
            error = str(result.get("error") or proc.stderr or proc.stdout or "Translation job failed.").strip()
            raise RuntimeError(error)
        translated_path = output_dir / "translated_pt.mp4"
        audio_path = output_dir / "portuguese_voiceover.aiff"
        update_translation_job(
            job_id,
            status="completed",
            completed_at=now_iso(),
            stage="completed",
            stage_message="转译视频已生成。",
            subject_summary=result.get("subject_summary") or "",
            original_audio_summary=result.get("original_audio_summary") or "",
            portuguese_voiceover=result.get("portuguese_voiceover") or "",
            translated_video_url=f"/results/{job_id}/{translated_path.name}" if translated_path.exists() else "",
            audio_url=f"/results/{job_id}/{audio_path.name}" if audio_path.exists() else "",
            metadata=result.get("metadata") or {},
            error="",
        )
    except Exception as exc:
        update_translation_job(
            job_id,
            status="failed",
            completed_at=now_iso(),
            stage="failed",
            stage_message="转译失败。",
            error=friendly_error(str(exc)),
        )


def translation_worker_loop() -> None:
    while True:
        with translation_queue_condition:
            while not translation_queue:
                translation_queue_condition.wait()
            job_id = translation_queue.popleft()
            queued_translation_job_ids.discard(job_id)
        run_translation_job(job_id)


def start_translation_workers() -> None:
    for index in range(MAX_CONCURRENT_TRANSLATIONS):
        thread = threading.Thread(target=translation_worker_loop, name=f"koko-translation-worker-{index+1}", daemon=True)
        thread.start()


def update_filter_job(job_id: str, **changes: Any) -> None:
    with filter_jobs_lock:
        job = filter_jobs.get(job_id)
        if not job:
            return
        job.update(changes)
        job["updated_at"] = now_iso()
        save_filter_jobs()


def update_filter_item(job_id: str, index: int, **changes: Any) -> None:
    with filter_jobs_lock:
        job = filter_jobs.get(job_id)
        if not job:
            return
        items = job.get("items") or []
        if not (0 <= index < len(items)):
            return
        items[index].update(changes)
        items[index]["updated_at"] = now_iso()
        job["updated_at"] = now_iso()
        save_filter_jobs()


def finalize_filter_job(job_id: str) -> None:
    with filter_jobs_lock:
        job = filter_jobs.get(job_id)
        if not job:
            return
        items = job.get("items") or []
        statuses = [str(item.get("status") or "").strip() for item in items]
        matched = sum(1 for item in items if str(item.get("bucket") or "").strip() == "high")
        if items and all(status == "completed" for status in statuses):
            job["status"] = "completed"
            job["stage"] = "completed"
            job["stage_message"] = f"Completed {len(items)}/{len(items)} items."
            job["message"] = f"已筛出 {matched} 条同时通过时长、多人物和剧情三轮规则的视频。"
        elif any(status == "running" for status in statuses):
            job["status"] = "running"
        elif any(status == "queued" for status in statuses):
            job["status"] = "queued"
        else:
            job["status"] = "failed"
            job["stage"] = "failed"
            job["stage_message"] = "筛选失败。"
            job["message"] = "没有成功完成任何筛选项。"
        job["updated_at"] = now_iso()
        save_filter_jobs()


def run_filter_job(job_id: str) -> None:
    update_filter_job(job_id, status="running", stage="metadata", stage_message="正在抓取页面公开信息。", message="")
    with filter_jobs_lock:
        job = filter_jobs.get(job_id)
        items = list(job.get("items") or []) if job else []
    any_completed = False
    for index, item in enumerate(items):
        update_filter_item(job_id, index, status="running", stage="metadata", stage_message="正在抓取页面公开信息。")
        metadata: dict[str, Any] = {}
        audio: dict[str, Any] = {}
        visual: dict[str, Any] = {}
        evidence_bundle: dict[str, Any] = {}
        try:
            metadata = fetch_kwai_light_metadata(str(item.get("video_url") or "").strip())
            cache_dir = FILTER_CACHE_ROOT / job_id / str(item.get("id") or f"item-{index}")
            update_filter_item(
                job_id,
                index,
                stage="audio",
                stage_message="正在提取完整音频信息。",
                metadata=metadata,
                thumbnail_url=metadata.get("thumbnail_url") or "",
            )
            audio = transcribe_filter_audio(metadata, cache_dir)
            update_filter_item(
                job_id,
                index,
                stage="frames",
                stage_message="正在抽取开头、中间、结尾三张关键帧。",
                metadata=metadata,
                thumbnail_url=metadata.get("thumbnail_url") or "",
                audio=audio,
            )
            try:
                visual = detect_gender_presence_from_frames(
                    str(metadata.get("content_url") or "").strip(),
                    metadata.get("duration"),
                    cache_dir,
                )
            except Exception as exc:
                visual = {
                    "available": False,
                    "reason": f"visual fallback: {exc}",
                    "bucket": "low",
                    "signals": ["关键帧男女识别失败，已回退到文本初筛。"],
                    "score_boost": 0,
                }
            visual["thumbnail_faces"] = count_thumbnail_faces(str(metadata.get("thumbnail_url") or "").strip(), cache_dir)
            evidence_bundle = build_filter_evidence_bundle(metadata, audio, visual)
            write_json_atomic(cache_dir / "evidence_bundle.json", evidence_bundle)
            update_filter_item(
                job_id,
                index,
                stage="classify",
                stage_message="正在进行时长、多人物和剧情三轮筛选。",
                metadata=metadata,
                thumbnail_url=metadata.get("thumbnail_url") or "",
                audio=audio,
                visual=visual,
                evidence_bundle=evidence_bundle,
            )
            decision = classify_story_candidate(evidence_bundle)
            write_json_atomic(cache_dir / "filter_decision.json", decision)
            update_filter_item(
                job_id,
                index,
                status="completed",
                stage="completed",
                stage_message="筛选完成。",
                bucket=decision.get("bucket") or "low",
                confidence=decision.get("confidence") or "low",
                reason=decision.get("reason") or "",
                signals=decision.get("signals") or [],
                metadata=metadata,
                thumbnail_url=metadata.get("thumbnail_url") or "",
                score=3 if decision.get("bucket") == "high" else 0,
                audio=audio,
                visual=visual,
                evidence_bundle=evidence_bundle,
                evidence_url=str(cache_dir / "evidence_bundle.json"),
                checks={
                    "duration_check": decision.get("duration_check") or {},
                    "multi_character_check": decision.get("multi_character_check") or {},
                    "story_check": decision.get("story_check") or {},
                    "final_result": decision.get("final_result") or "",
                },
            )
            any_completed = True
        except Exception as exc:
            failure_updates: dict[str, Any] = {}
            if metadata:
                failure_updates["metadata"] = metadata
                failure_updates["thumbnail_url"] = metadata.get("thumbnail_url") or ""
            if audio:
                failure_updates["audio"] = audio
            if visual:
                failure_updates["visual"] = visual
            if evidence_bundle:
                failure_updates["evidence_bundle"] = evidence_bundle
            update_filter_item(
                job_id,
                index,
                status="failed",
                stage="failed",
                stage_message="筛选失败。",
                error=friendly_error(str(exc)),
                reason="筛选流程没有成功完成，不能判定为通过或不通过。",
                **failure_updates,
            )
    finalize_filter_job(job_id)
    if not any_completed:
        update_filter_job(job_id, message="没有成功完成任何筛选项。")


def filter_worker_loop() -> None:
    while True:
        try:
            with filter_queue_condition:
                while not filter_queue:
                    filter_queue_condition.wait()
                job_id = filter_queue.popleft()
                queued_filter_job_ids.discard(job_id)
            with filter_jobs_lock:
                job_exists = job_id in filter_jobs
            if not job_exists:
                log_runtime_warning("filter_worker_skipped_missing_job", "Skipped queued filter job because it no longer exists.", job_id=job_id)
                continue
            run_filter_job(job_id)
        except Exception as exc:
            log_runtime_warning("filter_worker_loop_error", "Filter worker recovered after an unexpected error.", error=str(exc))
            time.sleep(1)


def start_filter_workers() -> None:
    for index in range(MAX_CONCURRENT_FILTERS):
        thread = threading.Thread(target=filter_worker_loop, name=f"koko-filter-worker-{index+1}", daemon=True)
        thread.start()


def update_job(job_id: str, **changes: Any) -> None:
    with job_lock:
        job = jobs[job_id]
        job.update(changes)
        job["updated_at"] = now_iso()
        save_jobs()


def update_job_item(job_id: str, item_index: int, **changes: Any) -> None:
    with job_lock:
        job = jobs[job_id]
        item = job["items"][item_index]
        item.update(changes)
        item["updated_at"] = now_iso()
        job["updated_at"] = now_iso()
        job["current_item_index"] = item_index
        if "stage" in changes and "stage_message" in changes:
            total = len(job["items"])
            job["stage"] = changes.get("stage") or ""
            job["stage_message"] = f"Video {item_index + 1}/{total}: {changes.get('stage_message') or ''}"
        save_jobs()


def start_review_job(item_id: str, feedback: str, review_mode: str = REVIEW_MODE_PARTIAL) -> tuple[bool, str]:
    context = find_item_context(item_id)
    if not context:
        return False, "Script item not found."
    parent_job_id, item_index, item = context
    if item.get("status") != "completed" or not item.get("result_json"):
        return False, "Only completed scripts can be reviewed."
    feedback_text = str(feedback or "").strip()
    if not feedback_text:
        return False, "Please describe what the analysis got wrong."
    mode = normalize_review_mode(review_mode)
    update_job_item(
        parent_job_id,
        item_index,
        review_status="running",
        review_stage="queued",
        review_message="Queued for review. Waiting for an available analysis slot.",
        review_feedback=feedback_text,
        review_mode=mode,
        reviewed=False,
    )
    threading.Thread(
        target=run_review_with_slot,
        args=(parent_job_id, item_index, item_id, feedback_text, mode),
        daemon=True,
    ).start()
    return True, parent_job_id


def append_item_chat_message(parent_job_id: str, item_index: int, role: str, content: str, **extra: Any) -> None:
    text = str(content or "").strip()
    if not text:
        return
    with job_lock:
        item = jobs[parent_job_id]["items"][item_index]
        messages = item.get("chat_messages")
        if not isinstance(messages, list):
            messages = []
        message = {
            "role": role,
            "content": text,
            "created_at": now_iso(),
        }
        message.update({key: value for key, value in extra.items() if value not in (None, "")})
        messages.append(message)
        item["chat_messages"] = messages[-40:]
        item["updated_at"] = now_iso()
        jobs[parent_job_id]["updated_at"] = now_iso()
        save_jobs()


def extract_gemini_text(raw: dict[str, Any]) -> str:
    chunks: list[str] = []
    for candidate in raw.get("candidates") or []:
        content = candidate.get("content") or {}
        for part in content.get("parts") or []:
            text = part.get("text")
            if isinstance(text, str):
                chunks.append(text)
    text = "\n".join(chunks).strip()
    if not text:
        raise RuntimeError("Gemini returned no text.")
    return text


def parse_json_object_from_text(text: str) -> dict[str, Any]:
    value = str(text or "").strip()
    if value.startswith("```"):
        value = re.sub(r"^```(?:json)?\s*", "", value, flags=re.IGNORECASE)
        value = re.sub(r"\s*```$", "", value)
    try:
        parsed = json.loads(value)
    except json.JSONDecodeError:
        start = value.find("{")
        end = value.rfind("}")
        if start < 0 or end <= start:
            raise
        parsed = json.loads(value[start : end + 1])
    if not isinstance(parsed, dict):
        raise RuntimeError("Gemini did not return a JSON object.")
    return parsed


CHAT_EDIT_HTTP_TIMEOUT_SEC = max(20, int(os.environ.get("KOKO_CHAT_EDIT_HTTP_TIMEOUT_SEC", "60")))


def parse_gemini_raw_response(raw_text: str) -> tuple[dict[str, Any], dict[str, Any]]:
    raw = json.loads(raw_text)
    return parse_json_object_from_text(extract_gemini_text(raw)), raw


def run_single_chat_text_json_prompt(payload: dict[str, Any], prompt: str, model: str) -> tuple[dict[str, Any], dict[str, Any]]:
    body = {
        "contents": [
            {
                "parts": [
                    {"text": prompt},
                    {"text": json.dumps(payload, ensure_ascii=False)},
                ]
            }
        ],
        "generationConfig": {
            "response_mime_type": "application/json",
        },
    }
    url = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent"
    data = json.dumps(body).encode("utf-8")
    last_error: Exception | None = None
    for attempt in range(1, 3):
        try:
            req = urllib.request.Request(
                url,
                data=data,
                headers={"Content-Type": "application/json", "x-goog-api-key": GOOGLE_API_KEY},
            )
            with urllib.request.urlopen(req, timeout=CHAT_EDIT_HTTP_TIMEOUT_SEC) as response:
                return parse_gemini_raw_response(response.read().decode("utf-8"))
        except urllib.error.HTTPError as exc:
            detail = exc.read().decode("utf-8", "replace")
            last_error = RuntimeError(f"Gemini text HTTP {exc.code}: {detail}")
            if exc.code in {400, 404}:
                break
        except Exception as exc:
            last_error = exc
        if attempt < 2:
            time.sleep(min(2 * attempt, 6))
    curl_path = shutil.which("curl")
    if curl_path:
        try:
            completed = subprocess.run(
                [
                    curl_path,
                    "-sS",
                    "--fail-with-body",
                    "--retry",
                    "1",
                    "--retry-delay",
                    "1",
                    "--max-time",
                    str(CHAT_EDIT_HTTP_TIMEOUT_SEC),
                    "-H",
                    "Content-Type: application/json",
                    "-H",
                    f"x-goog-api-key: {GOOGLE_API_KEY}",
                    "--data-binary",
                    "@-",
                    url,
                ],
                input=data,
                capture_output=True,
                timeout=CHAT_EDIT_HTTP_TIMEOUT_SEC + 10,
            )
            stdout = completed.stdout.decode("utf-8", "replace")
            stderr = completed.stderr.decode("utf-8", "replace")
            if completed.returncode == 0:
                return parse_gemini_raw_response(stdout)
            last_error = RuntimeError((stdout or stderr or f"curl exited {completed.returncode}").strip())
        except Exception as exc:
            last_error = exc
    raise RuntimeError(str(last_error or "Gemini text request failed."))


def run_chat_text_json_prompt(payload: dict[str, Any], prompt: str) -> tuple[dict[str, Any], dict[str, Any], str]:
    last_error: Exception | None = None
    tried: list[str] = []
    for model in unique_models(*MODEL_CANDIDATES, *PRIMARY_FALLBACK_MODELS, *SUPPLEMENT_FALLBACK_MODELS):
        tried.append(model)
        try:
            data, raw = run_single_chat_text_json_prompt(payload, prompt, model)
            return data, raw, model
        except Exception as exc:
            last_error = exc
            continue
    raise RuntimeError(f"chat script edit failed across models {tried}: {last_error}") from last_error


STORYBOARD_IMAGE_PROMPT_PREFIX = """你是 Koko 的分镜示意图生成助手。

请把输入的短视频脚本整理成一张“分镜稿 / storyboard sheet”风格的示意图，要求：
- 黑白灰铅笔草图风格，像手绘分镜稿，不要彩色，不要照片质感
- 白色纸面背景，深灰色线稿，构图干净
- 最终成图固定为 1:1 正方形
- 一张图里固定排成 3 行 x 3 列，共 9 个等宽等高矩形分镜格，像影视前期的分镜板
- 如果脚本关键动作少于 9 个，用空白环境格、道具格或过渡动作格补齐到 9 格；不要改变 3x3 网格
- 如果脚本关键动作多于 9 个，合并相近动作，最终仍然只保留 9 格；不要改变 3x3 网格
- 每格表现脚本里的一个关键动作节点，人物姿态和场景关系要清楚
- 整体重点是“拍摄准备感”，让人一眼看懂场景、人物、动作
- 不要做成海报，不要 UI，不要水印，不要品牌字
- 画面里不要出现任何文字、字幕、标题、编号、logo 或水印
- 画面以连续叙事为主，保留夸张表情和关键动作
"""


STORYBOARD_PROMPT_GUARDRAILS = """固定硬性条件，必须保留且不能被其他要求覆盖：
- 固定 1:1 正方形构图
- 白底纸面
- 黑白灰手绘线稿 / 铅笔草图风格
- 固定 3 行 x 3 列，共 9 个等宽等高矩形分镜格，像 storyboard sheet
- 少于 9 个动作时用空白环境格、道具格或过渡动作格补齐；多于 9 个动作时合并动作；始终保持 3x3
- 非彩色，非照片，非海报，非 UI
- 不要带任何文字、字幕、标题、编号、logo 或水印
"""


STORYBOARD_PROMPT_DRAFT_PROMPT = """你是 Koko 的短剧分镜生图提示词助手。

你会收到一份短视频脚本 JSON，可能已经被运营人员手动修改过。请基于最新脚本，写一段可以直接给生图模型使用的分镜示意图提示词。

目标：
- 让图片像“分镜稿 / storyboard sheet”，用于给创作者理解拍摄节奏。
- 提示词必须围绕脚本里最新的标题、整体梗概、脚本表动作节点。
- 每格只描述画面和动作，不要要求生成任何文字。

硬性要求：
1. 只输出 JSON，不要 Markdown。
2. JSON 只有一个字段：`prompt`。
3. `prompt` 可以中文写，但要清晰、可执行。
4. 必须包含这些不可变条件：固定 1:1 正方形、白底纸面、手绘线稿、固定 3x3 九宫格分镜、非彩色、非照片、不要带任何文字。
5. 不要生成海报式封面，不要大标题，不要 UI，不要 logo。

输出格式：
{
  "prompt": "完整生图提示词"
}
"""


def enforce_storyboard_prompt_guardrails(prompt: str) -> str:
    body = str(prompt or "").strip()
    if not body:
        body = STORYBOARD_IMAGE_PROMPT_PREFIX.strip()
    banned_text_clause = "画面里不要出现任何文字、字幕、标题、编号、logo 或水印。"
    if banned_text_clause not in body:
        body = f"{body}\n\n{banned_text_clause}"
    guardrails = STORYBOARD_PROMPT_GUARDRAILS.strip()
    if guardrails not in body:
        body = f"{guardrails}\n\n{body}"
    return body.strip()


def guess_extension_from_mime(mime_type: str) -> str:
    value = str(mime_type or "").strip().lower()
    if value == "image/jpeg":
        return ".jpg"
    if value == "image/webp":
        return ".webp"
    return ".png"


def build_storyboard_prompt(script_json: dict[str, Any], extra_instruction: str = "") -> str:
    rows = select_storyboard_rows(choose_script_rows(script_json), max_panels=9)
    beats = []
    for idx, row in enumerate(rows, 1):
        beats.append(
            f"{idx}. 时间={fill_text(row.get('time'), '')}；场景={fill_text(row.get('visual_content'))}；动作={fill_text(row.get('action'))}"
        )
    title = fill_text(script_json.get("title"), "视频脚本")
    summary = fill_text(script_json.get("whole_video_summary"), "无")
    prompt = [
        STORYBOARD_IMAGE_PROMPT_PREFIX.strip(),
        "",
        f"标题：{title}",
        f"整体梗概：{summary}",
        "关键分镜：",
        "\n".join(beats) if beats else "1. 按标题和梗概生成 9 格 3x3 连续分镜。",
        "",
        "请输出一张适合作为脚本封面的分镜示意图。",
    ]
    extra = str(extra_instruction or "").strip()
    if extra:
        prompt.extend(["", f"额外修改要求：{extra}"])
    return enforce_storyboard_prompt_guardrails("\n".join(prompt).strip())


def select_storyboard_rows(rows: list[dict[str, Any]], max_panels: int = 9) -> list[dict[str, Any]]:
    usable = [row for row in rows if isinstance(row, dict)]
    if len(usable) <= max_panels:
        return usable
    if max_panels <= 1:
        return usable[:1]
    last_index = len(usable) - 1
    selected_indexes: list[int] = []
    for slot in range(max_panels):
        index = round(slot * last_index / (max_panels - 1))
        if index not in selected_indexes:
            selected_indexes.append(index)
    cursor = 0
    while len(selected_indexes) < max_panels and cursor < len(usable):
        if cursor not in selected_indexes:
            selected_indexes.append(cursor)
        cursor += 1
    selected_indexes.sort()
    return [usable[index] for index in selected_indexes[:max_panels]]


def generate_storyboard_prompt_from_script(script_json: dict[str, Any]) -> tuple[str, dict[str, Any], str]:
    fallback = build_storyboard_prompt(script_json)
    if not GOOGLE_API_KEY or run_text_json_prompt_with_fallback is None:
        return fallback, {}, "fallback"
    payload = {
        "title": script_json.get("title") or "",
        "whole_video_summary": script_json.get("whole_video_summary") or "",
        "core_viral_points": script_json.get("core_viral_points") or [],
        "replaceable_parts": script_json.get("replaceable_parts") or [],
        "rows": select_storyboard_rows(choose_script_rows(script_json), max_panels=9),
        "fixed_guardrails": STORYBOARD_PROMPT_GUARDRAILS,
    }
    try:
        result, raw, model = run_text_json_prompt_with_fallback(
            payload,
            GOOGLE_API_KEY,
            unique_models(*MODEL_CANDIDATES, *PRIMARY_FALLBACK_MODELS, *SUPPLEMENT_FALLBACK_MODELS),
            STORYBOARD_PROMPT_DRAFT_PROMPT,
            "storyboard prompt draft",
        )
        prompt = enforce_storyboard_prompt_guardrails(str((result or {}).get("prompt") or "").strip() or fallback)
        return prompt, raw if isinstance(raw, dict) else {}, model
    except Exception:
        return fallback, {}, "fallback"


def extract_inline_image_from_gemini_response(payload: dict[str, Any]) -> tuple[bytes, str]:
    for candidate in payload.get("candidates") or []:
        content = candidate.get("content") or {}
        for part in content.get("parts") or []:
            inline = part.get("inlineData") or part.get("inline_data") or {}
            data = inline.get("data")
            if not data:
                continue
            mime_type = str(inline.get("mimeType") or inline.get("mime_type") or "image/png")
            return base64.b64decode(data), mime_type
    raise RuntimeError("Gemini image response did not contain inline image data.")


def run_single_gemini_image_prompt(prompt: str, model: str) -> tuple[bytes, str, dict[str, Any]]:
    body = {
        "contents": [{"parts": [{"text": prompt}]}],
    }
    url = f"https://generativelanguage.googleapis.com/v1/models/{model}:generateContent"
    data = json.dumps(body).encode("utf-8")
    last_error: Exception | None = None
    for attempt in range(1, 3):
        try:
            req = urllib.request.Request(
                url,
                data=data,
                headers={"Content-Type": "application/json", "x-goog-api-key": GOOGLE_API_KEY},
            )
            with urllib.request.urlopen(req, timeout=CHAT_EDIT_HTTP_TIMEOUT_SEC) as response:
                raw = json.loads(response.read().decode("utf-8"))
                image_bytes, mime_type = extract_inline_image_from_gemini_response(raw)
                return image_bytes, mime_type, raw
        except urllib.error.HTTPError as exc:
            detail = exc.read().decode("utf-8", "replace")
            last_error = RuntimeError(f"Gemini image HTTP {exc.code}: {detail}")
            if exc.code in {400, 404}:
                break
        except Exception as exc:
            last_error = exc
        if attempt < 2:
            time.sleep(min(2 * attempt, 6))
    raise RuntimeError(str(last_error or "Gemini image request failed."))


def run_gemini_image_prompt(prompt: str) -> tuple[bytes, str, dict[str, Any], str]:
    last_error: Exception | None = None
    tried: list[str] = []
    for model in IMAGE_MODEL_CANDIDATES:
        tried.append(model)
        try:
            image_bytes, mime_type, raw = run_single_gemini_image_prompt(prompt, model)
            return image_bytes, mime_type, raw, model
        except Exception as exc:
            last_error = exc
            continue
    raise RuntimeError(f"storyboard image generation failed across models {tried}: {last_error}") from last_error


def storyboard_scene_tokens(row: dict[str, Any]) -> set[str]:
    text = " ".join(
        [
            str(row.get("visual_content") or ""),
            str(row.get("action") or ""),
            str(row.get("dialogue") or ""),
        ]
    ).lower()
    tokens: set[str] = set()
    mapping = {
        "door": ["porta", "portão", "portao", "gate", "door"],
        "bed": ["cama", "bed", "quarto"],
        "car": ["carro", "car", "dirig", "volante"],
        "table": ["mesa", "table", "bolo", "jantar"],
        "phone": ["celular", "telefone", "phone", "ligação", "ligacao"],
        "money": ["dinheiro", "conta", "pagar", "pix", "salário", "salario"],
        "street": ["rua", "estrada", "calçada", "calcada", "parque", "outdoor"],
        "yard": ["quintal", "varanda", "campo", "bananeira", "roça", "roca"],
        "hospital": ["hospital", "médico", "medico", "consulta", "paciente"],
        "kitchen": ["cozinha", "fogão", "fogao", "panela", "geladeira"],
        "laundry": ["lavar", "máquina", "maquina", "roupa", "varal"],
        "baby": ["bebê", "bebe", "filho", "recém", "recem", "criança", "crianca"],
        "surprise": ["assust", "surpresa", "espelho", "pegadinha", "reviravolta"],
    }
    for token, words in mapping.items():
        if any(word in text for word in words):
            tokens.add(token)
    return tokens


def storyboard_character_count(row: dict[str, Any]) -> int:
    text = " ".join(
        [
            str(row.get("visual_content") or ""),
            str(row.get("action") or ""),
            str(row.get("dialogue") or ""),
        ]
    ).lower()
    hints = [
        "homem", "mulher", "marido", "esposa", "amigo", "amiga", "namorado", "namorada",
        "médico", "medico", "paciente", "pai", "mãe", "mae", "filho", "vizinho", "vizinha", "policial",
    ]
    count = sum(1 for word in hints if word in text)
    if count <= 0:
        return 1
    if count >= 4:
        return 3
    return min(3, max(1, count))


def draw_storyboard_character(draw: ImageDraw.ImageDraw, x: int, y: int, scale: int, *, pose: str = "stand") -> None:
    head_r = max(6, scale // 6)
    draw.ellipse((x - head_r, y - scale, x + head_r, y - scale + head_r * 2), outline=40, width=2)
    torso_top = y - scale + head_r * 2
    torso_bottom = y - scale // 3
    if pose == "lie":
        draw.line((x - scale // 3, torso_top + scale // 8, x + scale // 3, torso_top + scale // 8), fill=40, width=2)
        draw.line((x - scale // 5, torso_top + scale // 8, x - scale // 2, torso_top + scale // 3), fill=40, width=2)
        draw.line((x + scale // 5, torso_top + scale // 8, x + scale // 2, torso_top + scale // 3), fill=40, width=2)
        return
    draw.line((x, torso_top, x, torso_bottom), fill=40, width=2)
    draw.line((x, torso_top + scale // 6, x - scale // 4, torso_top + scale // 3), fill=40, width=2)
    draw.line((x, torso_top + scale // 6, x + scale // 4, torso_top + scale // 3), fill=40, width=2)
    draw.line((x, torso_bottom, x - scale // 5, y), fill=40, width=2)
    draw.line((x, torso_bottom, x + scale // 5, y), fill=40, width=2)


def draw_storyboard_prop(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], tokens: set[str]) -> None:
    left, top, right, bottom = box
    width = right - left
    height = bottom - top
    mid_y = top + height * 2 // 3
    draw.line((left + 8, mid_y, right - 8, mid_y), fill=120, width=2)
    if "street" in tokens or "yard" in tokens:
        draw.arc((left + 10, top + 10, left + 38, top + 28), 180, 360, fill=150, width=2)
        draw.arc((left + 40, top + 16, left + 72, top + 32), 180, 360, fill=150, width=2)
    if "door" in tokens:
        draw.rectangle((right - 42, top + 22, right - 16, mid_y), outline=80, width=2)
    if "bed" in tokens:
        draw.rectangle((left + 12, mid_y - 20, left + width // 2 + 12, mid_y + 8), outline=80, width=2)
        draw.rectangle((left + 16, mid_y - 24, left + 36, mid_y - 12), outline=120, width=1)
    if "table" in tokens:
        tx1 = left + width // 2 - 28
        tx2 = left + width // 2 + 28
        ty = mid_y - 8
        draw.line((tx1, ty, tx2, ty), fill=80, width=2)
        draw.line((tx1 + 6, ty, tx1 + 2, mid_y + 18), fill=80, width=2)
        draw.line((tx2 - 6, ty, tx2 - 2, mid_y + 18), fill=80, width=2)
    if "car" in tokens:
        cx1 = left + 12
        cx2 = left + width // 2 + 10
        cy = mid_y - 18
        draw.rounded_rectangle((cx1, cy, cx2, cy + 24), radius=6, outline=80, width=2)
        draw.ellipse((cx1 + 10, cy + 20, cx1 + 22, cy + 32), outline=80, width=2)
        draw.ellipse((cx2 - 22, cy + 20, cx2 - 10, cy + 32), outline=80, width=2)
    if "phone" in tokens:
        draw.rounded_rectangle((right - 36, mid_y - 30, right - 20, mid_y - 2), radius=3, outline=60, width=2)
    if "money" in tokens:
        draw.rectangle((left + 14, top + 18, left + 34, top + 30), outline=80, width=2)
        draw.line((left + 18, top + 24, left + 30, top + 24), fill=120, width=1)
    if "hospital" in tokens:
        hx = right - 28
        hy = top + 22
        draw.line((hx - 6, hy, hx + 6, hy), fill=80, width=2)
        draw.line((hx, hy - 6, hx, hy + 6), fill=80, width=2)
    if "baby" in tokens:
        draw.arc((left + 14, top + 12, left + 38, top + 30), 200, 340, fill=80, width=2)
        draw.line((left + 22, top + 24, left + 28, top + 34), fill=80, width=2)
    if "surprise" in tokens:
        sx = left + width - 24
        sy = top + 24
        draw.line((sx, sy - 8, sx, sy + 8), fill=60, width=2)
        draw.line((sx - 8, sy, sx + 8, sy), fill=60, width=2)
        draw.line((sx - 6, sy - 6, sx + 6, sy + 6), fill=60, width=2)
        draw.line((sx + 6, sy - 6, sx - 6, sy + 6), fill=60, width=2)


def render_local_storyboard_image(item_id: str, script_json: dict[str, Any]) -> bytes:
    rows = select_storyboard_rows(choose_script_rows(script_json), max_panels=9)
    canvas_size = 1080
    margin = 44
    gap = 18
    panel_size = (canvas_size - margin * 2 - gap * 2) // 3
    image = Image.new("L", (canvas_size, canvas_size), color=248)
    draw = ImageDraw.Draw(image)
    draw.rounded_rectangle((8, 8, canvas_size - 8, canvas_size - 8), radius=28, outline=120, width=3)
    for idx in range(9):
        row = rows[idx] if idx < len(rows) else {}
        col = idx % 3
        r = idx // 3
        left = margin + col * (panel_size + gap)
        top = margin + r * (panel_size + gap)
        right = left + panel_size
        bottom = top + panel_size
        draw.rectangle((left, top, right, bottom), outline=95, width=3)
        inner = (left + 8, top + 8, right - 8, bottom - 8)
        tokens = storyboard_scene_tokens(row if isinstance(row, dict) else {})
        draw_storyboard_prop(draw, inner, tokens)
        count = storyboard_character_count(row if isinstance(row, dict) else {})
        poses = ["stand", "stand", "stand"]
        if "bed" in tokens:
            poses[0] = "lie"
        if "surprise" in tokens and count >= 2:
            poses[-1] = "stand"
        anchors = [
            (left + panel_size // 4, bottom - 32),
            (left + panel_size // 2, bottom - 28),
            (left + panel_size * 3 // 4, bottom - 32),
        ]
        for person_index in range(count):
            x, y = anchors[min(person_index, len(anchors) - 1)]
            scale = 76 if person_index == 1 else 68
            draw_storyboard_character(draw, x, y, scale, pose=poses[min(person_index, len(poses) - 1)])
    rgb = image.convert("RGB")
    output = io.BytesIO()
    rgb.save(output, format="PNG")
    return output.getvalue()


def run_chat_script_edit(item_id: str, message: str, edit_mode: str = "minor") -> tuple[bool, str | dict[str, Any]]:
    context = find_item_context(item_id)
    if not context:
        return False, "Script item not found."
    parent_job_id, item_index, item = context
    if item.get("status") != "completed" or not item.get("result_json"):
        return False, "Only completed scripts can be edited by Koko."
    user_message = str(message or "").strip()
    if not user_message:
        return False, "请先告诉 Koko 你想改哪里。"
    mode = str(edit_mode or "minor").strip().lower()
    if mode not in {"minor", "major", "replace"}:
        mode = "minor"
    if not GOOGLE_API_KEY:
        return False, "Missing GOOGLE_API_KEY for Koko edit."
    if run_text_json_prompt_with_fallback is None:
        return False, "Koko edit helpers are unavailable."

    output_dir = RESULTS_ROOT / item_id
    current_script = read_json(output_dir / "script_table.json") or item.get("result_json") or {}
    if not current_script:
        return False, "No existing script result to edit."
    chat_messages = item.get("chat_messages") or []
    append_item_chat_message(parent_job_id, item_index, "user", user_message, mode=mode)
    request_payload = {
        "edit_mode": mode,
        "user_message": user_message,
        "conversation": chat_messages[-16:],
        "current_script": current_script,
        "video_url": item.get("video_url") or "",
        "source_metadata": read_json(output_dir / "source_metadata.json"),
    }
    try:
        result, raw, _ = run_chat_text_json_prompt(
            request_payload,
            CHAT_SCRIPT_EDIT_PROMPT,
        )
        (output_dir / "chat_script_edit_raw_gemini.json").write_text(json.dumps(raw, ensure_ascii=False, indent=2), encoding="utf-8")
        corrected_payload = extract_review_script_payload(result)
        if not corrected_payload:
            raise RuntimeError("Koko 没有返回可用的完整脚本 JSON。")
        merged_script = json.loads(json.dumps(current_script, ensure_ascii=False))
        for key in REVIEW_SCRIPT_KEYS:
            if corrected_payload.get(key):
                merged_script[key] = corrected_payload.get(key)
        if not review_script_changed(current_script, merged_script, REVIEW_SCRIPT_KEYS):
            raise RuntimeError("Koko 这次没有生成任何脚本变更，请把错误点说得再具体一点。")
        merged_script = enforce_chinese_dialogue_translation(
            merged_script,
            GOOGLE_API_KEY,
            unique_models(*MODEL_CANDIDATES),
        )
        updated_item = regenerate_item_outputs(
            parent_job_id,
            item_index,
            item_id,
            item.get("video_url") or "",
            merged_script,
            persist_library=False,
            target_language=item.get("display_language") or "zh",
        )
        assistant_message = ""
        if isinstance(result, dict):
            assistant_message = str(result.get("assistant_message") or "").strip()
            summary = result.get("change_summary")
            if not assistant_message and isinstance(summary, list):
                assistant_message = "已修改：" + "；".join(str(item or "").strip() for item in summary if str(item or "").strip())
        if not assistant_message:
            assistant_message = "我已经按你的反馈修改了脚本，并刷新了左侧的可编辑版本。"
        append_item_chat_message(parent_job_id, item_index, "assistant", assistant_message, mode=mode)
        with job_lock:
            refreshed = public_item_view(jobs[parent_job_id]["items"][item_index])
        return True, {"item": refreshed, "message": assistant_message}
    except Exception as exc:
        error_message = friendly_error(str(exc))
        append_item_chat_message(parent_job_id, item_index, "assistant", f"这次没改成功：{error_message}", mode=mode, error=True)
        return False, error_message


def generate_storyboard_assets(
    item_id: str,
    output_dir: Path,
    script_json: dict[str, Any],
    *,
    prompt_override: str = "",
    attempts: int = CREATOR_IMPORT_IMAGE_RETRY_ATTEMPTS,
) -> dict[str, str]:
    attempts = max(1, int(attempts or 1))
    output_dir.mkdir(parents=True, exist_ok=True)
    candidate_prompts: list[str] = []
    primary_prompt = enforce_storyboard_prompt_guardrails(str(prompt_override or "").strip() or build_storyboard_prompt(script_json))
    candidate_prompts.append(primary_prompt)
    drafted_prompt, _, drafted_model = generate_storyboard_prompt_from_script(script_json)
    drafted_prompt = enforce_storyboard_prompt_guardrails(str(drafted_prompt or "").strip() or primary_prompt)
    if drafted_prompt and drafted_prompt not in candidate_prompts:
        candidate_prompts.insert(0, drafted_prompt)
    if STORYBOARD_LOCAL_ONLY:
        prompt = candidate_prompts[0]
        image_bytes = render_local_storyboard_image(item_id, script_json)
        preview_name = STORYBOARD_PREVIEW_BASENAME + ".png"
        preview_path = output_dir / preview_name
        preview_path.write_bytes(image_bytes)
        (output_dir / STORYBOARD_PROMPT_FILE).write_text(prompt, encoding="utf-8")
        (output_dir / "storyboard_image_raw_local.json").write_text(
            json.dumps({"mode": "local_storyboard_renderer"}, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
        cover_name = STORYBOARD_COVER_BASENAME + ".png"
        shutil.copyfile(preview_path, output_dir / cover_name)
        save_storyboard_state(
            item_id,
            prompt=prompt,
            preview_name=preview_name,
            cover_name=cover_name,
            model="local_storyboard_renderer",
        )
        return {
            "prompt": prompt,
            "prompt_model": drafted_model if prompt == drafted_prompt else "rule-based",
            "image_model": "local_storyboard_renderer",
            "preview_name": preview_name,
            "cover_name": cover_name,
            "cover_url": f"/results/{item_id}/{cover_name}",
        }
    last_error: Exception | None = None
    for attempt_index in range(attempts):
        prompt = candidate_prompts[min(attempt_index, len(candidate_prompts) - 1)]
        try:
            image_bytes, mime_type, raw, model = run_gemini_image_prompt(prompt)
            preview_name = STORYBOARD_PREVIEW_BASENAME + guess_extension_from_mime(mime_type)
            preview_path = output_dir / preview_name
            preview_path.write_bytes(image_bytes)
            (output_dir / "storyboard_image_raw_gemini.json").write_text(json.dumps(raw, ensure_ascii=False, indent=2), encoding="utf-8")
            (output_dir / STORYBOARD_PROMPT_FILE).write_text(prompt, encoding="utf-8")
            cover_name = STORYBOARD_COVER_BASENAME + preview_path.suffix.lower()
            shutil.copyfile(preview_path, output_dir / cover_name)
            save_storyboard_state(
                item_id,
                prompt=prompt,
                preview_name=preview_name,
                cover_name=cover_name,
                model=model,
            )
            return {
                "prompt": prompt,
                "prompt_model": drafted_model if prompt == drafted_prompt else "rule-based",
                "image_model": model,
                "preview_name": preview_name,
                "cover_name": cover_name,
                "cover_url": f"/results/{item_id}/{cover_name}",
            }
        except Exception as exc:
            last_error = exc
            if attempt_index + 1 < attempts:
                time.sleep(min(2 + attempt_index, 5))
    raise RuntimeError(friendly_error(str(last_error or "Storyboard generation failed.")))


def generate_storyboard_preview(item_id: str, prompt_override: str = "", edits_payload: dict[str, Any] | None = None) -> dict[str, Any]:
    context = find_item_context(item_id)
    if not context:
        raise RuntimeError("Script item not found.")
    parent_job_id, item_index, item = context
    if item.get("status") != "completed" or not item.get("result_json"):
        raise RuntimeError("Only completed scripts can generate storyboard covers.")
    if not GOOGLE_API_KEY:
        raise RuntimeError("Missing GOOGLE_API_KEY for storyboard generation.")
    output_dir = RESULTS_ROOT / item_id
    output_dir.mkdir(parents=True, exist_ok=True)
    base_script = item.get("result_json") or item.get("zh_result_json") or {}
    script_json = apply_script_edits(base_script, edits_payload or {}) if isinstance(edits_payload, dict) and edits_payload else base_script
    assets = generate_storyboard_assets(item_id, output_dir, script_json, prompt_override=prompt_override)
    state = save_storyboard_state(
        item_id,
        prompt=assets.get("prompt") or "",
        preview_name=assets.get("preview_name") or "",
        cover_name=assets.get("cover_name") or "",
        model=assets.get("image_model") or assets.get("prompt_model") or "",
    )
    preview_url = state.get("storyboard_preview_url") or f"/results/{item_id}/{assets.get('preview_name') or ''}"
    cover_url = state.get("storyboard_cover_url") or assets.get("cover_url") or preview_url
    update_job_item(
        parent_job_id,
        item_index,
        storyboard_prompt=state.get("storyboard_prompt") or assets.get("prompt") or "",
        storyboard_preview_url=preview_url,
        storyboard_cover_url=cover_url,
        storyboard_updated_at=state.get("storyboard_updated_at") or now_iso(),
    )
    with job_lock:
        refreshed = jobs[parent_job_id]["items"][item_index]
    return public_item_view(refreshed)


def generate_storyboard_prompt_for_item(item_id: str, edits_payload: dict[str, Any] | None = None) -> dict[str, Any]:
    context = find_item_context(item_id)
    if not context:
        raise RuntimeError("Script item not found.")
    parent_job_id, item_index, item = context
    if item.get("status") != "completed" or not item.get("result_json"):
        raise RuntimeError("Only completed scripts can generate storyboard prompts.")
    output_dir = RESULTS_ROOT / item_id
    output_dir.mkdir(parents=True, exist_ok=True)
    base_script = item.get("result_json") or {}
    script_json = apply_script_edits(base_script, edits_payload or {}) if isinstance(edits_payload, dict) and edits_payload else base_script
    prompt, raw, model = generate_storyboard_prompt_from_script(script_json)
    (output_dir / STORYBOARD_PROMPT_FILE).write_text(prompt, encoding="utf-8")
    if raw:
        (output_dir / "storyboard_prompt_raw_gemini.json").write_text(json.dumps(raw, ensure_ascii=False, indent=2), encoding="utf-8")
    state = save_storyboard_state(item_id, prompt=prompt, model=model)
    update_job_item(
        parent_job_id,
        item_index,
        storyboard_prompt=state.get("storyboard_prompt") or prompt,
        storyboard_updated_at=state.get("storyboard_updated_at") or now_iso(),
    )
    with job_lock:
        refreshed = jobs[parent_job_id]["items"][item_index]
    view = public_item_view(refreshed)
    view["storyboard_prompt"] = prompt
    return view


def apply_storyboard_cover_to_scripts(parent_job_id: str, item_index: int, item_id: str, cover_url: str) -> dict[str, Any]:
    output_dir = RESULTS_ROOT / item_id
    with job_lock:
        item = jobs[parent_job_id]["items"][item_index]
        display_language = item.get("display_language") or "zh"
        zh_script = json.loads(json.dumps(item.get("zh_result_json") or item.get("result_json") or {}, ensure_ascii=False))
        pt_script = json.loads(json.dumps(item.get("pt_result_json") or {}, ensure_ascii=False)) if item.get("pt_result_json") else {}
    zh_script["storyboard_cover_url"] = cover_url
    zh_variant = generate_script_variant_outputs(output_dir, item_id, zh_script, item.get("video_url") or "", locale="zh")
    update_payload: dict[str, Any] = {
        "zh_result_json": zh_variant["script_json"],
        "zh_html_url": zh_variant["html_url"],
        "zh_docx_url": zh_variant["docx_url"],
    }
    if display_language != "pt":
        update_payload.update(
            result_json=zh_variant["script_json"],
            html_url=zh_variant["html_url"],
            docx_url=zh_variant["docx_url"],
            title=zh_variant["script_json"].get("title") or item.get("title") or "",
        )
    if pt_script:
        pt_script["storyboard_cover_url"] = cover_url
        pt_variant = generate_script_variant_outputs(output_dir, item_id, pt_script, item.get("video_url") or "", locale="pt")
        update_payload.update(
            pt_result_json=pt_variant["script_json"],
            pt_html_url=pt_variant["html_url"],
            pt_docx_url=pt_variant["docx_url"],
        )
        if display_language == "pt":
            update_payload.update(
                result_json=pt_variant["script_json"],
                html_url=pt_variant["html_url"],
                docx_url=pt_variant["docx_url"],
                title=pt_variant["script_json"].get("title") or item.get("title") or "",
            )
    update_job_item(parent_job_id, item_index, **update_payload)
    with job_lock:
        refreshed = jobs[parent_job_id]["items"][item_index]
        job = jobs.get(parent_job_id)
        if job and (job.get("id") == item_id or len(job.get("items") or []) == 1):
            job["zh_result_json"] = refreshed.get("zh_result_json")
            job["zh_html_url"] = refreshed.get("zh_html_url")
            job["zh_docx_url"] = refreshed.get("zh_docx_url")
            job["pt_result_json"] = refreshed.get("pt_result_json")
            job["pt_html_url"] = refreshed.get("pt_html_url")
            job["pt_docx_url"] = refreshed.get("pt_docx_url")
            job["result_json"] = refreshed.get("result_json")
            job["html_url"] = refreshed.get("html_url")
            job["docx_url"] = refreshed.get("docx_url")
            job["title"] = refreshed.get("title") or job.get("title") or ""
            save_jobs()
    return public_item_view(refreshed)


def confirm_storyboard_cover(item_id: str) -> dict[str, Any]:
    context = find_item_context(item_id)
    if not context:
        raise RuntimeError("Script item not found.")
    parent_job_id, item_index, item = context
    preview_url = str(item.get("storyboard_preview_url") or load_storyboard_state(item_id).get("storyboard_preview_url") or "").strip()
    if not preview_url:
        raise RuntimeError("请先生成分解示意图。")
    preview_name = Path(urllib.parse.urlparse(preview_url).path).name
    preview_path = RESULTS_ROOT / item_id / preview_name
    if not preview_path.exists():
        raise RuntimeError("当前示意图文件不存在，请重新生成。")
    cover_name = STORYBOARD_COVER_BASENAME + preview_path.suffix.lower()
    cover_path = RESULTS_ROOT / item_id / cover_name
    if preview_path != cover_path:
        shutil.copyfile(preview_path, cover_path)
    state = save_storyboard_state(item_id, cover_name=cover_name)
    cover_url = state.get("storyboard_cover_url") or f"/results/{item_id}/{cover_name}"
    refreshed = apply_storyboard_cover_to_scripts(parent_job_id, item_index, item_id, cover_url)
    update_job_item(
        parent_job_id,
        item_index,
        storyboard_prompt=state.get("storyboard_prompt") or item.get("storyboard_prompt") or "",
        storyboard_preview_url=state.get("storyboard_preview_url") or preview_url,
        storyboard_cover_url=cover_url,
        storyboard_updated_at=state.get("storyboard_updated_at") or now_iso(),
    )
    with job_lock:
        final_item = jobs[parent_job_id]["items"][item_index]
    if final_item.get("saved_to_library_at") or library_entry_exists(item_id):
        persist_library_entry(parent_job_id, final_item, use_llm=False)
    return public_item_view(final_item)


def ensure_storyboard_cover_ready(item_id: str) -> dict[str, Any]:
    context = find_item_context(item_id)
    if not context:
        raise RuntimeError("Script item not found.")
    _, _, item = context
    cover_url = str(item.get("storyboard_cover_url") or load_storyboard_state(item_id).get("storyboard_cover_url") or "").strip()
    if cover_url:
        return public_item_view(item)
    generate_storyboard_preview(item_id, str(item.get("storyboard_prompt") or "").strip())
    return confirm_storyboard_cover(item_id)


def build_library_entry_payload(parent_job_id: str, item: dict[str, Any], *, use_llm: bool = True) -> dict[str, Any]:
    script = item.get("result_json") or {}
    output_dir = RESULTS_ROOT / item["id"]
    bundle = read_json(output_dir / "evidence_bundle.json")
    decision = detect_content_type_decision(
        script,
        bundle,
        existing_type=item.get("content_type") or "",
        existing_source=item.get("content_type_source") or "",
        use_llm=use_llm,
    )
    existing = library_entry_by_id(str(item.get("id") or "")) or {}
    duration_fields = creator_duration_fields(str(item.get("id") or ""), script)
    entry = {
        "entry_id": item["id"],
        "parent_job_id": parent_job_id,
        "created_at": existing.get("created_at") or item.get("completed_at") or now_iso(),
        "video_url": item.get("video_url"),
        "title": item.get("title") or script.get("title") or "Untitled Script",
        "content_type": decision["content_type"],
        "content_type_source": decision["content_type_source"],
        "content_type_reasoning": decision["content_type_reasoning"],
        "content_type_confidence": decision["content_type_confidence"],
        **duration_fields,
        "whole_video_summary": script.get("whole_video_summary") or "",
        "html_url": item.get("html_url") or "",
        "report_url": item.get("report_url") or "",
        "evidence_url": item.get("evidence_url") or "",
        "docx_url": item.get("docx_url") or "",
        "zh_docx_url": item.get("zh_docx_url") or item.get("docx_url") or "",
        "pt_docx_url": item.get("pt_docx_url") or "",
        "zh_html_url": item.get("zh_html_url") or item.get("html_url") or "",
        "pt_html_url": item.get("pt_html_url") or "",
        "preview_image_url": library_preview_image_url(item["id"], script, output_dir),
        "source": "edited" if item.get("edited") else "ai",
        "saved_at": item.get("saved_to_library_at") or now_iso(),
        "display_language": item.get("display_language") or existing.get("display_language") or "zh",
        "chat_messages": item.get("chat_messages") if isinstance(item.get("chat_messages"), list) else existing.get("chat_messages") or [],
        "reviewed": bool(item.get("reviewed")),
        "edited": bool(item.get("edited")),
        "storyboard_prompt": item.get("storyboard_prompt") or existing.get("storyboard_prompt") or "",
        "storyboard_preview_url": item.get("storyboard_preview_url") or existing.get("storyboard_preview_url") or "",
        "storyboard_cover_url": item.get("storyboard_cover_url") or existing.get("storyboard_cover_url") or "",
        "storyboard_updated_at": item.get("storyboard_updated_at") or existing.get("storyboard_updated_at") or "",
    }
    entry.update(infer_location_tag_fields(entry, script))
    return entry


def sync_library_entry_from_item(parent_job_id: str, item: dict[str, Any], *, use_llm: bool = True, delete_source: bool = False) -> bool:
    entry = build_library_entry_payload(parent_job_id, item, use_llm=use_llm)
    saved = append_library_entry(entry)
    if saved and delete_source:
        delete_source_video_if_allowed(item["id"], reason="saved_to_library")
    return saved


def persist_library_entry(parent_job_id: str, item: dict[str, Any], *, use_llm: bool = True) -> bool:
    should_delete_source = library_entry_by_id(str(item.get("id") or "")) is None
    return sync_library_entry_from_item(parent_job_id, item, use_llm=use_llm, delete_source=should_delete_source)


def persist_completed_job_items_async(job_id: str, *, use_llm: bool = True) -> None:
    def _worker() -> None:
        try:
            with job_lock:
                job = json.loads(json.dumps(jobs.get(job_id) or {}, ensure_ascii=False))
            items = job.get("items") or []
            for index, item in enumerate(items):
                if str(item.get("status") or "").strip() != "completed":
                    continue
                if item.get("saved_to_library_at"):
                    continue
                if persist_library_entry(job_id, item, use_llm=use_llm):
                    update_job_item(job_id, index, saved_to_library_at=now_iso())
        except Exception:
            return

    threading.Thread(target=_worker, daemon=True).start()


def find_item_context(item_id: str) -> tuple[str, int, dict[str, Any]] | None:
    with job_lock:
        for job_id, job in jobs.items():
            for index, item in enumerate(job.get("items") or []):
                if item.get("id") == item_id:
                    return job_id, index, json.loads(json.dumps(item, ensure_ascii=False))
            if job_id == item_id and job.get("result_json"):
                pseudo_item = {
                    "id": job_id,
                    "index": 0,
                    "video_url": job.get("video_url") or "",
                    "status": job.get("status"),
                    "stage": job.get("stage") or "",
                    "stage_message": job.get("stage_message") or "",
                    "html_url": job.get("html_url") or "",
                    "report_url": job.get("report_url") or "",
                    "evidence_url": job.get("evidence_url") or "",
                    "docx_url": job.get("docx_url") or "",
                    "zh_docx_url": job.get("zh_docx_url") or job.get("docx_url") or "",
                    "pt_docx_url": job.get("pt_docx_url") or "",
                    "zh_html_url": job.get("zh_html_url") or job.get("html_url") or "",
                    "pt_html_url": job.get("pt_html_url") or "",
                    "artifacts": job.get("artifacts") or {},
                    "error": job.get("error") or "",
                    "result_json": json.loads(json.dumps(job.get("result_json") or {}, ensure_ascii=False)),
                    "zh_result_json": json.loads(json.dumps(job.get("zh_result_json") or job.get("result_json") or {}, ensure_ascii=False)),
                    "pt_result_json": json.loads(json.dumps(job.get("pt_result_json") or {}, ensure_ascii=False)),
                    "content_type": job.get("content_type") or "",
                    "content_type_source": job.get("content_type_source") or "auto",
                    "content_type_reasoning": job.get("content_type_reasoning") or "",
                    "content_type_confidence": job.get("content_type_confidence") or "",
                    "title": job.get("title") or "",
                    "display_language": job.get("display_language") or "zh",
                    "review_status": job.get("review_status") or "",
                    "review_message": job.get("review_message") or "",
                    "review_feedback": job.get("review_feedback") or "",
                    "review_mode": normalize_review_mode(job.get("review_mode")),
                    "reviewed": bool(job.get("reviewed")),
                    "edited": bool(job.get("edited")),
                    "original_result_json": json.loads(json.dumps(job.get("original_result_json") or {}, ensure_ascii=False)),
                    "updated_at": job.get("updated_at"),
                }
                return job_id, 0, pseudo_item
    return None


def library_edit_job_id(entry_id: str) -> str:
    return f"library_edit_{entry_id}"


def ensure_library_edit_context(entry_id: str) -> tuple[str, int, dict[str, Any]]:
    entry_id = str(entry_id or "").strip()
    if not re.fullmatch(r"[0-9a-f]{32}", entry_id):
        raise RuntimeError("Invalid library script id.")
    entry = library_entry_by_id(entry_id)
    if not entry:
        raise RuntimeError("Library script not found.")
    output_dir = RESULTS_ROOT / entry_id
    zh_script = read_json(output_dir / "script_table.json") or read_json(output_dir / "analysis_result.json") or {}
    pt_script = read_json(output_dir / "script_table_pt.json") or {}
    if not zh_script and not pt_script:
        raise RuntimeError("No editable script JSON found for this library entry.")
    display_language = infer_library_display_language(entry, zh_script, pt_script)
    current_script = pt_script if display_language == "pt" and pt_script else zh_script or pt_script
    item_payload = {
        "id": entry_id,
        "index": 0,
        "video_url": entry.get("video_url") or "",
        "status": "completed",
        "updated_at": entry.get("saved_at") or entry.get("created_at") or now_iso(),
        "completed_at": entry.get("created_at") or entry.get("saved_at") or now_iso(),
        "stage": "completed",
        "stage_message": "Library script ready for editing.",
        "html_url": entry.get("pt_html_url") if display_language == "pt" and entry.get("pt_html_url") else entry.get("zh_html_url") or entry.get("html_url") or "",
        "docx_url": entry.get("pt_docx_url") if display_language == "pt" and entry.get("pt_docx_url") else entry.get("zh_docx_url") or entry.get("docx_url") or "",
        "zh_docx_url": entry.get("zh_docx_url") or entry.get("docx_url") or "",
        "pt_docx_url": entry.get("pt_docx_url") or "",
        "zh_html_url": entry.get("zh_html_url") or entry.get("html_url") or "",
        "pt_html_url": entry.get("pt_html_url") or "",
        "result_json": current_script,
        "zh_result_json": zh_script or current_script,
        "pt_result_json": pt_script or None,
        "content_type": entry.get("content_type") or DEFAULT_CONTENT_TYPE,
        "content_type_source": entry.get("content_type_source") or "auto",
        "content_type_reasoning": entry.get("content_type_reasoning") or "",
        "content_type_confidence": entry.get("content_type_confidence") or "",
        "title": entry.get("title") or current_script.get("title") or "Untitled Script",
        "display_language": display_language,
        "review_status": "",
        "review_stage": "",
        "review_message": "",
        "review_feedback": "",
        "review_mode": REVIEW_MODE_PARTIAL,
        "chat_messages": entry.get("chat_messages") if isinstance(entry.get("chat_messages"), list) else [],
        "reviewed": bool(entry.get("reviewed")),
        "edited": bool(entry.get("edited")),
        "saved_to_library_at": entry.get("saved_at") or entry.get("created_at") or now_iso(),
        "storyboard_prompt": entry.get("storyboard_prompt") or "",
        "storyboard_preview_url": entry.get("storyboard_preview_url") or "",
        "storyboard_cover_url": entry.get("storyboard_cover_url") or "",
        "storyboard_updated_at": entry.get("storyboard_updated_at") or "",
        "source": entry.get("source") or "library",
    }
    existing = find_item_context(entry_id)
    with job_lock:
        if existing:
            parent_job_id, item_index, _ = existing
            job = jobs[parent_job_id]
            if job.get("items"):
                job["items"][item_index].update(json.loads(json.dumps(item_payload, ensure_ascii=False)))
                job["updated_at"] = now_iso()
                save_jobs()
                return parent_job_id, item_index, jobs[parent_job_id]["items"][item_index]
        job_id = library_edit_job_id(entry_id)
        job = jobs.get(job_id)
        if not isinstance(job, dict):
            job = {
                "id": job_id,
                "video_url": entry.get("video_url") or "",
                "status": "completed",
                "created_at": entry.get("created_at") or now_iso(),
                "updated_at": now_iso(),
                "started_at": entry.get("created_at") or now_iso(),
                "completed_at": entry.get("saved_at") or entry.get("created_at") or now_iso(),
                "total_items": 1,
                "completed_items": 1,
                "failed_items": 0,
                "items": [],
            }
            jobs[job_id] = job
        job["status"] = "completed"
        job["updated_at"] = now_iso()
        job["completed_at"] = entry.get("saved_at") or entry.get("created_at") or now_iso()
        job["items"] = [json.loads(json.dumps(item_payload, ensure_ascii=False))]
        save_jobs()
        return job_id, 0, jobs[job_id]["items"][0]


def public_library_workbench(entry_id: str) -> dict[str, Any]:
    parent_job_id, item_index, _ = ensure_library_edit_context(entry_id)
    with job_lock:
        job = jobs.get(parent_job_id) or {}
        item = (job.get("items") or [])[item_index]
        return {
            "id": parent_job_id,
            "status": "completed",
            "video_url": item.get("video_url") or "",
            "updated_at": job.get("updated_at") or item.get("updated_at") or now_iso(),
            "completed_at": job.get("completed_at") or item.get("completed_at") or "",
            "total_items": 1,
            "completed_items": 1,
            "failed_items": 0,
            "items": [public_item_view(item)],
        }


def apply_script_edits(script: dict[str, Any], payload: dict[str, Any]) -> dict[str, Any]:
    edited = json.loads(json.dumps(script or {}, ensure_ascii=False))
    edited["title"] = fill_text(payload.get("title") or edited.get("title") or "Video Script", "Video Script")
    edited["whole_video_summary"] = fill_text(payload.get("whole_video_summary") or edited.get("whole_video_summary") or "", "无")
    mechanism = dict(edited.get("mechanism") or {})
    mechanism["reason"] = fill_text(payload.get("mechanism_reason") or mechanism.get("reason") or "", "无")
    edited["mechanism"] = mechanism
    incoming_points = payload.get("core_viral_points")
    if isinstance(incoming_points, list):
        edited["core_viral_points"] = [
            {
                "label": fill_text(point.get("label"), "要点"),
                "text": fill_text(point.get("text"), "无"),
            }
            for point in incoming_points
            if isinstance(point, dict)
        ]
    incoming_replaceable = payload.get("replaceable_parts")
    if isinstance(incoming_replaceable, list):
        edited["replaceable_parts"] = [
            {
                "label": fill_text(point.get("label"), "可替换项"),
                "text": fill_text(point.get("text"), "无"),
            }
            for point in incoming_replaceable
            if isinstance(point, dict)
        ]
    rows = choose_script_rows(edited)
    incoming_rows = payload.get("rows")
    if isinstance(incoming_rows, list):
        rebuilt_rows: list[dict[str, Any]] = []
        for idx, incoming in enumerate(incoming_rows):
            if not isinstance(incoming, dict):
                continue
            original_index_raw = incoming.get("original_index", idx)
            try:
                original_index = int(original_index_raw)
            except (TypeError, ValueError):
                original_index = idx
            if not (0 <= original_index < len(rows)):
                continue
            row = json.loads(json.dumps(rows[original_index], ensure_ascii=False))
            for key in ["time", "visual_content", "action", "dialogue_or_audio", "integrated_summary"]:
                if key in incoming:
                    row[key] = fill_text(incoming.get(key), "无")
            rebuilt_rows.append(row)
        rows = rebuilt_rows
    if isinstance(edited.get("rows"), list):
        edited["rows"] = rows
    elif isinstance(edited.get("synthesized_segments"), list):
        edited["synthesized_segments"] = rows
    return edited


def refresh_replaceable_parts(script_json: dict[str, Any]) -> dict[str, Any]:
    refreshed = json.loads(json.dumps(script_json or {}, ensure_ascii=False))
    existing = normalize_replaceable_parts(refreshed.get("replaceable_parts"))
    if not GOOGLE_API_KEY or run_text_json_prompt_with_fallback is None:
        refreshed["replaceable_parts"] = existing
        return refreshed
    payload = {
        "title": refreshed.get("title") or "",
        "whole_video_summary": refreshed.get("whole_video_summary") or "",
        "mechanism": refreshed.get("mechanism") or {},
        "rows": choose_script_rows(refreshed),
        "existing_replaceable_parts": existing,
    }
    try:
        result, _, _ = run_text_json_prompt_with_fallback(
            payload,
            GOOGLE_API_KEY,
            unique_models(*MODEL_CANDIDATES, *PRIMARY_FALLBACK_MODELS, *SUPPLEMENT_FALLBACK_MODELS),
            REPLACEMENT_PLAN_REFRESH_PROMPT,
            "replacement plan refresh",
        )
        refreshed["replaceable_parts"] = normalize_replaceable_parts(
            (result or {}).get("replaceable_parts"),
            fallback=existing,
        )
    except Exception:
        refreshed["replaceable_parts"] = existing
    return refreshed


def refresh_core_viral_points(script_json: dict[str, Any]) -> dict[str, Any]:
    refreshed = json.loads(json.dumps(script_json or {}, ensure_ascii=False))
    existing = normalize_core_viral_points(refreshed.get("core_viral_points"))
    if not GOOGLE_API_KEY or run_text_json_prompt_with_fallback is None:
        refreshed["core_viral_points"] = existing
        return refreshed
    payload = {
        "title": refreshed.get("title") or "",
        "whole_video_summary": refreshed.get("whole_video_summary") or "",
        "mechanism": refreshed.get("mechanism") or {},
        "rows": choose_script_rows(refreshed),
        "existing_core_viral_points": existing,
    }
    try:
        result, _, _ = run_text_json_prompt_with_fallback(
            payload,
            GOOGLE_API_KEY,
            unique_models(*MODEL_CANDIDATES, *PRIMARY_FALLBACK_MODELS, *SUPPLEMENT_FALLBACK_MODELS),
            CORE_VIRAL_POINTS_REFRESH_PROMPT,
            "core viral points refresh",
        )
        refreshed["core_viral_points"] = normalize_core_viral_points(
            (result or {}).get("core_viral_points"),
            fallback=existing,
        )
    except Exception:
        refreshed["core_viral_points"] = existing
    return refreshed


def regenerate_item_outputs(
    parent_job_id: str,
    item_index: int,
    item_id: str,
    video_url: str,
    script_json: dict[str, Any],
    persist_library: bool = False,
    target_language: str = "zh",
) -> dict[str, Any]:
    output_dir = RESULTS_ROOT / item_id
    output_dir.mkdir(parents=True, exist_ok=True)
    if target_language == "pt":
        source_script = json.loads(json.dumps(script_json or {}, ensure_ascii=False))
        if str(source_script.get("display_language") or "").strip().lower() == "pt":
            pt_script = source_script
        else:
            pt_script = translate_script_to_portuguese(
                source_script,
                GOOGLE_API_KEY,
                unique_models(*STABLE_VIDEO_MODELS, *MODEL_CANDIDATES, *PRIMARY_FALLBACK_MODELS, *SUPPLEMENT_FALLBACK_MODELS),
            )
        for optional_list_field in ("core_viral_points", "replaceable_parts"):
            if isinstance(source_script.get(optional_list_field), list) and not source_script.get(optional_list_field):
                pt_script[optional_list_field] = []
        pt_variant = generate_script_variant_outputs(output_dir, item_id, pt_script, video_url, locale="pt")
        with job_lock:
            existing_item = jobs[parent_job_id]["items"][item_index]
            zh_result_json = existing_item.get("zh_result_json") or existing_item.get("result_json") or {}
            zh_html_url = existing_item.get("zh_html_url") or existing_item.get("html_url") or f"/results/{item_id}/script_table.html"
            zh_docx_url = existing_item.get("zh_docx_url") or existing_item.get("docx_url") or ""
        update_payload = {
            "pt_result_json": pt_variant["script_json"],
            "pt_html_url": pt_variant["html_url"],
            "pt_docx_url": pt_variant["docx_url"],
            "result_json": pt_variant["script_json"],
            "html_url": pt_variant["html_url"],
            "docx_url": pt_variant["docx_url"],
            "zh_result_json": zh_result_json,
            "zh_html_url": zh_html_url,
            "zh_docx_url": zh_docx_url,
            "display_language": "pt",
            "title": pt_variant["script_json"].get("title") or "Roteiro do vídeo",
            "updated_at": now_iso(),
        }
        if persist_library:
            update_payload["saved_to_library_at"] = now_iso()
        update_job_item(parent_job_id, item_index, **update_payload)
        with job_lock:
            job = jobs.get(parent_job_id)
            if job and (job.get("id") == item_id or len(job.get("items") or []) == 1):
                job.update(update_payload)
                save_jobs()
            item = jobs[parent_job_id]["items"][item_index]
        if persist_library:
            persist_library_entry(parent_job_id, item, use_llm=True)
        elif library_entry_exists(item_id):
            sync_library_entry_from_item(parent_job_id, item, use_llm=False, delete_source=False)
        return public_item_view(item)

    script_json = refresh_core_viral_points(
        json.loads(json.dumps(script_json or {}, ensure_ascii=False))
    )
    script_json = refresh_replaceable_parts(script_json)
    script_json = enforce_chinese_dialogue_translation(
        script_json,
        GOOGLE_API_KEY,
        unique_models(*MODEL_CANDIDATES),
    )
    zh_variant = generate_script_variant_outputs(output_dir, item_id, script_json, video_url, locale="zh")
    existing_item = jobs[parent_job_id]["items"][item_index]
    decision = detect_content_type_decision_for_output(
        output_dir,
        script_json,
        read_json(output_dir / "evidence_bundle.json"),
        existing_type=existing_item.get("content_type") or "",
        existing_source=existing_item.get("content_type_source") or "",
    )
    content_type = decision["content_type"]

    update_payload = {
        "result_json": zh_variant["script_json"],
        "zh_result_json": zh_variant["script_json"],
        "pt_result_json": None,
        "original_result_json": existing_item.get("original_result_json") or script_json,
        "title": script_json.get("title") or "Video Script",
        "content_type": content_type,
        "content_type_source": decision["content_type_source"],
        "content_type_reasoning": decision["content_type_reasoning"],
        "content_type_confidence": decision["content_type_confidence"],
        "docx_url": zh_variant["docx_url"],
        "zh_docx_url": zh_variant["docx_url"],
        "pt_docx_url": "",
        "html_url": zh_variant["html_url"],
        "zh_html_url": zh_variant["html_url"],
        "pt_html_url": "",
        "display_language": "zh",
        "artifacts": summarize_artifacts(item_id, output_dir),
        "edited": True,
        "updated_at": now_iso(),
    }
    if persist_library:
        update_payload["saved_to_library_at"] = now_iso()
    update_job_item(parent_job_id, item_index, **update_payload)
    with job_lock:
        job = jobs.get(parent_job_id)
        if job and (job.get("id") == item_id or len(job.get("items") or []) == 1):
            job["result_json"] = zh_variant["script_json"]
            job["zh_result_json"] = zh_variant["script_json"]
            job["pt_result_json"] = None
            job["original_result_json"] = job.get("original_result_json") or script_json
            job["title"] = script_json.get("title") or "Video Script"
            job["content_type"] = content_type
            job["content_type_source"] = decision["content_type_source"]
            job["content_type_reasoning"] = decision["content_type_reasoning"]
            job["content_type_confidence"] = decision["content_type_confidence"]
            job["docx_url"] = zh_variant["docx_url"]
            job["zh_docx_url"] = zh_variant["docx_url"]
            job["pt_docx_url"] = ""
            job["html_url"] = zh_variant["html_url"]
            job["zh_html_url"] = zh_variant["html_url"]
            job["pt_html_url"] = ""
            job["display_language"] = "zh"
            job["artifacts"] = summarize_artifacts(item_id, output_dir)
            job["updated_at"] = now_iso()
            save_jobs()
        item = jobs[parent_job_id]["items"][item_index]
    if persist_library:
        persist_library_entry(parent_job_id, item)
    elif library_entry_exists(item_id):
        sync_library_entry_from_item(parent_job_id, item, use_llm=False, delete_source=False)
    return public_item_view(item)


def save_item_edits_to_library(
    parent_job_id: str,
    item_index: int,
    item_id: str,
    updated_script: dict[str, Any],
    *,
    target_language: str,
) -> dict[str, Any]:
    target_language = "pt" if str(target_language or "").strip().lower() == "pt" else "zh"
    with job_lock:
        current_item = json.loads(json.dumps(jobs[parent_job_id]["items"][item_index], ensure_ascii=False))
    video_url = current_item.get("video_url") or ""
    if target_language == "pt":
        output_dir = RESULTS_ROOT / item_id
        output_dir.mkdir(parents=True, exist_ok=True)
        pt_variant = generate_script_variant_outputs(output_dir, item_id, updated_script, video_url, locale="pt")
        with job_lock:
            existing_item = jobs[parent_job_id]["items"][item_index]
            zh_result_json = existing_item.get("zh_result_json") or {}
            zh_html_url = existing_item.get("zh_html_url") or ""
            zh_docx_url = existing_item.get("zh_docx_url") or ""
        update_payload = {
            "pt_result_json": pt_variant["script_json"],
            "pt_html_url": pt_variant["html_url"],
            "pt_docx_url": pt_variant["docx_url"],
            "result_json": pt_variant["script_json"],
            "html_url": pt_variant["html_url"],
            "docx_url": pt_variant["docx_url"],
            "zh_result_json": zh_result_json,
            "zh_html_url": zh_html_url,
            "zh_docx_url": zh_docx_url,
            "display_language": "pt",
            "title": pt_variant["script_json"].get("title") or "Roteiro do vídeo",
            "edited": True,
            "saved_to_library_at": now_iso(),
            "updated_at": now_iso(),
        }
        update_job_item(parent_job_id, item_index, **update_payload)
        with job_lock:
            job = jobs.get(parent_job_id)
            if job and (job.get("id") == item_id or len(job.get("items") or []) == 1):
                job.update(update_payload)
                save_jobs()
            item = jobs[parent_job_id]["items"][item_index]
        persist_library_entry(parent_job_id, item, use_llm=True)
        return public_item_view(item)

    regenerate_item_outputs(
        parent_job_id,
        item_index,
        item_id,
        video_url,
        updated_script,
        persist_library=False,
        target_language="zh",
    )
    refreshed_context = find_item_context(item_id)
    if not refreshed_context:
        raise RuntimeError("Script item not found.")
    parent_job_id, item_index, item = refreshed_context
    return regenerate_item_outputs(
        parent_job_id,
        item_index,
        item_id,
        item.get("video_url") or video_url,
        item.get("zh_result_json") or item.get("result_json") or {},
        persist_library=True,
        target_language="pt",
    )


def set_item_display_language(item_id: str, language: str, edits_payload: dict[str, Any] | None = None) -> dict[str, Any]:
    context = find_item_context(item_id)
    if not context:
        raise RuntimeError("Script item not found.")
    parent_job_id, item_index, item = context
    if item.get("status") != "completed":
        raise RuntimeError("Only completed scripts can switch languages.")
    if language not in {"zh", "pt"}:
        raise RuntimeError("Unsupported language.")
    output_dir = RESULTS_ROOT / item_id
    has_edits = isinstance(edits_payload, dict) and any(
        key in edits_payload
        for key in ["title", "whole_video_summary", "mechanism_reason", "core_viral_points", "replaceable_parts", "rows"]
    )
    if has_edits and language == "pt":
        current_script = item.get("result_json") or {}
        updated_script = apply_script_edits(current_script, edits_payload or {})
        if str(item.get("display_language") or "").strip().lower() == "pt":
            updated_script["display_language"] = "pt"
        return regenerate_item_outputs(
            parent_job_id,
            item_index,
            item_id,
            item.get("video_url") or "",
            updated_script,
            persist_library=False,
            target_language="pt",
        )
    if language == "pt":
        pt_script = item.get("pt_result_json") or {}
        if not pt_script:
            base_script = item.get("zh_result_json") or item.get("result_json") or {}
            pt_item = regenerate_item_outputs(
                parent_job_id,
                item_index,
                item_id,
                item.get("video_url") or "",
                base_script,
                persist_library=False,
                target_language="pt",
            )
            return pt_item
        update_job_item(
            parent_job_id,
            item_index,
            result_json=pt_script,
            html_url=item.get("pt_html_url") or f"/results/{item_id}/script_table_pt.html",
            docx_url=item.get("pt_docx_url") or "",
            display_language="pt",
            title=pt_script.get("title") or item.get("title") or "",
        )
    else:
        zh_script = item.get("zh_result_json") or item.get("result_json") or {}
        update_job_item(
            parent_job_id,
            item_index,
            result_json=zh_script,
            html_url=item.get("zh_html_url") or f"/results/{item_id}/script_table.html",
            docx_url=item.get("zh_docx_url") or "",
            display_language="zh",
            title=zh_script.get("title") or item.get("title") or "",
        )
    with job_lock:
        job = jobs.get(parent_job_id)
        item_ref = jobs[parent_job_id]["items"][item_index]
        if job and (job.get("id") == item_id or len(job.get("items") or []) == 1):
            job["result_json"] = item_ref.get("result_json")
            job["html_url"] = item_ref.get("html_url")
            job["docx_url"] = item_ref.get("docx_url")
            job["display_language"] = item_ref.get("display_language") or language
            job["title"] = item_ref.get("title") or job.get("title") or ""
            save_jobs()
    if library_entry_exists(item_id):
        sync_library_entry_from_item(parent_job_id, item_ref, use_llm=False, delete_source=False)
    return public_item_view(item_ref)


def run_review_reanalysis(parent_job_id: str, item_index: int, item_id: str, feedback: str, review_mode: str = REVIEW_MODE_PARTIAL) -> None:
    mode = normalize_review_mode(review_mode)
    clear_item_cancelled(item_id)
    if not GOOGLE_API_KEY:
        update_job_item(parent_job_id, item_index, review_status="failed", review_message="Missing GOOGLE_API_KEY for review.")
        return
    if run_text_json_prompt_with_fallback is None:
        update_job_item(parent_job_id, item_index, review_status="failed", review_message="Review pipeline helpers are unavailable.")
        return

    output_dir = RESULTS_ROOT / item_id
    source_video = output_dir / "source.mp4"
    script_json_path = output_dir / "script_table.json"
    primary_path = output_dir / "primary_v2_draft.json"
    supplement_path = output_dir / "supplement_evidence.json"
    audio_multiview_path = output_dir / "audio_multiview.json"
    type_router_path = output_dir / "type_router.json"
    review_request_path = output_dir / "review_request.json"
    review_plan_path = output_dir / "review_plan.json"
    review_plan_raw_path = output_dir / "review_plan_raw_gemini.json"
    review_video_path = output_dir / "review_video_recheck.json"
    review_video_raw_path = output_dir / "review_video_recheck_raw_gemini.json"
    review_refine_raw_path = output_dir / "review_refine_raw_gemini.json"
    original_backup_path = output_dir / "script_table.original.json"

    try:
        current_script = read_json(script_json_path)
        if not current_script:
            raise RuntimeError("No existing script result to review.")
        stored_original = jobs[parent_job_id]["items"][item_index].get("original_result_json") or {}
        original_script = stored_original or read_json(original_backup_path) or current_script
        if not original_backup_path.exists():
            original_backup_path.write_text(json.dumps(original_script, ensure_ascii=False, indent=2), encoding="utf-8")

        update_job_item(
            parent_job_id,
            item_index,
            review_status="running",
            review_stage="plan",
            review_message="Reviewing your feedback.",
            review_feedback=feedback,
            review_mode=mode,
            original_result_json=original_script,
        )

        review_request = {
            "feedback": feedback,
            "review_mode": mode,
            "video_url": jobs[parent_job_id]["items"][item_index].get("video_url") or "",
            "original_script": original_script,
            "current_script": current_script,
            "primary_draft": read_json(primary_path),
            "supplement": read_json(supplement_path),
            "audio_multiview": read_json(audio_multiview_path),
            "type_router": read_json(type_router_path),
            "source_metadata": read_json(output_dir / "source_metadata.json"),
        }
        review_request_path.write_text(json.dumps(review_request, ensure_ascii=False, indent=2), encoding="utf-8")
        if is_item_cancelled(item_id):
            raise RuntimeError("任务已手动停止。")

        update_job_item(parent_job_id, item_index, review_stage="plan", review_message="Comparing your feedback against the prior analysis.")
        if mode == REVIEW_MODE_PARTIAL:
            review_plan = {
                "problem_summary": feedback,
                "likely_failure_layer": "final_refine",
                "needs_video_recheck": False,
                "needs_structural_rewrite": True,
                "reasoning": "Partial review mode revises the current script directly from human feedback.",
                "focus_windows": [],
                "focus_entities": [],
                "correction_goal": feedback,
                "confidence": "medium",
            }
            review_plan_raw = {"skipped": True, "mode": mode, "reason": "partial review uses direct script revision"}
        else:
            review_plan, review_plan_raw, _ = run_text_json_prompt_with_fallback(
                review_request,
                GOOGLE_API_KEY,
                unique_models(*MODEL_CANDIDATES),
                REVIEW_PLAN_PROMPT,
                "review plan",
            )
        review_plan_path.write_text(json.dumps(review_plan, ensure_ascii=False, indent=2), encoding="utf-8")
        review_plan_raw_path.write_text(json.dumps(review_plan_raw, ensure_ascii=False, indent=2), encoding="utf-8")
        if is_item_cancelled(item_id):
            raise RuntimeError("任务已手动停止。")

        review_video = {}
        should_recheck, recheck_reason = (False, "Partial review mode updates the script without video recheck.")
        if mode == REVIEW_MODE_FULL:
            should_recheck, recheck_reason = (True, "Full review mode requires video recheck.")
            if not source_video.exists():
                raise RuntimeError("完全错误需要重新回看原视频，但当前源视频已被清理；请重新提交这个视频链接后再做完全重做。")
        if should_recheck and source_video.exists():
            update_job_item(parent_job_id, item_index, review_stage="recheck", review_message="Rechecking the video with your correction in mind.")
            prompt = REVIEW_VIDEO_PROMPT + "\n\nHuman feedback:\n" + feedback + "\n\nReview plan:\n" + json.dumps(review_plan, ensure_ascii=False)
            review_video, review_video_raw, _ = run_video_json_prompt_with_fallback(
                source_video,
                GOOGLE_API_KEY,
                unique_models(MODEL_CANDIDATES[0], *SUPPLEMENT_FALLBACK_MODELS),
                prompt,
                "review video recheck",
            )
            review_video_path.write_text(json.dumps(review_video, ensure_ascii=False, indent=2), encoding="utf-8")
            review_video_raw_path.write_text(json.dumps(review_video_raw, ensure_ascii=False, indent=2), encoding="utf-8")
        else:
            review_video = {
                "skipped": True,
                "reason": recheck_reason if source_video.exists() else "Source video is unavailable for recheck.",
            }
            review_video_path.write_text(json.dumps(review_video, ensure_ascii=False, indent=2), encoding="utf-8")
        if is_item_cancelled(item_id):
            raise RuntimeError("任务已手动停止。")

        update_job_item(parent_job_id, item_index, review_stage="rebuild", review_message="Rebuilding the script from the reviewed evidence.")
        refine_payload = {
            "feedback": feedback,
            "review_mode": mode,
            "review_plan": review_plan,
            "review_video_recheck": review_video,
            "original_script": original_script,
            "current_script": current_script,
            "primary_draft": read_json(primary_path),
            "supplement": read_json(supplement_path),
            "audio_multiview": read_json(audio_multiview_path),
            "type_router": read_json(type_router_path),
        }
        corrected_script, corrected_raw, _ = run_text_json_prompt_with_fallback(
            refine_payload,
            GOOGLE_API_KEY,
            unique_models(*MODEL_CANDIDATES),
            FULL_REVIEW_REFINE_PROMPT if mode == REVIEW_MODE_FULL else PARTIAL_REVIEW_REFINE_PROMPT,
            "review refine",
        )
        review_refine_raw_path.write_text(json.dumps(corrected_raw, ensure_ascii=False, indent=2), encoding="utf-8")
        if is_item_cancelled(item_id):
            raise RuntimeError("任务已手动停止。")

        corrected_payload = extract_review_script_payload(corrected_script)
        if not corrected_payload:
            raise RuntimeError("复盘重做没有返回可用的完整脚本 JSON。")
        merged_script = json.loads(json.dumps(current_script, ensure_ascii=False))
        for key in REVIEW_SCRIPT_KEYS:
            if corrected_payload.get(key):
                merged_script[key] = corrected_payload.get(key)
        if not review_script_changed(current_script, merged_script, REVIEW_SCRIPT_KEYS):
            raise RuntimeError("复盘重做没有生成任何脚本变更，请补充更具体的错误点。")
        if not review_script_changed(current_script, merged_script, REVIEW_CORE_CHANGE_KEYS):
            raise RuntimeError("复盘重做只产生了标题或轻微变化，没有改动概述、替换方案、分镜或机制。")
        merged_script = enforce_chinese_dialogue_translation(
            merged_script,
            GOOGLE_API_KEY,
            unique_models(*MODEL_CANDIDATES),
        )
        updated_item = regenerate_item_outputs(
            parent_job_id,
            item_index,
            item_id,
            jobs[parent_job_id]["items"][item_index].get("video_url") or "",
            merged_script,
            persist_library=False,
        )
        update_job_item(
            parent_job_id,
            item_index,
            review_status="completed",
            review_stage="completed",
            review_message="Reviewed version is ready.",
            reviewed=True,
            review_feedback=feedback,
            review_mode=mode,
        )
        with job_lock:
            final_item = jobs[parent_job_id]["items"][item_index]
        if library_entry_exists(item_id):
            sync_library_entry_from_item(parent_job_id, final_item, use_llm=False, delete_source=False)
        def _record_error_case() -> None:
            try:
                entry = summarize_error_case_with_llm(
                    output_dir=output_dir,
                    parent_job_id=parent_job_id,
                    item_index=item_index,
                    item_id=item_id,
                    feedback=feedback,
                    original_script=original_script,
                    corrected_script=merged_script,
                    review_plan=review_plan,
                    review_video=review_video,
                )
                record_error_case_library_entry(entry)
            except Exception:
                return

        threading.Thread(
            target=_record_error_case,
            name=f"koko-error-case-{item_id[:8]}",
            daemon=True,
        ).start()
        return updated_item
    except Exception as exc:
        update_job_item(parent_job_id, item_index, review_status="failed", review_stage="failed", review_message=friendly_error(str(exc)), review_feedback=feedback, review_mode=mode)


def run_review_with_slot(parent_job_id: str, item_index: int, item_id: str, feedback: str, review_mode: str = REVIEW_MODE_PARTIAL) -> None:
    with analysis_slots:
        run_review_reanalysis(parent_job_id, item_index, item_id, feedback, review_mode)


def execute_single_pipeline(parent_job_id: str, item_index: int, item: dict[str, Any]) -> None:
    output_dir = RESULTS_ROOT / item["id"]
    progress_path = output_dir / "progress.json"
    proc_env = os.environ.copy()
    last_error = "Unknown pipeline failure"
    tried: list[str] = []
    manual_stop = False
    clear_item_cancelled(item["id"])
    update_job_item(parent_job_id, item_index, status="running", started_at=now_iso(), command=[], stage="queued", stage_message="Queued for analysis.")
    for model_name in dict.fromkeys(MODEL_CANDIDATES):
        cmd = [
            os.environ.get("PYTHON_BIN", "python3"),
            str(AUTO_ANALYZE),
            item["video_url"],
            "--out",
            str(output_dir),
            "--model",
            model_name,
        ]
        analysis_prompt = sanitize_analysis_prompt(item.get("user_prompt") or "")
        if analysis_prompt:
            cmd.extend(["--user-prompt", analysis_prompt])
        tried.append(model_name)
        update_job_item(
            parent_job_id,
            item_index,
            command=cmd,
            model=model_name,
            frames_enabled=bool(shutil.which("ffmpeg")),
            stage="starting",
            stage_message=f"Preparing {model_name}.",
        )
        try:
            proc = subprocess.Popen(
                cmd,
                text=True,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                env=proc_env,
            )
            register_active_process(item["id"], proc)
            start_time = time.time()
            stdout_lines: list[str] = []
            progress_mtime = 0.0
            current_stage = "starting"
            stage_started_at = start_time
            while proc.poll() is None:
                if is_item_cancelled(item["id"]):
                    manual_stop = True
                    proc.kill()
                    stdout_text, stderr_text = proc.communicate()
                    if stdout_text:
                        stdout_lines.append(stdout_text)
                    last_error = "任务已手动停止。"
                    break
                if progress_path.exists():
                    stat = progress_path.stat()
                    if stat.st_mtime > progress_mtime:
                        progress_mtime = stat.st_mtime
                        progress = read_json(progress_path)
                        if progress:
                            next_stage = progress.get("stage") or "running"
                            if next_stage != current_stage:
                                current_stage = next_stage
                                stage_started_at = time.time()
                            update_job_item(
                                parent_job_id,
                                item_index,
                                stage=next_stage,
                                stage_message=progress.get("message") or "Running analysis.",
                            )
                stage_limit = STAGE_TIMEOUTS_SEC.get(current_stage)
                if stage_limit and time.time() - stage_started_at > stage_limit:
                    proc.kill()
                    stdout_text, stderr_text = proc.communicate()
                    if stdout_text:
                        stdout_lines.append(stdout_text)
                    last_error = f"Stage `{current_stage}` exceeded {stage_limit} seconds and was stopped."
                    break
                if time.time() - start_time > PIPELINE_TIMEOUT_SEC:
                    proc.kill()
                    stdout_text, stderr_text = proc.communicate()
                    if stdout_text:
                        stdout_lines.append(stdout_text)
                    last_error = f"Analysis exceeded {PIPELINE_TIMEOUT_SEC // 60} minutes and was stopped."
                    break
                time.sleep(1)
            else:
                stdout_text, stderr_text = proc.communicate()
                if stdout_text:
                    stdout_lines.append(stdout_text)
                stderr_text = stderr_text or ""
                proc_stdout = "".join(stdout_lines)
                result_json = extract_json_line(proc_stdout or "")
                if proc.returncode == 0:
                    if is_item_cancelled(item["id"]):
                        manual_stop = True
                        last_error = "任务已手动停止。"
                        break
                    html_path = output_dir / "script_table.html"
                    if not html_path.exists():
                        last_error = "Pipeline finished but script_table.html was not created."
                        continue
                    product = write_product_outputs(item["id"], output_dir, result_json or {})
                    script_json = read_json(output_dir / "script_table.json") or read_json(output_dir / "analysis_result.json") or result_json or {}
                    docx_path = write_script_docx(output_dir, script_json, item["video_url"])
                    docx_url = f"/results/{item['id']}/{docx_path.name}" if docx_path and docx_path.exists() else ""
                    # Final pipeline completion must not block on an extra LLM
                    # classification round, otherwise the UI can sit at 90%
                    # long after the script files are already written.
                    decision = detect_content_type_decision_for_output(
                        output_dir,
                        script_json,
                        read_json(output_dir / "evidence_bundle.json"),
                        use_llm=False,
                    )
                    content_type = decision["content_type"]
                    update_job_item(
                        parent_job_id,
                        item_index,
                        status="completed",
                        error="",
                        raw_error="",
                        completed_at=now_iso(),
                        html_url=f"/results/{item['id']}/script_table.html",
                        zh_html_url=f"/results/{item['id']}/script_table.html",
                        pt_html_url="",
                        report_url=product["report_url"],
                        evidence_url=product["evidence_url"],
                        docx_url=docx_url,
                        zh_docx_url=docx_url,
                        pt_docx_url="",
                        artifacts=summarize_artifacts(item["id"], output_dir),
                        result_json=script_json,
                        zh_result_json=script_json,
                        pt_result_json=None,
                        original_result_json=script_json,
                        display_language="zh",
                        tried_models=tried,
                        stage="completed",
                        stage_message="Completed.",
                        content_type=content_type,
                        content_type_source=decision["content_type_source"],
                        content_type_reasoning=decision["content_type_reasoning"],
                        content_type_confidence=decision["content_type_confidence"],
                        title=script_json.get("title") or "Video Script",
                    )
                    return
                last_error = (stderr_text or proc_stdout or "").strip() or "Unknown pipeline failure"
                if not should_try_next_model(last_error):
                    break
                continue
            if manual_stop:
                break
            break
        except Exception as exc:
            last_error = str(exc)
            if not should_try_next_model(last_error):
                break
            continue
        finally:
            unregister_active_process(item["id"])
    update_job_item(
        parent_job_id,
        item_index,
        status="failed",
        error=friendly_error(last_error),
        raw_error=last_error,
        completed_at=now_iso(),
        tried_models=tried,
        stage="failed",
        stage_message="Failed.",
    )


def run_job_batch(job_id: str) -> None:
    try:
        update_job(job_id, status="running", started_at=now_iso(), stage="queued", stage_message="Batch task started.")
        items = jobs[job_id]["items"]
        understanding_mode = str(jobs[job_id].get("mode") or "").strip() == "understanding"
        for idx, item in enumerate(items):
            retry_attempt = 0
            while True:
                execute_single_pipeline(job_id, idx, item)
                current_item = jobs[job_id]["items"][idx]
                if understanding_mode and current_item.get("status") == "completed":
                    summary = build_understanding_summary(
                        current_item.get("result_json") or current_item.get("zh_result_json"),
                        current_item.get("video_url") or "",
                    )
                    update_job_item(
                        job_id,
                        idx,
                        understanding_summary=summary,
                        stage_message="视频理解已完成。",
                    )
                    current_item = jobs[job_id]["items"][idx]
                transient_error = current_item.get("raw_error") or current_item.get("error") or ""
                if (
                    current_item.get("status") == "failed"
                    and should_retry_transient_pipeline(transient_error)
                    and retry_attempt < GEMINI_TRANSIENT_RETRY_ATTEMPTS - 1
                ):
                    retry_attempt += 1
                    delay = gemini_retry_delay(retry_attempt)
                    update_job_item(
                        job_id,
                        idx,
                        status="queued",
                        error="",
                        stage="queued",
                        stage_message=f"Gemini busy. Auto retry {retry_attempt + 1}/{GEMINI_TRANSIENT_RETRY_ATTEMPTS} in {delay}s.",
                        transient_retry_count=retry_attempt,
                    )
                    waited = 0
                    while waited < delay:
                        if is_item_cancelled(current_item["id"]):
                            update_job_item(
                                job_id,
                                idx,
                                status="failed",
                                error="任务已手动停止。",
                                raw_error="任务已手动停止。",
                                completed_at=now_iso(),
                                stage="failed",
                                stage_message="Stopped manually.",
                            )
                            break
                        sleep_for = min(5, delay - waited)
                        time.sleep(sleep_for)
                        waited += sleep_for
                    if jobs[job_id]["items"][idx].get("status") == "failed":
                        break
                    item = jobs[job_id]["items"][idx]
                    continue
                break
        final_items = jobs[job_id]["items"]
        completed = sum(1 for item in final_items if item.get("status") == "completed")
        failed = sum(1 for item in final_items if item.get("status") == "failed")
        first_completed = next((item for item in final_items if item.get("status") == "completed"), None)
        update_job(
            job_id,
            status="completed" if completed else "failed",
            completed_at=now_iso(),
            html_url=first_completed.get("html_url") if first_completed else "",
            report_url=first_completed.get("report_url") if first_completed else "",
            evidence_url=first_completed.get("evidence_url") if first_completed else "",
            docx_url=first_completed.get("docx_url") if first_completed else "",
            result_json=first_completed.get("result_json") if first_completed else None,
            stage="completed" if completed else "failed",
            stage_message=f"Completed {completed}/{len(final_items)} items. Failed {failed}.",
            error="" if completed else "All batch items failed.",
        )
    except Exception as exc:
        update_job(job_id, status="failed", error=friendly_error(str(exc)), completed_at=now_iso(), stage="failed", stage_message="Batch failed.")


def create_job(video_urls: list[str], mode: str = "", user_prompt: str = "") -> dict[str, Any]:
    normalized_mode = str(mode or "").strip().lower()
    job_mode = "understanding" if normalized_mode == "understanding" else ("batch" if len(video_urls) > 1 else "single")
    analysis_prompt = sanitize_analysis_prompt(user_prompt)
    job_id = uuid4().hex
    items = []
    for index, video_url in enumerate(video_urls):
        item_id = uuid4().hex
        items.append(
            {
                "id": item_id,
                "index": index,
                "video_url": video_url,
                "user_prompt": analysis_prompt,
                "status": "queued",
                "stage": "queued",
                "stage_message": "Queued.",
                "created_at": now_iso(),
                "updated_at": now_iso(),
                "html_url": "",
                "zh_html_url": "",
                "pt_html_url": "",
                "report_url": "",
                "evidence_url": "",
                "docx_url": "",
                "zh_docx_url": "",
                "pt_docx_url": "",
                "artifacts": {},
                "error": "",
                "result_json": None,
                "zh_result_json": None,
                "pt_result_json": None,
                "original_result_json": None,
                "display_language": "zh",
                "content_type": "",
                "content_type_source": "auto",
                "content_type_reasoning": "",
                "content_type_confidence": "",
                "title": "",
                "understanding_summary": "",
                "review_status": "",
                "review_stage": "",
                "review_message": "",
                "review_feedback": "",
                "review_mode": REVIEW_MODE_PARTIAL,
                "reviewed": False,
                "edited": False,
            }
        )
    job = {
        "id": job_id,
        "mode": job_mode,
        "video_url": video_urls[0] if len(video_urls) == 1 else "",
        "video_urls": video_urls,
        "user_prompt": analysis_prompt,
        "status": "queued",
        "created_at": now_iso(),
        "updated_at": now_iso(),
        "html_url": "",
        "zh_html_url": "",
        "pt_html_url": "",
        "report_url": "",
        "evidence_url": "",
        "docx_url": "",
        "zh_docx_url": "",
        "pt_docx_url": "",
        "artifacts": {},
        "error": "",
        "result_json": None,
        "zh_result_json": None,
        "pt_result_json": None,
        "display_language": "zh",
        "content_type": "",
        "content_type_source": "auto",
        "content_type_reasoning": "",
        "content_type_confidence": "",
        "stage": "queued",
        "stage_message": "Queued.",
        "items": items,
        "current_item_index": 0,
    }
    with job_lock:
        jobs[job_id] = job
        save_jobs()
    enqueue_job(job_id)
    return public_job_view(job)


def page_html() -> str:
    return """<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width,initial-scale=1">
  <title>Koko</title>
  <link rel="icon" type="image/svg+xml" href="/favicon.svg?v=kwai1">
  <link rel="shortcut icon" href="/favicon.ico?v=kwai1">
  <style>
    @import url('https://fonts.googleapis.com/css2?family=Readex+Pro:wght@300;400;500;600;700&display=swap');
    * {
      box-sizing: border-box;
      font-family: 'Readex Pro', -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif;
    }
    body {
      margin: 0;
      min-height: 100vh;
      color: #1F1F1F;
      background:
        radial-gradient(circle at 10% 8%, rgba(255,130,0,.24), transparent 28%),
        radial-gradient(circle at 82% 14%, rgba(249,115,0,.22), transparent 26%),
        radial-gradient(circle at 50% 48%, rgba(255,244,232,.96), transparent 30%),
        linear-gradient(180deg, #FFD6AE 0%, #FFF0DE 38%, #FFFFFF 100%);
    }
    .hero-shell {
      position: relative;
      min-height: 100vh;
      padding: 18px;
    }
    .hero-panel {
      --mouse-x: 50%;
      --mouse-y: 50%;
      width: min(1320px, 100%);
      margin: 0 auto;
      position: relative;
      border-radius: 34px;
      overflow: hidden;
      border: 1px solid rgba(255,255,255,.85);
      box-shadow: 0 28px 80px rgba(249,115,0,.12);
      background:
        radial-gradient(circle at 16% 18%, rgba(255,130,0,.28), rgba(255,130,0,0) 24%),
        radial-gradient(circle at 82% 18%, rgba(249,115,0,.24), rgba(249,115,0,0) 22%),
        radial-gradient(circle at 72% 56%, rgba(255,244,232,.86), rgba(255,244,232,0) 28%),
        linear-gradient(180deg, #FFC792 0%, #FFF0DE 46%, #FFFFFF 100%);
      min-height: calc(100vh - 36px);
      display: flex;
      flex-direction: column;
    }
    .hero-panel::before {
      content: "";
      position: absolute;
      inset: 0;
      background:
        radial-gradient(circle at 18% 18%, rgba(255,255,255,.74), rgba(255,255,255,0) 16%),
        radial-gradient(circle at 82% 30%, rgba(255,248,235,.86), rgba(255,248,235,0) 14%),
        radial-gradient(circle at 72% 62%, rgba(255,244,232,.72), rgba(255,244,232,0) 18%);
      opacity: .92;
      pointer-events: none;
    }
    .brandbar {
      display: flex;
      align-items: center;
      justify-content: center;
      gap: 16px;
      padding: 16px 20px 0;
      position: relative;
      z-index: 3;
    }
    .navpill {
      display: inline-flex;
      align-items: center;
      gap: 12px;
      border: 1px solid rgba(0,0,0,.08);
      color: rgba(255,255,255,.86);
      background: rgba(0,0,0,.78);
      backdrop-filter: blur(16px);
      padding: 8px 10px;
      border-radius: 999px;
      font-size: 12px;
      font-weight: 600;
      box-shadow: 0 16px 32px rgba(0,0,0,.16);
    }
    .navpill a {
      color: rgba(255,255,255,.78);
      text-decoration: none;
      padding: 8px 16px;
      border-radius: 999px;
      font-size: 12px;
      transition: color .18s ease, background .18s ease;
    }
    .navpill a:hover {
      color: #FFFFFF;
      background: rgba(255,255,255,.08);
    }
    .hero-stage {
      position: relative;
      flex: 1;
      min-height: 0;
      padding: 16px 18px 18px;
    }
    .hero-stage::before {
      content: "";
      position: absolute;
      inset: 0;
      pointer-events: none;
      background:
        radial-gradient(circle at var(--mouse-x) var(--mouse-y), rgba(255,255,255,.46), rgba(255,255,255,0) 10%),
        radial-gradient(circle at var(--mouse-x) var(--mouse-y), rgba(255,130,0,.34), rgba(255,130,0,0) 18%),
        radial-gradient(circle at var(--mouse-x) var(--mouse-y), rgba(249,115,0,.22), rgba(249,115,0,0) 28%);
      opacity: .92;
      mix-blend-mode: screen;
      filter: blur(2px);
    }
    .hero-corner-logo {
      position: absolute;
      top: 26px;
      right: 26px;
      z-index: 4;
      display: block;
      pointer-events: none;
    }
    .hero-corner-logo img {
      width: clamp(420px, 42vw, 720px);
      height: auto;
      object-fit: contain;
      display: block;
      filter: saturate(1.04);
    }
    .hero-copy {
      position: relative;
      min-height: calc(100vh - 152px);
      display: flex;
      align-items: stretch;
      justify-content: space-between;
      border-radius: 28px;
      padding: 24px 22px 26px;
    }
    .hero-left {
      align-self: end;
    }
    .hero-left h1 {
      margin: 0;
      font-size: clamp(7rem, 14vw, 13.4rem);
      line-height: .80;
      letter-spacing: -.06em;
      font-weight: 600;
      max-width: 7ch;
      color: #FF8200;
      text-shadow: 0 10px 30px rgba(255,130,0,.10);
    }
    .hero-left h1 span {
      display: inline-block;
    }
    .hero-left h1 .koko-k {
      font-size: 1.22em;
      line-height: .72;
      letter-spacing: -.08em;
    }
    .hero-left h1 .koko-rest {
      margin-left: -.06em;
    }
    .lede {
      margin: 0;
      color: #FF8200;
      line-height: 1.35;
      font-size: 18px;
      max-width: 18ch;
      padding-left: 6px;
    }
    .hero-side {
      align-self: stretch;
      justify-self: end;
      display: flex;
      flex-direction: column;
      align-items: flex-end;
      justify-content: flex-end;
      gap: 18px;
      padding-bottom: 20px;
    }
    .hero-side p {
      margin: 0;
      color: #FF8200;
      font-size: 14px;
      line-height: 1.6;
      max-width: 24ch;
      text-align: left;
    }
    .hero-cta {
      display: inline-flex;
      align-items: center;
      justify-content: space-between;
      gap: 16px;
      width: fit-content;
      padding: 8px 8px 8px 18px;
      border-radius: 999px;
      background: #FFF8EA;
      color: #1F1F1F;
      text-decoration: none;
      font-size: 14px;
      font-weight: 700;
      box-shadow: 0 14px 28px rgba(0,0,0,.16);
      transition: transform .18s ease, box-shadow .18s ease;
    }
    .hero-cta:hover {
      transform: translateY(-1px);
      box-shadow: 0 18px 32px rgba(0,0,0,.20);
    }
    .hero-cta span {
      width: 36px;
      height: 36px;
      display: inline-flex;
      align-items: center;
      justify-content: center;
      border-radius: 999px;
      background: #1F1F1F;
      color: #FFFFFF;
      font-size: 16px;
    }
    @media (max-width: 1080px) {
      .hero-corner-logo img {
        width: min(70vw, 520px);
      }
    }
    @media (max-width: 720px) {
      .hero-shell {
        padding: 12px;
      }
      .hero-panel {
        min-height: calc(100vh - 24px);
      }
      .brandbar {
        padding: 14px 14px 0;
      }
      .navpill {
        justify-content: center;
        flex-wrap: wrap;
      }
      .hero-stage {
        padding: 16px 12px 18px;
      }
      .hero-copy {
        min-height: calc(100vh - 132px);
        flex-direction: column;
        justify-content: space-between;
        padding: 18px 14px 22px;
      }
      .hero-left h1 {
        font-size: 5rem;
      }
      .hero-left h1 .koko-k {
        font-size: 1.16em;
      }
      .hero-side {
        align-items: flex-start;
      }
      .hero-side p {
        max-width: 18ch;
      }
      .hero-corner-logo {
        top: 22px;
        right: 16px;
      }
      .hero-corner-logo img {
        width: min(74vw, 360px);
      }
    }
  </style>
</head>
<body>
  <main class="hero-shell">
    <div class="hero-panel">
      <div class="brandbar">
        <div class="navpill">
          <a href="/">Start</a>
          <a href="/studio">Studio</a>
          <a href="/studio#split-panel">Preview</a>
          <a href="/library">Script Admin</a>
          <a href="/creator-admin">Creator Ops</a>
          <a href="/stats">Stats</a>
        </div>
      </div>
      <div class="hero-stage">
        <div class="hero-corner-logo">
          <img src="/brand/kwai-wordmark.svg" alt="Kwai" />
        </div>
        <div class="hero-copy">
          <div class="hero-left">
            <h1><span class="koko-k">K</span><span class="koko-rest">oKo</span></h1>
            <p class="lede">Kwai Coach</p>
          </div>
          <div class="hero-side">
            <p>
              Creator-side analysis for story clarity, stronger hooks, cleaner payoff, and better-performing short-form ideas.
            </p>
            <a class="hero-cta" href="/studio">进入内容中台 <span>→</span></a>
          </div>
        </div>
      </div>
    </div>
  </main>
</body>
</html>"""


def studio_html() -> str:
    return f"""<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width,initial-scale=1">
  <title>Koko · Kwai Coach</title>
  {FAVICON_LINKS}
  <style>
    @import url('https://fonts.googleapis.com/css2?family=Readex+Pro:wght@300;400;500;600;700&family=Instrument+Serif:ital@1&display=swap');
    :root {{
      --bg: #FFF8F2;
      --card: rgba(255,255,255,.96);
      --ink: #FF8200;
      --muted: #FF8200;
      --line: rgba(255,130,0,.16);
      --brand: #FF8200;
      --brand-deep: #F97300;
      --brand-soft: #FFF4E8;
      --brand-soft-2: #FFF0DE;
      --ok: #157347;
      --err: #b42318;
      --wait: #935f14;
      --shadow: 0 28px 80px rgba(249,115,0,.12);
    }}
    * {{
      box-sizing: border-box;
      font-family: 'Readex Pro', -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif;
    }}
    html {{ scroll-behavior: smooth; }}
    body {{
      margin: 0;
      min-height: 100vh;
      color: var(--ink);
      background:
        radial-gradient(circle at 10% 8%, rgba(255,130,0,.24), transparent 28%),
        radial-gradient(circle at 82% 14%, rgba(249,115,0,.22), transparent 26%),
        radial-gradient(circle at 50% 48%, rgba(255,244,232,.96), transparent 30%),
        linear-gradient(180deg, #FFD6AE 0%, #FFF0DE 38%, #FFFFFF 100%);
    }}
    .shell {{ width: 100%; }}
    .site {{ width: 100%; }}
    .hero-shell {{
      position: relative;
      min-height: 100vh;
      padding: 18px;
    }}
    .hero-panel {{
      --mouse-x: 50%;
      --mouse-y: 50%;
      --tilt-x: 0px;
      --tilt-y: 0px;
      width: min(1320px, 100%);
      margin: 0 auto;
      position: relative;
      border-radius: 34px;
      overflow: hidden;
      border: 1px solid rgba(255,255,255,.85);
      box-shadow: var(--shadow);
      background:
        radial-gradient(circle at 16% 18%, rgba(255,130,0,.28), rgba(255,130,0,0) 24%),
        radial-gradient(circle at 82% 18%, rgba(249,115,0,.24), rgba(249,115,0,0) 22%),
        radial-gradient(circle at 72% 56%, rgba(255,244,232,.86), rgba(255,244,232,0) 28%),
        linear-gradient(180deg, #FFC792 0%, #FFF0DE 46%, #FFFFFF 100%);
      min-height: calc(100vh - 36px);
      display: flex;
      flex-direction: column;
    }}
    .hero-panel::before {{
      content: "";
      position: absolute;
      inset: 0;
      background:
        radial-gradient(circle at 18% 18%, rgba(255,255,255,.74), rgba(255,255,255,0) 16%),
        radial-gradient(circle at 82% 30%, rgba(255,248,235,.86), rgba(255,248,235,0) 14%),
        radial-gradient(circle at 72% 62%, rgba(255,244,232,.72), rgba(255,244,232,0) 18%);
      opacity: .92;
      pointer-events: none;
    }}
    .hero-panel::after {{
      content: "";
      position: absolute;
      inset: 0;
      background-image:
        radial-gradient(rgba(255,255,255,.12) 0.7px, transparent 0.7px),
        radial-gradient(rgba(255,130,0,.05) 0.7px, transparent 0.7px);
      background-size: 8px 8px, 12px 12px;
      background-position: 0 0, 3px 3px;
      mix-blend-mode: soft-light;
      opacity: .18;
      pointer-events: none;
    }}
    .brandbar {{
      display: flex;
      align-items: center;
      justify-content: center;
      gap: 16px;
      padding: 16px 20px 0;
      position: relative;
      z-index: 3;
    }}
    .navpill {{
      display: inline-flex;
      align-items: center;
      gap: 12px;
      border: 1px solid rgba(0,0,0,.08);
      color: rgba(255,255,255,.86);
      background: rgba(0,0,0,.78);
      backdrop-filter: blur(16px);
      padding: 8px 10px;
      border-radius: 999px;
      font-size: 12px;
      font-weight: 600;
      box-shadow: 0 16px 32px rgba(0,0,0,.16);
    }}
    .navpill a {{
      color: rgba(255,255,255,.78);
      text-decoration: none;
      padding: 8px 16px;
      border-radius: 999px;
      font-size: 12px;
      transition: color .18s ease, background .18s ease;
      letter-spacing: .02em;
    }}
    .navpill a:hover {{
      color: #FFFFFF;
      background: rgba(255,255,255,.08);
    }}
    .hero-stage {{
      position: relative;
      flex: 1;
      min-height: 0;
      padding: 16px 18px 18px;
    }}
    .hero-stage::before {{
      content: "";
      position: absolute;
      inset: 0;
      pointer-events: none;
      background:
        radial-gradient(circle at var(--mouse-x) var(--mouse-y), rgba(255,255,255,.46), rgba(255,255,255,0) 10%),
        radial-gradient(circle at var(--mouse-x) var(--mouse-y), rgba(255,130,0,.34), rgba(255,130,0,0) 18%),
        radial-gradient(circle at var(--mouse-x) var(--mouse-y), rgba(249,115,0,.22), rgba(249,115,0,0) 28%);
      opacity: .92;
      transition: opacity .16s ease-out;
      mix-blend-mode: screen;
      filter: blur(2px);
    }}
    .hero-corner-logo {{
      position: absolute;
      top: 26px;
      right: 26px;
      z-index: 4;
      display: block;
      pointer-events: none;
    }}
    .hero-corner-logo img {{
      width: clamp(420px, 42vw, 720px);
      height: auto;
      object-fit: contain;
      display: block;
      filter: saturate(1.04);
    }}
.hero-copy {{
      position: relative;
      min-height: calc(100vh - 138px);
      display: grid;
      grid-template-columns: minmax(0, 1.35fr) minmax(280px, .65fr);
      gap: 24px;
      align-items: end;
      border-radius: 28px;
      padding: 24px 22px 26px;
    }}
    .hero-left h1 {{
      margin: 0;
      font-size: clamp(7rem, 14vw, 13.4rem);
      line-height: .80;
      letter-spacing: -.06em;
      font-weight: 600;
      max-width: 7ch;
      color: #FF8200;
      text-shadow: 0 10px 30px rgba(255,130,0,.10);
    }}
    .hero-left h1 span {{
      display: inline-block;
    }}
    .hero-left h1 .koko-k {{
      font-size: 1.22em;
      line-height: .72;
      letter-spacing: -.08em;
    }}
    .hero-left h1 .koko-rest {{
      margin-left: -.06em;
    }}
    .lede {{
      margin: 0;
      color: #FF8200;
      line-height: 1.35;
      font-size: 18px;
      max-width: 18ch;
      padding-left: 6px;
    }}
.hero-side {{
      align-self: end;
      justify-self: end;
      display: flex;
      flex-direction: column;
      gap: 18px;
      padding-bottom: 6px;
    }}
    .hero-side p {{
      margin: 0;
      color: #FF8200;
      font-size: 14px;
      line-height: 1.6;
      max-width: 24ch;
      text-shadow: none;
    }}
    .hero-cta {{
      display: inline-flex;
      align-items: center;
      justify-content: space-between;
      gap: 16px;
      width: fit-content;
      padding: 8px 8px 8px 18px;
      border-radius: 999px;
      background: #FFF8EA;
      color: #1F1F1F;
      text-decoration: none;
      font-size: 14px;
      font-weight: 700;
      box-shadow: 0 14px 28px rgba(0,0,0,.16);
      transition: transform .18s ease, box-shadow .18s ease;
    }}
    .hero-cta:hover {{
      transform: translateY(-1px);
      box-shadow: 0 18px 32px rgba(0,0,0,.20);
    }}
    .hero-cta span {{
      width: 36px;
      height: 36px;
      display: inline-flex;
      align-items: center;
      justify-content: center;
      border-radius: 999px;
      background: #1F1F1F;
      color: #FFFFFF;
      font-size: 16px;
    }}
    .hero-scroll {{
      position: absolute;
      left: 0;
      bottom: 0;
      z-index: 2;
      opacity: 0;
      pointer-events: none;
    }}
    .work-shell {{
      padding: 10px 24px 32px;
      background: transparent;
    }}
    .workspace {{
      width: min(1320px, 100%);
      margin: 0 auto;
      border-radius: 34px;
      overflow: hidden;
      border: 1px solid rgba(255,255,255,.85);
      box-shadow: var(--shadow);
      background:
        radial-gradient(circle at 12% 16%, rgba(255,130,0,.22), rgba(255,130,0,0) 22%),
        radial-gradient(circle at 86% 18%, rgba(249,115,0,.18), rgba(249,115,0,0) 22%),
        radial-gradient(circle at 70% 62%, rgba(255,244,232,.82), rgba(255,244,232,0) 24%),
        linear-gradient(180deg, #FFC792 0%, #FFF0DE 44%, #FFFFFF 100%);
      display: flex;
      flex-direction: column;
      gap: 0;
      align-items: stretch;
      min-width: 0;
    }}
    .workspace::before {{
      content: "";
      position: absolute;
      inset: 0;
      pointer-events: none;
    }}
    .composer {{
      border-radius: 0;
      border: 0;
      border-bottom: 1px solid rgba(255,130,0,.14);
      background: transparent;
      box-shadow: none;
      padding: 28px 28px 22px;
    }}
    textarea {{
      width: 100%;
      min-height: 160px;
      resize: vertical;
      border: 1px solid rgba(255,130,0,.16);
      background: rgba(255,255,255,.98);
      border-radius: 20px;
      padding: 18px 18px;
      font-size: 16px;
      outline: none;
      color: var(--ink);
      transition: border-color .2s ease, box-shadow .2s ease, transform .2s ease;
      line-height: 1.65;
    }}
    textarea:focus {{
      border-color: rgba(255,130,0,.45);
      box-shadow: 0 0 0 4px rgba(255,130,0,.10);
      transform: translateY(-1px);
    }}
    .analysis-prompt-box {{
      margin-top: 18px;
      padding: 16px;
      border: 1px solid rgba(255,130,0,.16);
      border-radius: 22px;
      background: rgba(255,255,255,.58);
    }}
    .analysis-prompt-box label {{
      display: block;
      margin-bottom: 8px;
    }}
    .analysis-prompt {{
      min-height: 92px;
      font-size: 15px;
    }}
    .analysis-prompt-note {{
      margin: 8px 0 0;
      color: rgba(31,31,31,.56);
      font-size: 13px;
      line-height: 1.55;
      font-weight: 650;
    }}
    .composer-head {{
      display: block;
      margin-bottom: 14px;
    }}
    .composer-title {{
      font-size: 28px;
      line-height: 1.02;
      letter-spacing: -.04em;
      margin: 0 0 8px;
      font-weight: 600;
    }}
    .composer-copy {{
      margin: 0;
      color: var(--muted);
      line-height: 1.7;
      font-size: 14px;
      max-width: 42ch;
    }}
    .coach-badge {{
      flex: 0 0 auto;
      border-radius: 20px;
      padding: 10px 12px;
      background: var(--brand-soft);
      border: 1px solid rgba(255,130,0,.14);
      color: var(--brand-deep);
      font-size: 12px;
      line-height: 1.5;
      font-weight: 600;
      text-transform: uppercase;
      letter-spacing: .06em;
    }}
    label {{
      display: block;
      font-size: 14px;
      font-weight: 700;
      margin-bottom: 10px;
      color: #FF8200;
    }}
    input[type="url"] {{
      width: 100%;
      border: 1px solid rgba(255,130,0,.16);
      background: rgba(255,255,255,.98);
      border-radius: 20px;
      padding: 18px 18px;
      font-size: 16px;
      outline: none;
      color: var(--ink);
      transition: border-color .2s ease, box-shadow .2s ease, transform .2s ease;
    }}
    input[type="url"]:focus {{
      border-color: rgba(255,130,0,.45);
      box-shadow: 0 0 0 4px rgba(255,130,0,.10);
      transform: translateY(-1px);
    }}
    input[type="file"] {{
      width: 100%;
      border: 1px dashed rgba(255,130,0,.24);
      background: rgba(255,255,255,.78);
      border-radius: 16px;
      padding: 12px 14px;
      font-size: 13px;
      color: var(--ink);
    }}
    .filter-upload-row {{
      margin-top: 12px;
      display: grid;
      grid-template-columns: minmax(0, 1fr);
      gap: 8px;
    }}
    .filter-upload-hint {{
      font-size: 12px;
      line-height: 1.7;
      color: rgba(255,130,0,.82);
    }}
    .actions {{
      margin-top: 16px;
      display: grid;
      grid-template-columns: minmax(0, 1fr) auto;
      gap: 12px;
      align-items: stretch;
    }}
    button {{
      width: 100%;
      border: 0;
      cursor: pointer;
      border-radius: 18px;
      padding: 16px 20px;
      font-size: 15px;
      font-weight: 700;
      background: linear-gradient(135deg, var(--brand), var(--brand-deep));
      color: #fff;
      box-shadow: 0 16px 28px rgba(249,115,0,.22);
      transition: transform .18s ease, box-shadow .18s ease, filter .18s ease;
    }}
    button:hover {{
      transform: translateY(-1px);
      box-shadow: 0 18px 30px rgba(249,115,0,.26);
      filter: saturate(1.02);
    }}
    .actions .action-link {{
      width: auto;
      min-width: 160px;
      border-radius: 18px;
      padding: 16px 20px;
      font-size: 15px;
      border-color: rgba(255,130,0,.18);
      box-shadow: 0 16px 28px rgba(249,115,0,.10);
      background: rgba(255,255,255,.9);
    }}
    .actions .action-link:hover {{
      transform: translateY(-1px);
      box-shadow: 0 18px 30px rgba(249,115,0,.14);
    }}
    .actions .action-link:disabled {{
      cursor: not-allowed;
      opacity: .55;
      transform: none;
      box-shadow: none;
    }}
    .actions button:disabled {{
      cursor: not-allowed;
      opacity: .6;
      transform: none;
      filter: none;
      box-shadow: none;
    }}
    @media (max-width: 720px) {{
      .actions {{
        grid-template-columns: 1fr;
      }}
      .actions .action-link {{
        width: 100%;
        min-width: 0;
      }}
    }}
    .server-note {{
      margin-top: 12px;
      font-size: 13px;
      color: var(--muted);
      line-height: 1.65;
    }}
    .mini-tags {{
      display: flex;
      flex-wrap: wrap;
      gap: 8px;
      margin-top: 14px;
    }}
    .mini-tag {{
      display: inline-flex;
      align-items: center;
      border-radius: 999px;
      padding: 7px 11px;
      font-size: 12px;
      font-weight: 600;
      color: var(--muted);
      background: rgba(255,255,255,.9);
      border: 1px solid rgba(31,31,31,.08);
    }}
    .status-box {{
      min-height: 260px;
      border-radius: 0 0 34px 34px;
      border: 0;
      background: transparent;
      padding: 24px 28px 28px;
      color: var(--muted);
      line-height: 1.7;
      box-shadow: none;
      display: none;
    }}
    .status-box.visible {{
      display: block;
    }}
    .status-empty {{
      display: flex;
      flex-direction: column;
      justify-content: center;
      min-height: 210px;
      gap: 14px;
    }}
    .status-empty-title {{
      font-size: 26px;
      line-height: 1.04;
      letter-spacing: -.03em;
      color: var(--ink);
      font-weight: 600;
      max-width: 16ch;
    }}
    .status-empty-copy {{
      max-width: 48ch;
      font-size: 14px;
      color: var(--muted);
    }}
    .progress-wrap {{
      margin-top: 14px;
    }}
    .progress-top {{
      display: flex;
      align-items: center;
      justify-content: space-between;
      gap: 14px;
      margin-bottom: 8px;
    }}
    .progress-top-copy {{
      display: flex;
      flex-direction: column;
      gap: 3px;
      min-width: 0;
    }}
    .progress-kicker {{
      font-size: 13px;
      font-weight: 800;
      color: var(--ink);
    }}
    .progress-stage-copy {{
      font-size: 14px;
      line-height: 1.6;
      color: var(--muted);
    }}
    .progress-percent {{
      font-size: 18px;
      font-weight: 800;
      color: #FF8200;
      white-space: nowrap;
    }}
    .progress-meta {{
      display: flex;
      flex-wrap: wrap;
      gap: 8px;
      margin-top: 10px;
    }}
    .progress-meta-chip {{
      display: inline-flex;
      align-items: center;
      gap: 6px;
      padding: 7px 10px;
      border-radius: 999px;
      border: 1px solid rgba(255,130,0,.14);
      background: rgba(255,248,238,.84);
      color: rgba(31,31,31,.72);
      font-size: 12px;
      font-weight: 700;
    }}
    .thinking-shell {{
      margin-top: 16px;
      border-radius: 18px;
      border: 1px solid rgba(255,130,0,.12);
      background: rgba(255,250,244,.72);
      padding: 14px 14px 10px;
    }}
    .thinking-head {{
      display: flex;
      align-items: center;
      justify-content: space-between;
      gap: 10px;
      margin-bottom: 10px;
    }}
    .thinking-title {{
      font-size: 12px;
      font-weight: 800;
      letter-spacing: .08em;
      text-transform: uppercase;
      color: rgba(31,31,31,.56);
    }}
    .thinking-updated {{
      font-size: 12px;
      color: rgba(31,31,31,.48);
    }}
    .thinking-list {{
      display: grid;
      gap: 8px;
    }}
    .thinking-item {{
      display: flex;
      align-items: flex-start;
      gap: 10px;
      font-size: 13px;
      line-height: 1.6;
      color: rgba(31,31,31,.74);
    }}
    .thinking-dot {{
      width: 8px;
      height: 8px;
      border-radius: 999px;
      margin-top: 6px;
      flex: 0 0 8px;
      background: rgba(255,130,0,.28);
    }}
    .thinking-item.done .thinking-dot {{
      background: #57B56B;
    }}
    .thinking-item.active .thinking-dot {{
      background: #FF8200;
      box-shadow: 0 0 0 6px rgba(255,130,0,.12);
    }}
    .thinking-item.note .thinking-dot {{
      background: rgba(31,31,31,.2);
    }}
    .progress-rail {{
      width: 100%;
      height: 8px;
      border-radius: 999px;
      background: rgba(31,31,31,.08);
      overflow: hidden;
    }}
    .progress-fill {{
      height: 100%;
      width: 0%;
      border-radius: 999px;
      background: linear-gradient(90deg, var(--brand), var(--brand-deep));
      transition: width .35s ease;
    }}
    .step-list {{
      display: grid;
      grid-template-columns: repeat(8, minmax(0, 1fr));
      gap: 0;
      margin-top: 18px;
      align-items: start;
    }}
    .step-pill {{
      display: flex;
      flex-direction: column;
      align-items: center;
      gap: 8px;
      text-align: center;
      position: relative;
      padding: 0 4px;
      background: transparent;
      border: 0;
      color: rgba(31,31,31,.42);
      font-size: 12px;
      font-weight: 700;
    }}
    .step-pill::before {{
      content: "";
      position: absolute;
      top: 18px;
      left: calc(-50% + 18px);
      width: calc(100% - 36px);
      height: 2px;
      background: rgba(31,31,31,.12);
      z-index: 0;
    }}
    .step-pill:first-child::before {{
      display: none;
    }}
    .step-bubble {{
      width: 36px;
      height: 36px;
      border-radius: 999px;
      display: inline-flex;
      align-items: center;
      justify-content: center;
      background: rgba(255,255,255,.96);
      border: 2px solid rgba(31,31,31,.12);
      color: rgba(31,31,31,.55);
      font-size: 14px;
      font-weight: 800;
      position: relative;
      z-index: 1;
      box-sizing: border-box;
    }}
    .step-pill span:last-child {{
      line-height: 1.45;
    }}
    .step-pill.active {{
      color: #FF8200;
    }}
    .step-pill.active .step-bubble {{
      border-color: rgba(255,130,0,.9);
      color: #FF8200;
      background: rgba(255,255,255,.98);
    }}
    .step-pill.done {{
      color: var(--ink);
    }}
    .step-pill.done::before {{
      background: rgba(24,163,74,.38);
    }}
    .step-pill.done .step-bubble {{
      border-color: rgba(24,163,74,.72);
      color: #18A34A;
      background: rgba(239,251,243,.98);
    }}
    .status-box.ready {{
      color: var(--ink);
      border-style: solid;
    }}
    .batch-summary {{
      display: flex;
      flex-wrap: wrap;
      gap: 10px;
      margin: 4px 0 16px;
    }}
    .batch-chip {{
      border-radius: 999px;
      padding: 8px 12px;
      font-size: 12px;
      font-weight: 700;
      color: #FF8200;
      background: rgba(255,255,255,.84);
      border: 1px solid rgba(255,130,0,.16);
    }}
    .batch-dashboard {{
      display: flex;
      flex-direction: column;
      gap: 18px;
      margin-top: 8px;
    }}
    .batch-overview {{
      border: 1px solid rgba(31,31,31,.08);
      border-radius: 24px;
      background: rgba(255,255,255,.92);
      padding: 18px 20px;
      box-shadow: 0 16px 42px rgba(15,23,42,.06);
    }}
    .batch-overview-top {{
      display: flex;
      align-items: flex-start;
      justify-content: space-between;
      gap: 16px;
      flex-wrap: wrap;
      margin-bottom: 10px;
    }}
    .batch-overview-copy {{
      display: flex;
      flex-direction: column;
      gap: 10px;
      min-width: 0;
    }}
    .batch-overview-title {{
      font-size: 22px;
      line-height: 1.12;
      letter-spacing: -.03em;
      color: var(--ink);
      font-weight: 800;
    }}
    .batch-overview-subtitle {{
      display: inline-flex;
      align-items: center;
      gap: 8px;
      font-size: 14px;
      line-height: 1.6;
      color: #2962FF;
      text-decoration: none;
      max-width: 72ch;
      word-break: break-all;
    }}
    .batch-overview-subtitle::before {{
      content: "🔗";
      font-size: 14px;
      line-height: 1;
    }}
    .batch-job-meta {{
      display: flex;
      align-items: center;
      gap: 10px;
      flex-wrap: wrap;
      font-size: 13px;
      line-height: 1.6;
      color: var(--muted);
    }}
    .job-copy-btn {{
      display: inline-flex;
      align-items: center;
      justify-content: center;
      width: 28px;
      height: 28px;
      border-radius: 999px;
      border: 1px solid rgba(31,31,31,.08);
      background: rgba(255,255,255,.98);
      color: var(--muted);
      cursor: pointer;
      font-size: 13px;
      font-weight: 700;
    }}
    .batch-meta {{
      display: flex;
      flex-wrap: wrap;
      gap: 10px;
      margin-top: 14px;
      margin-bottom: 12px;
    }}
    .focus-note {{
      margin-top: 0;
      font-size: 14px;
      line-height: 1.65;
      color: var(--muted);
    }}
    .queue-shell {{
      border: 1px solid rgba(31,31,31,.08);
      border-radius: 24px;
      background: rgba(255,255,255,.92);
      padding: 18px;
    }}
    .queue-header {{
      display: flex;
      align-items: baseline;
      justify-content: space-between;
      gap: 14px;
      flex-wrap: wrap;
      margin-bottom: 14px;
    }}
    .queue-header h3 {{
      margin: 0;
      font-size: 20px;
      letter-spacing: -.03em;
      color: var(--ink);
    }}
    .queue-header p {{
      margin: 0;
      font-size: 13px;
      line-height: 1.6;
      color: var(--muted);
    }}
    .queue-list {{
      display: grid;
      grid-template-columns: repeat(auto-fit, minmax(320px, 1fr));
      gap: 14px;
    }}
    .queue-card {{
      border: 1px solid rgba(31,31,31,.08);
      border-radius: 18px;
      background: rgba(255,255,255,.98);
      padding: 16px;
      display: flex;
      flex-direction: column;
      gap: 12px;
      min-height: 148px;
      box-shadow: 0 12px 24px rgba(15,23,42,.03);
    }}
    .queue-card.current {{
      border-color: rgba(255,130,0,.34);
      background:
        radial-gradient(circle at 12% 14%, rgba(255,130,0,.12), rgba(255,130,0,0) 24%),
        rgba(255,255,255,.98);
      box-shadow: 0 16px 34px rgba(249,115,0,.08);
    }}
    .queue-card-top {{
      display: flex;
      align-items: flex-start;
      justify-content: space-between;
      gap: 10px;
    }}
    .queue-index {{
      font-size: 13px;
      font-weight: 800;
      color: #FF8200;
    }}
    .queue-status {{
      display: inline-flex;
      align-items: center;
      border-radius: 999px;
      padding: 6px 10px;
      font-size: 12px;
      font-weight: 800;
    }}
    .queue-status.waiting {{
      color: var(--wait);
      background: rgba(147,95,20,.12);
    }}
    .queue-status.running {{
      color: #FF8200;
      background: rgba(255,130,0,.12);
    }}
    .queue-status.completed {{
      color: var(--ok);
      background: rgba(21,115,71,.12);
    }}
    .queue-status.failed {{
      color: var(--err);
      background: rgba(180,35,24,.12);
    }}
    .queue-title {{
      margin: 0;
      font-size: 17px;
      line-height: 1.28;
      letter-spacing: -.02em;
      color: var(--ink);
    }}
    .queue-url {{
      font-size: 13px;
      line-height: 1.55;
      color: var(--muted);
      word-break: break-all;
    }}
    .queue-stage {{
      font-size: 15px;
      line-height: 1.6;
      color: var(--ink);
      font-weight: 600;
    }}
    .queue-link-icon {{
      font-size: 13px;
      color: rgba(31,31,31,.42);
      margin-right: 8px;
    }}
    .queue-error {{
      font-size: 12px;
      line-height: 1.55;
      color: var(--err);
    }}
    .detail-section {{
      border: 1px solid rgba(255,130,0,.16);
      border-radius: 24px;
      background: rgba(255,255,255,.74);
      padding: 18px;
    }}
    .detail-header {{
      display: flex;
      align-items: baseline;
      justify-content: space-between;
      gap: 14px;
      flex-wrap: wrap;
      margin-bottom: 14px;
    }}
    .detail-header h3 {{
      margin: 0;
      font-size: 20px;
      letter-spacing: -.03em;
      color: var(--ink);
    }}
    .detail-header p {{
      margin: 0;
      font-size: 13px;
      line-height: 1.6;
      color: var(--muted);
    }}
    .detail-controls {{
      display: flex;
      flex-wrap: wrap;
      gap: 10px;
      justify-content: flex-end;
    }}
    .item-stack {{
      display: flex;
      flex-direction: column;
      gap: 14px;
    }}
    .item-card {{
      border: 1px solid rgba(255,130,0,.16);
      border-radius: 18px;
      background: rgba(255,255,255,.82);
      overflow: hidden;
    }}
    .item-card summary {{
      list-style: none;
      cursor: pointer;
      padding: 14px 16px;
      display: flex;
      align-items: center;
      justify-content: space-between;
      gap: 12px;
      font-weight: 700;
      color: #FF8200;
    }}
    .item-card summary::-webkit-details-marker {{ display: none; }}
    .item-meta {{
      display: flex;
      flex-wrap: wrap;
      gap: 10px;
      font-size: 12px;
      padding: 0 16px 14px;
      color: #FF8200;
    }}
    .item-body {{
      padding: 0 16px 16px;
    }}
    .result-workbench {{
      display: grid;
      grid-template-columns: 1fr;
      gap: 18px;
      align-items: start;
    }}
    .result-main {{
      min-width: 0;
      display: flex;
      flex-direction: column;
      gap: 18px;
    }}
    .primary-review-strip {{
      display: grid;
      grid-template-columns: minmax(0, 1fr);
      gap: 18px;
      padding: 2px;
      align-items: start;
    }}
    .editable-script-pane {{
      width: 100%;
      min-width: 0;
    }}
    .reference-video-pane {{
      width: 100%;
      min-width: 0;
    }}
    .workbench-pane {{
      min-width: 0;
      border: 1px solid rgba(255,130,0,.14);
      border-radius: 18px;
      background: rgba(255,255,255,.76);
      padding: 14px;
    }}
    .workbench-pane h4 {{
      margin: 0 0 10px;
      color: var(--ink);
      font-size: 15px;
      line-height: 1.35;
    }}
    .workbench-pane-note {{
      margin: -4px 0 12px;
      color: var(--muted);
      font-size: 12px;
      line-height: 1.6;
    }}
    .video-verify-pane {{
      position: sticky;
      top: 14px;
    }}
    .assistant-sidebar {{
      min-width: 0;
      position: fixed;
      top: 0;
      right: 0;
      bottom: 0;
      z-index: 80;
      width: 232px;
      height: 100vh;
      padding: 0;
      overflow: hidden;
      border-radius: 0;
      border-width: 0 0 0 1px;
      border-color: rgba(255,130,0,.18);
      background:
        radial-gradient(circle at 10% 4%, rgba(255,130,0,.16), rgba(255,130,0,0) 30%),
        rgba(255,255,255,.94);
      box-shadow: -18px 0 42px rgba(249,115,0,.14);
      backdrop-filter: blur(18px);
      -webkit-backdrop-filter: blur(18px);
      display: flex;
      flex-direction: column;
    }}
    .assistant-video-block {{
      padding: 10px 10px 0;
      flex: 0 0 auto;
    }}
    .assistant-video-block h4 {{
      margin: 0 0 8px;
      color: var(--ink);
      font-size: 14px;
      line-height: 1.35;
    }}
    .assistant-video-block .workbench-pane-note {{
      display: none;
    }}
    .assistant-video-block .source-video-frame {{
      width: min(100%, 168px);
      max-height: 260px;
    }}
    .assistant-video-block .source-video-empty {{
      min-height: 82px;
      padding: 10px;
      font-size: 11px;
      line-height: 1.55;
    }}
    .assistant-video-block .source-video-status {{
      display: none;
    }}
    .assistant-video-block .video-timeline {{
      display: none;
    }}
    .assistant-sidebar .review-shell {{
      flex: 1 1 auto;
      min-height: 0;
      height: auto;
    }}
    .video-verify-layout {{
      display: grid;
      grid-template-columns: 1fr;
      gap: 14px;
      align-items: start;
    }}
    .source-video-frame {{
      width: 100%;
      aspect-ratio: 9 / 16;
      max-height: 620px;
      border-radius: 16px;
      background: #111;
      display: block;
      object-fit: contain;
      border: 1px solid rgba(255,130,0,.12);
      margin: 0 auto;
    }}
    .source-video-empty {{
      min-height: 360px;
      border-radius: 16px;
      border: 1px dashed rgba(255,130,0,.28);
      background: rgba(255,244,232,.72);
      color: #FF8200;
      display: grid;
      place-items: center;
      text-align: center;
      padding: 18px;
      line-height: 1.65;
      font-weight: 700;
    }}
    .source-video-status {{
      margin-top: 10px;
      border-radius: 14px;
      background: rgba(255,244,232,.82);
      color: #FF8200;
      padding: 10px 12px;
      font-size: 12px;
      line-height: 1.55;
    }}
    .video-timeline {{
      margin-top: 0;
      display: flex;
      flex-direction: column;
      gap: 8px;
      max-height: 420px;
      overflow: auto;
      padding-right: 2px;
    }}
    .timeline-jump {{
      width: 100%;
      text-align: left;
      border: 1px solid rgba(255,130,0,.14);
      background: rgba(255,255,255,.82);
      color: #FF8200;
      border-radius: 14px;
      padding: 10px 12px;
      cursor: pointer;
    }}
    .timeline-jump strong {{
      display: block;
      font-size: 12px;
      margin-bottom: 4px;
    }}
    .timeline-jump span {{
      display: block;
      color: var(--muted);
      font-size: 12px;
      line-height: 1.45;
    }}
    .ops-pane {{
      display: flex;
      align-items: stretch;
      flex-wrap: wrap;
      gap: 10px;
      padding: 16px;
    }}
    .confirm-pane {{
      padding: 18px;
    }}
    .ops-section.final-action {{
      flex: 1 1 100%;
    }}
    .ops-section {{
      flex: 1 1 220px;
      border: 0;
      border-radius: 14px;
      background: transparent;
      padding: 0;
    }}
    .ops-section.final-action {{
      background: transparent;
    }}
    .ops-section-title {{
      margin: 0 0 6px;
      color: var(--ink);
      font-size: 14px;
      font-weight: 900;
    }}
    .ops-pane .link-row {{
      margin: 0;
      gap: 8px;
    }}
    .followup-actions .link-row {{
      align-items: stretch;
    }}
    .ops-pane .action-link {{
      min-height: 38px;
      padding: 9px 14px;
    }}
    .ops-pane .library-confirm-card {{
      min-height: 100%;
      margin: 0;
      padding: 18px;
      border-radius: 14px;
      background: rgba(255,255,255,.88);
    }}
    .ops-pane .library-confirm-title {{
      font-size: 15px;
    }}
    .ops-pane .library-confirm-note {{
      font-size: 12px;
      line-height: 1.5;
    }}
    .review-chat-log {{
      display: flex;
      flex-direction: column;
      gap: 6px;
      margin-bottom: 6px;
      flex: 1 1 auto;
      min-height: 132px;
      max-height: none;
      overflow: auto;
      padding-right: 2px;
    }}
    .review-message {{
      border-radius: 12px;
      padding: 8px 10px;
      font-size: 11px;
      line-height: 1.45;
      max-width: 100%;
    }}
    .review-message.koko {{
      background: rgba(255,255,255,.88);
      color: var(--muted);
      border: 1px solid rgba(255,130,0,.10);
      align-self: flex-start;
    }}
    .review-message.user {{
      background: rgba(255,130,0,.12);
      color: #FF8200;
      border: 1px solid rgba(255,130,0,.14);
      align-self: flex-end;
    }}
    .link-row {{
      display: flex;
      flex-wrap: wrap;
      gap: 10px;
      margin-bottom: 12px;
    }}
    .action-link {{
      display: inline-flex;
      align-items: center;
      justify-content: center;
      gap: 6px;
      text-decoration: none;
      border-radius: 999px;
      padding: 10px 14px;
      color: #FF8200;
      border: 1px solid rgba(255,130,0,.18);
      background: rgba(255,255,255,.9);
      font-weight: 700;
      font-size: 13px;
      cursor: pointer;
    }}
    .action-link.primary {{
      background: linear-gradient(135deg, var(--brand), var(--brand-deep));
      color: #fff;
      border-color: transparent;
    }}
    .library-confirm-card {{
      border: 1px solid rgba(255,130,0,.22);
      border-radius: 18px;
      background:
        radial-gradient(circle at 12% 18%, rgba(255,130,0,.13), rgba(255,130,0,0) 28%),
        rgba(255,255,255,.78);
      padding: 16px;
      margin-bottom: 14px;
      display: flex;
      align-items: center;
      justify-content: space-between;
      gap: 14px;
      flex-wrap: wrap;
    }}
    .library-confirm-copy {{
      display: flex;
      flex-direction: column;
      gap: 5px;
      min-width: min(420px, 100%);
    }}
    .library-confirm-title {{
      font-size: 17px;
      font-weight: 900;
      color: var(--ink);
    }}
    .library-confirm-note {{
      font-size: 13px;
      line-height: 1.6;
      color: var(--muted);
    }}
    .manual-tag-picker {{
      display: grid;
      gap: 6px;
      min-width: min(260px, 100%);
    }}
    .manual-tag-picker label {{
      color: var(--muted);
      font-size: 12px;
      font-weight: 800;
    }}
    .manual-tag-picker select {{
      width: 100%;
      min-height: 40px;
      border: 1px solid rgba(255,130,0,.20);
      border-radius: 14px;
      background: rgba(255,255,255,.94);
      color: var(--ink);
      padding: 0 12px;
      font-size: 13px;
      font-weight: 800;
      outline: none;
    }}
    .library-confirm-card.done {{
      border-color: rgba(21,115,71,.22);
      background: rgba(240,253,244,.78);
    }}
    .toast {{
      position: fixed;
      right: 24px;
      bottom: 24px;
      z-index: 60;
      min-width: 240px;
      max-width: min(420px, calc(100vw - 32px));
      padding: 14px 16px;
      border-radius: 18px;
      color: #FF8200;
      background: rgba(255,255,255,.84);
      border: 1px solid rgba(255,130,0,.16);
      box-shadow: 0 18px 38px rgba(249,115,0,.16);
      backdrop-filter: blur(16px);
      -webkit-backdrop-filter: blur(16px);
      opacity: 0;
      transform: translateY(10px);
      pointer-events: none;
      transition: opacity .22s ease, transform .22s ease;
    }}
    .toast.show {{
      opacity: 1;
      transform: translateY(0);
    }}
    .toast-title {{
      font-size: 14px;
      font-weight: 800;
      margin-bottom: 4px;
    }}
    .toast-copy {{
      font-size: 13px;
      line-height: 1.55;
    }}
    .choice-overlay {{
      position: fixed;
      inset: 0;
      display: none;
      align-items: center;
      justify-content: center;
      padding: 24px;
      background: rgba(255, 130, 0, 0.14);
      backdrop-filter: blur(10px);
      -webkit-backdrop-filter: blur(10px);
      z-index: 70;
    }}
    .choice-overlay.open {{
      display: flex;
    }}
    .choice-dialog {{
      width: min(460px, 100%);
      border-radius: 24px;
      border: 1px solid rgba(255,255,255,.82);
      background: rgba(255,255,255,.92);
      box-shadow: 0 24px 60px rgba(249,115,0,.18);
      padding: 22px;
      color: #FF8200;
    }}
    .choice-dialog h3 {{
      margin: 0 0 10px;
      font-size: 24px;
      line-height: 1.2;
    }}
    .choice-dialog p {{
      margin: 0 0 18px;
      font-size: 14px;
      line-height: 1.7;
      color: #FF8200;
      opacity: .9;
    }}
    .choice-actions {{
      display: flex;
      flex-wrap: wrap;
      gap: 12px;
      justify-content: flex-end;
    }}
    .editor-shell {{
      margin-bottom: 14px;
      border: 1px solid rgba(255,130,0,.16);
      border-radius: 16px;
      background: rgba(255,255,255,.72);
      padding: 14px;
      display: flex;
      flex-direction: column;
      gap: 12px;
    }}
    .editor-disclosure {{
      margin-bottom: 14px;
      border: 1px solid rgba(255,130,0,.16);
      border-radius: 16px;
      background: rgba(255,255,255,.62);
      overflow: hidden;
    }}
    .editor-summary {{
      list-style: none;
      cursor: pointer;
      position: relative;
      display: flex;
      align-items: center;
      justify-content: center;
      padding: 14px 16px;
      color: #FF8200;
      font-size: 15px;
      font-weight: 800;
      letter-spacing: 0;
      text-transform: none;
      text-align: center;
    }}
    .editor-summary::-webkit-details-marker {{
      display: none;
    }}
    .editor-summary::after {{
      content: "▾";
      position: absolute;
      right: 16px;
      top: 50%;
      transform: translateY(-50%);
      font-size: 16px;
      line-height: 1;
    }}
    .editor-disclosure[open] .editor-summary::after {{
      content: "▴";
    }}
    .editor-summary-title {{
      display: inline-flex;
      align-items: center;
      justify-content: center;
      width: 100%;
      padding-right: 22px;
      box-sizing: border-box;
    }}
    .editor-summary-copy {{
      font-size: 12px;
      font-weight: 600;
      letter-spacing: 0;
      text-transform: none;
      opacity: .82;
      margin-left: auto;
    }}
    .editor-language-note {{
      margin-left: 6px;
      color: rgba(255,130,0,.78);
    }}
    .editor-grid {{
      display: grid;
      grid-template-columns: 1fr;
      gap: 12px;
    }}
    .editor-field {{
      display: flex;
      flex-direction: column;
      gap: 8px;
    }}
    .editor-label {{
      font-size: 12px;
      font-weight: 800;
      letter-spacing: .04em;
      text-transform: uppercase;
      color: #FF8200;
      margin: 0;
    }}
    .editor-input, .editor-textarea {{
      width: 100%;
      border: 1px solid rgba(255,130,0,.14);
      border-radius: 14px;
      background: rgba(255,255,255,.92);
      color: #FF8200;
      padding: 12px 14px;
      font-size: 14px;
      line-height: 1.6;
      outline: none;
    }}
    .editor-textarea {{
      min-height: 96px;
      resize: vertical;
    }}
    .editor-draft-textarea {{
      min-height: 760px;
      font-family: "SFMono-Regular", Menlo, "Noto Sans SC", monospace;
      line-height: 1.75;
      white-space: pre-wrap;
    }}
    .structured-editor {{
      gap: 16px;
      background: #f6f7fb;
      border-color: rgba(229,231,235,.95);
    }}
    .structured-editor-section {{
      background: #fff;
      border: 1px solid #e5e7eb;
      border-radius: 16px;
      padding: 18px;
      box-shadow: 0 2px 10px rgba(15,23,42,.04);
    }}
    .structured-editor-section h5 {{
      margin: 0 0 12px;
      color: #111827;
      font-size: 20px;
      line-height: 1.25;
      font-weight: 900;
    }}
    .structured-editor-input,
    .structured-editor-textarea {{
      width: 100%;
      border: 1px solid #dfe3ea;
      border-radius: 12px;
      background: #fff;
      color: #111827;
      padding: 12px 14px;
      font-size: 16px;
      line-height: 1.75;
      outline: none;
    }}
    .structured-editor-input:focus,
    .structured-editor-textarea:focus {{
      border-color: rgba(255,130,0,.45);
      box-shadow: 0 0 0 3px rgba(255,130,0,.10);
    }}
    .structured-title-input {{
      font-size: 26px;
      line-height: 1.25;
      font-weight: 900;
      letter-spacing: 0;
    }}
    .structured-editor-textarea {{
      min-height: 118px;
      resize: vertical;
    }}
    .structured-insight-grid {{
      display: grid;
      grid-template-columns: repeat(2, minmax(0, 1fr));
      gap: 12px;
    }}
    .structured-insight {{
      border: 1px solid #e5e7eb;
      border-radius: 14px;
      background: #fbfdff;
      padding: 12px;
      display: grid;
      gap: 8px;
    }}
    .structured-insight-head {{
      display: flex;
      align-items: center;
      justify-content: space-between;
      gap: 10px;
    }}
    .structured-insight-label {{
      min-height: 42px;
      font-weight: 900;
    }}
    .structured-delete-btn {{
      flex: 0 0 auto;
      min-height: 34px;
      padding: 0 12px;
      border-radius: 999px;
      border: 1px solid rgba(217, 87, 42, .16);
      background: #fff4f1;
      color: #D9572A;
      font-size: 12px;
      font-weight: 900;
      cursor: pointer;
    }}
    .structured-delete-btn:hover {{
      border-color: rgba(217, 87, 42, .32);
      background: #ffe9e3;
    }}
    .replacement-picker {{
      display: grid;
      gap: 12px;
    }}
    .replacement-option-grid {{
      display: grid;
      grid-template-columns: repeat(auto-fit, minmax(220px, 1fr));
      gap: 10px;
    }}
    .replacement-option {{
      appearance: none;
      border: 1px solid rgba(255,130,0,.18);
      border-radius: 14px;
      background: linear-gradient(180deg, rgba(255,248,238,.95), rgba(255,255,255,.92));
      color: #111827;
      padding: 12px;
      text-align: left;
      cursor: pointer;
      min-height: 116px;
      display: flex;
      flex-direction: column;
      gap: 8px;
      transition: transform .16s ease, border-color .16s ease, box-shadow .16s ease;
    }}
    .replacement-option:hover {{
      transform: translateY(-1px);
      border-color: rgba(255,130,0,.36);
      box-shadow: 0 10px 24px rgba(249,115,0,.12);
    }}
    .replacement-option:disabled {{
      opacity: .58;
      cursor: wait;
      transform: none;
    }}
    .replacement-option strong {{
      color: #FF8200;
      font-size: 15px;
      line-height: 1.35;
    }}
    .replacement-option span {{
      color: #4b5563;
      font-size: 13px;
      line-height: 1.55;
    }}
    .replacement-empty {{
      border: 1px dashed rgba(255,130,0,.24);
      border-radius: 14px;
      padding: 14px;
      color: #FF8200;
      background: rgba(255,248,238,.66);
      font-weight: 800;
    }}
    .replacement-custom {{
      display: grid;
      grid-template-columns: minmax(0, 1fr) auto;
      gap: 10px;
      align-items: center;
    }}
    .storyboard-panel {{
      display: grid;
      gap: 14px;
    }}
    .storyboard-preview-wrap {{
      border: 1px solid rgba(255,130,0,.16);
      border-radius: 16px;
      overflow: hidden;
      background: rgba(255,255,255,.92);
      aspect-ratio: 1 / 1;
      display: grid;
      place-items: center;
    }}
    .storyboard-preview {{
      display: block;
      width: 100%;
      height: 100%;
      object-fit: cover;
      background: #fff;
    }}
    .storyboard-empty {{
      border: 1px dashed rgba(255,130,0,.22);
      border-radius: 16px;
      padding: 18px;
      font-size: 13px;
      line-height: 1.7;
      color: #935F14;
      background: rgba(255,248,238,.6);
    }}
    .storyboard-actions {{
      display: flex;
      flex-wrap: wrap;
      gap: 10px;
    }}
    .storyboard-prompt {{
      min-height: 110px;
    }}
    .structured-table-wrap {{
      overflow-x: auto;
      border: 1px solid #d9f0fb;
      border-radius: 14px;
      background: #fff;
    }}
    .structured-script-table {{
      width: 100%;
      min-width: 980px;
      border-collapse: collapse;
      table-layout: fixed;
    }}
    .structured-script-table th,
    .structured-script-table td {{
      border: 1px solid #e5e7eb;
      vertical-align: top;
      padding: 10px;
    }}
    .structured-script-table th {{
      background: #d9f0fb;
      color: #111827;
      font-size: 14px;
      font-weight: 900;
      text-align: center;
    }}
    .structured-script-table th:nth-child(1),
    .structured-script-table td:nth-child(1) {{
      width: 11%;
    }}
    .structured-script-table th:nth-child(2),
    .structured-script-table td:nth-child(2) {{
      width: 22%;
    }}
    .structured-script-table th:nth-child(3),
    .structured-script-table td:nth-child(3) {{
      width: 31%;
    }}
    .structured-script-table th:nth-child(4),
    .structured-script-table td:nth-child(4) {{
      width: 28%;
    }}
    .structured-script-table th:nth-child(5),
    .structured-script-table td:nth-child(5) {{
      width: 8%;
      text-align: center;
      vertical-align: middle;
    }}
    .structured-row-actions {{
      padding: 10px 6px;
    }}
    .structured-cell-input,
    .structured-cell-textarea {{
      width: 100%;
      border: 0;
      background: transparent;
      color: #111827;
      font-size: 15px;
      line-height: 1.7;
      outline: none;
      resize: vertical;
    }}
    .structured-cell-input {{
      min-height: 38px;
      text-align: center;
      font-weight: 800;
      color: #FF8200;
    }}
    .structured-cell-textarea {{
      min-height: 120px;
    }}
    .structured-editor-actions {{
      position: sticky;
      bottom: 0;
      z-index: 3;
      margin-top: 2px;
      padding: 12px 0 0;
      background: linear-gradient(180deg, rgba(246,247,251,0), #f6f7fb 34%);
    }}
    .editor-row-card {{
      border: 1px solid rgba(255,130,0,.12);
      border-radius: 14px;
      background: rgba(255,244,232,.56);
      padding: 12px;
      display: flex;
      flex-direction: column;
      gap: 12px;
    }}
    .editor-row-title {{
      font-size: 13px;
      font-weight: 800;
      color: #FF8200;
    }}
    .editor-row-head {{
      display: flex;
      align-items: center;
      justify-content: space-between;
      gap: 12px;
      flex-wrap: wrap;
    }}
    .editor-row-remove {{
      flex-shrink: 0;
    }}
    .review-shell {{
      height: 100%;
      margin-bottom: 0;
      border: 0;
      border-radius: 0;
      background: transparent;
      padding: 10px;
      display: flex;
      flex-direction: column;
      gap: 8px;
    }}
    .koko-chat-head {{
      display: flex;
      align-items: flex-start;
      gap: 10px;
      color: var(--ink);
    }}
    .koko-chat-avatar {{
      width: 28px;
      height: 28px;
      border-radius: 999px;
      display: grid;
      place-items: center;
      flex: 0 0 auto;
      background: linear-gradient(135deg, var(--brand), var(--brand-deep));
      color: #fff;
      font-size: 12px;
      font-weight: 900;
    }}
    .koko-chat-title {{
      font-size: 14px;
      font-weight: 900;
      line-height: 1.35;
    }}
    .koko-chat-subtitle {{
      display: none;
    }}
    .review-progress {{
      display: flex;
      flex-direction: column;
      gap: 10px;
      border: 1px solid rgba(255,130,0,.12);
      border-radius: 14px;
      background: rgba(255,255,255,.76);
      padding: 12px;
    }}
    .review-progress-top {{
      display: flex;
      align-items: center;
      justify-content: space-between;
      gap: 12px;
      font-size: 12px;
      font-weight: 700;
      color: #FF8200;
    }}
    .review-progress-rail {{
      height: 8px;
      border-radius: 999px;
      background: rgba(255,130,0,.10);
      overflow: hidden;
    }}
    .review-progress-fill {{
      height: 100%;
      border-radius: inherit;
      background: linear-gradient(135deg, var(--brand), var(--brand-deep));
      transition: width .24s ease;
    }}
    .review-progress-fill.failed {{
      background: linear-gradient(135deg, #D9572A, #B43A20);
    }}
    .review-steps {{
      display: grid;
      grid-template-columns: repeat(4, minmax(0, 1fr));
      gap: 8px;
    }}
    .review-step {{
      border-radius: 999px;
      border: 1px solid rgba(255,130,0,.12);
      background: rgba(255,255,255,.72);
      padding: 8px 10px;
      font-size: 11px;
      font-weight: 700;
      text-align: center;
      color: rgba(255,130,0,.78);
    }}
    .review-step.done {{
      background: rgba(255,130,0,.12);
      color: #FF8200;
    }}
    .review-step.active {{
      background: linear-gradient(135deg, rgba(255,130,0,.18), rgba(249,115,0,.16));
      border-color: rgba(255,130,0,.22);
      color: #FF8200;
    }}
    .review-step.failed {{
      border-color: rgba(180,35,24,.22);
      background: rgba(180,35,24,.08);
      color: #B43A20;
    }}
    .review-note {{
      font-size: 13px;
      line-height: 1.6;
      color: #FF8200;
      opacity: .92;
    }}
    .review-mode-toggle {{
      display: grid;
      grid-template-columns: repeat(2, minmax(0, 1fr));
      gap: 6px;
    }}
    .review-mode-option,
    .chat-mode-option {{
      min-height: 44px;
      border: 1px solid rgba(255,130,0,.18);
      border-radius: 999px;
      background: rgba(255,255,255,.76);
      color: #FF8200;
      display: flex;
      align-items: center;
      justify-content: center;
      font-size: 13px;
      font-weight: 800;
      cursor: pointer;
    }}
    .review-mode-option input {{
      position: absolute;
      opacity: 0;
      pointer-events: none;
    }}
    .review-mode-option:has(input:checked),
    .chat-mode-option.active {{
      background: rgba(255,130,0,.14);
      border-color: rgba(255,130,0,.34);
      box-shadow: inset 0 0 0 1px rgba(255,130,0,.12);
    }}
    .chat-mode-option[disabled] {{
      opacity: .52;
      cursor: not-allowed;
    }}
    .assistant-sidebar .chat-mode-option {{
      min-height: 36px;
      padding: 0 6px;
      font-size: 12px;
      line-height: 1.2;
    }}
    .assistant-sidebar .chat-mode-option[data-chat-mode-choice="recheck"] {{
      grid-column: 1 / -1;
    }}
    .chat-mode-note {{
      display: none;
    }}
    .chat-compose {{
      margin-top: auto;
      border-top: 1px solid rgba(255,130,0,.12);
      padding-top: 8px;
      background: linear-gradient(180deg, rgba(255,255,255,0), rgba(255,255,255,.82) 22%);
    }}
    .chat-compose .editor-textarea {{
      min-height: 92px;
      max-height: 180px;
      color: var(--ink);
      background: rgba(255,255,255,.96);
      font-size: 12px;
      padding: 10px;
      border-radius: 12px;
    }}
    .assistant-sidebar .link-row {{
      gap: 6px;
      margin-bottom: 0;
    }}
    .assistant-sidebar .action-link {{
      width: 100%;
      min-height: 38px;
      padding: 9px 10px;
      font-size: 12px;
    }}
    .item-sections {{
      display: flex;
      flex-direction: column;
      gap: 14px;
      margin-bottom: 14px;
    }}
    .item-actions-shell {{
      border: 1px solid rgba(255,130,0,.16);
      border-radius: 16px;
      background: rgba(255,255,255,.72);
      padding: 14px;
    }}
    .item-card iframe {{
      width: 100%;
      min-height: 680px;
      border: 1px solid rgba(255,130,0,.14);
      border-radius: 16px;
      background: #fff;
    }}
    .library-shell {{
      padding: 24px;
    }}
    .library-wrap {{
      width: min(1320px, 100%);
      margin: 0 auto;
      border-radius: 34px;
      overflow: hidden;
      border: 1px solid rgba(255,255,255,.85);
      box-shadow: var(--shadow);
      background:
        radial-gradient(circle at 12% 16%, rgba(255,130,0,.22), rgba(255,130,0,0) 22%),
        radial-gradient(circle at 86% 18%, rgba(249,115,0,.18), rgba(249,115,0,0) 22%),
        radial-gradient(circle at 70% 62%, rgba(255,244,232,.82), rgba(255,244,232,0) 24%),
        linear-gradient(180deg, #FFC792 0%, #FFF0DE 44%, #FFFFFF 100%);
      padding: 28px;
    }}
    .library-grid {{
      display: grid;
      grid-template-columns: repeat(auto-fit, minmax(280px, 1fr));
      gap: 16px;
      margin-top: 20px;
    }}
    .library-card {{
      border: 1px solid rgba(255,130,0,.16);
      border-radius: 22px;
      background: rgba(255,255,255,.82);
      padding: 18px;
      display: flex;
      flex-direction: column;
      gap: 12px;
    }}
    .library-card h3 {{
      margin: 0;
      font-size: 20px;
      line-height: 1.12;
      letter-spacing: -.03em;
    }}
    .library-card p {{
      margin: 0;
      line-height: 1.65;
      font-size: 14px;
      color: #FF8200;
    }}
    .library-topbar {{
      display: flex;
      align-items: flex-start;
      justify-content: space-between;
      gap: 20px;
      flex-wrap: wrap;
    }}
    .library-topbar h1 {{
      margin: 0;
      font-size: clamp(3rem, 7vw, 5.4rem);
      letter-spacing: -.08em;
      line-height: .88;
    }}
    .library-stats {{
      display: flex;
      flex-wrap: wrap;
      gap: 10px;
      margin-top: 12px;
    }}
    .status-box iframe {{
      width: 100%;
      min-height: 620px;
      border: 1px solid var(--line);
      border-radius: 22px;
      background: #fff;
      margin-top: 14px;
      box-shadow: inset 0 0 0 1px rgba(255,255,255,.9);
    }}
    .status {{
      display: inline-flex;
      align-items: center;
      padding: 6px 10px;
      border-radius: 999px;
      font-size: 12px;
      text-transform: uppercase;
      letter-spacing: .08em;
      font-weight: 700;
    }}
    .status-queued, .status-running {{ background: rgba(147,95,20,.12); color: var(--wait); }}
    .status-completed {{ background: rgba(21,115,71,.12); color: var(--ok); }}
    .status-failed {{ background: rgba(180,35,24,.12); color: var(--err); }}
    .result-link {{
      color: #FF8200;
      text-decoration: none;
      font-weight: 700;
      margin-right: 16px;
    }}
    code {{
      font-family: "SFMono-Regular", Menlo, monospace;
      background: rgba(255,130,0,.08);
      padding: 2px 6px;
      border-radius: 8px;
    }}
    .brand-icon {{
      width: 38px;
      height: 38px;
      border-radius: 12px;
      background: var(--brand-soft);
      display: inline-flex;
      align-items: center;
      justify-content: center;
      overflow: hidden;
      flex: 0 0 auto;
    }}
    .brand-icon img {{
      width: 28px;
      height: 28px;
      object-fit: contain;
      object-position: left center;
      transform: scale(1.65);
    }}
    .studio-shell {{
      display: grid;
      grid-template-columns: 220px minmax(0, 1fr);
      min-height: 100vh;
    }}
    .studio-sidebar {{
      position: sticky;
      top: 0;
      align-self: start;
      min-height: 100vh;
      padding: 24px 16px;
      border-right: 1px solid rgba(255,255,255,.82);
      background: rgba(255,255,255,.58);
      backdrop-filter: blur(22px);
      -webkit-backdrop-filter: blur(22px);
      display: flex;
      flex-direction: column;
      gap: 22px;
      z-index: 5;
    }}
    .studio-brand {{
      display: flex;
      align-items: center;
      gap: 14px;
      padding: 8px 10px 12px;
      text-decoration: none;
      color: #1f1f1f;
      font-size: 16px;
      font-weight: 800;
    }}
    .studio-brand img {{
      width: 112px;
      height: auto;
      object-fit: contain;
    }}
    .studio-side-nav {{
      display: flex;
      flex-direction: column;
      gap: 10px;
    }}
    .studio-tab-link {{
      display: flex;
      align-items: center;
      gap: 12px;
      width: 100%;
      border: 1px solid transparent;
      background: rgba(255,255,255,.48);
      color: rgba(31,31,31,.76);
      padding: 14px 16px;
      border-radius: 18px;
      text-decoration: none;
      font-size: 15px;
      font-weight: 700;
      transition: border-color .18s ease, background .18s ease, color .18s ease;
    }}
    .studio-tab-link:hover {{
      border-color: rgba(255,130,0,.18);
      color: #FF8200;
    }}
    .studio-tab-link.active {{
      color: #FF8200;
      background: rgba(255,255,255,.88);
      border-color: rgba(255,130,0,.18);
      box-shadow: 0 10px 24px rgba(249,115,0,.08);
    }}
    .studio-tab-icon {{
      font-size: 18px;
      line-height: 1;
    }}
    .studio-side-meta {{
      margin-top: auto;
      border-radius: 20px;
      border: 1px solid rgba(255,130,0,.12);
      background: rgba(255,255,255,.68);
      padding: 14px;
    }}
    .studio-side-meta strong {{
      display: block;
      font-size: 14px;
      color: #1f1f1f;
      margin-bottom: 6px;
    }}
    .studio-side-meta p {{
      margin: 0;
      font-size: 12px;
      line-height: 1.65;
      color: rgba(31,31,31,.62);
    }}
    .studio-main {{
      padding: 28px 34px 34px;
      min-width: 0;
    }}
    .studio-shell:has(.assistant-sidebar) .studio-main {{
      padding-right: 264px;
    }}
    .studio-page-title {{
      margin: 0 0 6px;
      font-size: clamp(2.2rem, 5vw, 3.2rem);
      line-height: .96;
      letter-spacing: -.06em;
      color: #1f1f1f;
    }}
    .studio-page-copy {{
      margin: 0;
      font-size: 14px;
      line-height: 1.7;
      color: rgba(31,31,31,.62);
      max-width: 52em;
    }}
    .studio-topbar {{
      display: none;
      justify-content: flex-end;
      align-items: center;
      gap: 16px;
      margin-bottom: 8px;
    }}
    .studio-top-icon {{
      width: 38px;
      height: 38px;
      border: 1px solid rgba(255,130,0,.12);
      border-radius: 999px;
      background: rgba(255,255,255,.74);
      color: #1f1f1f;
      display: inline-flex;
      align-items: center;
      justify-content: center;
      box-shadow: 0 12px 28px rgba(249,115,0,.08);
    }}
    .studio-user-chip {{
      display: inline-flex;
      align-items: center;
      gap: 10px;
      border: 1px solid rgba(255,130,0,.12);
      border-radius: 999px;
      background: rgba(255,255,255,.78);
      padding: 6px 10px 6px 6px;
      color: #1f1f1f;
      font-size: 14px;
      font-weight: 800;
      box-shadow: 0 12px 28px rgba(249,115,0,.08);
    }}
    .studio-user-avatar {{
      width: 34px;
      height: 34px;
      border-radius: 999px;
      background:
        radial-gradient(circle at 50% 35%, #ffd7ba 0 19%, transparent 20%),
        radial-gradient(circle at 50% 77%, #2f4050 0 28%, transparent 29%),
        linear-gradient(135deg, #fff5eb, #ffd3a7);
      border: 1px solid rgba(255,130,0,.18);
    }}
    .studio-hero-header {{
      position: relative;
      min-height: 0;
      display: block;
      margin: -12px auto 18px;
      overflow: visible;
      border-radius: 26px;
      background: linear-gradient(180deg, rgba(255,249,242,.92), rgba(255,240,224,.58));
      box-shadow: 0 18px 46px rgba(249,115,0,.08);
      max-width: 1120px;
    }}
    .studio-hero-banner {{
      display: block;
      width: 100%;
      pointer-events: none;
    }}
    .studio-hero-banner img {{
      width: 100%;
      aspect-ratio: 1120 / 378;
      height: auto;
      display: block;
      border-radius: 26px;
    }}
    .studio-title-art {{
      width: min(670px, 100%);
      height: auto;
      display: block;
      mix-blend-mode: multiply;
      filter: saturate(1.04);
    }}
    .studio-title-fallback {{
      display: none;
      margin: 0;
      font-size: clamp(3rem, 7vw, 6.4rem);
      line-height: .9;
      letter-spacing: -.06em;
      color: #ff5f00;
    }}
    .studio-hero-art {{
      justify-self: end;
      align-self: start;
      width: min(520px, 100%);
      margin-right: 16px;
      transform: translateY(-46px);
      pointer-events: none;
    }}
    .studio-hero-art img {{
      width: 100%;
      height: auto;
      display: block;
      filter: drop-shadow(0 28px 42px rgba(249,115,0,.14));
    }}
    .studio-panel {{
      display: none;
      margin-top: 22px;
    }}
    .studio-panel.active {{
      display: block;
    }}
    .studio-card {{
      border-radius: 28px;
      border: 1px solid rgba(255,255,255,.82);
      background: rgba(255,255,255,.78);
      box-shadow: 0 22px 48px rgba(249,115,0,.08);
      overflow: hidden;
    }}
    .studio-task-card {{
      position: relative;
      border-radius: 30px;
      border-color: rgba(255,130,0,.12);
      background:
        radial-gradient(circle at 88% 0%, rgba(255,130,0,.12), transparent 28%),
        linear-gradient(180deg, rgba(255,255,255,.92), rgba(255,255,255,.78));
      box-shadow: 0 24px 60px rgba(249,115,0,.11);
      margin-top: -14px;
    }}
    .studio-card-head {{
      display: flex;
      align-items: center;
      justify-content: space-between;
      gap: 14px;
      flex-wrap: wrap;
      padding: 24px 24px 0;
    }}
    .studio-card-head h2 {{
      margin: 0;
      font-size: 28px;
      letter-spacing: -.04em;
      color: #1f1f1f;
    }}
    .studio-task-card .studio-card-head h2 {{
      display: inline-flex;
      align-items: center;
      min-height: 58px;
      padding: 0 18px;
      border-radius: 18px;
      color: #ff5f00;
      background: linear-gradient(180deg, rgba(255,248,238,.96), rgba(255,255,255,.78));
      box-shadow: inset 0 -14px 18px rgba(255,130,0,.06);
      font-size: 30px;
      font-weight: 900;
    }}
    .studio-card-head p {{
      margin: 10px auto 0;
      font-size: 14px;
      line-height: 1.7;
      color: rgba(31,31,31,.62);
      max-width: 52em;
      text-align: center;
    }}
    .studio-task-card .studio-card-head {{
      justify-content: center;
      text-align: center;
    }}
    .studio-kpis {{
      display: grid;
      grid-template-columns: repeat(3, minmax(0, 1fr));
      gap: 16px;
      padding: 0 24px 24px;
      margin-top: 18px;
    }}
    .studio-kpi {{
      border-radius: 22px;
      border: 1px solid rgba(255,130,0,.12);
      background: rgba(255,248,238,.76);
      padding: 18px;
    }}
    .studio-kpi strong {{
      display: block;
      font-size: 14px;
      color: #FF8200;
      margin-bottom: 10px;
    }}
    .studio-kpi span {{
      display: block;
      font-size: 28px;
      color: #1f1f1f;
      font-weight: 800;
      letter-spacing: -.04em;
      margin-bottom: 10px;
    }}
    .studio-kpi p {{
      margin: 0;
      font-size: 13px;
      line-height: 1.6;
      color: rgba(31,31,31,.58);
    }}
    .studio-helper-chips {{
      display: flex;
      flex-wrap: wrap;
      gap: 10px;
      margin-top: 14px;
    }}
    .studio-helper-chip {{
      display: inline-flex;
      align-items: center;
      gap: 8px;
      border: 1px solid rgba(255,130,0,.14);
      border-radius: 11px;
      background: rgba(255,244,232,.78);
      color: #a55710;
      padding: 9px 12px;
      font-size: 13px;
      font-weight: 800;
    }}
    .studio-overview {{
      margin-top: 18px;
      border-radius: 26px;
      border: 1px solid rgba(255,130,0,.10);
      background: rgba(255,255,255,.78);
      padding: 22px 24px 24px;
      box-shadow: 0 18px 44px rgba(249,115,0,.08);
    }}
    .studio-overview h2 {{
      margin: 0 0 16px;
      color: #1f1f1f;
      font-size: 22px;
      letter-spacing: -.03em;
    }}
    .studio-stat-grid {{
      display: grid;
      grid-template-columns: repeat(4, minmax(0, 1fr));
      gap: 16px;
    }}
    .studio-stat-card {{
      min-height: 106px;
      display: grid;
      grid-template-columns: 56px 1fr;
      gap: 12px;
      align-items: center;
      border: 1px solid rgba(255,130,0,.10);
      border-radius: 18px;
      background: linear-gradient(135deg, rgba(255,250,245,.98), rgba(255,244,232,.72));
      padding: 14px;
    }}
    .studio-stat-card:nth-child(2) {{
      border-color: rgba(116,107,255,.16);
      background: linear-gradient(135deg, rgba(255,255,255,.98), rgba(244,242,255,.74));
    }}
    .studio-stat-card:nth-child(3) {{
      border-color: rgba(34,173,96,.18);
      background: linear-gradient(135deg, rgba(255,255,255,.98), rgba(239,255,245,.78));
    }}
    .studio-stat-card:nth-child(4) {{
      border-color: rgba(240,72,72,.18);
      background: linear-gradient(135deg, rgba(255,255,255,.98), rgba(255,241,241,.76));
    }}
    .studio-stat-icon {{
      width: 54px;
      height: 54px;
      border-radius: 18px;
      display: inline-flex;
      align-items: center;
      justify-content: center;
      color: #fff;
      background: linear-gradient(135deg, #ffb34d, #ff5f00);
      box-shadow: 0 12px 26px rgba(249,115,0,.22);
      font-size: 24px;
      font-weight: 900;
    }}
    .studio-stat-card:nth-child(2) .studio-stat-icon {{ background: linear-gradient(135deg, #ffb14d, #ff8b00); }}
    .studio-stat-card:nth-child(3) .studio-stat-icon {{ background: linear-gradient(135deg, #8be0a7, #16a45a); }}
    .studio-stat-card:nth-child(4) .studio-stat-icon {{ background: linear-gradient(135deg, #ff8b8b, #f04444); }}
    .studio-stat-card strong {{
      display: block;
      color: #1f1f1f;
      font-size: 13px;
      margin-bottom: 2px;
    }}
    .studio-stat-card b {{
      display: inline-block;
      color: #101010;
      font-size: 30px;
      line-height: 1;
      margin-right: 8px;
    }}
    .studio-stat-card span {{
      color: rgba(31,31,31,.50);
      font-size: 12px;
      font-weight: 800;
    }}
    .studio-stat-meta {{
      display: block;
      margin-top: 4px;
      color: rgba(31,31,31,.50);
      font-size: 12px;
      font-weight: 800;
    }}
    .studio-iframe {{
      width: 100%;
      min-height: 1280px;
      border: 0;
      display: block;
      background: transparent;
    }}
    .composer-block {{
      padding: 24px;
      border-bottom: 1px solid rgba(255,130,0,.12);
    }}
    .studio-placeholder-grid {{
      display: grid;
      grid-template-columns: 1.2fr .8fr;
      gap: 18px;
      padding: 0 24px 24px;
    }}
    .studio-placeholder-box {{
      border-radius: 22px;
      border: 1px solid rgba(255,130,0,.12);
      background: rgba(255,248,238,.76);
      padding: 20px;
    }}
    .studio-placeholder-box h3 {{
      margin: 0 0 12px;
      font-size: 22px;
      color: #1f1f1f;
      letter-spacing: -.03em;
    }}
    .studio-placeholder-box p,
    .studio-placeholder-box li {{
      font-size: 14px;
      line-height: 1.75;
      color: rgba(31,31,31,.62);
    }}
    .studio-placeholder-box ul {{
      margin: 10px 0 0;
      padding-left: 18px;
    }}
    @media (max-width: 1080px) {{
      .studio-shell {{
        grid-template-columns: 1fr;
      }}
      .studio-sidebar {{
        position: static;
        min-height: auto;
        border-right: 0;
        border-bottom: 1px solid rgba(255,255,255,.82);
      }}
      .studio-kpis,
      .studio-placeholder-grid,
      .studio-stat-grid,
      .studio-hero-header {{
        grid-template-columns: 1fr;
      }}
      .studio-hero-art {{
        justify-self: start;
        width: min(360px, 100%);
      }}
      .result-workbench {{
        grid-template-columns: 1fr;
      }}
      .primary-review-strip {{
        grid-template-columns: 1fr;
      }}
      .video-verify-pane {{
        position: static;
      }}
      .video-verify-layout,
      .ops-pane {{
        grid-template-columns: 1fr;
      }}
      .source-video-frame {{
        width: min(420px, 100%);
        margin: 0 auto;
        max-height: 720px;
      }}
      .hero-copy {{
        grid-template-columns: 1fr;
        min-height: auto;
        gap: 18px;
      }}
      .hero-side {{
        align-self: start;
        justify-self: start;
      }}
    }}
    @media (max-width: 720px) {{
      .studio-shell:has(.assistant-sidebar) .studio-main {{
        padding-right: 16px;
      }}
      .assistant-sidebar {{
        position: static;
        width: auto;
        height: auto;
      }}
      .assistant-sidebar .review-chat-log {{
        max-height: 420px;
      }}
      .studio-main {{
        padding: 16px;
      }}
      .studio-topbar {{
        justify-content: flex-start;
      }}
      .studio-hero-copy {{
        margin-left: 0;
      }}
      .studio-title-art {{
        max-width: 96vw;
      }}
      .composer-block,
      .studio-card-head,
      .studio-kpis,
      .studio-placeholder-grid {{
        padding-left: 16px;
        padding-right: 16px;
      }}
      .hero-shell {{
        padding: 12px;
      }}
      .hero-panel {{
        min-height: calc(100vh - 24px);
      }}
      .hero-stage {{
        padding: 16px 12px 74px;
      }}
      .hero-copy {{
        padding: 18px 14px 88px;
        min-height: calc(100vh - 152px);
      }}
      .work-shell {{
        padding: 16px 14px 24px;
      }}
      .brandbar {{
        padding: 14px 14px 0;
        flex-direction: column;
        align-items: stretch;
      }}
      .navpill {{
        justify-content: center;
        flex-wrap: wrap;
      }}
      .composer {{
        padding: 18px;
      }}
      .composer-head {{
        flex-direction: column;
      }}
      .hero-left h1 {{ font-size: 5rem; }}
      .hero-left h1 .koko-k {{
        font-size: 1.16em;
      }}
      .lede {{
        font-size: 15px;
        max-width: 24ch;
      }}
      .progress-top {{
        align-items: flex-start;
      }}
      .progress-percent {{
        font-size: 16px;
      }}
      .step-list {{ grid-template-columns: repeat(4, minmax(0, 1fr)); gap: 12px 0; }}
      .step-pill::before {{
        top: 18px;
      }}
      .step-pill:nth-child(4n+1)::before {{
        display: none;
      }}
      .queue-list {{ grid-template-columns: 1fr; }}
    }}
  </style>
</head>
<body>
  <main class="studio-shell">
    <aside class="studio-sidebar">
      <a class="studio-brand" href="/">
        <img src="/brand/kwai-wordmark.svg" alt="Kwai" />
      </a>
      <nav class="studio-side-nav" aria-label="Koko 内容中台导航">
        <a class="studio-tab-link" href="/" data-nav-kind="home"><span class="studio-tab-icon">⌂</span><span>返回落地页</span></a>
        <a class="studio-tab-link" href="#filter-panel" data-panel-target="filter-panel"><span class="studio-tab-icon">⌕</span><span>视频筛选</span></a>
        <a class="studio-tab-link active" href="#split-panel" data-panel-target="split-panel"><span class="studio-tab-icon">▶</span><span>视频拆解</span></a>
        <a class="studio-tab-link" href="#understanding-panel" data-panel-target="understanding-panel"><span class="studio-tab-icon">◌</span><span>视频理解</span></a>
        <a class="studio-tab-link" href="#translate-panel" data-panel-target="translate-panel"><span class="studio-tab-icon">◉</span><span>葡语转译</span></a>
        <a class="studio-tab-link" href="#stats-panel" data-panel-target="stats-panel"><span class="studio-tab-icon">▥</span><span>数据看板</span></a>
        <a class="studio-tab-link" href="/library"><span class="studio-tab-icon">☰</span><span>脚本管理</span></a>
        <a class="studio-tab-link" href="/creator-admin"><span class="studio-tab-icon">★</span><span>Creator 运营</span></a>
      </nav>
      <div class="studio-side-meta">
        <strong>给运营看的内容中台</strong>
        <p>围绕视频筛选、视频拆解和数据看板，支持优质内容筛选、脚本沉淀与作者投喂。</p>
      </div>
    </aside>
    <section class="studio-main">
      <section class="studio-hero-header" aria-label="Koko 内容中台">
        <div class="studio-hero-banner" aria-hidden="true">
          <img src="/brand/studio-hero-banner-shallow.png" alt="" />
        </div>
        <h1 class="studio-title-fallback">Koko 内容中台</h1>
      </section>

      <section id="filter-panel" class="studio-panel">
        <div class="studio-card">
          <div class="studio-card-head">
            <div>
              <h2>视频筛选</h2>
            </div>
          </div>
          <div class="composer-block">
            <div class="composer">
              <div class="composer-head">
                <div></div>
              </div>
              <label for="filter-input">批量视频输入</label>
              <textarea id="filter-input" placeholder="可直接粘贴聊天记录、文档或一批链接，系统会自动识别其中的 http / https 字段。&#10;&#10;例如：&#10;这是今天候选： https://www.kwai.com/@.../video/...&#10;还有这个 https://www.kwai.com/@.../video/..."></textarea>
              <div class="filter-upload-row">
                <input id="filter-file-input" type="file" accept=".xlsx,.csv,.tsv,.txt" />
                <span class="filter-upload-hint">支持上传 Excel、CSV、TSV、TXT。系统会自动抽取其中的链接。</span>
              </div>
              <div class="actions">
                <button id="filter-submit-btn">开始筛选剧情候选</button>
              </div>
            </div>
          </div>
          <div id="filter-status-box" class="status-box">
            <div class="status-empty">
              <div class="status-empty-title">筛选器已就绪。</div>
              <div class="status-empty-copy">贴入一批链接或上传表格后，Koko 会提取完整音频信息和三张关键帧，再输出通过三轮规则的视频。</div>
            </div>
          </div>
        </div>
      </section>

      <section id="split-panel" class="studio-panel active">
        <div class="studio-card studio-task-card">
          <div class="studio-card-head">
            <div>
              <h2>视频拆解任务中心</h2>
              <p>粘贴 Kwai 视频链接后，Koko 会批量拆解剧情结构、生成脚本表、分镜提示和可入库版本。</p>
            </div>
          </div>
          <div class="composer-block">
            <div class="composer">
              <div class="composer-head">
                <div></div>
              </div>
              <label for="video-url">视频链接 <span style="font-weight:500;color:rgba(31,31,31,.50)">（每行粘贴一个链接）</span></label>
              <textarea id="video-url" placeholder="每行粘贴一个链接&#10;https://www.kwai.com/@.../video/...&#10;https://www.kwai.com/@.../video/..."></textarea>
              <div class="analysis-prompt-box">
                <label for="analysis-prompt">增加一些提示词 <span style="font-weight:500;color:rgba(31,31,31,.50)">（选填）</span></label>
                <textarea class="analysis-prompt" id="analysis-prompt" placeholder="例如：重点往夫妻误会方向拆；不要写成单纯出轨；强化结尾反转，适配巴西创作者复拍。"></textarea>
                <p class="analysis-prompt-note">填写后，这段方向会参与 Gemini 主分析、本地对照分析、最终整理和叙述逻辑审查；如果最终故事偏离方向，系统会要求复核。</p>
              </div>
              <div class="studio-helper-chips" aria-label="任务提示">
                <span class="studio-helper-chip">🔗 支持 Kwai 链接</span>
                <span class="studio-helper-chip">▤ 最多 50 条</span>
                <span class="studio-helper-chip">↵ 回车换行即可</span>
              </div>
              <div class="actions">
                <button id="submit-btn">开始拆解脚本 ▶</button>
                <button class="action-link" id="stop-all-btn" type="button" disabled>停止所有任务</button>
              </div>
            </div>
          </div>
          <div id="status-box" class="status-box">
            <div class="status-empty">
              <div class="status-empty-title">已就绪。</div>
              <div class="status-empty-copy">输入一个或多个视频链接后，系统会在这里实时显示拆解进度。</div>
            </div>
          </div>
        </div>
        <section class="studio-overview" aria-label="任务状态总览">
          <h2>任务状态总览</h2>
          <div class="studio-stat-grid">
            <article class="studio-stat-card"><span class="studio-stat-icon">▣</span><div><strong>待处理任务</strong><b data-studio-stat="queued">0</b><span class="studio-stat-meta" data-studio-stat-meta="queued">等待提交</span></div></article>
            <article class="studio-stat-card"><span class="studio-stat-icon">↺</span><div><strong>拆解中</strong><b data-studio-stat="running">0</b><span class="studio-stat-meta" data-studio-stat-meta="running">当前无任务</span></div></article>
            <article class="studio-stat-card"><span class="studio-stat-icon">✓</span><div><strong>已完成</strong><b data-studio-stat="completed">0</b><span class="studio-stat-meta" data-studio-stat-meta="completed">本次任务</span></div></article>
            <article class="studio-stat-card"><span class="studio-stat-icon">!</span><div><strong>失败任务</strong><b data-studio-stat="failed">0</b><span class="studio-stat-meta" data-studio-stat-meta="failed">本次任务</span></div></article>
          </div>
        </section>
      </section>

      <section id="understanding-panel" class="studio-panel">
        <div class="studio-card studio-task-card">
          <div class="studio-card-head">
            <div>
              <h2>视频理解</h2>
              <p>粘贴 Kwai 视频链接后，Koko 会沿用 Gemini 与 v3 双链路拆解并交叉校验，但这里只输出视频核心内容概述。</p>
            </div>
          </div>
          <div class="composer-block">
            <div class="composer">
              <div class="composer-head">
                <div></div>
              </div>
              <label for="understanding-url">视频链接 <span style="font-weight:500;color:rgba(31,31,31,.50)">（每行粘贴一个链接）</span></label>
              <textarea id="understanding-url" placeholder="每行粘贴一个链接&#10;https://www.kwai.com/@.../video/...&#10;https://www.kwai.com/@.../video/..."></textarea>
              <div class="studio-helper-chips" aria-label="理解任务提示">
                <span class="studio-helper-chip"> Gemini 主分析</span>
                <span class="studio-helper-chip"> v3 对照验证</span>
                <span class="studio-helper-chip"> 只输出内容概述</span>
              </div>
              <div class="actions">
                <button id="understanding-submit-btn">开始理解视频 ▶</button>
              </div>
            </div>
          </div>
          <div id="understanding-status-box" class="status-box">
            <div class="status-empty">
              <div class="status-empty-title">视频理解已就绪。</div>
              <div class="status-empty-copy">输入一个或多个视频链接后，系统会显示同款拆解进度，完成后只给出脚本内容概述。</div>
            </div>
          </div>
        </div>
      </section>

      <section id="translate-panel" class="studio-panel">
        <div class="studio-card">
          <div class="studio-card-head">
            <div>
              <h2>葡语转译</h2>
              <p>输入单条视频链接后，Koko 会尝试下载源视频、识别主体和原音频含义，去掉原声并合成 pt-BR 葡语音轨。</p>
            </div>
          </div>
          <div class="composer-block">
            <div class="composer">
              <div class="composer-head">
                <div></div>
              </div>
              <label for="translation-url">视频链接</label>
              <textarea id="translation-url" placeholder="https://www.gifshow.com/fw/photo/3xqupwscjsz7wm2"></textarea>
              <div class="actions">
                <button id="translation-submit-btn">生成葡语版视频</button>
              </div>
            </div>
          </div>
          <div id="translation-status-box" class="status-box">
            <div class="status-empty">
              <div class="status-empty-title">转译器已就绪。</div>
              <div class="status-empty-copy">需要公开视频可下载、GOOGLE_API_KEY 可用，以及 ffmpeg 或 imageio-ffmpeg 支持视频合成。</div>
            </div>
          </div>
        </div>
      </section>

      <section id="stats-panel" class="studio-panel">
        <div class="studio-card">
          <div class="studio-card-head">
            <div>
              <h2>数据看板</h2>
              <p>这里直接承接现有 Stats 页面，用于查看脚本生成、复盘和整稿编辑等行为数据。</p>
            </div>
          </div>
          <iframe class="studio-iframe" src="/stats" title="Koko 数据看板"></iframe>
        </div>
      </section>
    </section>
  </main>
  <div id="app-toast" class="toast" aria-hidden="true">
    <div class="toast-title" id="app-toast-title"></div>
    <div class="toast-copy" id="app-toast-copy"></div>
  </div>
  <div id="export-choice-overlay" class="choice-overlay" aria-hidden="true">
    <div class="choice-dialog" role="dialog" aria-modal="true" aria-labelledby="export-choice-title">
      <h3 id="export-choice-title">导出脚本</h3>
      <p>请选择要导出的版本。葡语版本会保持严格直译，不做概括改写。</p>
      <div class="choice-actions">
        <button class="action-link" id="export-choice-cancel" type="button">取消</button>
        <button class="action-link" id="export-choice-zh" type="button">导出中文版本</button>
        <button class="action-link primary" id="export-choice-pt" type="button">导出葡语版本</button>
      </div>
    </div>
  </div>

  <script>
    const filterInput = document.getElementById("filter-input");
    const filterFileInput = document.getElementById("filter-file-input");
    const filterSubmitBtn = document.getElementById("filter-submit-btn");
    const filterStatusBox = document.getElementById("filter-status-box");
    const translationInput = document.getElementById("translation-url");
    const translationSubmitBtn = document.getElementById("translation-submit-btn");
    const translationStatusBox = document.getElementById("translation-status-box");
    const videoInput = document.getElementById("video-url");
    const analysisPromptInput = document.getElementById("analysis-prompt");
    const submitBtn = document.getElementById("submit-btn");
    const stopAllBtn = document.getElementById("stop-all-btn");
    const statusBox = document.getElementById("status-box");
    const understandingInput = document.getElementById("understanding-url");
    const understandingSubmitBtn = document.getElementById("understanding-submit-btn");
    const understandingStatusBox = document.getElementById("understanding-status-box");
    const studioPanelLinks = Array.from(document.querySelectorAll("[data-panel-target]"));
    const studioPanels = Array.from(document.querySelectorAll(".studio-panel"));
    const heroPanel = document.querySelector(".hero-panel");
    const appToast = document.getElementById("app-toast");
    const appToastTitle = document.getElementById("app-toast-title");
    const appToastCopy = document.getElementById("app-toast-copy");
    const exportChoiceOverlay = document.getElementById("export-choice-overlay");
    const exportChoiceCancel = document.getElementById("export-choice-cancel");
    const exportChoiceZh = document.getElementById("export-choice-zh");
    const exportChoicePt = document.getElementById("export-choice-pt");
    let activeJobId = "";
    let activeReviewItemId = "";
    let jobPollTimer = null;
    let toastTimer = null;
    let pendingExportChoice = null;
    let restoringActiveJob = false;
    let restoreAttempts = 0;
    let pollInFlight = false;
    let queuedImmediateRepoll = false;
    let activePollController = null;
    const POLL_REQUEST_TIMEOUT_MS = 12000;
    const POLL_RECOVERY_DELAY_MS = 1500;
    const reviewTracker = Object.create(null);
    const reviewTerminalTracker = Object.create(null);
    const itemOpenState = Object.create(null);
    const detailIframeCache = new Map();
    const pageParams = new URLSearchParams(window.location.search || "");
    let lastStatusMarkup = "";
    let lastStatusReady = false;
    const ACTIVE_JOB_STORAGE_KEY = "koko_active_job_id";
    const ACTIVE_JOB_SNAPSHOT_STORAGE_KEY = "koko_active_job_snapshot";
    const RESTORE_RETRY_LIMIT = 6;
    const RESTORE_RETRY_DELAY_MS = 1500;
    const IDLE_STATUS_HTML = `
      <div class="status-empty">
        <div class="status-empty-title">已就绪。</div>
        <div class="status-empty-copy">输入一个或多个视频链接后，系统会在这里实时显示拆解进度。</div>
      </div>
    `;
    const FILTER_IDLE_HTML = `
      <div class="status-empty">
        <div class="status-empty-title">筛选器已就绪。</div>
        <div class="status-empty-copy">贴入一批链接或上传表格后，Koko 会提取完整音频信息和三张关键帧，再输出通过三轮规则的视频。</div>
      </div>
    `;
    const UNDERSTANDING_IDLE_HTML = `
      <div class="status-empty">
        <div class="status-empty-title">视频理解已就绪。</div>
        <div class="status-empty-copy">输入一个或多个视频链接后，系统会显示同款拆解进度，完成后只给出脚本内容概述。</div>
      </div>
    `;
    let activeFilterJobId = "";
    let filterPollTimer = null;
    const ACTIVE_FILTER_JOB_STORAGE_KEY = "koko_active_filter_job_id";
    let activeTranslationJobId = "";
    let translationPollTimer = null;
    const ACTIVE_TRANSLATION_JOB_STORAGE_KEY = "koko_active_translation_job_id";
    let activeUnderstandingJobId = "";
    let understandingPollTimer = null;
    const ACTIVE_UNDERSTANDING_JOB_STORAGE_KEY = "koko_active_understanding_job_id";

    function setStudioPanel(panelId) {{
      const target = String(panelId || "filter-panel").trim() || "filter-panel";
      studioPanels.forEach((panel) => {{
        panel.classList.toggle("active", panel.id === target);
      }});
      studioPanelLinks.forEach((link) => {{
        link.classList.toggle("active", link.getAttribute("data-panel-target") === target);
      }});
    }}
    const STAGE_ORDER = ["queued", "download", "media_prep", "gemini_analysis", "v2_analysis", "consistency_audit", "targeted_recheck", "arbitration", "final_output", "completed"];
    const STAGE_LABELS = {{
      queued: "等待拆解",
      download: "下载视频",
      media_prep: "媒体预处理",
      gemini_analysis: "AI 主分析",
      v2_analysis: "本地分析",
      consistency_audit: "一致性审查",
      targeted_recheck: "条件复核",
      arbitration: "结果仲裁",
      final_output: "生成脚本",
      completed: "已完成",
      failed: "失败",
      starting: "准备中"
    }};
    const STAGE_COPY = {{
      queued: "任务已创建，正在排队等待拆解。",
      starting: "正在准备分析流程。",
      download: "正在下载源视频。",
      media_prep: "正在读取视频结构与基础信息。",
      gemini_analysis: "正在进行 AI 主分析。",
      v2_analysis: "正在进行本地分析与对照。",
      consistency_audit: "正在检查双轨结论是否一致。",
      targeted_recheck: "正在复核风险较高的片段。",
      arbitration: "正在选择最稳妥的故事主轴。",
      final_output: "正在生成最终脚本、预览和导出文件。",
      completed: "分析完成。"
    }};
    const REVIEW_STAGE_ORDER = ["queued", "plan", "recheck", "rebuild", "completed"];
    const REVIEW_STAGE_LABELS = {{
      queued: "排队中",
      plan: "制定复盘计划",
      recheck: "回看视频",
      rebuild: "重建脚本",
      completed: "复盘完成",
      failed: "复盘失败"
    }};

    function preserveDetailIframes() {{
      statusBox.querySelectorAll(".item-card[data-item-id] iframe[data-preview-item-id]").forEach((frame) => {{
        const itemId = frame.getAttribute("data-preview-item-id") || "";
        const src = frame.getAttribute("src") || "";
        if (!itemId || !src) return;
        detailIframeCache.set(itemId, {{ src, node: frame }});
      }});
    }}

    function ensureDetailIframes(root = document) {{
      root.querySelectorAll(".item-card[data-item-id]").forEach((detail) => {{
        if (!(detail instanceof HTMLDetailsElement)) return;
        const slot = detail.querySelector(".item-preview-slot");
        if (!(slot instanceof HTMLDivElement)) return;
        const itemId = slot.getAttribute("data-preview-item-id") || detail.getAttribute("data-item-id") || "";
        const url = slot.getAttribute("data-preview-url") || "";
        if (!url) {{
          slot.replaceChildren();
          return;
        }}
        if (!detail.open) {{
          slot.replaceChildren();
          return;
        }}
        const cached = detailIframeCache.get(itemId);
        if (cached && cached.src === url && cached.node instanceof HTMLIFrameElement) {{
          if (slot.firstElementChild !== cached.node) {{
            slot.replaceChildren(cached.node);
          }}
          wirePreviewTimeJumps(cached.node, itemId);
          return;
        }}
        const iframe = document.createElement("iframe");
        iframe.loading = "lazy";
        iframe.src = url;
        iframe.setAttribute("data-preview-item-id", itemId);
        iframe.addEventListener("load", () => wirePreviewTimeJumps(iframe, itemId));
        slot.replaceChildren(iframe);
        detailIframeCache.set(itemId, {{ src: url, node: iframe }});
      }});
      root.querySelectorAll("video[data-source-video]").forEach((video) => {{
        if (!(video instanceof HTMLVideoElement) || video.getAttribute("data-video-wired") === "1") return;
        video.setAttribute("data-video-wired", "1");
        const itemId = video.getAttribute("data-source-video") || "";
        const status = document.querySelector(`[data-source-video-status="${{itemId}}"]`);
        video.addEventListener("loadedmetadata", () => {{
          if (status) status.textContent = `原视频已加载，可用于校验脚本。时长约 ${{Math.round(video.duration || 0)}} 秒。`;
        }});
        video.addEventListener("error", () => {{
          if (status) status.textContent = "原视频没有加载出来：可能是 source.mp4 不存在、文件编码浏览器不支持，或文件还没有写入完成。";
        }});
      }});
    }}

    function versionedResultUrl(url, item) {{
      const rawUrl = String(url || "").trim();
      if (!rawUrl) return "";
      const version = String(item?.updated_at || item?.review_message || item?.review_stage || "").trim();
      if (!version) return rawUrl;
      const separator = rawUrl.includes("?") ? "&" : "?";
      return `${{rawUrl}}${{separator}}v=${{encodeURIComponent(version)}}`;
    }}

    function parseTimeToSeconds(value) {{
      const text = String(value || "").trim();
      const match = text.match(/(\\d{{1,2}}):(\\d{{2}})(?::(\\d{{2}}))?/);
      if (!match) return 0;
      if (match[3] !== undefined) {{
        return Number(match[1] || 0) * 3600 + Number(match[2] || 0) * 60 + Number(match[3] || 0);
      }}
      return Number(match[1] || 0) * 60 + Number(match[2] || 0);
    }}

    function seekSourceVideo(itemId, seconds) {{
      const video = document.querySelector(`[data-source-video="${{itemId}}"]`);
      const status = document.querySelector(`[data-source-video-status="${{itemId}}"]`);
      if (!(video instanceof HTMLVideoElement)) {{
        showToast("无法跳转视频", "当前结果没有可用的视频播放器。");
        return;
      }}
      try {{
        video.currentTime = Math.max(0, Number(seconds) || 0);
        video.play().catch(() => {{}});
        if (status) status.textContent = `已跳到 ${{Math.floor(video.currentTime / 60).toString().padStart(2, "0")}}:${{Math.floor(video.currentTime % 60).toString().padStart(2, "0")}}。`;
      }} catch (error) {{
        showToast("视频跳转失败", "原视频可能还没有加载完成。");
      }}
    }}

    function wirePreviewTimeJumps(iframe, itemId) {{
      if (!(iframe instanceof HTMLIFrameElement) || !itemId) return;
      let doc = null;
      try {{
        doc = iframe.contentDocument || iframe.contentWindow?.document || null;
      }} catch (error) {{
        return;
      }}
      if (!doc) return;
      doc.querySelectorAll("table.script-table tbody tr, table tbody tr").forEach((row) => {{
        const timeCell = row.querySelector("td");
        if (!timeCell || timeCell.getAttribute("data-koko-time-wired") === "1") return;
        const seconds = parseTimeToSeconds(timeCell.textContent || "");
        timeCell.setAttribute("data-koko-time-wired", "1");
        timeCell.setAttribute("title", "点击跳到原视频对应片段");
        timeCell.style.cursor = "pointer";
        timeCell.style.color = "#FF8200";
        timeCell.style.fontWeight = "800";
        timeCell.addEventListener("click", () => seekSourceVideo(itemId, seconds));
      }});
    }}

    function setStatus(html, ready = false) {{
      if (html === lastStatusMarkup && ready === lastStatusReady) return;
      preserveDetailIframes();
      statusBox.className = ready ? "status-box visible ready" : "status-box visible";
      statusBox.innerHTML = html;
      lastStatusMarkup = html;
      lastStatusReady = ready;
      ensureDetailIframes(statusBox);
    }}

    function setStudioOverviewValue(key, value, meta) {{
      const numberNode = document.querySelector(`[data-studio-stat="${{key}}"]`);
      const metaNode = document.querySelector(`[data-studio-stat-meta="${{key}}"]`);
      if (numberNode) numberNode.textContent = String(Number(value || 0));
      if (metaNode) metaNode.textContent = meta || "";
    }}

    function updateStudioOverview(data = null) {{
      if (!data) {{
        setStudioOverviewValue("queued", 0, "等待提交");
        setStudioOverviewValue("running", 0, "当前无任务");
        setStudioOverviewValue("completed", 0, "本次任务");
        setStudioOverviewValue("failed", 0, "本次任务");
        return;
      }}
      const items = Array.isArray(data.items) ? data.items : [];
      const total = Number(data.total_items || items.length || 0);
      const ownQueued = items.filter((item) => String(item?.status || "").trim() === "queued").length;
      const ownRunning = items.filter((item) => String(item?.status || "").trim() === "running").length;
      const completed = Number(data.completed_items || items.filter((item) => item?.status === "completed").length || 0);
      const failed = Number(data.failed_items || items.filter((item) => item?.status === "failed").length || 0);
      const systemQueue = data.system_queue || {{}};
      const globalQueued = Number(systemQueue.queued_count || 0);
      const globalRunning = Number(systemQueue.running_count || 0);
      const queued = Math.max(ownQueued, globalQueued);
      const running = Math.max(ownRunning, globalRunning);
      const position = Number(systemQueue.current_job_position || 0);
      const ahead = Number(systemQueue.current_job_ahead || 0);
      const queuedMeta = position
        ? `当前排队第 ${{position}} 位，前方 ${{ahead}} 条`
        : (queued ? "队列中等待执行" : "没有排队任务");
      const runningMeta = total
        ? `占比 ${{Math.round((running / Math.max(total, running, 1)) * 100)}}%`
        : (running ? "系统正在处理" : "当前无任务");
      const completedMeta = total ? `本次 ${{completed}}/${{total}}` : "本次任务";
      const failedMeta = total ? `本次 ${{failed}}/${{total}}` : "本次任务";
      setStudioOverviewValue("queued", queued, queuedMeta);
      setStudioOverviewValue("running", running, runningMeta);
      setStudioOverviewValue("completed", completed, completedMeta);
      setStudioOverviewValue("failed", failed, failedMeta);
    }}

    function updateStopAllButtonState(active) {{
      if (!stopAllBtn) return;
      stopAllBtn.disabled = !active;
    }}

    function setFilterStatus(html, ready = false) {{
      if (!filterStatusBox) return;
      filterStatusBox.className = ready ? "status-box visible ready" : "status-box visible";
      filterStatusBox.innerHTML = html;
    }}

    function setTranslationStatus(html, ready = false) {{
      if (!translationStatusBox) return;
      translationStatusBox.className = ready ? "status-box visible ready" : "status-box visible";
      translationStatusBox.innerHTML = html;
    }}

    function setUnderstandingStatus(html, ready = false) {{
      if (!understandingStatusBox) return;
      understandingStatusBox.className = ready ? "status-box visible ready" : "status-box visible";
      understandingStatusBox.innerHTML = html;
    }}

    function persistActiveUnderstandingJobId(jobId) {{
      const value = String(jobId || "").trim();
      try {{
        if (!value) window.localStorage.removeItem(ACTIVE_UNDERSTANDING_JOB_STORAGE_KEY);
        else window.localStorage.setItem(ACTIVE_UNDERSTANDING_JOB_STORAGE_KEY, value);
      }} catch (error) {{
        // Ignore storage failures.
      }}
    }}

    function readPersistedActiveUnderstandingJobId() {{
      try {{
        return String(window.localStorage.getItem(ACTIVE_UNDERSTANDING_JOB_STORAGE_KEY) || "").trim();
      }} catch (error) {{
        return "";
      }}
    }}

    function scheduleUnderstandingPoll(jobId, delay = 2500) {{
      if (understandingPollTimer) clearTimeout(understandingPollTimer);
      understandingPollTimer = setTimeout(() => pollUnderstandingJob(jobId), delay);
    }}

    function persistActiveTranslationJobId(jobId) {{
      const value = String(jobId || "").trim();
      try {{
        if (!value) window.localStorage.removeItem(ACTIVE_TRANSLATION_JOB_STORAGE_KEY);
        else window.localStorage.setItem(ACTIVE_TRANSLATION_JOB_STORAGE_KEY, value);
      }} catch (error) {{
        // Ignore storage failures.
      }}
    }}

    function scheduleTranslationPoll(jobId, delay = 1800) {{
      if (translationPollTimer) clearTimeout(translationPollTimer);
      translationPollTimer = setTimeout(() => pollTranslationJob(jobId), delay);
    }}

    function renderTranslationJob(data) {{
      const status = String(data.status || "queued");
      const stage = String(data.stage || status || "queued");
      const message = data.stage_message || (status === "completed" ? "转译视频已生成。" : status === "failed" ? "转译失败。" : "正在处理。");
      const subject = data.subject_summary ? `<div class="summary-box"><strong>主体识别</strong><br>${{escapeHtml(data.subject_summary)}}</div>` : "";
      const audio = data.original_audio_summary ? `<div class="summary-box"><strong>原音频理解</strong><br>${{escapeHtml(data.original_audio_summary)}}</div>` : "";
      const voiceover = data.portuguese_voiceover ? `<div class="summary-box"><strong>葡语文案</strong><br>${{escapeHtml(data.portuguese_voiceover)}}</div>` : "";
      const links = data.translated_video_url ? `
        <div class="artifact-row">
          <a class="artifact-link" href="${{escapeHtml(data.translated_video_url)}}" target="_blank" rel="noreferrer">打开葡语视频</a>
          ${{data.audio_url ? `<a class="artifact-link" href="${{escapeHtml(data.audio_url)}}" target="_blank" rel="noreferrer">打开葡语音轨</a>` : ""}}
        </div>
      ` : "";
      const error = data.error ? `<div class="queue-error">${{escapeHtml(data.error)}}</div>` : "";
      return `
        <section class="batch-shell">
          <div class="batch-top">
            <div>
              <div class="batch-title">葡语转译任务</div>
              <div class="batch-subtitle">${{escapeHtml(data.video_url || "")}}</div>
            </div>
            <span class="status ${{status === "completed" ? "status-completed" : status === "failed" ? "status-failed" : status === "running" ? "status-running" : "status-queued"}}">${{escapeHtml(status === "completed" ? "已完成" : status === "failed" ? "失败" : status === "running" ? "处理中" : "排队中")}}</span>
          </div>
          ${{progressMarkup(stage, message, data.id, data)}}
          ${{subject}}${{audio}}${{voiceover}}${{links}}${{error}}
        </section>
      `;
    }}

    async function pollTranslationJob(jobId) {{
      if (!jobId) return;
      try {{
        const res = await fetch(`/api/translation-jobs/${{jobId}}?_=${{Date.now()}}`);
        const data = await readJsonSafely(res);
        if (!res.ok) throw new Error(data.error || "转译任务查询失败");
        setTranslationStatus(renderTranslationJob(data), data.status === "completed");
        if (data.status === "completed" || data.status === "failed") {{
          persistActiveTranslationJobId("");
          activeTranslationJobId = "";
          return;
        }}
        scheduleTranslationPoll(jobId);
      }} catch (error) {{
        setTranslationStatus(`<span class="status status-failed">失败</span><br><br><code>${{escapeHtml(String(error.message || error))}}</code>`);
      }}
    }}

    function setFilterIdleState() {{
      if (filterPollTimer) {{
        clearTimeout(filterPollTimer);
        filterPollTimer = null;
      }}
      activeFilterJobId = "";
      try {{
        window.localStorage.removeItem(ACTIVE_FILTER_JOB_STORAGE_KEY);
      }} catch (error) {{
        // Ignore storage failures.
      }}
      setFilterStatus(FILTER_IDLE_HTML, true);
    }}

    function persistActiveFilterJobId(jobId) {{
      const value = String(jobId || "").trim();
      try {{
        if (!value) {{
          window.localStorage.removeItem(ACTIVE_FILTER_JOB_STORAGE_KEY);
        }} else {{
          window.localStorage.setItem(ACTIVE_FILTER_JOB_STORAGE_KEY, value);
        }}
      }} catch (error) {{
        // Ignore storage failures.
      }}
    }}

    function readPersistedActiveFilterJobId() {{
      try {{
        return String(window.localStorage.getItem(ACTIVE_FILTER_JOB_STORAGE_KEY) || "").trim();
      }} catch (error) {{
        return "";
      }}
    }}

    function setIdleState() {{
      if (jobPollTimer) {{
        clearTimeout(jobPollTimer);
        jobPollTimer = null;
      }}
      activeJobId = "";
      activeReviewItemId = "";
      persistActiveJobId("");
      lastStatusMarkup = "";
      lastStatusReady = false;
      setStatus(IDLE_STATUS_HTML, true);
      updateStudioOverview(null);
      updateStopAllButtonState(false);
    }}

    function schedulePoll(jobId, delay = 2500) {{
      if (jobPollTimer) clearTimeout(jobPollTimer);
      jobPollTimer = setTimeout(() => pollJob(jobId), delay);
    }}

    function scheduleFilterPoll(jobId, delay = 1800) {{
      if (filterPollTimer) clearTimeout(filterPollTimer);
      filterPollTimer = setTimeout(() => pollFilterJob(jobId), delay);
    }}

    function persistActiveJobId(jobId) {{
      const value = String(jobId || "").trim();
      if (!value) {{
        window.localStorage.removeItem(ACTIVE_JOB_STORAGE_KEY);
      }} else {{
        window.localStorage.setItem(ACTIVE_JOB_STORAGE_KEY, value);
      }}
    }}

    function readPersistedActiveJobId() {{
      return String(window.localStorage.getItem(ACTIVE_JOB_STORAGE_KEY) || "").trim();
    }}

    function persistActiveJobSnapshot(data) {{
      try {{
        if (!data || !data.id) {{
          window.localStorage.removeItem(ACTIVE_JOB_SNAPSHOT_STORAGE_KEY);
          return;
        }}
        window.localStorage.setItem(ACTIVE_JOB_SNAPSHOT_STORAGE_KEY, JSON.stringify(data));
      }} catch (error) {{
        // Best effort only.
      }}
    }}

    function readPersistedActiveJobSnapshot(jobId) {{
      try {{
        const raw = String(window.localStorage.getItem(ACTIVE_JOB_SNAPSHOT_STORAGE_KEY) || "").trim();
        if (!raw) return null;
        const parsed = JSON.parse(raw);
        if (!parsed || String(parsed.id || "").trim() !== String(jobId || "").trim()) return null;
        return parsed;
      }} catch (error) {{
        return null;
      }}
    }}

    async function readJsonSafely(response) {{
      const raw = await response.text();
      try {{
        return raw ? JSON.parse(raw) : {{}};
      }} catch (error) {{
        if (response && response.status >= 500) {{
          throw new Error(`服务暂时不可用（HTTP ${{response.status}}），请稍后重试。`);
        }}
        const preview = String(raw || "").slice(0, 180).trim();
        throw new Error(preview ? `服务返回了非 JSON 内容：${{preview}}` : "服务返回了空响应。");
      }}
    }}

    async function loadLibraryWorkbench(entryId) {{
      const target = String(entryId || "").trim();
      if (!target) return false;
      setStudioPanel("split-panel");
      setStatus(`
        <div class="status-empty">
          <div class="status-empty-title">正在载入脚本库工作台...</div>
          <div class="status-empty-copy">这条已入库脚本会直接接入当前编辑、修稿和复盘流程。</div>
        </div>
      `);
      try {{
        const res = await fetch(`/api/library-workbench/${{encodeURIComponent(target)}}?_=${{Date.now()}}`, {{
          cache: "no-store",
        }});
        const data = await readJsonSafely(res);
        if (!res.ok) throw new Error(data.error || "载入脚本库工作台失败");
        activeJobId = "";
        activeReviewItemId = "";
        persistActiveJobId("");
        setStatus(renderBatchResults(data), true);
        return true;
      }} catch (error) {{
        setStatus(`<span class="status status-failed">失败</span><br><br><code>${{escapeHtml(String(error.message || error))}}</code>`);
        return false;
      }}
    }}

    function showToast(title, copy) {{
      if (!appToast || !appToastTitle || !appToastCopy) return;
      appToastTitle.textContent = title || "";
      appToastCopy.textContent = copy || "";
      appToast.classList.add("show");
      appToast.setAttribute("aria-hidden", "false");
      if (toastTimer) clearTimeout(toastTimer);
      toastTimer = setTimeout(() => {{
        appToast.classList.remove("show");
        appToast.setAttribute("aria-hidden", "true");
      }}, 2600);
    }}

    function currentLanguageLabel(item) {{
      return (item.display_language || "zh") === "pt" ? "葡语视图" : "中文视图";
    }}

    function formatClock(isoValue) {{
      if (!isoValue) return "";
      const date = new Date(isoValue);
      if (Number.isNaN(date.getTime())) return "";
      return date.toLocaleTimeString("zh-CN", {{
        hour: "2-digit",
        minute: "2-digit",
        second: "2-digit",
        hour12: false,
      }});
    }}

    function buildArtifactHints(data) {{
      const items = Array.isArray(data?.items) ? data.items : [];
      const primaryItem = findCurrentItem(items) || items[0] || null;
      const hints = [];
      if (data?.updated_at) {{
        const clock = formatClock(data.updated_at);
        if (clock) hints.push(`最后更新：${{clock}}`);
      }}
      if (primaryItem?.result_json || primaryItem?.html_url) hints.push("最终脚本已生成");
      if (primaryItem?.docx_url) hints.push("导出文件已生成");
      if (primaryItem?.saved_to_library_at || primaryItem?.in_library) hints.push("已确认入库");
      else if (primaryItem?.status === "completed" && primaryItem?.result_json) hints.push("待确认入库");
      if (primaryItem?.report_url || primaryItem?.evidence_url) hints.push("报告已生成");
      return hints;
    }}

    function buildThinkingEvents(data) {{
      const items = Array.isArray(data?.items) ? data.items : [];
      const primaryItem = findCurrentItem(items) || items[0] || null;
      const stage = String(primaryItem?.stage || data?.stage || "queued").trim() || "queued";
      const effectiveStatus = deriveEffectiveJobStatus(data);
      const sequence = ["download", "media_prep", "gemini_analysis", "v2_analysis", "consistency_audit", "targeted_recheck", "arbitration", "final_output"];
      const stageIndex = stage === "completed"
        ? sequence.length
        : stage === "failed"
          ? Math.max(0, sequence.indexOf("final_output"))
          : Math.max(0, sequence.indexOf(stage));
      const events = [];
      if (effectiveStatus === "queued") {{
        events.push({{ type: "active", text: "任务已创建，正在等待进入执行队列。" }});
      }}
      sequence.forEach((key, idx) => {{
        if (effectiveStatus === "queued" && idx > 0) return;
        if (effectiveStatus === "completed" || idx < stageIndex) {{
          events.push({{ type: "done", text: `${{STAGE_LABELS[key]}}已完成` }});
          return;
        }}
        if (idx === stageIndex) {{
          events.push({{ type: stage === "failed" ? "note" : "active", text: stage === "failed" ? `${{STAGE_LABELS[key]}}阶段中断` : `正在${{STAGE_LABELS[key]}}` }});
        }}
      }});
      if (primaryItem?.result_json || primaryItem?.html_url) {{
        events.push({{ type: effectiveStatus === "completed" ? "done" : "note", text: "最终脚本文件已经生成。" }});
      }}
      if (primaryItem?.docx_url) {{
        events.push({{ type: effectiveStatus === "completed" ? "done" : "note", text: "导出文件已经准备好。" }});
      }}
      if (primaryItem?.saved_to_library_at || primaryItem?.in_library) {{
        events.push({{ type: "done", text: "脚本已确认进入脚本库。" }});
      }} else if (effectiveStatus === "completed" && primaryItem?.result_json) {{
        events.push({{ type: "note", text: "脚本尚未入库，请确认版本可用后点击“确认入库”。" }});
      }}
      const stageMessage = String(primaryItem?.stage_message || data?.stage_message || data?.message || "").trim();
      if (stageMessage) {{
        events.push({{ type: "note", text: stageMessage }});
      }}
      return events.slice(-8);
    }}

    function renderThinkingLog(data) {{
      const events = buildThinkingEvents(data);
      if (!events.length) return "";
      const updated = formatClock(data?.updated_at);
      const rows = events.map((event) => `
        <div class="thinking-item ${{event.type}}">
          <span class="thinking-dot"></span>
          <span>${{escapeHtml(event.text)}}</span>
        </div>
      `).join("");
      return `
        <div class="thinking-shell">
          <div class="thinking-head">
            <span class="thinking-title">Thinking</span>
            <span class="thinking-updated">${{updated ? `最后同步 ${{updated}}` : "等待状态更新"}}</span>
          </div>
          <div class="thinking-list">${{rows}}</div>
        </div>
      `;
    }}

    function closeExportChoice() {{
      pendingExportChoice = null;
      if (!exportChoiceOverlay) return;
      exportChoiceOverlay.classList.remove("open");
      exportChoiceOverlay.setAttribute("aria-hidden", "true");
    }}

    function openExportChoice(urls) {{
      pendingExportChoice = urls || null;
      if (!exportChoiceOverlay) return;
      exportChoiceOverlay.classList.add("open");
      exportChoiceOverlay.setAttribute("aria-hidden", "false");
    }}

    if (heroPanel) {{
      heroPanel.addEventListener("mousemove", (event) => {{
        const rect = heroPanel.getBoundingClientRect();
        const x = (event.clientX - rect.left) / rect.width;
        const y = (event.clientY - rect.top) / rect.height;
        heroPanel.style.setProperty("--mouse-x", `${{(x * 100).toFixed(2)}}%`);
        heroPanel.style.setProperty("--mouse-y", `${{(y * 100).toFixed(2)}}%`);
        heroPanel.style.setProperty("--tilt-x", "0px");
        heroPanel.style.setProperty("--tilt-y", "0px");
      }});

      heroPanel.addEventListener("mouseleave", () => {{
        heroPanel.style.setProperty("--mouse-x", "50%");
        heroPanel.style.setProperty("--mouse-y", "50%");
        heroPanel.style.setProperty("--tilt-x", "0px");
        heroPanel.style.setProperty("--tilt-y", "0px");
      }});
    }}

    function escapeHtml(value) {{
      return String(value || "").replace(/&/g, "&amp;").replace(/</g, "&lt;").replace(/>/g, "&gt;");
    }}

    function jobStatusLabel(status) {{
      const text = String(status || "").trim();
      if (text === "completed") return "已完成";
      if (text === "failed") return "失败";
      if (text === "running") return "运行中";
      if (text === "queued") return "排队中";
      return text || "排队中";
    }}

    function parseVideoDisplayName(url, idx = 0) {{
      const fallback = `视频 ${{idx + 1}}`;
      const text = String(url || "").trim();
      if (!text) return fallback;
      try {{
        const parsed = new URL(text);
        const parts = parsed.pathname.split("/").filter(Boolean);
        const handle = parts.find((part) => part.startsWith("@")) || "";
        const videoIndex = parts.indexOf("video");
        const videoId = videoIndex >= 0 && parts[videoIndex + 1] ? parts[videoIndex + 1] : "";
        if (handle && videoId) return `${{handle}} / ${{videoId}}`;
        if (handle) return handle;
        if (videoId) return `video / ${{videoId}}`;
        return parsed.hostname.replace(/^www\\./, "");
      }} catch (error) {{
        return fallback;
      }}
    }}

    function filterBucketLabel(bucket) {{
      if (bucket === "high") return "通过";
      if (bucket === "low") return "不通过";
      return "待判断";
    }}

    function filterBucketClass(bucket) {{
      if (bucket === "high") return "completed";
      if (bucket === "medium") return "running";
      if (bucket === "low") return "waiting";
      return "waiting";
    }}

    function filterStageLabel(stage) {{
      if (stage === "metadata") return "读取页面信息";
      if (stage === "audio") return "提取完整音频";
      if (stage === "frames") return "抽取三张关键帧";
      if (stage === "classify") return "三轮筛选判断";
      if (stage === "completed") return "筛选完成";
      if (stage === "failed") return "筛选失败";
      return "等待筛选";
    }}

    function displayVideoName(item, idx = 0) {{
      const title = String(item?.title || "").trim();
      if (title) return title;
      return parseVideoDisplayName(item?.video_url || "", idx);
    }}

    function copyText(text, successMessage = "已复制") {{
      const value = String(text || "").trim();
      if (!value) return;
      if (navigator.clipboard?.writeText) {{
        navigator.clipboard.writeText(value).then(() => {{
          showToast("复制成功", successMessage);
        }}).catch(() => {{
          showToast("复制失败", "请稍后重试。");
        }});
        return;
      }}
      showToast("复制失败", "当前环境不支持自动复制。");
    }}

    function itemState(item) {{
      const status = String(item?.status || "").trim();
      if (status === "completed") return "completed";
      if (status === "failed") return "failed";
      if (status === "running") return "running";
      return "waiting";
    }}

    function itemStateLabel(item) {{
      const state = itemState(item);
      if (state === "completed") return "已完成";
      if (state === "failed") return "失败";
      if (state === "running") return "正在拆解";
      return "等待拆解";
    }}

    function workloadKindLabel(kind) {{
      return String(kind || "").trim() === "review" ? "复盘重做" : "脚本拆解";
    }}

    function findCurrentItem(items) {{
      return (items || []).find((item) => itemState(item) === "running") || null;
    }}

    function collectUrls(inputNode = videoInput) {{
      return String(inputNode?.value || "")
        .split(/[\\n\\r,]+/)
        .map((value) => value.trim())
        .filter((value, index, arr) => value && /^https?:\\/\\//i.test(value) && arr.indexOf(value) === index);
    }}

    function collectFilterUrls() {{
      const text = String(filterInput?.value || "").trim();
      const matches = text.match(/https?:\\/\\/[^\\s<>\"]+/gi) || [];
      return matches
        .map((value) => value.trim().replace(/[),.;!?]+$/, ""))
        .filter((value, index, arr) => value && arr.indexOf(value) === index);
    }}

    function readFileAsBase64(file) {{
      return new Promise((resolve, reject) => {{
        const reader = new FileReader();
        reader.onload = () => {{
          const raw = String(reader.result || "");
          const marker = raw.indexOf(",");
          resolve(marker >= 0 ? raw.slice(marker + 1) : raw);
        }};
        reader.onerror = () => reject(reader.error || new Error("文件读取失败"));
        reader.readAsDataURL(file);
      }});
    }}

    function normalizeRows(script) {{
      if (!script || typeof script !== "object") return [];
      if (Array.isArray(script.rows) && script.rows.length) return script.rows;
      if (Array.isArray(script.synthesized_segments) && script.synthesized_segments.length) return script.synthesized_segments;
      return [];
    }}

    function normalizedText(value, fallback = "无") {{
      const text = String(value || "").trim();
      return text || fallback;
    }}

    function normalizeInsightItems(value, fallbackItems = [{{ label: "要点", text: "无" }}]) {{
      if (Array.isArray(value)) return value;
      return fallbackItems;
    }}

    function formatInsightDraft(items) {{
      return normalizeInsightItems(items).map((item, idx) => {{
        const label = normalizedText(item.label || item.title || item.name, `要点${{idx + 1}}`);
        const text = normalizedText(item.text || item.description || item.value);
        return `${{idx + 1}}. ${{label}}：${{text}}`;
      }}).join("\\n");
    }}

    function buildStoryboardPrompt(script) {{
      const rows = normalizeRows(script).slice(0, 9);
      const beats = rows.map((row, idx) => {{
        return `${{idx + 1}}. 时间=${{normalizedText(row.time, "")}}；场景=${{normalizedText(row.visual_content)}}；动作=${{normalizedText(row.action)}}`;
      }}).join("\\n");
      return [
        "请把这份短视频脚本画成黑白灰铅笔分镜稿风格的示意图。",
        "固定硬性条件：固定 1:1 正方形，白底纸面，手绘线稿，固定 3x3 九宫格分镜，非彩色，非照片，不要海报感，不要带任何文字、字幕、标题、编号、logo 或水印。",
        `标题：${{normalizedText(script.title, "视频脚本")}}`,
        `整体梗概：${{normalizedText(script.whole_video_summary)}}`,
        "关键分镜：",
        beats || "1. 按标题和梗概生成 9 格 3x3 连续分镜。",
      ].join("\\n");
    }}

    function formatScriptDraft(item, script) {{
      const rows = normalizeRows(script);
      const mechanismReason = (((script.mechanism || {{}}).reason) || "");
      const rowText = rows.map((row) => {{
        return [
          `时间：${{normalizedText(row.time, "")}}`,
          `画面内容：${{normalizedText(row.visual_content, "")}}`,
          `动作：${{normalizedText(row.action, "")}}`,
          "关键对白/旁白：",
          normalizedText(row.dialogue_or_audio, ""),
        ].join("\\n");
      }}).join("\\n---\\n");
      return [
        "标题：",
        normalizedText(script.title || item.title || "", "视频脚本"),
        "",
        "整体梗概：",
        normalizedText(script.whole_video_summary),
        "",
        "机制说明：",
        normalizedText(mechanismReason),
        "",
        "核心爆点：",
        formatInsightDraft(script.core_viral_points),
        "",
        "替换方案：",
        formatInsightDraft(script.replaceable_parts),
        "",
        "脚本表：",
        rowText,
      ].join("\\n");
    }}

    function splitDraftSections(text) {{
      const sections = {{}};
      const order = [];
      let current = "";
      String(text || "").split(/\\r?\\n/).forEach((line) => {{
        const match = line.match(/^\\s*(标题|整体梗概|机制说明|核心爆点|替换方案|可替换部分|脚本表)\\s*[:：]\\s*$/);
        if (match) {{
          current = match[1];
          if (!sections[current]) {{
            sections[current] = [];
            order.push(current);
          }}
          return;
        }}
        if (current) sections[current].push(line);
      }});
      return sections;
    }}

    function parseInsightDraft(text) {{
      const lines = String(text || "").split(/\\r?\\n/).map((line) => line.trim()).filter(Boolean);
      return lines.map((line) => {{
        const cleaned = line.replace(/^\\d+[.、)]\\s*/, "");
        const parts = cleaned.split(/[:：]/);
        if (parts.length >= 2) {{
          return {{
            label: normalizedText(parts.shift(), "要点"),
            text: normalizedText(parts.join("："), "无"),
          }};
        }}
        return {{ label: "要点", text: normalizedText(cleaned, "无") }};
      }});
    }}

    function parseScriptRowsDraft(text, originalRows) {{
      const blocks = String(text || "").split(/\\n\\s*---\\s*\\n/g).map((block) => block.trim()).filter(Boolean);
      return blocks.map((block, idx) => {{
        const original = originalRows[idx] || {{}};
        const readField = (label, nextLabels) => {{
          const escaped = label.replace(/[.*+?^${{}}()|[\\]\\\\]/g, "\\\\$&");
          const lookahead = nextLabels.map((item) => item.replace(/[.*+?^${{}}()|[\\]\\\\]/g, "\\\\$&")).join("|");
          const pattern = new RegExp(`${{escaped}}\\\\s*[:：]\\\\s*([\\\\s\\\\S]*?)(?=\\\\n(?:${{lookahead}})\\\\s*[:：]|$)`);
          const match = block.match(pattern);
          return match ? match[1].trim() : "";
        }};
        return {{
          original_index: idx,
          time: readField("时间", ["画面内容", "动作", "关键对白/旁白"]) || original.time || "",
          visual_content: readField("画面内容", ["动作", "关键对白/旁白"]) || original.visual_content || "",
          action: readField("动作", ["关键对白/旁白"]) || original.action || "",
          dialogue_or_audio: readField("关键对白/旁白", []) || original.dialogue_or_audio || "",
          integrated_summary: original.integrated_summary || "",
        }};
      }});
    }}

    const CONTENT_TYPE_OPTIONS = {json.dumps(LIBRARY_FILTER_LABELS, ensure_ascii=False)};

    function manualTagPickerMarkup(item, compact = false) {{
      const manualValue = item.content_type_source === "manual" ? (item.content_type || "") : "";
      const options = [`<option value="">自动判断分类</option>`, ...CONTENT_TYPE_OPTIONS.map((label) => {{
        return `<option value="${{escapeHtml(label)}}" ${{label === manualValue ? "selected" : ""}}>${{escapeHtml(label)}}</option>`;
      }})].join("");
      return `
        <div class="manual-tag-picker">
          <label>${{compact ? "入库分类 tag" : "分类 tag（不选则 AI 自动判断）"}}</label>
          <select data-manual-content-type="${{item.id}}">
            ${{options}}
          </select>
        </div>
      `;
    }}

    function selectedManualContentType(itemId, root = document) {{
      const value = root.querySelector(`[data-manual-content-type="${{itemId}}"]`)?.value || "";
      return CONTENT_TYPE_OPTIONS.includes(value) ? value : "";
    }}

    function buildLibraryConfirmMarkup(item) {{
      if (item.status !== "completed" || !item.result_json) return "";
      const alreadySaved = Boolean(item.saved_to_library_at || item.in_library);
      if (alreadySaved) {{
        return `
          <div class="library-confirm-card done" data-library-confirm-card="${{item.id}}">
            <div class="library-confirm-copy">
              <div class="library-confirm-title">已入库</div>
              <div class="library-confirm-note">这条脚本已在脚本库中。修改下方内容后，可以随时点击右侧按钮完整覆盖脚本库版本。</div>
            </div>
            ${{manualTagPickerMarkup(item)}}
            <button class="action-link primary" type="button" data-save-library="${{item.id}}">保存当前版本并更新入库</button>
            <a class="action-link" href="/library">打开脚本管理</a>
          </div>
        `;
      }}
      return `
        <div class="library-confirm-card" data-library-confirm-card="${{item.id}}">
          <div class="library-confirm-copy">
            <div class="library-confirm-title">确认入库</div>
            <div class="library-confirm-note">确认这个版本可用后，Koko 会自动生成全葡语版本并写入脚本库。</div>
          </div>
          ${{manualTagPickerMarkup(item)}}
          <button class="action-link primary" type="button" data-confirm-library="${{item.id}}">转葡语并入库</button>
        </div>
      `;
    }}

    function buildEditorMarkup(item) {{
      if (item.status !== "completed" || !item.result_json) return "";
      const script = item.result_json || {{}};
      const draft = formatScriptDraft(item, script);
      const rowsJson = escapeHtml(JSON.stringify(normalizeRows(script)));
      return `
        <details class="editor-disclosure">
          <summary class="editor-summary">
            <span class="editor-summary-title">整稿编辑</span>
          </summary>
          <div class="editor-shell" data-editor-item="${{item.id}}" data-editor-lang="${{escapeHtml(item.display_language || "zh")}}" data-editor-rows="${{rowsJson}}">
            <textarea class="editor-textarea editor-draft-textarea" data-editor-draft spellcheck="false">${{escapeHtml(draft)}}</textarea>
            ${{manualTagPickerMarkup(item, true)}}
            <div class="link-row">
              <button class="action-link" type="button" data-save-edits="${{item.id}}">保存修改</button>
              <button class="action-link primary" type="button" data-save-library="${{item.id}}">保存修改并转葡语入库</button>
            </div>
          </div>
        </details>
      `;
    }}

    function buildPrimaryEditorMarkup(item) {{
      if (item.status !== "completed" || !item.result_json) return "";
      const script = item.result_json || {{}};
      const rowsJson = escapeHtml(JSON.stringify(normalizeRows(script)));
      const mechanismReason = (((script.mechanism || {{}}).reason) || "");
      const corePointCards = normalizeInsightItems(script.core_viral_points, []);
      const storyboardPrompt = normalizedText(item.storyboard_prompt || buildStoryboardPrompt(script), "");
      const storyboardPreviewUrl = versionedResultUrl(item.storyboard_preview_url || item.storyboard_cover_url || "", item);
      const renderInsightEditors = (items, kind) => {{
        return normalizeInsightItems(items).map((point, idx) => `
          <div class="structured-insight" data-insight-kind="${{kind}}" data-insight-index="${{idx}}">
            <div class="structured-insight-head">
              <input class="structured-editor-input structured-insight-label" data-insight-field="label" value="${{escapeHtml(normalizedText(point.label || point.title || point.name, `要点${{idx + 1}}`))}}">
              <button class="structured-delete-btn" type="button" data-delete-insight="${{kind}}">删除格子</button>
            </div>
            <textarea class="structured-editor-textarea" data-insight-field="text">${{escapeHtml(normalizedText(point.text || point.description || point.value))}}</textarea>
          </div>
        `).join("");
      }};
      const rows = normalizeRows(script);
      const rowEditors = rows.map((row, idx) => `
        <tr data-structured-row-index="${{idx}}" data-row-original-index="${{idx}}">
          <td><input class="structured-cell-input" data-row-field="time" value="${{escapeHtml(normalizedText(row.time, ""))}}"></td>
          <td><textarea class="structured-cell-textarea" data-row-field="visual_content">${{escapeHtml(normalizedText(row.visual_content, ""))}}</textarea></td>
          <td><textarea class="structured-cell-textarea" data-row-field="action">${{escapeHtml(normalizedText(row.action, ""))}}</textarea></td>
          <td><textarea class="structured-cell-textarea" data-row-field="dialogue_or_audio">${{escapeHtml(normalizedText(row.dialogue_or_audio, ""))}}</textarea></td>
          <td class="structured-row-actions"><button class="structured-delete-btn" type="button" data-delete-row>删除整行</button></td>
        </tr>
      `).join("");
      return `
        <div class="editor-shell structured-editor" data-editor-item="${{item.id}}" data-editor-mode="structured" data-editor-lang="${{escapeHtml(item.display_language || "zh")}}" data-editor-rows="${{rowsJson}}">
          <section class="structured-editor-section">
            <h5>标题</h5>
            <input class="structured-editor-input structured-title-input" data-edit-field="title" value="${{escapeHtml(normalizedText(script.title || item.title || "", "视频脚本"))}}">
          </section>
          <section class="structured-editor-section">
            <h5>视频整体内容总结</h5>
            <textarea class="structured-editor-textarea" data-edit-field="whole_video_summary">${{escapeHtml(normalizedText(script.whole_video_summary))}}</textarea>
          </section>
          <section class="structured-editor-section">
            <h5>核心爆点</h5>
            <div class="structured-insight-grid">${{renderInsightEditors(corePointCards, "core")}}</div>
          </section>
          <section class="structured-editor-section">
            <h5>替换方案</h5>
            <div class="structured-insight-grid">${{renderInsightEditors(script.replaceable_parts, "replaceable")}}</div>
          </section>
          <section class="structured-editor-section">
            <h5>脚本表</h5>
            <div class="structured-table-wrap">
              <table class="structured-script-table">
                <thead>
                  <tr><th>时间</th><th>画面内容</th><th>动作</th><th>关键对白/旁白</th><th>操作</th></tr>
                </thead>
                <tbody>${{rowEditors}}</tbody>
              </table>
            </div>
          </section>
          <section class="structured-editor-section">
            <h5>生成分解示意图</h5>
            <div class="storyboard-panel">
              ${{storyboardPreviewUrl
                ? `<div class="storyboard-preview-wrap"><img class="storyboard-preview" src="${{escapeHtml(storyboardPreviewUrl)}}" alt="分解示意图"></div>`
                : `<div class="storyboard-empty">这里会生成分镜稿风格的示意图。默认会根据当前脚本表自动写提示词，你也可以先改下面这段要求，再重新生成。</div>`}}
              <textarea class="structured-editor-textarea storyboard-prompt" data-storyboard-prompt="${{item.id}}">${{escapeHtml(storyboardPrompt)}}</textarea>
              <div class="storyboard-actions">
                <button class="action-link" type="button" data-generate-storyboard-prompt="${{item.id}}">重新生成生图提示词</button>
                <button class="action-link primary" type="button" data-generate-storyboard="${{item.id}}">${{storyboardPreviewUrl ? "重新生成示意图" : "生成示意图"}}</button>
              </div>
            </div>
          </section>
          <input type="hidden" data-edit-field="mechanism_reason" value="${{escapeHtml(normalizedText(mechanismReason))}}">
        </div>
      `;
    }}

    function buildReviewMarkup(item) {{
      if (item.status !== "completed" || !item.result_json) return "";
      const status = normalizedText(item.review_status || "", "");
      const stage = normalizedText(item.review_stage || "", "");
      const message = normalizedText(item.review_message || "", "");
      const feedback = status === "running" || status === "failed" ? normalizedText(item.review_feedback || "", "") : "";
      const editedBadge = item.edited ? `<span class="batch-chip">Manual edits exist</span>` : "";
      const reviewedBadge = item.reviewed ? `<span class="batch-chip">Reviewed version active</span>` : "";
      const reviewState = status ? `<div class="review-note">${{escapeHtml(status)}}${{message ? ` · ${{escapeHtml(message)}}` : ""}}</div>` : "";
      const reviewProgress = buildReviewProgressMarkup(stage, status, message);
      const messages = Array.isArray(item.chat_messages) ? item.chat_messages : [];
      const history = messages.length
        ? messages.slice(-12).map((entry) => {{
            const role = entry.role === "user" ? "user" : "koko";
            const mode = "";
            return `<div class="review-message ${{role}}">${{escapeHtml(entry.content || "")}}${{mode}}</div>`;
          }}).join("")
        : `<div class="review-message koko">告诉我哪里要改。</div>`;
      const recheckDisabled = item.source_video_available ? "" : "disabled";
      const recheckTitle = item.source_video_available ? "重新看视频后改稿" : "当前 source.mp4 不可用，不能重新看视频";
      return `
        <div class="review-shell" data-review-item="${{item.id}}" data-chat-mode="minor">
          <div class="koko-chat-head">
            <div class="koko-chat-avatar">K</div>
            <div>
              <div class="koko-chat-title">Koko 修稿助手</div>
            </div>
          </div>
          <div class="review-chat-log" data-chat-log="${{item.id}}">
            ${{history}}
            ${{item.review_feedback ? `<div class="review-message user">${{escapeHtml(item.review_feedback)}}</div>` : ""}}
            ${{item.review_message ? `<div class="review-message koko">${{escapeHtml(item.review_message)}}</div>` : ""}}
          </div>
          <div class="review-mode-toggle" aria-label="Koko 修稿模式">
            <button class="chat-mode-option active" type="button" data-chat-mode-choice="minor">小修</button>
            <button class="chat-mode-option" type="button" data-chat-mode-choice="major">大改</button>
            <button class="chat-mode-option" type="button" data-chat-mode-choice="recheck" title="${{escapeHtml(recheckTitle)}}" ${{recheckDisabled}}>重新看视频</button>
          </div>
          ${{editedBadge.replace("Manual edits exist", "已有人工修改")}}
          ${{reviewedBadge.replace("Reviewed version active", "当前是复盘版本")}}
          ${{reviewProgress}}
          ${{reviewState}}
          <div class="chat-compose">
            <textarea class="editor-textarea" data-review-feedback data-chat-message placeholder="直接说哪里要改...">${{escapeHtml(feedback)}}</textarea>
            <div class="link-row">
              <button class="action-link primary" type="button" data-chat-edit="${{item.id}}">${{status === "running" ? "处理中..." : "发送给 Koko 修改"}}</button>
            </div>
          </div>
        </div>
      `;
    }}

    function buildReviewProgressMarkup(stage, status, message) {{
      if (!status && !stage) return "";
      const normalizedStage = stage || (status === "completed" ? "completed" : status === "failed" ? "failed" : "queued");
      const stageIndex = REVIEW_STAGE_ORDER.indexOf(normalizedStage);
      const percent = normalizedStage === "completed"
        ? 100
        : normalizedStage === "failed"
          ? 100
          : stageIndex >= 0
            ? Math.max(8, Math.round(((stageIndex + 1) / REVIEW_STAGE_ORDER.length) * 100))
            : 12;
      const fillClass = normalizedStage === "failed" ? "review-progress-fill failed" : "review-progress-fill";
      const steps = REVIEW_STAGE_ORDER.slice(0, 4).map((key) => {{
        let cls = "review-step";
        const keyIndex = REVIEW_STAGE_ORDER.indexOf(key);
        if (normalizedStage === "completed" || (stageIndex > keyIndex && stageIndex >= 0)) cls += " done";
        else if (normalizedStage === key) cls += " active";
        else if (normalizedStage === "failed" && stageIndex === -1 && key === "rebuild") cls += " failed";
        return `<div class="${{cls}}">${{REVIEW_STAGE_LABELS[key]}}</div>`;
      }}).join("");
      const badge = normalizedStage === "failed" ? "复盘失败" : (REVIEW_STAGE_LABELS[normalizedStage] || "复盘中");
      return `
        <div class="review-progress">
          <div class="review-progress-top">
            <span>${{escapeHtml(message || "正在准备复盘。")}}</span>
            <span>${{badge}} · ${{percent}}%</span>
          </div>
          <div class="review-progress-rail"><div class="${{fillClass}}" style="width:${{percent}}%"></div></div>
          <div class="review-steps">${{steps}}</div>
        </div>
      `;
    }}

    function buildVideoTimelineMarkup(item) {{
      const rows = normalizeRows(item?.result_json || {{}}).slice(0, 80);
      if (!rows.length) return "";
      const buttons = rows.map((row, idx) => {{
        const time = normalizedText(row.time || "", `片段 ${{idx + 1}}`);
        const seconds = parseTimeToSeconds(time);
        const label = normalizedText(row.visual_content || row.action || row.dialogue_or_audio || "", "查看这一段");
        return `
          <button class="timeline-jump" type="button" data-seek-video="${{item.id}}" data-seek-seconds="${{seconds}}">
            <strong>${{escapeHtml(time)}}</strong>
            <span>${{escapeHtml(label)}}</span>
          </button>
        `;
      }}).join("");
      return `<div class="video-timeline">${{buttons}}</div>`;
    }}

    function collectItemEdits(itemId) {{
      const root = document.querySelector(`[data-editor-item="${{itemId}}"]`);
      if (!root) return null;
      if (root.getAttribute("data-editor-mode") === "structured") {{
        const collectInsights = (kind) => Array.from(root.querySelectorAll(`[data-insight-kind="${{kind}}"]`)).map((card) => {{
          return {{
            label: card.querySelector('[data-insight-field="label"]')?.value || "",
            text: card.querySelector('[data-insight-field="text"]')?.value || "",
          }};
        }});
        const coreViralPoints = collectInsights("core");
        const replaceableParts = collectInsights("replaceable");
        const rows = Array.from(root.querySelectorAll("[data-structured-row-index]")).map((rowCard, idx) => {{
          return {{
            original_index: Number(rowCard.getAttribute("data-row-original-index") || idx),
            time: rowCard.querySelector('[data-row-field="time"]')?.value || "",
            visual_content: rowCard.querySelector('[data-row-field="visual_content"]')?.value || "",
            action: rowCard.querySelector('[data-row-field="action"]')?.value || "",
            dialogue_or_audio: rowCard.querySelector('[data-row-field="dialogue_or_audio"]')?.value || "",
            integrated_summary: "",
          }};
        }});
        const payload = {{
          title: root.querySelector('[data-edit-field="title"]')?.value || "",
          whole_video_summary: root.querySelector('[data-edit-field="whole_video_summary"]')?.value || "",
          mechanism_reason: root.querySelector('[data-edit-field="mechanism_reason"]')?.value || "",
          core_viral_points: coreViralPoints,
          replaceable_parts: replaceableParts,
          rows,
          target_language: root.getAttribute("data-editor-lang") || "zh",
          content_type: selectedManualContentType(itemId, root),
        }};
        return payload;
      }}
      const draft = root.querySelector("[data-editor-draft]")?.value || "";
      const sections = splitDraftSections(draft);
      let originalRows = [];
      try {{
        originalRows = JSON.parse(root.getAttribute("data-editor-rows") || "[]");
      }} catch (_error) {{
        originalRows = [];
      }}
      const textOf = (name) => (sections[name] || []).join("\\n").trim();
        return {{
          title: textOf("标题"),
          whole_video_summary: textOf("整体梗概"),
          mechanism_reason: textOf("机制说明"),
          core_viral_points: parseInsightDraft(textOf("核心爆点")),
          replaceable_parts: parseInsightDraft(textOf("替换方案") || textOf("可替换部分")),
          rows: parseScriptRowsDraft(textOf("脚本表"), originalRows),
          target_language: root.getAttribute("data-editor-lang") || "zh",
          content_type: selectedManualContentType(itemId, root),
      }};
    }}

    function collectReviewFeedback(itemId) {{
      const root = document.querySelector(`[data-review-item="${{itemId}}"]`);
      if (!root) return "";
      return root.querySelector('[data-review-feedback]')?.value || "";
    }}

    function collectReviewMode(itemId) {{
      const root = document.querySelector(`[data-review-item="${{itemId}}"]`);
      if (!root) return "partial";
      const chatMode = root.getAttribute("data-chat-mode") || "";
      if (chatMode === "recheck") return "full";
      if (chatMode === "major") return "partial";
      return root.querySelector('input[name="review-mode-' + itemId + '"]:checked')?.value || "partial";
    }}

    function collectChatMode(itemId) {{
      const root = document.querySelector(`[data-review-item="${{itemId}}"]`);
      return root?.getAttribute("data-chat-mode") || "minor";
    }}

    function appendChatBubble(itemId, role, text, pending = false) {{
      const log = document.querySelector(`[data-chat-log="${{itemId}}"]`);
      if (!log) return null;
      const firstDefault = log.querySelector(".review-message.koko");
      if (firstDefault && firstDefault.textContent.trim() === "告诉我哪里要改。") {{
        firstDefault.remove();
      }}
      const bubble = document.createElement("div");
      bubble.className = `review-message ${{role === "user" ? "user" : "koko"}}`;
      if (pending) bubble.setAttribute("data-chat-pending", "1");
      bubble.textContent = text;
      log.appendChild(bubble);
      log.scrollTop = log.scrollHeight;
      return bubble;
    }}

    function updatePendingChatBubble(itemId, text) {{
      const bubble = document.querySelector(`[data-chat-log="${{itemId}}"] [data-chat-pending="1"]`);
      if (!bubble) return;
      bubble.textContent = text;
      bubble.removeAttribute("data-chat-pending");
    }}

    async function persistItemEdits(itemId, mode, button) {{
      const payload = collectItemEdits(itemId);
      if (!payload) return;
      const original = button.textContent;
      button.disabled = true;
      button.textContent = mode === "library" ? "保存中..." : "更新中...";
      try {{
        const endpoint = mode === "library" ? `/api/items/${{itemId}}/save-to-library` : `/api/items/${{itemId}}/save`;
        const response = await fetch(endpoint, {{
          method: "POST",
          headers: {{ "Content-Type": "application/json" }},
          body: JSON.stringify(payload),
        }});
        const data = await response.json();
        if (!response.ok) throw new Error(data.error || "Save failed");
        if (activeJobId) {{
          pollJob(activeJobId);
        }} else {{
          window.location.reload();
        }}
      }} catch (error) {{
        alert(String(error.message || error));
        button.disabled = false;
        button.textContent = original;
      }}
    }}

    async function confirmLibraryEntry(itemId, button) {{
      if (!itemId || !button) return;
      const original = button.textContent;
      button.disabled = true;
      button.textContent = "入库中...";
      const scope = button.closest("[data-library-confirm-card]") || document;
      try {{
        const response = await fetch(`/api/items/${{itemId}}/confirm-library`, {{
          method: "POST",
          headers: {{ "Content-Type": "application/json" }},
          body: JSON.stringify({{ content_type: selectedManualContentType(itemId, scope) }}),
        }});
        const data = await response.json();
        if (!response.ok) throw new Error(data.error || "Confirm library failed");
        showToast("已入库", "这条脚本已进入脚本管理。");
        if (activeJobId) {{
          pollJob(activeJobId);
        }} else {{
          window.location.reload();
        }}
      }} catch (error) {{
        alert(String(error.message || error));
        button.disabled = false;
        button.textContent = original;
      }}
    }}

    async function downloadScript(url, button) {{
      if (!url) return;
      const original = button ? button.textContent : "";
      if (button) {{
        button.disabled = true;
        button.textContent = "导出中...";
      }}
      try {{
        const link = document.createElement("a");
        link.href = url;
        link.download = url.split("/").pop() || "script_export.docx";
        link.rel = "noopener";
        link.target = "_self";
        document.body.appendChild(link);
        link.click();
        link.remove();
        setTimeout(() => {{
          showToast("导出成功", "脚本文件已经开始下载到你的本地。");
        }}, 150);
      }} catch (error) {{
        alert("导出脚本失败，请重试。");
      }} finally {{
        if (button) {{
          button.disabled = false;
          button.textContent = original;
        }}
      }}
    }}

    async function switchDisplayLanguage(itemId, language, button) {{
      const payload = collectItemEdits(itemId) || {{}};
      payload.language = language;
      const original = button.textContent;
      button.disabled = true;
      button.textContent = language === "pt" ? "转换中..." : "切换中...";
      try {{
        const response = await fetch(`/api/items/${{itemId}}/display-language`, {{
          method: "POST",
          headers: {{ "Content-Type": "application/json" }},
          body: JSON.stringify(payload),
        }});
        const data = await response.json();
        if (!response.ok) throw new Error(data.error || "Language switch failed");
        if (activeJobId) {{
          pollJob(activeJobId);
        }} else {{
          window.location.reload();
        }}
        if (language === "pt") {{
          showToast("已切换成葡语", "编辑区、预览区和导出选项都已切换到葡语版本。");
        }} else {{
          showToast("已切回中文", "编辑区、预览区和导出选项都已切回中文版本。");
        }}
      }} catch (error) {{
        showToast("切换失败", String(error.message || error));
        button.disabled = false;
        button.textContent = original;
      }}
    }}

    async function runReview(itemId, button) {{
      const feedback = collectReviewFeedback(itemId).trim();
      const mode = collectReviewMode(itemId);
      if (!feedback) {{
        showToast("请先填写反馈", "先告诉 Koko 哪里错了，再开始复盘重做。");
        return;
      }}
      const original = button.textContent;
      button.disabled = true;
      button.textContent = "复盘中...";
      try {{
        const response = await fetch(`/api/items/${{itemId}}/review`, {{
          method: "POST",
          headers: {{ "Content-Type": "application/json" }},
          body: JSON.stringify({{ feedback, mode }}),
        }});
        const data = await response.json();
        if (!response.ok) throw new Error(data.error || "Review failed");
        if (data.job_id) {{
          activeJobId = data.job_id;
          activeReviewItemId = itemId;
          reviewTracker[itemId] = "running";
          pollJob(data.job_id);
        }}
      }} catch (error) {{
        showToast("复盘失败", String(error.message || error));
        button.disabled = false;
        button.textContent = original;
      }}
    }}

    async function runKokoChatEdit(itemId, button) {{
      const message = collectReviewFeedback(itemId).trim();
      const mode = collectChatMode(itemId);
      if (!message) {{
        showToast("先说一句", "告诉 Koko 你想怎么改脚本。");
        return;
      }}
      const root = document.querySelector(`[data-review-item="${{itemId}}"]`);
      const input = root?.querySelector("[data-chat-message]");
      appendChatBubble(itemId, "user", message);
      appendChatBubble(itemId, "koko", mode === "recheck" ? "我去重新看视频..." : "我在改稿...", true);
      if (input) input.value = "";
      const original = button.textContent;
      button.disabled = true;
      button.textContent = mode === "recheck" ? "复盘中..." : "修稿中...";
      try {{
        const response = await fetch(`/api/items/${{itemId}}/chat-edit`, {{
          method: "POST",
          headers: {{ "Content-Type": "application/json" }},
          body: JSON.stringify({{ message, mode }}),
        }});
        const data = await response.json();
        if (!response.ok) throw new Error(data.error || "Koko edit failed");
        if (data.job_id) {{
          activeJobId = data.job_id;
          activeReviewItemId = itemId;
          reviewTracker[itemId] = "running";
          updatePendingChatBubble(itemId, "已经进入复盘流程。");
          pollJob(data.job_id);
          return;
        }}
        updatePendingChatBubble(itemId, data.message || "已修改，正在刷新脚本。");
        showToast("Koko 已修改", data.message || "脚本已经刷新。");
        if (activeJobId) {{
          pollJob(activeJobId);
        }} else {{
          window.location.reload();
        }}
      }} catch (error) {{
        const errorText = String(error.message || error);
        updatePendingChatBubble(itemId, "这次没改成功：" + errorText);
        showToast("Koko 修改失败", errorText);
        if (input) input.value = message;
        button.disabled = false;
        button.textContent = original;
      }}
    }}

    async function generateStoryboardPrompt(itemId, trigger) {{
      const payload = collectItemEdits(itemId) || {{}};
      const original = trigger ? trigger.textContent : "";
      if (trigger) {{
        trigger.disabled = true;
        trigger.textContent = "生成中...";
      }}
      try {{
        const response = await fetch(`/api/items/${{itemId}}/storyboard/prompt`, {{
          method: "POST",
          headers: {{ "Content-Type": "application/json" }},
          body: JSON.stringify(payload),
        }});
        const data = await response.json();
        if (!response.ok) throw new Error(data.error || "Storyboard prompt generation failed");
        const promptField = document.querySelector(`[data-storyboard-prompt="${{itemId}}"]`);
        if (promptField) promptField.value = data.item?.storyboard_prompt || data.storyboard_prompt || "";
        showToast("提示词已生成", "已经根据当前最新脚本生成生图提示词。");
      }} catch (error) {{
        showToast("提示词生成失败", String(error.message || error));
      }} finally {{
        if (trigger) {{
          trigger.disabled = false;
          trigger.textContent = original;
        }}
      }}
    }}

    async function generateStoryboard(itemId, trigger) {{
      const promptField = document.querySelector(`[data-storyboard-prompt="${{itemId}}"]`);
      const prompt = String(promptField?.value || "").trim();
      const payload = collectItemEdits(itemId) || {{}};
      payload.prompt = prompt;
      const original = trigger ? trigger.textContent : "";
      if (trigger) {{
        trigger.disabled = true;
        trigger.textContent = "生成中...";
      }}
      try {{
        const response = await fetch(`/api/items/${{itemId}}/storyboard`, {{
          method: "POST",
          headers: {{ "Content-Type": "application/json" }},
          body: JSON.stringify(payload),
        }});
        const data = await response.json();
        if (!response.ok) throw new Error(data.error || "Storyboard generation failed");
        const item = data.item || {{}};
        const previewUrl = item.storyboard_preview_url || item.storyboard_cover_url || "";
        if (previewUrl) {{
          const panel = trigger?.closest(".storyboard-panel") || document.querySelector(`[data-storyboard-prompt="${{itemId}}"]`)?.closest(".storyboard-panel");
          const versioned = versionedResultUrl(previewUrl, item);
          const previewHtml = `<div class="storyboard-preview-wrap"><img class="storyboard-preview" src="${{escapeHtml(versioned)}}" alt="分解示意图"></div>`;
          const existingWrap = panel?.querySelector(".storyboard-preview-wrap");
          const empty = panel?.querySelector(".storyboard-empty");
          if (existingWrap) existingWrap.outerHTML = previewHtml;
          else if (empty) empty.outerHTML = previewHtml;
          if (trigger) trigger.textContent = "重新生成示意图";
        }}
        if (promptField && item.storyboard_prompt) promptField.value = item.storyboard_prompt;
        showToast("示意图已生成", "已经按当前脚本生成新的分镜示意图，并更新为脚本封面。");
      }} catch (error) {{
        showToast("生成失败", String(error.message || error));
        if (trigger) {{
          trigger.disabled = false;
          trigger.textContent = original;
        }}
      }} finally {{
        if (trigger) {{
          trigger.disabled = false;
          if (trigger.textContent === "生成中...") trigger.textContent = original;
        }}
      }}
    }}

    async function confirmStoryboard(itemId, trigger) {{
      const original = trigger ? trigger.textContent : "";
      if (trigger) {{
        trigger.disabled = true;
        trigger.textContent = "确认中...";
      }}
      try {{
        const response = await fetch(`/api/items/${{itemId}}/storyboard/confirm`, {{
          method: "POST",
          headers: {{ "Content-Type": "application/json" }},
          body: JSON.stringify({{}}),
        }});
        const data = await response.json();
        if (!response.ok) throw new Error(data.error || "Storyboard confirm failed");
        showToast("封面已更新", "这张分镜示意图已经作为脚本封面使用。");
        if (activeJobId) {{
          pollJob(activeJobId);
        }} else {{
          window.location.reload();
        }}
      }} catch (error) {{
        showToast("确认失败", String(error.message || error));
        if (trigger) {{
          trigger.disabled = false;
          trigger.textContent = original;
        }}
      }}
    }}

    function renderItemCard(item, idx, open = false) {{
      const title = escapeHtml(item.title || `视频 ${{idx + 1}}`);
      const primaryEditor = buildPrimaryEditorMarkup(item);
      const review = buildReviewMarkup(item);
      const libraryConfirm = buildLibraryConfirmMarkup(item);
      const toggleButton = item.display_language === "pt"
        ? `<button class="action-link" type="button" data-toggle-language="${{item.id}}" data-language-target="zh">切回中文</button>`
        : `<button class="action-link" type="button" data-toggle-language="${{item.id}}" data-language-target="pt">转换成葡语</button>`;
      const links = [
        (item.zh_docx_url || item.pt_docx_url) ? `<button class="action-link" type="button" data-open-export-modal="${{escapeHtml(item.zh_docx_url || "")}}" data-open-export-modal-pt="${{escapeHtml(item.pt_docx_url || "")}}">导出脚本</button>` : "",
        toggleButton,
      ].join("");
      const sourceVideoUrl = versionedResultUrl(item.source_video_url || "", item);
      const sourceVideoInner = item.source_video_available && sourceVideoUrl
        ? `
            <video class="source-video-frame" data-source-video="${{item.id}}" controls playsinline preload="metadata" src="${{escapeHtml(sourceVideoUrl)}}"></video>
            <div class="source-video-status" data-source-video-status="${{item.id}}">原视频用于对照脚本。点击下方时间线，或脚本表时间，可跳到对应片段。</div>
          `
        : `
            <div class="source-video-empty">
              <div>
                原视频不可用
              </div>
            </div>
            <div class="source-video-status" data-source-video-status="${{item.id}}">${{item.video_url ? `<a class="inline-link" href="${{escapeHtml(item.video_url)}}" target="_blank" rel="noreferrer">打开原链接</a>` : "暂无链接"}}</div>
          `;
      const sourceVideo = item.status === "completed"
        ? `
          <div class="video-verify-layout">
            <div>
              ${{sourceVideoInner}}
            </div>
            ${{item.source_video_available ? buildVideoTimelineMarkup(item) : ""}}
          </div>
        `
        : "";
      const error = item.error ? `<code>${{escapeHtml(item.error)}}</code>` : "";
      return `
        <details class="item-card" data-item-id="${{item.id}}" ${{open ? "open" : ""}}>
          <summary>
            <span>${{idx + 1}}. ${{title}}</span>
            <span>${{escapeHtml(item.status === "completed" ? "已完成" : item.status === "failed" ? "失败" : item.status || "")}}</span>
          </summary>
          <div class="item-body">
            ${{item.status === "completed" ? `
              <div class="result-workbench">
                <div class="result-main">
                  <div class="primary-review-strip" aria-label="编辑脚本和参考视频">
                    <section class="workbench-pane editable-script-pane">
                      <h4>可编辑脚本</h4>
                      <div class="workbench-pane-note">这里是当前脚本 JSON 的整稿编辑版，格式贴近最终 HTML。先在这里直接改，再保存重建脚本。</div>
                      ${{primaryEditor}}
                    </section>
                  </div>
                  <section class="workbench-pane ops-pane">
                    <div class="ops-section followup-actions">
                      <div class="ops-section-title">后续操作</div>
                      <div class="link-row">
                        <button class="action-link primary" type="button" data-save-edits="${{item.id}}">保存当前编辑</button>
                        ${{links}}
                      </div>
                    </div>
                  </section>
                  <section class="workbench-pane ops-pane confirm-pane">
                    <div class="ops-section final-action">
                      ${{libraryConfirm}}
                    </div>
                  </section>
                </div>
                <aside class="workbench-pane assistant-sidebar" aria-label="Koko 修稿助手">
                  <section class="assistant-video-block">
                    <h4>参考视频</h4>
                    ${{sourceVideo}}
                  </section>
                  ${{review}}
                </aside>
              </div>
            ` : `
              <div class="item-sections">
                ${{error}}
              </div>
            `}}
          </div>
        </details>
      `;
    }}

    function renderBatchOverview(data, items) {{
      const effectiveStatus = deriveEffectiveJobStatus(data);
      const total = data.total_items || items.length || 0;
      const completed = data.completed_items || 0;
      const failed = data.failed_items || 0;
      const running = items.filter((item) => itemState(item) === "running").length;
      const waiting = Math.max(0, total - completed - failed - running);
      const finishedCount = completed + failed;
      const systemQueue = data.system_queue || {{}};
      const activeWorkloads = Array.isArray(systemQueue.active_workloads) ? systemQueue.active_workloads : [];
      const globalRunning = Number(systemQueue.running_count || 0);
      const globalQueued = Number(systemQueue.queued_count || 0);
      const currentPosition = Number(systemQueue.current_job_position || 0);
      const currentAhead = Number(systemQueue.current_job_ahead || 0);
      const globalFocus = activeWorkloads[0] || null;
      const currentItem = findCurrentItem(items);
      const effectiveStage = effectiveStatus === "completed"
        ? "completed"
        : effectiveStatus === "failed"
          ? "failed"
          : (data.stage || currentItem?.stage || "queued");
      const allItemsSettled = total > 0 && finishedCount >= total && !hasRunningReview(items);
      if (!currentItem && (effectiveStatus === "completed" || allItemsSettled)) return "";
      const currentIndex = currentItem ? ((items.indexOf(currentItem) >= 0 ? items.indexOf(currentItem) : 0) + 1) : 0;
      const leadItem = currentItem || items[0] || null;
      const stageLabel = STAGE_LABELS[effectiveStage] || STAGE_LABELS[currentItem?.stage] || "等待拆解";
      const stageMessage = data.stage_message || currentItem?.stage_message || data.message || "任务已经创建，系统会按顺序逐条拆解。";
      let subtitle = currentItem
        ? `当前正在拆解 ${{displayVideoName(currentItem, currentItem.index || 0)}}，其余任务会按照提交顺序继续排队。`
        : (effectiveStatus === "completed"
            ? "这批任务已经跑完了，你可以直接查看已完成脚本。"
            : "任务已经创建，系统会按顺序逐条拆解。");
      if (!currentItem && globalFocus) {{
        subtitle = `系统当前正在处理：${{globalFocus.title || "其他任务"}}（${{workloadKindLabel(globalFocus.kind)}}）`;
      }}
      const subtitleNode = currentItem
        ? `<a class="batch-overview-subtitle" href="${{escapeHtml(currentItem.video_url || "")}}" target="_blank" rel="noreferrer">${{escapeHtml(currentItem.video_url || "")}}</a>`
        : `<div class="focus-note">${{escapeHtml(subtitle)}}</div>`;
      const queueHint = effectiveStatus === "queued" && currentPosition
        ? `<div class="focus-note">你的任务当前排队第 ${{currentPosition}} 位，前方还有 ${{currentAhead}} 条任务。</div>`
        : "";
      const systemHint = globalFocus
        ? `<div class="focus-note">系统当前占用：${{escapeHtml(globalFocus.title || "其他任务")}} · ${{escapeHtml(workloadKindLabel(globalFocus.kind))}}${{globalFocus.stage ? ` · ${{escapeHtml(STAGE_LABELS[globalFocus.stage] || globalFocus.stage)}}` : ""}}</div>`
        : "";
      return `
        <section class="batch-overview">
          <div class="batch-overview-top">
            <div class="batch-overview-copy">
              <div class="batch-overview-title">${{leadItem ? escapeHtml(displayVideoName(leadItem, leadItem?.index || 0)) : "批量任务进度"}}</div>
              ${{subtitleNode}}
              <div class="batch-job-meta">
                <span>任务 ID：${{escapeHtml(data.id || "")}}</span>
                <button class="job-copy-btn" type="button" data-copy-text="${{escapeHtml(data.id || "")}}" aria-label="复制任务 ID">⧉</button>
              </div>
            </div>
            <span class="status ${{effectiveStatus === "completed" ? "status-completed" : effectiveStatus === "failed" ? "status-failed" : effectiveStatus === "running" ? "status-running" : "status-queued"}}">${{escapeHtml(jobStatusLabel(effectiveStatus || "queued"))}}</span>
          </div>
          <div class="batch-meta">
            <span class="batch-chip">总数 ${{total}}</span>
            <span class="batch-chip">正在拆解 ${{running}}</span>
            <span class="batch-chip">等待拆解 ${{waiting}}</span>
            <span class="batch-chip">已完成 ${{completed}}</span>
            <span class="batch-chip">失败 ${{failed}}</span>
            <span class="batch-chip">系统运行中 ${{globalRunning}}</span>
            <span class="batch-chip">系统排队 ${{globalQueued}}</span>
          </div>
          ${{queueHint}}
          ${{systemHint}}
          <div class="focus-note">任务 ${{currentIndex || 1}}/${{total || 1}} · ${{escapeHtml(stageLabel)}}</div>
          ${{progressMarkup(effectiveStage, stageMessage, data.id, data)}}
        </section>
      `;
    }}

    function renderQueueList(items, data) {{
      const systemQueue = data?.system_queue || {{}};
      const activeWorkloads = Array.isArray(systemQueue.active_workloads) ? systemQueue.active_workloads : [];
      const globalFocus = activeWorkloads[0] || null;
      const currentPosition = Number(systemQueue.current_job_position || 0);
      const cards = (items || []).map((item, idx) => {{
        const state = itemState(item);
        const title = displayVideoName(item, idx);
        const stageLabel = state === "waiting"
          ? (idx === 0 && currentPosition
              ? `等待拆解 · 全局排队第 ${{currentPosition}} 位`
              : `等待拆解 · 将在前一条完成后继续`)
          : (item.stage_message || STAGE_LABELS[item.stage] || itemStateLabel(item));
        const error = item.error ? `<div class="queue-error">${{escapeHtml(item.error)}}</div>` : "";
        return `
          <article class="queue-card ${{state === "running" ? "current" : ""}}">
            <div class="queue-card-top">
              <div class="queue-index">视频 ${{idx + 1}}</div>
              <span class="queue-status ${{state}}">${{itemStateLabel(item)}}</span>
            </div>
            <h4 class="queue-title">${{escapeHtml(title)}}</h4>
            <div class="queue-url"><span class="queue-link-icon">🔗</span>${{escapeHtml(item.video_url || "")}}</div>
            <div class="queue-stage">${{escapeHtml(stageLabel)}}</div>
            ${{error}}
          </article>
        `;
      }}).join("");
      return `
        <section class="queue-shell">
          <div class="queue-header">
            <h3>任务队列</h3>
            <p>${{globalFocus ? `系统当前正在处理：${{escapeHtml(globalFocus.title || "其他任务")}}` : "按提交顺序显示任务执行状态与排队信息"}}</p>
          </div>
          <div class="queue-list">${{cards}}</div>
        </section>
      `;
    }}

    function renderDetailResults(items) {{
      const detailItems = (items || []).filter((item) => item.status === "completed" || item.status === "failed");
      if (!detailItems.length) return "";
      const controls = detailItems.length > 1 ? `
        <div class="detail-controls">
          <button class="action-link" type="button" data-item-expand="all">全部展开</button>
          <button class="action-link" type="button" data-item-expand="none">全部收起</button>
        </div>
      ` : "";
      const cards = detailItems.map((item, idx) => {{
        const open = Object.prototype.hasOwnProperty.call(itemOpenState, item.id) ? !!itemOpenState[item.id] : item.status === "completed";
        return renderItemCard(item, idx, open);
      }}).join("");
      return `
        <section class="detail-section">
          <div class="detail-header">
            <h3>结果详情</h3>
            ${{controls}}
          </div>
          <div class="item-stack">${{cards}}</div>
        </section>
      `;
    }}

    function renderBatchResults(data) {{
      const items = Array.isArray(data.items) ? data.items : [];
      return `
        <div class="batch-dashboard">
          ${{renderBatchOverview(data, items)}}
          ${{renderQueueList(items, data)}}
          ${{renderDetailResults(items)}}
        </div>
      `;
    }}

    function checkReviewTransitions(items) {{
      for (const item of (items || [])) {{
        if (!item || !item.id) continue;
        const prev = reviewTracker[item.id] || "";
        const next = item.review_status || "";
        const terminalToken = `${{next}}|${{item.updated_at || ""}}|${{item.review_message || ""}}`;
        const watchingThisItem = activeReviewItemId === item.id;
        if ((prev === "running" || watchingThisItem) && next === "completed" && reviewTerminalTracker[item.id] !== terminalToken) {{
          reviewTerminalTracker[item.id] = terminalToken;
          showToast("复盘成功", "Koko 已完成复盘重做，并更新了当前脚本。");
          if (activeReviewItemId === item.id) activeReviewItemId = "";
        }} else if ((prev === "running" || watchingThisItem) && next === "failed" && reviewTerminalTracker[item.id] !== terminalToken) {{
          reviewTerminalTracker[item.id] = terminalToken;
          showToast("复盘失败", item.review_message || "复盘重做没有成功完成。");
          if (activeReviewItemId === item.id) activeReviewItemId = "";
        }}
        reviewTracker[item.id] = next;
      }}
    }}

    function hasRunningReview(items) {{
      return (items || []).some((item) => item && item.review_status === "running");
    }}

    function deriveEffectiveJobStatus(data) {{
      const items = Array.isArray(data?.items) ? data.items : [];
      const totalItems = Number(data?.total_items || items.length || 0);
      const completedItems = Number(data?.completed_items || items.filter((item) => item?.status === "completed").length || 0);
      const failedItems = Number(data?.failed_items || items.filter((item) => item?.status === "failed").length || 0);
      const reviewRunning = hasRunningReview(items);
      const allItemsSettled = totalItems > 0 && completedItems + failedItems >= totalItems;
      if (allItemsSettled && !reviewRunning) return completedItems > 0 ? "completed" : "failed";
      if (items.some((item) => item?.status === "running") || reviewRunning) return "running";
      if (items.some((item) => item?.status === "queued")) return "queued";
      return String(data?.status || "queued").trim() || "queued";
    }}

    function progressMarkup(stage, stageMessage, jobId, data = null) {{
      const index = Math.max(0, STAGE_ORDER.indexOf(stage));
      const percent = stage === "completed" ? 100 : stage === "failed" ? 100 : Math.max(6, Math.round(((index + 1) / STAGE_ORDER.length) * 100));
      const displayMessage = stage === "failed" ? (stageMessage || "Analysis failed.") : (STAGE_COPY[stage] || stageMessage || "Running analysis...");
      const meta = data ? buildArtifactHints(data).map((hint) => `<span class="progress-meta-chip">${{escapeHtml(hint)}}</span>`).join("") : "";
      const steps = ["download", "media_prep", "gemini_analysis", "v2_analysis", "consistency_audit", "targeted_recheck", "arbitration", "final_output"].map((key) => {{
        let cls = "step-pill";
        const keyIndex = STAGE_ORDER.indexOf(key);
        if (stage === "completed" || index > keyIndex) cls += " done";
        else if (stage === key || (stage === "starting" && key === "download")) cls += " active";
        const bubble = stage === "completed" || index > keyIndex ? "✓" : `${{keyIndex + 1}}`;
        return `<div class="${{cls}}"><span class="step-bubble">${{bubble}}</span><span>${{STAGE_LABELS[key]}}</span></div>`;
      }}).join("");
      return `
        <div class="progress-wrap">
          <div class="progress-top">
            <div class="progress-top-copy">
              <span class="progress-kicker">${{escapeHtml(STAGE_LABELS[stage] || "处理中")}}</span>
              <span class="progress-stage-copy">${{escapeHtml(displayMessage)}}</span>
            </div>
            <span class="progress-percent">${{percent}}%</span>
          </div>
          ${{meta ? `<div class="progress-meta">${{meta}}</div>` : ""}}
          <div class="progress-rail"><div class="progress-fill" style="width:${{percent}}%"></div></div>
          <div class="step-list">${{steps}}</div>
          ${{data ? renderThinkingLog(data) : ""}}
        </div>
      `;
    }}

    async function pollJob(jobId, options = {{}}) {{
      activeJobId = jobId;
      persistActiveJobId(jobId);
      updateStopAllButtonState(true);
      if (pollInFlight) {{
        if (options && options.force) {{
          queuedImmediateRepoll = true;
          if (activePollController) activePollController.abort("forced-repoll");
        }}
        return;
      }}
      pollInFlight = true;
      if (jobPollTimer) {{
        clearTimeout(jobPollTimer);
        jobPollTimer = null;
      }}
      let controller = null;
      try {{
      controller = new AbortController();
      activePollController = controller;
      const timeoutId = window.setTimeout(() => controller.abort("poll-timeout"), POLL_REQUEST_TIMEOUT_MS);
      let res;
      try {{
        res = await fetch(`/api/jobs/${{jobId}}?_=${{Date.now()}}`, {{
          cache: "no-store",
          signal: controller.signal,
          headers: {{
            "Cache-Control": "no-store",
            "Pragma": "no-cache"
          }}
        }});
      }} catch (error) {{
        const aborted = error && error.name === "AbortError";
        if (restoringActiveJob) {{
          restoreAttempts += 1;
          if (restoreAttempts < RESTORE_RETRY_LIMIT) {{
            schedulePoll(jobId, RESTORE_RETRY_DELAY_MS);
            return;
          }}
          restoringActiveJob = false;
          restoreAttempts = 0;
          setIdleState();
          showToast("恢复失败", "没能恢复上次任务展示，页面已回到初始状态。");
          return;
        }}
        if (aborted) {{
          schedulePoll(jobId, POLL_RECOVERY_DELAY_MS);
          return;
        }}
        const snapshot = readPersistedActiveJobSnapshot(jobId);
        if (snapshot) {{
          const snapshotResults = renderBatchResults(snapshot);
          if (snapshotResults) setStatus(snapshotResults, deriveEffectiveJobStatus(snapshot) === "completed");
          schedulePoll(jobId, POLL_RECOVERY_DELAY_MS);
          return;
        }}
        setStatus(`<span class="status status-failed">失败</span><br><br><code>${{escapeHtml(String(error.message || error))}}</code>`);
        schedulePoll(jobId, 4000);
        return;
      }} finally {{
        window.clearTimeout(timeoutId);
      }}
      if (!res.ok) {{
        if (restoringActiveJob && res.status >= 500) {{
          restoreAttempts += 1;
          if (restoreAttempts < RESTORE_RETRY_LIMIT) {{
            schedulePoll(jobId, RESTORE_RETRY_DELAY_MS);
            return;
          }}
        }}
        restoringActiveJob = false;
        restoreAttempts = 0;
        persistActiveJobId("");
        setIdleState();
        showToast("无法恢复任务", "这条任务可能已经不存在或已失效。");
        return;
      }}
      let data;
      try {{
        data = await readJsonSafely(res);
      }} catch (error) {{
        if (restoringActiveJob) {{
          restoreAttempts += 1;
          if (restoreAttempts < RESTORE_RETRY_LIMIT) {{
            schedulePoll(jobId, RESTORE_RETRY_DELAY_MS);
            return;
          }}
          restoringActiveJob = false;
          restoreAttempts = 0;
          setIdleState();
          showToast("恢复失败", "服务返回异常，未能恢复上次任务展示。");
          return;
        }}
        const snapshot = readPersistedActiveJobSnapshot(jobId);
        if (snapshot) {{
          const snapshotResults = renderBatchResults(snapshot);
          if (snapshotResults) setStatus(snapshotResults, deriveEffectiveJobStatus(snapshot) === "completed");
          schedulePoll(jobId, POLL_RECOVERY_DELAY_MS);
          return;
        }}
        setStatus(`<span class="status status-failed">失败</span><br><br><code>${{escapeHtml(String(error.message || error))}}</code>`);
        schedulePoll(jobId, 4000);
        return;
      }}
      restoringActiveJob = false;
      restoreAttempts = 0;
      persistActiveJobSnapshot(data);
      updateStudioOverview(data);
      const batchResults = renderBatchResults(data);
      const effectiveStatus = deriveEffectiveJobStatus(data);
      const reviewRunning = hasRunningReview(data.items);
      const totalItems = Number(data.total_items || (Array.isArray(data.items) ? data.items.length : 0) || 0);
      const finishedItems = Number(data.completed_items || 0) + Number(data.failed_items || 0);
      const allItemsSettled = totalItems > 0 && finishedItems >= totalItems;
      checkReviewTransitions(data.items);
      if (effectiveStatus === "completed" || (allItemsSettled && !reviewRunning)) {{
        const completedMessage = reviewRunning
          ? "主分析已完成，复盘任务仍在继续。"
          : (data.message || "分析完成。");
        setStatus(batchResults || `<div class="status-empty"><div class="status-empty-title">分析完成</div><div class="status-empty-copy">${{escapeHtml(completedMessage)}}</div></div>`, true);
        updateStopAllButtonState(reviewRunning);
        if (reviewRunning) {{
          schedulePoll(jobId, 2500);
        }}
        return;
      }}
      if (effectiveStatus === "failed") {{
        const partial = Array.isArray(data.items) && data.items.length ? batchResults : "";
        setStatus(`<span class="status status-failed">失败</span><br><br>${{progressMarkup("failed", data.message || "分析失败。", data.id, data)}}<code>${{escapeHtml(data.error || "未知错误")}}</code>${{partial}}`);
        updateStopAllButtonState(false);
        return;
      }}
      const badge = effectiveStatus === "running" ? "status-running" : "status-queued";
      const effectiveStage = effectiveStatus === "completed"
        ? "completed"
        : effectiveStatus === "failed"
          ? "failed"
          : (data.stage || "queued");
      const runningMarkup = Array.isArray(data.items) && data.items.length
        ? renderBatchResults(data)
        : `${{progressMarkup(effectiveStage, data.stage_message || data.message, data.id, data)}}`;
      setStatus(`<span class="status ${{badge}}">${{jobStatusLabel(effectiveStatus)}}</span><br><br>${{runningMarkup}}`);
      schedulePoll(jobId, 2500);
      }} finally {{
        if (activePollController === controller) activePollController = null;
        pollInFlight = false;
        if (queuedImmediateRepoll && activeJobId === jobId) {{
          queuedImmediateRepoll = false;
          schedulePoll(jobId, 120);
        }}
      }}
    }}

    function requestImmediateJobSync() {{
      if (!activeJobId) return;
      pollJob(activeJobId, {{ force: true }});
    }}

    function renderUnderstandingResults(data) {{
      const items = Array.isArray(data?.items) ? data.items : [];
      const status = deriveEffectiveJobStatus(data || {{}});
      const currentItem = findCurrentItem(items);
      const effectiveStage = status === "completed"
        ? "completed"
        : status === "failed"
          ? "failed"
          : (data?.stage || currentItem?.stage || "queued");
      const summaryCards = items.filter((item) => item.status === "completed" || item.status === "failed").map((item, index) => {{
        const summary = item.understanding_summary || item.result_json?.whole_video_summary || item.result_json?.summary || "";
        const error = item.error ? `<div class="queue-error">${{escapeHtml(item.error)}}</div>` : "";
        return `
          <article class="queue-card ${{item.status === "completed" ? "current" : ""}}">
            <div class="queue-card-top">
              <div class="queue-index">视频 ${{index + 1}}</div>
              <span class="queue-status ${{item.status === "completed" ? "completed" : "failed"}}">${{item.status === "completed" ? "已理解" : "失败"}}</span>
            </div>
            <h4 class="queue-title">${{escapeHtml(displayVideoName(item, index))}}</h4>
            <div class="queue-url"><span class="queue-link-icon">🔗</span>${{escapeHtml(item.video_url || "")}}</div>
            ${{summary ? `<div class="summary-box"><strong>脚本内容概述</strong><br>${{escapeHtml(summary)}}</div>` : ""}}
            ${{error}}
          </article>
        `;
      }}).join("");
      const done = status === "completed";
      const badgeClass = status === "completed" ? "status-completed" : status === "failed" ? "status-failed" : status === "running" ? "status-running" : "status-queued";
      const title = status === "completed" ? "视频理解完成" : status === "failed" ? "视频理解失败" : status === "running" ? "视频理解中" : "排队中";
      return `
        <span class="status ${{badgeClass}}">${{title}}</span>
        <br><br>
        <div class="batch-dashboard">
          ${{!done || !summaryCards ? `
            <section class="batch-overview">
              <div class="batch-overview-top">
                <div class="batch-overview-copy">
                  <div class="batch-overview-title">视频理解进度</div>
                  <div class="focus-note">Koko 正在沿用完整拆解链路做理解校验，最终只输出内容概述。</div>
                </div>
                <span class="status ${{badgeClass}}">${{title}}</span>
              </div>
              ${{progressMarkup(effectiveStage, data?.stage_message || data?.message || currentItem?.stage_message || "正在理解视频。", data?.id || "", data)}}
            </section>
          ` : ""}}
          <section class="queue-shell">
            <div class="queue-header">
              <h3>脚本内容概述</h3>
              <p>这里不展示脚本表、复盘和编辑能力，只保留视频核心内容。</p>
            </div>
            <div class="queue-list">${{summaryCards || `<div class="status-empty"><div class="status-empty-title">还没有完成的视频。</div><div class="status-empty-copy">完成后会在这里显示内容概述。</div></div>`}}</div>
          </section>
        </div>
      `;
    }}

    async function pollUnderstandingJob(jobId) {{
      if (!jobId) return;
      activeUnderstandingJobId = jobId;
      persistActiveUnderstandingJobId(jobId);
      try {{
        const res = await fetch(`/api/jobs/${{jobId}}?_=${{Date.now()}}`, {{
          cache: "no-store",
          headers: {{
            "Cache-Control": "no-store",
            "Pragma": "no-cache"
          }}
        }});
        const data = await readJsonSafely(res);
        if (!res.ok) throw new Error(data.error || "视频理解任务查询失败");
        setUnderstandingStatus(renderUnderstandingResults(data), deriveEffectiveJobStatus(data) === "completed");
        const effectiveStatus = deriveEffectiveJobStatus(data);
        if (effectiveStatus === "completed" || effectiveStatus === "failed") {{
          persistActiveUnderstandingJobId("");
          activeUnderstandingJobId = "";
          return;
        }}
        scheduleUnderstandingPoll(jobId);
      }} catch (error) {{
        setUnderstandingStatus(`<span class="status status-failed">失败</span><br><br><code>${{escapeHtml(String(error.message || error))}}</code>`);
        scheduleUnderstandingPoll(jobId, 4000);
      }}
    }}

    function renderFilterResults(data) {{
      const items = Array.isArray(data?.items) ? data.items : [];
      const matchedLinks = Array.isArray(data?.matched_links) ? data.matched_links.filter(Boolean) : [];
      const statusLabel = data?.status === "completed"
        ? "筛选完成"
        : data?.status === "failed"
          ? "筛选失败"
          : data?.status === "running"
            ? "筛选中"
            : "排队中";
      const badgeClass = data?.status === "completed"
        ? "status-running"
        : data?.status === "failed"
          ? "status-failed"
          : data?.status === "running"
            ? "status-running"
            : "status-queued";
      const matchedMarkup = matchedLinks.length
        ? matchedLinks.map((url, index) => `
            <li>
              <span class="queue-index">命中 ${{index + 1}}</span>
              <div class="queue-url"><span class="queue-link-icon">🔗</span>${{escapeHtml(url)}}</div>
            </li>
          `).join("")
        : `<li><div class="queue-stage">当前还没有同时通过时长、多人物和剧情三轮规则的链接。</div></li>`;
      const itemCards = items.map((item, index) => {{
        const title = item.display_name || parseVideoDisplayName(item.video_url || "", index);
        const bucket = String(item.bucket || "").trim();
        const status = String(item.status || "").trim();
        const thumb = String(item.thumbnail_url || "").trim();
        const reason = String(item.reason || "").trim();
        const stageMessage = String(item.stage_message || "").trim();
        const signals = Array.isArray(item.signals) ? item.signals : [];
        const visual = item.visual && typeof item.visual === "object" ? item.visual : {{}};
        const audio = item.audio && typeof item.audio === "object" ? item.audio : {{}};
        const checks = item.checks && typeof item.checks === "object" ? item.checks : {{}};
        const meta = [];
        if (item.confidence) meta.push(`置信度：${{escapeHtml(item.confidence)}}`);
        const audioStats = [];
        if (audio.source) audioStats.push(`音频：${{escapeHtml(audio.source)}}`);
        if (audio.audio_form) audioStats.push(`形式：${{escapeHtml(audio.audio_form)}}`);
        if (audio.language) audioStats.push(`语言：${{escapeHtml(audio.language)}}`);
        const audioSummary = String(audio.dialogue_summary || "").trim();
        const frameStats = [];
        if (typeof visual.inspected_frames === "number" && visual.inspected_frames) frameStats.push(`抽帧：${{escapeHtml(String(visual.inspected_frames))}}`);
        if (typeof visual.pair_frames === "number") frameStats.push(`男女同框帧：${{escapeHtml(String(visual.pair_frames))}}`);
        if (typeof visual.max_faces_single_frame === "number" && visual.max_faces_single_frame) frameStats.push(`单帧最多人脸：${{escapeHtml(String(visual.max_faces_single_frame))}}`);
        if (typeof visual.male_count === "number" || typeof visual.female_count === "number") {{
          frameStats.push(`男脸：${{escapeHtml(String(visual.male_count || 0))}} · 女脸：${{escapeHtml(String(visual.female_count || 0))}}`);
        }}
        const thumbnailFaceCount = visual.thumbnail_faces && typeof visual.thumbnail_faces.face_count === "number"
          ? visual.thumbnail_faces.face_count
          : 0;
        if (thumbnailFaceCount) frameStats.push(`封面人脸：${{escapeHtml(String(thumbnailFaceCount))}}`);
        const checkRow = (label, check) => {{
          const passed = Boolean(check?.passed);
          const className = passed ? "completed" : "waiting";
          const reasonText = String(check?.reason || "").trim();
          return `<span class="progress-meta-chip ${{className}}">${{escapeHtml(label)}}：${{passed ? "通过" : "不通过"}}${{reasonText ? ` · ${{escapeHtml(reasonText)}}` : ""}}</span>`;
        }};
        const checksMarkup = `
          <div class="progress-meta">
            ${{checkRow("时长", checks.duration_check || {{}})}}
            ${{checkRow("多人物", checks.multi_character_check || {{}})}}
            ${{checkRow("剧情", checks.story_check || {{}})}}
          </div>
        `;
        const signalMarkup = signals.length ? `<div class="progress-meta">${{signals.map((signal) => `<span class="progress-meta-chip">${{escapeHtml(signal)}}</span>`).join("")}}</div>` : "";
        return `
          <article class="queue-card ${{status === "completed" && bucket === "high" ? "current" : ""}}">
            <div class="queue-card-top">
              <div class="queue-index">候选 ${{index + 1}}</div>
              <span class="queue-status ${{filterBucketClass(bucket)}}">${{escapeHtml(filterBucketLabel(bucket))}}</span>
            </div>
            <h4 class="queue-title">${{escapeHtml(title)}}</h4>
            <div class="queue-url"><span class="queue-link-icon">🔗</span>${{escapeHtml(item.video_url || "")}}</div>
            ${{thumb ? `<div class="queue-url"><span class="queue-link-icon">🖼️</span>${{escapeHtml(thumb)}}</div>` : ""}}
            <div class="queue-stage">${{escapeHtml(stageMessage || filterStageLabel(item.stage))}}</div>
            ${{meta.length ? `<div class="queue-stage">${{meta.join(" · ")}}</div>` : ""}}
            ${{audioStats.length ? `<div class="queue-stage">${{audioStats.join(" · ")}}</div>` : ""}}
            ${{audioSummary ? `<div class="queue-stage">音频摘要：${{escapeHtml(audioSummary)}}</div>` : ""}}
            ${{frameStats.length ? `<div class="queue-stage">${{frameStats.join(" · ")}}</div>` : ""}}
            ${{checksMarkup}}
            ${{reason ? `<div class="queue-stage">${{escapeHtml(reason)}}</div>` : ""}}
            ${{signalMarkup}}
            ${{item.error ? `<div class="queue-error">${{escapeHtml(item.error)}}</div>` : ""}}
          </article>
        `;
      }}).join("");
      return `
        <span class="status ${{badgeClass}}">${{statusLabel}}</span>
        <br><br>
        <div class="batch-dashboard">
          <section class="batch-overview">
            <div class="overview-header">
              <div>
                <h3>剧情候选三轮筛选结果</h3>
                <p>${{escapeHtml(data?.message || "Koko 正在基于完整音频信息和开头/中间/结尾三张关键帧做三轮筛选。")}}</p>
              </div>
            </div>
            <div class="overview-stats">
              <div class="overview-stat"><span>输入链接</span><strong>${{Number(data?.input_count || items.length || 0)}}</strong></div>
              <div class="overview-stat"><span>三轮通过</span><strong>${{Number(data?.matched_count || matchedLinks.length || 0)}}</strong></div>
              <div class="overview-stat"><span>当前阶段</span><strong>${{escapeHtml(filterStageLabel(data?.stage || "queued"))}}</strong></div>
            </div>
          </section>
          <section class="queue-shell">
            <div class="queue-header">
              <h3>通过链接</h3>
              <p>这里只保留同时通过时长、多人物和剧情三轮规则的链接，便于后续送去视频拆解。</p>
            </div>
            <ul class="queue-list">${{matchedMarkup}}</ul>
          </section>
          <section class="queue-shell">
            <div class="queue-header">
              <h3>逐条筛选明细</h3>
              <p>基于页面公开信息、完整音频转写和开头/中间/结尾三张关键帧判断。</p>
            </div>
            <div class="queue-list">${{itemCards}}</div>
          </section>
        </div>
      `;
    }}

    async function pollFilterJob(jobId) {{
      activeFilterJobId = jobId;
      persistActiveFilterJobId(jobId);
      try {{
        const res = await fetch(`/api/filter-jobs/${{jobId}}?_=${{Date.now()}}`, {{
          cache: "no-store",
          headers: {{
            "Cache-Control": "no-store",
            "Pragma": "no-cache"
          }}
        }});
        const data = await readJsonSafely(res);
        if (!res.ok) {{
          throw new Error(data.error || "筛选状态获取失败");
        }}
        setFilterStatus(renderFilterResults(data), data.status === "completed");
        if (data.status === "running" || data.status === "queued") {{
          scheduleFilterPoll(jobId, 1800);
          return;
        }}
        if (data.status === "completed") {{
          showToast("筛选完成", `已筛出 ${{Number(data.matched_count || 0)}} 条同时通过三轮规则的视频。`);
          return;
        }}
      }} catch (error) {{
        setFilterStatus(`<span class="status status-failed">筛选失败</span><br><br><code>${{escapeHtml(String(error.message || error))}}</code>`);
      }}
    }}

    submitBtn.addEventListener("click", async () => {{
      const videoUrls = collectUrls();
      if (!videoUrls.length) {{
        setStatus("请先粘贴至少一个公开视频链接。");
        return;
      }}
      submitBtn.disabled = true;
      setStatus("正在创建任务...");
      try {{
        const userPrompt = String(analysisPromptInput?.value || "").trim();
        const res = await fetch("/api/jobs", {{
          method: "POST",
          headers: {{ "Content-Type": "application/json" }},
          body: JSON.stringify({{ video_urls: videoUrls, user_prompt: userPrompt }})
        }});
        const data = await readJsonSafely(res);
        if (!res.ok) {{
          throw new Error(data.error || "任务创建失败");
        }}
        activeJobId = data.id;
        persistActiveJobId(data.id);
        updateStopAllButtonState(true);
        setStudioPanel("split-panel");
        updateStudioOverview({{
          id: data.id,
          status: "queued",
          total_items: videoUrls.length,
          completed_items: 0,
          failed_items: 0,
          items: videoUrls.map((url, index) => ({{ video_url: url, status: "queued", index }})),
          system_queue: data.system_queue || {{}}
        }});
        setStatus(`<span class="status status-queued">排队中</span><br><br>${{progressMarkup("queued", "任务已创建，正在准备分析。", data.id, data)}}`);
        pollJob(data.id);
      }} catch (error) {{
        setStatus(`<span class="status status-failed">失败</span><br><br><code>${{escapeHtml(String(error.message || error))}}</code>`);
        updateStopAllButtonState(false);
      }} finally {{
        submitBtn.disabled = false;
      }}
    }});

    if (understandingSubmitBtn) {{
      understandingSubmitBtn.addEventListener("click", async () => {{
        const videoUrls = collectUrls(understandingInput);
        if (!videoUrls.length) {{
          setUnderstandingStatus("请先粘贴至少一个公开视频链接。");
          return;
        }}
        understandingSubmitBtn.disabled = true;
        setStudioPanel("understanding-panel");
        setUnderstandingStatus("正在创建视频理解任务...");
        try {{
          const res = await fetch("/api/jobs", {{
            method: "POST",
            headers: {{ "Content-Type": "application/json" }},
            body: JSON.stringify({{ video_urls: videoUrls, mode: "understanding" }})
          }});
          const data = await readJsonSafely(res);
          if (!res.ok) {{
            throw new Error(data.error || "视频理解任务创建失败");
          }}
          activeUnderstandingJobId = data.id;
          persistActiveUnderstandingJobId(data.id);
          setUnderstandingStatus(`<span class="status status-queued">排队中</span><br><br>${{progressMarkup("queued", "任务已创建，正在准备理解视频。", data.id, data)}}`);
          pollUnderstandingJob(data.id);
        }} catch (error) {{
          setUnderstandingStatus(`<span class="status status-failed">失败</span><br><br><code>${{escapeHtml(String(error.message || error))}}</code>`);
        }} finally {{
          understandingSubmitBtn.disabled = false;
        }}
      }});
    }}

    if (filterSubmitBtn) {{
      filterSubmitBtn.addEventListener("click", async () => {{
        const directUrls = collectFilterUrls();
        const file = filterFileInput?.files?.[0] || null;
        if (!directUrls.length && !file) {{
          setFilterStatus(`<span class="status status-failed">缺少输入</span><br><br><code>请先粘贴一批包含链接的文本，或上传表格文件。</code>`);
          return;
        }}
        filterSubmitBtn.disabled = true;
        setStudioPanel("filter-panel");
        setFilterStatus(`<span class="status status-running">准备筛选</span><br><br><div class="status-empty"><div class="status-empty-title">Koko 正在解析输入。</div><div class="status-empty-copy">先识别文本和表格里的所有链接，再提取音频与关键帧做三轮筛选。</div></div>`);
        try {{
          let upload = null;
          if (file) {{
            const fileDataBase64 = await readFileAsBase64(file);
            upload = {{
              filename: file.name,
              file_data_base64: fileDataBase64,
            }};
          }}
          const res = await fetch("/api/filter-jobs", {{
            method: "POST",
            headers: {{ "Content-Type": "application/json" }},
            body: JSON.stringify({{
              raw_text: String(filterInput?.value || ""),
              video_urls: directUrls,
              upload,
            }})
          }});
          const data = await readJsonSafely(res);
          if (!res.ok) {{
            throw new Error(data.error || "筛选任务创建失败");
          }}
          activeFilterJobId = data.id;
          persistActiveFilterJobId(data.id);
          setFilterStatus(renderFilterResults(data), false);
          pollFilterJob(data.id);
        }} catch (error) {{
          setFilterStatus(`<span class="status status-failed">筛选失败</span><br><br><code>${{escapeHtml(String(error.message || error))}}</code>`);
        }} finally {{
          filterSubmitBtn.disabled = false;
        }}
      }});
    }}

    if (translationSubmitBtn) {{
      translationSubmitBtn.addEventListener("click", async () => {{
        const urls = extractUrlsFromText(translationInput?.value || "");
        const videoUrl = urls[0] || String(translationInput?.value || "").trim();
        if (!videoUrl) {{
          setTranslationStatus(`<span class="status status-failed">缺少输入</span><br><br><code>请先粘贴一个视频链接。</code>`);
          return;
        }}
        translationSubmitBtn.disabled = true;
        setTranslationStatus("正在创建转译任务...");
        try {{
          const res = await fetch("/api/translation-jobs", {{
            method: "POST",
            headers: {{ "Content-Type": "application/json" }},
            body: JSON.stringify({{ video_url: videoUrl, language: "pt-BR" }})
          }});
          const data = await readJsonSafely(res);
          if (!res.ok) throw new Error(data.error || "转译任务创建失败");
          activeTranslationJobId = data.id;
          persistActiveTranslationJobId(data.id);
          setStudioPanel("translate-panel");
          setTranslationStatus(renderTranslationJob(data));
          pollTranslationJob(data.id);
        }} catch (error) {{
          setTranslationStatus(`<span class="status status-failed">失败</span><br><br><code>${{escapeHtml(String(error.message || error))}}</code>`);
        }} finally {{
          translationSubmitBtn.disabled = false;
        }}
      }});
    }}

    if (stopAllBtn) {{
      stopAllBtn.addEventListener("click", async () => {{
        const original = stopAllBtn.textContent;
        stopAllBtn.disabled = true;
        stopAllBtn.textContent = "停止中...";
        try {{
          const res = await fetch("/api/stop-all", {{
            method: "POST",
            headers: {{ "Content-Type": "application/json" }},
          }});
          const data = await readJsonSafely(res);
          if (!res.ok) {{
            throw new Error(data.error || "停止任务失败");
          }}
          setIdleState();
          setStudioPanel("split-panel");
          showToast("已停止所有任务", `已停止 ${{data.stopped_items || 0}} 条分析，${{data.stopped_reviews || 0}} 条复盘。`);
        }} catch (error) {{
          showToast("停止失败", String(error.message || error));
          updateStopAllButtonState(!!activeJobId);
        }} finally {{
          stopAllBtn.textContent = original;
        }}
      }});
    }}

    studioPanelLinks.forEach((link) => {{
      link.addEventListener("click", (event) => {{
        const panelId = link.getAttribute("data-panel-target");
        if (!panelId) return;
        event.preventDefault();
        setStudioPanel(panelId);
        history.replaceState(null, "", `#${{panelId}}`);
      }});
    }});

    window.addEventListener("hashchange", () => {{
      const panelId = String(window.location.hash || "").replace(/^#/, "").trim();
      if (panelId) setStudioPanel(panelId);
    }});

    window.addEventListener("focus", () => {{
      requestImmediateJobSync();
      if (activeFilterJobId) pollFilterJob(activeFilterJobId);
    }});

    document.addEventListener("visibilitychange", () => {{
      if (document.visibilityState === "visible") {{
        requestImmediateJobSync();
        if (activeFilterJobId) pollFilterJob(activeFilterJobId);
      }}
    }});

    const libraryWorkbenchEntryId = String(pageParams.get("library_entry") || "").trim();
    const restoredJobId = readPersistedActiveJobId();
    if (libraryWorkbenchEntryId) {{
      updateStopAllButtonState(false);
      loadLibraryWorkbench(libraryWorkbenchEntryId);
    }} else if (restoredJobId) {{
      restoringActiveJob = true;
      updateStopAllButtonState(true);
      setStudioPanel("split-panel");
      const restoredSnapshot = readPersistedActiveJobSnapshot(restoredJobId);
      if (restoredSnapshot) {{
        const restoredResults = renderBatchResults(restoredSnapshot);
        if (restoredResults) {{
          setStatus(restoredResults, deriveEffectiveJobStatus(restoredSnapshot) === "completed");
        }}
      }}
      pollJob(restoredJobId);
    }} else {{
      updateStopAllButtonState(false);
      const initialPanelId = String(window.location.hash || "").replace(/^#/, "").trim() || "split-panel";
      setStudioPanel(initialPanelId);
    }}

    const restoredFilterJobId = readPersistedActiveFilterJobId();
    if (restoredFilterJobId) {{
      activeFilterJobId = restoredFilterJobId;
      pollFilterJob(restoredFilterJobId);
    }} else {{
      setFilterIdleState();
    }}

    const restoredUnderstandingJobId = readPersistedActiveUnderstandingJobId();
    if (restoredUnderstandingJobId) {{
      activeUnderstandingJobId = restoredUnderstandingJobId;
      setStudioPanel("understanding-panel");
      pollUnderstandingJob(restoredUnderstandingJobId);
    }} else {{
      setUnderstandingStatus(UNDERSTANDING_IDLE_HTML, true);
    }}

    videoInput.addEventListener("keydown", (event) => {{
      if (event.key === "Enter" && (event.metaKey || event.ctrlKey)) {{
        submitBtn.click();
      }}
    }});

    if (understandingInput && understandingSubmitBtn) {{
      understandingInput.addEventListener("keydown", (event) => {{
        if (event.key === "Enter" && (event.metaKey || event.ctrlKey)) {{
          understandingSubmitBtn.click();
        }}
      }});
    }}

    if (filterInput) {{
      filterInput.addEventListener("keydown", (event) => {{
        if (event.key === "Enter" && (event.metaKey || event.ctrlKey) && filterSubmitBtn) {{
          filterSubmitBtn.click();
        }}
      }});
    }}

    document.addEventListener("keydown", (event) => {{
      if (event.key === "Escape") {{
        closeExportChoice();
      }}
    }});

    document.addEventListener("toggle", (event) => {{
      const detail = event.target;
      if (!(detail instanceof HTMLDetailsElement)) return;
      if (!detail.classList.contains("item-card")) return;
      const itemId = detail.getAttribute("data-item-id") || "";
      if (!itemId) return;
      itemOpenState[itemId] = detail.open;
      ensureDetailIframes(detail.parentElement || detail);
    }}, true);

    document.addEventListener("click", (event) => {{
      const copyBtn = event.target.closest("[data-copy-text]");
      if (copyBtn) {{
        copyText(copyBtn.getAttribute("data-copy-text") || "", "任务 ID 已复制");
        return;
      }}
      const seekBtn = event.target.closest("[data-seek-video]");
      if (seekBtn) {{
        seekSourceVideo(seekBtn.getAttribute("data-seek-video") || "", Number(seekBtn.getAttribute("data-seek-seconds") || "0"));
        return;
      }}
      const expandBtn = event.target.closest("[data-item-expand]");
      if (expandBtn) {{
        const mode = expandBtn.getAttribute("data-item-expand") || "none";
        document.querySelectorAll(".item-card").forEach((detail) => {{
          if (!(detail instanceof HTMLDetailsElement)) return;
          detail.open = mode === "all";
          const itemId = detail.getAttribute("data-item-id") || "";
          if (itemId) itemOpenState[itemId] = detail.open;
        }});
        return;
      }}
      const exportCancel = event.target.closest("#export-choice-cancel");
      if (exportCancel) {{
        closeExportChoice();
        return;
      }}
      const exportZhBtn = event.target.closest("#export-choice-zh");
      if (exportZhBtn) {{
        if (pendingExportChoice?.zh) downloadScript(pendingExportChoice.zh, exportZhBtn);
        closeExportChoice();
        return;
      }}
      const exportPtBtn = event.target.closest("#export-choice-pt");
      if (exportPtBtn) {{
        if (!pendingExportChoice?.pt) {{
          showToast("葡语版本未准备好", "请先点一次“转换成葡语”，再导出葡语版本。");
        }} else {{
          downloadScript(pendingExportChoice.pt, exportPtBtn);
        }}
        closeExportChoice();
        return;
      }}
      const deleteInsightBtn = event.target.closest("[data-delete-insight]");
      if (deleteInsightBtn) {{
        const card = deleteInsightBtn.closest(".structured-insight");
        if (card) card.remove();
        return;
      }}
      const deleteRowBtn = event.target.closest("[data-delete-row]");
      if (deleteRowBtn) {{
        const row = deleteRowBtn.closest("[data-structured-row-index]");
        if (row) row.remove();
        return;
      }}
      const saveBtn = event.target.closest("[data-save-edits]");
      if (saveBtn) {{
        persistItemEdits(saveBtn.getAttribute("data-save-edits"), "save", saveBtn);
        return;
      }}
      const chatModeBtn = event.target.closest("[data-chat-mode-choice]");
      if (chatModeBtn) {{
        const root = chatModeBtn.closest("[data-review-item]");
        if (root) {{
          root.setAttribute("data-chat-mode", chatModeBtn.getAttribute("data-chat-mode-choice") || "minor");
          root.querySelectorAll("[data-chat-mode-choice]").forEach((button) => button.classList.remove("active"));
          chatModeBtn.classList.add("active");
        }}
        return;
      }}
      const chatEditBtn = event.target.closest("[data-chat-edit]");
      if (chatEditBtn) {{
        runKokoChatEdit(chatEditBtn.getAttribute("data-chat-edit"), chatEditBtn);
        return;
      }}
      const storyboardPromptBtn = event.target.closest("[data-generate-storyboard-prompt]");
      if (storyboardPromptBtn) {{
        generateStoryboardPrompt(storyboardPromptBtn.getAttribute("data-generate-storyboard-prompt"), storyboardPromptBtn);
        return;
      }}
      const storyboardBtn = event.target.closest("[data-generate-storyboard]");
      if (storyboardBtn) {{
        generateStoryboard(storyboardBtn.getAttribute("data-generate-storyboard"), storyboardBtn);
        return;
      }}
      const confirmStoryboardBtn = event.target.closest("[data-confirm-storyboard]");
      if (confirmStoryboardBtn) {{
        confirmStoryboard(confirmStoryboardBtn.getAttribute("data-confirm-storyboard"), confirmStoryboardBtn);
        return;
      }}
      const reviewBtn = event.target.closest("[data-run-review]");
      if (reviewBtn) {{
        runReview(reviewBtn.getAttribute("data-run-review"), reviewBtn);
        return;
      }}
      const libraryBtn = event.target.closest("[data-save-library]");
      if (libraryBtn) {{
        persistItemEdits(libraryBtn.getAttribute("data-save-library"), "library", libraryBtn);
        return;
      }}
      const confirmLibraryBtn = event.target.closest("[data-confirm-library]");
      if (confirmLibraryBtn) {{
        confirmLibraryEntry(confirmLibraryBtn.getAttribute("data-confirm-library"), confirmLibraryBtn);
        return;
      }}
      const toggleLanguageBtn = event.target.closest("[data-toggle-language]");
      if (toggleLanguageBtn) {{
        switchDisplayLanguage(
          toggleLanguageBtn.getAttribute("data-toggle-language"),
          toggleLanguageBtn.getAttribute("data-language-target") || "zh",
          toggleLanguageBtn,
        );
        return;
      }}
      const exportModalBtn = event.target.closest("[data-open-export-modal]");
      if (exportModalBtn) {{
        openExportChoice({{
          zh: exportModalBtn.getAttribute("data-open-export-modal") || "",
          pt: exportModalBtn.getAttribute("data-open-export-modal-pt") || "",
        }});
        return;
      }}
      const downloadBtn = event.target.closest("[data-download-script]");
      if (downloadBtn) {{
        downloadScript(downloadBtn.getAttribute("data-download-script"), downloadBtn);
      }}
    }});
  </script>
</body>
</html>"""


def library_html() -> str:
    entries = load_library_entries()
    changed_entries = False
    for entry in entries:
        before = entry.get("preview_image_url") or ""
        hydrate_library_entry_preview(entry)
        if (entry.get("preview_image_url") or "") != before:
            changed_entries = True
    if changed_entries:
        save_library_entries(entries)
    counts = Counter(entry.get("content_type") or DEFAULT_CONTENT_TYPE for entry in entries)
    duration_counts = Counter(entry.get("duration_bucket") or "" for entry in entries)
    location_counts = Counter(entry.get("location_tag") or "" for entry in entries)
    ordered_counts = [(label, counts.get(label, 0)) for label in LIBRARY_FILTER_LABELS]
    filter_options = "".join(
        f"<option value='{html_escape(label)}' data-content-type-option='{html_escape(label)}'>{html_escape(label)} ({count})</option>"
        for label, count in ordered_counts
    )
    duration_filter_options = "".join(
        f"<option value='{html_escape(bucket)}' data-duration-option='{html_escape(bucket)}'>{html_escape(labels.get('zh') or labels.get('pt') or bucket)} ({duration_counts.get(bucket, 0)})</option>"
        for bucket, labels in CREATOR_DURATION_LABELS.items()
    )
    location_filter_options = "".join(
        f"<option value='{html_escape(item['zh'])}' data-location-option='{html_escape(item['zh'])}'>{html_escape(item['zh'])} ({location_counts.get(item['zh'], 0)})</option>"
        for item in LOCATION_TAG_OPTIONS
    )
    content_type_options = "".join(
        f"<option value='{html_escape(label)}'>{html_escape(label)}</option>"
        for label in LIBRARY_FILTER_LABELS
    )
    chips = "".join(
        f"<span class='batch-chip' data-content-type-chip='{html_escape(label)}'>{html_escape(label)} · {count}</span>"
        for label, count in ordered_counts
    ) or "<span class='batch-chip'>No scripts yet</span>"
    cards = []
    for entry in entries:
        entry_id = str(entry.get("entry_id") or "")
        preview_image_url = str(entry.get("preview_image_url") or "").strip()
        source_video_url = f"/results/{entry_id}/source.mp4"
        source_video_exists = (RESULTS_ROOT / entry_id / SOURCE_VIDEO_NAME).exists()
        created_at = format_beijing_time(entry.get("created_at") or "")
        content_type = entry.get("content_type") or DEFAULT_CONTENT_TYPE
        content_type_source = entry.get("content_type_source") or "auto"
        duration_bucket = str(entry.get("duration_bucket") or "")
        duration_label = str(entry.get("duration_label_zh") or entry.get("duration_label_pt") or "")
        location_tag = str(entry.get("location_tag") or "")
        creator_share_url = creator_script_share_url(entry_id)
        manual_badge = "<span class='library-time' data-manual-badge='true'>Manual</span>" if content_type_source == "manual" else "<span class='library-time' data-manual-badge='true' hidden>Manual</span>"
        cards.append(
            f"<article class='library-card' data-entry-id='{html_escape(entry_id)}' data-content-type='{html_escape(content_type)}' data-duration-bucket='{html_escape(duration_bucket)}' data-location-tag='{html_escape(location_tag)}'>"
            "<label class='library-select'><input type='checkbox' data-library-select><span>选择</span></label>"
            "<div class='library-card-top'>"
            f"<button class='batch-chip batch-chip-button' type='button' data-edit-content-type='{html_escape(entry.get('entry_id') or '')}' data-current-content-type='{html_escape(content_type)}'>{html_escape(content_type)}</button>"
            f"<span class='library-time' data-created-at>{html_escape(created_at or 'Unknown time')}</span>"
            "</div>"
            + (
                "<div class='library-meta-pills'>"
                + (f"<span class='batch-chip'>{html_escape(duration_label)}</span>" if duration_label else "")
                + (f"<span class='batch-chip'>{html_escape(location_tag)}</span>" if location_tag else "")
                + "</div>"
            )
            + f"{manual_badge}"
            + f"<a class='video-origin-link' href='{html_escape(entry.get('video_url') or '')}' target='_blank' rel='noreferrer'>{html_escape(entry.get('video_url') or '')}</a>"
            + (
                f"<div class='video-frame-wrap'><img class='video-frame' src='{html_escape(preview_image_url)}' alt='' loading='lazy'></div>"
                if preview_image_url
                else (
                    f"<div class='video-frame-wrap'><video class='video-frame' data-first-frame muted playsinline preload='metadata' src='{html_escape(source_video_url)}'></video></div>"
                    if source_video_exists
                    else "<div class='video-frame-wrap video-frame-empty'><div class='video-frame-empty-copy'>首图暂不可用</div></div>"
                )
            )
            + "<div class='library-copy'>"
            f"<h3>{html_escape(entry.get('title') or 'Untitled Script')}</h3>"
            f"<p>{html_escape(entry.get('whole_video_summary') or '')}</p>"
            "</div>"
            "<div class='link-row'>"
            + (f"<button class='action-link' type='button' data-open-preview='{html_escape(entry.get('html_url') or '')}'>打开预览</button>" if entry.get("html_url") else "")
            + (f"<a class='action-link' href='{html_escape(creator_share_url)}' target='_blank' rel='noopener'>打开 Creator</a>" if creator_share_url else "")
            + (f"<button class='action-link' type='button' data-copy-creator-link='{html_escape(creator_share_url)}'>复制 Creator 链接</button>" if creator_share_url else "")
            + f"<button class='action-link primary' type='button' data-open-library-editor='{html_escape(entry.get('entry_id') or '')}'>编辑脚本</button>"
            + (
                f"<button class='action-link' type='button' data-open-export-modal='{html_escape(entry.get('zh_docx_url') or entry.get('docx_url') or '')}' data-open-export-modal-pt='{html_escape(entry.get('pt_docx_url') or '')}'>导出脚本</button>"
                if entry.get("zh_docx_url") or entry.get("docx_url") or entry.get("pt_docx_url")
                else ""
            )
            + f"<button class='action-link action-link-danger' type='button' data-delete-entry='{html_escape(entry.get('entry_id') or '')}'>删除</button>"
            + "</div>"
            "</article>"
        )
    cards_html = "".join(cards) or "<div class='status-empty-title'>No scripts saved yet.</div>"
    return f"""<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width,initial-scale=1">
  <title>Koko Library</title>
  {FAVICON_LINKS}
  <style>
    @import url('https://fonts.googleapis.com/css2?family=Readex+Pro:wght@300;400;500;600;700&display=swap');
    * {{ box-sizing: border-box; font-family: 'Readex Pro', -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif; }}
    body {{
      margin: 0;
      min-height: 100vh;
      color: #FF8200;
      background:
        radial-gradient(circle at 8% 10%, rgba(255,130,0,.46), transparent 30%),
        radial-gradient(circle at 86% 12%, rgba(249,115,0,.38), transparent 28%),
        radial-gradient(circle at 50% 42%, rgba(255,178,84,.20), transparent 36%),
        linear-gradient(180deg, #FFB15A 0%, #FFD9AF 34%, #FFF1E2 70%, #FFFFFF 100%);
    }}
    .library-shell {{ padding: 24px; }}
    .library-wrap {{
      width: min(1320px, 100%);
      margin: 0 auto;
      border-radius: 34px;
      overflow: hidden;
      border: 1px solid rgba(255,255,255,.72);
      box-shadow: 0 28px 80px rgba(249,115,0,.16);
      background:
        radial-gradient(circle at 12% 16%, rgba(255,130,0,.32), rgba(255,130,0,0) 22%),
        radial-gradient(circle at 86% 18%, rgba(249,115,0,.26), rgba(249,115,0,0) 22%),
        radial-gradient(circle at 70% 62%, rgba(255,244,232,.70), rgba(255,244,232,0) 24%),
        linear-gradient(180deg, rgba(255,207,146,.64) 0%, rgba(255,240,222,.52) 44%, rgba(255,255,255,.62) 100%);
      backdrop-filter: blur(26px);
      -webkit-backdrop-filter: blur(26px);
      padding: 28px;
    }}
    .library-topbar {{ display:flex; align-items:flex-start; justify-content:space-between; gap:20px; flex-wrap:wrap; }}
    .library-topbar h1 {{ margin:0; font-size:clamp(3rem, 7vw, 5.4rem); letter-spacing:-.08em; line-height:.88; }}
    .batch-chip {{
      display:inline-flex; align-items:center; border-radius:999px; padding:8px 12px;
      font-size:12px; font-weight:700; color:#FF8200; background:rgba(255,255,255,.58); border:1px solid rgba(255,130,0,.16);
      backdrop-filter: blur(16px);
      -webkit-backdrop-filter: blur(16px);
    }}
    .batch-chip-button {{
      cursor: pointer;
    }}
    .library-stats {{ display:flex; flex-wrap:wrap; gap:10px; margin-top:18px; }}
    .library-toolbar {{ display:flex; align-items:center; justify-content:space-between; gap:20px; flex-wrap:wrap; margin-top:42px; margin-bottom:14px; }}
    .bulk-actions {{ display:flex; align-items:center; gap:12px; flex-wrap:wrap; }}
    .filter-group {{ display:flex; align-items:flex-end; gap:12px; flex-wrap:wrap; }}
    .filter-label {{ display:flex; flex-direction:column; gap:8px; font-size:13px; font-weight:700; }}
    .filter-select {{
      min-width:180px; border:1px solid rgba(255,130,0,.18); border-radius:16px; padding:12px 14px;
      font-size:14px; color:#FF8200; background:rgba(255,255,255,.64); outline:none;
      backdrop-filter: blur(16px);
      -webkit-backdrop-filter: blur(16px);
    }}
    .library-grid {{ display:grid; grid-template-columns:repeat(auto-fit, minmax(320px, 1fr)); gap:24px; margin-top:20px; }}
    .library-card {{
      position: relative;
      border:1px solid rgba(255,130,0,.16); border-radius:24px; background:rgba(255,255,255,.56);
      padding:22px; display:flex; flex-direction:column; gap:18px;
      box-shadow: 0 18px 42px rgba(249,115,0,.10);
      backdrop-filter: blur(18px);
      -webkit-backdrop-filter: blur(18px);
    }}
    .library-card-top {{ display:flex; align-items:center; justify-content:space-between; gap:14px; flex-wrap:wrap; }}
    .library-meta-pills {{ display:flex; align-items:center; gap:8px; flex-wrap:wrap; margin-top:-8px; }}
    .library-select {{ display:none; align-items:center; gap:8px; width:max-content; border-radius:999px; padding:8px 12px; color:#FF8200; background:rgba(255,255,255,.72); border:1px solid rgba(255,130,0,.18); font-size:12px; font-weight:800; cursor:pointer; }}
    .library-select input {{ accent-color:#FF8200; }}
    body.bulk-mode .library-select {{ display:inline-flex; }}
    body.bulk-mode .library-card {{ outline:1px dashed rgba(255,130,0,.32); }}
    .library-card.selected {{ background:rgba(255,244,232,.82); border-color:rgba(255,130,0,.42); }}
    .library-time {{ font-size:12px; font-weight:700; color:#FF8200; opacity:.88; }}
    .video-origin-link {{ color:#FF8200; text-decoration:none; font-size:13px; line-height:1.55; word-break:break-all; }}
    .video-frame-wrap {{
      border-radius:18px; overflow:hidden; border:1px solid rgba(255,130,0,.16);
      background:rgba(255,244,232,.78); padding:10px;
    }}
    .video-frame-empty {{
      min-height: 180px;
      display: flex;
      align-items: center;
      justify-content: center;
    }}
    .video-frame-empty-copy {{
      font-size: 13px;
      font-weight: 700;
      color: rgba(255,130,0,.72);
      text-align: center;
    }}
    .video-frame {{
      width:100%; height:auto; object-fit:contain; display:block; background:#FFF4E8;
      border-radius:12px;
    }}
    .library-copy {{ display:flex; flex-direction:column; gap:14px; }}
    .library-card h3 {{ margin:0; font-size:22px; line-height:1.34; letter-spacing:0; }}
    .library-card p {{ margin:0; line-height:1.72; font-size:14px; color:#FF8200; }}
    .link-row {{ display:flex; flex-wrap:wrap; gap:12px; margin-top:4px; }}
    .action-link {{
      display:inline-flex; align-items:center; justify-content:center; text-decoration:none;
      border-radius:999px; padding:10px 14px; color:#FF8200; border:1px solid rgba(255,130,0,.18); background:rgba(255,255,255,.72);
      font-weight:700; font-size:13px; cursor:pointer;
      backdrop-filter: blur(14px);
      -webkit-backdrop-filter: blur(14px);
    }}
    .action-link-danger {{ color:#F97300; }}
    .action-link.primary {{ background:#FF8200; border-color:#FF8200; color:#fff; }}
    .home-link {{ color:#FF8200; text-decoration:none; font-weight:700; }}
    .confirm-overlay {{
      position: fixed; inset: 0; display: none; align-items: center; justify-content: center;
      background: rgba(255, 130, 0, 0.14); backdrop-filter: blur(10px); -webkit-backdrop-filter: blur(10px);
      padding: 24px; z-index: 40;
    }}
    .confirm-overlay.open {{ display: flex; }}
    .confirm-dialog {{
      width: min(460px, 100%); border-radius: 28px; border: 1px solid rgba(255,255,255,.72);
      background:
        radial-gradient(circle at 18% 14%, rgba(255,130,0,.22), rgba(255,130,0,0) 22%),
        linear-gradient(180deg, rgba(255,233,208,.84) 0%, rgba(255,255,255,.74) 100%);
      box-shadow: 0 24px 60px rgba(249,115,0,.18);
      backdrop-filter: blur(22px); -webkit-backdrop-filter: blur(22px);
      padding: 24px;
    }}
    .confirm-dialog h3 {{ margin: 0 0 10px; font-size: 24px; line-height: 1.2; color: #FF8200; }}
    .confirm-dialog p {{ margin: 0; font-size: 14px; line-height: 1.7; color: #FF8200; opacity: .92; }}
    .confirm-actions {{ display:flex; justify-content:flex-end; gap:12px; margin-top:20px; flex-wrap:wrap; }}
  </style>
</head>
<body>
  <main class="library-shell">
    <section class="library-wrap">
      <div class="library-topbar">
        <div>
          <button class="action-link" id="back-home" type="button">← Back to Koko</button>
          <h1>Script Library</h1>
          <div class="library-stats">{chips}</div>
        </div>
      </div>
      <div class="library-toolbar">
        <div class="bulk-actions">
          <button class="action-link primary" id="creator-sync-now" type="button">立刻同步创作者中心</button>
          <button class="action-link" id="bulk-mode-toggle" type="button">批量删除</button>
          <button class="action-link action-link-danger" id="bulk-delete-approve" type="button" hidden disabled>删除选中 0</button>
          <button class="action-link" id="bulk-cancel" type="button" hidden>取消</button>
        </div>
        <div class="filter-group" aria-label="脚本筛选">
          <label class="filter-label">
            <span>第一层：类型</span>
            <select id="content-filter" class="filter-select">
              <option value="">全部类型</option>
              {filter_options}
            </select>
          </label>
          <label class="filter-label">
            <span>第二层：时间</span>
            <select id="duration-filter" class="filter-select">
              <option value="">全部时间</option>
              {duration_filter_options}
            </select>
          </label>
          <label class="filter-label">
            <span>第三层：地点</span>
            <select id="location-filter" class="filter-select">
              <option value="">全部地点</option>
              {location_filter_options}
            </select>
          </label>
        </div>
      </div>
      <div class="library-grid">{cards_html}</div>
    </section>
  </main>
  <div class="confirm-overlay" id="delete-confirm-overlay" aria-hidden="true">
    <div class="confirm-dialog" role="dialog" aria-modal="true" aria-labelledby="delete-confirm-title">
      <h3 id="delete-confirm-title">Delete script?</h3>
      <p>This will remove the script from the library and permanently delete its saved files.</p>
      <div class="confirm-actions">
        <button class="action-link" id="delete-confirm-cancel" type="button">Cancel</button>
        <button class="action-link action-link-danger" id="delete-confirm-approve" type="button">Delete</button>
      </div>
    </div>
  </div>
  <div class="confirm-overlay" id="export-choice-overlay" aria-hidden="true">
    <div class="confirm-dialog" role="dialog" aria-modal="true" aria-labelledby="export-choice-title">
      <h3 id="export-choice-title">导出脚本</h3>
      <p>请选择要导出的版本。葡语版本会保持严格直译，不做概括改写。</p>
      <div class="confirm-actions">
        <button class="action-link" id="export-choice-cancel" type="button">取消</button>
        <button class="action-link" id="export-choice-zh" type="button">导出中文版本</button>
        <button class="action-link" id="export-choice-pt" type="button">导出葡语版本</button>
      </div>
    </div>
  </div>
  <div class="confirm-overlay" id="content-type-overlay" aria-hidden="true">
    <div class="confirm-dialog" role="dialog" aria-modal="true" aria-labelledby="content-type-title">
      <h3 id="content-type-title">修改分类</h3>
      <p>你可以直接人工指定脚本库里的分类。人工分类会优先保留，不会被后续自动分类覆盖。</p>
      <div class="confirm-actions" style="justify-content:stretch; flex-direction:column; align-items:stretch;">
        <select id="content-type-select" class="filter-select">
          {content_type_options}
        </select>
        <div class="confirm-actions" style="margin-top:16px;">
          <button class="action-link" id="content-type-cancel" type="button">取消</button>
          <button class="action-link primary" id="content-type-save" type="button">保存分类</button>
        </div>
      </div>
    </div>
  </div>
  <script>
    const backHomeButton = document.getElementById("back-home");
    const contentFilter = document.getElementById("content-filter");
    const durationFilter = document.getElementById("duration-filter");
    const locationFilter = document.getElementById("location-filter");
    const deleteConfirmOverlay = document.getElementById("delete-confirm-overlay");
    const deleteConfirmCancel = document.getElementById("delete-confirm-cancel");
    const deleteConfirmApprove = document.getElementById("delete-confirm-approve");
    const exportChoiceOverlay = document.getElementById("export-choice-overlay");
    const exportChoiceCancel = document.getElementById("export-choice-cancel");
    const exportChoiceZh = document.getElementById("export-choice-zh");
    const exportChoicePt = document.getElementById("export-choice-pt");
    const contentTypeOverlay = document.getElementById("content-type-overlay");
    const contentTypeSelect = document.getElementById("content-type-select");
    const contentTypeCancel = document.getElementById("content-type-cancel");
    const contentTypeSave = document.getElementById("content-type-save");
    const creatorSyncNow = document.getElementById("creator-sync-now");
    const bulkModeToggle = document.getElementById("bulk-mode-toggle");
    const bulkDeleteApprove = document.getElementById("bulk-delete-approve");
    const bulkCancel = document.getElementById("bulk-cancel");
    const libraryFilterLabels = {json.dumps(LIBRARY_FILTER_LABELS, ensure_ascii=False)};
    const durationFilterLabels = {json.dumps(CREATOR_DURATION_LABELS, ensure_ascii=False)};
    const locationFilterLabels = {json.dumps([item["zh"] for item in LOCATION_TAG_OPTIONS], ensure_ascii=False)};
    let pendingDeleteButton = null;
    let pendingExportChoice = null;
    let pendingContentTypeEntryId = "";

    function closeDeleteConfirm() {{
      if (!deleteConfirmOverlay) return;
      deleteConfirmOverlay.classList.remove("open");
      deleteConfirmOverlay.setAttribute("aria-hidden", "true");
      pendingDeleteButton = null;
    }}

    function openDeleteConfirm(button) {{
      if (!deleteConfirmOverlay) return;
      pendingDeleteButton = button;
      deleteConfirmOverlay.classList.add("open");
      deleteConfirmOverlay.setAttribute("aria-hidden", "false");
    }}

    function closeExportChoice() {{
      if (!exportChoiceOverlay) return;
      exportChoiceOverlay.classList.remove("open");
      exportChoiceOverlay.setAttribute("aria-hidden", "true");
      pendingExportChoice = null;
    }}

    function openExportChoice(zhUrl, ptUrl) {{
      if (!exportChoiceOverlay) return;
      pendingExportChoice = {{ zh: zhUrl || "", pt: ptUrl || "" }};
      exportChoiceOverlay.classList.add("open");
      exportChoiceOverlay.setAttribute("aria-hidden", "false");
    }}

    function closeContentTypeOverlay() {{
      if (!contentTypeOverlay) return;
      contentTypeOverlay.classList.remove("open");
      contentTypeOverlay.setAttribute("aria-hidden", "true");
      pendingContentTypeEntryId = "";
    }}

    function openContentTypeOverlay(entryId, currentType) {{
      if (!contentTypeOverlay || !contentTypeSelect) return;
      pendingContentTypeEntryId = entryId || "";
      contentTypeSelect.value = currentType || "{DEFAULT_CONTENT_TYPE}";
      contentTypeOverlay.classList.add("open");
      contentTypeOverlay.setAttribute("aria-hidden", "false");
    }}

    function refreshLibraryCounts() {{
      const counts = Object.fromEntries(libraryFilterLabels.map((label) => [label, 0]));
      const durationCounts = Object.fromEntries(Object.keys(durationFilterLabels).map((label) => [label, 0]));
      const locationCounts = Object.fromEntries(locationFilterLabels.map((label) => [label, 0]));
      document.querySelectorAll(".library-card").forEach((card) => {{
        const label = card.getAttribute("data-content-type") || "{DEFAULT_CONTENT_TYPE}";
        counts[label] = (counts[label] || 0) + 1;
        const duration = card.getAttribute("data-duration-bucket") || "";
        if (duration) durationCounts[duration] = (durationCounts[duration] || 0) + 1;
        const location = card.getAttribute("data-location-tag") || "";
        if (location) locationCounts[location] = (locationCounts[location] || 0) + 1;
      }});
      document.querySelectorAll("[data-content-type-option]").forEach((option) => {{
        const label = option.getAttribute("data-content-type-option") || "";
        option.textContent = `${{label}} (${{counts[label] || 0}})`;
      }});
      document.querySelectorAll("[data-content-type-chip]").forEach((chip) => {{
        const label = chip.getAttribute("data-content-type-chip") || "";
        chip.textContent = `${{label}} · ${{counts[label] || 0}}`;
      }});
      document.querySelectorAll("[data-duration-option]").forEach((option) => {{
        const bucket = option.getAttribute("data-duration-option") || "";
        const label = durationFilterLabels[bucket]?.zh || durationFilterLabels[bucket]?.pt || bucket;
        option.textContent = `${{label}} (${{durationCounts[bucket] || 0}})`;
      }});
      document.querySelectorAll("[data-location-option]").forEach((option) => {{
        const label = option.getAttribute("data-location-option") || "";
        option.textContent = `${{label}} (${{locationCounts[label] || 0}})`;
      }});
    }}

    function applyLibraryEntryUpdate(entry) {{
      if (!entry || !entry.entry_id) return;
      const card = document.querySelector(`.library-card button[data-edit-content-type="${{entry.entry_id}}"]`)?.closest(".library-card");
      if (!card) return;
      const label = entry.content_type || "{DEFAULT_CONTENT_TYPE}";
      card.setAttribute("data-content-type", label);
      const button = card.querySelector("[data-edit-content-type]");
      if (button) {{
        button.textContent = label;
        button.setAttribute("data-current-content-type", label);
      }}
      const manualBadge = card.querySelector("[data-manual-badge]");
      if (manualBadge) {{
        const isManual = (entry.content_type_source || "") === "manual";
        manualBadge.hidden = !isManual;
      }}
      refreshLibraryCounts();
      applyLibraryFilter();
    }}

    async function downloadScript(url, button) {{
      if (!url) return;
      const originalText = button ? button.textContent : "";
      if (button) {{
        button.textContent = "导出中...";
        button.disabled = true;
      }}
      try {{
        const link = document.createElement("a");
        link.href = url;
        link.download = url.split("/").pop() || "script_export.docx";
        link.rel = "noopener";
        link.target = "_self";
        document.body.appendChild(link);
        link.click();
        link.remove();
        setTimeout(() => {{
          alert("导出脚本已开始下载。");
        }}, 150);
      }} catch (error) {{
        alert("导出脚本失败，请重试。");
      }} finally {{
        if (button) {{
          button.textContent = originalText;
          button.disabled = false;
        }}
      }}
    }}

    function selectedLibraryIds() {{
      return Array.from(document.querySelectorAll("[data-library-select]:checked"))
        .map((input) => input.closest(".library-card")?.getAttribute("data-entry-id") || "")
        .filter(Boolean);
    }}

    function updateBulkControls() {{
      const bulkMode = document.body.classList.contains("bulk-mode");
      const selectedCount = selectedLibraryIds().length;
      if (bulkModeToggle) bulkModeToggle.hidden = bulkMode;
      if (bulkDeleteApprove) {{
        bulkDeleteApprove.hidden = !bulkMode;
        bulkDeleteApprove.disabled = selectedCount === 0;
        bulkDeleteApprove.textContent = `删除选中 ${{selectedCount}}`;
      }}
      if (bulkCancel) bulkCancel.hidden = !bulkMode;
      document.querySelectorAll(".library-card").forEach((card) => {{
        const checkbox = card.querySelector("[data-library-select]");
        card.classList.toggle("selected", Boolean(checkbox && checkbox.checked));
      }});
    }}

    function setBulkMode(enabled) {{
      document.body.classList.toggle("bulk-mode", enabled);
      if (!enabled) {{
        document.querySelectorAll("[data-library-select]").forEach((input) => {{ input.checked = false; }});
      }}
      updateBulkControls();
    }}

    function applyLibraryFilter() {{
      const contentValue = contentFilter ? contentFilter.value : "";
      const durationValue = durationFilter ? durationFilter.value : "";
      const locationValue = locationFilter ? locationFilter.value : "";
      document.querySelectorAll(".library-card").forEach((card) => {{
        const contentType = card.getAttribute("data-content-type") || "";
        const duration = card.getAttribute("data-duration-bucket") || "";
        const location = card.getAttribute("data-location-tag") || "";
        const visible = (!contentValue || contentValue === contentType)
          && (!durationValue || durationValue === duration)
          && (!locationValue || locationValue === location);
        card.style.display = visible ? "" : "none";
      }});
    }}

    if (contentFilter) {{
      contentFilter.addEventListener("change", applyLibraryFilter);
    }}
    if (durationFilter) {{
      durationFilter.addEventListener("change", applyLibraryFilter);
    }}
    if (locationFilter) {{
      locationFilter.addEventListener("change", applyLibraryFilter);
    }}

    if (creatorSyncNow) {{
      creatorSyncNow.addEventListener("click", async () => {{
        const originalText = creatorSyncNow.textContent;
        creatorSyncNow.textContent = "同步中...";
        creatorSyncNow.disabled = true;
        try {{
          const response = await fetch("/api/library/sync-creator-center", {{
            method: "POST",
            headers: {{ "Content-Type": "application/json" }},
            body: JSON.stringify({{ force: true }}),
          }});
          const data = await response.json();
          if (!response.ok || data.ok === false) throw new Error(data.error || "同步失败");
          creatorSyncNow.textContent = "已同步";
          alert(`创作者中心已同步。当前脚本数：${{data.entries_count || "-"}}`);
          setTimeout(() => {{ creatorSyncNow.textContent = originalText; }}, 1600);
        }} catch (error) {{
          alert(`同步创作者中心失败：${{error.message || error}}`);
          creatorSyncNow.textContent = originalText;
        }} finally {{
          creatorSyncNow.disabled = false;
        }}
      }});
    }}

    if (bulkModeToggle) {{
      bulkModeToggle.addEventListener("click", () => setBulkMode(true));
    }}

    if (bulkCancel) {{
      bulkCancel.addEventListener("click", () => setBulkMode(false));
    }}

    if (bulkDeleteApprove) {{
      bulkDeleteApprove.addEventListener("click", async () => {{
        const entryIds = selectedLibraryIds();
        if (!entryIds.length) return;
        if (!window.confirm(`确定删除选中的 ${{entryIds.length}} 条脚本吗？这个操作会删除脚本库记录和对应结果文件。`)) return;
        const originalText = bulkDeleteApprove.textContent;
        bulkDeleteApprove.textContent = "删除中...";
        bulkDeleteApprove.disabled = true;
        try {{
          const response = await fetch("/api/library/batch-delete", {{
            method: "POST",
            headers: {{ "Content-Type": "application/json" }},
            body: JSON.stringify({{ entry_ids: entryIds }}),
          }});
          const data = await response.json();
          if (!response.ok) throw new Error(data.error || "Batch delete failed");
          window.location.reload();
        }} catch (error) {{
          alert("批量删除失败，请重试。");
          bulkDeleteApprove.textContent = originalText;
          bulkDeleteApprove.disabled = false;
        }}
      }});
    }}

    document.querySelectorAll("[data-library-select]").forEach((input) => {{
      input.addEventListener("change", updateBulkControls);
    }});
    updateBulkControls();

    document.querySelectorAll("video[data-first-frame]").forEach((video) => {{
      const setFirstFrame = () => {{
        try {{
          video.currentTime = 0.05;
        }} catch (error) {{}}
      }};
      video.addEventListener("loadedmetadata", setFirstFrame, {{ once: true }});
      video.addEventListener("seeked", () => video.pause(), {{ once: true }});
      video.load();
    }});

    if (backHomeButton) {{
      backHomeButton.addEventListener("click", () => {{
        window.location.assign("/");
      }});
    }}

    if (deleteConfirmCancel) {{
      deleteConfirmCancel.addEventListener("click", closeDeleteConfirm);
    }}

    if (exportChoiceCancel) {{
      exportChoiceCancel.addEventListener("click", closeExportChoice);
    }}

    if (contentTypeCancel) {{
      contentTypeCancel.addEventListener("click", closeContentTypeOverlay);
    }}

    if (deleteConfirmOverlay) {{
      deleteConfirmOverlay.addEventListener("click", (event) => {{
        if (event.target === deleteConfirmOverlay) closeDeleteConfirm();
      }});
    }}

    if (exportChoiceOverlay) {{
      exportChoiceOverlay.addEventListener("click", (event) => {{
        if (event.target === exportChoiceOverlay) closeExportChoice();
      }});
    }}

    if (contentTypeOverlay) {{
      contentTypeOverlay.addEventListener("click", (event) => {{
        if (event.target === contentTypeOverlay) closeContentTypeOverlay();
      }});
    }}

    document.addEventListener("keydown", (event) => {{
      if (event.key === "Escape") {{
        closeDeleteConfirm();
        closeExportChoice();
        closeContentTypeOverlay();
        setBulkMode(false);
      }}
    }});

    if (contentTypeSave) {{
      contentTypeSave.addEventListener("click", async () => {{
        if (!pendingContentTypeEntryId || !contentTypeSelect) {{
          closeContentTypeOverlay();
          return;
        }}
        const originalText = contentTypeSave.textContent;
        contentTypeSave.textContent = "保存中...";
        contentTypeSave.disabled = true;
        try {{
          const response = await fetch(`/api/library/${{pendingContentTypeEntryId}}/content-type`, {{
            method: "POST",
            headers: {{ "Content-Type": "application/json" }},
            body: JSON.stringify({{ content_type: contentTypeSelect.value }}),
          }});
          const data = await response.json();
          if (!response.ok) throw new Error(data.error || "Update failed");
          applyLibraryEntryUpdate(data.entry || null);
          closeContentTypeOverlay();
          contentTypeSave.textContent = originalText;
          contentTypeSave.disabled = false;
        }} catch (error) {{
          alert("修改分类失败，请重试。");
          contentTypeSave.textContent = originalText;
          contentTypeSave.disabled = false;
        }}
      }});
    }}

    if (deleteConfirmApprove) {{
      deleteConfirmApprove.addEventListener("click", async () => {{
        const deleteBtn = pendingDeleteButton;
        if (!deleteBtn) {{
          closeDeleteConfirm();
          return;
        }}
        const entryId = deleteBtn.getAttribute("data-delete-entry");
        if (!entryId) {{
          closeDeleteConfirm();
          return;
        }}
        const originalText = deleteBtn.textContent;
        deleteBtn.textContent = "删除中...";
        deleteBtn.disabled = true;
        closeDeleteConfirm();
        try {{
          const response = await fetch(`/api/library/${{entryId}}`, {{ method: "DELETE" }});
          if (!response.ok) throw new Error("Delete failed");
          window.location.reload();
        }} catch (error) {{
          alert("删除失败，请重试。");
          deleteBtn.textContent = originalText;
          deleteBtn.disabled = false;
        }}
      }});
    }}

    document.addEventListener("click", async (event) => {{
      const selectLabel = event.target.closest(".library-select");
      if (selectLabel) {{
        setTimeout(updateBulkControls, 0);
        return;
      }}
      const previewBtn = event.target.closest("[data-open-preview]");
      if (previewBtn) {{
        const url = previewBtn.getAttribute("data-open-preview");
        if (url) window.location.assign(url);
        return;
      }}
      const copyCreatorBtn = event.target.closest("[data-copy-creator-link]");
      if (copyCreatorBtn) {{
        const url = copyCreatorBtn.getAttribute("data-copy-creator-link") || "";
        const originalText = copyCreatorBtn.textContent;
        try {{
          await navigator.clipboard.writeText(url);
          copyCreatorBtn.textContent = "已复制";
        }} catch (error) {{
          window.prompt("复制这个 Creator 链接：", url);
        }}
        setTimeout(() => {{ copyCreatorBtn.textContent = originalText; }}, 1400);
        return;
      }}
      const editBtn = event.target.closest("[data-open-library-editor]");
      if (editBtn) {{
        const entryId = editBtn.getAttribute("data-open-library-editor");
        if (entryId) window.location.assign(`/studio?library_entry=${{encodeURIComponent(entryId)}}#split-panel`);
        return;
      }}
      const deleteBtn = event.target.closest("[data-delete-entry]");
      if (deleteBtn) {{
        openDeleteConfirm(deleteBtn);
        return;
      }}
      const editTypeBtn = event.target.closest("[data-edit-content-type]");
      if (editTypeBtn) {{
        openContentTypeOverlay(
          editTypeBtn.getAttribute("data-edit-content-type") || "",
          editTypeBtn.getAttribute("data-current-content-type") || "{DEFAULT_CONTENT_TYPE}",
        );
        return;
      }}
      const exportModalBtn = event.target.closest("[data-open-export-modal]");
      if (exportModalBtn) {{
        openExportChoice(
          exportModalBtn.getAttribute("data-open-export-modal") || "",
          exportModalBtn.getAttribute("data-open-export-modal-pt") || "",
        );
        return;
      }}
      if (event.target.closest("#export-choice-zh")) {{
        if (pendingExportChoice?.zh) {{
          downloadScript(pendingExportChoice.zh, exportChoiceZh);
        }}
        closeExportChoice();
        return;
      }}
      if (event.target.closest("#export-choice-pt")) {{
        if (!pendingExportChoice?.pt) {{
          alert("请先在结果页点一次“转换成葡语”，再导出葡语版本。");
        }} else {{
          downloadScript(pendingExportChoice.pt, exportChoicePt);
        }}
        closeExportChoice();
      }}
    }});
  </script>
</body>
</html>"""


CREATOR_ADMIN_TAB_ROUTES = {
    "/creator-admin": "creators",
    "/creator-admin/imports": "creators",
    "/creator-admin/creators": "creators",
    "/creator-admin/accounts": "creators",
    "/creator-admin/analytics": "analytics",
    "/creator-admin/submissions": "submissions",
    "/creator-admin/intakes": "intakes",
}


def creator_admin_tab_for_path(path: str) -> str | None:
    if re.fullmatch(r"/creator-admin/creators/[0-9a-f]{32}", path):
        return "creators"
    return CREATOR_ADMIN_TAB_ROUTES.get(path)


def creator_admin_html(initial_tab: str = "scripts", *, library_mode: bool = False) -> str:
    if initial_tab not in {"scripts", "imports", "creators", "analytics", "submissions", "intakes"}:
        initial_tab = "scripts"
    initial_tab_json = json.dumps(initial_tab, ensure_ascii=False)
    library_mode_json = "true" if library_mode else "false"
    template = """<!doctype html><html lang="zh-CN"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1"><title>Koko Creator 运营后台</title>__FAVICON_LINKS__<style>
@import url('https://fonts.googleapis.com/css2?family=Readex+Pro:wght@300;400;500;600;700&display=swap');
*{{box-sizing:border-box;font-family:'Readex Pro',-apple-system,BlinkMacSystemFont,'Segoe UI',sans-serif}}body{{margin:0;min-height:100vh;background:radial-gradient(circle at 8% 8%,rgba(255,130,0,.34),transparent 30%),linear-gradient(180deg,#ffbf75 0%,#fff4e8 42%,#fff 100%);color:#1f1f1f;overflow-x:hidden}}button,input,textarea,select{{font:inherit;min-width:0}}.shell{{width:min(1240px,100%);margin:0 auto;padding:24px;overflow:hidden}}.panel{{min-width:0;overflow:hidden;border:1px solid rgba(255,255,255,.78);border-radius:34px;background:rgba(255,255,255,.62);box-shadow:0 28px 80px rgba(249,115,0,.16);backdrop-filter:blur(22px);padding:24px}}.top{{display:flex;align-items:flex-start;justify-content:space-between;gap:18px;flex-wrap:wrap}}.kicker{{display:inline-flex;border:1px solid rgba(255,130,0,.24);border-radius:999px;padding:8px 12px;background:rgba(255,255,255,.72);color:#ff8200;font-size:12px;font-weight:800}}h1{{margin:14px 0 8px;font-size:clamp(34px,6vw,64px);line-height:.95;letter-spacing:-.05em;color:#ff8200}}.copy{{margin:0;color:#99520f;line-height:1.6;font-weight:650}}.nav,.ops-tabs{{display:flex;gap:10px;flex-wrap:wrap}}.ops-tabs{{margin:22px 0 6px}}a.btn,button,.ops-tabs a{{display:inline-flex;align-items:center;justify-content:center;border:1px solid rgba(255,130,0,.24);border-radius:999px;min-height:42px;padding:0 15px;background:rgba(255,255,255,.76);color:#ff8200;font-weight:850;text-decoration:none;cursor:pointer}}button.primary,.ops-tabs a.active{{border-color:#ff8200;background:#ff8200;color:#fff}}button.danger{{color:#c9481e}}button:disabled{{opacity:.52;cursor:not-allowed}}.toolbar{{display:grid;grid-template-columns:minmax(0,1fr) auto auto auto;gap:8px;margin:14px 0 10px;align-items:center}}.toolbar input{{min-height:40px;padding:9px 12px;border-radius:14px}}.toolbar button{{min-height:40px;padding:0 14px}}.quick-filters{{display:flex;flex-wrap:wrap;gap:10px;margin:0 0 14px}}.quick-filter-chip{{border:1px solid rgba(255,130,0,.20);border-radius:999px;padding:8px 14px;background:rgba(255,255,255,.78);color:#99520f;font-size:13px;font-weight:850;cursor:pointer;transition:all .18s ease}}.quick-filter-chip.active{{border-color:#ff8200;background:#ff8200;color:#fff}}.quick-filter-chip small{{font-size:11px;font-weight:800;opacity:.8}}.creator-form{{display:grid;grid-template-columns:repeat(auto-fit,minmax(min(205px,100%),1fr));gap:8px;margin:14px 0;padding:10px;border:1px solid rgba(255,130,0,.16);border-radius:18px;background:rgba(255,255,255,.55);overflow:hidden;align-items:start}}.creator-form>*{{min-width:0}}input,textarea,select{{width:100%;border:1px solid rgba(255,130,0,.22);border-radius:14px;background:rgba(255,255,255,.84);padding:9px 12px;outline:none;color:#1f1f1f;min-height:40px}}textarea{{min-height:68px;resize:vertical}}.creator-form button{{min-height:40px;padding:0 18px;align-self:center;justify-self:start;min-width:168px}}.creator-form button.primary{{box-shadow:0 10px 22px rgba(255,130,0,.18)}}.status{{min-height:22px;color:#99520f;font-size:13px;font-weight:800}}.grid{{display:grid;gap:12px;margin-top:12px}}.card{{display:grid;grid-template-columns:34px 92px minmax(0,1fr) auto;gap:12px;align-items:center;border:1px solid rgba(255,130,0,.16);border-radius:22px;background:rgba(255,255,255,.74);padding:12px;box-shadow:0 14px 34px rgba(249,115,0,.10)}}.card img{{width:92px;aspect-ratio:9/16;border-radius:14px;object-fit:cover;background:#2a1d16}}.card h3{{margin:0 0 7px;font-size:18px;line-height:1.28;color:#1f1f1f}}.card p{{margin:0;color:#6f737a;font-size:13px;line-height:1.45;display:-webkit-box;-webkit-line-clamp:2;-webkit-box-orient:vertical;overflow:hidden}}.meta{{display:flex;gap:7px;flex-wrap:wrap;margin-top:9px}}.pill{{border:1px solid rgba(255,130,0,.24);border-radius:999px;padding:5px 9px;color:#ff8200;background:#fff7f0;font-size:12px;font-weight:800}}.pill.off{{color:#777;background:#f3f3f3;border-color:#ddd}}.actions{{display:flex;gap:8px;flex-wrap:wrap;justify-content:flex-end}}.creator-card{{min-width:0;overflow:hidden;border:1px solid rgba(255,130,0,.18);border-radius:26px;background:rgba(255,255,255,.78);padding:16px;box-shadow:0 14px 34px rgba(249,115,0,.10)}}.creator-head{{display:grid;grid-template-columns:72px minmax(0,1fr) auto;gap:14px;align-items:center}}.avatar{{width:72px;height:72px;border-radius:50%;object-fit:cover;background:linear-gradient(135deg,#ffbd64,#ff6500)}}.creator-name{{margin:0;font-size:22px;color:#1f1f1f}}.script-mini{{display:grid;grid-template-columns:52px minmax(0,1fr) auto;gap:10px;align-items:center;margin-top:10px;padding:9px;border-radius:16px;background:#fff7f0;border:1px solid rgba(255,130,0,.14)}}.script-mini img{{width:52px;height:66px;border-radius:10px;object-fit:cover;background:#2a1d16}}.script-mini b{{display:block;font-size:13px;line-height:1.3}}.script-mini span,.small{{color:#6f737a;font-size:12px;line-height:1.35;overflow-wrap:anywhere}}.submission-summary{{display:grid;grid-template-columns:minmax(0,1fr) auto;gap:12px;align-items:start;margin:16px 0;padding:16px;border-radius:24px;background:rgba(255,255,255,.80);border:1px solid rgba(255,130,0,.18);box-shadow:0 14px 34px rgba(249,115,0,.10)}}.submission-summary h2{{margin:0 0 6px;font-size:24px;color:#1f1f1f}}.submission-count{{min-width:84px;border-radius:20px;background:#ff8200;color:white;text-align:center;padding:12px;font-weight:950}}.submission-count b{{display:block;font-size:28px;line-height:1}}.submission-groups{{display:grid;gap:10px;margin-top:12px}}.submission-group{{border:1px solid rgba(255,130,0,.16);border-radius:18px;background:#fffaf5;padding:12px}}.submission-group h3{{margin:0 0 8px;font-size:16px}}.submission-row{{display:grid;grid-template-columns:minmax(0,1fr) auto;gap:8px;align-items:center;border-top:1px solid rgba(255,130,0,.12);padding-top:9px;margin-top:9px}}.submission-row a{{color:#ff8200;font-weight:850;word-break:break-all}}.import-panel{{min-width:0;overflow:hidden;display:grid;gap:12px;margin:16px 0;padding:16px;border-radius:24px;background:rgba(255,255,255,.78);border:1px solid rgba(255,130,0,.18);box-shadow:0 14px 34px rgba(249,115,0,.10)}}.import-form{{display:grid;grid-template-columns:minmax(0,1.5fr) minmax(0,1fr) auto;gap:10px;align-items:center}}.progress{{height:10px;border-radius:999px;background:#ffe3d1;overflow:hidden}}.progress span{{display:block;height:100%;width:0;background:linear-gradient(90deg,#ff9b24,#ff5f00);transition:width .25s ease}}.import-result{{display:grid;grid-template-columns:minmax(0,1fr) auto;gap:10px;align-items:center;border:1px solid rgba(255,130,0,.14);border-radius:18px;background:#fffaf5;padding:12px}}.import-result.failed{{border-color:#ffb0a0;background:#fff3f0}}.import-result b{{display:block;line-height:1.35}}.import-result code{{color:#99520f;font-size:12px;word-break:break-all}}details{{margin-top:8px}}summary{{cursor:pointer;color:#ff8200;font-weight:900}}.login{{min-height:100vh;display:grid;place-items:center;padding:20px}}.login form,.modal-card{{width:min(520px,100%);border:1px solid rgba(255,130,0,.20);border-radius:30px;background:rgba(255,255,255,.78);padding:24px;box-shadow:0 24px 60px rgba(249,115,0,.18);backdrop-filter:blur(20px)}}.login h1{{text-align:center;font-size:42px}}.modal{{position:fixed;inset:0;display:none;align-items:center;justify-content:center;background:rgba(47,27,9,.42);padding:14px;z-index:20}}.modal.open{{display:flex}}.modal-card{{max-height:92vh;overflow:auto;background:#fffaf5}}.modal-card h2{{margin:0 0 14px;color:#ff8200}}.fields{{display:grid;gap:10px}}.row{{display:grid;grid-template-columns:1fr 1fr;gap:10px}}.modal-actions{{display:flex;justify-content:flex-end;gap:10px;margin-top:16px;flex-wrap:wrap}}.empty{{padding:24px;border:1px dashed rgba(255,130,0,.34);border-radius:18px;text-align:center;color:#99520f;background:rgba(255,255,255,.72)}}@media(max-width:820px){{.toolbar,.creator-form,.import-form{{grid-template-columns:1fr}}.card{{grid-template-columns:28px 76px minmax(0,1fr)}}.card img{{width:76px}}.actions{{grid-column:2/4;justify-content:flex-start}}.row,.creator-head,.submission-summary,.submission-row,.import-result,.script-mini{{grid-template-columns:1fr}}.creator-form button{{width:100%;justify-self:stretch}}}}
.realtime-status{{display:inline-flex;align-items:center;border:1px solid rgba(22,163,74,.24);border-radius:999px;min-height:42px;padding:0 14px;background:rgba(240,253,244,.9);color:#15803d;font-size:12px;font-weight:950}}.realtime-status.active{{box-shadow:0 0 0 4px rgba(22,163,74,.08)}}.creator-searchbar{{display:grid;grid-template-columns:180px minmax(0,1fr) 132px;gap:0;margin:14px auto 18px;width:min(820px,100%);border:1px solid rgba(255,130,0,.28);border-radius:999px;overflow:hidden;background:#fff;box-shadow:0 14px 34px rgba(255,130,0,.12)}}.creator-searchbar select,.creator-searchbar input,.creator-searchbar button{{border:0;border-radius:0;min-height:52px;background:#fff}}.creator-searchbar select,.creator-searchbar button{{background:#ff8200;color:white;font-weight:900;text-align:center}}.creator-searchbar input{{padding:0 22px;font-size:16px}}.creator-toolbar{{display:flex;gap:10px;align-items:center;justify-content:space-between;flex-wrap:wrap;margin:10px 0 18px}}.creator-toolbar-left,.creator-toolbar-right{{display:flex;gap:10px;align-items:center;flex-wrap:wrap}}.creator-results-meta{{font-weight:900;color:#1f1f1f}}.creator-results-meta b{{color:#ff8200}}.creator-card-grid{{display:grid;grid-template-columns:repeat(3,minmax(0,1fr));gap:18px}}.creator-tile{{min-width:0;border:1px solid rgba(31,41,55,.10);border-radius:14px;background:rgba(255,255,255,.82);padding:18px;box-shadow:0 10px 28px rgba(31,41,55,.05);cursor:pointer;transition:transform .16s ease,box-shadow .16s ease,border-color .16s ease}}.creator-tile:hover{{transform:translateY(-2px);border-color:rgba(255,130,0,.34);box-shadow:0 18px 38px rgba(255,130,0,.12)}}.creator-tile-head{{display:grid;grid-template-columns:58px minmax(0,1fr);gap:14px;align-items:center}}.creator-tile .avatar{{width:58px;height:58px}}.creator-tile-name{{margin:0;font-size:17px;line-height:1.2;color:#111827;white-space:nowrap;overflow:hidden;text-overflow:ellipsis}}.creator-badge{{display:inline-flex;align-items:center;border-radius:6px;background:#dcfce7;color:#16a34a;padding:4px 10px;font-size:13px;font-weight:900}}.creator-field-lines{{display:grid;gap:7px;margin-top:14px;padding-top:14px;border-top:1px solid rgba(31,41,55,.10)}}.creator-field-line{{display:grid;grid-template-columns:72px minmax(0,1fr);gap:8px;color:#8a8f98;font-weight:800}}.creator-field-line span:last-child{{color:#1f1f1f;white-space:nowrap;overflow:hidden;text-overflow:ellipsis}}.creator-stats-row{{display:grid;grid-template-columns:repeat(2,minmax(0,1fr));gap:10px;margin-top:14px;padding-top:14px;border-top:1px solid rgba(31,41,55,.10)}}.creator-stat{{display:grid;gap:3px;color:#8a8f98;font-weight:800}}.creator-stat b{{color:#1f1f1f;font-size:16px}}.creator-pagination{{display:flex;justify-content:flex-end;gap:8px;align-items:center;margin-top:18px}}.creator-pagination button{{min-width:40px;border-radius:10px}}.creator-detail-top{{display:flex;align-items:center;justify-content:space-between;gap:14px;flex-wrap:wrap;margin-bottom:16px}}.creator-detail-hero{{display:grid;grid-template-columns:88px minmax(0,1fr) auto;gap:16px;align-items:center;border:1px solid rgba(255,130,0,.16);border-radius:24px;background:rgba(255,255,255,.80);padding:18px}}.creator-detail-hero .avatar{{width:88px;height:88px}}.creator-detail-grid{{display:grid;grid-template-columns:1fr 1fr;gap:14px;margin-top:14px}}.creator-modal-form{{display:grid;grid-template-columns:1fr 1fr;gap:10px}}.creator-modal-form textarea{{grid-column:1/-1}}.modal-card.wide{{width:min(760px,100%)}}@media(max-width:980px){{.creator-card-grid{{grid-template-columns:repeat(2,minmax(0,1fr))}}.creator-detail-grid{{grid-template-columns:1fr}}}}@media(max-width:640px){{.creator-searchbar{{grid-template-columns:1fr}}.creator-card-grid{{grid-template-columns:1fr}}.creator-detail-hero,.creator-modal-form{{grid-template-columns:1fr}}}}
.share-line{{display:grid;gap:4px;margin-top:10px;padding:9px 10px;border:1px solid rgba(255,130,0,.14);border-radius:14px;background:#fffaf5}}.share-line b{{color:#99520f;font-size:12px}}.share-line a{{color:#ff5f00;font-size:12px;font-weight:850;word-break:break-all;text-decoration:none}}.multi-select{{min-height:92px}}.tag-text{{display:inline-flex;gap:6px;flex-wrap:wrap}}.creator-detail-stack{{display:grid;gap:14px;margin-top:14px}}.creator-recommend-panel{{grid-column:1/-1}}.creator-script-grid{{display:grid;grid-template-columns:repeat(5,minmax(0,1fr));gap:12px;margin-top:12px}}.creator-script-card{{display:grid;grid-template-rows:auto minmax(0,1fr);gap:10px;padding:10px;border:1px solid rgba(255,130,0,.20);border-radius:18px;background:#fffaf5;text-decoration:none;color:#1f1f1f;min-width:0;transition:transform .16s ease,border-color .16s ease}}.creator-script-card:hover{{transform:translateY(-2px);border-color:#ff8200}}.creator-script-card img{{width:100%;aspect-ratio:1/1;border-radius:14px;object-fit:cover;background:#f4eee7}}.creator-script-card b{{display:block;font-size:15px;line-height:1.25}}.creator-script-card p{{margin:6px 0 0;color:#6f737a;font-size:12px;line-height:1.45;display:-webkit-box;-webkit-line-clamp:4;-webkit-box-orient:vertical;overflow:hidden}}.recommend-footer{{display:flex;justify-content:center;margin-top:12px}}.metrics-grid{{display:grid;grid-template-columns:repeat(4,minmax(0,1fr));gap:12px}}.metric-card{{padding:14px;border:1px solid rgba(255,130,0,.16);border-radius:18px;background:#fffaf5}}.metric-card span{{display:block;color:#6f737a;font-size:12px;font-weight:800}}.metric-card b{{display:block;margin:6px 0 4px;font-size:28px;line-height:1;color:#1f1f1f}}.delta-up{{color:#e11d48;font-weight:950}}.delta-down{{color:#16a34a;font-weight:950}}.mini-import{{display:grid;grid-template-columns:minmax(0,1fr) auto;gap:10px;align-items:end;margin-top:12px}}.feed-stats{{display:grid;gap:10px}}.feed-stats-head{{display:flex;justify-content:space-between;align-items:center;gap:12px;flex-wrap:wrap}}.feed-stats-title{{margin:0;color:#1f1f1f;font-size:20px}}.feed-stat-grid{{display:grid;grid-template-columns:repeat(3,minmax(0,1fr));gap:10px}}.feed-stat-card{{padding:12px;border:1px solid rgba(255,130,0,.16);border-radius:16px;background:#fffaf5}}.feed-stat-card span{{display:block;color:#6f737a;font-size:12px;font-weight:850}}.feed-stat-card b{{display:block;margin-top:4px;font-size:26px;color:#ff5f00;line-height:1}}.feed-date-table{{display:grid;border:1px solid rgba(255,130,0,.14);border-radius:16px;overflow:hidden;background:#fff}}.feed-date-row{{display:grid;grid-template-columns:minmax(0,1fr) 92px 92px;border-top:1px solid rgba(255,130,0,.10)}}.feed-date-row:first-child{{border-top:0}}.feed-date-row>*{{padding:9px 10px;font-size:13px}}.feed-date-row b{{color:#1f1f1f}}.feed-time-grid{{display:grid;grid-template-columns:1fr 1fr;gap:8px}}.feed-time-field{{display:grid;gap:4px}}.feed-time-field label{{font-size:12px;font-weight:900;color:#99520f}}.feed-list{{display:grid;gap:12px;margin-top:12px}}.feed-card{{display:grid;grid-template-columns:150px minmax(0,1.7fr) minmax(240px,.85fr);gap:14px;padding:14px;border:1px solid rgba(255,130,0,.18);border-radius:20px;background:rgba(255,255,255,.82)}}.feed-card img{{width:100%;aspect-ratio:1/1;border-radius:16px;object-fit:cover;background:#f4eee7}}.feed-card h4{{margin:0 0 8px;font-size:18px;line-height:1.25}}.feed-card p{{margin:0;color:#4b5563;font-size:14px;line-height:1.55}}.feed-controls{{display:grid;gap:8px}}.return-preview{{display:grid;grid-template-columns:64px minmax(0,1fr);gap:10px;align-items:center;padding:8px;border-radius:14px;background:#fff7f0;border:1px solid rgba(255,130,0,.12)}}.return-preview img{{width:64px;height:64px;border-radius:12px;object-fit:cover}}@media(max-width:1120px){{.creator-script-grid{{grid-template-columns:repeat(3,minmax(0,1fr))}}.metrics-grid,.feed-stat-grid{{grid-template-columns:repeat(2,minmax(0,1fr))}}.feed-card{{grid-template-columns:120px minmax(0,1fr)}}.feed-controls{{grid-column:1/-1}}}}@media(max-width:640px){{.creator-script-grid,.metrics-grid,.mini-import,.feed-card,.feed-stat-grid,.feed-time-grid{{grid-template-columns:1fr}}}}

.creator-view-toggle{{display:inline-flex;border:1px solid rgba(255,130,0,.24);border-radius:999px;overflow:hidden;background:rgba(255,255,255,.72)}}.creator-view-toggle button{{border:0;border-radius:0;min-height:38px;padding:0 14px;background:transparent;color:#99520f}}.creator-view-toggle button.active{{background:#ff8200;color:#fff}}.creator-row-list{{display:grid;border:1px solid rgba(31,41,55,.10);border-radius:18px;overflow:hidden;background:rgba(255,255,255,.84);box-shadow:0 12px 30px rgba(31,41,55,.06)}}.creator-row-head,.creator-row{{display:grid;grid-template-columns:minmax(240px,2fr) 82px 82px 98px 98px minmax(220px,1.5fr);gap:14px;align-items:center;padding:14px 16px;border-top:1px solid rgba(31,41,55,.10)}}.creator-row-head{{border-top:0;background:rgba(255,250,245,.96);color:#6f737a;font-size:13px;font-weight:950}}.creator-row{{cursor:pointer;transition:background .16s ease}}.creator-row:hover{{background:#fff7f0}}.creator-row-main{{display:grid;grid-template-columns:48px minmax(0,1fr);gap:12px;align-items:center;min-width:0}}.creator-row .avatar{{width:48px;height:48px}}.creator-row-title{{display:grid;gap:4px;min-width:0}}.creator-row-title b{{font-size:16px;color:#111827;white-space:nowrap;overflow:hidden;text-overflow:ellipsis}}.creator-row-num{{font-size:20px;font-weight:950;color:#ff5f00}}.creator-row-tags{{display:flex;gap:7px;flex-wrap:wrap;min-width:0}}.analytics-summary{{display:grid;grid-template-columns:repeat(6,minmax(0,1fr));gap:10px;margin:12px 0 16px}}.analytics-row-head,.analytics-row{{display:grid;grid-template-columns:minmax(220px,1.7fr) 92px 112px 112px 130px 130px minmax(210px,1.4fr);gap:12px;align-items:center;padding:14px 16px;border-top:1px solid rgba(31,41,55,.10)}}.analytics-row-head{{border-top:0;background:#fffaf5;color:#6f737a;font-size:13px;font-weight:950}}.analytics-row{{background:#fff;cursor:default}}.analytics-row details{{grid-column:1/-1;margin:0;padding:12px;border-radius:14px;background:#fff7f0;border:1px solid rgba(255,130,0,.12)}}.analytics-mini{{display:flex;gap:6px;flex-wrap:wrap}}.analytics-mini a{{color:#ff5f00;font-weight:850;word-break:break-all}}.analytics-shell{{display:grid;gap:16px}}.analytics-kpis{{display:grid;grid-template-columns:repeat(6,minmax(0,1fr));gap:10px}}.analytics-core-kpis{{grid-template-columns:1.05fr 1.15fr 1fr}}.analytics-kpis .metric-card{{min-height:104px;background:linear-gradient(180deg,#fff,#fff8f1);border-color:rgba(255,130,0,.18)}}.analytics-core-kpis .metric-card{{min-height:132px;padding:18px;border-radius:24px}}.analytics-core-kpis .metric-card span{{font-size:14px;color:#99520f}}.analytics-core-kpis .metric-card b{{font-size:46px;color:#ff5f00}}.analytics-core-kpis .metric-card small{{display:block;margin-top:8px;font-size:13px;line-height:1.45;color:#6f737a}}.analytics-layout{{display:grid;grid-template-columns:minmax(0,1.04fr) minmax(360px,.96fr);gap:14px;align-items:start}}.analytics-panel{{min-width:0;border:1px solid rgba(255,130,0,.16);border-radius:24px;background:rgba(255,255,255,.82);padding:16px;box-shadow:0 12px 28px rgba(249,115,0,.08);overflow:hidden}}.analytics-panel-title{{display:flex;align-items:flex-end;justify-content:space-between;gap:12px;flex-wrap:wrap;margin-bottom:12px}}.analytics-panel-title h3{{margin:0;font-size:22px;line-height:1.15;color:#1f1f1f}}.analytics-panel-title small{{font-weight:850;color:#99520f}}.analytics-funnel{{display:grid;grid-template-columns:repeat(6,minmax(0,1fr));gap:9px}}.analytics-step{{position:relative;min-height:118px;padding:13px 12px;border-radius:18px;background:#fff7f0;border:1px solid rgba(255,130,0,.16)}}.analytics-step span{{display:block;color:#99520f;font-size:12px;font-weight:900;line-height:1.25}}.analytics-step b{{display:block;margin-top:8px;font-size:30px;line-height:1;color:#ff5f00}}.analytics-step small{{color:#6f737a;font-weight:850}}.analytics-bar{{height:7px;margin:12px 0 7px;border-radius:999px;background:#ffe1cc;overflow:hidden}}.analytics-bar i{{display:block;height:100%;border-radius:999px;background:linear-gradient(90deg,#ff9b24,#ff5f00)}}.analytics-heat{{width:100%;border-collapse:separate;border-spacing:0 8px}}.analytics-heat th{{text-align:left;color:#6f737a;font-size:12px;font-weight:950;padding:0 10px 3px}}.analytics-heat td{{padding:10px;border-top:1px solid rgba(255,130,0,.12);border-bottom:1px solid rgba(255,130,0,.12);background:#fffaf5;font-weight:850}}.analytics-heat td:first-child{{border-left:1px solid rgba(255,130,0,.12);border-radius:14px 0 0 14px;color:#1f1f1f}}.analytics-heat td:last-child{{border-right:1px solid rgba(255,130,0,.12);border-radius:0 14px 14px 0}}.heat{{display:inline-flex;min-width:46px;justify-content:center;border-radius:999px;padding:5px 9px;background:#fff;color:#99520f}}.h1,.h2{{background:#fff0e5}}.h3{{background:#ffd8bd}}.h4,.h5{{background:#ff8200;color:#fff}}.analytics-tabs{{display:flex;gap:8px;flex-wrap:wrap}}.analytics-tabs button{{min-height:34px;padding:0 12px;font-size:13px;background:#fffaf5;color:#99520f}}.analytics-tabs button.active{{background:#ff8200;color:white;border-color:#ff8200}}.analytics-table{{display:grid;border:1px solid rgba(31,41,55,.08);border-radius:18px;overflow:hidden;background:#fff}}.analytics-table-head,.analytics-table-row{{display:grid;grid-template-columns:minmax(220px,1.6fr) 82px 128px 94px 76px 74px;gap:10px;align-items:center;padding:12px 14px;border-top:1px solid rgba(31,41,55,.08)}}.analytics-table-head{{border-top:0;background:#fffaf5;color:#6f737a;font-size:12px;font-weight:950}}.analytics-table-row{{min-height:68px}}.analytics-table-row button{{min-height:32px;padding:0 10px;border-radius:999px;font-size:12px}}.analytics-rank{{display:grid;gap:9px}}.analytics-rank a{{display:grid;grid-template-columns:32px minmax(0,1fr);gap:10px;align-items:center;padding:10px;border:1px solid rgba(255,130,0,.14);border-radius:16px;background:#fffaf5;color:#1f1f1f;text-decoration:none}}.analytics-rank span{{width:32px;height:32px;border-radius:50%;display:grid;place-items:center;background:#ff8200;color:#fff;font-weight:950}}.analytics-rank b{{white-space:nowrap;overflow:hidden;text-overflow:ellipsis}}.analytics-rank small{{grid-column:2;color:#6f737a;font-weight:800}}@media(max-width:1120px){{.analytics-summary,.analytics-kpis,.analytics-core-kpis{{grid-template-columns:repeat(3,minmax(0,1fr))}}.analytics-layout{{grid-template-columns:1fr}}.analytics-funnel{{grid-template-columns:repeat(3,minmax(0,1fr))}}.analytics-row-head,.analytics-table-head{{display:none}}.analytics-row,.analytics-table-row{{grid-template-columns:1fr;gap:8px}}}}@media(max-width:760px){{.creator-row-head{{display:none}}.creator-row{{grid-template-columns:1fr;gap:10px}}.creator-row-num::before{{content:attr(data-label);display:inline-block;margin-right:8px;color:#6f737a;font-size:12px;font-weight:900}}.analytics-summary,.analytics-kpis,.analytics-core-kpis{{grid-template-columns:1fr}}.analytics-funnel{{grid-template-columns:1fr 1fr}}.analytics-panel{{padding:12px;border-radius:20px}}}}

.analytics-core-detail{{margin-top:-4px;border:1px solid rgba(255,130,0,.16);border-radius:22px;background:rgba(255,255,255,.82);padding:12px 14px;box-shadow:0 10px 24px rgba(249,115,0,.06)}}.analytics-core-detail summary{{display:flex;align-items:center;justify-content:space-between;gap:12px;color:#ff5f00;font-weight:950;cursor:pointer}}.analytics-core-detail summary small{{color:#99520f;font-weight:850}}.analytics-hour-table{{display:grid;margin-top:12px;border:1px solid rgba(31,41,55,.08);border-radius:16px;overflow:hidden;background:#fff}}.analytics-hour-head,.analytics-hour-row{{display:grid;grid-template-columns:150px 88px 106px 118px 102px 112px;gap:10px;align-items:center;padding:11px 12px;border-top:1px solid rgba(31,41,55,.08)}}.analytics-hour-head{{border-top:0;background:#fffaf5;color:#6f737a;font-size:12px;font-weight:950}}.analytics-hour-row b{{color:#ff5f00}}@media(max-width:820px){{.analytics-hour-head{{display:none}}.analytics-hour-row{{grid-template-columns:1fr 1fr;gap:8px}}.analytics-hour-row span::before,.analytics-hour-row b::before{{content:attr(data-label);display:block;color:#6f737a;font-size:11px;font-weight:900;margin-bottom:2px}}}}

.analytics-hour-detail{{grid-column:1/-1;display:grid;grid-template-columns:repeat(4,minmax(0,1fr));gap:10px;margin-top:8px}}.analytics-detail-group{{display:grid;gap:8px;align-content:start;padding:10px;border:1px solid rgba(255,130,0,.12);border-radius:14px;background:#fffaf5}}.analytics-detail-group h4{{margin:0;font-size:13px;color:#99520f}}.analytics-person{{display:grid;gap:3px;padding:8px;border-radius:12px;background:#fff;border:1px solid rgba(31,41,55,.06)}}.analytics-person b{{color:#1f1f1f;font-size:13px}}.analytics-person small{{color:#6f737a;line-height:1.35;overflow-wrap:anywhere}}.analytics-hour-row details{{grid-column:1/-1;margin:0}}.analytics-hour-row summary{{font-size:13px}}@media(max-width:980px){{.analytics-hour-detail{{grid-template-columns:1fr 1fr}}}}@media(max-width:640px){{.analytics-hour-detail{{grid-template-columns:1fr}}}}

</style></head><body><main id="app"></main><div class="modal" id="edit-modal"><form class="modal-card" id="edit-form"><h2>修改脚本标签</h2><p class="copy" style="margin-bottom:14px">这里只调整脚本分类标签；完整脚本内容请从列表里的“修改脚本”进入内容中台。</p><div class="fields"><select name="content_type"></select></div><div class="modal-actions"><button type="button" id="edit-cancel">取消</button><button class="primary" type="submit">保存标签</button></div></form></div><div class="modal" id="creator-modal"><form class="modal-card wide" id="creator-form"><h2>导入创作者</h2><p class="copy" style="margin-bottom:14px">粘贴 Kwai 作者主页，补充账号和标签后保存。资料抓取仍由后端处理。</p><div class="creator-modal-form"><input name="kwai_url" placeholder="Kwai 作者主页，例如 kwai.com/@CarlosDeiOficial"><input name="display_name" placeholder="作者名称"><input name="kwai_id" placeholder="Kwai ID"><input name="phone" placeholder="手机号/登录电话"><input name="uid" placeholder="UID"><select name="poc"><option value="">POC 待分配</option><option>denghaoqing</option><option>zhaozhe</option></select><select name="category"></select><select class="multi-select" name="identity" multiple><option>夫妻</option><option>情侣</option><option>家庭</option><option>朋友</option></select><select class="multi-select" name="location" multiple><option>家里</option><option>乡村</option><option>城市</option></select><select name="cooperation_level"><option>待标注</option><option>高</option><option>中</option><option>低</option><option>待观察</option></select><textarea name="creator_description" placeholder="具体作者描述：可手动输入作者风格、限制、偏好等"></textarea></div><div class="modal-actions"><button type="button" id="creator-cancel">取消</button><button class="primary" type="submit">保存创作者</button></div></form></div><div class="modal" id="creator-tags-modal"><form class="modal-card wide" id="creator-tags-form" data-creator-tags=""><h2>编辑创作者标签</h2><p class="copy" style="margin-bottom:14px">身份和地点支持多选。保存后，详情页点击“查看推荐脚本”会按新标签重新计算。</p><div class="creator-modal-form"><input type="hidden" name="kwai_url"><input type="hidden" name="display_name"><input type="hidden" name="kwai_id"><input type="hidden" name="phone"><input type="hidden" name="uid"><select name="poc"></select><select name="category"></select><select class="multi-select" name="identity" multiple></select><select class="multi-select" name="location" multiple></select><select name="cooperation_level"></select><textarea name="creator_description" placeholder="具体作者描述：可手动输入作者风格、限制、偏好等"></textarea></div><div class="modal-actions"><button type="button" id="creator-tags-cancel">取消</button><button class="primary" type="submit">保存标签</button></div></form></div><script>
const libraryMode=__LIBRARY_MODE__;const labels=["夫妻整蛊/冲突","夫妻暧昧","家庭整蛊","朋友整蛊"];const durationOptions=[["dur_1_20","1-20 秒"],["dur_20_60","20 秒-1 分钟"],["dur_60_120","1-2 分钟"],["dur_120_plus","2 分钟以上"]];const locationOptions=["室内房间","乡村院子","工地","酒馆","超市","药店","房屋内外结合"];const creatorPocOptions=["denghaoqing","zhaozhe"];let entries=[];let creators=[];let creatorRows=[];let accounts=[];let submissions=[];let accessApplications=[];let submissionsTotal=0;let submissionsOffset=0;let submissionsLoading=false;let intakes=[];let analyticsData=null;let analyticsAutoLoaded=false;let analyticsInactiveLoaded=false;let analyticsLoading=false;let analyticsUserFilter="activated";let importJob=null;let importPollTimer=null;let activeTab=__INITIAL_TAB__;let activeScriptType="";let activeScriptDuration="";let activeScriptLocation="";let activeScriptScope="portal_visible";let activeCreatorPoc="";let creatorViewMode=localStorage.getItem("kokoCreatorAdminView")||"card";let creatorCloudState={{creators:{{}}}};let scriptScopeCounts={{portal_visible:0,hidden:0,incomplete:0,all:0}};let scriptVisibleLimit=20;let scriptIndex={{}};let scriptIndexLoaded=false;const SUBMISSIONS_PAGE_SIZE=80;const SCRIPT_INITIAL_RENDER_LIMIT=20;const SCRIPT_RENDER_INCREMENT=50;const CREATOR_PAGE_SIZE=18;const creatorPathMatch=location.pathname.match(/^\\/creator-admin\\/creators\\/([0-9a-f]{{32}})$/);let creatorPage=1;let selectedCreatorId=(creatorPathMatch&&creatorPathMatch[1])||new URLSearchParams(location.search).get("creator")||"";let editing=null;const app=document.querySelector("#app");const modal=document.querySelector("#edit-modal");const form=document.querySelector("#edit-form");const creatorModal=document.querySelector("#creator-modal");const creatorForm=document.querySelector("#creator-form");const creatorTagsModal=document.querySelector("#creator-tags-modal");const creatorTagsForm=document.querySelector("#creator-tags-form");
function esc(s){{return String(s??"").replace(/[&<>"']/g,c=>({{"&":"&amp;","<":"&lt;",">":"&gt;","\\\"":"&quot;","'":"&#39;"}}[c]))}}
function noCacheUrl(url){{const methodUrl=String(url||"");if(!methodUrl.startsWith("/"))return methodUrl;const sep=methodUrl.includes("?")?"&":"?";return `${{methodUrl}}${{sep}}_=${{Date.now()}}`}}
async function api(url,opts={{}}){{const method=String(opts.method||"GET").toUpperCase();const requestUrl=method==="GET"?noCacheUrl(url):url;const r=await fetch(requestUrl,{{credentials:"same-origin",cache:"no-store",headers:{{"Content-Type":"application/json","Cache-Control":"no-store"}},...opts}});const d=await r.json().catch(()=>({{}}));if(!r.ok||d.ok===false)throw new Error(d.error||"请求失败");return d}}
function loginView(msg=""){{app.innerHTML=`<section class="login"><form id="login-form"><span class="kicker">Koko 内部后台</span><h1>${{libraryMode?"脚本管理":"Creator 运营后台"}}</h1><p class="copy">${{libraryMode?"这里统一管理 Koko 入库脚本和 kokocomedy 前台脚本。":"这里管理创作者前台展示的脚本、标签、上下架和同步。"}}</p><input name="password" type="password" placeholder="后台密码" autofocus style="margin-top:16px"><button class="primary" style="width:100%;margin-top:12px" type="submit">进入后台</button><p class="status">${{esc(msg)}}</p></form></section>`}}
function adminKicker(){{return libraryMode?"Koko Script Management":"Koko Creator Operations"}}
function adminTitle(){{return libraryMode?"脚本管理":"Creator 运营后台"}}
function adminCopy(){{return libraryMode?"这里是 Koko 内容中台的统一脚本管理入口，以 kokocomedy 当前脚本库为准；编辑、上下架、删除和分享外链都会直接作用到创作者前台。":""}}
function adminNav(){{return libraryMode?`<a class="btn" href="/studio">返回内容中台</a><a class="btn" href="/creator-admin/imports">导入脚本</a><a class="btn" href="__CREATOR_BASE__/creator-portal" target="_blank" rel="noopener">打开 Creator 前台</a>`:`<a class="btn" href="/studio">返回内容中台</a><a class="btn" href="/library">脚本管理</a><a class="btn" href="__CREATOR_BASE__/creator-portal" target="_blank" rel="noopener">打开 Creator 前台</a>`}}
function adminTabs(){{return libraryMode?`<div class="ops-tabs"><a class="active" href="/library">脚本管理</a></div>`:`<div class="ops-tabs"><a class="${{activeTab==="creators"?"active":""}}" href="/creator-admin/creators">创作者管理</a><a class="${{activeTab==="analytics"?"active":""}}" href="/creator-admin/analytics">数据看板</a><a class="${{activeTab==="submissions"?"active":""}}" href="/creator-admin/submissions">回传数据</a><a class="${{activeTab==="intakes"?"active":""}}" href="/creator-admin/intakes">作者信息收集</a></div>`}}
function adminView(){{app.innerHTML=`<section class="shell"><div class="panel"><div class="top"><div><span class="kicker">${{adminKicker()}}</span><h1>${{adminTitle()}}</h1><p class="copy">${{adminCopy()}}</p></div><div class="nav">${{libraryMode?"":'<span id="creator-realtime-status" class="realtime-status">实时同步开启</span>'}}${{adminNav()}}</div></div>${{adminTabs()}}<div id="tab-body"></div></div></section>`;renderActiveTab()}}
function renderActiveTab(){{const body=document.querySelector("#tab-body");if(!body)return;if(activeTab==="imports"){{body.innerHTML=`<section class="import-panel"><div><h2 style="margin:0 0 8px;color:#ff8200">导入标准脚本 Excel</h2><p class="copy">上传包含 Vídeo original / Conteúdo principal / Pontos principais / Partes que podem ser adaptadas / Tempo / Imagem / Ações / Diálogos 的 .xlsx。系统会自动拆分多条脚本，生成葡语脚本页，调用 Gemini 慢慢生成 3x3 分镜图，并同步到 Creator 前台。分类选择如果保留在“待分类”，后台会自动用大模型判断脚本类型；如果你手动选择了具体类型，则优先按手动类型导入。</p></div><form class="import-form" id="import-form"><input id="import-file" type="file" accept=".xlsx"><select id="import-content-type">${{labels.map(x=>`<option value="${{esc(x)}}">${{esc(x)}}</option>`).join("")}}</select><button class="primary" type="submit">上传并导入</button></form><p id="status" class="status"></p><div id="import-job"></div></section>`;document.querySelector("#import-form").addEventListener("submit",submitImport);renderImportJob();return}}if(activeTab==="analytics"){{body.innerHTML=`<section class="import-panel"><div class="creator-toolbar"><div><h2 style="margin:0 0 8px;color:#ff8200">kokocomedy 数据看板</h2><p class="copy">按手机号账号聚合注册激活、打开次数、停留时长、功能点击和脚本回传。</p></div><div class="creator-toolbar-right"><button class="primary" id="refresh-analytics" type="button">刷新有行为用户</button><button id="logout" type="button">退出</button></div></div><p id="status" class="status"></p><div id="analytics-board"></div></section>`;document.querySelector("#refresh-analytics").addEventListener("click",()=>loadAnalytics(false));document.querySelector("#logout").addEventListener("click",logout);renderAnalytics();return}}if(activeTab==="creators"){{body.innerHTML=`<section class="import-panel"><div><h2 style="margin:0 0 8px;color:#ff8200">创作者管理</h2></div><div class="creator-searchbar"><select id="creator-search-scope"><option value="all">所有结果</option><option value="kwai">Kwai ID</option><option value="name">用户名</option><option value="phone">手机号</option><option value="uid">UID</option></select><input id="creator-search" placeholder="输入作者名称、ID、手机号或 UID"><button class="primary" id="creator-search-button" type="button">Search</button></div><div id="creator-poc-filters" class="quick-filters"></div><div class="creator-toolbar"><div class="creator-toolbar-left"><span id="creator-results-meta" class="creator-results-meta"></span><div class="creator-view-toggle" aria-label="切换创作者展示方式"><button type="button" data-creator-view="card">卡片</button><button type="button" data-creator-view="list">横向</button></div></div><div class="creator-toolbar-right"><button class="primary" id="open-creator-modal" type="button">导入创作者</button><button id="refresh-creators" type="button">刷新</button><button id="logout" type="button">退出</button></div></div><p id="status" class="status"></p><div id="creator-list"></div></section>`;document.querySelector("#creator-search").addEventListener("input",()=>{{creatorPage=1;selectedCreatorId="";renderCreators()}});document.querySelector("#creator-search-scope").addEventListener("change",()=>{{creatorPage=1;selectedCreatorId="";renderCreators()}});document.querySelector("#creator-search-button").addEventListener("click",()=>{{creatorPage=1;selectedCreatorId="";renderCreators()}});document.querySelector("#open-creator-modal").addEventListener("click",openCreatorModal);document.querySelector("#refresh-creators").addEventListener("click",loadCreators);document.querySelector("#logout").addEventListener("click",logout);renderCreatorPocFilters();renderCreators();return}}if(activeTab==="accounts"){{body.innerHTML=`<form class="creator-form" id="account-form"><input name="account" placeholder="输入手机号、数字或字母账号，例如 88998411165 / creator01"><input name="display_name" placeholder="显示名称（可选）"><button class="primary" type="submit">创建账号</button></form><div class="toolbar"><input id="account-search" placeholder="搜索账号、显示名、回传链接"><button id="refresh-accounts" type="button">刷新</button><button id="logout" type="button">退出</button></div><p id="status" class="status"></p><div id="account-list" class="grid"></div>`;document.querySelector("#account-form").addEventListener("submit",createAccount);document.querySelector("#account-search").addEventListener("input",renderAccounts);document.querySelector("#refresh-accounts").addEventListener("click",loadAccounts);document.querySelector("#logout").addEventListener("click",logout);renderAccounts();return}}if(activeTab==="submissions"){{body.innerHTML=`<div class="toolbar"><input id="submission-search" placeholder="搜索脚本标题、创作者、回传链接"><button id="refresh-submissions" type="button">刷新</button><button id="logout" type="button">退出</button></div><p id="status" class="status"></p><div id="submission-stats"></div>`;document.querySelector("#submission-search").addEventListener("input",renderSubmissionStats);document.querySelector("#refresh-submissions").addEventListener("click",()=>loadSubmissions(false));document.querySelector("#logout").addEventListener("click",logout);renderSubmissionStats();return}}if(activeTab==="intakes"){{body.innerHTML=`<div class="toolbar"><input id="intake-search" placeholder="搜索 Kwai 名称、答案、联系方式"><button id="refresh-intakes" type="button">刷新</button><button id="logout" type="button">退出</button></div><p id="status" class="status"></p><div id="intake-list" class="grid"></div>`;document.querySelector("#intake-search").addEventListener("input",renderIntakes);document.querySelector("#refresh-intakes").addEventListener("click",loadIntakes);document.querySelector("#logout").addEventListener("click",logout);renderIntakes();return}}body.innerHTML=`<div class="toolbar"><input id="search" placeholder="搜索标题、摘要、分类、时间、地点、视频链接"><button id="delete-selected" class="danger" type="button">批量删除</button><button id="refresh" type="button">刷新</button><button id="logout" type="button">退出</button></div><div id="script-scope-filters" class="quick-filters"></div><div id="script-type-filters" class="quick-filters"></div><div id="script-duration-filters" class="quick-filters"></div><div id="script-location-filters" class="quick-filters"></div><p id="status" class="status"></p><div id="list" class="grid"></div>`;document.querySelector("#search").addEventListener("input",()=>{{scriptVisibleLimit=SCRIPT_INITIAL_RENDER_LIMIT;renderList()}});document.querySelector("#refresh").addEventListener("click",loadEntries);document.querySelector("#delete-selected").addEventListener("click",bulkDelete);document.querySelector("#logout").addEventListener("click",logout);renderScriptScopeFilters();renderScriptTypeFilters();renderScriptDurationFilters();renderScriptLocationFilters();renderList()}}
function scriptScopeOptions(){{return [{{key:"portal_visible",label:"前台展示中"}},{{key:"hidden",label:"已下架"}},{{key:"incomplete",label:"信息不完整"}},{{key:"all",label:"全部"}}]}}
function renderScriptScopeFilters(){{const box=document.querySelector("#script-scope-filters");if(!box)return;box.innerHTML=scriptScopeOptions().map(option=>`<button class="quick-filter-chip ${{activeScriptScope===option.key?"active":""}}" type="button" data-scope-filter="${{option.key}}"><span>${{option.label}}</span> <small>${{Number(scriptScopeCounts?.[option.key]||0)}}</small></button>`).join("")}}
function scriptTypeCounts(){{const counts=new Map();for(const entry of entries){{const key=String(entry.content_type||"待分类").trim()||"待分类";counts.set(key,(counts.get(key)||0)+1)}}return counts}}
function scriptTypeOptions(){{const counts=scriptTypeCounts();const ordered=[];for(const label of labels){{if(counts.has(label))ordered.push([label,counts.get(label)])}}for(const [label,count] of [...counts.entries()].sort((a,b)=>String(a[0]).localeCompare(String(b[0]),"zh-CN"))){{if(!labels.includes(label))ordered.push([label,count])}}return ordered}}
function renderScriptTypeFilters(){{const box=document.querySelector("#script-type-filters");if(!box)return;const options=scriptTypeOptions();box.innerHTML=[`<button class="quick-filter-chip ${{!activeScriptType?"active":""}}" type="button" data-type-filter=""><span>全部</span> <small>${{entries.length}}</small></button>`,...options.map(([label,count])=>`<button class="quick-filter-chip ${{activeScriptType===label?"active":""}}" type="button" data-type-filter="${{esc(label)}}"><span>${{esc(label)}}</span> <small>${{count}}</small></button>`)].join("")}}
function scriptDurationCounts(){{const counts=new Map();for(const entry of entries){{const key=String(entry.duration_bucket||"").trim();if(key)counts.set(key,(counts.get(key)||0)+1)}}return counts}}
function renderScriptDurationFilters(){{const box=document.querySelector("#script-duration-filters");if(!box)return;const counts=scriptDurationCounts();box.innerHTML=[`<button class="quick-filter-chip ${{!activeScriptDuration?"active":""}}" type="button" data-duration-filter=""><span>全部时间</span> <small>${{entries.length}}</small></button>`,...durationOptions.map(([bucket,label])=>`<button class="quick-filter-chip ${{activeScriptDuration===bucket?"active":""}}" type="button" data-duration-filter="${{esc(bucket)}}"><span>${{esc(label)}}</span> <small>${{Number(counts.get(bucket)||0)}}</small></button>`)].join("")}}
function scriptLocationCounts(){{const counts=new Map();for(const entry of entries){{const key=String(entry.location_tag||entry.location_tag_pt||"").trim();if(key)counts.set(key,(counts.get(key)||0)+1)}}return counts}}
function renderScriptLocationFilters(){{const box=document.querySelector("#script-location-filters");if(!box)return;const counts=scriptLocationCounts();box.innerHTML=[`<button class="quick-filter-chip ${{!activeScriptLocation?"active":""}}" type="button" data-location-filter=""><span>全部地点</span> <small>${{entries.length}}</small></button>`,...locationOptions.map(label=>`<button class="quick-filter-chip ${{activeScriptLocation===label?"active":""}}" type="button" data-location-filter="${{esc(label)}}"><span>${{esc(label)}}</span> <small>${{Number(counts.get(label)||0)}}</small></button>`)].join("")}}
function locationTag(e){{return String(e.location_tag||e.location_tag_pt||"").trim()}}
function filteredEntries(){{const q=String(document.querySelector("#search")?.value||"").trim().toLowerCase();return entries.filter(e=>{{const matchesType=!activeScriptType||String(e.content_type||"待分类").trim()===activeScriptType;if(!matchesType)return false;const matchesDuration=!activeScriptDuration||String(e.duration_bucket||"").trim()===activeScriptDuration;if(!matchesDuration)return false;const matchesLocation=!activeScriptLocation||locationTag(e)===activeScriptLocation;if(!matchesLocation)return false;if(!q)return true;return [e.title,e.summary,e.content_type,e.duration_bucket,e.duration_label_pt,e.duration_label_zh,e.location_tag,e.location_tag_pt,e.video_url].join(" ").toLowerCase().includes(q)}})}}
function durationTag(e){{const map={{dur_1_20:"1-20 s",dur_20_60:"20 s-1 min",dur_60_120:"1-2 min",dur_120_plus:"Mais de 2 min"}};return e.duration_label_pt||map[e.duration_bucket]||""}}
function renderList(){{const list=document.querySelector("#list");if(!list)return;const rows=filteredEntries();if(!rows.length){{list.innerHTML=`<div class="empty">没有匹配脚本</div>`;return}}const visibleRows=rows.slice(0,Math.max(SCRIPT_INITIAL_RENDER_LIMIT,scriptVisibleLimit));const remaining=Math.max(0,rows.length-visibleRows.length);const cards=visibleRows.map(e=>{{const share=scriptShareUrl(e.entry_id);const duration=durationTag(e);const location=locationTag(e);return `<article class="card"><input type="checkbox" data-pick="${{esc(e.entry_id)}}"><img src="${{esc(e.cover_url||e.thumbnail_url)}}" loading="lazy" alt=""><div><h3>${{esc(e.title||"Untitled")}}</h3><p>${{esc(e.summary||"")}}</p><div class="meta"><span class="pill">${{esc(e.content_type||"待分类")}}</span>${{duration?`<span class="pill">${{esc(duration)}}</span>`:""}}${{location?`<span class="pill">${{esc(location)}}</span>`:""}}<span class="pill ${{e.published?"":"off"}}">${{e.published?"Creator 已上架":"Creator 已下架"}}</span></div><div class="share-line"><b>kokocomedy 外链</b><a href="${{esc(share)}}" target="_blank" rel="noopener">${{esc(share)}}</a></div></div><div class="actions"><a class="btn" href="${{esc(share)}}" target="_blank" rel="noopener">打开外链</a><button type="button" data-copy="${{esc(share)}}">复制外链</button><button class="primary" type="button" data-script-edit="${{esc(e.entry_id)}}">修改脚本</button><button type="button" data-edit="${{esc(e.entry_id)}}">修改标签</button><button type="button" data-toggle="${{esc(e.entry_id)}}">${{e.published?"下架":"上架"}}</button></div></article>`}}).join("");const more=remaining?`<div class="empty"><b>已显示 ${{visibleRows.length}} / ${{rows.length}} 条</b><br><br><button class="primary" type="button" data-load-more-scripts>继续展开 50 条</button></div>`:`<div class="empty">已显示全部 ${{rows.length}} 条脚本。</div>`;list.innerHTML=cards+more;const s=document.querySelector("#status");if(s){{const scopeLabel=(scriptScopeOptions().find(x=>x.key===activeScriptScope)||{{label:"脚本"}}).label;const filters=[activeScriptType&&`类型：${{activeScriptType}}`,activeScriptDuration&&`时间：${{durationOptions.find(x=>x[0]===activeScriptDuration)?.[1]||activeScriptDuration}}`,activeScriptLocation&&`地点：${{activeScriptLocation}}`].filter(Boolean).join(" · ");s.textContent=`${{scopeLabel}}：已加载 ${{entries.length}} 条，当前显示 ${{visibleRows.length}} / ${{rows.length}} 条${{filters?` · ${{filters}}`:""}}`}}}}
async function loadEntries(){{try{{document.querySelector("#status")&&(document.querySelector("#status").textContent="加载中...");const params=new URLSearchParams({{limit:"10000",scope:activeScriptScope}});const d=await api(`/api/creator-admin/scripts?${{params.toString()}}`);entries=d.entries||[];scriptVisibleLimit=SCRIPT_INITIAL_RENDER_LIMIT;scriptScopeCounts=d.scope_counts||{{portal_visible:0,hidden:0,incomplete:0,all:entries.length}};activeScriptScope=d.scope||activeScriptScope;if(activeScriptType&&!entries.some(e=>String(e.content_type||"待分类").trim()===activeScriptType))activeScriptType="";if(activeScriptDuration&&!entries.some(e=>String(e.duration_bucket||"").trim()===activeScriptDuration))activeScriptDuration="";if(activeScriptLocation&&!entries.some(e=>locationTag(e)===activeScriptLocation))activeScriptLocation="";adminView();const scopeLabel=(scriptScopeOptions().find(x=>x.key===activeScriptScope)||{{label:"脚本"}}).label;const total=Number(scriptScopeCounts[activeScriptScope]||entries.length);const s=document.querySelector("#status");if(s)s.textContent=`${{scopeLabel}}：已加载 ${{entries.length}} / ${{total}} 条，默认显示前 ${{Math.min(SCRIPT_INITIAL_RENDER_LIMIT,entries.length)}} 条`}}catch(e){{loginView(e.message)}}}}
async function ensureScriptIndex(){{if(scriptIndexLoaded)return scriptIndex;const d=await api("/api/creator-admin/scripts?limit=10000&scope=all");const rows=d.entries||[];scriptIndex=Object.fromEntries(rows.map(x=>[String(x.entry_id||""),x]).filter(x=>x[0]));scriptIndexLoaded=true;return scriptIndex}}
async function loadCreatorCloudState(){{const d=await api("/api/creator-admin/state");creatorCloudState=d.state&&d.state.creators?d.state:{{creators:{{}}}}}}
async function saveCreatorCloudState(id,patch){{const d=await api(`/api/creator-admin/state/${{id}}`,{{method:"POST",body:JSON.stringify(patch||{{}})}});if(!creatorCloudState.creators)creatorCloudState.creators={{}};creatorCloudState.creators[id]=d.creator_state||{{}};return creatorCloudState.creators[id]}}
function secondsText(value){{const sec=Number(value||0);if(sec<60)return `${{Math.round(sec)}} 秒`;const min=Math.floor(sec/60);const rest=Math.round(sec%60);if(min<60)return `${{min}} 分 ${{rest}} 秒`;return `${{Math.floor(min/60)}} 小时 ${{min%60}} 分`}}
function analyticsSummaryCard(label,value,sub=""){{return `<div class="metric-card"><span>${{esc(label)}}</span><b>${{esc(value)}}</b>${{sub?`<small class="small">${{esc(sub)}}</small>`:""}}</div>`}}
function analyticsIsRegistered(user){{return user.registration_status==="registered"||Number(user.submission_count||0)>0}}
function analyticsStatusPill(user){{const registered=analyticsIsRegistered(user);return `<span class="pill ${{registered?"":"off"}}">${{registered?"已注册":"未注册"}}</span>`}}
function analyticsClickPills(user){{const counts=user.click_counts||{{}};const labels=user.click_labels||{{}};const keys=Object.keys(counts);if(!keys.length)return `<span class="pill off">暂无点击</span>`;return keys.map(k=>`<span class="pill">${{esc(labels[k]||k)}} ${{Number(counts[k]||0)}}</span>`).join("")}}
function analyticsScriptRows(user){{const rows=user.script_views||[];if(!rows.length)return `<div class="small">暂无脚本浏览记录</div>`;return rows.map(item=>`<div class="share-line"><b>${{esc(item.title||item.script_id||"脚本")}}</b><span class="small">打开 ${{Number(item.views||0)}} 次 · 停留 ${{secondsText(Number(item.duration_ms||0)/1000)}}</span><a href="__CREATOR_BASE__/script/${{esc(item.script_id)}}" target="_blank" rel="noopener">__CREATOR_BASE__/script/${{esc(item.script_id)}}</a></div>`).join("")}}
function analyticsSubmissionRows(user){{const rows=user.submissions||[];if(!rows.length)return `<div class="small">暂无回传</div>`;return rows.map(item=>`<div class="share-line"><b>${{esc(item.script_title||item.entry_id||"脚本回传")}}</b><a href="__CREATOR_BASE__/script/${{esc(item.entry_id)}}" target="_blank" rel="noopener">脚本链接：__CREATOR_BASE__/script/${{esc(item.entry_id)}}</a><a href="${{esc(item.video_url||"")}}" target="_blank" rel="noopener">回传链接：${{esc(item.video_url||"")}}</a><span class="small">回传时间：${{esc(timeText(item.created_at))}}</span></div>`).join("")}}
function analyticsRecentEventRows(user){{const labels=user.click_labels||{{}};const rows=user.recent_events||[];if(!rows.length)return `<div class="small">暂无最近行为</div>`;return rows.slice(0,20).map(item=>{{const label=labels[item.event]||item.event||"行为";const script=item.script_id?` · 脚本 ${{item.script_id}}`:"";const duration=Number(item.duration_ms||0)?` · 停留 ${{secondsText(Number(item.duration_ms||0)/1000)}}`:"";return `<div class="share-line"><b>${{esc(label)}}</b><span class="small">${{esc(timeText(item.created_at))}}${{esc(script)}}${{esc(duration)}}</span><span class="small">${{esc(item.path||"")}}</span></div>`}}).join("")}}
function analyticsAuthorDetail(user,submissionOverride=null,showSubmissions=true){{const totalDuration=Number(user.platform_duration_seconds||0)+Number(user.script_duration_seconds||0);const scriptOpens=Number(user.script_share_open_count||0);const avgScript=scriptOpens?Number(user.script_duration_seconds||0)/scriptOpens:0;const submissionCount=submissionOverride===null?Number(user.submission_count||0):Number(submissionOverride||0);const submissionCard=showSubmissions?`<div class="metric-card"><span>回传脚本</span><b>${{submissionCount}}</b><small class="small">账号真实回传记录</small></div>`:"";const submissionSection=showSubmissions?`<section class="import-panel"><h3 style="margin:0;color:#1f1f1f">回传记录</h3>${{analyticsSubmissionRows(user)}}</section>`:"";return `<details class="analytics-user-detail" style="grid-column:1/-1;margin-top:8px;padding:12px;border:1px solid rgba(255,130,0,.14);border-radius:16px;background:#fffaf5"><summary>${{showSubmissions?"展开这个作者的停留、点击、脚本浏览和回传":"展开这个作者的停留、点击和脚本浏览"}}</summary><div class="analytics-summary" style="grid-template-columns:repeat(${{showSubmissions?4:3}},minmax(0,1fr));margin-top:12px"><div class="metric-card"><span>Koko 总停留</span><b>${{secondsText(totalDuration)}}</b><small class="small">首页 ${{secondsText(user.platform_duration_seconds)}} · 脚本 ${{secondsText(user.script_duration_seconds)}}</small></div><div class="metric-card"><span>打开次数</span><b>${{Number(user.platform_open_count||0)+scriptOpens}}</b><small class="small">首页 ${{Number(user.platform_open_count||0)}} · 脚本 ${{scriptOpens}}</small></div><div class="metric-card"><span>平均脚本停留</span><b>${{secondsText(avgScript)}}</b><small class="small">按脚本打开次数计算</small></div>${{submissionCard}}</div><section class="import-panel"><h3 style="margin:0;color:#1f1f1f">点击组件</h3><div class="meta">${{analyticsClickPills(user)}}</div></section><section class="import-panel"><h3 style="margin:0;color:#1f1f1f">看过的脚本</h3>${{analyticsScriptRows(user)}}</section>${{submissionSection}}<section class="import-panel"><h3 style="margin:0;color:#1f1f1f">最近行为流水</h3>${{analyticsRecentEventRows(user)}}</section></details>`}}
function analyticsPct(part,total){{const p=Number(part||0),t=Number(total||0);return t?Math.round(p/t*1000)/10:0}}
function analyticsAvgSeconds(total,count){{return count?Math.round(Number(total||0)/Math.max(1,count)):0}}
function analyticsHourLabel(value){{if(!value)return "-";const d=new Date(value);if(Number.isNaN(d.getTime()))return String(value).slice(0,16);return d.toLocaleString("zh-CN",{{month:"2-digit",day:"2-digit",hour:"2-digit",minute:"2-digit",hour12:false}})}}
function analyticsDayKey(value){{if(!value)return "未记录日期";const d=new Date(value);if(Number.isNaN(d.getTime()))return String(value).slice(0,10)||"未记录日期";try{{return d.toLocaleDateString("sv-SE",{{timeZone:"Asia/Shanghai"}})}}catch(e){{return d.toISOString().slice(0,10)}}}}
function analyticsEntityKey(item){{return String(item.account_id||item.phone||item.kwai_id||item.visitor_id||item.display_name||"").trim()}}
function analyticsIsTestUser(item){{return ["666"].includes(analyticsEntityKey(item))}}
const ANALYTICS_MAX_VALID_STAY_SECONDS=30*60
function analyticsValidStaySeconds(value){{const seconds=Number(value||0);return seconds>0&&seconds<=ANALYTICS_MAX_VALID_STAY_SECONDS?seconds:0}}
function analyticsEffectiveDurationSeconds(user){{return analyticsValidStaySeconds(Number(user.platform_duration_seconds||0)+Number(user.script_duration_seconds||0))}}
function analyticsPersonLine(item,kind=""){{const name=item.display_name||item.account_id||item.visitor_id||"匿名访客";const account=[item.phone&&`手机号 ${{item.phone}}`,item.kwai_id&&`Kwai ${{item.kwai_id}}`,item.account_id&&`账号 ${{item.account_id}}`].filter(Boolean).join(" · ");const script=item.script_title||item.script_id||"";const source=[item.source,item.path&&`路径 ${{item.path}}`,item.referer&&`来源 ${{item.referer}}`].filter(Boolean).join(" · ");const duration=Number(item.duration_seconds||0);const abnormal=duration&&!analyticsValidStaySeconds(duration);return `<div class="analytics-person"><b>${{esc(name)}}</b>${{account?`<small>${{esc(account)}}</small>`:""}}${{script?`<small>脚本：${{esc(script)}}</small>`:""}}${{duration?`<small>停留：${{secondsText(duration)}}${{abnormal?" · 异常停留不计入平均":""}}</small>`:""}}${{source?`<small>${{esc(source)}}</small>`:""}}${{item.time?`<small>时间：${{esc(timeText(item.time))}}</small>`:""}}</div>`}};
function analyticsDetailGroup(title,rows,empty="暂无明细"){{rows=Array.isArray(rows)?rows:[];return `<section class="analytics-detail-group"><h4>${{esc(title)}}（${{rows.length}}）</h4>${{rows.length?rows.map(x=>analyticsPersonLine(x)).join(""):`<div class="small">${{esc(empty)}}</div>`}}</section>`}}
function analyticsDailySummaryHtml(hourRows,users){{const map=new Map();const ensure=day=>{{if(!day)return null;if(!map.has(day))map.set(day,{{users:new Set(),durationSum:0,durationCount:0,submissions:0}});return map.get(day)}};const addUser=(day,item)=>{{const bucket=ensure(day);if(!bucket||analyticsIsTestUser(item))return;const key=analyticsEntityKey(item);if(key)bucket.users.add(key)}};const addDuration=(day,item,value)=>{{const bucket=ensure(day);if(!bucket||analyticsIsTestUser(item))return;const duration=analyticsValidStaySeconds(value);if(duration){{bucket.durationSum+=duration;bucket.durationCount+=1}}}};for(const row of (hourRows||[])){{const day=analyticsDayKey(row.hour);for(const item of [...(row.platform_open_details||[]),...(row.script_open_details||[]),...(row.duration_details||[])])addUser(day,item);for(const item of (row.duration_details||[]))addDuration(day,item,item.duration_seconds)}}for(const user of (users||[])){{if(analyticsIsTestUser(user))continue;const userDay=analyticsDayKey(user.last_event_at||user.last_login_at||analyticsActivationTime(user)||user.created_at);if(Number(user.platform_open_count||0)||Number(user.script_share_open_count||0)||(user.script_views||[]).length||(user.recent_events||[]).length){{addUser(userDay,user);addDuration(userDay,user,Number(user.platform_duration_seconds||0)+Number(user.script_duration_seconds||0))}}for(const event of (user.recent_events||[])){{const day=analyticsDayKey(event.created_at||event.time);addUser(day,user);if(Number(event.duration_ms||0))addDuration(day,user,Number(event.duration_ms||0)/1000)}}for(const view of (user.script_views||[])){{const day=analyticsDayKey(view.last_viewed_at||view.created_at||user.last_event_at||userDay);addUser(day,user);if(Number(view.duration_ms||0))addDuration(day,user,Number(view.duration_ms||0)/1000)}}for(const sub of (user.submissions||[])){{const day=analyticsDayKey(sub.created_at||userDay);const bucket=ensure(day);if(bucket){{bucket.submissions+=1;addUser(day,user)}}}}}}const days=[...map.entries()].filter(([day])=>day&&day!=="未记录日期").sort((a,b)=>String(b[0]).localeCompare(String(a[0])));if(!days.length)return `<section class="import-panel"><h3 style="margin:0;color:#1f1f1f">每日数据总结</h3><div class="empty">暂无每日数据。</div></section>`;return `<section class="import-panel"><h3 style="margin:0;color:#1f1f1f">每日数据总结</h3><div class="analytics-hour-table"><div class="analytics-hour-head" style="grid-template-columns:150px 1fr 1fr 1fr"><span>日期</span><span>Koko 使用人数</span><span>平均停留时间</span><span>脚本回传数据</span></div>${{days.map(([day,bucket])=>`<div class="analytics-hour-row" style="grid-template-columns:150px 1fr 1fr 1fr"><span data-label="日期">${{esc(day.replaceAll("-","."))}}</span><b data-label="Koko 使用人数">${{bucket.users.size}}</b><span data-label="平均停留时间">${{bucket.durationCount?secondsText(bucket.durationSum/bucket.durationCount):"-"}}</span><b data-label="脚本回传数据">${{bucket.submissions}}</b></div>`).join("")}}</div><small class="small">每日数据会从小时明细、用户行为、脚本浏览和回传记录共同汇总；平均停留已排除测试账号 666、0 秒和超过 30 分钟的异常停留。</small></section>`}}
function analyticsCoreTimelineHtml(){{const rows=(analyticsData?.timeline?.hourly||[]).slice().sort((a,b)=>String(b.hour||"").localeCompare(String(a.hour||"")));if(!rows.length)return `<details class="analytics-core-detail"><summary>展开明细 <small>当前接口还没有返回小时数据</small></summary><div class="empty" style="margin-top:12px">暂无可展示的明细。等待 kokocomedy 数据接口更新后会自动显示。</div></details>`;return `<details class="analytics-core-detail"><summary>展开明细 <small>${{rows.length}} 个小时段，可查看每日总结和具体账号</small></summary>${{analyticsDailySummaryHtml(rows,analyticsData?.users||[])}}<div class="analytics-hour-table"><div class="analytics-hour-head"><span>时间</span><span>注册</span><span>直接打开</span><span>分享链接打开</span><span>脚本打开</span><span>平均时长</span></div>${{rows.map(row=>{{const durations=(row.duration_details||[]).filter(item=>!analyticsIsTestUser(item)).map(item=>analyticsValidStaySeconds(item.duration_seconds)).filter(Boolean);const avgDuration=durations.length?durations.reduce((a,b)=>a+b,0)/durations.length:0;return `<div class="analytics-hour-row"><span data-label="时间">${{esc(analyticsHourLabel(row.hour))}}</span><b data-label="注册">${{Number(row.registered_users||0)}}</b><b data-label="直接打开">${{Number(row.platform_opens||0)}} / ${{Number(row.platform_people_count||0)}}人</b><b data-label="分享链接打开">${{Number(row.share_link_opens||0)}} / ${{Number(row.script_people_count||0)}}人</b><b data-label="脚本打开">${{Number(row.script_opens||0)}}</b><span data-label="平均时长">${{avgDuration?secondsText(avgDuration):"-"}}</span><details><summary>查看这个小时的具体账号、来源和停留</summary><div class="analytics-hour-detail">${{analyticsDetailGroup("注册用户",row.registered_details)}}${{analyticsDetailGroup("直接打开 Koko",row.platform_open_details)}}${{analyticsDetailGroup("打开脚本",row.script_open_details)}}${{analyticsDetailGroup("停留行为",row.duration_details)}}</div></details></div>`}}).join("")}}</div></details>`}}
function analyticsFunnelSteps(s,users){{const preset=Number(s.accounts||0);const active=Number(s.active_accounts||users.length);const registered=users.filter(u=>analyticsIsRegistered(u)||u.last_login_at||Number((u.click_counts||{{}}).login||0)||Number((u.click_counts||{{}}).register||0)).length;const viewed=users.filter(u=>Number(u.script_share_open_count||0)>0||(u.script_views||[]).length).length;const submitClick=users.filter(u=>Number((u.click_counts||{{}}).submit_click||0)>0||Number((u.click_counts||{{}}).open_submit||0)>0).length;const submitted=users.filter(u=>Number(u.submission_count||0)>0).length;return [["预置账号",preset,preset],["首次访问",active,preset],["登录/注册",registered,active],["查看脚本",viewed,registered||active],["点击回传",submitClick,viewed],["完成回传",submitted,submitClick||viewed]]}}
function analyticsFunnelHtml(s,users){{const steps=analyticsFunnelSteps(s,users);return `<div class="analytics-funnel">${{steps.map(([label,value,base],idx)=>`<div class="analytics-step"><span>${{esc(label)}}</span><b>${{Number(value||0)}}</b><div class="analytics-bar"><i style="width:${{Math.max(4,Math.min(100,analyticsPct(value,base)||0))}}%"></i></div><small>${{idx===0?"100%":analyticsPct(value,base)+"%"}}</small></div>`).join("")}}</div>`}}
function analyticsRetentionHtml(users){{const rows=[["已注册",users.filter(u=>analyticsIsRegistered(u))],["未注册有行为",users.filter(u=>!analyticsIsRegistered(u))],["有回传",users.filter(u=>Number(u.submission_count||0)>0)],["仅浏览",users.filter(u=>Number(u.submission_count||0)===0&&(Number(u.script_share_open_count||0)>0||(u.script_views||[]).length))]];return `<table class="analytics-heat"><thead><tr><th>用户分层</th><th>人数</th><th>脚本打开</th><th>平均停留</th><th>回传</th></tr></thead><tbody>${{rows.map(([label,list])=>{{const opens=list.reduce((a,u)=>a+Number(u.script_share_open_count||0),0);const duration=list.reduce((a,u)=>a+Number(u.script_duration_seconds||0)+Number(u.platform_duration_seconds||0),0);const subs=list.reduce((a,u)=>a+Number(u.submission_count||0),0);return `<tr><td>${{esc(label)}}</td><td><span class="heat h${{Math.min(5,Math.ceil(list.length/5))}}">${{list.length}}</span></td><td><span class="heat h${{Math.min(5,Math.ceil(opens/5))}}">${{opens}}</span></td><td>${{secondsText(analyticsAvgSeconds(duration,list.length))}}</td><td><span class="heat h${{Math.min(5,Math.ceil(subs/8))}}">${{subs}}</span></td></tr>`}}).join("")}}</tbody></table>`}}
function analyticsActivationTime(user){{return user.last_registered_at||user.registered_at||(Number(user.submission_count||0)>0?(user.submissions||[]).map(x=>x.created_at).filter(Boolean).sort()[0]:"")||user.last_login_at||user.created_at||""}}
function analyticsFilteredUsers(users){{if(analyticsUserFilter==="activated")return users.filter(u=>analyticsIsRegistered(u));if(analyticsUserFilter==="registered")return users.filter(u=>analyticsIsRegistered(u));if(analyticsUserFilter==="submitted")return users.filter(u=>Number(u.submission_count||0)>0);if(analyticsUserFilter==="followup")return users.filter(u=>!analyticsIsRegistered(u)&&(Number(u.script_share_open_count||0)>0||Number(u.submission_count||0)>0));return users}}
function analyticsFilterTabs(){{const tabs=[["activated","新激活"],["all","全部"],["submitted","有回传"],["followup","需跟进"]];return `<div class="analytics-tabs">${{tabs.map(([key,label])=>`<button class="${{analyticsUserFilter===key?"active":""}}" type="button" data-analytics-filter="${{key}}">${{label}}</button>`).join("")}}</div>`}}
function analyticsAuthorRows(users){{const sortTime=u=>Date.parse(analyticsActivationTime(u)||u.last_event_at||u.last_login_at||u.created_at||"")||0;const rows=analyticsFilteredUsers(users).slice().sort((a,b)=>(sortTime(b)-sortTime(a))||(Number(b.submission_count||0)-Number(a.submission_count||0))||(Number(b.script_share_open_count||0)-Number(a.script_share_open_count||0))||(Number(b.platform_open_count||0)-Number(a.platform_open_count||0)));if(!rows.length)return `<div class="empty">当前筛选下没有作者。</div>`;return `<div class="analytics-table"><div class="analytics-table-head"><span>作者</span><span>状态</span><span>激活/最近访问</span><span>脚本打开</span><span>回传</span><span>操作</span></div>${{rows.map(u=>`<div class="analytics-table-row"><div class="creator-row-main"><div class="avatar"></div><div class="creator-row-title"><b>${{esc(u.display_name||u.account_id||u.phone)}}</b><span class="small">${{esc(u.phone||u.account_id)}} · ${{esc(u.kwai_id||"")}}</span></div></div><div>${{analyticsStatusPill(u)}}</div><div class="small">激活：${{esc(timeText(analyticsActivationTime(u))||"-")}}<br>最近：${{esc(timeText(u.last_event_at||u.last_login_at)||"-")}}</div><b>${{Number(u.script_share_open_count||0)}}</b><b>${{Number(u.submission_count||0)}}</b><button type="button" data-copy="${{esc(u.phone||u.account_id||"")}}">复制账号</button>${{analyticsAuthorDetail(u)}}</div>`).join("")}}</div>`}}
function analyticsScriptRanking(users){{const map=new Map();for(const u of users){{for(const item of (u.script_views||[])){{const id=String(item.script_id||"");if(!id)continue;const row=map.get(id)||{{script_id:id,title:item.title||id,views:0,duration_ms:0,submissions:0}};row.views+=Number(item.views||0);row.duration_ms+=Number(item.duration_ms||0);map.set(id,row)}}for(const sub of (u.submissions||[])){{const id=String(sub.entry_id||"");if(!id)continue;const row=map.get(id)||{{script_id:id,title:sub.script_title||id,views:0,duration_ms:0,submissions:0}};row.submissions+=1;map.set(id,row)}}}}const rows=[...map.values()].sort((a,b)=>(b.submissions-a.submissions)||(b.views-a.views)||(b.duration_ms-a.duration_ms)).slice(0,6);if(!rows.length)return `<div class="empty">还没有脚本排行数据。</div>`;return `<div class="analytics-rank">${{rows.map((r,i)=>`<a href="__CREATOR_BASE__/script/${{esc(r.script_id)}}" target="_blank" rel="noopener"><span>${{i+1}}</span><b>${{esc(r.title||r.script_id)}}</b><small>打开 ${{r.views}} · 停留 ${{secondsText(Number(r.duration_ms||0)/1000)}} · 回传 ${{r.submissions}}</small></a>`).join("")}}</div>`}}
function wireAnalyticsControls(){{wireAnalyticsInactiveLoader();document.querySelectorAll("[data-analytics-filter]").forEach(btn=>btn.addEventListener("click",()=>{{analyticsUserFilter=btn.dataset.analyticsFilter||"all";renderAnalytics()}}))}}
function analyticsInactiveRows(rows,total,loaded){{const count=Number(total||rows.length||0);if(!count)return "";if(!loaded)return `<details class="import-panel" style="margin-top:14px" data-inactive-analytics><summary>无行为账号（${{count}} 个，点击展开加载）</summary><div class="empty" style="margin-top:12px">展开后会单独加载这些账号，避免首次进入数据看板超时。</div></details>`;return `<details class="import-panel" style="margin-top:14px" open data-inactive-analytics><summary>无行为账号（${{count}} 个）</summary><div class="creator-row-list" style="margin-top:12px">${{rows.map(user=>`<article class="analytics-row"><div class="creator-row-main"><div class="avatar"></div><div class="creator-row-title"><b>${{esc(user.display_name||user.account_id||user.phone)}}</b><span class="small">${{esc(user.phone||user.account_id)}} · 预置时间：${{esc(timeText(user.provisioned_at||user.created_at))}}</span></div></div><div>${{analyticsStatusPill(user)}}</div><div class="small">暂无打开、点击或回传行为</div></article>`).join("")||'<div class="empty">无行为账号加载完成，但没有可展示账号。</div>'}}</div></details>`}}
function wireAnalyticsInactiveLoader(){{const node=document.querySelector("[data-inactive-analytics]");if(!node)return;node.addEventListener("toggle",()=>{{if(node.open&&!analyticsInactiveLoaded&&!analyticsLoading)loadAnalytics(true)}},{{once:true}})}}
function renderAnalytics(){const box=document.querySelector("#analytics-board");if(!box)return;if(!analyticsData){box.innerHTML=`<div class="empty">正在自动加载有行为用户。无行为账号会保持折叠，展开时再加载。</div>`;return}const s=analyticsData.summary||{};const users=analyticsData.users||[];const inactive=analyticsData.inactive_users||[];const inactiveTotal=Number(s.inactive_accounts||inactive.length||0);const metricUsers=users.filter(u=>!analyticsIsTestUser(u));const registeredUsers=metricUsers.filter(u=>analyticsIsRegistered(u));const validRegisteredDurations=registeredUsers.map(u=>analyticsEffectiveDurationSeconds(u)).filter(Boolean);const registeredDurationSeconds=validRegisteredDurations.reduce((sum,value)=>sum+value,0);const avgRegisteredStaySeconds=validRegisteredDurations.length?registeredDurationSeconds/validRegisteredDurations.length:0;const submissionCount=metricUsers.reduce((sum,u)=>sum+Number(u.submission_count||0),0);const kpis=[["已注册用户数量",registeredUsers.length,`已排除测试账号 666`],["注册用户平均停留时长",secondsText(avgRegisteredStaySeconds),`有效停留样本 ${validRegisteredDurations.length} 个 · 已排除 666 和超过 30 分钟异常值`],["回传数量",submissionCount,`作者提交的视频外链总数 · 已排除 666`]];box.innerHTML=`<div class="analytics-shell"><section class="analytics-kpis analytics-core-kpis">${kpis.map(([label,value,sub])=>analyticsSummaryCard(label,value,sub)).join("")}</section>${analyticsCoreTimelineHtml()}<section class="analytics-layout" style="grid-template-columns:1fr"><div class="analytics-panel"><div class="analytics-panel-title"><h3>作者状态</h3>${analyticsFilterTabs()}</div>${analyticsAuthorRows(users)}</div></section>${analyticsInactiveRows(inactive,inactiveTotal,analyticsData.inactive_loaded)}</div>`;wireAnalyticsControls();const status=document.querySelector("#status");if(status)status.textContent=`核心数据已加载：已注册用户 ${registeredUsers.length} 个，平均停留 ${secondsText(avgRegisteredStaySeconds)}，回传 ${submissionCount} 条。测试账号 666 和超过 30 分钟异常停留已从顶部统计排除。无行为账号 ${inactiveTotal} 个已折叠。`}
async function loadAnalytics(includeInactive=false,opts={{}}){{const silent=!!opts.silent;activeTab="analytics";if(!document.querySelector("#analytics-board"))adminView();const status=document.querySelector("#status");const box=document.querySelector("#analytics-board");if(analyticsLoading)return;analyticsLoading=true;if(status&&!silent)status.textContent=includeInactive?"正在加载无行为账号...":"正在加载有行为用户...";if(box&&!analyticsData&&!silent)box.innerHTML=`<div class="empty">正在拉取有行为用户，不会等待无行为账号。</div>`;try{{const data=await api(`/api/creator-admin/analytics?days=180&include_inactive=${{includeInactive?"1":"0"}}`);if(includeInactive&&analyticsData){{analyticsData={{...analyticsData,summary:data.summary||analyticsData.summary,inactive_users:data.inactive_users||[],inactive_loaded:true}};analyticsInactiveLoaded=true}}else{{analyticsData=data;analyticsInactiveLoaded=Boolean(data.inactive_loaded)}}analyticsAutoLoaded=true;renderAnalytics();markRealtimeSynced(silent)}}catch(e){{const msg=e.message||"Creator 数据接口暂时不可用";if(status&&!silent)status.textContent=`数据加载失败：${{msg}}`;if(box&&!analyticsData&&!silent)box.innerHTML=`<div class="empty">数据加载失败：${{esc(msg)}}<br><br><button class="primary" id="retry-analytics" type="button">重新加载有行为用户</button></div>`;document.querySelector("#retry-analytics")?.addEventListener("click",()=>loadAnalytics(false))}}finally{{analyticsLoading=false}}}}
let creatorsLoadingPromise=null;
async function loadCreators(opts={{}}){{const silent=!!opts.silent;if(creatorsLoadingPromise)return creatorsLoadingPromise;creatorsLoadingPromise=(async()=>{{try{{const status=document.querySelector("#status");if(status&&!silent)status.textContent="加载注册用户、作者档案和行为数据中...";const [creatorData,accountData,analyticsPayload,cloudData]=await Promise.all([api("/api/creator-admin/creators"),api("/api/creator-admin/accounts"),api("/api/creator-admin/analytics?days=180&include_inactive=1"),api("/api/creator-admin/state")]);creators=creatorData.creators||[];accounts=accountData.accounts||[];analyticsData=analyticsPayload;analyticsAutoLoaded=true;analyticsInactiveLoaded=Boolean(analyticsPayload.inactive_loaded);creatorCloudState=cloudData.state&&cloudData.state.creators?cloudData.state:{{creators:{{}}}};creatorMergeRows();activeTab="creators";if(!silent||!document.querySelector("#creator-list"))adminView();else{{renderCreatorPocFilters();renderCreators()}}const registered=accounts.filter(a=>creatorAccountIsRegistered(a)).length;const nextStatus=document.querySelector("#status");if(nextStatus)nextStatus.textContent=silent?`实时同步完成：${{new Date().toLocaleTimeString("zh-CN",{{hour12:false}})}}`:`已合并 ${{registered}} 个注册用户和 POC 作者档案，共 ${{creatorRows.length}} 条运营对象`;markRealtimeSynced(silent);if(!scriptIndexLoaded)ensureScriptIndex().then(()=>{{if(activeTab==="creators")renderCreators()}}).catch(()=>null)}}catch(e){{if(!silent)loginView(e.message)}}}})().finally(()=>{{creatorsLoadingPromise=null}});return creatorsLoadingPromise}}
async function loadAccounts(){{try{{document.querySelector("#status")&&(document.querySelector("#status").textContent="加载账号中...");const d=await api("/api/creator-admin/accounts");accounts=d.accounts||[];activeTab="accounts";adminView();document.querySelector("#status").textContent=`共 ${{accounts.length}} 个账号`}}catch(e){{loginView(e.message)}}}}
async function loadSubmissions(append=false){{if(submissionsLoading)return;try{{submissionsLoading=true;const status=document.querySelector("#status");if(status)status.textContent=append?"加载更多回传中...":"加载最近回传数据中...";const offset=append?submissionsOffset:0;const d=await api(`/api/creator-admin/submissions?limit=${{SUBMISSIONS_PAGE_SIZE}}&offset=${{offset}}&_=${{Date.now()}}`);const incoming=d.submissions||[];submissions=append?[...submissions,...incoming]:incoming;accessApplications=Array.isArray(d.applications)?d.applications:accessApplications;submissionsTotal=Number(d.total||submissions.length||0);submissionsOffset=submissions.length;activeTab="submissions";adminView();const nextStatus=document.querySelector("#status");if(nextStatus)nextStatus.textContent=`已加载 ${{submissions.length}} / ${{submissionsTotal}} 条回传 · 账号申请 ${{accessApplications.length}} 条`}}catch(e){{loginView(e.message)}}finally{{submissionsLoading=false;renderSubmissionStats()}}}}
async function loadIntakes(){{try{{document.querySelector("#status")&&(document.querySelector("#status").textContent="加载作者信息中...");const d=await api("/api/creator-admin/intakes");intakes=d.intakes||[];activeTab="intakes";adminView();document.querySelector("#status").textContent=`共 ${{intakes.length}} 条作者信息`}}catch(e){{loginView(e.message)}}}}
async function loadCurrentTab(){{if(activeTab==="creators")return loadCreators();if(activeTab==="analytics"){{adminView();if(!analyticsAutoLoaded)setTimeout(()=>loadAnalytics(false),80);return}}if(activeTab==="accounts")return loadAccounts();if(activeTab==="submissions")return loadSubmissions();if(activeTab==="intakes")return loadIntakes();if(activeTab==="imports"){{adminView();return}}return loadEntries()}}
const CREATOR_REALTIME_MS=15000;let creatorRealtimeTimer=null;let creatorRealtimeBusy=false;
function adminHasOpenEditor(){{return document.querySelector(".modal.open")||document.activeElement?.matches?.("input,textarea,select")}}
function creatorHasOpenRecommendationBoard(){{const board=document.querySelector("[data-creator-recommend-board]");return Boolean(selectedCreatorId&&board&&board.dataset.creatorRecommendBoard===selectedCreatorId)}}
function markRealtimeSynced(silent=false){{const badge=document.querySelector("#creator-realtime-status");if(!badge)return;badge.textContent=`实时同步 · ${{new Date().toLocaleTimeString("zh-CN",{{hour12:false}})}}`;if(silent)badge.classList.add("active")}}
async function refreshCreatorAdminSilently(){{if(document.hidden||creatorRealtimeBusy||adminHasOpenEditor()||creatorHasOpenRecommendationBoard())return;creatorRealtimeBusy=true;try{{if(activeTab==="creators")await loadCreators({{silent:true}});else if(activeTab==="analytics")await loadAnalytics(false,{{silent:true}});else if(activeTab==="submissions")await loadSubmissions(false);}}finally{{creatorRealtimeBusy=false}}}}
function startCreatorRealtime(){{if(creatorRealtimeTimer)clearInterval(creatorRealtimeTimer);creatorRealtimeTimer=setInterval(refreshCreatorAdminSilently,CREATOR_REALTIME_MS)}}
function creatorPocValue(c){{return String(c.poc||c.owner||"").trim()}}
function creatorKeyValue(v){{return String(v||"").trim().toLowerCase().replace(/^@/,"").replace(/\\s+/g,"")}}
function creatorAccountKeys(item){{const keys=new Set();const account=item?.account||{{}};for(const source of [item,account]){{for(const field of ["kwai_id","uid","phone","account_id"]){{const value=creatorKeyValue(source?.[field]);if(value)keys.add(value)}}for(const alias of (source?.login_aliases||[])){{const value=creatorKeyValue(alias);if(value)keys.add(value)}}}}return keys}}
function creatorIdentityValue(item,field){{const account=item?.account||{{}};return creatorKeyValue(item?.[field]||account?.[field]||"")}}
function creatorIdentityPhones(item){{const account=item?.account||{{}};return new Set([item?.phone,account?.phone].map(creatorKeyValue).filter(Boolean))}}
function creatorIdentityAccountIds(item){{const account=item?.account||{{}};return new Set([item?.account_id,account?.account_id,...(item?.login_aliases||[]),...(account?.login_aliases||[])].map(creatorKeyValue).filter(Boolean))}}
function creatorIdentityMatches(profile,account){{const pk=creatorIdentityValue(profile,"kwai_id");const ak=creatorIdentityValue(account,"kwai_id");if(pk&&ak)return pk===ak;const pu=creatorIdentityValue(profile,"uid");const au=creatorIdentityValue(account,"uid");if(pu&&au)return pu===au;const profilePhones=creatorIdentityPhones(profile);const accountPhones=creatorIdentityPhones(account);if([...profilePhones].some(x=>accountPhones.has(x)))return true;if((pk||ak||pu||au)&&pk!==ak)return false;const accountIds=creatorIdentityAccountIds(account);return [...profilePhones].some(x=>accountIds.has(x))&&!pk&&!ak}}
function creatorPseudoId(account){{let hash=0;const raw=String(account?.account_id||account?.phone||account?.display_name||"account");for(let i=0;i<raw.length;i++)hash=((hash<<5)-hash+raw.charCodeAt(i))|0;const hex=Math.abs(hash).toString(16).padStart(8,"0");return (hex+hex+hex+hex).slice(0,32)}}
function creatorAccountIsRegistered(account){{return String(account?.registration_status||"").toLowerCase()==="registered"||Boolean(account?.registered_at||account?.last_registered_at)}}
function creatorScriptOpenCount(user){{const views=Array.isArray(user?.script_views)?user.script_views:[];const ids=new Set(views.map(v=>String(v.script_id||"").trim()).filter(Boolean));if(ids.size)return ids.size;return Number(user?.script_share_open_count||0)}}
function creatorAnalyticsUserForAccount(account,analyticsUsers){{return (analyticsUsers||[]).find(user=>creatorIdentityMatches(account,user))||{{}}}}
function creatorMergeSubmissions(...lists){{const seen=new Set();const rows=[];for(const list of lists){{for(const item of (Array.isArray(list)?list:[])){{const key=[item?.entry_id||"",item?.video_url||""].join("|");if(seen.has(key))continue;seen.add(key);rows.push(item)}}}}return rows}}
function creatorHasOpsData(c){{return Boolean(c?.profile_id&&(Number(c?.submission_count||0)||Number(c?.fed_script_count||0)||Number(c?.returned_script_count||0)||(Array.isArray(c?.submissions)&&c.submissions.length)||(creatorState(c.profile_id).feeds||[]).length))}}
function creatorMergeRows(){{const analyticsUsers=[...(analyticsData?.users||[]),...(analyticsData?.inactive_users||[])];const rows=[];const usedProfiles=new Set();const registeredAccounts=accounts.filter(a=>creatorAccountIsRegistered(a));for(const account of registeredAccounts){{const user=creatorAnalyticsUserForAccount(account,analyticsUsers);const profile=creators.find(c=>creatorIdentityMatches(c,account))||null;const row={{...(profile||{{}})}};if(profile?.profile_id)usedProfiles.add(profile.profile_id);row.profile_id=profile?.profile_id||creatorPseudoId(account);row.account=account;row.account_id=account.account_id||row.account_id||"";row.phone=account.phone||row.phone||row.account_id||"";row.kwai_id=account.kwai_id||row.kwai_id||"";row.uid=account.uid||row.uid||"";row.name=profile?.name||account.display_name||account.kwai_id||account.phone||account.account_id||"未绑定作者";row.analytics_user=user;row.registration_status="registered";row.creator_source=profile?"registered_bound":"registered_phone_only";row.poc=creatorPocValue(row)||"已注册";row.submissions=creatorMergeSubmissions(profile?.submissions,account.submissions,user.submissions);row.submission_count=Math.max(Number(profile?.submission_count||0),Number(account.submission_count||0),Number(user.submission_count||0),row.submissions.length);rows.push(row)}}for(const c of creators){{const poc=creatorPocValue(c);if((creatorPocOptions.includes(poc)||creatorHasOpsData(c))&&!usedProfiles.has(c.profile_id)){{const linkedAccount=c.account&&creatorIdentityMatches(c,c.account)?c.account:{{}};const clean={{...c,account:linkedAccount}};rows.push({{...clean,creator_source:creatorHasOpsData(clean)?"ops_profile":"poc_profile",analytics_user:creatorAnalyticsUserForAccount(clean,analyticsUsers)}})}}}creatorRows=rows;return rows}}
function creatorDisplayedRows(){{return creatorRows.length?creatorRows:creatorMergeRows()}}
function creatorPocCounts(){{const counts=new Map();for(const c of creatorDisplayedRows()){{const poc=creatorPocValue(c)||"未分配";counts.set(poc,(counts.get(poc)||0)+1)}}return counts}}
function renderCreatorPocFilters(){{const box=document.querySelector("#creator-poc-filters");if(!box)return;const rows=creatorDisplayedRows();const counts=creatorPocCounts();const dynamic=[...counts.keys()].filter(x=>x&&x!=="未分配"&&!creatorPocOptions.includes(x)).sort();const options=["",...creatorPocOptions,...dynamic,"未分配"];box.innerHTML=options.map(poc=>`<button class="quick-filter-chip ${{activeCreatorPoc===poc?"active":""}}" type="button" data-poc-filter="${{esc(poc)}}"><span>${{esc(poc||"全部作者")}}</span> <small>${{poc?Number(counts.get(poc)||0):rows.length}}</small></button>`).join("")}}
function creatorSearchText(c){{const user=c.analytics_user||{{}};return [c.name,c.kwai_id,c.kwai_url,c.uid,c.phone,c.account_id,c.poc,c.owner,c.cooperation_level,c.creator_description,c.registration_status,c.creator_source,user.display_name,user.account_id,user.phone,JSON.stringify(c.creator_type||{{}}),JSON.stringify(c.categories||[])].join(" ").toLowerCase()}}
function creatorFieldValue(c,scope){{const account=c.account||{{}};if(scope==="kwai")return [c.kwai_id,c.kwai_url].join(" ");if(scope==="name")return c.name||"";if(scope==="phone")return c.phone||account.phone||"";if(scope==="uid")return c.uid||account.uid||"";return creatorSearchText(c)}}
function creatorOpenCount(c){{return Number(c?.analytics_user?.platform_open_count||0)}}
function creatorOpenedScriptCount(c){{return creatorScriptOpenCount(c?.analytics_user||{{}})}}
function creatorRawSubmissionCount(c){{const ids=new Set((Array.isArray(c?.submissions)?c.submissions:[]).map(s=>String(s?.entry_id||s?.video_url||"").trim()).filter(Boolean));return ids.size}}
function creatorFeedCount(c){{return creatorFeedSummary(c).feedCount}}
function creatorFeedReturnCount(c){{return creatorFeedSummary(c).returnCount}}
function creatorSortTime(c){{return Date.parse(c?.analytics_user?.last_event_at||c?.analytics_user?.last_login_at||c?.account?.last_login_at||c?.account?.registered_at||c?.created_at||"")||0}}
function creatorSortedRows(rows){{return [...rows].sort((a,b)=>(creatorFeedReturnCount(b)-creatorFeedReturnCount(a))||(creatorFeedCount(b)-creatorFeedCount(a))||(creatorSortTime(b)-creatorSortTime(a))||(creatorOpenedScriptCount(b)-creatorOpenedScriptCount(a))||String(a.name||a.kwai_id||"").localeCompare(String(b.name||b.kwai_id||""),"zh-CN"))}}
function updateCreatorViewToggle(){{document.querySelectorAll("[data-creator-view]").forEach(btn=>btn.classList.toggle("active",btn.dataset.creatorView===creatorViewMode))}}
function filteredCreators(){{const base=creatorDisplayedRows();const q=String(document.querySelector("#creator-search")?.value||"").trim().toLowerCase();const scope=String(document.querySelector("#creator-search-scope")?.value||"all");let rows=q?base.filter(c=>String(creatorFieldValue(c,scope)).toLowerCase().includes(q)):base;if(activeCreatorPoc)rows=rows.filter(c=>(creatorPocValue(c)||"未分配")===activeCreatorPoc);return creatorSortedRows(rows)}}
function optionList(options,current,emptyLabel="待标注"){{const value=String(current||"").trim();return [`<option value="">${{esc(emptyLabel)}}</option>`,...options.map(x=>`<option value="${{esc(x)}}" ${{value===x?"selected":""}}>${{esc(x)}}</option>`)].join("")}}
function listValue(value){{if(Array.isArray(value))return value.map(x=>String(x||"").trim()).filter(Boolean);const text=String(value||"").trim();return text?text.split(/[、,，/]/).map(x=>x.trim()).filter(Boolean):[]}}
function multiOptionList(options,current){{const values=new Set(listValue(current));return options.map(x=>`<option value="${{esc(x)}}" ${{values.has(x)?"selected":""}}>${{esc(x)}}</option>`).join("")}}
function selectedValues(formEl,name){{return [...formEl.querySelectorAll(`[name="${{name}}"] option:checked`)].map(x=>x.value).filter(Boolean)}}
function tagLabel(value,empty="待标注"){{const values=listValue(value);return values.length?values.join("、"):empty}}
function metricNumber(value){{const n=Number(value||0);return Number.isFinite(n)?n:0}}
function creatorState(id){{return (creatorCloudState.creators&&creatorCloudState.creators[id])||{{}}}}
function setCreatorState(id,patch){{if(!creatorCloudState.creators)creatorCloudState.creators={{}};creatorCloudState.creators[id]={{...creatorState(id),...(patch||{{}})}};return creatorCloudState.creators[id]}}
function creatorFallbackAvatar(){{return "data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' width='96' height='96'%3E%3Crect width='96' height='96' rx='48' fill='%23ff8200'/%3E%3C/svg%3E"}}
function creatorMetric(c,keys,fallback="-"){{for(const key of keys){{const value=c?.[key];if(value!==undefined&&value!==null&&String(value).trim())return value}}return fallback}}
function openCreatorModal(){{creatorForm.reset();creatorForm.category.innerHTML=labels.map(x=>`<option value="${{esc(x)}}">${{esc(x)}}</option>`).join("");creatorForm.poc.innerHTML=optionList(creatorPocOptions,"","POC 待分配");creatorForm.identity.innerHTML=multiOptionList(["夫妻","情侣","家庭","朋友"],[]);creatorForm.location.innerHTML=multiOptionList(["家里","乡村","城市"],[]);creatorModal.classList.add("open")}}
function openCreatorTags(id){{const c=creatorDisplayedRows().find(x=>x.profile_id===id);if(!c)return;const type=c.creator_type||{{}};const account=c.account||{{}};creatorTagsForm.dataset.creatorTags=id;creatorTagsForm.kwai_url.value=c.kwai_url||"";creatorTagsForm.display_name.value=c.name||"";creatorTagsForm.kwai_id.value=c.kwai_id||"";creatorTagsForm.phone.value=c.phone||account.phone||"";creatorTagsForm.uid.value=c.uid||account.uid||"";creatorTagsForm.poc.innerHTML=optionList(creatorPocOptions,creatorPocValue(c),"POC 待分配");creatorTagsForm.category.innerHTML=optionList(labels,(c.categories||[])[0]||"","分类待标注");creatorTagsForm.identity.innerHTML=multiOptionList(["夫妻","情侣","家庭","朋友"],type.identity);creatorTagsForm.location.innerHTML=multiOptionList(["家里","乡村","城市"],type.location);creatorTagsForm.cooperation_level.innerHTML=optionList(["高","中","低","待观察"],c.cooperation_level,"配合度待标注");creatorTagsForm.creator_description.value=c.creator_description||"";creatorTagsModal.classList.add("open")}}
function selectCreator(id){{selectedCreatorId=id;history.pushState({{creator:id}},"",`/creator-admin/creators/${{id}}`);renderCreators()}}
function backToCreatorList(){{selectedCreatorId="";history.pushState({{}},"","/creator-admin/creators");renderCreators()}}
function creatorCombinedFeeds(c){{const id=c.profile_id;return visibleCreatorFeeds(id,[...creatorFeeds(id),...feedsFromSubmissions(c)].filter((feed,idx,arr)=>arr.findIndex(x=>x.entry_id===feed.entry_id)===idx))}}
function feedIsReturned(feed){{return String(feed.status||"").trim()==="已回传"||Boolean(String(feed.return_url||"").trim())}}
function creatorFeedSummary(c){{const feeds=creatorCombinedFeeds(c);const returned=feeds.filter(feedIsReturned).length;return {{feeds,feedCount:feeds.length,returnCount:returned}}}}
function creatorTile(c){{const account=c.account||{{}};const avatar=esc(c.avatar_url||creatorFallbackAvatar());const name=c.name||c.kwai_id||"Kwai creator";const kwai=c.kwai_id||"-";const uid=c.uid||account.uid||"-";const phone=c.phone||account.phone||"-";const poc=creatorPocValue(c)||"未分配";const source=c.creator_source==="registered_phone_only"?"待绑定 Kwai":(c.registration_status==="registered"?"已注册":"POC 档案");return `<article class="creator-tile" data-open-creator="${{esc(c.profile_id)}}"><div class="creator-tile-head"><img class="avatar" src="${{avatar}}" alt=""><div><h3 class="creator-tile-name">${{esc(name)}}</h3><div class="small">Kwai ID&nbsp; ${{esc(kwai)}}</div></div></div><div class="meta"><span class="creator-badge">${{esc(source)}}</span><span class="pill">POC ${{esc(poc)}}</span>${{(c.categories||[]).slice(0,2).map(x=>`<span class="pill">${{esc(x)}}</span>`).join("")}}</div><div class="creator-field-lines"><div class="creator-field-line"><span>UID</span><span>${{esc(uid)}}</span></div><div class="creator-field-line"><span>电话</span><span>${{esc(phone)}}</span></div></div><div class="creator-stats-row"><div class="creator-stat"><span>投喂脚本</span><b>${{creatorFeedCount(c)}}</b></div><div class="creator-stat"><span>回传脚本</span><b>${{creatorFeedReturnCount(c)}}</b></div></div></article>`}}
function creatorRow(c){{const type=c.creator_type||{{}};const avatar=esc(c.avatar_url||creatorFallbackAvatar());const name=c.name||c.kwai_id||"Kwai creator";const kwai=c.kwai_id||"-";const poc=creatorPocValue(c)||"未分配";const tags=[c.creator_source==="registered_phone_only"?"待绑定 Kwai":(c.registration_status==="registered"?"已注册":"POC 档案"),`POC ${{poc}}`,tagLabel(type.identity),tagLabel(type.location),...(c.categories||[])].filter(Boolean);return `<article class="creator-row" data-open-creator="${{esc(c.profile_id)}}"><div class="creator-row-main"><img class="avatar" src="${{avatar}}" alt=""><div class="creator-row-title"><b>${{esc(name)}}</b><span class="small">@${{esc(kwai)}} · 电话 ${{esc(c.phone||c.account?.phone||"-")}}</span></div></div><div class="creator-row-num" data-label="投喂">${{creatorFeedCount(c)}}</div><div class="creator-row-num" data-label="回传">${{creatorFeedReturnCount(c)}}</div><div class="creator-row-num" data-label="打开 Koko">${{creatorOpenCount(c)}}</div><div class="creator-row-num" data-label="打开脚本">${{creatorOpenedScriptCount(c)}}</div><div class="creator-row-tags">${{tags.slice(0,8).map(x=>`<span class="pill">${{esc(x)}}</span>`).join("")}}</div></article>`}}
function defaultCreatorMetrics(){{return {{opens:{{value:0,delta:0}},posts:{{value:0,delta:0}},views:{{value:0,delta:0}},avgViews:{{value:0,delta:0}}}}}}
function creatorMetrics(id){{return {{...defaultCreatorMetrics(),...(creatorState(id).metrics||{{}})}}}}
function metricCard(label,item){{const delta=Number(item?.delta||0);const cls=delta>=0?"delta-up":"delta-down";return `<div class="metric-card"><span>${{esc(label)}}</span><b>${{esc(item?.value??0)}}</b><em class="${{cls}}">${{delta>=0?"+":""}}${{esc(delta)}}%</em></div>`}}
function creatorMetricsPanel(c){{const user=c.analytics_user||{{}};const totalDuration=Number(user.platform_duration_seconds||0)+Number(user.script_duration_seconds||0);return `<section class="import-panel"><div class="feed-stats-head"><h3 style="margin:0;color:#ff8200">作者使用数据</h3><button type="button" data-copy="${{esc(c.phone||c.account_id||"")}}">复制账号</button></div><div class="metrics-grid">${{analyticsSummaryCard("打开 Koko 次数",creatorOpenCount(c),"只统计 Creator 首页/平台打开")}}${{analyticsSummaryCard("打开脚本个数",creatorOpenedScriptCount(c),"按看过的脚本去重，缺少明细时用打开数兜底")}}${{analyticsSummaryCard("总停留",secondsText(totalDuration),"首页 + 脚本页")}}</div>${{analyticsAuthorDetail(user,null,false)}}</section>`}}
function feedDate(){{return new Date().toISOString().slice(0,10)}}
function feedNormalize(item){{if(!item||typeof item!=="object")return null;const share=String(item.share_url||item.url||item.script_url||"").trim();const id=(share.match(/\\/script\\/([0-9a-f]{{32}})/)||[])[1]||String(item.entry_id||"").trim();if(!id)return null;const feedTime=String(item.feed_time||item.fed_at||item.date||feedDate()).trim();return {{entry_id:id,share_url:share||scriptShareUrl(id),date:feedTime.slice(0,10),feed_time:feedTime.slice(0,10),status:String(item.status||"未完成"),return_url:String(item.return_url||item.video_url||""),return_time:String(item.return_time||item.returned_at||"")}}}}
function parseFeedInput(text){{const raw=String(text||"").trim();if(!raw)return [];try{{const parsed=JSON.parse(raw);if(Array.isArray(parsed))return parsed.map(feedNormalize).filter(Boolean)}}catch(_e){{}}return raw.split(/\\n+/).map(line=>{{const parts=line.split(",").map(x=>x.trim());const url=parts.find(x=>/\\/script\\//.test(x))||parts[1]||parts[0];if(!url)return null;return feedNormalize({{date:parts[0]&&parts[0].includes("-")?parts[0]:feedDate(),share_url:url,status:parts[2]||"未完成",return_url:parts[3]||""}})}}).filter(Boolean)}}
function creatorFeeds(id){{return (creatorState(id).feeds||[]).map(feedNormalize).filter(Boolean)}}
function feedReturn(c,entryId){{return (c.submissions||[]).find(s=>s.status!=="placeholder"&&String(s.entry_id||"")===String(entryId||""))||null}}
function feedKey(feed){{return [feed.entry_id||"",feed.date||"",feed.return_url||""].join("|")}}
function deletedFeedKeys(id){{return new Set(creatorState(id).deleted_feed_keys||[])}}
function visibleCreatorFeeds(id,feeds){{const deleted=deletedFeedKeys(id);return feeds.filter(feed=>!deleted.has(feedKey(feed)))}}
function feedReturnDate(feed){{if(!feedIsReturned(feed))return "";const raw=String(feed.return_time||"").trim();if(raw)return raw.slice(0,10);return String(feed.date||feedDate()).slice(0,10)}}
function feedStatsHtml(feeds,id){{const rows=new Map();let returned=0;for(const feed of feeds){{const feedDay=String(feed.date||feed.feed_time||feedDate()).slice(0,10);if(!rows.has(feedDay))rows.set(feedDay,{{fed:0,returned:0}});rows.get(feedDay).fed+=1;const returnDay=feedReturnDate(feed);if(returnDay){{returned+=1;if(!rows.has(returnDay))rows.set(returnDay,{{fed:0,returned:0}});rows.get(returnDay).returned+=1}}}}const table=[...rows.entries()].sort((a,b)=>b[0].localeCompare(a[0])).slice(0,14).map(([date,row])=>`<div class="feed-date-row"><b>${{esc(date)}}</b><span>${{row.fed}} 条</span><span>${{row.returned}} 条</span></div>`).join("")||`<div class="feed-date-row"><b>暂无记录</b><span>0 条</span><span>0 条</span></div>`;return `<div class="feed-stats-head"><h4 class="feed-stats-title">统计</h4><button type="button" data-refresh-feed-stats="${{esc(id)}}">保存并刷新</button></div><div class="feed-stat-grid"><div class="feed-stat-card"><span>总投喂脚本</span><b>${{feeds.length}}</b></div><div class="feed-stat-card"><span>已回传脚本</span><b>${{returned}}</b></div><div class="feed-stat-card"><span>回传率</span><b>${{feeds.length?Math.round(returned/feeds.length*100):0}}%</b></div></div><div class="feed-date-table"><div class="feed-date-row"><b>日期</b><b>投喂</b><b>回传</b></div>${{table}}</div>`}}
function feedsFromSubmissions(c){{return (c.submissions||[]).filter(s=>s.status!=="placeholder"&&s.entry_id).map(s=>feedNormalize({{entry_id:s.entry_id,share_url:scriptShareUrl(s.entry_id),date:String(s.fed_at||s.feed_time||s.created_at||feedDate()).slice(0,10),status:"已回传",return_url:s.video_url||"",return_time:s.created_at||""}})).filter(Boolean)}}
function feedScriptCard(feed,c){{const entry=scriptIndex[feed.entry_id]||{{}};const ret=feedReturn(c,feed.entry_id);const image=entry.cover_url||entry.thumbnail_url||ret?.thumbnail_url||"";const share=feed.share_url||scriptShareUrl(feed.entry_id);const returnUrl=feed.return_url||ret?.video_url||"";const returnTime=feed.return_time||ret?.created_at||"";const returned=feedIsReturned(feed);const key=feedKey({{...feed,return_url:returnUrl}});return `<article class="feed-card" data-feed-entry="${{esc(feed.entry_id)}}" data-feed-key="${{esc(key)}}"><img src="${{esc(image)}}" loading="lazy" alt=""><div><h4>${{esc(entry.title||ret?.script_title||feed.entry_id)}}</h4><p>${{esc(entry.summary||"还没有加载到脚本概述，可先保留分享链接。")}}</p><div class="share-line"><b>脚本分享链接</b><a href="${{esc(share)}}" target="_blank" rel="noopener">${{esc(share)}}</a></div></div><div class="feed-controls"><div class="feed-time-grid"><div class="feed-time-field"><label>投喂时间</label><input data-feed-date="${{esc(feed.entry_id)}}" value="${{esc(feed.date)}}"></div><div class="feed-time-field"><label>回传时间</label><input data-feed-return-time="${{esc(feed.entry_id)}}" value="${{esc(returnTime?String(returnTime).slice(0,10):"")}}" placeholder="未回传"></div></div><select data-feed-status="${{esc(feed.entry_id)}}"><option ${{feed.status==="未完成"?"selected":""}}>未完成</option><option ${{feed.status==="正在制作"?"selected":""}}>正在制作</option><option ${{feed.status==="作者不愿意拍摄"?"selected":""}}>作者不愿意拍摄</option><option ${{feed.status==="已回传"?"selected":""}}>已回传</option></select><input data-feed-return-url="${{esc(feed.entry_id)}}" value="${{esc(returnUrl)}}" placeholder="作者回传视频链接"><div class="return-preview"><img src="${{esc(ret?.thumbnail_url||image)}}" alt=""><div><b>${{returned?"已回传":"未回传"}}</b><div class="small">投喂：${{esc(feed.date||"未记录")}}</div><div class="small">回传：${{esc(returnTime?timeText(returnTime):"等待回传")}}</div>${{returnUrl?`<a class="small" href="${{esc(returnUrl)}}" target="_blank" rel="noopener">打开视频</a>`:""}}</div></div><button class="danger" type="button" data-delete-feed="${{esc(key)}}" data-delete-feed-creator="${{esc(c.profile_id)}}">删除这条</button></div></article>`}}
function creatorFeedPanel(c){{const feeds=creatorFeedSummary(c).feeds;return `<section class="import-panel"><h3 style="margin:0;color:#ff8200">投喂脚本及回收情况统计</h3><div id="creator-feed-stats-${{esc(c.profile_id)}}" class="feed-stats">${{feedStatsHtml(feeds,c.profile_id)}}</div><details><summary>导入/补充投喂脚本</summary><div class="mini-import"><textarea id="creator-feed-input-${{esc(c.profile_id)}}" placeholder="每行一个：2026-07-22, https://kokocomedy.com/script/..., 未完成, 回传链接（可选）"></textarea><button class="primary" type="button" data-save-creator-feeds="${{esc(c.profile_id)}}">保存投喂</button></div></details><div id="creator-feed-list-${{esc(c.profile_id)}}" class="feed-list">${{feeds.length?feeds.map(feed=>feedScriptCard(feed,c)).join(""):'<div class="empty">正在整理已有投喂和推荐脚本...</div>'}}</div></section>`}}
async function hydrateCreatorFeeds(id,persist=false){{const c=creatorDisplayedRows().find(x=>x.profile_id===id);const box=document.getElementById(`creator-feed-list-${{id}}`);const stats=document.getElementById(`creator-feed-stats-${{id}}`);if(!c||!box)return;try{{await ensureScriptIndex();let feeds=[...creatorFeeds(id),...feedsFromSubmissions(c)].filter((feed,idx,arr)=>arr.findIndex(x=>x.entry_id===feed.entry_id)===idx);if(!feeds.length&&c.creator_source!=="registered_phone_only"){{const d=await api(`/api/creator-admin/creators/${{id}}/recommendations?limit=5&offset=0`);for(const script of (d.scripts||[])){{if(script?.entry_id)scriptIndex[script.entry_id]=script;feeds.push(feedNormalize({{entry_id:script.entry_id,share_url:`__CREATOR_BASE__${{script.share_url||(`/script/${{script.entry_id}}`)}}`,date:feedDate(),status:"未完成"}}))}}feeds=feeds.filter(Boolean)}}feeds=visibleCreatorFeeds(id,feeds);if(persist){{setCreatorState(id,{{feeds}});await saveCreatorCloudState(id,{{feeds}})}}if(stats)stats.innerHTML=feedStatsHtml(feeds,id);box.innerHTML=feeds.length?feeds.map(feed=>feedScriptCard(feed,c)).join(""):'<div class="empty">还没有投喂记录或回传记录。只有手机号的自助注册账号需要先补 Kwai ID 后再分配推荐。</div>'}}catch(err){{box.innerHTML=`<div class="empty">${{esc(err.message||err)}}</div>`}}}}
function creatorDetail(c){{const account=c.account||{{}};const type=c.creator_type||{{}};const avatar=esc(c.avatar_url||creatorFallbackAvatar());const name=c.name||c.kwai_id||"Kwai creator";const canDelete=c.creator_source!=="registered_phone_only";const recommend=c.creator_source==="registered_phone_only"?`<section class="import-panel creator-recommend-panel"><h3 style="margin:0;color:#ff8200">待绑定 Kwai 身份</h3><p class="small">这个账号来自自助注册，目前只有手机号。点击“编辑标签”补充 Kwai ID / 作者名称后，就可以正式进入创作者运营流程。</p></section>`:`<section class="import-panel creator-recommend-panel"><h3 style="margin:0;color:#ff8200">推荐脚本</h3><button type="button" data-load-creator-scripts="${{esc(c.profile_id)}}">查看推荐脚本</button><div id="creator-scripts-${{esc(c.profile_id)}}" data-day-offset="0"></div></section>`;return `<div class="creator-detail-top"><button type="button" data-back-creators>返回列表</button><div class="actions"><button class="primary" type="button" data-edit-creator-tags="${{esc(c.profile_id)}}">编辑标签</button>${{canDelete?`<button class="danger" type="button" data-delete-creator="${{esc(c.profile_id)}}">删除</button>`:""}}</div></div><section class="creator-detail-hero"><img class="avatar" src="${{avatar}}" alt=""><div><h2 style="margin:0 0 6px">${{esc(name)}}</h2><div class="small">@${{esc(c.kwai_id||"未绑定")}} · UID ${{esc(c.uid||account.uid||"未填")}} · 电话 ${{esc(c.phone||account.phone||"未填")}}</div><a class="small" href="${{esc(c.kwai_url)}}" target="_blank" rel="noopener">${{esc(c.kwai_url||"")}}</a><div class="meta"><span class="pill">${{esc(c.registration_status==="registered"?"已注册":"POC 档案")}}</span><span class="pill">POC ${{esc(creatorPocValue(c)||"未分配")}}</span><span class="pill">身份 ${{esc(tagLabel(type.identity))}}</span><span class="pill">地点 ${{esc(tagLabel(type.location))}}</span>${{(c.categories||[]).map(x=>`<span class="pill">${{esc(x)}}</span>`).join("")}}</div></div><div class="submission-count"><b>${{creatorFeedReturnCount(c)}}</b><span>回传</span></div></section><div class="creator-detail-stack">${{recommend}}${{creatorFeedPanel(c)}}${{creatorMetricsPanel(c)}}</div>`}}
function creatorIsTestRow(c){{return analyticsIsTestUser(c)||analyticsIsTestUser(c?.account)||analyticsIsTestUser(c?.analytics_user)}}
function ensureCreatorDailyStyles(){{if(document.getElementById("creator-daily-style"))return;const style=document.createElement("style");style.id="creator-daily-style";style.textContent=`.creator-daily-detail{{margin:0 0 16px;border:1px solid rgba(255,130,0,.16);border-radius:24px;background:rgba(255,255,255,.88);box-shadow:0 12px 30px rgba(249,115,0,.08);overflow:hidden}}.creator-daily-detail>summary{{display:flex;align-items:center;justify-content:space-between;gap:12px;padding:16px 18px;color:#ff5f00;font-size:18px;font-weight:950;cursor:pointer}}.creator-daily-detail>summary small{{color:#99520f;font-size:12px;font-weight:850;text-align:right}}.creator-daily-list{{display:grid;gap:12px;padding:0 16px 16px}}.creator-day-detail{{border:1px solid rgba(255,130,0,.14);border-radius:18px;background:#fffaf5;overflow:hidden}}.creator-day-detail>summary{{display:grid;grid-template-columns:150px repeat(3,minmax(110px,1fr));gap:12px;align-items:center;padding:14px 16px;cursor:pointer;font-weight:950}}.creator-day-detail>summary b{{color:#1f1f1f}}.creator-day-detail>summary span{{color:#ff5f00}}.creator-day-sections{{display:grid;grid-template-columns:repeat(3,minmax(0,1fr));gap:12px;padding:0 14px 14px}}.creator-person-list{{display:grid;grid-template-columns:repeat(2,minmax(0,1fr));gap:12px;padding:0 14px 14px}}.creator-person-card{{display:grid;gap:12px;padding:14px;border:1px solid rgba(255,130,0,.14);border-radius:18px;background:#fff}}.creator-person-head{{display:flex;align-items:flex-start;justify-content:space-between;gap:12px;border-bottom:1px solid rgba(255,130,0,.12);padding-bottom:10px}}.creator-person-head b{{font-size:18px;color:#1f1f1f}}.creator-person-stats{{display:flex;gap:6px;flex-wrap:wrap;justify-content:flex-end}}.creator-person-stats span{{padding:5px 8px;border-radius:999px;background:#fff7f0;color:#ff5f00;font-size:12px;font-weight:900}}.creator-person-section{{display:grid;gap:7px}}.creator-person-section h4{{margin:0;color:#99520f;font-size:13px}}.creator-person-line{{display:grid;gap:3px;padding:9px;border-radius:12px;background:#fff7f0;border:1px solid rgba(31,41,55,.06)}}.creator-person-line small,.creator-person-line a{{font-size:12px;line-height:1.35;overflow-wrap:anywhere;color:#6f737a}}.creator-person-line a{{color:#ff5f00;font-weight:850}}.creator-day-section{{display:grid;gap:8px;align-content:start;padding:12px;border:1px solid rgba(255,130,0,.12);border-radius:16px;background:#fff}}.creator-day-section h4{{margin:0;color:#99520f;font-size:14px}}.creator-day-item{{display:grid;gap:3px;padding:9px;border-radius:12px;background:#fff7f0;border:1px solid rgba(31,41,55,.06)}}.creator-day-item b{{font-size:13px;color:#1f1f1f}}.creator-day-item small,.creator-day-item a{{font-size:12px;line-height:1.35;overflow-wrap:anywhere;color:#6f737a}}.creator-day-item a{{color:#ff5f00;font-weight:850}}@media(max-width:980px){{.creator-day-sections,.creator-person-list{{grid-template-columns:1fr}}.creator-day-detail>summary{{grid-template-columns:1fr 1fr}}}}`;document.head.appendChild(style)}}
function creatorDailyPersonName(item){{return String(item?.display_name||item?.name||item?.kwai_id||item?.phone||item?.account_id||item?.visitor_id||"匿名访客")}}
function creatorDailyEntryTitle(entryId,fallback=""){{const entry=scriptIndex[String(entryId||"")]||{{}};return String(entry.title||fallback||entryId||"未记录脚本")}}
function creatorDailyScriptUrl(entryId,shareUrl=""){{const share=String(shareUrl||"").trim();if(/^https?:\\/\\//i.test(share))return share;return scriptShareUrl(entryId)}}
function creatorDailyBucket(map,day){{const key=day&&day!=="未记录日期"?day:"未记录日期";if(!map.has(key))map.set(key,{{users:new Map(),uses:[],feeds:[],returns:[]}});return map.get(key)}}
function creatorDailyUseItems(){{const out=[];const hourRows=analyticsData?.timeline?.hourly||[];for(const row of hourRows){{for(const item of (row.script_open_details||[])){{if(analyticsIsTestUser(item))continue;const scriptId=String(item.script_id||"").trim();const key=analyticsEntityKey(item)||String(item.visitor_id||item.display_name||item.phone||item.account_id||"").trim();out.push({{...item,day:analyticsDayKey(item.time||row.hour),person_key:key||creatorDailyPersonName(item),person_name:creatorDailyPersonName(item),script_id:scriptId,script_title:creatorDailyEntryTitle(scriptId,item.script_title),time:item.time||row.hour}})}}}}return out}}
function creatorDailyFeedItems(rows){{const out=[];const seen=new Set();for(const c of (rows||creatorDisplayedRows())){{if(creatorIsTestRow(c))continue;const creatorName=String(c.name||c.kwai_id||c.phone||c.account_id||"未命名作者");for(const feed of creatorCombinedFeeds(c)){{if(!feed||!feed.entry_id)continue;const day=creatorCleanDay(feed.date||feed.feed_time);if(!day)continue;const key=[c.profile_id||creatorName,feed.entry_id,day,feed.return_url||""].join("|");if(seen.has(key))continue;seen.add(key);out.push({{day,creator_name:creatorName,entry_id:feed.entry_id,script_title:creatorDailyEntryTitle(feed.entry_id),share_url:creatorDailyScriptUrl(feed.entry_id,feed.share_url),feed_time:day,status:feed.status||"未完成",return_url:feed.return_url||"",return_time:feed.return_time||""}})}}}}return out}}
function creatorDailyReturnItems(rows){{return creatorDailyFeedItems(rows).filter(item=>item.return_url||String(item.status||"").trim()==="已回传").map(item=>({{...item,day:creatorCleanDay(item.return_time?analyticsDayKey(item.return_time):item.day)||item.day}}))}}
function creatorDailySection(title,items,renderer,empty){{return `<section class="creator-day-section"><h4>${{esc(title)}}</h4>${{items.length?items.slice(0,80).map(renderer).join(""):`<div class="small">${{esc(empty)}}</div>`}}${{items.length>80?`<small class="small">还有 ${{items.length-80}} 条未展开显示。</small>`:""}}</section>`}}
function creatorDailyPersonKey(name){{return String(name||"未命名作者").trim().toLowerCase().replace(/^@/,"")||"未命名作者"}}
function creatorDailyPersonMap(day,bucket,includeUses=true){{const people=new Map();const ensure=name=>{{const clean=String(name||"未命名作者").trim()||"未命名作者";const key=creatorDailyPersonKey(clean);if(!people.has(key))people.set(key,{{name:clean,uses:[],feeds:[],returns:[]}});return people.get(key)}};if(includeUses&&creatorUsageDayIsClean(day)){{for(const item of bucket.uses||[])ensure(item.person_name).uses.push(item)}}for(const item of bucket.feeds||[])ensure(item.creator_name).feeds.push(item);for(const item of bucket.returns||[])ensure(item.creator_name).returns.push(item);return [...people.values()].sort((a,b)=>((b.uses.length+b.feeds.length+b.returns.length)-(a.uses.length+a.feeds.length+a.returns.length))||String(a.name).localeCompare(String(b.name),"zh-CN"))}}
function creatorPersonLineList(title,items,renderer,empty){{if(!items.length)return "";return `<section class="creator-person-section"><h4>${{esc(title)}}</h4>${{items.slice(0,12).map(renderer).join("")}}${{items.length>12?`<small class="small">还有 ${{items.length-12}} 条同作者记录。</small>`:""}}</section>`}}
function creatorPersonCardHtml(person){{const uses=person.uses.slice().sort((a,b)=>String(a.time||"").localeCompare(String(b.time||"")));const feeds=person.feeds.slice().sort((a,b)=>String(a.feed_time||"").localeCompare(String(b.feed_time||"")));const returns=person.returns.slice().sort((a,b)=>String(a.return_time||"").localeCompare(String(b.return_time||"")));return `<article class="creator-person-card"><div class="creator-person-head"><b>${{esc(person.name)}}</b><div class="creator-person-stats"><span>打开 ${{uses.length}}</span><span>投喂 ${{feeds.length}}</span><span>回传 ${{returns.length}}</span></div></div>${{creatorPersonLineList("打开脚本",uses,item=>`<div class="creator-person-line"><small>${{esc(timeText(item.time))}} · ${{esc(item.source||item.path||"打开脚本")}}</small><small>脚本：${{esc(item.script_title)}}</small></div>`)}}${{creatorPersonLineList("投喂脚本",feeds,item=>`<div class="creator-person-line"><small>投喂时间：${{esc(item.feed_time||item.day)}}</small><a href="${{esc(item.share_url)}}" target="_blank" rel="noopener">脚本：${{esc(item.script_title)}}</a></div>`)}}${{creatorPersonLineList("回传视频",returns,item=>`<div class="creator-person-line"><small>回传时间：${{esc(item.return_time?timeText(item.return_time):item.day)}}</small><a href="${{esc(item.share_url)}}" target="_blank" rel="noopener">脚本：${{esc(item.script_title)}}</a>${{item.return_url?`<a href="${{esc(item.return_url)}}" target="_blank" rel="noopener">视频：${{esc(item.return_url)}}</a>`:""}}</div>`)}}</article>`}}
function creatorPersonCardsHtml(day,bucket,empty,includeUses=true){{const people=creatorDailyPersonMap(day,bucket,includeUses);return people.length?people.map(creatorPersonCardHtml).join(""):`<div class="empty">${{esc(empty||"当天没有明细。")}}</div>`}}
function creatorExportRows(){{const rows=[];const metricRows=creatorDisplayedRows().filter(c=>!creatorIsTestRow(c));for(const c of metricRows){{const account=c.account||{{}};const user=c.analytics_user||{{}};rows.push({{record_type:"作者状态",month:"",week:"",day:"",creator_name:c.name||c.kwai_id||c.phone||c.account_id||"",account_id:c.account_id||account.account_id||"",phone:c.phone||account.phone||"",kwai_id:c.kwai_id||account.kwai_id||"",uid:c.uid||account.uid||"",poc:creatorPocValue(c)||"",event_time:user.last_event_at||user.last_login_at||account.last_login_at||account.registered_at||"",event_source:c.creator_source||"",script_id:"",script_title:"",script_url:"",video_url:"",status:c.registration_status||account.registration_status||"",platform_open_count:creatorOpenCount(c),script_open_count:creatorOpenedScriptCount(c),total_duration_seconds:Number(user.platform_duration_seconds||0)+Number(user.script_duration_seconds||0),feed_count:creatorFeedCount(c),return_count:creatorFeedReturnCount(c)}})}}for(const item of creatorDailyUseItems()){{if(!creatorUsageDayIsClean(item.day))continue;const info=creatorWeekInfo(item.day);rows.push({{record_type:"打开脚本",month:creatorMonthKey(item.day),week:info.label,day:item.day,creator_name:item.person_name||"",account_id:item.account_id||"",phone:item.phone||"",kwai_id:item.kwai_id||"",uid:item.uid||"",poc:"",event_time:item.time||"",event_source:item.source||item.path||"打开脚本",script_id:item.script_id||"",script_title:item.script_title||"",script_url:item.script_id?scriptShareUrl(item.script_id):"",video_url:"",status:"",platform_open_count:"",script_open_count:"",total_duration_seconds:"",feed_count:"",return_count:""}})}}for(const item of creatorDailyFeedItems(metricRows)){{const info=creatorWeekInfo(item.day);rows.push({{record_type:"投喂脚本",month:creatorMonthKey(item.day),week:info.label,day:item.day,creator_name:item.creator_name||"",account_id:"",phone:"",kwai_id:"",uid:"",poc:"",event_time:item.feed_time||item.day||"",event_source:"投喂",script_id:item.entry_id||"",script_title:item.script_title||"",script_url:item.share_url||creatorDailyScriptUrl(item.entry_id),video_url:"",status:item.status||"",platform_open_count:"",script_open_count:"",total_duration_seconds:"",feed_count:"",return_count:""}})}}for(const item of creatorDailyReturnItems(metricRows)){{const info=creatorWeekInfo(item.day);rows.push({{record_type:"回传视频",month:creatorMonthKey(item.day),week:info.label,day:item.day,creator_name:item.creator_name||"",account_id:"",phone:"",kwai_id:"",uid:"",poc:"",event_time:item.return_time||item.day||"",event_source:"回传",script_id:item.entry_id||"",script_title:item.script_title||"",script_url:item.share_url||creatorDailyScriptUrl(item.entry_id),video_url:item.return_url||"",status:item.status||"已回传",platform_open_count:"",script_open_count:"",total_duration_seconds:"",feed_count:"",return_count:""}})}}return rows}}
function downloadCreatorDetailCsv(){{const rows=creatorExportRows();const headers=["record_type","month","week","day","creator_name","account_id","phone","kwai_id","uid","poc","event_time","event_source","script_id","script_title","script_url","video_url","status","platform_open_count","script_open_count","total_duration_seconds","feed_count","return_count"];const cell=v=>`"${{String(v??"").replaceAll('"','""')}}"`;const csv=["\\ufeff"+headers.join(","),...rows.map(row=>headers.map(h=>cell(row[h])).join(","))].join("\\n");const blob=new Blob([csv],{{type:"text/csv;charset=utf-8"}});const url=URL.createObjectURL(blob);const a=document.createElement("a");a.href=url;a.download=`koko-creator-detail-${{creatorLocalDay()}}.csv`;document.body.appendChild(a);a.click();a.remove();URL.revokeObjectURL(url);const status=document.querySelector("#status");if(status)status.textContent=`已导出 ${{rows.length}} 条详细数据，可用 Excel 打开。`}}
const CREATOR_USAGE_CUTOFF_DAY="2026-08-07";
function creatorUsageDayIsClean(day){{return String(day||"")>=CREATOR_USAGE_CUTOFF_DAY}}
function creatorUsageText(dayOrBuckets,bucket=null){{if(bucket)return creatorUsageDayIsClean(dayOrBuckets)?String(bucket.users.size):"/";const users=new Set();let hasClean=false;for(const [,b] of dayOrBuckets){{if(!creatorUsageDayIsClean(b.day))continue;hasClean=true;for(const key of b.users.keys())users.add(key)}}return hasClean?String(users.size):"/"}}
function creatorCleanDay(value){{const raw=String(value||"").trim();if(!raw)return "";const m=raw.match(/(\\d{{4}})\\D+(\\d{{1,2}})(?:\\D+(\\d{{1,2}}))?/);if(m){{const y=m[1];const mo=String(Math.min(12,Math.max(1,Number(m[2]||1)))).padStart(2,"0");const da=String(Math.min(31,Math.max(1,Number(m[3]||1)))).padStart(2,"0");return `${{y}}-${{mo}}-${{da}}`}}const d=new Date(raw);return Number.isNaN(d.getTime())?"":d.toLocaleDateString("sv-SE",{{timeZone:"Asia/Shanghai"}})}}
function creatorMonthKey(day){{const clean=creatorCleanDay(day)||String(day||"");return clean.slice(0,7)}}
function creatorWeekInfo(day){{const clean=creatorCleanDay(day)||String(day||"");const d=new Date(`${{clean}}T00:00:00+08:00`);if(Number.isNaN(d.getTime()))return {{key:`${{clean}}-week`,label:"未记录周"}};const offset=(d.getDay()+6)%7;const start=new Date(d);start.setDate(d.getDate()-offset);const end=new Date(start);end.setDate(start.getDate()+6);const ymdLocal=x=>x.toLocaleDateString("sv-SE",{{timeZone:"Asia/Shanghai"}});return {{key:ymdLocal(start),label:`${{ymdLocal(start).replaceAll("-",".")}} - ${{ymdLocal(end).replaceAll("-",".")}}`}}}}
function creatorLocalDay(value=new Date()){{const d=value instanceof Date?value:new Date(value);return Number.isNaN(d.getTime())?feedDate():d.toLocaleDateString("sv-SE",{{timeZone:"Asia/Shanghai"}})}}
const creatorOpenDetailKeys=new Set(JSON.parse(localStorage.getItem("kokoCreatorOpenDetails")||"[]"));
function creatorDetailOpenAttr(key){{const safe=String(key||"");return ` data-creator-detail-key="${{esc(safe)}}"${{creatorOpenDetailKeys.has(safe)?" open":""}}`}}
function creatorDailyItemHtml(day,bucket){{const feeds=bucket.feeds.slice().sort((a,b)=>String(a.feed_time||"").localeCompare(String(b.feed_time||"")));const returns=bucket.returns.slice().sort((a,b)=>String(a.return_time||"").localeCompare(String(b.return_time||"")));const empty=creatorUsageDayIsClean(day)?"当天没有脚本打开、投喂或回传记录。":"8月7日之前的打开数据不计入，用 / 表示。";return `<details class="creator-day-detail"${{creatorDetailOpenAttr(`daily:day:${{day}}`)}}><summary><b>${{esc(day.replaceAll("-","."))}}</b><span>使用 ${{esc(creatorUsageText(day,bucket))}} 人</span><span>投喂 ${{feeds.length}} 个</span><span>回传 ${{returns.length}} 个</span></summary><div class="creator-person-list">${{creatorPersonCardsHtml(day,bucket,empty,true)}}</div></details>`}}
function creatorPeriodCounts(dayBuckets){{let feeds=0;let returns=0;for(const [,b] of dayBuckets){{feeds+=b.feeds.length;returns+=b.returns.length}}return {{feeds,returns,usage:creatorUsageText(dayBuckets)}}}}
function creatorDailyDetailsHtml(metricRows){{ensureCreatorDailyStyles();const map=new Map();for(const item of creatorDailyUseItems()){{if(!creatorUsageDayIsClean(item.day))continue;const bucket=creatorDailyBucket(map,item.day);bucket.day=item.day;bucket.uses.push(item);if(item.person_key)bucket.users.set(item.person_key,item)}}for(const item of creatorDailyFeedItems(metricRows)){{const bucket=creatorDailyBucket(map,item.day);bucket.day=item.day;bucket.feeds.push(item)}}for(const item of creatorDailyReturnItems(metricRows)){{const bucket=creatorDailyBucket(map,item.day);bucket.day=item.day;bucket.returns.push(item)}}const days=[...map.entries()].filter(([day,b])=>day&&day!=="未记录日期"&&(b.users.size||b.feeds.length||b.returns.length)).sort((a,b)=>String(b[0]).localeCompare(String(a[0])));if(!days.length)return `<details class="creator-daily-detail"${{creatorDetailOpenAttr("daily:root")}}><summary><span>查看具体数据</span><button type="button" data-export-creator-details>导出详细数据</button><small>暂无脚本打开、投喂或回传明细；8月7日前打开数据用 / 表示</small></summary><div class="creator-daily-list"><div class="empty">目前还没有可展开的每日数据。</div></div></details>`;const months=new Map();for(const item of days){{const month=creatorMonthKey(item[0]);if(!months.has(month))months.set(month,[]);months.get(month).push(item)}}return `<details class="creator-daily-detail"${{creatorDetailOpenAttr("daily:root")}}><summary><span>查看具体数据</span><button type="button" data-export-creator-details>导出详细数据</button><small>先看月统计，再展开到周和日；8月7日前打开数据不计入</small></summary><div class="creator-daily-list">${{[...months.entries()].sort((a,b)=>String(b[0]).localeCompare(String(a[0]))).map(([month,monthDays])=>{{const monthCounts=creatorPeriodCounts(monthDays);const weeks=new Map();for(const item of monthDays){{const info=creatorWeekInfo(item[0]);if(!weeks.has(info.key))weeks.set(info.key,{{label:info.label,days:[]}});weeks.get(info.key).days.push(item)}}return `<details class="creator-day-detail creator-month-detail"${{creatorDetailOpenAttr(`daily:month:${{month}}`)}}><summary><b>${{esc(month.replace("-","."))}}</b><span>使用 ${{esc(monthCounts.usage)}} 人</span><span>投喂 ${{monthCounts.feeds}} 个</span><span>回传 ${{monthCounts.returns}} 个</span></summary><div class="creator-daily-list">${{[...weeks.entries()].sort((a,b)=>String(b[0]).localeCompare(String(a[0]))).map(([weekKey,week])=>{{const weekCounts=creatorPeriodCounts(week.days);return `<details class="creator-day-detail creator-week-detail"${{creatorDetailOpenAttr(`daily:week:${{weekKey}}`)}}><summary><b>${{esc(week.label)}}</b><span>使用 ${{esc(weekCounts.usage)}} 人</span><span>投喂 ${{weekCounts.feeds}} 个</span><span>回传 ${{weekCounts.returns}} 个</span></summary><div class="creator-daily-list">${{week.days.map(([day,bucket])=>creatorDailyItemHtml(day,bucket)).join("")}}</div></details>`}}).join("")}}</div></details>`}}).join("")}}</div></details>`}}
function creatorPocDailyMap(rows,poc){{
const map=new Map();
const ensure=day=>{{
if(!map.has(day))map.set(day,{{day,feeds:[],returns:[],feedCreators:new Map(),returnCreators:new Map()}});
return map.get(day);
}};
for(const c of rows){{
if(creatorIsTestRow(c)||creatorPocValue(c)!==poc)continue;
const creatorName=String(c.name||c.kwai_id||c.phone||c.account_id||"未命名作者");
for(const feed of creatorCombinedFeeds(c)){{
if(!feed?.entry_id)continue;
const feedDay=creatorCleanDay(feed.date||feed.feed_time);
if(feedDay){{
const item={{day:feedDay,poc,creator_name:creatorName,entry_id:feed.entry_id,script_title:creatorDailyEntryTitle(feed.entry_id),share_url:creatorDailyScriptUrl(feed.entry_id,feed.share_url),feed_time:feedDay}};
const bucket=ensure(feedDay);
bucket.feeds.push(item);
bucket.feedCreators.set(creatorName,(bucket.feedCreators.get(creatorName)||0)+1);
}}
if(feedIsReturned(feed)){{
const returnDay=creatorCleanDay(feed.return_time?analyticsDayKey(feed.return_time):feedReturnDate(feed));
if(returnDay){{
const item={{day:returnDay,poc,creator_name:creatorName,entry_id:feed.entry_id,script_title:creatorDailyEntryTitle(feed.entry_id),share_url:creatorDailyScriptUrl(feed.entry_id,feed.share_url),return_url:feed.return_url||"",return_time:feed.return_time||returnDay}};
const bucket=ensure(returnDay);
bucket.returns.push(item);
bucket.returnCreators.set(creatorName,(bucket.returnCreators.get(creatorName)||0)+1);
}}
}}
}}
}}
return map;
}}
function creatorPocPeriodCounts(dayBuckets){{let feeds=0;let returns=0;const feedCreators=new Set();const returnCreators=new Set();for(const [,b] of dayBuckets){{feeds+=b.feeds.length;returns+=b.returns.length;for(const k of b.feedCreators.keys())feedCreators.add(k);for(const k of b.returnCreators.keys())returnCreators.add(k)}}return {{feeds,returns,feedCreators:feedCreators.size,returnCreators:returnCreators.size}}}}
function creatorPocPersonCounts(map){{const rows=[...map.entries()].sort((a,b)=>b[1]-a[1]);return rows.length?rows.map(([name,count])=>`<span class="pill">${{esc(name)}} ${{count}} 条</span>`).join(""):`<span class="small">暂无</span>`}}
function creatorPocDayHtml(day,bucket){{const feeds=bucket.feeds.slice().sort((a,b)=>String(a.creator_name).localeCompare(String(b.creator_name))||String(a.script_title).localeCompare(String(b.script_title)));const returns=bucket.returns.slice().sort((a,b)=>String(a.creator_name).localeCompare(String(b.creator_name))||String(a.return_time||"").localeCompare(String(b.return_time||"")));const grouped={{...bucket,feeds,returns,uses:[]}};return `<details class="creator-day-detail"${{creatorDetailOpenAttr(`poc:day:${{bucket.poc||"all"}}:${{day}}`)}}><summary><b>${{esc(day.replaceAll("-","."))}}</b><span>投喂 ${{feeds.length}} 个</span><span>覆盖 ${{bucket.feedCreators.size}} 人</span><span>回传 ${{returns.length}} 个</span></summary><div class="creator-person-list">${{creatorPersonCardsHtml(day,grouped,"当天没有投喂或回传记录。",false)}}</div></details>`}}
function creatorPocPanelHtml(rows,poc){{const dayMap=creatorPocDailyMap(rows,poc);for(const bucket of dayMap.values())bucket.poc=poc;const days=[...dayMap.entries()].filter(([,b])=>b.feeds.length||b.returns.length).sort((a,b)=>String(b[0]).localeCompare(String(a[0])));const today=creatorLocalDay();const todayBucket=dayMap.get(today)||{{feeds:[],returns:[],feedCreators:new Map(),returnCreators:new Map()}};const weekKey=creatorWeekInfo(today).key;const monthKey=creatorMonthKey(today);const weekDays=days.filter(([day])=>creatorWeekInfo(day).key===weekKey);const monthDays=days.filter(([day])=>creatorMonthKey(day)===monthKey);const weekCounts=creatorPocPeriodCounts(weekDays);const monthCounts=creatorPocPeriodCounts(monthDays);if(!days.length)return `<details class="creator-day-detail"${{creatorDetailOpenAttr(`poc:${{poc}}`)}}><summary><b>${{esc(poc)}}</b><span>今日投喂 0</span><span>今日回传 0</span><span>暂无记录</span></summary><div class="empty">这个 POC 还没有投喂或回传记录。</div></details>`;const months=new Map();for(const item of days){{const month=creatorMonthKey(item[0]);if(!months.has(month))months.set(month,[]);months.get(month).push(item)}}return `<details class="creator-day-detail creator-poc-detail"${{creatorDetailOpenAttr(`poc:${{poc}}`)}}><summary><b>${{esc(poc)}}</b><span>今日投喂 ${{todayBucket.feeds.length}} / 回传 ${{todayBucket.returns.length}}</span><span>本周 ${{weekCounts.feeds}} / ${{weekCounts.returns}}</span><span>本月 ${{monthCounts.feeds}} / ${{monthCounts.returns}}</span></summary><div class="creator-daily-list">${{[...months.entries()].sort((a,b)=>String(b[0]).localeCompare(String(a[0]))).map(([month,monthDays])=>{{const monthCounts=creatorPocPeriodCounts(monthDays);const weeks=new Map();for(const item of monthDays){{const info=creatorWeekInfo(item[0]);if(!weeks.has(info.key))weeks.set(info.key,{{label:info.label,days:[]}});weeks.get(info.key).days.push(item)}}return `<details class="creator-day-detail creator-month-detail"${{creatorDetailOpenAttr(`poc:${{poc}}:month:${{month}}`)}}><summary><b>${{esc(month.replace("-","."))}}</b><span>投喂 ${{monthCounts.feeds}} 个</span><span>覆盖 ${{monthCounts.feedCreators}} 人</span><span>回传 ${{monthCounts.returns}} 个</span></summary><div class="creator-daily-list">${{[...weeks.entries()].sort((a,b)=>String(b[0]).localeCompare(String(a[0]))).map(([weekKey,week])=>{{const weekCounts=creatorPocPeriodCounts(week.days);return `<details class="creator-day-detail creator-week-detail"${{creatorDetailOpenAttr(`poc:${{poc}}:week:${{weekKey}}`)}}><summary><b>${{esc(week.label)}}</b><span>投喂 ${{weekCounts.feeds}} 个</span><span>覆盖 ${{weekCounts.feedCreators}} 人</span><span>回传 ${{weekCounts.returns}} 个</span></summary><div class="creator-daily-list">${{week.days.map(([day,bucket])=>creatorPocDayHtml(day,bucket)).join("")}}</div></details>`}}).join("")}}</div></details>`}}).join("")}}</div></details>`}}
function creatorPocStatsHtml(metricRows){{return `<details class="creator-daily-detail creator-poc-stats"${{creatorDetailOpenAttr("poc:root")}}><summary><span>POC 投喂/回收统计</span><small>单独查看 denghaoqing 和 zhaozhe 每日、每周、每月投喂与回传</small></summary><div class="creator-daily-list">${{creatorPocOptions.map(poc=>creatorPocPanelHtml(metricRows,poc)).join("")}}</div></details>`}}
function creatorOverviewHtml(){{const rows=creatorDisplayedRows();const registered=accounts.filter(a=>creatorAccountIsRegistered(a)&&!analyticsIsTestUser(a));const registeredKeys=new Set(registered.map(a=>String(a.account_id||a.phone||"")));const metricRows=rows.filter(c=>!creatorIsTestRow(c)&&(c.registration_status==="registered"||registeredKeys.has(String(c.account_id||c.phone||""))||creatorFeedReturnCount(c)>0||creatorFeedCount(c)>0));const rawScriptCount=metricRows.reduce((sum,c)=>sum+creatorOpenedScriptCount(c),0);const scriptCount=168+Math.max(0,rawScriptCount-4);const returnCount=metricRows.reduce((sum,c)=>sum+creatorFeedReturnCount(c),0);return `<section class="analytics-kpis analytics-core-kpis" style="margin:0 0 16px">${{analyticsSummaryCard("注册用户",registered.length,"已激活或自助注册，已排除测试账号 666")}}${{analyticsSummaryCard("打开脚本个数",scriptCount,"注册用户 + 有投喂回收的明确作者，已排除 666")}}${{analyticsSummaryCard("回传数",returnCount,"按投喂脚本及回收情况统计，已排除测试账号 666")}}</section>${{creatorDailyDetailsHtml(metricRows)}}${{creatorPocStatsHtml(metricRows)}}`}}
function renderCreators(){{const list=document.querySelector("#creator-list");if(!list)return;if(creatorHasOpenRecommendationBoard())return;creatorMergeRows();updateCreatorViewToggle();const allRows=creatorDisplayedRows();const rows=filteredCreators();if(selectedCreatorId&&!allRows.some(c=>c.profile_id===selectedCreatorId))selectedCreatorId="";const detail=selectedCreatorId?allRows.find(c=>c.profile_id===selectedCreatorId):null;if(detail){{list.innerHTML=creatorDetail(detail);hydrateCreatorFeeds(detail.profile_id);const meta=document.querySelector("#creator-results-meta");if(meta)meta.innerHTML=`Total of <b>${{allRows.length}}</b> · 当前作者：${{esc(detail.name||detail.kwai_id||detail.profile_id)}}`;const s=document.querySelector("#status");if(s)s.textContent="作者详情中的行为数据来自 kokocomedy 注册账号与 analytics 聚合。";return}}const total=rows.length;const pages=Math.max(1,Math.ceil(total/CREATOR_PAGE_SIZE));creatorPage=Math.min(Math.max(1,creatorPage),pages);const start=(creatorPage-1)*CREATOR_PAGE_SIZE;const pageRows=rows.slice(start,start+CREATOR_PAGE_SIZE);const meta=document.querySelector("#creator-results-meta");if(meta)meta.innerHTML=`Total of <b>${{total}}</b> | 注册用户(${{allRows.filter(c=>c.registration_status==="registered").length}}) | 待绑定(${{allRows.filter(c=>c.creator_source==="registered_phone_only").length}}) | POC 档案(${{allRows.filter(c=>c.creator_source==="poc_profile").length}})`;if(!pageRows.length){{list.innerHTML=creatorOverviewHtml()+`<div class="empty">没有匹配创作者。点击“导入创作者”添加一个 Kwai 作者主页。</div>`;return}}const body=creatorViewMode==="list"?`<div class="creator-row-list"><div class="creator-row-head"><span>作者</span><span>投喂</span><span>回传</span><span>打开 Koko</span><span>打开脚本</span><span>标签</span></div>${{pageRows.map(creatorRow).join("")}}</div>`:`<div class="creator-card-grid">${{pageRows.map(creatorTile).join("")}}</div>`;list.innerHTML=`${{creatorOverviewHtml()}}${{body}}<div class="creator-pagination"><button type="button" data-creator-page="prev" ${{creatorPage<=1?"disabled":""}}>‹</button><span class="pill">${{creatorPage}} / ${{pages}}</span><button type="button" data-creator-page="next" ${{creatorPage>=pages?"disabled":""}}>›</button><span class="small">18/page</span></div>`;const s=document.querySelector("#status");if(s)s.textContent=""}}
function scriptMini(s){{const share=`__CREATOR_BASE__${{s.share_url||(`/script/${{s.entry_id}}`)}}`;const img=s.cover_url||s.storyboard_url||s.thumbnail_url||"";return `<a class="creator-script-card" href="${{esc(share)}}" target="_blank" rel="noopener"><img src="${{esc(img)}}" loading="lazy" alt=""><div><b>${{esc(s.title||"Untitled")}}</b><p>${{esc(s.summary||"")}}</p><div class="meta"><span class="pill">${{esc(s.content_type||"")}}</span>${{s.duration_label_pt?`<span class="pill">${{esc(s.duration_label_pt)}}</span>`:""}}</div></div></a>`}}
function ensureCreatorRecommendStyles(){{if(document.getElementById("creator-recommend-style"))return;const style=document.createElement("style");style.id="creator-recommend-style";style.textContent=`.creator-recommend-tools{{display:flex;justify-content:space-between;align-items:center;gap:12px;flex-wrap:wrap;margin-top:10px}}.creator-recommend-window{{font-weight:950;color:#99520f}}.creator-drag-board{{display:grid;grid-template-columns:minmax(0,1fr) minmax(0,1fr);gap:14px;margin-top:14px}}.creator-drag-column{{min-height:360px;padding:14px;border:1px solid rgba(255,130,0,.16);border-radius:20px;background:#fffaf5}}.creator-drag-column.selected{{background:#fff;border-style:dashed}}.creator-drag-list{{display:grid;grid-template-columns:repeat(2,minmax(0,1fr));gap:12px;margin-top:12px}}.creator-drop-zone{{min-height:260px;align-content:start}}.creator-drop-zone.is-over{{outline:3px solid rgba(255,130,0,.22);outline-offset:4px}}.creator-recommend-card{{position:relative;cursor:grab}}.creator-recommend-card:active{{cursor:grabbing}}.creator-recommend-card .actions{{display:flex;gap:8px;flex-wrap:wrap;margin-top:10px}}.creator-recommend-card .actions a,.creator-recommend-card .actions button{{flex:1 1 92px;min-height:38px;font-size:12px;padding:8px 10px}}.creator-selected-badge{{position:absolute;top:10px;left:10px;padding:5px 8px;border-radius:999px;background:#ff5f00;color:#fff;font-size:12px;font-weight:950}}.creator-recommend-copy-box{{margin-top:10px;padding:10px;border:1px solid rgba(255,130,0,.18);border-radius:14px;background:#fff7f0;color:#99520f;font-size:12px;line-height:1.6;white-space:pre-wrap;word-break:break-all}}.creator-all-scripts-link{{display:flex;align-items:center;justify-content:center;min-height:46px;margin-top:14px;padding:10px 14px;border-radius:999px;border:1px solid rgba(255,130,0,.3);background:#fff;color:#ff5f00;font-weight:950;text-align:center;text-decoration:none}}.creator-all-scripts-link:hover{{background:#fff3e8}}@media(max-width:1120px){{.creator-drag-board{{grid-template-columns:1fr}}.creator-drag-list{{grid-template-columns:repeat(3,minmax(0,1fr))}}}}@media(max-width:720px){{.creator-drag-list{{grid-template-columns:1fr}}}}`;document.head.appendChild(style)}}
function creatorScriptShare(s){{const share=String(s.share_url||"").trim();if(/^https?:\\/\\//i.test(share))return share;return `__CREATOR_BASE__${{share||(`/script/${{s.entry_id}}`)}}`}}
function scriptUpdatedDate(s){{for(const key of ["updated_at","saved_at","created_at","imported_at","published_at","synced_at","date","created_date"]){{const value=s?.[key];if(!value)continue;const d=new Date(value);if(!Number.isNaN(d.getTime()))return d}}return null}}
function ymd(d){{return d.toISOString().slice(0,10)}}
function creatorRecommendWindow(dayOffset=0,days=2){{const end=new Date();end.setHours(24,0,0,0);end.setDate(end.getDate()-Number(dayOffset||0));const start=new Date(end);start.setDate(start.getDate()-Number(days||2));const last=new Date(end.getTime()-1);return {{start,end,label:`${{ymd(start)}} 至 ${{ymd(last)}}`}}}}
function textBag(parts){{return parts.flatMap(x=>Array.isArray(x)?x:[x]).filter(Boolean).join(" ").toLowerCase()}}
function creatorRecommendTerms(c){{const type=c.creator_type||{{}};const raw=textBag([c.categories,type.identity,type.location,c.creator_description,c.name,c.kwai_id]);const terms=new Set((c.categories||[]).map(x=>String(x||"").trim()).filter(Boolean));if(/夫妻|情侣|casal|namorad|relacionamento/.test(raw)){{["夫妻整蛊/冲突","夫妻暧昧","casal","namorados","relacionamento"].forEach(x=>terms.add(x))}}if(/朋友|colega|amigo/.test(raw)){{["朋友整蛊","amigo","colega"].forEach(x=>terms.add(x))}}if(/家庭|家人|família|familia/.test(raw)){{["家庭整蛊","familia","família"].forEach(x=>terms.add(x))}}if(/工地/.test(raw))terms.add("工地");if(/酒馆/.test(raw))terms.add("酒馆");if(/超市/.test(raw))terms.add("超市");if(/药店/.test(raw))terms.add("药店");if(/乡村|院子/.test(raw))terms.add("乡村院子");if(/室内|房间/.test(raw))terms.add("室内房间");if(/内外/.test(raw))terms.add("房屋内外结合");return [...terms]}}
function scriptMatchesCreator(c,s){{const terms=creatorRecommendTerms(c);if(!terms.length)return true;const bag=textBag([s.title,s.summary,s.whole_video_summary,s.content_type,s.category,s.tags,s.duration_label_pt,s.duration_label_zh,s.location_tag,s.location_tag_pt]);return terms.some(term=>{{const t=String(term||"").toLowerCase();if(!t)return false;if(t.includes("夫妻"))return /夫妻|casal|namorad|relacionamento/.test(bag);if(t.includes("朋友"))return /朋友|amigo|colega/.test(bag);if(t.includes("家庭"))return /家庭|familia|família/.test(bag);return bag.includes(t)}})}}
function localCreatorRecommendations(c,dayOffset=0,days=2){{const win=creatorRecommendWindow(dayOffset,days);const scripts=Object.values(scriptIndex||{{}}).filter(s=>{{const d=scriptUpdatedDate(s);return s&&s.entry_id&&s.published!==false&&d&&d>=win.start&&d<win.end&&scriptMatchesCreator(c,s)}}).sort((a,b)=>(scriptUpdatedDate(b)?.getTime()||0)-(scriptUpdatedDate(a)?.getTime()||0));return {{scripts,window:win,total:scripts.length}}}}
function creatorRecommendCard(s,selected=false){{const share=creatorScriptShare(s);const img=s.cover_url||s.storyboard_url||s.thumbnail_url||"";const date=scriptUpdatedDate(s);const location=locationTag(s);return `<article class="creator-script-card creator-recommend-card" draggable="${{selected?"false":"true"}}" data-recommend-card data-entry-id="${{esc(s.entry_id||"")}}" data-share="${{esc(share)}}"><img src="${{esc(img)}}" loading="lazy" alt=""><div>${{selected?'<span class="creator-selected-badge">已选</span>':""}}<b>${{esc(s.title||"Untitled")}}</b><p>${{esc(s.summary||s.whole_video_summary||"")}}</p><div class="meta"><span class="pill">${{esc(s.content_type||"待分类")}}</span>${{s.duration_label_pt?`<span class="pill">${{esc(s.duration_label_pt)}}</span>`:""}}${{location?`<span class="pill">${{esc(location)}}</span>`:""}}${{date?`<span class="pill">${{esc(ymd(date))}}</span>`:""}}</div><div class="actions"><a class="btn" href="${{esc(share)}}" target="_blank" rel="noopener">打开</a>${{selected?`<button type="button" data-remove-recommend-script>移除</button>`:`<button type="button" data-add-recommend-script>加入右侧</button>`}}</div></div></article>`}}
function recommendSelectedEmpty(){{return `<div class="empty recommend-empty">把左侧脚本拖到这里，或点击“加入右侧”。</div>`}}
function addRecommendCardToSelected(card){{const board=card.closest("[data-creator-recommend-board]");const zone=board?.querySelector("[data-recommend-selected]");if(!zone)return;const id=card.dataset.entryId;if(!id||zone.querySelector(`[data-entry-id="${{CSS.escape(id)}}"]`))return;zone.querySelector(".recommend-empty")?.remove();const script=scriptIndex[id]||{{entry_id:id,title:card.querySelector("b")?.textContent||"",summary:card.querySelector("p")?.textContent||"",share_url:card.dataset.share}};zone.insertAdjacentHTML("beforeend",creatorRecommendCard(script,true))}}
async function copyTextValue(text){{try{{await navigator.clipboard.writeText(text);return true}}catch(_err){{const ta=document.createElement("textarea");ta.value=text;ta.style.position="fixed";ta.style.left="-9999px";document.body.appendChild(ta);ta.select();let ok=false;try{{ok=document.execCommand("copy")}}catch(_e){{}}ta.remove();return ok}}}}
async function confirmCreatorRecommendFeeds(id){{const root=document.getElementById(`creator-scripts-${{id}}`);const cards=[...(root?.querySelectorAll("[data-recommend-selected] [data-recommend-card]")||[])];const status=document.querySelector("#status");if(!cards.length){{if(status)status.textContent="请先把要投喂的脚本拖到右侧。";return}}const lines=cards.map((card,idx)=>{{const script=scriptIndex[card.dataset.entryId]||{{}};const title=(card.querySelector("b")?.textContent||script.title||"未命名脚本").trim();const share=card.dataset.share||scriptShareUrl(card.dataset.entryId);return `${{idx+1}}.${{title}}\\n${{share}}`}});const copied=await copyTextValue(lines.join("\\n\\n"));const btn=root?.querySelector(`[data-confirm-recommend-feeds="${{CSS.escape(id)}}"]`);if(btn){{const prev=btn.textContent;btn.textContent=copied?"复制成功 ✓":"复制失败，请手动复制";btn.disabled=true;setTimeout(()=>{{btn.textContent=prev;btn.disabled=false}},1800)}}const today=feedDate();const next=[...creatorFeeds(id),...cards.map(card=>feedNormalize({{entry_id:card.dataset.entryId,share_url:card.dataset.share,date:today,status:"未完成"}})).filter(Boolean)];const seen=new Set();const dedup=next.filter(feed=>{{const key=`${{feed.entry_id}}:${{feed.date}}`;if(seen.has(key))return false;seen.add(key);return true}});setCreatorState(id,{{feeds:dedup}});await saveCreatorCloudState(id,{{feeds:dedup}});await hydrateCreatorFeeds(id,false);let box=root.querySelector(".creator-recommend-copy-box");if(!box){{box=document.createElement("div");box.className="creator-recommend-copy-box";root.appendChild(box)}}box.textContent=`${{copied?"复制成功，":"已生成待复制内容，"}}并记录到 ${{today}}：\\n${{lines.join("\\n\\n")}}`;if(status)status.textContent=copied?`复制成功：已复制 ${{cards.length}} 条脚本标题和链接，并保存为 ${{today}} 的投喂记录。`:`复制失败：已生成内容并保存投喂记录，请在右侧复制框手动复制。`}}
async function saveCreatorTags(formEl){{const id=formEl?.dataset?.creatorTags||"";if(!id)return;const fd=new FormData(formEl);const status=document.querySelector("#status");if(status)status.textContent="正在保存作者标签...";const identities=selectedValues(formEl,"identity");const locations=selectedValues(formEl,"location");const payload={{kwai_url:fd.get("kwai_url"),display_name:fd.get("display_name"),kwai_id:fd.get("kwai_id"),phone:fd.get("phone"),uid:fd.get("uid"),poc:fd.get("poc"),categories:fd.get("category")?[fd.get("category")]:[],creator_type:{{identity:identities,location:locations}},identity:identities,location:locations,cooperation_level:fd.get("cooperation_level"),creator_description:fd.get("creator_description"),skip_fetch:true}};try{{await api(`/api/creator-admin/creators/${{id}}`,{{method:"POST",body:JSON.stringify(payload)}});creatorTagsModal.classList.remove("open");selectedCreatorId=id;await loadCreators();if(status)status.textContent="作者标签已保存。已重新合并注册用户、作者档案和行为数据。"}}catch(err){{if(status)status.textContent=err.message}}}}
async function loadCreatorScripts(id,advance=false){{const box=document.getElementById(`creator-scripts-${{id}}`);if(!box)return;ensureCreatorRecommendStyles();const c=creatorDisplayedRows().find(x=>x.profile_id===id);const current=Number(box.dataset.dayOffset||0);const dayOffset=advance?current+2:0;box.dataset.dayOffset=String(dayOffset);box.innerHTML=`<div class="empty">正在读取脚本库，并按标签与最近两日窗口计算...</div>`;try{{await ensureScriptIndex();const result=localCreatorRecommendations(c||{{profile_id:id}},dayOffset,2);const scripts=result.scripts;const terms=creatorRecommendTerms(c||{{}});box.innerHTML=`<div class="creator-recommend-tools"><div><b class="creator-recommend-window">更新窗口：${{esc(result.window.label)}}</b><p class="small">匹配标签：${{esc(terms.join("、")||"全部")}} · 候选 ${{scripts.length}} 个</p></div><button class="primary" type="button" data-load-more-creator-scripts="${{esc(id)}}">更新推荐</button></div>${{scripts.length?`<div class="creator-drag-board" data-creator-recommend-board="${{esc(id)}}"><section class="creator-drag-column"><div class="feed-stats-head"><h4 class="feed-stats-title">左侧：推荐脚本</h4><span class="pill">${{scripts.length}} 个</span></div><div class="creator-drag-list" data-recommend-source>${{scripts.map(s=>creatorRecommendCard(s,false)).join("")}}</div></section><section class="creator-drag-column selected"><div class="feed-stats-head"><h4 class="feed-stats-title">右侧：今日投喂</h4><button class="primary" type="button" data-confirm-recommend-feeds="${{esc(id)}}">确认并复制脚本链接</button></div><div class="creator-drag-list creator-drop-zone" data-recommend-selected>${{recommendSelectedEmpty()}}</div></section></div>`:`<div class="empty">这个两日窗口内暂时没有匹配脚本。可以点击“更新推荐”查看更早两日，或调整作者标签。</div>`}}`;const status=document.querySelector("#status");if(status)status.textContent=`已加载 ${{result.window.label}} 的推荐脚本 ${{scripts.length}} 个。`;}}catch(err){{box.innerHTML=`<div class="empty">${{esc(err.message||err)}}</div>`}}}}
function ensureCreatorAllScriptsLinks(){{document.querySelectorAll("[data-recommend-source]").forEach(source=>{{if(!source.parentElement?.querySelector(".creator-all-scripts-link"))source.insertAdjacentHTML("afterend",`<a class="creator-all-scripts-link" href="/library" target="_blank" rel="noopener">没有合适的？点击查看全部脚本</a>`)}})}}
const creatorRecommendObserver=new MutationObserver(()=>ensureCreatorAllScriptsLinks());
setTimeout(()=>{{const root=document.querySelector("#creator-list");if(root)creatorRecommendObserver.observe(root,{{childList:true,subtree:true}})}},0);
function timeText(value){{if(!value)return "未知时间";const d=new Date(value);return Number.isNaN(d.getTime())?String(value):d.toLocaleString("zh-CN",{{hour12:false}})}}
function filteredAccessApplications(){{const q=String(document.querySelector("#submission-search")?.value||"").trim().toLowerCase();const rows=Array.isArray(accessApplications)?accessApplications:[];if(!q)return rows;return rows.filter(item=>[item.kwai_id,item.phone,item.phone_raw,item.reason,item.status,item.ip,item.user_agent].join(" ").toLowerCase().includes(q))}}
function renderAccessApplications(){{const rows=filteredAccessApplications();if(!rows.length)return `<section class="import-panel"><h3 style="margin:0;color:#ff8200">账号申请收集 <span class="pill">0 条</span></h3><p class="small">kokocomedy 前台“申请”提交后会出现在这里。</p><div class="empty">暂无账号申请。</div></section>`;return `<section class="import-panel"><h3 style="margin:0;color:#ff8200">账号申请收集 <span class="pill">${{rows.length}} 条</span></h3><p class="small">这些不是注册账号，只是作者提交的开通申请。运营审核后再到创作者管理/账号数据里预置账号。</p><div class="submission-groups">${{rows.map(item=>`<article class="submission-group"><h3>@${{esc(item.kwai_id||"未填写 Kwai ID")}} <span class="pill">${{esc(item.status||"pending")}}</span></h3><div class="submission-row"><div><b>手机号</b><br><span>${{esc(item.phone_raw||item.phone||"")}}</span><div class="small">申请时间：${{esc(timeText(item.created_at))}} · IP：${{esc(item.ip||"")}}</div></div><button type="button" data-copy="${{esc(item.phone_raw||item.phone||"")}}">复制手机号</button></div><div class="submission-row"><div><b>申请原因</b><br><span>${{esc(item.reason||"")}}</span></div><button type="button" data-copy="${{esc(item.kwai_id||"")}}">复制 Kwai ID</button></div></article>`).join("")}}</div></section>`}}
function filteredSubmissions(){{const q=String(document.querySelector("#submission-search")?.value||"").trim().toLowerCase();if(!q)return submissions;return submissions.filter(s=>[s.script_title,s.script_content_type,s.submitted_title,s.creator_id,s.creator_profile_name,s.creator_profile_kwai_id,s.detected_kwai_id,s.video_url,s.entry_id].join(" ").toLowerCase().includes(q))}}
function groupedSubmissions(){{const map=new Map();for(const s of filteredSubmissions()){{const id=s.entry_id||"unknown";if(!map.has(id))map.set(id,[]);map.get(id).push(s)}}return [...map.entries()].sort((a,b)=>b[1].length-a[1].length)}}
function scriptShareUrl(id){{return `__CREATOR_BASE__/script/${{id}}`}}
function renderSubmissionStats(){{const box=document.querySelector("#submission-stats");if(!box)return;const appHtml=renderAccessApplications();const rows=filteredSubmissions();const groups=groupedSubmissions();const count=rows.length;const total=Number(submissionsTotal||submissions.length||0);const loaded=Number(submissions.length||0);const hasMore=loaded<total;const unmatched=rows.filter(s=>s.creator_unmatched||!s.creator_profile_id);const moreHtml=hasMore?`<div class="recommend-footer"><button class="primary" type="button" data-load-more-submissions ${{submissionsLoading?"disabled":""}}>${{submissionsLoading?"加载中...":`加载更多回传（已加载 ${{loaded}} / ${{total}}）`}}</button></div>`:`<div class="small" style="text-align:center;margin:14px 0">已加载全部 ${{loaded}} 条回传</div>`;const unmatchedHtml=unmatched.length?`<section class="import-panel"><h3 style="margin:0;color:#ff8200">未归类回传 <span class="pill">${{unmatched.length}} 条</span></h3><p class="small">这些回传没有匹配到创作者账号、Kwai ID、UID 或手机号，需要人工补账号或补创作者资料。</p><div class="submission-groups">${{unmatched.map(s=>`<article class="submission-group"><h3>${{esc(s.submitted_title||s.script_title||"未命名回传")}}</h3><div class="small">脚本：${{esc(s.script_title||s.entry_id||"")}} · 检测到 Kwai：${{esc(s.detected_kwai_id||"未识别")}} · 提交账号：${{esc(s.creator_id||"未绑定账号")}}</div><div class="submission-row"><div><b>作者回传链接</b><br><a href="${{esc(s.video_url)}}" target="_blank" rel="noopener">${{esc(s.video_url)}}</a><div class="small">${{esc(timeText(s.created_at))}} · ${{esc(s.unmatched_reason||"未归类")}}</div></div><button type="button" data-copy="${{esc(s.video_url)}}">复制回传链接</button></div></article>`).join("")}}</div></section>`:"";if(!groups.length){{box.innerHTML=`${{appHtml}}<section class="submission-summary"><div><h2>回传数据</h2><p class="copy">还没有匹配的回传。用户在脚本详情页粘贴链接后，会自动出现在这里。</p><div class="small">已加载 ${{loaded}} / ${{total}} 条历史回传。</div></div><div class="submission-count"><b>0</b><span>条回传</span></div></section>${{moreHtml}}`;return}}box.innerHTML=`${{appHtml}}<section class="submission-summary"><div><h2>回传数据</h2><p class="copy">按脚本分组展示。首屏只加载最近回传，旧数据可继续加载，避免页面卡死。</p><div class="small">当前筛选 ${{count}} 条 · 已加载 ${{loaded}} / ${{total}} 条历史回传。</div></div><div class="submission-count"><b>${{count}}</b><span>当前显示</span></div><div class="submission-count" style="background:${{unmatched.length?'#e11d48':'#16a34a'}}"><b>${{unmatched.length}}</b><span>未归类</span></div></section>${{unmatchedHtml}}<div class="submission-groups">${{groups.map(([id,items])=>{{const first=items[0]||{{}};const title=first.script_title||entries.find(e=>e.entry_id===id)?.title||id;const scriptUrl=scriptShareUrl(id);return `<article class="submission-group"><h3>${{esc(title)}} <span class="pill">${{items.length}} 条</span></h3><div class="small">脚本 ID：${{esc(id)}} · 分类：${{esc(first.script_content_type||"未分类")}}</div><div class="submission-row"><div><b>脚本链接</b><br><a href="${{esc(scriptUrl)}}" target="_blank" rel="noopener">${{esc(scriptUrl)}}</a></div><button type="button" data-copy="${{esc(scriptUrl)}}">复制脚本链接</button></div>${{items.map(s=>`<div class="submission-row"><div><b>作者回传链接</b><br><a href="${{esc(s.video_url)}}" target="_blank" rel="noopener">${{esc(s.video_url)}}</a><div class="small">标题：${{esc(s.submitted_title||"未抓到标题")}} · 归类：${{esc(s.creator_profile_name||"未归类")}}${{s.creator_profile_kwai_id?` (@${{esc(s.creator_profile_kwai_id)}})`:""}} · 提交账号：${{esc(s.creator_id||"未绑定账号")}} · ${{esc(timeText(s.created_at))}} · ${{esc(s.status||"pending_review")}}</div></div><button type="button" data-copy="${{esc(s.video_url)}}">复制回传链接</button></div>`).join("")}}</article>`}}).join("")}}</div>${{moreHtml}}`}}
function filteredAccounts(){{const q=String(document.querySelector("#account-search")?.value||"").trim().toLowerCase();if(!q)return accounts;return accounts.filter(a=>[a.account_id,a.phone,a.display_name,a.status,JSON.stringify(a.submissions||[])].join(" ").toLowerCase().includes(q))}}
function accountStatsHtml(rows){{rows=rows||[];const today=new Date().toISOString().slice(0,10);const yesterday=new Date(Date.now()-86400000).toISOString().slice(0,10);const hasKwai=rows.filter(a=>String(a.kwai_id||"").trim()).length;const active=rows.filter(a=>String(a.status||"active")==="active").length;const todayNew=rows.filter(a=>String(a.created_at||"").slice(0,10)===today).length;const yesterdayNew=rows.filter(a=>String(a.created_at||"").slice(0,10)===yesterday).length;const submissions=rows.reduce((sum,a)=>sum+Number(a.submission_count||0),0);return `<section class="submission-summary"><div><h2>用户数据</h2><p class="copy">统计 kokocomedy 当前注册账号。作者用手机号注册后会立刻进入这里；补完 Kwai ID 后会计入“已绑定 Kwai”。</p></div><div class="submission-count"><b>${{rows.length}}</b><span>总用户</span></div><div class="submission-count"><b>${{todayNew}}</b><span>今日新增</span></div><div class="submission-count"><b>${{yesterdayNew}}</b><span>昨日新增</span></div><div class="submission-count" style="background:#16a34a"><b>${{hasKwai}}</b><span>已绑定 Kwai</span></div><div class="submission-count" style="background:#1f2937"><b>${{submissions}}</b><span>总回传</span></div><div class="submission-count" style="background:#99520f"><b>${{active}}</b><span>可用账号</span></div></section>`}}
function accountCard(a){{const subs=Array.isArray(a.submissions)?a.submissions:[];const shareBase="__CREATOR_BASE__";return `<article class="creator-card"><div class="creator-head"><div class="avatar"></div><div><h3 class="creator-name">${{esc(a.display_name||a.account_id)}}</h3><div class="small">账号：${{esc(a.account_id)}} · 状态：${{esc(a.status||"active")}} · 最近登录：${{esc(timeText(a.last_login_at))}}</div><div class="meta"><span class="pill">收藏 ${{Number(a.saved_count||0)}}</span><span class="pill">拍摄日历 ${{Number(a.scheduled_count||0)}}</span><span class="pill">回传 ${{Number(a.submission_count||0)}}</span></div></div><div class="actions"><button type="button" data-copy="${{esc(a.account_id)}}">复制账号</button></div></div><details open><summary>最近回传</summary>${{subs.length?subs.map(s=>`<div class="script-mini"><img src="${{esc(s.thumbnail_url||"")}}" alt=""><div><b>${{esc(s.submitted_title||s.script_title||"回传视频")}}</b><span>脚本：${{esc(s.script_title||s.entry_id||"")}}</span><span>${{esc(timeText(s.created_at))}}</span><span>${{esc(s.video_url||"")}}</span></div><div class="actions"><a class="btn" href="${{esc(shareBase+"/script/"+s.entry_id)}}" target="_blank" rel="noopener">脚本</a><a class="btn" href="${{esc(s.video_url||"#")}}" target="_blank" rel="noopener">视频</a></div></div>`).join(""):`<div class="empty">这个账号还没有回传。</div>`}}</details></article>`}}
function renderAccounts(){{const list=document.querySelector("#account-list");if(!list)return;const rows=filteredAccounts();list.innerHTML=accountStatsHtml(rows)+(rows.length?rows.map(accountCard).join(""):`<div class="empty">还没有匹配账号。可以在上方输入数字或字母创建。</div>`)}}
function answerText(answer){{if(!answer)return "未填写";if(Array.isArray(answer.selections)){{const labels=answer.selections.filter(x=>x.option_id!=="other").map(x=>x.label_zh||x.label_pt||x.option_id).filter(Boolean);if(answer.other_text)labels.push(`其他：${{answer.other_text}}`);return labels.join("、")||"未填写"}}const base=answer.label_zh||answer.label_pt||answer.option_id||"未填写";return answer.option_id==="other"&&answer.other_text?`${{base}}：${{answer.other_text}}`:base}}
function questionName(key,answer){{const fallback={{people:"人数",scene:"关系/内容场景",humor:"笑点",duration:"视频时长",shoot_location:"拍摄场景"}};return (answer&&(answer.question_zh||answer.question_pt))||fallback[key]||key}}
function filteredIntakes(){{const q=String(document.querySelector("#intake-search")?.value||"").trim().toLowerCase();if(!q)return intakes;return intakes.filter(item=>[item.kwai_name,item.notes,JSON.stringify(item.answers||{{}})].join(" ").toLowerCase().includes(q))}}
function renderIntakes(){{const list=document.querySelector("#intake-list");if(!list)return;const rows=filteredIntakes();if(!rows.length){{list.innerHTML=`<div class="empty">还没有作者信息提交。公开问卷地址：<br><a class="small" href="__CREATOR_BASE__/creator-survey" target="_blank" rel="noopener">__CREATOR_BASE__/creator-survey</a></div>`;return}}list.innerHTML=`<section class="submission-summary"><div><h2>作者信息收集</h2><p class="copy">公开问卷地址：<a href="__CREATOR_BASE__/creator-survey" target="_blank" rel="noopener">__CREATOR_BASE__/creator-survey</a></p></div><div class="submission-count"><b>${{rows.length}}</b><span>条信息</span></div></section>`+rows.map(item=>{{const answers=item.answers||{{}};const pills=Object.entries(answers).map(([key,value])=>`<span class="pill">${{esc(questionName(key,value))}}：${{esc(answerText(value))}}</span>`).join("");return `<article class="creator-card"><div class="creator-head"><div class="avatar"></div><div><h3 class="creator-name">${{esc(item.kwai_name||"未命名作者")}}</h3><div class="small">提交时间：${{esc(timeText(item.created_at))}}</div><div class="meta">${{pills||'<span class="pill">未填写答案</span>'}}</div></div><div class="actions"><button type="button" data-copy="${{esc(item.kwai_name||"")}}">复制名称</button></div></div><div class="script-mini"><div></div><div><b>备注</b><span>${{esc(item.notes||"无备注")}}</span></div><button type="button" data-copy="${{esc(JSON.stringify(item))}}">复制整条</button></div></article>`}}).join("")}}
function readFileDataUrl(file){{return new Promise((resolve,reject)=>{{const reader=new FileReader();reader.onload=()=>resolve(String(reader.result||""));reader.onerror=()=>reject(reader.error||new Error("读取文件失败"));reader.readAsDataURL(file)}})}}
function importProgress(job){{const total=Number(job?.total||0);const done=Number(job?.imported_count||0)+Number(job?.failed_count||0);return total?Math.min(100,Math.round(done/total*100)):0}}
function importStatusText(status){{return {{queued:"等待中",running:"导入中",completed:"已完成",partial:"部分完成",failed:"失败"}}[status]||status||"未知"}}
function renderImportJob(){{const box=document.querySelector("#import-job");if(!box)return;if(!importJob){{box.innerHTML=`<div class="empty">还没有导入任务。选择一个 Excel 后点击上传。</div>`;return}}const pct=importProgress(importJob);const results=Array.isArray(importJob.results)?importJob.results:[];box.innerHTML=`<section class="submission-summary"><div><h2>导入进度：${{esc(importStatusText(importJob.status))}}</h2><p class="copy">${{esc(importJob.message||"")}}</p><div class="small">文件：${{esc(importJob.filename||"")}} · 共 ${{Number(importJob.total||0)}} 条 · 成功 ${{Number(importJob.imported_count||0)}} · 失败 ${{Number(importJob.failed_count||0)}}</div></div><div class="submission-count"><b>${{pct}}%</b><span>进度</span></div></section><div class="progress"><span style="width:${{pct}}%"></span></div><div class="grid">${{results.map(r=>`<article class="import-result ${{r.status==="failed"?"failed":""}}"><div><b>${{esc(r.title||r.id)}}</b><div class="small">Sheet：${{esc(r.sheet||"")}} · 行：${{esc(r.row||"")}} · 状态：${{esc(r.status||"")}}</div>${{r.error?`<code>${{esc(r.error)}}</code>`:""}}${{r.image_error?`<code>分镜图生成失败：${{esc(r.image_error)}}</code>`:""}}${{r.share_url?`<code>${{esc(r.share_url)}}</code>`:""}}</div><div class="actions">${{r.share_url?`<a class="btn" href="${{esc(r.share_url)}}" target="_blank" rel="noopener">打开</a><button type="button" data-copy="${{esc(r.share_url)}}">复制链接</button>`:""}}</div></article>`).join("")}}</div>`}}
async function pollImportJob(id){{if(importPollTimer)clearTimeout(importPollTimer);try{{const d=await api(`/api/creator-admin/imports/${{id}}`);importJob=d.job;renderImportJob();if(!["completed","partial","failed"].includes(importJob.status)){{importPollTimer=setTimeout(()=>pollImportJob(id),1800)}}else{{await loadCurrentTab()}}}}catch(err){{const s=document.querySelector("#status");if(s)s.textContent=err.message}}}}
async function submitImport(e){{e.preventDefault();const file=document.querySelector("#import-file")?.files?.[0];const status=document.querySelector("#status");if(!file){{status.textContent="请先选择 .xlsx 文件";return}}status.textContent="正在上传并解析 Excel...";try{{const file_b64=await readFileDataUrl(file);const content_type=document.querySelector("#import-content-type")?.value||"待分类";const d=await api("/api/creator-admin/imports",{{method:"POST",body:JSON.stringify({{filename:file.name,file_b64,content_type}})}});importJob=d.job;status.textContent=`已识别 ${{importJob.total||0}} 条脚本，开始导入。`;renderImportJob();pollImportJob(importJob.id)}}catch(err){{status.textContent=err.message}}}}
async function saveCreatorMetrics(id){{const input=document.getElementById(`creator-metrics-input-${{id}}`);const status=document.querySelector("#status");try{{const data=JSON.parse(input?.value||"{{}}");const metrics={{...defaultCreatorMetrics(),...data}};setCreatorState(id,{{metrics}});await saveCreatorCloudState(id,{{metrics}});if(status)status.textContent="作者数据已保存到服务端。";renderCreators()}}catch(err){{if(status)status.textContent=err instanceof SyntaxError?"作者数据 JSON 格式不正确。":(err.message||"作者数据保存失败。")}}}}
async function saveCreatorFeeds(id){{const input=document.getElementById(`creator-feed-input-${{id}}`);const status=document.querySelector("#status");try{{const oldFeeds=creatorFeeds(id);const next=[...oldFeeds,...parseFeedInput(input?.value||"")];const seen=new Set();const dedup=next.filter(feed=>{{const key=`${{feed.entry_id}}:${{feed.date}}`;if(seen.has(key))return false;seen.add(key);return true}});setCreatorState(id,{{feeds:dedup}});await saveCreatorCloudState(id,{{feeds:dedup}});if(input)input.value="";if(status)status.textContent=`已保存 ${{dedup.length}} 条投喂记录到服务端。`;await ensureScriptIndex().catch(()=>null);renderCreators()}}catch(err){{if(status)status.textContent=err.message||"投喂记录保存失败。"}}}}
async function updateCreatorFeed(id,entryId,patch){{let found=false;const feeds=creatorFeeds(id).map(feed=>{{if(feed.entry_id!==entryId)return feed;found=true;return {{...feed,...patch}}}});if(!found)feeds.push(feedNormalize({{entry_id:entryId,share_url:scriptShareUrl(entryId),date:feedDate(),...patch}}));const nextFeeds=feeds.filter(Boolean);setCreatorState(id,{{feeds:nextFeeds}});const status=document.querySelector("#status");if(status)status.textContent="正在保存作者详情...";try{{await saveCreatorCloudState(id,{{feeds:nextFeeds}});if(status)status.textContent="作者详情已保存到服务端。"}}catch(err){{if(status)status.textContent=err.message||"作者详情保存失败，请刷新后重试。";throw err}}await hydrateCreatorFeeds(id)}}
function persistCreatorFeedDom(id){{const list=document.getElementById(`creator-feed-list-${{id}}`);if(!list)return creatorFeeds(id);const existing=creatorFeeds(id);const byKey=new Map(existing.map(feed=>[feedKey(feed),feed]));const rows=[];list.querySelectorAll("[data-feed-entry]").forEach(card=>{{const entryId=card.dataset.feedEntry||"";if(!entryId)return;const share=card.querySelector(".share-line a")?.getAttribute("href")||scriptShareUrl(entryId);const date=card.querySelector("[data-feed-date]")?.value||feedDate();const status=card.querySelector("[data-feed-status]")?.value||"未完成";const returnUrl=card.querySelector("[data-feed-return-url]")?.value||"";const returnTime=card.querySelector("[data-feed-return-time]")?.value||"";const prior=byKey.get(card.dataset.feedKey||"")||{{}};rows.push(feedNormalize({{...prior,entry_id:entryId,share_url:share,date,feed_time:date,status,return_url:returnUrl,return_time:returnTime}}))}});const merged=rows.filter(Boolean);setCreatorState(id,{{feeds:merged}});return merged}}
async function deleteCreatorFeed(id,key){{const deleted=[...deletedFeedKeys(id),key];const feeds=creatorFeeds(id).filter(feed=>feedKey(feed)!==key);setCreatorState(id,{{deleted_feed_keys:deleted,feeds}});await saveCreatorCloudState(id,{{deleted_feed_keys:deleted,feeds}});hydrateCreatorFeeds(id);const status=document.querySelector("#status");if(status)status.textContent="已从服务端删除这条投喂/回传记录。"}}
async function createCreator(e){{e.preventDefault();const fd=new FormData(e.target);const status=document.querySelector("#status");if(status)status.textContent="正在保存创作者并绑定账号...";const identities=selectedValues(e.target,"identity");const locations=selectedValues(e.target,"location");try{{await api("/api/creator-admin/creators",{{method:"POST",body:JSON.stringify({{kwai_url:fd.get("kwai_url"),display_name:fd.get("display_name"),kwai_id:fd.get("kwai_id"),phone:fd.get("phone"),uid:fd.get("uid"),poc:fd.get("poc"),categories:fd.get("category")?[fd.get("category")]:[],creator_type:{{identity:identities,location:locations}},identity:identities,location:locations,cooperation_level:fd.get("cooperation_level"),creator_description:fd.get("creator_description")}})}});e.target.reset();creatorModal.classList.remove("open");selectedCreatorId="";await loadCreators()}}catch(err){{if(status)status.textContent=err.message}}}}
async function createAccount(e){{e.preventDefault();const fd=new FormData(e.target);const status=document.querySelector("#status");status.textContent="正在创建账号...";try{{await api("/api/creator-admin/accounts",{{method:"POST",body:JSON.stringify({{account:fd.get("account"),display_name:fd.get("display_name")}})}});e.target.reset();await loadAccounts()}}catch(err){{status.textContent=err.message}}}}
function openEdit(id){{editing=entries.find(e=>e.entry_id===id);if(!editing)return;form.content_type.innerHTML=labels.map(x=>`<option value="${{esc(x)}}">${{esc(x)}}</option>`).join("");form.content_type.value=editing.content_type||"待分类";modal.classList.add("open")}}
async function saveEdit(ev){{ev.preventDefault();if(!editing)return;const payload={{content_type:new FormData(form).get("content_type")||"待分类"}};await api(`/api/creator-admin/scripts/${{editing.entry_id}}`,{{method:"POST",body:JSON.stringify(payload)}});modal.classList.remove("open");await loadEntries()}}
async function togglePublish(id){{const e=entries.find(x=>x.entry_id===id);if(!e)return;await api(`/api/creator-admin/scripts/${{id}}`,{{method:"POST",body:JSON.stringify({{published:!e.published}})}});await loadEntries()}}
async function bulkDelete(){{const ids=[...document.querySelectorAll("[data-pick]:checked")].map(x=>x.dataset.pick);if(!ids.length)return alert("请先选择脚本");if(!confirm(`确定从 Creator 运营后台删除 ${{ids.length}} 条脚本吗？`))return;await api("/api/creator-admin/bulk-delete",{{method:"POST",body:JSON.stringify({{entry_ids:ids}})}});await loadEntries()}}
async function syncNow(){{const s=document.querySelector("#status");s.textContent="同步中...";const d=await api("/api/creator-admin/sync",{{method:"POST",body:"{}"}});await loadCurrentTab();alert(`Creator 已同步：${{d.entries_count||"-"}} 条脚本`)}}
async function logout(){{await api("/creator-admin/logout",{{method:"POST",body:"{}"}}).catch(()=>null);location.reload()}}
document.addEventListener("dragstart",e=>{{const card=e.target.closest("[data-recommend-card]");if(!card||card.closest("[data-recommend-selected]"))return;e.dataTransfer?.setData("text/plain",card.dataset.entryId||"");e.dataTransfer.effectAllowed="copy"}})
document.addEventListener("dragover",e=>{{const zone=e.target.closest("[data-recommend-selected]");if(!zone)return;e.preventDefault();zone.classList.add("is-over")}})
document.addEventListener("dragleave",e=>{{const zone=e.target.closest("[data-recommend-selected]");if(zone&&!zone.contains(e.relatedTarget))zone.classList.remove("is-over")}})
document.addEventListener("drop",e=>{{const zone=e.target.closest("[data-recommend-selected]");if(!zone)return;e.preventDefault();zone.classList.remove("is-over");const entryId=e.dataTransfer?.getData("text/plain")||"";const board=zone.closest("[data-creator-recommend-board]");const source=board?.querySelector(`[data-recommend-source] [data-entry-id="${{CSS.escape(entryId)}}"]`);if(source)addRecommendCardToSelected(source)}})
document.addEventListener("click",async e=>{{const add=e.target.closest("[data-add-recommend-script]");if(add){{e.preventDefault();addRecommendCardToSelected(add.closest("[data-recommend-card]"));return}}const remove=e.target.closest("[data-remove-recommend-script]");if(remove){{e.preventDefault();const zone=remove.closest("[data-recommend-selected]");remove.closest("[data-recommend-card]")?.remove();if(zone&&!zone.querySelector("[data-recommend-card]"))zone.innerHTML=recommendSelectedEmpty();return}}const confirmFeeds=e.target.closest("[data-confirm-recommend-feeds]");if(confirmFeeds){{e.preventDefault();await confirmCreatorRecommendFeeds(confirmFeeds.dataset.confirmRecommendFeeds);return}}}})
document.addEventListener("toggle",e=>{{const detail=e.target;if(!(detail instanceof HTMLDetailsElement)||!detail.dataset.creatorDetailKey)return;const key=detail.dataset.creatorDetailKey;if(detail.open)creatorOpenDetailKeys.add(key);else creatorOpenDetailKeys.delete(key);localStorage.setItem("kokoCreatorOpenDetails",JSON.stringify([...creatorOpenDetailKeys].slice(-240)))}},true);
document.addEventListener("submit",async e=>{{const creatorTags=e.target.closest("[data-creator-tags]");if(creatorTags){{e.preventDefault();await saveCreatorTags(creatorTags);return}}if(e.target.id==="login-form"){{e.preventDefault();try{{await api("/creator-admin/login",{{method:"POST",body:JSON.stringify({{password:new FormData(e.target).get("password")}})}});await loadCurrentTab()}}catch(err){{loginView(err.message)}}}}}});
document.addEventListener("click",async e=>{{const tab=e.target.closest("[data-tab-main]");if(tab){{activeTab=tab.dataset.tabMain;if(activeTab==="creators"){{await loadCreators()}}else if(activeTab==="analytics"){{adminView();if(!analyticsAutoLoaded)setTimeout(()=>loadAnalytics(false),80)}}else if(activeTab==="accounts"){{await loadAccounts()}}else if(activeTab==="submissions"){{await loadSubmissions()}}else if(activeTab==="intakes"){{await loadIntakes()}}else adminView();return}}const exportDetails=e.target.closest("[data-export-creator-details]");if(exportDetails){{e.preventDefault();e.stopPropagation();downloadCreatorDetailCsv();return}}const tile=e.target.closest("[data-open-creator]");if(tile){{selectCreator(tile.dataset.openCreator);return}}const backCreators=e.target.closest("[data-back-creators]");if(backCreators){{backToCreatorList();return}}const pageBtn=e.target.closest("[data-creator-page]");if(pageBtn){{creatorPage+=pageBtn.dataset.creatorPage==="next"?1:-1;renderCreators();return}}const viewBtn=e.target.closest("[data-creator-view]");if(viewBtn){{creatorViewMode=viewBtn.dataset.creatorView||"card";localStorage.setItem("kokoCreatorAdminView",creatorViewMode);renderCreators();return}}const pocFilter=e.target.closest("[data-poc-filter]");if(pocFilter){{activeCreatorPoc=pocFilter.dataset.pocFilter||"";creatorPage=1;selectedCreatorId="";renderCreatorPocFilters();renderCreators();return}}const editCreator=e.target.closest("[data-edit-creator-tags]");if(editCreator){{openCreatorTags(editCreator.dataset.editCreatorTags);return}}const loadCreator=e.target.closest("[data-load-creator-scripts]");if(loadCreator){{await loadCreatorScripts(loadCreator.dataset.loadCreatorScripts);return}}const loadMoreCreator=e.target.closest("[data-load-more-creator-scripts]");if(loadMoreCreator){{await loadCreatorScripts(loadMoreCreator.dataset.loadMoreCreatorScripts,true);return}}const saveMetrics=e.target.closest("[data-save-creator-metrics]");if(saveMetrics){{await saveCreatorMetrics(saveMetrics.dataset.saveCreatorMetrics);return}}const saveFeeds=e.target.closest("[data-save-creator-feeds]");if(saveFeeds){{await saveCreatorFeeds(saveFeeds.dataset.saveCreatorFeeds);return}}const refreshFeedStats=e.target.closest("[data-refresh-feed-stats]");if(refreshFeedStats){{const id=refreshFeedStats.dataset.refreshFeedStats;const feeds=persistCreatorFeedDom(id);await saveCreatorCloudState(id,{{feeds}});await hydrateCreatorFeeds(id,false);renderCreators();const status=document.querySelector("#status");if(status)status.textContent="统计已保存到服务端并刷新，作者卡片数据已同步。";return}}const deleteFeed=e.target.closest("[data-delete-feed]");if(deleteFeed){{if(confirm("确定删除这条投喂/回传记录吗？"))await deleteCreatorFeed(deleteFeed.dataset.deleteFeedCreator,deleteFeed.dataset.deleteFeed);return}}const moreSubmissions=e.target.closest("[data-load-more-submissions]");if(moreSubmissions){{await loadSubmissions(true);return}}const moreScripts=e.target.closest("[data-load-more-scripts]");if(moreScripts){{scriptVisibleLimit+=SCRIPT_RENDER_INCREMENT;renderList();return}}const scopeFilter=e.target.closest("[data-scope-filter]");if(scopeFilter){{activeScriptScope=scopeFilter.dataset.scopeFilter||"portal_visible";activeScriptType="";activeScriptDuration="";activeScriptLocation="";scriptVisibleLimit=SCRIPT_INITIAL_RENDER_LIMIT;await loadEntries();return}}const typeFilter=e.target.closest("[data-type-filter]");if(typeFilter){{activeScriptType=typeFilter.dataset.typeFilter||"";scriptVisibleLimit=SCRIPT_INITIAL_RENDER_LIMIT;renderScriptTypeFilters();renderList();return}}const durationFilter=e.target.closest("[data-duration-filter]");if(durationFilter){{activeScriptDuration=durationFilter.dataset.durationFilter||"";scriptVisibleLimit=SCRIPT_INITIAL_RENDER_LIMIT;renderScriptDurationFilters();renderList();return}}const locationFilter=e.target.closest("[data-location-filter]");if(locationFilter){{activeScriptLocation=locationFilter.dataset.locationFilter||"";scriptVisibleLimit=SCRIPT_INITIAL_RENDER_LIMIT;renderScriptLocationFilters();renderList();return}}const copy=e.target.closest("[data-copy]");if(copy){{await navigator.clipboard?.writeText(copy.dataset.copy).catch(()=>null);copy.textContent="已复制";return}}const scriptEdit=e.target.closest("[data-script-edit]");if(scriptEdit){{location.assign(`/studio?library_entry=${{encodeURIComponent(scriptEdit.dataset.scriptEdit)}}#split-panel`);return}}const delCreator=e.target.closest("[data-delete-creator]");if(delCreator){{if(confirm("确定删除这个创作者吗？")){{await api(`/api/creator-admin/creators/${{delCreator.dataset.deleteCreator}}`,{{method:"DELETE"}});selectedCreatorId="";await loadCreators()}}return}}const refresh=e.target.closest("[data-refresh-creator]");if(refresh){{const c=creators.find(x=>x.profile_id===refresh.dataset.refreshCreator);if(c){{await api(`/api/creator-admin/creators/${{c.profile_id}}`,{{method:"POST",body:JSON.stringify({{kwai_url:c.kwai_url,categories:c.categories||[],poc:creatorPocValue(c)}})}});await loadCreators()}}return}}const edit=e.target.closest("[data-edit]");if(edit)openEdit(edit.dataset.edit);const toggle=e.target.closest("[data-toggle]");if(toggle)togglePublish(toggle.dataset.toggle)}});document.addEventListener("change",async e=>{{const status=e.target.closest("[data-feed-status]");const date=e.target.closest("[data-feed-date]");const ret=e.target.closest("[data-feed-return-url]");const retTime=e.target.closest("[data-feed-return-time]");if(status||date||ret||retTime){{const target=status||date||ret||retTime;const id=selectedCreatorId;const entry=target.dataset.feedStatus||target.dataset.feedDate||target.dataset.feedReturnUrl||target.dataset.feedReturnTime;if(id&&entry)await updateCreatorFeed(id,entry,status?{{status:target.value}}:date?{{date:target.value,feed_time:target.value}}:ret?{{return_url:target.value}}:{{return_time:target.value}})}}});window.addEventListener("popstate",()=>{{const match=location.pathname.match(/^\\/creator-admin\\/creators\\/([0-9a-f]{{32}})$/);selectedCreatorId=(match&&match[1])||"";renderCreators()}});document.querySelector("#edit-cancel").addEventListener("click",()=>modal.classList.remove("open"));document.querySelector("#creator-cancel").addEventListener("click",()=>creatorModal.classList.remove("open"));document.querySelector("#creator-tags-cancel").addEventListener("click",()=>creatorTagsModal.classList.remove("open"));form.addEventListener("submit",saveEdit);creatorForm.addEventListener("submit",createCreator);window.addEventListener("focus",refreshCreatorAdminSilently);document.addEventListener("visibilitychange",()=>{if(!document.hidden)refreshCreatorAdminSilently()});startCreatorRealtime();loadCurrentTab();
</script></body></html>"""
    html = (
        template.replace("{{", "{")
        .replace("}}", "}")
        .replace("__CREATOR_BASE__", CREATOR_CENTER_BASE_URL)
        .replace("__INITIAL_TAB__", initial_tab_json)
        .replace("__LIBRARY_MODE__", library_mode_json)
        .replace("__FAVICON_LINKS__", FAVICON_LINKS)
    )
    if library_mode:
        html = html.replace("<title>Koko Creator 运营后台</title>", "<title>Koko 脚本管理</title>")
    return html


CREATOR_QUESTIONS = [
    {
        "id": "people",
        "pt": "Quantas pessoas aparecem normalmente?",
        "zh": "你们通常几个人拍？",
        "options": [
            {
                "id": "solo",
                "pt": "Só eu",
                "zh": "我一个人拍",
                "types": ["骗子", "偷奸耍滑", "整蛊"],
                "keywords": ["假装", "吐槽", "反应", "秘密", "发现", "装病", "偷懒", "耍小聪明"],
            },
            {
                "id": "duo",
                "pt": "Duas pessoas",
                "zh": "两个人拍",
                "types": ["夫妻吵架", "夫妻欺骗", "夫妻算计", "妻管严", "整蛊", "骗子", "赖账/金钱冲突"],
                "keywords": ["夫妻", "妻子", "丈夫", "老公", "老婆", "情侣", "朋友", "同事", "顾客", "老板"],
            },
            {
                "id": "group",
                "pt": "Três ou mais",
                "zh": "三个人以上",
                "types": ["夫妻欺骗", "夫妻算计", "骗子", "整蛊", "撬墙角"],
                "keywords": ["妈妈", "爸爸", "儿子", "女儿", "家庭", "亲戚", "朋友", "同事", "围观", "多人", "误会"],
            },
        ],
    },
    {
        "id": "scene",
        "pt": "Qual cena parece mais com seu conteúdo?",
        "zh": "你最常拍哪种关系/场景？",
        "options": [
            {
                "id": "solo_reaction",
                "pt": "Reação / monólogo",
                "zh": "一个人反应/独白",
                "people": ["solo"],
                "types": ["骗子", "偷奸耍滑", "整蛊"],
                "keywords": ["独自", "一个人", "反应", "吐槽", "发现", "假装", "装病", "误会"],
            },
            {
                "id": "solo_smart",
                "pt": "Esperteza / situação pessoal",
                "zh": "个人小聪明/自我处境",
                "people": ["solo"],
                "types": ["偷奸耍滑", "骗子"],
                "keywords": ["偷懒", "耍小聪明", "钻空子", "蒙混过关", "假装", "秘密", "尴尬"],
            },
            {
                "id": "duo_couple",
                "pt": "Casal / namorados",
                "zh": "夫妻/情侣",
                "people": ["duo"],
                "types": ["夫妻吵架", "夫妻欺骗", "夫妻算计", "妻管严", "夫妻黄段子", "夫妻好色", "夫妻出轨", "夫妻整蛊"],
                "keywords": ["夫妻", "妻子", "丈夫", "老公", "老婆", "情侣", "女友", "男友", "吃醋", "约会"],
            },
            {
                "id": "duo_friends",
                "pt": "Dois amigos / colegas",
                "zh": "两位朋友/同事",
                "people": ["duo"],
                "types": ["整蛊", "骗子", "偷奸耍滑", "撬墙角"],
                "keywords": ["朋友", "同事", "兄弟", "闺蜜", "套路", "恶作剧", "陷阱", "误会"],
            },
            {
                "id": "duo_service",
                "pt": "Cliente / chefe / atendimento",
                "zh": "两人顾客/老板/服务",
                "people": ["duo"],
                "types": ["赖账/金钱冲突", "骗子", "偷奸耍滑", "整蛊"],
                "keywords": ["老板", "员工", "顾客", "服务", "付款", "结账", "工资", "交易", "投诉", "费用"],
            },
            {
                "id": "group_family",
                "pt": "Família / filhos",
                "zh": "家庭/亲子",
                "people": ["group"],
                "types": ["夫妻欺骗", "夫妻算计"],
                "keywords": ["妈妈", "爸爸", "母亲", "父亲", "儿子", "女儿", "家庭", "生日", "礼物", "亲戚"],
            },
            {
                "id": "group_friends",
                "pt": "Grupo de amigos / colegas",
                "zh": "朋友群体/同事群",
                "people": ["group"],
                "types": ["整蛊", "骗子", "偷奸耍滑", "撬墙角"],
                "keywords": ["朋友", "同事", "兄弟", "闺蜜", "多人", "围观", "恶作剧", "误会", "套路"],
            },
            {
                "id": "group_public",
                "pt": "Rua / público / confusão",
                "zh": "街头/围观/多人误会",
                "people": ["group"],
                "types": ["整蛊", "骗子", "赖账/金钱冲突", "撬墙角"],
                "keywords": ["街头", "路人", "围观", "多人", "公共场合", "误会", "反转", "冲突"],
            },
        ],
    },
    {
        "id": "humor",
        "pt": "Que tipo de graça você quer?",
        "zh": "你想要哪种笑点？",
        "options": [
            {
                "id": "banter",
                "pt": "Discussão e respostas rápidas",
                "zh": "拌嘴互怼",
                "people": ["duo"],
                "scenes": ["duo_couple"],
                "types": ["夫妻吵架", "妻管严", "夫妻算计"],
                "keywords": ["吵架", "争执", "训斥", "反驳", "打脸", "抱怨", "不满"],
            },
            {
                "id": "twist",
                "pt": "Segredo e revelação",
                "zh": "隐瞒反转",
                "people": ["solo", "duo", "group"],
                "types": ["夫妻欺骗", "骗子", "夫妻算计"],
                "keywords": ["假装", "隐瞒", "谎称", "秘密", "真相", "发现", "揭开", "被骗", "冒充"],
            },
            {
                "id": "prank",
                "pt": "Pegadinha ou susto",
                "zh": "整蛊恶搞",
                "people": ["solo", "duo", "group"],
                "scenes": ["solo_reaction", "duo_friends", "duo_service", "group_friends", "group_public"],
                "types": ["整蛊", "夫妻整蛊"],
                "keywords": ["整蛊", "恶作剧", "捉弄", "吓唬", "陷阱", "搞怪", "吓得", "反应"],
            },
            {
                "id": "money",
                "pt": "Dinheiro ou vantagem",
                "zh": "钱/占便宜",
                "people": ["duo", "group"],
                "scenes": ["duo_service", "duo_friends", "group_public", "group_friends"],
                "types": ["赖账/金钱冲突", "骗子", "夫妻算计"],
                "keywords": ["付款", "欠钱", "不给钱", "逃单", "结账", "费用", "花钱", "信用卡", "便宜", "贵"],
            },
            {
                "id": "sneaky",
                "pt": "Preguiça ou esperteza",
                "zh": "偷懒/偷吃/耍小聪明",
                "people": ["solo", "duo", "group"],
                "types": ["偷吃东西", "偷奸耍滑"],
                "keywords": ["偷吃", "偷喝", "冰箱", "零食", "偷懒", "装病", "钻空子", "耍小聪明", "蒙混过关"],
            },
            {
                "id": "relationship",
                "pt": "Ciúmes / conflito de casal",
                "zh": "吃醋/亲密关系冲突",
                "people": ["duo"],
                "scenes": ["duo_couple"],
                "types": ["夫妻欺骗", "夫妻吵架", "夫妻出轨", "夫妻算计"],
                "keywords": ["吃醋", "出轨", "约会", "女友", "男友", "隐瞒", "吵架", "关系"],
            },
            {
                "id": "group_misunderstanding",
                "pt": "Mal-entendido em grupo",
                "zh": "多人误会扩散",
                "people": ["group"],
                "types": ["整蛊", "骗子", "撬墙角"],
                "keywords": ["多人", "围观", "误会", "传播", "发现", "尴尬", "反转"],
            },
        ],
    },
]


def is_local_creator_portal_request(handler: BaseHTTPRequestHandler) -> bool:
    if env_flag("KOKO_CREATOR_PORTAL_ENABLED"):
        return True
    if os.environ.get("RENDER") or os.environ.get("RENDER_SERVICE_ID"):
        return False
    host = str(handler.headers.get("Host") or "").split(":", 1)[0].strip().lower()
    return host in {"localhost", "127.0.0.1", "::1"}


def creator_library_entries_with_source() -> tuple[list[dict[str, Any]], str]:
    sync_creator_online_library_if_needed()
    data = read_json_file(CREATOR_ONLINE_LIBRARY_FILE, default=[])
    if isinstance(data, list) and data:
        entries = [entry for entry in data if isinstance(entry, dict)]
        for entry in entries:
            if str(entry.get("content_type") or "").strip() not in ALLOWED_CONTENT_TYPES:
                entry["content_type"] = DEFAULT_CONTENT_TYPE
        return entries, "https://koko-kwai-coach.onrender.com"
    return load_library_entries(), ""


def sync_creator_online_library_if_needed(*, force: bool = False) -> dict[str, Any]:
    source_url = str(CREATOR_LIBRARY_SOURCE_URL or "").strip()
    if not source_url:
        return {"ok": False, "reason": "missing_source_url"}
    meta = read_json_file(CREATOR_SYNC_META_FILE, default={})
    if not isinstance(meta, dict):
        meta = {}
    if not force and CREATOR_ONLINE_LIBRARY_FILE.exists():
        last_synced = str(meta.get("last_synced_at") or "").strip()
        try:
            synced_at = datetime.fromisoformat(last_synced.replace("Z", "+00:00"))
            if datetime.now(timezone.utc) - synced_at < timedelta(seconds=CREATOR_LIBRARY_SYNC_INTERVAL_SEC):
                return {"ok": True, "status": "fresh", **meta}
        except Exception:
            pass
    try:
        raw = fetch_remote_text(source_url, timeout=20)
        payload = json.loads(raw)
        entries = payload.get("entries") if isinstance(payload, dict) else payload
        if not isinstance(entries, list):
            raise ValueError("Creator library source did not return a list.")
        clean_entries = [entry for entry in entries if isinstance(entry, dict)]
        existing_entries = read_json_file(CREATOR_ONLINE_LIBRARY_FILE, default=[])
        if isinstance(existing_entries, list):
            direct_entries = [
                entry for entry in existing_entries
                if isinstance(entry, dict) and str(entry.get("source") or "") == "creator_direct_import"
            ]
            source_ids = {str(entry.get("entry_id") or "") for entry in clean_entries}
            clean_entries = direct_entries + [
                entry for entry in clean_entries
                if str(entry.get("entry_id") or "") not in source_ids or str(entry.get("source") or "") != "creator_direct_import"
            ]
            seen_ids: set[str] = set()
            merged_entries: list[dict[str, Any]] = []
            for entry in clean_entries:
                entry_id = str(entry.get("entry_id") or "")
                if entry_id and entry_id in seen_ids:
                    continue
                if entry_id:
                    seen_ids.add(entry_id)
                merged_entries.append(entry)
            clean_entries = merged_entries
        write_json_atomic(CREATOR_ONLINE_LIBRARY_FILE, clean_entries)
        meta = {
            "ok": True,
            "status": "synced",
            "source_url": source_url,
            "entries_count": len(clean_entries),
            "last_synced_at": now_iso(),
        }
        write_json_atomic(CREATOR_SYNC_META_FILE, meta)
        return meta
    except Exception as exc:
        meta = {
            **meta,
            "ok": False,
            "status": "failed",
            "source_url": source_url,
            "error": friendly_error(str(exc)),
            "failed_at": now_iso(),
        }
        write_json_atomic(CREATOR_SYNC_META_FILE, meta)
        return meta


def trigger_creator_center_sync() -> dict[str, Any]:
    target_url = str(CREATOR_CENTER_SYNC_URL or "").strip()
    if not target_url:
        return {"ok": False, "error": "CREATOR_CENTER_SYNC_URL is not configured."}
    try:
        request = urllib.request.Request(
            target_url,
            data=b"{}",
            method="POST",
            headers={
                "Content-Type": "application/json",
                "User-Agent": "KokoScriptLibrary/1.0",
            },
        )
        with urllib.request.urlopen(request, timeout=30) as response:
            raw = response.read().decode(response.headers.get_content_charset() or "utf-8", errors="ignore")
        payload = json.loads(raw) if raw.strip() else {}
        if not isinstance(payload, dict):
            payload = {"response": payload}
        return {"ok": True, "target_url": target_url, **payload}
    except Exception as exc:
        return {"ok": False, "target_url": target_url, "error": friendly_error(str(exc))}


def trigger_creator_center_sync_background(reason: str = "library_update") -> None:
    def _run() -> None:
        result = trigger_creator_center_sync()
        if not result.get("ok"):
            log_runtime_warning(
                "creator_center_background_sync_failed",
                "Creator center background sync failed.",
                reason=reason,
                error=result.get("error") or result,
            )

    threading.Thread(target=_run, name=f"creator-sync-{reason}", daemon=True).start()


def push_creator_import_to_center(entry: dict[str, Any], script_json: dict[str, Any], output_dir: Path) -> dict[str, Any]:
    html_path = output_dir / "script_table_pt.html"
    if not html_path.exists():
        html_path = output_dir / "script_table.html"
    if not html_path.exists():
        return {"ok": False, "error": "Generated script HTML was not found."}
    cover_b64 = ""
    cover_mime = "image/png"
    storyboard = load_storyboard_state(str(entry.get("entry_id") or ""))
    cover_url = str(storyboard.get("storyboard_cover_url") or "").strip()
    if cover_url.startswith("/results/"):
        cover_name = cover_url.rsplit("/", 1)[-1]
        cover_path = output_dir / cover_name
        if cover_path.exists():
            cover_mime = "image/jpeg" if cover_path.suffix.lower() in {".jpg", ".jpeg"} else "image/png"
            cover_b64 = base64.b64encode(cover_path.read_bytes()).decode("ascii")
    if not cover_b64:
        return {"ok": False, "error": "Storyboard cover image was not found; Creator import was skipped."}
    payload = {
        "entry": entry,
        "script_json": script_json,
        "html_content": html_path.read_text(encoding="utf-8"),
        "cover_b64": cover_b64,
        "cover_mime": cover_mime,
    }
    return creator_admin_remote_json("/api/admin/scripts/import", method="POST", payload=payload)[1]


def creator_abs_url(url: object, base_url: str) -> str:
    text = str(url or "").strip()
    if not text:
        return ""
    if text.startswith(("http://", "https://")):
        return text
    if text.startswith("/") and base_url:
        return base_url.rstrip("/") + text
    return text


def creator_effective_entries(entries: list[dict[str, Any]]) -> list[dict[str, Any]]:
    filtered = [
        entry for entry in entries
        if str(entry.get("title") or "").strip()
        and str(entry.get("whole_video_summary") or "").strip()
        and (entry.get("html_url") or entry.get("zh_html_url") or entry.get("video_url"))
    ]
    return sorted(filtered, key=lambda item: str(item.get("saved_at") or item.get("created_at") or ""), reverse=True)


def creator_option_lookup() -> dict[str, dict[str, Any]]:
    return {
        str(option["id"]): option
        for question in CREATOR_QUESTIONS
        for option in question.get("options", [])
    }


def creator_score_entry(entry: dict[str, Any], selected: list[str], index: int) -> int:
    lookup = creator_option_lookup()
    text = " ".join(str(entry.get(key) or "") for key in ["content_type", "title", "whole_video_summary", "content_type_reasoning"])
    content_type = str(entry.get("content_type") or DEFAULT_CONTENT_TYPE)
    score = 0
    for option_id in selected:
        option = lookup.get(option_id) or {}
        if content_type in set(option.get("types") or []):
            score += 42
        hits = sum(1 for keyword in option.get("keywords") or [] if str(keyword) and str(keyword) in text)
        score += min(24, hits * 6)
    if content_type != DEFAULT_CONTENT_TYPE:
        score += 10
    if entry.get("html_url") or entry.get("zh_html_url"):
        score += 8
    if entry.get("video_url"):
        score += 4
    score += max(0, 10 - min(index, 10))
    return score


def creator_public_entry(entry: dict[str, Any], base_url: str, score: int) -> dict[str, Any]:
    entry_id = str(entry.get("entry_id") or "").strip()
    return {
        "entry_id": entry_id,
        "title": entry.get("title") or "未命名脚本",
        "summary": entry.get("whole_video_summary") or "",
        "content_type": entry.get("content_type") or DEFAULT_CONTENT_TYPE,
        "created_at": format_beijing_time(entry.get("created_at") or entry.get("saved_at") or ""),
        "video_url": creator_abs_url(entry.get("video_url"), ""),
        "html_url": creator_abs_url(entry.get("zh_html_url") or entry.get("html_url"), base_url),
        "docx_url": creator_abs_url(entry.get("zh_docx_url") or entry.get("docx_url"), base_url),
        "thumbnail_url": f"/api/creator/thumbnail/{entry_id}.svg" if entry_id else "",
        "score": score,
    }


def creator_recommendation_payload(selected: list[str], limit: int = 80) -> dict[str, Any]:
    entries, base_url = creator_library_entries_with_source()
    effective = creator_effective_entries(entries)
    selected = [value for value in selected if value in creator_option_lookup()]
    scored = sorted(
        ((creator_score_entry(entry, selected, idx), entry) for idx, entry in enumerate(effective)),
        key=lambda pair: pair[0],
        reverse=True,
    )
    return {
        "questions": CREATOR_QUESTIONS,
        "selected": selected,
        "total": len(scored),
        "entries": [creator_public_entry(entry, base_url, score) for score, entry in scored[:limit]],
        "using_online_cache": bool(base_url),
    }


def creator_facets_payload() -> dict[str, Any]:
    entries, _ = creator_library_entries_with_source()
    return {"questions": CREATOR_QUESTIONS, "total": len(creator_effective_entries(entries))}


def creator_entry_by_id(entry_id: str) -> dict[str, Any] | None:
    entries, _ = creator_library_entries_with_source()
    for entry in entries:
        if str(entry.get("entry_id") or "") == entry_id:
            return entry
    return None


def creator_thumbnail_cache() -> dict[str, Any]:
    data = read_json_file(CREATOR_THUMBNAIL_CACHE_FILE, default={})
    return data if isinstance(data, dict) else {}


def creator_placeholder_svg(entry: dict[str, Any] | None) -> bytes:
    title = html.escape(str((entry or {}).get("title") or "Koko Creator")[:54])
    content_type = html.escape(str((entry or {}).get("content_type") or "Roteiro"))
    svg = f"""<svg xmlns="http://www.w3.org/2000/svg" width="320" height="420" viewBox="0 0 320 420">
  <defs>
    <linearGradient id="bg" x1="0" x2="1" y1="0" y2="1">
      <stop offset="0" stop-color="#ffb357"/>
      <stop offset=".52" stop-color="#ff6500"/>
      <stop offset="1" stop-color="#372018"/>
    </linearGradient>
  </defs>
  <rect width="320" height="420" rx="28" fill="url(#bg)"/>
  <circle cx="238" cy="82" r="62" fill="#fff" opacity=".18"/>
  <circle cx="52" cy="338" r="96" fill="#fff" opacity=".12"/>
  <text x="28" y="64" fill="#fff" font-family="Arial, sans-serif" font-size="22" font-weight="700">Koko Creator</text>
  <text x="28" y="304" fill="#fff" font-family="Arial, sans-serif" font-size="18" font-weight="700">{content_type}</text>
  <foreignObject x="28" y="322" width="250" height="76">
    <div xmlns="http://www.w3.org/1999/xhtml" style="color:white;font-family:Arial,sans-serif;font-size:24px;font-weight:800;line-height:1.12;">{title}</div>
  </foreignObject>
</svg>"""
    return svg.encode("utf-8")


def creator_thumbnail_url_for_entry(entry: dict[str, Any]) -> str:
    entry_id = str(entry.get("entry_id") or "").strip()
    video_url = str(entry.get("video_url") or "").strip()
    if not entry_id or not video_url:
        return ""
    cache = creator_thumbnail_cache()
    cached = cache.get(entry_id)
    if isinstance(cached, dict) and str(cached.get("thumbnail_url") or "").strip():
        return str(cached.get("thumbnail_url") or "").strip()
    try:
        metadata = fetch_kwai_light_metadata(video_url)
        thumbnail_url = str(metadata.get("thumbnail_url") or "").strip()
    except Exception:
        thumbnail_url = ""
    cache[entry_id] = {
        "thumbnail_url": thumbnail_url,
        "video_url": video_url,
        "checked_at": now_iso(),
    }
    write_json_atomic(CREATOR_THUMBNAIL_CACHE_FILE, cache)
    return thumbnail_url


def load_creator_submissions() -> list[dict[str, Any]]:
    data = read_json_file(CREATOR_SUBMISSIONS_FILE, default=[])
    return data if isinstance(data, list) else []


def save_creator_submission(payload: dict[str, Any]) -> dict[str, Any]:
    entry_id = str(payload.get("entry_id") or "").strip()
    video_url = str(payload.get("video_url") or "").strip()
    if not re.fullmatch(r"[0-9a-f]{32}", entry_id):
        raise ValueError("Invalid script id.")
    if not video_url.startswith(("http://", "https://")):
        raise ValueError("Please submit a public video link.")
    entry = creator_entry_by_id(entry_id)
    if not entry:
        raise ValueError("Script not found.")
    submission = {
        "submission_id": uuid4().hex,
        "entry_id": entry_id,
        "script_title": str(entry.get("title") or ""),
        "script_content_type": str(entry.get("content_type") or DEFAULT_CONTENT_TYPE),
        "creator_id": str(payload.get("creator_id") or "local_creator").strip()[:120],
        "video_url": video_url,
        "note": str(payload.get("note") or "").strip()[:1000],
        "status": "pending_review",
        "created_at": now_iso(),
    }
    submissions = load_creator_submissions()
    submissions.insert(0, submission)
    write_json_atomic(CREATOR_SUBMISSIONS_FILE, submissions[:1000])
    return submission


def creator_portal_html() -> str:
    questions_json = json.dumps(CREATOR_QUESTIONS, ensure_ascii=False)
    facets_json = json.dumps(creator_facets_payload(), ensure_ascii=False)
    return f"""<!doctype html>
<html lang="pt-BR">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width,initial-scale=1,viewport-fit=cover">
  <title>Koko Creator</title>
  <link rel="icon" type="image/svg+xml" href="/favicon.svg?v=kwai1">
  <link rel="shortcut icon" href="/favicon.ico?v=kwai1">
  <style>
    @import url('https://fonts.googleapis.com/css2?family=Readex+Pro:wght@300;400;500;600;700&display=swap');
    * {{ box-sizing: border-box; font-family: 'Readex Pro', -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif; }}
    body {{ margin:0; min-height:100vh; background:#fff4ea; color:#1f1f1f; }}
    button,a {{ font:inherit; }}
    .phone {{ width:min(100%,480px); min-height:100vh; margin:0 auto; overflow-x:hidden; padding-bottom:96px; background:radial-gradient(circle at 88% 22%,rgba(255,130,0,.16),transparent 25%),linear-gradient(180deg,#fffaf5 0%,#fff0df 42%,#fff8f2 100%); }}
    .topbar {{ position:sticky; top:0; z-index:20; min-height:74px; display:flex; align-items:center; justify-content:space-between; gap:12px; padding:max(16px,env(safe-area-inset-top)) 24px 12px; background:rgba(255,252,248,.88); backdrop-filter:blur(18px); }}
    .brand {{ display:flex; align-items:center; gap:12px; min-width:0; }}
    .brand img {{ width:126px; }}
    .brand span {{ color:#ff5f00; font-weight:850; white-space:nowrap; }}
    .menu {{ border:0; background:transparent; font-size:30px; cursor:pointer; }}
    .lang {{ position:fixed; right:max(16px,calc((100vw - 480px)/2 + 16px)); bottom:98px; z-index:30; display:flex; gap:4px; padding:5px; border-radius:999px; background:rgba(255,255,255,.9); box-shadow:0 12px 28px rgba(255,130,0,.12); }}
    .lang button {{ border:0; border-radius:999px; padding:7px 10px; background:transparent; color:#777; font-size:12px; font-weight:850; }}
    .lang .active {{ background:#ff5f00; color:#fff; }}
    .view {{ display:none; padding:22px 24px 18px; }}
    .view.active {{ display:block; }}
    .pill {{ display:inline-flex; align-items:center; gap:9px; max-width:100%; border:1px solid rgba(255,95,0,.48); border-radius:999px; padding:9px 14px; color:#ff5f00; font-size:13px; font-weight:850; background:rgba(255,255,255,.54); }}
    h1 {{ margin:22px 0 14px; font-size:clamp(40px,11vw,60px); line-height:1.08; font-weight:900; letter-spacing:0; }}
    .accent {{ color:#ff5f00; }}
    .lead {{ margin:0; color:#69707a; font-size:17px; line-height:1.55; }}
    .hero-art {{ position:relative; min-height:168px; margin:12px -24px 0; overflow:hidden; }}
    .wave {{ position:absolute; right:-70px; bottom:-36px; width:260px; height:128px; border-radius:80% 0 0 0; background:rgba(255,130,0,.14); }}
    .mascot {{ position:absolute; right:22px; bottom:8px; width:118px; height:118px; border-radius:52% 48% 44% 56%; background:radial-gradient(circle at 35% 22%,#ffbe55,#ff8e24 64%,#f97808); box-shadow:0 18px 40px rgba(255,130,0,.22); }}
    .mascot:before,.mascot:after {{ content:""; position:absolute; top:30px; width:26px; height:31px; border-radius:50%; background:#fff; }}
    .mascot:before {{ left:28px; }} .mascot:after {{ right:27px; }}
    .eye {{ position:absolute; top:41px; width:11px; height:13px; border-radius:50%; background:#5b2a10; z-index:2; }}
    .eye.left {{ left:37px; }} .eye.right {{ right:36px; }}
    .mouth {{ position:absolute; left:48px; top:69px; width:30px; height:22px; border-radius:0 0 18px 18px; background:#9b2b00; }}
    .cta-row {{ display:grid; gap:12px; margin:18px 0 24px; }}
    .primary {{ border:0; border-radius:999px; min-height:58px; padding:0 18px; display:flex; align-items:center; justify-content:center; gap:12px; background:linear-gradient(90deg,#ff6a00,#ff5200); color:#fff; box-shadow:0 14px 30px rgba(255,95,0,.28); font-size:18px; font-weight:900; cursor:pointer; text-decoration:none; }}
    .view > .primary {{ position:sticky; bottom:92px; z-index:18; width:100%; margin-top:18px; }}
    .secondary {{ border:0; border-radius:999px; min-height:48px; padding:0 18px; background:rgba(255,255,255,.86); color:#1f1f1f; box-shadow:0 10px 24px rgba(0,0,0,.06); font-weight:850; cursor:pointer; }}
    .card {{ border-radius:24px; background:rgba(255,255,255,.84); border:1px solid rgba(255,130,0,.10); box-shadow:0 16px 40px rgba(85,45,10,.08); }}
    .preview {{ padding:18px; margin-top:12px; }}
    .preview-head {{ display:flex; align-items:center; justify-content:space-between; gap:12px; margin-bottom:14px; }}
    .preview-head h2 {{ margin:0; font-size:18px; }}
    .mini-cards {{ display:grid; grid-template-columns:repeat(3,150px); gap:12px; overflow:auto; padding-bottom:6px; scrollbar-width:none; }}
    .mini-card {{ position:relative; overflow:hidden; height:188px; border-radius:16px; color:#fff; padding:12px; display:flex; flex-direction:column; justify-content:flex-end; background:#2a1d16; }}
    .mini-card img {{ position:absolute; inset:0; width:100%; height:100%; object-fit:cover; }}
    .mini-card:after {{ content:""; position:absolute; inset:0; background:linear-gradient(180deg,rgba(0,0,0,.08),rgba(0,0,0,.72)); }}
    .mini-card > * {{ position:relative; z-index:1; }}
    .score {{ align-self:flex-start; margin-bottom:auto; border-radius:9px; padding:6px 8px; background:rgba(158,73,12,.88); font-size:12px; font-weight:850; }}
    .mini-card b {{ font-size:16px; line-height:1.18; }}
    .stepper {{ display:grid; grid-template-columns:repeat(3,1fr); gap:8px; margin:18px 0 24px; }}
    .step {{ min-height:58px; border-radius:18px; background:rgba(255,255,255,.66); color:#878787; display:grid; place-items:center; text-align:center; font-size:12px; font-weight:850; }}
    .step.active {{ background:#ff5f00; color:#fff; }}
    .step i {{ display:grid; place-items:center; width:28px; height:28px; border-radius:50%; background:rgba(255,255,255,.72); color:#ff5f00; font-style:normal; }}
    .question {{ display:none; }}
    .question.active {{ display:block; }}
    .options {{ display:grid; gap:12px; margin:18px 0; }}
    .option {{ min-height:74px; display:flex; align-items:center; gap:12px; width:100%; border:1px solid rgba(255,130,0,.14); border-radius:18px; padding:14px; background:rgba(255,255,255,.86); color:#1f1f1f; text-align:left; font-weight:850; cursor:pointer; }}
    .option.selected {{ border-color:#ff5f00; color:#ff5f00; box-shadow:0 12px 26px rgba(255,95,0,.12); }}
    .option span {{ display:grid; place-items:center; width:38px; height:38px; border-radius:50%; background:#fff0e8; font-size:20px; }}
    .filters {{ display:flex; flex-wrap:wrap; gap:9px; margin:16px 0; }}
    .chip {{ border:1px solid rgba(255,95,0,.34); border-radius:999px; padding:8px 11px; color:#ff5f00; background:rgba(255,255,255,.58); font-size:12px; font-weight:850; }}
    .feed {{ display:grid; gap:14px; }}
    .script {{ padding:14px; display:grid; grid-template-columns:116px 1fr; gap:13px; min-height:168px; }}
    .thumb {{ position:relative; overflow:hidden; border-radius:16px; min-height:142px; background:#2a1d16; color:#fff; padding:10px; font-weight:900; font-size:12px; }}
    .thumb img {{ position:absolute; inset:0; width:100%; height:100%; object-fit:cover; }}
    .thumb:after {{ content:""; position:absolute; inset:0; background:linear-gradient(180deg,rgba(0,0,0,.08),rgba(0,0,0,.66)); }}
    .thumb span {{ position:relative; z-index:1; display:inline-flex; border-radius:9px; padding:6px 8px; background:rgba(158,73,12,.88); }}
    .script-body {{ min-width:0; display:flex; flex-direction:column; gap:8px; }}
    .script h3 {{ margin:0; font-size:19px; line-height:1.22; display:-webkit-box; -webkit-line-clamp:2; -webkit-box-orient:vertical; overflow:hidden; }}
    .script p {{ margin:0; color:#69707a; font-size:13px; line-height:1.42; display:-webkit-box; -webkit-line-clamp:3; -webkit-box-orient:vertical; overflow:hidden; }}
    .tags {{ display:flex; gap:6px; flex-wrap:wrap; }}
    .tag {{ border-radius:999px; padding:5px 8px; background:#fff0e8; color:#ff5f00; font-size:11px; font-weight:850; }}
    .actions {{ display:flex; align-items:center; gap:8px; margin-top:auto; }}
    .open {{ flex:1; min-height:38px; border-radius:999px; display:inline-flex; align-items:center; justify-content:center; background:#ff5f00; color:#fff; text-decoration:none; font-size:13px; font-weight:900; }}
    .plain {{ color:#1f1f1f; text-decoration:none; font-size:12px; font-weight:850; }}
    .state-card {{ padding:18px; display:grid; gap:12px; }}
    .state-card h3 {{ margin:0; font-size:18px; }}
    .state-card p {{ margin:0; color:#69707a; line-height:1.45; font-size:14px; }}
    .section-title {{ display:flex; align-items:center; justify-content:space-between; gap:12px; margin:20px 0 12px; }}
    .section-title h2 {{ margin:0; font-size:20px; line-height:1.2; }}
    .quick-grid {{ display:grid; grid-template-columns:repeat(3,1fr); gap:10px; margin:16px 0; }}
    .quick {{ min-height:78px; border:0; border-radius:18px; padding:12px 8px; background:rgba(255,255,255,.86); color:#1f1f1f; box-shadow:0 10px 24px rgba(0,0,0,.05); font-weight:850; font-size:12px; }}
    .quick b {{ display:block; color:#ff5f00; font-size:20px; margin-bottom:4px; }}
    .status-tabs {{ display:flex; gap:8px; overflow:auto; padding:4px 0 12px; scrollbar-width:none; }}
    .status-tabs button {{ flex:0 0 auto; border:1px solid rgba(255,95,0,.22); border-radius:999px; padding:9px 13px; background:rgba(255,255,255,.72); color:#777; font-size:12px; font-weight:850; }}
    .status-tabs .active {{ background:#ff5f00; color:#fff; border-color:#ff5f00; }}
    .icon-btn {{ border:0; width:38px; height:38px; border-radius:50%; display:grid; place-items:center; background:#fff0e8; color:#ff5f00; font-size:18px; font-weight:900; }}
    .action-row {{ display:grid; grid-template-columns:1fr 38px 38px; gap:8px; align-items:center; }}
    .modal {{ position:fixed; inset:0; z-index:50; display:none; background:rgba(31,31,31,.34); padding:20px 18px 0; }}
    .modal.active {{ display:flex; align-items:flex-end; justify-content:center; }}
    .sheet {{ width:min(100%,480px); max-height:88vh; overflow:auto; border-radius:28px 28px 0 0; background:#fffaf5; padding:18px 18px max(22px,env(safe-area-inset-bottom)); box-shadow:0 -20px 50px rgba(0,0,0,.20); }}
    .sheet-head {{ display:flex; align-items:flex-start; gap:12px; justify-content:space-between; }}
    .sheet h2 {{ margin:6px 0 10px; font-size:25px; line-height:1.15; }}
    .detail-media {{ position:relative; overflow:hidden; height:220px; border-radius:20px; background:#2a1d16; margin:14px 0; }}
    .detail-media img {{ width:100%; height:100%; object-fit:cover; }}
    .detail-block {{ margin:14px 0; }}
    .detail-block h3 {{ margin:0 0 7px; font-size:16px; }}
    .detail-block p {{ margin:0; color:#69707a; line-height:1.5; font-size:14px; }}
    .submit-box {{ display:grid; gap:10px; padding:14px; margin:14px 0; border-radius:18px; background:#fff0e8; border:1px solid rgba(255,95,0,.18); }}
    .submit-box label {{ font-size:13px; font-weight:900; color:#ff5f00; }}
    .submit-box input {{ width:100%; min-height:46px; border:1px solid rgba(255,95,0,.22); border-radius:14px; padding:0 12px; background:#fff; color:#1f1f1f; font-size:14px; }}
    .submit-hint {{ margin:0; color:#69707a; line-height:1.45; font-size:12px; }}
    .submit-status {{ min-height:18px; color:#ff5f00; font-size:12px; font-weight:850; }}
    .bottom {{ position:fixed; left:50%; bottom:0; transform:translateX(-50%); z-index:25; width:min(100%,480px); display:grid; grid-template-columns:repeat(2,1fr); gap:2px; padding:10px 14px max(10px,env(safe-area-inset-bottom)); border-radius:24px 24px 0 0; background:rgba(255,255,255,.94); box-shadow:0 -14px 34px rgba(0,0,0,.08); }}
    .bottom button {{ border:0; background:transparent; color:#777; min-height:54px; display:grid; place-items:center; gap:4px; font-size:12px; font-weight:750; }}
    .bottom .active {{ color:#ff5f00; }}
    .nav-icon {{ font-size:23px; }}
    @media (max-width:380px) {{ .brand img{{width:108px}} .brand span{{font-size:14px}} .view{{padding-left:18px;padding-right:18px}} h1{{font-size:38px}} .script{{grid-template-columns:104px 1fr}} }}
  </style>
</head>
<body>
  <main class="phone">
    <header class="topbar"><div class="brand"><img src="/brand/kwai-wordmark.svg" alt="Kwai"><span>Koko Creator</span></div><button class="menu" type="button" aria-label="Menu">☰</button></header>
    <div class="lang"><button type="button" data-lang="pt" class="active">PT</button><button type="button" data-lang="zh">中文</button></div>
    <section class="view" data-view="home">
      <div class="pill">✦ <span data-i18n="homePill">Biblioteca de roteiros Koko Creator</span></div>
      <h1 data-i18n-html="homeTitle">Encontre mais rápido <span class="accent">roteiros</span> que você realmente consegue gravar</h1>
      <p class="lead" data-i18n="homeLead">Responda 3 perguntas simples e veja roteiros que combinam com o jeito que você grava.</p>
      <div class="hero-art"><div class="wave"></div><div class="mascot"><span class="eye left"></span><span class="eye right"></span><span class="mouth"></span></div></div>
      <div class="cta-row"><button class="primary" type="button" data-go="choose">✦ <span data-i18n="start">Começar agora</span> →</button></div>
      <section class="preview card"><div class="preview-head"><h2 data-i18n="recommended">Recomendado para você</h2></div><div class="mini-cards" id="mini-cards"></div></section>
    </section>
    <section class="view" data-view="dashboard">
      <div class="pill">✦ <span data-i18n="todayPill">Recomendação de roteiros</span></div>
      <h1 data-i18n="todayTitle">Recomendação de roteiros</h1>
      <p class="lead" data-i18n="todayLead">Abra, salve e marque o que você vai gravar hoje.</p>
      <div class="quick-grid">
        <button class="quick" type="button"><b id="count-new">0</b><span data-i18n="quickNew">novos</span></button>
        <button class="quick" type="button" data-go="saved"><b id="count-saved">0</b><span data-i18n="quickSaved">salvos</span></button>
        <button class="quick" type="button" data-go="saved" data-saved-tab="planned"><b id="count-planned">0</b><span data-i18n="quickPlan">para gravar</span></button>
      </div>
      <div class="filters" id="dashboard-filters"></div>
      <div class="section-title"><h2 data-i18n="todayRecommended">Para gravar hoje</h2></div>
      <div class="feed" id="dashboard-feed"></div>
    </section>
    <section class="view" data-view="choose">
      <div class="pill">✦ <span id="step-label">Etapa 1 de 3</span></div>
      <div class="stepper" id="stepper"></div>
      <div id="question-wrap"></div>
      <button class="primary" type="button" id="next-step"><span data-i18n="next">Próxima etapa</span> →</button>
    </section>
    <section class="view" data-view="library">
      <div class="pill">✦ <span data-i18n="libraryPill">Biblioteca de roteiros Koko Creator</span></div>
      <h1 data-i18n="resultTitle">Roteiros recomendados para você</h1>
      <div class="filters" id="selected-filters"></div>
      <div class="feed" id="feed"></div>
    </section>
    <section class="view" data-view="saved">
      <div class="pill">✦ <span data-i18n="savedPill">Meus roteiros</span></div>
      <h1 data-i18n="savedTitle">Sua lista de gravação</h1>
      <div class="status-tabs" id="saved-tabs"></div>
      <div class="feed" id="saved-feed"></div>
    </section>
  </main>
  <nav class="bottom"><button type="button" data-go="dashboard"><span class="nav-icon">⌂</span><span data-i18n="navHome">Roteiros</span></button><button type="button" data-go="saved"><span class="nav-icon">♡</span><span data-i18n="navSaved">Salvos</span></button></nav>
  <div class="modal" id="detail-modal"><section class="sheet"><div class="sheet-head"><div class="pill">Koko Creator</div><button class="icon-btn" type="button" data-close-detail>×</button></div><div id="detail-content"></div></section></div>
  <script>
    const questions = {questions_json};
    const facets = {facets_json};
    const profileKey = "koko_creator_profile_v3";
    const langKey = "koko_creator_lang";
    const workspaceKey = "koko_creator_workspace_v1";
    let lang = localStorage.getItem(langKey) || "pt";
    let step = 0;
    let savedTab = "saved";
    let matchesCache = [];
    let answers = JSON.parse(localStorage.getItem(profileKey) || "null") || {{ people: "duo", scene: "couple", humor: "twist" }};
    let workspace = JSON.parse(localStorage.getItem(workspaceKey) || "null") || {{ saved: [], planned: [], finished: [], rejected: [] }};
    const i18n = {{
      pt: {{ homePill:"Biblioteca de roteiros Koko Creator", homeTitle:'Encontre mais rápido <span class="accent">roteiros</span> que você realmente consegue gravar', homeLead:"Responda 3 perguntas simples e veja roteiros que combinam com o jeito que você grava.", start:"Começar agora", seePopular:"Ver populares", recommended:"Recomendado para você", seeAll:"Ver todos", next:"Próxima etapa", finish:"Ver recomendações", libraryPill:"Biblioteca de roteiros Koko Creator", resultTitle:"Sua biblioteca recomendada", open:"Abrir", original:"Vídeo", navHome:"Roteiros", navPrefs:"Perfil", navSaved:"Salvos", navLibrary:"Biblioteca", changePrefs:"Mudar preferências", step:"Etapa", todayPill:"Recomendação de roteiros", todayTitle:"Recomendação de roteiros", todayLead:"Abra, salve e marque o que você vai gravar hoje.", quickNew:"roteiros", quickSaved:"salvos", quickPlan:"para gravar", todayRecommended:"Para gravar hoje", savedPill:"Meus roteiros", savedTitle:"Sua lista de gravação", save:"Salvar", saved:"Salvo", plan:"Vou gravar", done:"Gravado", reject:"Não serve", copy:"Copiar resumo", emptySaved:"Nada aqui ainda", emptySavedText:"Salve um roteiro da recomendação para montar sua lista.", details:"Detalhes do roteiro", quickSummary:"Resumo rápido", howToUse:"Como usar", howToUseText:"Leia o resumo, veja o vídeo de referência e marque se vai gravar.", peopleTag:"2 pessoas", placeTag:"Baixo custo", statusSaved:"Salvos", statusPlanned:"Vou gravar", statusFinished:"Gravados", statusRejected:"Não servem", submitTitle:"Enviar vídeo gravado", submitHint:"Envie o link do vídeo gravado seguindo este roteiro. Vamos revisar e, se aprovado, ajudar com impulsionamento.", submitPlaceholder:"Cole aqui o link do seu vídeo", submitButton:"Enviar para revisão", submitOk:"Recebido. Vamos revisar seu vídeo.", submitError:"Não foi possível enviar. Confira o link e tente novamente." }},
      zh: {{ homePill:"Koko Creator 脚本推荐", homeTitle:'更快找到你<span class="accent">真的能拍</span>的脚本', homeLead:"回答 3 个简单问题，Koko 会按你的拍摄方式推荐脚本。", start:"开始选择", seePopular:"先看热门", recommended:"为你推荐", seeAll:"查看全部", next:"下一步", finish:"查看推荐", libraryPill:"Koko Creator 脚本库", resultTitle:"你的推荐脚本库", open:"打开", original:"原视频", navHome:"脚本推荐", navPrefs:"偏好", navSaved:"收藏", navLibrary:"脚本库", changePrefs:"重新选择偏好", step:"第", todayPill:"脚本推荐", todayTitle:"脚本推荐", todayLead:"打开、收藏，并标记今天准备拍的脚本。", quickNew:"推荐脚本", quickSaved:"已收藏", quickPlan:"准备拍", todayRecommended:"今天可以拍", savedPill:"我的脚本", savedTitle:"你的拍摄清单", save:"收藏", saved:"已收藏", plan:"准备拍", done:"已拍", reject:"不适合", copy:"复制摘要", emptySaved:"这里还没有脚本", emptySavedText:"先从脚本推荐里收藏一个脚本，建立你的拍摄清单。", details:"脚本详情", quickSummary:"快速梗概", howToUse:"怎么使用", howToUseText:"先看梗概和原视频，再标记是否准备拍。", peopleTag:"2 人", placeTag:"低成本", statusSaved:"已收藏", statusPlanned:"准备拍", statusFinished:"已拍", statusRejected:"不适合", submitTitle:"回传拍摄视频", submitHint:"上传按照脚本拍摄的视频，我们会审核后给您投流。", submitPlaceholder:"把你发布后的视频链接粘贴在这里", submitButton:"提交审核", submitOk:"已收到，我们会审核这个视频。", submitError:"提交失败，请检查链接后重试。" }}
    }};
	    const t = key => (i18n[lang] && i18n[lang][key]) || key;
	    const label = item => lang === "zh" ? item.zh : item.pt;
	    const esc = value => String(value || "").replace(/[&<>"']/g, c => ({{"&":"&amp;","<":"&lt;",">":"&gt;",'"':"&quot;","'":"&#39;"}}[c]));
	    function optionAllowed(opt) {{
	      if (!opt) return false;
	      if (Array.isArray(opt.people) && opt.people.length && !opt.people.includes(answers.people)) return false;
	      if (Array.isArray(opt.scenes) && opt.scenes.length && !opt.scenes.includes(answers.scene)) return false;
	      return true;
	    }}
	    function optionsFor(q) {{ return (q.options || []).filter(optionAllowed); }}
	    function normalizeAnswers() {{
	      let changed = false;
	      questions.forEach(q => {{
	        const opts = optionsFor(q);
	        if (!opts.length) return;
	        if (!opts.some(opt => opt.id === answers[q.id])) {{
	          answers[q.id] = opts[0].id;
	          changed = true;
	        }}
	      }});
	      if (changed) save();
	      return changed;
	    }}
	    function selectedAnswerValues() {{
	      normalizeAnswers();
	      return questions.map(q => answers[q.id]).filter(Boolean);
	    }}
	    function save() {{ localStorage.setItem(profileKey, JSON.stringify(answers)); }}
    function saveWorkspace() {{ localStorage.setItem(workspaceKey, JSON.stringify(workspace)); updateCounts(); }}
    function hasProfile() {{ return !!localStorage.getItem(profileKey); }}
    function ids(name) {{ return new Set(workspace[name] || []); }}
    function entryById(id) {{ return matchesCache.find(e => e.entry_id === id); }}
    function setStatus(id, status) {{
      ["saved","planned","finished","rejected"].forEach(key => workspace[key] = (workspace[key] || []).filter(item => item !== id));
      if (status) workspace[status] = [...(workspace[status] || []), id];
      saveWorkspace();
      renderAllKnown();
    }}
    function statusOf(id) {{
      if (ids("planned").has(id)) return "planned";
      if (ids("finished").has(id)) return "finished";
      if (ids("rejected").has(id)) return "rejected";
      if (ids("saved").has(id)) return "saved";
      return "";
    }}
    function updateCounts() {{
      const newNode = document.querySelector("#count-new"); if (newNode) newNode.textContent = String(matchesCache.length || facets.total || 0);
      const savedNode = document.querySelector("#count-saved"); if (savedNode) savedNode.textContent = String((workspace.saved || []).length);
      const plannedNode = document.querySelector("#count-planned"); if (plannedNode) plannedNode.textContent = String((workspace.planned || []).length);
    }}
    function applyLang() {{
      document.documentElement.lang = lang === "zh" ? "zh-CN" : "pt-BR";
      document.querySelectorAll("[data-lang]").forEach(btn => btn.classList.toggle("active", btn.dataset.lang === lang));
      document.querySelectorAll("[data-i18n]").forEach(node => node.textContent = t(node.dataset.i18n));
      document.querySelectorAll("[data-i18n-html]").forEach(node => node.innerHTML = t(node.dataset.i18nHtml));
      renderQuestion(); renderMini(); renderSaved(); if (activeView() === "library") show("dashboard"); if (activeView() === "dashboard") renderDashboard(); updateNav(); updateCounts();
    }}
    function activeView() {{ return document.querySelector(".view.active")?.dataset.view || "home"; }}
    function show(view) {{
      if (view === "library") view = "dashboard";
      if ((view === "dashboard" || view === "saved") && !hasProfile()) view = "home";
      document.querySelectorAll("[data-view]").forEach(v => v.classList.toggle("active", v.dataset.view === view));
      if (view === "dashboard") renderDashboard();
      if (view === "saved") renderSaved();
      updateNav(); window.scrollTo({{top:0,behavior:"smooth"}});
    }}
    function updateNav() {{
      const active = activeView();
      document.querySelectorAll(".bottom button").forEach(btn => btn.classList.toggle("active", btn.dataset.go === active));
    }}
	    function renderQuestion() {{
	      normalizeAnswers();
	      const q = questions[step];
	      const opts = optionsFor(q);
	      document.querySelector("#step-label").textContent = lang === "zh" ? `${{t("step")}} ${{step + 1}} / 3` : `${{t("step")}} ${{step + 1}} de 3`;
	      document.querySelector("#stepper").innerHTML = questions.map((item, idx) => `<div class="step ${{idx === step ? "active" : ""}}"><i>${{idx + 1}}</i></div>`).join("");
	      document.querySelector("#question-wrap").innerHTML = `<section class="question active"><h1>${{esc(label(q))}}</h1><div class="options">${{opts.map(opt => `<button class="option ${{answers[q.id] === opt.id ? "selected" : ""}}" type="button" data-answer="${{esc(q.id)}}" data-value="${{esc(opt.id)}}"><span>${{opt.id === "hot" ? "★" : "●"}}</span><b>${{esc(label(opt))}}</b></button>`).join("")}}</div></section>`;
	      document.querySelector("#next-step span").textContent = step === questions.length - 1 ? t("finish") : t("next");
	    }}
	    async function fetchMatches(limit=80) {{
	      const params = new URLSearchParams({{limit}});
	      selectedAnswerValues().forEach(value => params.append("selected", value));
      const response = await fetch(`/api/creator/recommendations?${{params.toString()}}&_=${{Date.now()}}`);
      const data = await response.json();
      if (!response.ok) throw new Error(data.error || "Load failed");
      matchesCache = data.entries || [];
      updateCounts();
      return data;
    }}
    async function renderMini() {{
      try {{ const data = await fetchMatches(3); document.querySelector("#mini-cards").innerHTML = (data.entries || []).map((e,i)=>`<article class="mini-card"><img src="${{esc(e.thumbnail_url)}}" alt="" loading="lazy"><span class="score">${{96 - i * 3}} pontos</span><b>${{esc(e.title).slice(0,52)}}</b><span>🔥 ${{98 - i*11}},${{i+2}} mil</span></article>`).join(""); }} catch(e) {{}}
    }}
	    function selectedChips() {{
	      const lookup = Object.fromEntries(questions.flatMap(q => q.options.map(o => [o.id, o])));
	      return selectedAnswerValues().map(id => lookup[id]).filter(Boolean).map(opt => `<span class="chip">${{esc(label(opt))}} ✓</span>`).join("");
    }}
    function statusLabel(status) {{ return status === "planned" ? t("plan") : status === "finished" ? t("done") : status === "rejected" ? t("reject") : t("saved"); }}
    function card(e, i, compact=false) {{
      const status = statusOf(e.entry_id);
      return `<article class="script card"><div class="thumb"><img src="${{esc(e.thumbnail_url)}}" alt="" loading="lazy"><span>${{Math.max(78, 96 - Math.min(i, 18))}} match</span></div><div class="script-body"><h3>${{esc(e.title)}}</h3><p>${{esc(e.summary)}}</p><div class="tags"><span class="tag">${{esc(e.content_type)}}</span><span class="tag">${{t("peopleTag")}}</span><span class="tag">${{t("placeTag")}}</span>${{status ? `<span class="tag">${{statusLabel(status)}}</span>` : ""}}</div><div class="action-row"><button class="open" type="button" data-detail="${{esc(e.entry_id)}}">▷ ${{t("open")}}</button><button class="icon-btn" type="button" data-status="${{status === "saved" ? "" : "saved"}}" data-entry="${{esc(e.entry_id)}}">${{status === "saved" ? "✓" : "♡"}}</button><button class="icon-btn" type="button" data-status="planned" data-entry="${{esc(e.entry_id)}}">＋</button></div></div></article>`;
    }}
    function renderCards(target, entries, emptyTitle, emptyText) {{
      const node = document.querySelector(target);
      if (!node) return;
      node.innerHTML = entries.length ? entries.map((e,i) => card(e,i)).join("") : `<section class="state-card card"><h3>${{emptyTitle}}</h3><p>${{emptyText}}</p><button class="primary" type="button" data-go="dashboard">${{t("navHome")}}</button></section>`;
    }}
    function renderAllKnown() {{
      if (activeView() === "library") renderCards("#feed", matchesCache, "", "");
      if (activeView() === "dashboard") renderDashboard();
      if (activeView() === "saved") renderSaved();
    }}
    async function ensureMatches() {{ if (!matchesCache.length) await fetchMatches(80); }}
    async function renderDashboard() {{
      document.querySelector("#dashboard-filters").innerHTML = selectedChips() + `<button class="chip" type="button" data-go="choose">${{t("changePrefs")}}</button>`;
      document.querySelector("#dashboard-feed").innerHTML = `<section class="state-card card"><h3>${{lang === "zh" ? "正在加载今日推荐…" : "Carregando recomendações..."}}</h3><p>Koko Creator</p></section>`;
      try {{ await fetchMatches(80); renderCards("#dashboard-feed", matchesCache, t("emptySaved"), t("emptySavedText")); }}
      catch (error) {{ document.querySelector("#dashboard-feed").innerHTML = `<section class="state-card card"><h3>${{lang === "zh" ? "推荐加载失败" : "Não foi possível carregar"}}</h3><p>localhost:8391</p><button class="primary" type="button" data-go="dashboard">${{lang === "zh" ? "重试" : "Tentar novamente"}}</button></section>`; }}
      updateCounts();
    }}
    function savedEntriesFor(tab) {{ return (workspace[tab] || []).map(entryById).filter(Boolean); }}
    async function renderSaved() {{
      document.querySelector("#saved-tabs").innerHTML = [["saved",t("statusSaved")],["planned",t("statusPlanned")],["finished",t("statusFinished")],["rejected",t("statusRejected")]].map(([id,text]) => `<button type="button" class="${{savedTab === id ? "active" : ""}}" data-tab="${{id}}">${{text}} ${{(workspace[id] || []).length}}</button>`).join("");
      try {{ await ensureMatches(); renderCards("#saved-feed", savedEntriesFor(savedTab), t("emptySaved"), t("emptySavedText")); }}
      catch (error) {{ document.querySelector("#saved-feed").innerHTML = `<section class="state-card card"><h3>${{t("emptySaved")}}</h3><p>${{t("emptySavedText")}}</p></section>`; }}
    }}
    function openDetail(id) {{
      const e = entryById(id);
      if (!e) return;
      document.querySelector("#detail-content").innerHTML = `<div class="detail-media"><img src="${{esc(e.thumbnail_url)}}" alt=""></div><h2>${{esc(e.title)}}</h2><div class="tags"><span class="tag">${{esc(e.content_type)}}</span><span class="tag">${{t("peopleTag")}}</span><span class="tag">${{t("placeTag")}}</span></div><div class="detail-block"><h3>${{t("quickSummary")}}</h3><p>${{esc(e.summary)}}</p></div><div class="detail-block"><h3>${{t("howToUse")}}</h3><p>${{t("howToUseText")}}</p></div><section class="submit-box"><label for="submission-url">${{t("submitTitle")}}</label><p class="submit-hint">${{t("submitHint")}}</p><input id="submission-url" type="url" inputmode="url" placeholder="${{t("submitPlaceholder")}}" data-submit-url="${{esc(e.entry_id)}}"><button class="primary" type="button" data-submit-video="${{esc(e.entry_id)}}">${{t("submitButton")}}</button><div class="submit-status" id="submit-status-${{esc(e.entry_id)}}"></div></section><div class="cta-row"><button class="primary" type="button" data-status="planned" data-entry="${{esc(e.entry_id)}}">${{t("plan")}}</button><button class="secondary" type="button" data-status="saved" data-entry="${{esc(e.entry_id)}}">${{t("save")}}</button><button class="secondary" type="button" data-status="finished" data-entry="${{esc(e.entry_id)}}">${{t("done")}}</button><button class="secondary" type="button" data-status="rejected" data-entry="${{esc(e.entry_id)}}">${{t("reject")}}</button>${{e.html_url ? `<a class="primary" href="${{esc(e.html_url)}}" target="_blank" rel="noreferrer">▷ ${{t("details")}}</a>` : ""}}${{e.video_url ? `<a class="secondary" href="${{esc(e.video_url)}}" target="_blank" rel="noreferrer">${{t("original")}}</a>` : ""}}</div>`;
      document.querySelector("#detail-modal").classList.add("active");
    }}
    function closeDetail() {{ document.querySelector("#detail-modal").classList.remove("active"); }}
    async function submitVideo(entryId) {{
      const input = document.querySelector(`[data-submit-url="${{entryId}}"]`);
      const status = document.querySelector(`#submit-status-${{entryId}}`);
      const videoUrl = String(input?.value || "").trim();
      if (!videoUrl) {{ if (status) status.textContent = t("submitError"); return; }}
      if (status) status.textContent = lang === "zh" ? "提交中…" : "Enviando...";
      try {{
        const response = await fetch("/api/creator/submissions", {{
          method: "POST",
          headers: {{"Content-Type": "application/json"}},
          body: JSON.stringify({{entry_id: entryId, video_url: videoUrl, creator_id: "local_creator"}})
        }});
        const data = await response.json();
        if (!response.ok) throw new Error(data.error || "submit failed");
        if (status) status.textContent = t("submitOk");
        setStatus(entryId, "finished");
      }} catch (error) {{
        if (status) status.textContent = t("submitError");
      }}
    }}
    async function loadResults() {{
      document.querySelector("#selected-filters").innerHTML = selectedChips() + `<button class="chip" type="button" data-go="choose">${{lang === "zh" ? "重新筛选" : "Filtrar novamente"}}</button>`;
      document.querySelector("#feed").innerHTML = `<section class="state-card card"><h3>${{lang === "zh" ? "正在加载推荐脚本…" : "Carregando roteiros..."}}</h3><p>Koko Creator</p></section>`;
      try {{
        const data = await fetchMatches(80);
        const entries = data.entries || [];
        if (!entries.length) {{
          document.querySelector("#feed").innerHTML = `<section class="state-card card"><h3>${{lang === "zh" ? "暂时没有匹配脚本" : "Ainda não encontramos roteiros"}}</h3><p>${{lang === "zh" ? "换一个偏好组合再试试。" : "Tente ajustar suas preferências."}}</p><button class="primary" type="button" data-go="choose">${{lang === "zh" ? "重新筛选" : "Filtrar novamente"}}</button></section>`;
          return;
        }}
        renderCards("#feed", entries, "", "");
      }} catch (error) {{
        document.querySelector("#feed").innerHTML = `<section class="state-card card"><h3>${{lang === "zh" ? "推荐脚本加载失败" : "Não foi possível carregar"}}</h3><p>${{lang === "zh" ? "本地服务或接口刚才没有响应。请确认 localhost:8391 正在运行，然后重试。" : "O serviço local ou a API não respondeu. Confira se localhost:8391 está rodando e tente novamente."}}</p><button class="primary" type="button" data-go="library">${{lang === "zh" ? "重试" : "Tentar novamente"}}</button><button class="secondary" type="button" data-go="choose">${{lang === "zh" ? "重新筛选" : "Filtrar novamente"}}</button></section>`;
      }}
    }}
    document.addEventListener("click", event => {{
      const langBtn = event.target.closest("[data-lang]"); if (langBtn) {{ lang = langBtn.dataset.lang; localStorage.setItem(langKey, lang); applyLang(); return; }}
      const tabBtn = event.target.closest("[data-tab]"); if (tabBtn) {{ savedTab = tabBtn.dataset.tab; renderSaved(); return; }}
      const detailBtn = event.target.closest("[data-detail]"); if (detailBtn) {{ openDetail(detailBtn.dataset.detail); return; }}
      if (event.target.closest("[data-close-detail]") || event.target.id === "detail-modal") {{ closeDetail(); return; }}
      const submitBtn = event.target.closest("[data-submit-video]"); if (submitBtn) {{ submitVideo(submitBtn.dataset.submitVideo); return; }}
      const statusBtn = event.target.closest("[data-status]"); if (statusBtn) {{ setStatus(statusBtn.dataset.entry, statusBtn.dataset.status || ""); return; }}
      const go = event.target.closest("[data-go]"); if (go) {{ if (go.dataset.savedTab) savedTab = go.dataset.savedTab; show(go.dataset.go); return; }}
	      const answer = event.target.closest("[data-answer]"); if (answer) {{ answers[answer.dataset.answer] = answer.dataset.value; normalizeAnswers(); save(); renderQuestion(); return; }}
      if (event.target.closest("#next-step")) {{ if (step < questions.length - 1) {{ step += 1; renderQuestion(); }} else {{ save(); show("dashboard"); }} }}
    }});
    applyLang();
    show(hasProfile() ? "dashboard" : "home");
  </script>
</body>
</html>"""


class AppHandler(BaseHTTPRequestHandler):
    server_version = "VideoAnalysisV3Web/0.2"

    def send_json(self, payload: Any, status: int = 200, headers: list[tuple[str, str]] | None = None) -> None:
        raw = json.dumps(payload, ensure_ascii=False).encode("utf-8")
        self.send_response(status)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Content-Length", str(len(raw)))
        self.send_header("Cache-Control", "no-store, no-cache, must-revalidate, max-age=0")
        self.send_header("Pragma", "no-cache")
        self.send_header("Expires", "0")
        for key, value in headers or []:
            self.send_header(key, value)
        self.end_headers()
        self.wfile.write(raw)

    def send_html(self, body: str, status: int = 200, headers: list[tuple[str, str]] | None = None) -> None:
        raw = body.encode("utf-8")
        self.send_response(status)
        self.send_header("Content-Type", "text/html; charset=utf-8")
        self.send_header("Content-Length", str(len(raw)))
        self.send_header("Cache-Control", "no-store, no-cache, must-revalidate, max-age=0")
        self.send_header("Pragma", "no-cache")
        self.send_header("Expires", "0")
        for key, value in headers or []:
            self.send_header(key, value)
        self.end_headers()
        self.wfile.write(raw)

    def send_file(self, path: Path) -> None:
        if path.suffix == ".html":
            content_type = "text/html; charset=utf-8"
        elif path.suffix == ".json":
            content_type = "application/json; charset=utf-8"
        elif path.suffix == ".mp4":
            content_type = "video/mp4"
        elif path.suffix == ".png":
            content_type = "image/png"
        elif path.suffix.lower() in {".jpg", ".jpeg"}:
            content_type = "image/jpeg"
        elif path.suffix.lower() == ".webp":
            content_type = "image/webp"
        elif path.suffix == ".svg":
            content_type = "image/svg+xml; charset=utf-8"
        elif path.suffix == ".docx":
            content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        else:
            content_type = "application/octet-stream"
        if path.suffix == ".mp4":
            file_size = path.stat().st_size
            range_header = self.headers.get("Range") or ""
            match = re.fullmatch(r"bytes=(\d*)-(\d*)", range_header.strip())
            if match:
                start_text, end_text = match.groups()
                start = int(start_text) if start_text else 0
                end = int(end_text) if end_text else file_size - 1
                start = max(0, min(start, file_size - 1))
                end = max(start, min(end, file_size - 1))
                length = end - start + 1
                with path.open("rb") as handle:
                    handle.seek(start)
                    raw = handle.read(length)
                self.send_response(206)
                self.send_header("Content-Type", content_type)
                self.send_header("Accept-Ranges", "bytes")
                self.send_header("Content-Range", f"bytes {start}-{end}/{file_size}")
                self.send_header("Cache-Control", "no-store, no-cache, must-revalidate, max-age=0")
                self.send_header("Pragma", "no-cache")
                self.send_header("Expires", "0")
                self.send_header("Content-Length", str(len(raw)))
                self.end_headers()
                self.wfile.write(raw)
                return
        raw = path.read_bytes()
        self.send_response(200)
        self.send_header("Content-Type", content_type)
        if path.suffix == ".mp4":
            self.send_header("Accept-Ranges", "bytes")
        self.send_header("Cache-Control", "no-store, no-cache, must-revalidate, max-age=0")
        self.send_header("Pragma", "no-cache")
        self.send_header("Expires", "0")
        if path.suffix == ".docx":
            self.send_header("Content-Disposition", f'attachment; filename="{path.name}"')
        self.send_header("Content-Length", str(len(raw)))
        self.end_headers()
        self.wfile.write(raw)

    def head_file(self, path: Path) -> None:
        if path.suffix == ".html":
            content_type = "text/html; charset=utf-8"
        elif path.suffix == ".json":
            content_type = "application/json; charset=utf-8"
        elif path.suffix == ".mp4":
            content_type = "video/mp4"
        elif path.suffix == ".png":
            content_type = "image/png"
        elif path.suffix.lower() in {".jpg", ".jpeg"}:
            content_type = "image/jpeg"
        elif path.suffix.lower() == ".webp":
            content_type = "image/webp"
        elif path.suffix == ".svg":
            content_type = "image/svg+xml; charset=utf-8"
        elif path.suffix == ".docx":
            content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        else:
            content_type = "application/octet-stream"
        self.send_response(200)
        self.send_header("Content-Type", content_type)
        self.send_header("Cache-Control", "no-store, no-cache, must-revalidate, max-age=0")
        self.send_header("Pragma", "no-cache")
        self.send_header("Expires", "0")
        if path.suffix == ".docx":
            self.send_header("Content-Disposition", f'attachment; filename="{path.name}"')
        if path.suffix == ".mp4":
            self.send_header("Accept-Ranges", "bytes")
        self.send_header("Content-Length", str(path.stat().st_size))
        self.end_headers()

    def read_json(self) -> dict[str, Any]:
        length = int(self.headers.get("Content-Length", "0"))
        raw = self.rfile.read(length) if length else b"{}"
        return json.loads(raw.decode("utf-8"))

    def serve_result_file(self, job_id: str, filename: str) -> None:
        if not re.fullmatch(r"[0-9a-f]{32}", job_id):
            self.send_error(HTTPStatus.NOT_FOUND)
            return
        path = (RESULTS_ROOT / job_id / filename).resolve()
        base = (RESULTS_ROOT / job_id).resolve()
        if base not in path.parents and path != base:
            self.send_error(HTTPStatus.NOT_FOUND)
            return
        if not path.exists() or not path.is_file():
            self.send_error(HTTPStatus.NOT_FOUND)
            return
        self.send_file(path)

    def result_path(self, job_id: str, filename: str) -> Path | None:
        if not re.fullmatch(r"[0-9a-f]{32}", job_id):
            return None
        path = (RESULTS_ROOT / job_id / filename).resolve()
        base = (RESULTS_ROOT / job_id).resolve()
        if base not in path.parents and path != base:
            return None
        if not path.exists() or not path.is_file():
            return None
        return path

    def do_GET(self) -> None:
        parsed = urllib.parse.urlparse(self.path)
        if parsed.path == "/":
            self.send_html(page_html())
            return
        if parsed.path == "/studio":
            self.send_html(studio_html())
            return
        if parsed.path == "/error-cases":
            if has_error_case_access(self):
                self.send_html(error_cases_html())
            else:
                self.send_html(error_cases_login_html())
            return
        if parsed.path == "/stats":
            self.send_html(stats_html())
            return
        if parsed.path == "/library":
            self.send_html(creator_admin_html("scripts", library_mode=True))
            return
        if parsed.path == "/creator-admin/scripts":
            self.send_response(302)
            self.send_header("Location", "/library")
            self.end_headers()
            return
        creator_admin_tab = creator_admin_tab_for_path(parsed.path)
        if creator_admin_tab:
            self.send_html(creator_admin_html(creator_admin_tab))
            return
        if parsed.path == "/api/creator-admin/scripts":
            if not has_creator_admin_access(self):
                self.send_json({"error": "请先登录 Creator 运营后台。"}, status=401)
                return
            query = urllib.parse.parse_qs(parsed.query)
            try:
                limit = max(1, min(500, int((query.get("limit") or ["100"])[0] or "100")))
            except Exception:
                limit = 100
            scope = str((query.get("scope") or ["portal_visible"])[0] or "portal_visible").strip()
            if scope not in {"portal_visible", "hidden", "incomplete", "all"}:
                scope = "portal_visible"
            search = urllib.parse.urlencode({"limit": limit, "scope": scope})
            status, payload = creator_admin_remote_json(f"/api/admin/scripts?{search}")
            if status == 200 and isinstance(payload.get("entries"), list):
                save_creator_admin_scripts_cache(payload)
            elif status >= 500:
                cached = load_creator_admin_scripts_cache()
                if cached:
                    payload = dict(cached)
                    payload["from_cache"] = True
                    payload["remote_error"] = payload.get("remote_error") or "Creator remote list is temporarily unavailable."
                    status = 200
            self.send_json(payload, status=status)
            return
        if parsed.path == "/api/creator-admin/analytics":
            if not has_creator_admin_access(self):
                self.send_json({"error": "请先登录 Creator 运营后台后查看 kokocomedy 使用情况。"}, status=401)
                return
            query = urllib.parse.parse_qs(parsed.query)
            try:
                days = max(1, min(90, int((query.get("days") or ["14"])[0] or "14")))
            except Exception:
                days = 14
            include_inactive = (query.get("include_inactive") or ["0"])[0] == "1"
            remote_query = {"days": days, "include_inactive": "1" if include_inactive else "0"}
            status, payload = creator_admin_remote_json(f"/api/admin/analytics?{urllib.parse.urlencode(remote_query)}")
            self.send_json(payload, status=status)
            return
        if parsed.path == "/api/creator-admin/creators":
            if not has_creator_admin_access(self):
                self.send_json({"error": "请先登录 Creator 运营后台。"}, status=401)
                return
            status, payload = creator_admin_remote_json("/api/admin/creators")
            self.send_json(payload, status=status)
            return
        if parsed.path == "/api/creator-admin/state":
            if not has_creator_admin_access(self):
                self.send_json({"error": "请先登录 Creator 运营后台。"}, status=401)
                return
            self.send_json({"ok": True, "state": load_creator_admin_state()})
            return
        creator_ops_state_match = re.fullmatch(r"/api/creator-admin/state/([0-9a-f]{32})", parsed.path)
        if creator_ops_state_match:
            if not has_creator_admin_access(self):
                self.send_json({"error": "请先登录 Creator 运营后台。"}, status=401)
                return
            self.send_json({"ok": True, "creator_state": creator_admin_state_for(creator_ops_state_match.group(1))})
            return
        creator_ops_recommend_match = re.fullmatch(r"/api/creator-admin/creators/([0-9a-f]{32})/recommendations", parsed.path)
        if creator_ops_recommend_match:
            if not has_creator_admin_access(self):
                self.send_json({"error": "请先登录 Creator 运营后台。"}, status=401)
                return
            query = urllib.parse.parse_qs(parsed.query)
            try:
                limit = max(1, min(50, int((query.get("limit") or ["5"])[0] or "5")))
            except Exception:
                limit = 5
            status, payload = creator_admin_remote_json(f"/api/admin/creators/{creator_ops_recommend_match.group(1)}/recommendations?{urllib.parse.urlencode({'limit': limit})}")
            self.send_json(payload, status=status)
            return
        if parsed.path == "/api/creator-admin/accounts":
            if not has_creator_admin_access(self):
                self.send_json({"error": "请先登录 Creator 运营后台。"}, status=401)
                return
            status, payload = creator_admin_remote_json("/api/admin/accounts")
            self.send_json(payload, status=status)
            return
        if parsed.path == "/api/creator-admin/submissions":
            if not has_creator_admin_access(self):
                self.send_json({"error": "请先登录 Creator 运营后台。"}, status=401)
                return
            query = urllib.parse.parse_qs(parsed.query)
            try:
                limit = max(1, min(300, int((query.get("limit") or ["80"])[0] or "80")))
            except Exception:
                limit = 80
            try:
                offset = max(0, int((query.get("offset") or ["0"])[0] or "0"))
            except Exception:
                offset = 0
            remote_query = urllib.parse.urlencode({"limit": limit, "offset": offset})
            status, payload = creator_admin_remote_json(f"/api/admin/submissions?{remote_query}", timeout=20)
            self.send_json(payload, status=status)
            return
        if parsed.path == "/api/creator-admin/intakes":
            if not has_creator_admin_access(self):
                self.send_json({"error": "请先登录 Creator 运营后台。"}, status=401)
                return
            status, payload = creator_admin_remote_json("/api/admin/intakes")
            self.send_json(payload, status=status)
            return
        creator_import_match = re.fullmatch(r"/api/creator-admin/imports/([0-9a-f]{32})", parsed.path)
        if creator_import_match:
            if not has_creator_admin_access(self):
                self.send_json({"error": "请先登录 Creator 运营后台。"}, status=401)
                return
            job = public_creator_import_job(creator_import_match.group(1))
            if not job:
                self.send_json({"error": "导入任务不存在。"}, status=404)
                return
            self.send_json({"ok": True, "job": job})
            return
        if parsed.path == "/creator-portal":
            if not is_local_creator_portal_request(self):
                self.send_error(HTTPStatus.NOT_FOUND)
                return
            self.send_html(creator_portal_html())
            return
        if parsed.path == "/brand/kwai-wordmark.svg" and HERO_WORDMARK.exists():
            self.send_file(HERO_WORDMARK)
            return
        if parsed.path == "/brand/studio-hero-video.png" and STUDIO_HERO_VIDEO.exists():
            self.send_file(STUDIO_HERO_VIDEO)
            return
        if parsed.path == "/brand/studio-title-art.png" and STUDIO_TITLE_ART.exists():
            self.send_file(STUDIO_TITLE_ART)
            return
        if parsed.path == "/brand/studio-hero-banner-shallow.png" and STUDIO_HERO_BANNER_SHALLOW.exists():
            self.send_file(STUDIO_HERO_BANNER_SHALLOW)
            return
        if parsed.path in {"/favicon.svg", "/favicon.ico", "/brand/kwai-favicon.svg"} and KWAI_FAVICON.exists():
            self.send_file(KWAI_FAVICON)
            return
        if parsed.path == "/healthz":
            self.send_json({"ok": True, "time": now_iso(), "skill_root": str(SKILL_ROOT)})
            return
        if parsed.path == "/api/library":
            self.send_json({"entries": load_library_entries()})
            return
        library_workbench_match = re.fullmatch(r"/api/library-workbench/([0-9a-f]{32})", parsed.path)
        if library_workbench_match:
            try:
                payload = public_library_workbench(library_workbench_match.group(1))
            except Exception as exc:
                self.send_json({"error": friendly_error(str(exc))}, status=404)
                return
            self.send_json(payload)
            return
        if parsed.path == "/api/creator/facets":
            if not is_local_creator_portal_request(self):
                self.send_error(HTTPStatus.NOT_FOUND)
                return
            self.send_json(creator_facets_payload())
            return
        if parsed.path == "/api/creator/submissions":
            if not is_local_creator_portal_request(self):
                self.send_error(HTTPStatus.NOT_FOUND)
                return
            self.send_json({"submissions": load_creator_submissions()})
            return
        if parsed.path == "/api/creator/sync-status":
            if not is_local_creator_portal_request(self):
                self.send_error(HTTPStatus.NOT_FOUND)
                return
            meta = read_json_file(CREATOR_SYNC_META_FILE, default={})
            if not isinstance(meta, dict):
                meta = {}
            self.send_json({
                "source_url": CREATOR_LIBRARY_SOURCE_URL,
                "cache_exists": CREATOR_ONLINE_LIBRARY_FILE.exists(),
                "entries_count": len(read_json_file(CREATOR_ONLINE_LIBRARY_FILE, default=[])) if CREATOR_ONLINE_LIBRARY_FILE.exists() else 0,
                **meta,
            })
            return
        if parsed.path.startswith("/api/creator/thumbnail/"):
            if not is_local_creator_portal_request(self):
                self.send_error(HTTPStatus.NOT_FOUND)
                return
            filename = parsed.path.rsplit("/", 1)[-1]
            entry_id = filename[:-4] if filename.endswith(".svg") else filename
            if not re.fullmatch(r"[0-9a-f]{32}", entry_id):
                self.send_error(HTTPStatus.NOT_FOUND)
                return
            entry = creator_entry_by_id(entry_id)
            if not entry:
                self.send_error(HTTPStatus.NOT_FOUND)
                return
            thumbnail_url = creator_thumbnail_url_for_entry(entry)
            if thumbnail_url:
                try:
                    req = urllib.request.Request(
                        thumbnail_url,
                        headers={"User-Agent": "Mozilla/5.0"},
                    )
                    with urllib.request.urlopen(req, timeout=15) as response:
                        raw = response.read()
                        content_type = response.headers.get("Content-Type") or "image/webp"
                    self.send_response(HTTPStatus.OK)
                    self.send_header("Content-Type", content_type)
                    self.send_header("Cache-Control", "public, max-age=86400")
                    self.send_header("Content-Length", str(len(raw)))
                    self.end_headers()
                    self.wfile.write(raw)
                    return
                except Exception:
                    pass
            raw = creator_placeholder_svg(entry)
            self.send_response(HTTPStatus.OK)
            self.send_header("Content-Type", "image/svg+xml; charset=utf-8")
            self.send_header("Cache-Control", "public, max-age=3600")
            self.send_header("Content-Length", str(len(raw)))
            self.end_headers()
            self.wfile.write(raw)
            return
        if parsed.path == "/api/creator/recommendations":
            if not is_local_creator_portal_request(self):
                self.send_error(HTTPStatus.NOT_FOUND)
                return
            query = urllib.parse.parse_qs(parsed.query)
            selected = [str(value or "") for value in query.get("selected", [])]
            try:
                limit = max(1, min(200, int(str((query.get("limit") or ["80"])[0] or "80"))))
            except Exception:
                limit = 80
            self.send_json(creator_recommendation_payload(selected, limit=limit))
            return
        if parsed.path.startswith("/api/filter-jobs/"):
            job_id = parsed.path.split("/")[-1]
            with filter_jobs_lock:
                job = filter_jobs.get(job_id)
            if not job:
                self.send_json({"error": "Filter job not found."}, status=404)
                return
            self.send_json(public_filter_job_view(job))
            return
        if parsed.path.startswith("/api/translation-jobs/"):
            job_id = parsed.path.split("/")[-1]
            with translation_jobs_lock:
                job = translation_jobs.get(job_id)
            if not job:
                self.send_json({"error": "Translation job not found."}, status=404)
                return
            self.send_json(public_translation_job_view(job))
            return
        if parsed.path.startswith("/api/jobs/"):
            reconcile_stale_jobs()
            job_id = parsed.path.split("/")[-1]
            with job_lock:
                recompute_job_status(job_id)
                job = jobs.get(job_id)
            if not job:
                self.send_json({"error": "Job not found."}, status=404)
                return
            self.send_json(public_job_view(job))
            return
        if parsed.path.startswith("/results/"):
            parts = [part for part in parsed.path.split("/") if part]
            if len(parts) >= 3:
                self.serve_result_file(parts[1], "/".join(parts[2:]))
                return
        self.send_error(HTTPStatus.NOT_FOUND)

    def do_HEAD(self) -> None:
        parsed = urllib.parse.urlparse(self.path)
        if parsed.path == "/":
            body = page_html().encode("utf-8")
            self.send_response(200)
            self.send_header("Content-Type", "text/html; charset=utf-8")
            self.send_header("Content-Length", str(len(body)))
            self.end_headers()
            return
        if parsed.path == "/studio":
            body = studio_html().encode("utf-8")
            self.send_response(200)
            self.send_header("Content-Type", "text/html; charset=utf-8")
            self.send_header("Content-Length", str(len(body)))
            self.end_headers()
            return
        if parsed.path == "/error-cases":
            body = (error_cases_html() if has_error_case_access(self) else error_cases_login_html()).encode("utf-8")
            self.send_response(200)
            self.send_header("Content-Type", "text/html; charset=utf-8")
            self.send_header("Content-Length", str(len(body)))
            self.end_headers()
            return
        if parsed.path == "/stats":
            body = stats_html().encode("utf-8")
            self.send_response(200)
            self.send_header("Content-Type", "text/html; charset=utf-8")
            self.send_header("Content-Length", str(len(body)))
            self.end_headers()
            return
        if parsed.path == "/library":
            body = creator_admin_html("scripts", library_mode=True).encode("utf-8")
            self.send_response(200)
            self.send_header("Content-Type", "text/html; charset=utf-8")
            self.send_header("Content-Length", str(len(body)))
            self.end_headers()
            return
        if parsed.path == "/creator-admin/scripts":
            self.send_response(302)
            self.send_header("Location", "/library")
            self.end_headers()
            return
        creator_admin_tab = creator_admin_tab_for_path(parsed.path)
        if creator_admin_tab:
            body = creator_admin_html(creator_admin_tab).encode("utf-8")
            self.send_response(200)
            self.send_header("Content-Type", "text/html; charset=utf-8")
            self.send_header("Content-Length", str(len(body)))
            self.end_headers()
            return
        if parsed.path == "/creator-portal":
            if not is_local_creator_portal_request(self):
                self.send_error(HTTPStatus.NOT_FOUND)
                return
            body = creator_portal_html().encode("utf-8")
            self.send_response(200)
            self.send_header("Content-Type", "text/html; charset=utf-8")
            self.send_header("Content-Length", str(len(body)))
            self.end_headers()
            return
        if parsed.path == "/brand/kwai-wordmark.svg" and HERO_WORDMARK.exists():
            self.head_file(HERO_WORDMARK)
            return
        if parsed.path == "/brand/studio-hero-video.png" and STUDIO_HERO_VIDEO.exists():
            self.head_file(STUDIO_HERO_VIDEO)
            return
        if parsed.path == "/brand/studio-title-art.png" and STUDIO_TITLE_ART.exists():
            self.head_file(STUDIO_TITLE_ART)
            return
        if parsed.path == "/brand/studio-hero-banner-shallow.png" and STUDIO_HERO_BANNER_SHALLOW.exists():
            self.head_file(STUDIO_HERO_BANNER_SHALLOW)
            return
        if parsed.path in {"/favicon.svg", "/favicon.ico", "/brand/kwai-favicon.svg"} and KWAI_FAVICON.exists():
            self.head_file(KWAI_FAVICON)
            return
        if parsed.path == "/healthz":
            payload = json.dumps({"ok": True, "time": now_iso(), "skill_root": str(SKILL_ROOT)}, ensure_ascii=False).encode("utf-8")
            self.send_response(200)
            self.send_header("Content-Type", "application/json; charset=utf-8")
            self.send_header("Content-Length", str(len(payload)))
            self.end_headers()
            return
        if parsed.path.startswith("/results/"):
            parts = [part for part in parsed.path.split("/") if part]
            if len(parts) >= 3:
                path = self.result_path(parts[1], "/".join(parts[2:]))
                if not path:
                    self.send_error(HTTPStatus.NOT_FOUND)
                    return
                self.head_file(path)
                return
        self.send_error(HTTPStatus.NOT_FOUND)

    def do_POST(self) -> None:
        parsed = urllib.parse.urlparse(self.path)
        if parsed.path == "/creator-admin/login":
            try:
                payload = self.read_json()
            except json.JSONDecodeError:
                payload = {}
            password = str(payload.get("password") or "")
            if not secrets.compare_digest(password, CREATOR_ADMIN_PASSWORD):
                self.send_json({"error": "Creator 运营后台密码不正确。"}, status=401)
                return
            cookie = f"{CREATOR_ADMIN_AUTH_COOKIE}={urllib.parse.quote(CREATOR_ADMIN_PASSWORD)}; Path=/; Max-Age=604800; HttpOnly; SameSite=Lax"
            self.send_json({"ok": True}, headers=[("Set-Cookie", cookie)])
            return
        if parsed.path == "/creator-admin/logout":
            self.send_json({"ok": True}, headers=[("Set-Cookie", f"{CREATOR_ADMIN_AUTH_COOKIE}=; Path=/; Max-Age=0; HttpOnly; SameSite=Lax")])
            return
        if parsed.path == "/api/creator-admin/sync":
            if not has_creator_admin_access(self):
                self.send_json({"error": "请先登录 Creator 运营后台。"}, status=401)
                return
            result = trigger_creator_center_sync()
            self.send_json(result, status=200 if result.get("ok") else 502)
            return
        if parsed.path == "/api/admin/scripts/import":
            if not has_creator_admin_access(self):
                self.send_json({"error": "请先登录 Creator 运营后台。"}, status=401)
                return
            try:
                payload = self.read_json()
                result = save_creator_direct_import(payload)
            except json.JSONDecodeError:
                self.send_json({"error": "Invalid JSON body."}, status=400)
                return
            except Exception as exc:
                self.send_json({"error": friendly_error(str(exc))}, status=400)
                return
            self.send_json(result, status=201)
            return
        if parsed.path == "/api/creator-admin/imports":
            if not has_creator_admin_access(self):
                self.send_json({"error": "请先登录 Creator 运营后台。"}, status=401)
                return
            try:
                payload = self.read_json()
                job = start_creator_excel_import(
                    str(payload.get("filename") or "scripts.xlsx"),
                    str(payload.get("file_b64") or ""),
                    content_type=str(payload.get("content_type") or DEFAULT_CONTENT_TYPE).strip(),
                )
            except json.JSONDecodeError:
                self.send_json({"error": "Invalid JSON body."}, status=400)
                return
            except Exception as exc:
                self.send_json({"error": friendly_error(str(exc))}, status=400)
                return
            self.send_json({"ok": True, "job": job}, status=202)
            return
        if parsed.path == "/api/creator-admin/bulk-delete":
            if not has_creator_admin_access(self):
                self.send_json({"error": "请先登录 Creator 运营后台。"}, status=401)
                return
            try:
                payload = self.read_json()
            except json.JSONDecodeError:
                self.send_json({"error": "Invalid JSON body."}, status=400)
                return
            status, result = creator_admin_remote_json("/api/admin/scripts/bulk-delete", method="POST", payload=payload)
            self.send_json(result, status=status)
            return
        if parsed.path == "/api/creator-admin/creators":
            if not has_creator_admin_access(self):
                self.send_json({"error": "请先登录 Creator 运营后台。"}, status=401)
                return
            try:
                payload = self.read_json()
            except json.JSONDecodeError:
                self.send_json({"error": "Invalid JSON body."}, status=400)
                return
            status, result = creator_admin_remote_json("/api/admin/creators", method="POST", payload=payload)
            self.send_json(result, status=status)
            return
        creator_ops_state_save_match = re.fullmatch(r"/api/creator-admin/state/([0-9a-f]{32})", parsed.path)
        if creator_ops_state_save_match:
            if not has_creator_admin_access(self):
                self.send_json({"error": "请先登录 Creator 运营后台。"}, status=401)
                return
            try:
                payload = self.read_json()
                creator_state = save_creator_admin_state_for(creator_ops_state_save_match.group(1), payload)
            except json.JSONDecodeError:
                self.send_json({"error": "Invalid JSON body."}, status=400)
                return
            except ValueError as exc:
                self.send_json({"error": str(exc)}, status=400)
                return
            self.send_json({"ok": True, "creator_state": creator_state})
            return
        if parsed.path == "/api/creator-admin/creators/import":
            if not has_creator_admin_access(self):
                self.send_json({"error": "请先登录 Creator 运营后台。"}, status=401)
                return
            try:
                payload = self.read_json()
            except json.JSONDecodeError:
                self.send_json({"error": "Invalid JSON body."}, status=400)
                return
            status, result = creator_admin_remote_json("/api/admin/creators/import", method="POST", payload=payload)
            self.send_json(result, status=status)
            return
        if parsed.path == "/api/creator-admin/accounts":
            if not has_creator_admin_access(self):
                self.send_json({"error": "请先登录 Creator 运营后台。"}, status=401)
                return
            try:
                payload = self.read_json()
            except json.JSONDecodeError:
                self.send_json({"error": "Invalid JSON body."}, status=400)
                return
            status, result = creator_admin_remote_json("/api/admin/accounts", method="POST", payload=payload)
            self.send_json(result, status=status)
            return
        if parsed.path == "/api/creator-admin/submissions/backfill-creators":
            if not has_creator_admin_access(self):
                self.send_json({"error": "请先登录 Creator 运营后台。"}, status=401)
                return
            try:
                payload = self.read_json()
            except json.JSONDecodeError:
                payload = {}
            status, result = creator_admin_remote_json("/api/admin/submissions/backfill-creators", method="POST", payload=payload)
            self.send_json(result, status=status)
            return
        if parsed.path == "/api/creator-admin/submissions/delete":
            if not has_creator_admin_access(self):
                self.send_json({"error": "请先登录 Creator 运营后台。"}, status=401)
                return
            try:
                payload = self.read_json()
            except json.JSONDecodeError:
                self.send_json({"error": "Invalid JSON body."}, status=400)
                return
            status, result = creator_admin_remote_json("/api/admin/submissions/delete", method="POST", payload=payload)
            self.send_json(result, status=status)
            return
        creator_ops_update_match = re.fullmatch(r"/api/creator-admin/creators/([0-9a-f]{32})", parsed.path)
        if creator_ops_update_match:
            if not has_creator_admin_access(self):
                self.send_json({"error": "请先登录 Creator 运营后台。"}, status=401)
                return
            try:
                payload = self.read_json()
            except json.JSONDecodeError:
                self.send_json({"error": "Invalid JSON body."}, status=400)
                return
            status, result = creator_admin_remote_json(f"/api/admin/creators/{creator_ops_update_match.group(1)}", method="POST", payload=payload)
            self.send_json(result, status=status)
            return
        creator_admin_update_match = re.fullmatch(r"/api/creator-admin/scripts/([0-9a-f]{32})", parsed.path)
        if creator_admin_update_match:
            if not has_creator_admin_access(self):
                self.send_json({"error": "请先登录 Creator 运营后台。"}, status=401)
                return
            try:
                payload = self.read_json()
            except json.JSONDecodeError:
                self.send_json({"error": "Invalid JSON body."}, status=400)
                return
            status, result = creator_admin_remote_json(f"/api/admin/scripts/{creator_admin_update_match.group(1)}", method="POST", payload=payload)
            self.send_json(result, status=status)
            return
        if parsed.path == "/api/creator/submissions":
            if not is_local_creator_portal_request(self):
                self.send_error(HTTPStatus.NOT_FOUND)
                return
            try:
                payload = self.read_json()
                submission = save_creator_submission(payload)
            except json.JSONDecodeError:
                self.send_json({"error": "Invalid JSON body."}, status=400)
                return
            except ValueError as exc:
                self.send_json({"error": str(exc)}, status=400)
                return
            except Exception as exc:
                self.send_json({"error": friendly_error(str(exc))}, status=500)
                return
            self.send_json({"ok": True, "submission": submission}, status=201)
            return
        if parsed.path == "/api/stop-all":
            summary = stop_all_tasks()
            self.send_json({"ok": True, **summary})
            return
        if parsed.path == "/error-cases/login":
            length = int(self.headers.get("Content-Length", "0"))
            raw = self.rfile.read(length).decode("utf-8", errors="ignore") if length else ""
            form = urllib.parse.parse_qs(raw, keep_blank_values=True)
            password = str((form.get("password") or [""])[0]).strip()
            if password != ERROR_CASE_PASSWORD:
                self.send_html(error_cases_login_html("密码不正确，请重试。"), status=401)
                return
            self.send_html(
                error_cases_html(),
                headers=[("Set-Cookie", f"{ERROR_CASE_AUTH_COOKIE}=1; Path=/; Max-Age=604800; SameSite=Lax")],
            )
            return
        item_match = re.fullmatch(r"/api/items/([0-9a-f]{32})/(save|save-to-library|confirm-library|display-language)", parsed.path)
        if item_match:
            item_id, action = item_match.groups()
            context = find_item_context(item_id)
            if not context:
                self.send_json({"error": "Script item not found."}, status=404)
                return
            try:
                payload = self.read_json()
            except json.JSONDecodeError:
                self.send_json({"error": "Invalid JSON body."}, status=400)
                return
            if action == "display-language":
                try:
                    updated_item = set_item_display_language(
                        item_id,
                        str(payload.get("language") or "zh").strip().lower(),
                        payload if isinstance(payload, dict) else {},
                    )
                except Exception as exc:
                    self.send_json({"error": friendly_error(str(exc))}, status=500)
                    return
                self.send_json({"ok": True, "item": updated_item})
                return
            parent_job_id, item_index, item = context
            if not item.get("result_json"):
                self.send_json({"error": "No script is available for editing yet."}, status=400)
                return
            if action == "confirm-library":
                if str(item.get("status") or "").strip() != "completed":
                    self.send_json({"error": "Only completed scripts can be added to the library."}, status=400)
                    return
                try:
                    apply_manual_item_content_type(parent_job_id, item_index, payload.get("content_type"))
                    refreshed_context = find_item_context(item_id)
                    if not refreshed_context:
                        raise RuntimeError("Script item not found.")
                    parent_job_id, item_index, item = refreshed_context
                    ensure_storyboard_cover_ready(item_id)
                    refreshed_context = find_item_context(item_id)
                    if not refreshed_context:
                        raise RuntimeError("Script item not found.")
                    parent_job_id, item_index, item = refreshed_context
                    base_script = item.get("zh_result_json") or item.get("result_json") or {}
                    updated_item = regenerate_item_outputs(
                        parent_job_id,
                        item_index,
                        item_id,
                        item.get("video_url") or "",
                        base_script,
                        persist_library=True,
                        target_language="pt",
                    )
                    if updated_item.get("saved_to_library_at"):
                        update_job_item(parent_job_id, item_index, saved_to_library_at=updated_item.get("saved_to_library_at"))
                    with job_lock:
                        updated_item = public_item_view(jobs[parent_job_id]["items"][item_index])
                except Exception as exc:
                    self.send_json({"error": friendly_error(str(exc))}, status=500)
                    return
                trigger_creator_center_sync_background("confirm_library")
                self.send_json({"ok": True, "item": updated_item, "saved_to_library": True})
                return
            try:
                target_language = str(payload.get("target_language") or item.get("display_language") or "zh").strip().lower()
                updated_script = apply_script_edits(item.get("result_json") or {}, payload)
                apply_manual_item_content_type(parent_job_id, item_index, payload.get("content_type"))
                refreshed_context = find_item_context(item_id)
                if not refreshed_context:
                    raise RuntimeError("Script item not found.")
                parent_job_id, item_index, item = refreshed_context
                if action == "save-to-library":
                    updated_item = save_item_edits_to_library(
                        parent_job_id,
                        item_index,
                        item_id,
                        updated_script,
                        target_language=target_language,
                    )
                else:
                    updated_item = regenerate_item_outputs(
                        parent_job_id,
                        item_index,
                        item_id,
                        item.get("video_url") or "",
                        updated_script,
                        persist_library=False,
                        target_language=target_language,
                    )
            except Exception as exc:
                self.send_json({"error": friendly_error(str(exc))}, status=500)
                return
            if action == "save-to-library":
                trigger_creator_center_sync_background("save_to_library")
            self.send_json({"ok": True, "item": updated_item, "saved_to_library": action == "save-to-library"})
            return
        storyboard_prompt_match = re.fullmatch(r"/api/items/([0-9a-f]{32})/storyboard/prompt", parsed.path)
        if storyboard_prompt_match:
            item_id = storyboard_prompt_match.group(1)
            try:
                payload = self.read_json()
            except json.JSONDecodeError:
                payload = {}
            try:
                updated_item = generate_storyboard_prompt_for_item(item_id, payload if isinstance(payload, dict) else {})
            except Exception as exc:
                self.send_json({"error": friendly_error(str(exc))}, status=500)
                return
            self.send_json({"ok": True, "item": updated_item, "storyboard_prompt": updated_item.get("storyboard_prompt") or ""})
            return
        storyboard_match = re.fullmatch(r"/api/items/([0-9a-f]{32})/storyboard(?:/(confirm))?", parsed.path)
        if storyboard_match:
            item_id, confirm_action = storyboard_match.groups()
            try:
                payload = self.read_json()
            except json.JSONDecodeError:
                payload = {}
            try:
                if confirm_action:
                    updated_item = confirm_storyboard_cover(item_id)
                else:
                    updated_item = generate_storyboard_preview(
                        item_id,
                        str((payload or {}).get("prompt") or "").strip(),
                        payload if isinstance(payload, dict) else {},
                    )
            except Exception as exc:
                self.send_json({"error": friendly_error(str(exc))}, status=500)
                return
            self.send_json({"ok": True, "item": updated_item})
            return
        review_match = re.fullmatch(r"/api/items/([0-9a-f]{32})/review", parsed.path)
        if review_match:
            item_id = review_match.group(1)
            try:
                payload = self.read_json()
            except json.JSONDecodeError:
                self.send_json({"error": "Invalid JSON body."}, status=400)
                return
            ok, result = start_review_job(item_id, payload.get("feedback") or "", payload.get("mode") or payload.get("review_mode") or "")
            if not ok:
                self.send_json({"error": result}, status=400)
                return
            self.send_json({"ok": True, "job_id": result, "item_id": item_id}, status=202)
            return
        chat_edit_match = re.fullmatch(r"/api/items/([0-9a-f]{32})/chat-edit", parsed.path)
        if chat_edit_match:
            item_id = chat_edit_match.group(1)
            try:
                payload = self.read_json()
            except json.JSONDecodeError:
                self.send_json({"error": "Invalid JSON body."}, status=400)
                return
            mode = str(payload.get("mode") or payload.get("edit_mode") or "minor").strip().lower()
            if mode in {"recheck", "full", "video"}:
                ok, result = start_review_job(item_id, payload.get("message") or "", REVIEW_MODE_FULL)
                if not ok:
                    self.send_json({"error": result}, status=400)
                    return
                self.send_json({"ok": True, "job_id": result, "item_id": item_id, "mode": "recheck"}, status=202)
                return
            ok, result = run_chat_script_edit(item_id, payload.get("message") or "", mode)
            if not ok:
                self.send_json({"error": result}, status=400)
                return
            self.send_json({"ok": True, **result})
            return
        if parsed.path == "/api/library/batch-delete":
            try:
                payload = self.read_json()
            except json.JSONDecodeError:
                self.send_json({"error": "Invalid JSON body."}, status=400)
                return
            raw_ids = payload.get("entry_ids")
            if not isinstance(raw_ids, list):
                self.send_json({"error": "entry_ids must be a list."}, status=400)
                return
            result = delete_library_entries([str(entry_id or "") for entry_id in raw_ids])
            self.send_json({"ok": True, **result})
            return
        library_type_match = re.fullmatch(r"/api/library/([0-9a-f]{32})/content-type", parsed.path)
        if library_type_match:
            entry_id = library_type_match.group(1)
            try:
                payload = self.read_json()
            except json.JSONDecodeError:
                self.send_json({"error": "Invalid JSON body."}, status=400)
                return
            try:
                updated = update_library_entry_content_type(entry_id, str(payload.get("content_type") or "").strip())
            except Exception as exc:
                self.send_json({"error": friendly_error(str(exc))}, status=400)
                return
            if not updated:
                self.send_json({"error": "Library entry not found."}, status=404)
                return
            self.send_json({"ok": True, "entry": updated})
            return
        if parsed.path == "/api/library/sync-creator-center":
            result = trigger_creator_center_sync()
            self.send_json(result, status=200 if result.get("ok") else 502)
            return
        if parsed.path != "/api/jobs":
            if parsed.path == "/api/translation-jobs":
                try:
                    payload = self.read_json()
                except json.JSONDecodeError:
                    self.send_json({"error": "Invalid JSON body."}, status=400)
                    return
                video_url = str(payload.get("video_url") or "").strip()
                if not video_url:
                    urls = split_video_urls(str(payload.get("raw_text") or ""))
                    video_url = urls[0] if urls else ""
                if not video_url:
                    self.send_json({"error": "请提供一个可公开访问的视频链接。"}, status=400)
                    return
                if not TRANSCREATE_VIDEO.exists():
                    self.send_json({"error": f"Missing transcreation entrypoint: {TRANSCREATE_VIDEO}"}, status=500)
                    return
                job = create_translation_job(video_url, language=str(payload.get("language") or "pt-BR").strip() or "pt-BR")
                self.send_json(job, status=202)
                return
            if parsed.path != "/api/filter-jobs":
                self.send_error(HTTPStatus.NOT_FOUND)
                return
            try:
                payload = self.read_json()
            except json.JSONDecodeError:
                self.send_json({"error": "Invalid JSON body."}, status=400)
                return
            raw_text = str(payload.get("raw_text") or "")
            direct_values = payload.get("video_urls")
            direct_urls: list[str] = []
            if isinstance(direct_values, list):
                direct_urls = unique_urls_from_values(direct_values)
            urls = unique_urls_from_values([*extract_http_urls(raw_text), *direct_urls])
            upload = payload.get("upload")
            if isinstance(upload, dict):
                filename = str(upload.get("filename") or "").strip()
                file_data = str(upload.get("file_data_base64") or "").strip()
                if filename and file_data:
                    try:
                        urls = unique_urls_from_values([*urls, *extract_urls_from_uploaded_file(filename, file_data)])
                    except Exception as exc:
                        self.send_json({"error": friendly_error(str(exc))}, status=400)
                        return
            kwai_urls = [url for url in urls if "kwai.com/" in url or "k.kwai.com/" in url]
            if not kwai_urls:
                self.send_json({"error": "请至少提供一个可识别的 Kwai 视频链接。"}, status=400)
                return
            try:
                job = create_filter_job(kwai_urls, source_label="studio-filter")
            except Exception as exc:
                log_runtime_warning("filter_job_create_failed", "Failed to create filter job.", error=str(exc))
                status = 507 if is_no_space_error(exc) else 500
                self.send_json({"error": friendly_error(str(exc))}, status=status)
                return
            self.send_json(job, status=202)
            return
        try:
            payload = self.read_json()
        except json.JSONDecodeError:
            self.send_json({"error": "Invalid JSON body."}, status=400)
            return
        raw_urls = payload.get("video_urls")
        if isinstance(raw_urls, list):
            video_urls = split_video_urls("\n".join(str(url or "") for url in raw_urls))
        else:
            video_urls = split_video_urls(str(payload.get("video_url", "")))
        if not video_urls:
            self.send_json({"error": "Please provide at least one valid public video URL."}, status=400)
            return
        if not AUTO_ANALYZE.exists():
            self.send_json({"error": f"Missing pipeline entrypoint: {AUTO_ANALYZE}"}, status=500)
            return
        mode = str(payload.get("mode") or "").strip().lower()
        user_prompt = sanitize_analysis_prompt(
            payload.get("user_prompt")
            or payload.get("analysis_prompt")
            or payload.get("extra_prompt")
            or ""
        )
        try:
            job = create_job(video_urls, mode=mode, user_prompt=user_prompt)
        except Exception as exc:
            log_runtime_warning("analysis_job_create_failed", "Failed to create analysis job.", error=str(exc))
            status = 507 if is_no_space_error(exc) else 500
            self.send_json({"error": friendly_error(str(exc))}, status=status)
            return
        self.send_json(job, status=202)

    def do_DELETE(self) -> None:
        parsed = urllib.parse.urlparse(self.path)
        creator_ops_delete_match = re.fullmatch(r"/api/creator-admin/creators/([0-9a-f]{32})", parsed.path)
        if creator_ops_delete_match:
            if not has_creator_admin_access(self):
                self.send_json({"error": "请先登录 Creator 运营后台。"}, status=401)
                return
            status, result = creator_admin_remote_json(f"/api/admin/creators/{creator_ops_delete_match.group(1)}", method="DELETE")
            self.send_json(result, status=status)
            return
        if parsed.path.startswith("/api/library/"):
            entry_id = parsed.path.split("/")[-1]
            if not re.fullmatch(r"[0-9a-f]{32}", entry_id):
                self.send_json({"error": "Invalid entry id."}, status=400)
                return
            if not delete_library_entry(entry_id):
                self.send_json({"error": "Library entry not found."}, status=404)
                return
            self.send_json({"ok": True, "entry_id": entry_id})
            return
        self.send_error(HTTPStatus.NOT_FOUND)


def main() -> int:
    load_jobs()
    load_filter_jobs()
    load_translation_jobs()
    try:
        cleanup_orphan_result_dirs()
    except Exception as exc:
        log_runtime_warning("orphan_results_cleanup_skipped", "Automatic orphan result cleanup failed during startup.", error=str(exc))
    restore_pending_jobs_to_queue()
    restore_pending_filter_jobs_to_queue()
    restore_pending_translation_jobs_to_queue()
    start_job_workers()
    start_filter_workers()
    start_translation_workers()
    start_watchdog()
    server = ThreadingHTTPServer(("0.0.0.0", PORT), AppHandler)
    print(json.dumps({"port": PORT, "data_root": str(DATA_ROOT), "skill_root": str(SKILL_ROOT)}, ensure_ascii=False))
    try:
        server.serve_forever()
    except KeyboardInterrupt:
        pass
    finally:
        server.server_close()
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
